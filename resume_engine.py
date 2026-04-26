import os
os.environ["STREAMLIT_WATCHDOG"] = "false"
import json
import random
import string
import re
import asyncio
import io
from concurrent.futures import ThreadPoolExecutor, as_completed
import urllib.parse
import base64
from io import BytesIO
from collections import Counter
from datetime import datetime
import time

import streamlit as st
import streamlit.components.v1 as components
from base64 import b64encode
import requests
import fitz
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import altair as alt
from PIL import Image
from pdf2image import convert_from_path
from dotenv import load_dotenv
from nltk.stem import WordNetLemmatizer
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from xhtml2pdf import pisa
from pydantic import BaseModel
from streamlit_pdf_viewer import pdf_viewer
import torch
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from llm_manager import (
    call_llm, load_groq_api_keys, get_healthy_keys, increment_key_usage,
    mark_key_failure, _mem_record_failure, _mem_clear_failure,
    _mem_increment_usage, _async_mark_failure, _async_increment_usage,
    _async_clear_failure,
)
from db_manager import (
    db_manager, insert_candidate, get_top_domains_by_score,
    get_database_stats, detect_domain_from_title_and_description,
    get_domain_similarity, build_resume_domain_prompt, build_jd_domain_prompt,
    DOMAIN_VALID_LIST,
)
from user_login import (
    create_user_table, add_user, complete_registration, verify_user,
    get_logins_today, get_total_registered_users, log_user_action,
    username_exists, email_exists, is_valid_email, save_user_api_key,
    get_user_api_key, get_all_user_logs, generate_otp, send_email_otp,
    get_user_by_email, update_password_by_email, is_strong_password,
    domain_has_mx_record, send_login_link, verify_login_token,
    cleanup_expired_login_tokens, check_and_gate_feature,
    record_feature_usage, get_usage_count_last_hour, check_brute_force,
)

# ── resume_engine.py ────────────────────────────────────────────────────────
gender_words = {
    "masculine": [
        # Dominance / aggression-coded
        "active", "aggressive", "ambitious", "assertive", "autonomous", "boast", "bold",
        "challenging", "competitive", "confident", "courageous", "decisive", "determined", "dominant", "driven",
        "dynamic", "forceful", "independent", "individualistic", "intellectual", "lead", "leader", "objective",
        "outspoken", "persistent", "principled", "proactive", "resilient", "self-reliant", "self-sufficient",
        "strong", "superior", "tenacious", "guru", "tech guru", "technical guru", "visionary", "manpower",
        "strongman", "command", "assert", "headstrong", "rockstar", "superstar", "go-getter", "trailblazer",
        "results-driven", "fast-paced", "determination", "competitive spirit",
        # Additional research-backed masculine-coded terms (Gaucher et al., 2011)
        "analytical", "backbone", "challenge", "champion", "combat", "conquer", "courageous",
        "crusade", "debate", "fearless", "fight", "grit", "hustle", "impact", "ninja",
        "power", "ruthless", "self-starter", "sharp", "warrior", "win", "wrestler",
        "alpha", "beast", "brutally honest", "cutting-edge", "dominate", "edge", "elite",
        "fearless", "grind", "hardcore", "hero", "high-performance", "intense",
        "kill it", "relentless", "savage", "slayer", "tiger", "tough", "uncompromising"
    ],
    
    "feminine": [
        # Communal / warmth-coded
        "affectionate", "agreeable", "attentive", "collaborative", "committed", "compassionate", "considerate",
        "cooperative", "dependable", "dependent", "emotional", "empathetic", "enthusiastic", "friendly", "gentle",
        "honest", "inclusive", "interpersonal", "kind", "loyal", "modest", "nurturing", "pleasant", "polite",
        "sensitive", "supportive", "sympathetic", "tactful", "tender", "trustworthy", "understanding", "warm",
        "yield", "adaptable", "communal", "helpful", "dedicated", "respectful", "nurture", "sociable",
        "relationship-oriented", "team player", "people-oriented", "empathetic listener",
        "gentle communicator", "open-minded",
        # Additional research-backed feminine-coded terms
        "balance", "caring", "child-friendly", "connect", "connection", "flexible hours",
        "harmony", "heart", "humanize", "mindful", "patience", "patient", "peace",
        "personal touch", "responsive", "share", "sharing", "together", "unite",
        "welcoming", "wholesome", "connect with", "feeling", "feelings", "giving back",
        "heartfelt", "humanity", "inspire", "inspired", "passion", "passionate",
        "personable", "relate", "relatable", "soften", "soft skills", "spread",
        "thrive", "togetherness", "transparent", "uplift", "vulnerable"
    ]
}

def detect_bias(text):
    # Split into sentences using simple delimiters
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())

    masc_set, fem_set = set(), set()
    masculine_found, feminine_found = [] , []

    masculine_words = sorted(gender_words["masculine"], key=len, reverse=True)
    feminine_words = sorted(gender_words["feminine"], key=len, reverse=True)

    for sent in sentences:
        sent_text = sent.strip()
        sent_lower = sent_text.lower()
        matched_spans = []

        def is_overlapping(start, end):
            return any(start < e and end > s for s, e in matched_spans)

        # 🔵 Highlight masculine words in blue
        for word in masculine_words:
            pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
            for match in pattern.finditer(sent_lower):
                start, end = match.span()
                if not is_overlapping(start, end):
                    matched_spans.append((start, end))
                    key = (word.lower(), sent_text)
                    if key not in masc_set:
                        masc_set.add(key)
                        highlighted = re.sub(
                            rf'\b({re.escape(word)})\b',
                            r'<span style="color:blue;">\1</span>',
                            sent_text,
                            flags=re.IGNORECASE
                        )
                        masculine_found.append({
                            "word": word,
                            "sentence": highlighted
                        })

        # 🔴 Highlight feminine words in red
        for word in feminine_words:
            pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
            for match in pattern.finditer(sent_lower):
                start, end = match.span()
                if not is_overlapping(start, end):
                    matched_spans.append((start, end))
                    key = (word.lower(), sent_text)
                    if key not in fem_set:
                        fem_set.add(key)
                        highlighted = re.sub(
                            rf'\b({re.escape(word)})\b',
                            r'<span style="color:red;">\1</span>',
                            sent_text,
                            flags=re.IGNORECASE
                        )
                        feminine_found.append({
                            "word": word,
                            "sentence": highlighted
                        })

    masc = len(masculine_found)
    fem = len(feminine_found)
    total = masc + fem
    bias_score = min(total / 20, 1.0) if total > 0 else 0.0

    return round(bias_score, 2), masc, fem, masculine_found, feminine_found

replacement_mapping = {
    "masculine": {
        "active": "engaged",
        "aggressive": "proactive",
        "ambitious": "motivated",
        "analytical": "detail-oriented",
        "assertive": "direct",
        "autonomous": "self-directed",
        "boast": "highlight",
        "bold": "confident",
        "challenging": "demanding",
        "competitive": "goal-oriented",
        "confident": "self-assured",
        "courageous": "bold",
        "decisive": "action-oriented",
        "determined": "focused",
        "dominant": "influential",
        "driven": "committed",
        "dynamic": "adaptable",
        "forceful": "persuasive",
        "guru": "technical expert",
        "independent": "self-sufficient",
        "individualistic": "self-motivated",
        "intellectual": "knowledgeable",
        "lead": "guide",
        "leader": "team lead",
        "objective": "unbiased",
        "outspoken": "expressive",
        "persistent": "tenacious",
        "principled": "ethical",
        "proactive": "initiative-taking",
        "resilient": "adaptable",
        "self-reliant": "resourceful",
        "self-sufficient": "capable",
        "strong": "capable",
        "superior": "exceptional",
        "tenacious": "determined",
        "technical guru": "technical expert",
        "visionary": "forward-thinking",
        "manpower": "workforce",
        "strongman": "resilient individual",
        "command": "direct",
        "assert": "state clearly",
        "headstrong": "determined",
        "rockstar": "top performer",
        "superstar": "outstanding contributor",
        "go-getter": "initiative-taker",
        "trailblazer": "innovator",
        "results-driven": "outcome-focused",
        "fast-paced": "dynamic",
        "determination": "commitment",
        "competitive spirit": "goal-oriented mindset",
        # New additions
        "ninja": "specialist",
        "warrior": "dedicated professional",
        "alpha": "senior",
        "beast": "high performer",
        "dominate": "excel in",
        "elite": "high-performing",
        "relentless": "persistent",
        "savage": "highly skilled",
        "hustle": "work efficiently",
        "grit": "resilience",
        "hardcore": "rigorous",
        "hero": "key contributor",
        "ruthless": "highly focused",
        "kill it": "excel",
        "champion": "advocate",
        "conquer": "achieve",
        "fight": "address",
        "win": "achieve success",
        "crush": "exceed targets",
        "unstoppable": "highly motivated",
        "fearless": "courageous",
        "power": "capability",
        "backbone": "core strength",
        "sharp": "perceptive"
    },
    
    "feminine": {
        "affectionate": "approachable",
        "agreeable": "cooperative",
        "attentive": "observant",
        "collaborative": "team-oriented",
        "collaborate": "team-oriented",
        "collaborated": "worked together",
        "committed": "dedicated",
        "compassionate": "caring",
        "considerate": "thoughtful",
        "cooperative": "supportive",
        "dependable": "reliable",
        "dependent": "team-oriented",
        "emotional": "passionate",
        "empathetic": "perceptive",
        "enthusiastic": "energized",
        "gentle": "respectful",
        "honest": "transparent",
        "inclusive": "open-minded",
        "interpersonal": "people-focused",
        "kind": "respectful",
        "loyal": "dedicated",
        "modest": "measured",
        "nurturing": "supportive",
        "pleasant": "professional",
        "polite": "courteous",
        "sensitive": "perceptive",
        "supportive": "enabling",
        "sympathetic": "understanding",
        "tactful": "diplomatic",
        "tender": "considerate",
        "trustworthy": "reliable",
        "understanding": "empathetic",
        "warm": "welcoming",
        "yield": "adjust",
        "adaptable": "flexible",
        "communal": "team-centered",
        "helpful": "contributive",
        "dedicated": "committed",
        "respectful": "professional",
        "nurture": "develop",
        "sociable": "collegial",
        "relationship-oriented": "team-focused",
        "team player": "collaborative member",
        "people-oriented": "stakeholder-focused",
        "empathetic listener": "active listener",
        "gentle communicator": "considerate communicator",
        "open-minded": "inclusive",
        # New additions
        "passionate": "highly motivated",
        "inspired": "driven by purpose",
        "inspire": "motivate",
        "vulnerable": "transparent",
        "heartfelt": "sincere",
        "harmony": "alignment",
        "caring": "attentive",
        "patient": "thorough",
        "wholesome": "balanced",
        "togetherness": "team cohesion",
        "soft skills": "professional competencies",
        "personal touch": "tailored approach",
        "feeling": "assessment",
        "feelings": "perspectives",
        "transparent": "accountable",
        "uplift": "elevate",
        "thrive": "excel",
        "welcoming": "inclusive",
        "relatable": "accessible",
        "connect": "engage",
        "together": "collaboratively",
        "sharing": "distributing",
        "mindful": "deliberate",
        "balance": "manage effectively"
    }
}

def rewrite_and_optimize_resume(text, replacement_mapping, user_location):
    """
    ⚡ MERGED FUNCTION — replaces rewrite_text_with_llm() + optimize_resume_to_json()
    Single LLM call that returns BOTH:
      - rewritten_text : plain-text ATS-optimised resume + job title suggestions (for UI display)
      - json_str       : strict JSON object (for DOCX generation)
      - rewrite_ok     : bool — False when LLM was unavailable and original text was returned
    Saves 1 API key call per resume analysis (6 → 5 calls total).
    """

    # ── ALL CAPS normalization ────────────────────────────────────────────────
    # Resumes written entirely in uppercase confuse the LLM's section detection
    # and bias replacement. Normalize body lines that are fully uppercase.
    # Section headers (short ALL-CAPS labels like "EXPERIENCE") are intentionally
    # preserved — only longer uppercase body lines (sentences, bullet content) are
    # title-cased so the LLM receives readable prose.
    _normalized_lines = []
    for _line in text.split('\n'):
        _stripped = _line.strip()
        # Normalize only if: fully uppercase, more than 15 chars, no leading emoji/symbol
        if (
            _stripped
            and _stripped.isupper()
            and len(_stripped) > 15
            and _stripped[0].isalpha()
        ):
            # Title-case the content, preserving original leading whitespace
            _leading = _line[: len(_line) - len(_line.lstrip())]
            _normalized_lines.append(_leading + _stripped.title())
        else:
            _normalized_lines.append(_line)
    text = '\n'.join(_normalized_lines)
    # ─────────────────────────────────────────────────────────────────────────

    formatted_mapping = "\n".join(
        [f'- "{key}" → "{value}"' for key, value in replacement_mapping.items()]
    )

    prompt = f"""You are an enterprise-grade ATS resume optimization engine and bias-removal specialist.

Your task is to process the resume below and return TWO outputs in a single response, separated by the exact delimiter shown.

════════════════════════════════════════════════════════
OUTPUT STRUCTURE (return EXACTLY this — no deviation):
════════════════════════════════════════════════════════

===REWRITTEN_RESUME_START===
<full plain-text ATS-optimised resume here>
<followed by job title suggestions block>
===REWRITTEN_RESUME_END===

===JSON_START===
<strict JSON object here — no markdown fences, no explanation>
===JSON_END===

════════════════════════════════════════════════════════
PART 1 — PLAIN TEXT REWRITE (inside REWRITTEN_RESUME tags)
════════════════════════════════════════════════════════

ABSOLUTE RULES — NEVER VIOLATE:
• DO NOT fabricate companies, job titles, degrees, institutions, or dates.
• DO NOT invent statistics or metrics not implied by the resume content.
• DO NOT add certifications, tools, or skills absent from the resume.
• DO NOT use personal pronouns ANYWHERE — not in summary, bullets, descriptions, or any section.
  BANNED WORDS/PHRASES (zero tolerance across the ENTIRE resume):
  ✗ "I", "I am", "I have", "I worked", "I built", "I led", "I developed"
  ✗ "My", "Me", "We", "Our", "Us"
  ✗ "As a [role]..." — NEVER start any sentence with "As a"
  ✗ "With a passion for...", "Passionate about...", "I am passionate..."
  ✗ "Highly motivated", "Dynamic professional", "Results-driven", "Go-getter"
  ✗ "Seasoned", "Veteran", "Proven track record", "Rockstar", "Ninja", "Guru"
  ✗ "I believe", "I feel", "I think", "In my experience"
  ✗ "Looking for", "Seeking to", "Hoping to" — in summary (use in objective only if truly fresher)
• VOICE: Write in THIRD-PERSON IMPLICIT (omit the subject entirely).
  ✓ CORRECT: "Full Stack Developer with 3 years of experience building scalable web apps."
  ✓ CORRECT: "Developed RESTful APIs using Django and PostgreSQL, reducing latency by 40%."
  ✓ CORRECT: "Aspiring Data Scientist with hands-on ML project experience in Python and TensorFlow."
  ✗ WRONG:   "As a Full Stack Developer, I have 3 years of experience..."
  ✗ WRONG:   "I developed RESTful APIs using Django..."
  ✗ WRONG:   "I am an aspiring Data Scientist..."
• DO NOT repeat the same phrase or word across multiple sections.
• EVERY section must contain unique, non-overlapping content.
• DO NOT exaggerate experience level. If the candidate has no work experience or is a fresher/student, the summary MUST use honest framing: "Aspiring", "Entry-level", "Recent graduate". NEVER use "seasoned", "veteran", "proven track record", or imply years of experience that do not exist in the resume.
• PROFESSIONAL SUMMARY must reflect ONLY what is actually present in the resume. Do not upgrade a fresher to a professional or a junior to a senior.

YOU MAY:
✓ Strengthen bullet points with stronger action verbs and tighter phrasing.
✓ Reconstruct missing sections when clear evidence exists in the resume.
✓ Consolidate skills scattered across experience/projects into the Skills section.
✓ Infer tool proficiency when strongly implied (e.g., "built Flask API" → Python, Flask).
✓ Add plausible impact framing using "~" when the role implies measurable output.

SECTION ORDER: Contact Header → Professional Summary → Core Skills →
Work Experience → Projects → Education → Certifications & Links

CONTACT HEADER: Full Name | Job Title | Email | Phone | Location | LinkedIn URL | GitHub/Portfolio URL

PROFESSIONAL SUMMARY (2–3 sentences):
  FIRST — assess the candidate's actual experience level from the resume:
  • No experience / student / fresher → use "Aspiring [Role]" or "Recent [Degree] graduate"
  • 0–2 years → "Entry-level" or "Junior"
  • 2–5 years → "Mid-level" or just state domain + years honestly
  • 5+ years → "Senior" or "Lead" only if the resume clearly supports it

  SENTENCE STRUCTURE — STRICTLY FOLLOW:
  Sentence 1: [Seniority label] + [Job Title/Domain] + [experience framing] + [core domain/stack]
              NEVER start with "As a", "I am", "I have", or any pronoun.
              ⚠️ FRESHER RULE — if candidate has 0 or no work experience:
                NEVER write "with 0 years of experience" — this is BANNED.
                Use "specializing in" OR "with hands-on project experience in" instead.
                ✓ "Aspiring Software Developer specializing in Python, Django, and problem-solving."
                ✓ "Entry-level Java Developer with hands-on project experience in backend development."
              NEVER start with "As a", "I am", "I have", or any pronoun.
              ✓ "Mid-level Full Stack Developer with 3 years of experience building scalable web applications."
              ✓ "Entry-level Python Developer specializing in backend development and RESTful API design."
              ✓ "Aspiring Data Scientist with hands-on project experience in ML pipelines and NLP."
  Sentence 2: [Top 2–3 specific technical strengths from resume — tools, frameworks, proven skills]
              Start with a strong action noun or skill cluster. Never start with "I" or "As a".
              ✓ "Proficient in React.js, Node.js, and MongoDB with demonstrated experience in full-stack deployment."
              ✓ "Skilled in LangChain, FAISS, and LLaMA with hands-on RAG pipeline development."
  Sentence 3: [Career value proposition — what they bring or are seeking]
              ✓ "Committed to building efficient, scalable solutions that drive measurable business impact."
              ✓ "Seeking to contribute strong backend expertise to a growth-focused engineering team."

CORE SKILLS: labeled lines — Technical Skills: [...] and Professional Skills: [...]

WORK EXPERIENCE (reverse chronological):
  Job Title | Company Name | MMM YYYY – MMM YYYY
  [1-sentence role scope — NO pronouns, NO "As a", starts with a noun or action word]
  • [Action Verb] + [Task] + [Technology] + [Quantified impact]
  (3–5 bullets per role)
  Strong verbs ONLY: Architected, Engineered, Developed, Implemented, Optimized, Automated,
  Spearheaded, Deployed, Designed, Reduced, Increased, Streamlined, Led, Built, Delivered, Launched.
  NEVER: helped, assisted, worked on, involved in, responsible for, I did, I built, I led.
  EVERY bullet must start with a strong past-tense action verb — NEVER with "I", "We", or "As a".

  ⚠️ ZERO EXPERIENCE / FRESHER RULE — CRITICAL:
  If the candidate has NO work experience (no jobs, no internships):
  • DO NOT write a "Work Experience" section at all — omit it entirely.
  • DO NOT write "0 years of experience", "No experience", "N/A", or any placeholder.
  • Instead, elevate the PROJECTS section — give it prominence immediately after Core Skills.
  • In the Professional Summary, use honest framing: "Aspiring [Role]" or "Recent graduate
    with hands-on project experience in [domain]" — never imply professional tenure.
  If the candidate has only internship(s) but NO full-time roles:
  • Label the section "Internship Experience" instead of "Work Experience".
  • Present internships with the same bullet format (3–5 achievement bullets each).
  • DO NOT write "0 years of full-time experience" anywhere.

PROJECTS: Name | Tech Stack | Duration
  [1-sentence purpose — starts with a noun/verb, NO pronouns, NO "As a", NO "I built"]
  • [Achievement bullet — starts with strong action verb, NO pronouns]
  (3–5 bullets)
  ✓ "Built a real-time chat application using Socket.io and React.js."
  ✗ "I built a real-time chat application..."
  ✗ "As a developer, I created..."

EDUCATION: Degree, Major | Institution | Graduation Year | CGPA/Percentage

  ⚠️ EDUCATION FORMAT HANDLING — resumes use wildly different layouts. Handle ALL of these:

  LAYOUT VARIANTS (extract correctly from every format):
  • Standard inline:       "B.Tech (CSE) | Techno India University | 2021-2024 | CGPA: 7.94"
  • Degree first:          "B.Tech Computer Science\n  XYZ University\n  2020-2024"
  • Institution first:     "XYZ University\n  Bachelor of Technology in CSE\n  2020-2024"
  • Institution as heading + degree as bullet:
                           "Behala Aryya Vidyamandir(H.S) 2019-2021\n  ● WBCHSE (Class XII)"
                           → degree = "WBCHSE (Class XII)", institution = "Behala Aryya Vidyamandir(H.S)"
  • Board style (Class X/XII):
                           "XYZ School | CBSE | Class X | 2018 | 95%"
                           "CBSE - Class X (2019) — 91%"
                           → degree = "CBSE Class X", institution = school name if present
  • All on one line:       "B.Tech CSE, ABC University, 2024, CGPA 8.5"
  • Pursuing/Expected:     "B.Tech (CSE) — Expected 2025" / "Currently pursuing MBA from XYZ"
  • Multiple degrees same institution:
                           "ABC University\n  M.Tech 2022-2024 CGPA 8.2\n  B.Tech 2018-2022 CGPA 7.8"
                           → treat as TWO separate education entries
  • Short forms: B.E, B.Sc, M.Sc, MCA, BCA, MBA, Ph.D, Diploma — all valid degrees
  • Honours/Distinction:   "B.Tech (Hons) CSE" — preserve exactly including (Hons)

  YEAR EXTRACTION (CRITICAL — scan the ENTIRE education block):
    - Year can appear ANYWHERE: above, below, beside, or after the degree/institution
    - Accept ANY of these formats: "Oct 2021 – Jul 2024", "2021-2024", "October 2021 - July 2024",
      "Batch: 2024", "Passout: 2024", "Expected: 2025", "graduating 2025", "2024", "May 2023",
      right-aligned dates, dates below GPA line, dates on a separate line entirely
    - If only one year found → use it as graduation year
    - If a range found → preserve the full range as written (e.g. "October 2021 - July 2024")
    - NEVER leave year blank if ANY date pattern exists anywhere near the education block

  CGPA/SGPA/Percentage: preserve exactly as written, character for character.
  Also accept: "Marks: 456/500", "First Class", "Distinction", "Pass" as valid score formats.
  Include honors, distinctions, or relevant coursework if mentioned.
CERTIFICATIONS: • Name | Issuing Body | MMM YYYY

ATS FORMATTING:
• Single-column structure — no tables, columns, text boxes.
• Bullet points: "•" only.
• Section headings: ALL CAPS (e.g. PROFESSIONAL SUMMARY, CORE SKILLS, WORK EXPERIENCE).
• ⚠️ ALL CAPS applies to SECTION HEADINGS ONLY — NEVER apply ALL CAPS to body text,
  summary sentences, bullet points, skill names, or any other content.
• No emojis, no personal pronouns.

GRADE/GPA FORMATTING RULES (CRITICAL — applies to Education section):
• Preserve EXACTLY as written in the original resume — do NOT relabel or convert.
• If resume says "SGPA" → write SGPA. If resume says "CGPA" → write CGPA. NEVER swap them.
• If resume says "SGPA - 7.4" → write exactly "SGPA - 7.4". Do NOT write "CGPA - 7.4 SGPA" or "CGPA - 7.0 GPA".
• NEVER add a second label — if value already has SGPA/CGPA/GPA/Percentage prefix, do NOT add another.
• Bare numbers (e.g. "8.44") → write "CGPA: 8.44". Numbers with % → write "Percentage: 78.3%".

BIAS REPLACEMENT RULES — APPLY EXACTLY:
{formatted_mapping}

MANDATORY JOB TITLE SUGGESTIONS (append after the resume text):

### 🎯 Suggested Job Titles (Based on Resume)

Provide EXACTLY 5 job titles suited for a candidate in {user_location}.
FORMAT (STRICT — follow exactly, no extra lines, no URLs, no links):
1. **[Job Title]** — [Specific reason tied to resume evidence]
2. **[Job Title]** — [Specific reason tied to resume evidence]
3. **[Job Title]** — [Specific reason tied to resume evidence]
4. **[Job Title]** — [Specific reason tied to resume evidence]
5. **[Job Title]** — [Specific reason tied to resume evidence]

IMPORTANT: Do NOT include any URLs, hyperlinks, or 🔗 emoji. Do NOT add anything after the 5 entries.

════════════════════════════════════════════════════════
PART 2 — JSON OBJECT (inside JSON tags)
════════════════════════════════════════════════════════

Return ONLY valid JSON. No preamble, no explanation, no markdown fences.

CONTENT REWRITING — same ATS rules as Part 1 apply to all bullet fields.
Strong verbs only. Quantified impact. No pronouns. No repetition across sections.

RETURN ONLY THIS EXACT JSON STRUCTURE:
{{
  "contact": {{
    "name": "",
    "title": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin": "",
    "github": "",
    "portfolio": ""
  }},
  "summary": "",
  "skills": [],
  "soft_skills": [],
  "languages": [],
  "interests": [],
  "experience": [
    {{
      "role": "",
      "company": "",
      "duration": "",
      "description": "",
      "bullets": []
    }}
  ],
  "projects": [
    {{
      "name": "",
      "duration": "",
      "tech_stack": "",
      "url": "",
      "description": "",
      "bullets": []
    }}
  ],
  "education": [
    {{
      "degree": "",
      "institution": "",
      "year": "",
      "cgpa": "",
      "bullets": []
    }}
  ],
  "certifications": [
    {{
      "name": "",
      "issuer": "",
      "duration": ""
    }}
  ],
  "additional": [
    {{
      "name": "",
      "description": "",
      "duration": ""
    }}
  ]
}}

FIELD RULES:

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
GOLDEN RULE — APPLIES TO EVERY FIELD IN EVERY SECTION:
  • NEVER fabricate names, companies, degrees, institutions, skills, URLs, or metrics.
  • NEVER write "[Not Provided]", "N/A", "Unknown" anywhere — use "" for missing text, [] for missing arrays.
  • NEVER invent a date with zero context — if no clue exists anywhere → store "".
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
3-TIER DATE INFERENCE RULE — APPLIES TO ALL duration/year FIELDS:

  TIER 1 — EXPLICIT (highest priority):
    Date is clearly written in the resume near this entry.
    → Extract and store EXACTLY as written. Never reformat.
    → Examples: "Aug 2021 – Dec 2024", "October 2021 - July 2024", "Jul 2025", "2023", "45 days"

  TIER 2 — INFERRED FROM CONTEXT (use ONLY when strong evidence exists):
    Date is NOT written for this entry BUT can be confidently derived from surrounding resume context.
    Strong context signals (ALL must be true to infer):
      ✓ Project is explicitly described as part of an internship/job role that HAS dates → use those dates
      ✓ Project description mentions the company/internship period → use that period
      ✓ Certification is clearly tied to a dated training program mentioned elsewhere
      ✓ Education year can be derived from a "Batch of YYYY" or "Passout YYYY" written elsewhere
    → Store with "~" prefix to signal inference: e.g. "~Aug 2022 – Dec 2023"
    → NEVER infer just because two entries are near each other on the page.
    → NEVER infer from education dates for a project unless project explicitly says "college project" + education has clear dates.

  TIER 3 — UNKNOWN (no context at all):
    No date written, no strong context to infer from.
    → Store "" (empty string). The DOCX template will silently skip it.
    → NEVER guess. NEVER copy a random date from elsewhere.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

── CONTACT ──
- "contact.*" = extract exactly as written. Use "" not null for missing fields. Never invent email, phone, or URLs.

── SUMMARY ──
- "summary" = COPY THE PROFESSIONAL SUMMARY FROM PART 1 VERBATIM — every word, every sentence, nothing omitted.
  ⚠️ DO NOT rewrite, shorten, paraphrase, or summarize. The 2–3 sentence / 80-word guidance was for Part 1 ONLY.
  ⚠️ DO NOT stop after the first sentence — copy ALL sentences from Part 1's Professional Summary.
  ⚠️ If Part 1 has 3 sentences, this field MUST contain all 3 sentences.
  ⚠️ STRICT LIMIT: "summary" must contain ONLY the Professional Summary — maximum 3 sentences.
     NEVER include skills, projects, education, certifications, job titles, or any other section content here.
     If the resume is messy or unstructured, extract ONLY the summary sentences — do NOT dump the entire resume here.
     If no clear summary exists in the resume → write a clean 2–3 sentence summary from the candidate's experience.
  ⚠️ NEVER include "### 🎯 Suggested Job Titles" or any job title suggestions in this field.
  NO pronouns anywhere. No "I", "My", "As a", "I am", "I have".
  MUST be a single unbroken string — no newlines inside the JSON value.
  ✓ Correct: 2–3 sentences about the candidate's role, skills, and value proposition.
  ✗ Wrong: dumping the entire resume content, skills list, or job titles into this field.

── SKILLS ──
- "skills" = flat array of individual skill strings. Minimum 8. No duplicates. Only extract skills actually present in resume.
- "soft_skills" = professional competency phrases. Must NOT duplicate items in "skills".

── EXPERIENCE ──
- "experience" = if NO work experience AND no internships → set to [] (empty array). Never populate with placeholder roles.
  If only internships exist → include them with role/company/duration/bullets. Label role accurately (e.g. "Intern").
- "experience[].role" = exact job title as written. Never upgrade title (e.g. do NOT change "Intern" to "Developer").
- "experience[].company" = exact company name as written.
- "experience[].duration" = Apply 3-TIER DATE INFERENCE RULE.
    SCAN THE ENTIRE EXPERIENCE BLOCK for date patterns near this role (dates can appear above, below, or beside).
    Accepted formats: "Aug 2021 – Dec 2024", "August 2021 - December 2024", "Jan 2023 – Present",
    "2022–2023", "Feb 2024 - Nov 2025", month+year ranges, standalone years.
    If end date says "Present" / "present" / "current" / "Ongoing" → store as written (e.g. "Dec 2024 – Present").
    Tier 1: date found → store exactly. Tier 2: infer with "~" prefix. Tier 3: store "".
- "experience[].description" = 1-sentence role scope, unique from bullets. NO pronouns. NO "As a". NO "I".
    ✓ "Full-stack role focused on building responsive web applications using React.js and Node.js."
    ✗ "As a developer, I was responsible for building web applications."
- "experience[].bullets" = 3–5 bullets each. Strong past-tense verb + task + tech + quantified impact. NO pronouns. NO "I". NO "As a".

── PROJECTS ──
- "projects[].name" = exact project name as written.
- "projects[].duration" = Apply 3-TIER DATE INFERENCE RULE.
    SCAN THE ENTIRE PROJECT BLOCK for any date/time pattern near this project.
    Accepted formats: "Mar 2025 – Aug 2025", "October 2021 – December 2021", "3 months", "45 days",
    "Jan 2024", standalone year "2023", any time range written near the project.
    Tier 1: date found → store exactly as written (e.g. "October 2021 – December 2021").
    Tier 2: project is explicitly described as part of a dated internship/job → infer with "~" prefix
            e.g. project says "built during internship at XYZ (Aug 2022–Dec 2023)" → "~Aug 2022 – Dec 2023"
    Tier 3: no date, no context → store "".
- "projects[].tech_stack" = comma-separated technologies mentioned for this project.
- "projects[].url" = project URL/GitHub link if present. Use "" if not present.
- "projects[].description" = 1-sentence project purpose. NO pronouns. NO "I built". NO "As a". Unique, not repeated in bullets.
    ✓ "Real-time food delivery platform built with React.js, Node.js, and MongoDB."
    ✗ "I built a food delivery platform..."
- "projects[].bullets" = 3–5 bullets. Must NOT restate experience bullets. Strong past-tense verb + task + tech + impact. NO pronouns.

── EDUCATION ──
⚠️ EDUCATION IS THE MOST FORMAT-VARIABLE SECTION. Handle every layout — never skip an entry.

LAYOUT RECOGNITION RULES (apply before extracting any field):

  A) Institution-as-heading + degree-as-bullet (VERY COMMON in Indian resumes):
     "Behala Aryya Vidyamandir(H.S)  2019-2021"   ← this is the INSTITUTION
     "● WBCHSE (Class XII)"                         ← this is the DEGREE
     → degree = "WBCHSE (Class XII)", institution = "Behala Aryya Vidyamandir(H.S)", year = "2019-2021"
     NEVER skip this entry. NEVER treat the institution line as the degree.

  B) Degree-first + institution on next line:
     "B.Tech Computer Science"  → degree
     "XYZ University"           → institution

  C) Institution-first + degree on next line:
     "XYZ University"                          → institution
     "Bachelor of Technology in CSE"           → degree

  D) All inline (comma or pipe separated):
     "B.Tech CSE, ABC University, 2024, CGPA 8.5" → parse all fields from single line

  E) Multiple degrees at same institution → create SEPARATE education entries for each

  F) Class X / Class XII board entries:
     "XYZ School | CBSE | Class X | 2018 | 95%"
     → degree = "CBSE Class X", institution = "XYZ School", year = "2018", cgpa = "Percentage: 95%"
     "CBSE - Class X (2019) — 91%" (no school name)
     → degree = "CBSE Class X", institution = "", year = "2019", cgpa = "Percentage: 91%"

  G) Pursuing / Expected:
     "Currently pursuing B.Tech (CSE) from XYZ University, Expected 2026"
     → degree = "B.Tech (CSE)", institution = "XYZ University", year = "Expected 2026"

  H) Short degree forms — ALL valid, extract as-is:
     B.E, B.Sc, M.Sc, MCA, BCA, MBA, Ph.D, Diploma, Polytechnic, ITI

- "education[].degree" = extract the FULL degree name including type AND major/subject.
    If degree type and subject are on separate lines → combine them (e.g. "B.SC" + "Computer Science" → "B.SC Computer Science").
    If degree is in a bullet under the institution → extract from the bullet (Layout A above).
    NEVER leave blank if any degree-related text exists in the education block.
    NEVER treat a school/university name as the degree.
- "education[].institution" = university/college/school name exactly as written. Full name, not abbreviation.
    If institution appears as a heading above the degree bullet → still extract it correctly (Layout A).
- "education[].year" = Apply 3-TIER DATE INFERENCE RULE.
    SCAN THE ENTIRE EDUCATION BLOCK — year can appear ANYWHERE (above, below, beside, far from degree line).
    Accepted formats: "October 2021 - July 2024", "2021–2024", "Oct 2021 – Jul 2024",
    "Batch: 2024", "Passout: 2024", "Expected: 2025", "graduating 2025", standalone "2024", "May 2023",
    dates below CGPA line, dates to the right of institution, dates on completely separate lines.
    Tier 1: date found → store FULL range exactly as written (e.g. "October 2021 - July 2024").
    Tier 2: "Batch YYYY" or "Passout YYYY" found elsewhere in resume clearly tied to this degree → infer with "~" prefix.
    Tier 3: absolutely zero date/year exists anywhere → store "".
- "education[].cgpa" = Apply 3-TIER DATE INFERENCE RULE for grade (Tier 1 only — NEVER infer grades).
    SCAN THE ENTIRE EDUCATION BLOCK for any grade/score pattern.
    Normalize using THESE EXACT RULES — one label only, no duplicates:
      "7.0 GPA"              → store as "GPA: 7.0"
      "8.2 SGPA"             → store as "SGPA: 8.2"
      "SGPA 7.4"             → store as "SGPA: 7.4"
      "SGPA - 7.4"           → store as "SGPA: 7.4"
      "SGPA - 7.4 (1st Sem)" → store as "SGPA: 7.4 (1st Sem)"   ← preserve semester suffix
      "SGPA - 7.4 (Sem 2)"   → store as "SGPA: 7.4 (Sem 2)"     ← preserve semester suffix
      "CGPA 8.5"             → store as "CGPA: 8.5"
      "CGPA - 8.44"          → store as "CGPA: 8.44"
      "CGPA: 7.0 GPA"        → store as "CGPA: 7.0"   ← strip the trailing duplicate label
      "8.5/10"               → store as "CGPA: 8.5/10"
      "GPA: 3.9/4.0"         → store as "GPA: 3.9/4.0"
      "Percentage - 78.3%"   → store as "Percentage: 78.3%"
      "Percentage - 87.4%"   → store as "Percentage: 87.4%"
      "78.3%"                → store as "Percentage: 78.3%"
      "87%"                  → store as "Percentage: 87%"
      "87.4 percent"         → store as "Percentage: 87.4%"
      "8.44" (no label)      → store as "CGPA: 8.44"
    ⚠️ DASH FORMAT RULE: "LABEL - value" and "LABEL: value" are the same — both are valid.
       "SGPA - 7.4 (1st Sem)" is a valid SGPA score — NEVER drop it, NEVER treat dash as a separator meaning absence of score.
    NEVER convert SGPA to CGPA. NEVER convert GPA to CGPA. NEVER convert percentage to CGPA.
    NEVER produce duplicate labels like "CGPA: 7.0 GPA". Use "" if no score present.
- "education[].bullets" = honors, distinctions, relevant coursework, or industrial training if mentioned. Use [] if none.

── CERTIFICATIONS ──
- "certifications[].name" = exact certification name as written.
- "certifications[].issuer" = issuing organization as written. Use "" if not present.
- "certifications[].duration" = Apply 3-TIER DATE INFERENCE RULE.
    SCAN THE ENTIRE CERTIFICATION BLOCK for any date near this certification.
    Accepted formats: "July 2025 – October 2025", "Oct 2024", "2023", month+year, date ranges.
    Tier 1: date found → store exactly as written (full range if present).
    Tier 2: certification is explicitly tied to a dated training/internship program → infer with "~" prefix.
    Tier 3: no date, no context → store "".

── ADDITIONAL ──
- "additional" items MUST use object format: {{"name":"","description":"","duration":""}}.
- "additional[].duration" = Apply 3-TIER DATE INFERENCE RULE. Use "" if no context exists.

RESUME TEXT:
\"\"\"{text[:8000]}\"\"\"
"""

    # ── Smart throttle: if only 1 admin key is healthy, give it breathing room ──
    try:
        from llm_manager import load_groq_api_keys, get_healthy_keys
        _all_keys = load_groq_api_keys()
        _healthy  = get_healthy_keys(_all_keys)
        if len(_healthy) <= 1:
            time.sleep(3)   # 3-second pause to let the per-minute window recover
    except Exception:
        pass

    try:
        raw_response = call_llm(prompt, session=st.session_state)
    except Exception as _e:
        raw_response = ""

    # ── Guard: if LLM returned an error string or empty, return safe fallback ──
    _ERROR_PREFIXES = ("❌", "⚠️", "Error", "LLM unavailable", "No healthy", "rate limit", "quota")
    if not raw_response or any(raw_response.strip().startswith(p) for p in _ERROR_PREFIXES):
        # Return original text as-is + empty JSON + failure flag
        # Callers MUST check the 3rd value to know the rewrite did not happen
        return text, "", False

    # ── Parse the two sections out of the combined response ──────────────
    rewritten_text = ""
    json_str = ""

    rewrite_match = re.search(
        r"===REWRITTEN_RESUME_START===(.*?)===REWRITTEN_RESUME_END===",
        raw_response, re.DOTALL
    )
    json_match = re.search(
        r"===JSON_START===(.*?)===JSON_END===",
        raw_response, re.DOTALL
    )

    if rewrite_match:
        rewritten_text = rewrite_match.group(1).strip()
    else:
        # fallback: use everything before JSON block
        rewritten_text = raw_response.split("===JSON_START===")[0].strip()

    if json_match:
        json_str = json_match.group(1).strip()
    else:
        # fallback: try to extract JSON object from anywhere in the response
        json_fallback = re.search(r'\{[\s\S]*\}', raw_response)
        json_str = json_fallback.group(0).strip() if json_fallback else ""

    # ── Summary rescue: patch JSON summary from Part 1 if LLM truncated it ──
    # If the JSON summary is shorter than Part 1's summary, replace it with
    # the full Part 1 version. Wrapped in try/except — never breaks main flow.
    try:
        # Strip job title suggestions block from rewritten_text before rescue
        # so it never leaks into the JSON summary field
        _rewritten_clean = rewritten_text
        for _marker in ["### 🎯 Suggested Job Titles", "### Suggested Job Titles"]:
            if _marker in _rewritten_clean:
                _rewritten_clean = _rewritten_clean.split(_marker)[0].strip()
                break

        _summary_match = re.search(
            r'PROFESSIONAL SUMMARY\s*\n?(.*?)(?=\n[A-Z][A-Z\s&/]{3,}\n|\Z)',
            _rewritten_clean, re.DOTALL | re.IGNORECASE
        )
        if _summary_match:
            _part1_summary = _summary_match.group(1).strip()
            # Collapse internal newlines to single space (safe for JSON string)
            _part1_summary = re.sub(r'\s*\n\s*', ' ', _part1_summary).strip()
            # Strip leading punctuation artifacts (": ", "- " etc.) from Part 1 summary
            _part1_summary = re.sub(r'^[:\-–—|•·]+\s*', '', _part1_summary).strip()

            if _part1_summary and json_str:
                _json_summary_match = re.search(
                    r'"summary"\s*:\s*"(.*?)"(?=\s*,|\s*\})',
                    json_str, re.DOTALL
                )
                if _json_summary_match:
                    _json_summary = _json_summary_match.group(1).strip()
                    _json_words   = len(_json_summary.split())
                    _part1_words  = len(_part1_summary.split())
                    # Only patch if Part 1 has meaningfully more words → truncation happened
                    if _part1_words > _json_words + 5:
                        _escaped = _part1_summary.replace('\\', '\\\\').replace('"', '\\"')
                        json_str = re.sub(
                            r'("summary"\s*:\s*)".*?"',
                            lambda m: m.group(1) + '"' + _escaped + '"',
                            json_str, count=1, flags=re.DOTALL
                        )
    except Exception:
        pass  # Best-effort only — never break the main flow

    # ── Location fallback: inject sidebar user_location into JSON when resume has none ──
    # The LLM stores "" in contact.location when the resume has no location.
    # report_generator._clean() then renders it as "Not Provided" in the DOCX header.
    # If the user gave a preferred job location in the sidebar, patch it directly
    # into the raw JSON string here so extract_resume_json picks it up correctly.
    # This is done on the raw string (not parsed dict) to keep a single source of truth
    # before the dict is built — avoids double-patching in two places.
    if user_location and user_location.strip() and json_str:
        try:
            _loc_match = re.search(r'"location"\s*:\s*"([^"]*)"', json_str)
            if _loc_match:
                _existing_loc = _loc_match.group(1).strip()
                _is_empty = not _existing_loc or _existing_loc.lower() in (
                    "", "not provided", "[not provided]", "null", "none", "n/a", "na", "undefined"
                )
                if _is_empty:
                    _escaped_loc = user_location.strip().replace('\\', '\\\\').replace('"', '\\"')
                    json_str = re.sub(
                        r'("location"\s*:\s*)"[^"]*"',
                        lambda _m: _m.group(1) + '"' + _escaped_loc + '"',
                        json_str, count=1
                    )
        except Exception:
            pass  # never break main flow

    # ── Plain-text rewrite location patch ────────────────────────────────────
    # Only inject sidebar location when the resume genuinely has NO location.
    # Two sub-cases:
    #   1. Pipe slot exists but is empty/Not Provided  → replace with sidebar value
    #   2. Location field entirely absent from header  → insert before first URL
    # IMPORTANT: never fire if the header already contains a real location value.
    if user_location and user_location.strip() and rewritten_text:
        try:
            _header_break = rewritten_text.find('\n\n')
            _header_end   = _header_break if _header_break != -1 else min(500, len(rewritten_text))
            _header       = rewritten_text[:_header_end]
            _loc_val      = user_location.strip()

            # Sub-case 1: slot exists but is explicitly empty or "Not Provided"
            _LOC_EMPTY = re.compile(r'\|\s*(Not Provided|N/A|n/a|)\s*(\|)', re.IGNORECASE)
            if _LOC_EMPTY.search(_header):
                _header = _LOC_EMPTY.sub(
                    lambda _m: '| ' + _loc_val + ' ' + _m.group(2),
                    _header, count=1
                )
                rewritten_text = _header + rewritten_text[_header_end:]

            # Sub-case 2: location entirely absent from header.
            # Uses smart detection — checks all fields for non-URL/email/phone text,
            # checks field[0] for embedded location (e.g. "Kiran Rao, Bangalore"),
            # and skips injection if the header isn't pipe-structured at all.
            else:
                def _has_location_in_header(hdr):
                    # If fewer than 3 original pipe chars → not a real pipe header → skip
                    if hdr.count('|') < 3:
                        return True  # treat as "has location" to prevent unsafe injection
                    # Normalize newlines within header so multi-line headers split correctly
                    hdr_norm    = re.sub(r'\n', ' | ', hdr)
                    pipe_fields = [f.strip() for f in hdr_norm.split('|')]
                    # Too few fields after normalization → skip
                    if len(pipe_fields) < 4:
                        return True
                    # Check field[0] for embedded location: "Kiran Rao, Bangalore"
                    # Heuristic: comma present AND text after comma is 2+ alpha chars
                    f0 = pipe_fields[0]
                    if ',' in f0:
                        _after = f0.split(',', 1)[1].strip()
                        if len(_after) >= 2 and re.search(r'[a-zA-Z]', _after):
                            return True
                    # Check fields from index 2 onward for any non-URL/email/phone text
                    for f in pipe_fields[2:]:
                        if not f:
                            continue
                        if f.startswith('http'):
                            continue
                        if '@' in f:
                            continue
                        if re.match(r'^[\d\s\+\-\(\)\.]{7,}$', f):
                            continue
                        return True  # found a real text field → location present
                    return False

                if not _has_location_in_header(_header):
                    _url_pos = re.search(r'https?://', _header)
                    if _url_pos:
                        _insert_at = _url_pos.start()
                        _header = _header[:_insert_at] + _loc_val + ' | ' + _header[_insert_at:]
                    else:
                        _last_pipe = _header.rfind('|')
                        if _last_pipe != -1:
                            _header = _header[:_last_pipe + 1] + ' ' + _loc_val + _header[_last_pipe + 1:]
                    rewritten_text = _header + rewritten_text[_header_end:]
        except Exception:
            pass  # never break main flow

    return rewritten_text, json_str, True


# ── Thin compatibility wrappers — keep callers working without changes ────────

def rewrite_text_with_llm(text, replacement_mapping, user_location):
    """Compatibility wrapper — calls merged rewrite_and_optimize_resume()."""
    rewritten_text, _, _ok = rewrite_and_optimize_resume(text, replacement_mapping, user_location)
    return rewritten_text


def optimize_resume_to_json(raw_text: str) -> str:
    """Compatibility wrapper — calls merged rewrite_and_optimize_resume()."""
    _, json_str, _ok = rewrite_and_optimize_resume(raw_text, {}, "")
    return json_str


def _salvage_additional_str(s):
    """
    Extract name/description/duration from a leaked dict/JSON string.
    Handles JSON double-quoted, Python single-quoted, and mixed formats.
    Returns a clean dict or None.
    """
    import json as _j, re as _r
    if not s or not s.strip():
        return None
    s = s.strip()
    # Attempt 1: valid JSON double-quoted keys
    try:
        sub = _j.loads(s)
        if isinstance(sub, dict):
            n = str(sub.get("name", "") or "").strip()
            d = str(sub.get("description", "") or "").strip()
            r = str(sub.get("duration", "") or "").strip()
            if n or d:
                return {"name": n, "description": d, "duration": r}
    except Exception:
        pass
    # Attempt 2: replace single quotes -> double quotes and try JSON parse
    try:
        # Preserve escaped single quotes, swap bare ones to double quotes
        converted = s.replace("\\'", "\x01").replace("'", '"').replace("\x01", "\\'")
        sub = _j.loads(converted)
        if isinstance(sub, dict):
            n = str(sub.get("name", "") or "").strip()
            d = str(sub.get("description", "") or "").strip()
            r = str(sub.get("duration", "") or "").strip()
            if n or d:
                return {"name": n, "description": d, "duration": r}
    except Exception:
        pass
    # Attempt 3: brute-force regex — key can be single OR double quoted
    def _extract(key, text):
        for q in ('"', "'"):
            pat = q + key + q + r"\s*:\s*" + q + r"(.*?)" + q + r'(?=\s*,\s*[\'"{]|\s*\})'
            m = _r.search(pat, text, _r.DOTALL)
            if m:
                return m.group(1).strip()
        return ""
    n = _extract("name", s)
    d = _extract("description", s)
    r = _extract("duration", s)
    if n or d:
        return {"name": n, "description": d, "duration": r}
    return None


def _normalize_cgpa(raw: str) -> str:
    """
    Normalize academic score strings to a clean 'LABEL: value' format.
    Handles all input variants the LLM or raw PDF text might produce.
    Rules:
      - Never convert between score types (SGPA != CGPA != GPA != Percentage)
      - Never duplicate labels (e.g. 'CGPA: 7.0 GPA' -> 'CGPA: 7.0')
      - Numeric-only values -> assumed CGPA
      - Semester suffixes are preserved (e.g. 'SGPA: 7.4 (Semester 1)')
    """
    if not raw:
        return ""
    s = raw.strip()
    if not s:
        return ""

    # -- Pre-normalise: collapse spaces around slash e.g. "8.5 / 10" -> "8.5/10"
    s = re.sub(r'\s*/\s*', '/', s)

    # -- Pre-normalise: collapse "LABEL : value" (space before colon) ------
    s = re.sub(r'^(cgpa|sgpa|gpa|percentage)\s*:\s*', lambda m: m.group(1).upper() + ': ', s, flags=re.IGNORECASE)

    # -- Pre-normalise: collapse "LABEL - value" (dash format) → "LABEL: value"
    # e.g. "SGPA - 7.4 (1st Sem)" → "SGPA: 7.4 (1st Sem)"
    # e.g. "CGPA - 8.44" → "CGPA: 8.44"
    # e.g. "Percentage - 78.3%" → "Percentage: 78.3%"
    s = re.sub(r'^(cgpa|sgpa|gpa|percentage)\s*-\s*', lambda m: m.group(1).upper() + ': ', s, flags=re.IGNORECASE)

    # -- Already clean: CGPA/GPA/SGPA prefixes -----------------------------
    # Strip only trailing duplicate word labels, preserve semester suffixes
    for prefix in ("CGPA:", "GPA:", "SGPA:"):
        if s.upper().startswith(prefix.upper()):
            val = s[len(prefix):].strip()
            val = re.sub(r'\s+(cgpa|gpa|sgpa)\s*$', '', val, flags=re.IGNORECASE).strip()
            return f"{prefix} {val}"

    # -- Already clean: Percentage: prefix ---------------------------------
    if s.upper().startswith("PERCENTAGE:"):
        val = s[len("Percentage:"):].strip()
        val = re.sub(r'\s+percent(?:age)?\s*$', '', val, flags=re.IGNORECASE).strip()
        if not val.endswith('%'):
            val = val.rstrip('%').strip() + '%'
        return f"Percentage: {val}"

    # -- Pattern: value then label + optional semester suffix --------------
    # e.g. "8.2 SGPA", "7.4 SGPA (Semester 1)", "7.0 GPA"
    m = re.match(
        r'^(\d+(?:\.\d+)?(?:/\d+(?:\.\d+)?)?)\s*(cgpa|sgpa|gpa)((?:\s*[\(\[].*?[\)\]]|\s+Sem(?:ester)?\s*\d+)?)$',
        s, re.IGNORECASE
    )
    if m:
        val, label, suffix = m.group(1), m.group(2).upper(), m.group(3).strip()
        return f"{label}: {val}{(' ' + suffix) if suffix else ''}"

    # -- Pattern: label then value + optional semester suffix --------------
    # e.g. "CGPA 8.5", "SGPA 7.9", "SGPA 7.4 Semester 1", "SGPA 7.4 (Sem 2)"
    m = re.match(
        r'^(cgpa|sgpa|gpa)\s+(\d+(?:\.\d+)?(?:/\d+(?:\.\d+)?)?)((?:\s*[\(\[].*?[\)\]]|\s+Sem(?:ester)?\s*\d+)?)$',
        s, re.IGNORECASE
    )
    if m:
        label, val, suffix = m.group(1).upper(), m.group(2), m.group(3).strip()
        return f"{label}: {val}{(' ' + suffix) if suffix else ''}"

    # -- Pattern: percentage  e.g. "78%", "78.3%", "87.4 percent", "83 %" -
    m = re.match(r'^(\d+(?:\.\d+)?)\s*(%|percent(?:age)?)$', s, re.IGNORECASE)
    if m:
        return f"Percentage: {m.group(1)}%"

    # -- Pattern: x/y  e.g. "8.5/10", "3.9/4.0", "3.8/4", "85/100" --------
    m = re.match(r'^(\d+(?:\.\d+)?)/(\d+(?:\.\d+)?)$', s)
    if m:
        try:
            label = 'Percentage' if float(m.group(1)) > 10 else 'CGPA'
        except Exception:
            label = 'CGPA'
        return f"{label}: {m.group(1)}/{m.group(2)}"

    # -- Pattern: bare numeric (decimal or integer) e.g. "8.44", "7.0", "7", "85" --
    # Integers > 10 are treated as percentage, <= 10 as CGPA
    m = re.match(r'^(\d+(?:\.\d+)?)$', s)
    if m:
        try:
            numeric = float(m.group(1))
            return f"{'Percentage' if numeric > 10 else 'CGPA'}: {m.group(1)}"
        except Exception:
            pass

    # -- Fallback: return as-is (unknown format, don't corrupt) -----------
    return s
def extract_resume_json(llm_response: str) -> dict:
    """
    Safely extracts and parses JSON from LLM response.
    Handles markdown fences, leading/trailing text, and partial JSON.
    Returns a dict. Falls back to empty skeleton on any parse failure.
    """
    EMPTY = {
        "contact": {
            "name": "", "title": "", "email": "", "phone": "",
            "location": "", "linkedin": "", "github": "", "portfolio": ""
        },
        "summary": "",
        "skills": [],
        "soft_skills": [],
        "languages": [],
        "interests": [],
        "experience": [],
        "projects": [],
        "education": [],
        "certifications": [],
        "additional": [],
    }
    CONTACT_DEFAULTS = {
        "name": "", "title": "", "email": "", "phone": "",
        "location": "", "linkedin": "", "github": "", "portfolio": ""
    }
    if not llm_response:
        return EMPTY
    text = llm_response.strip()
    # Strip markdown fences
    text = re.sub(r'^```(?:json)?\s*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*```$', '', text)
    text = text.strip()
    # Find first { and last }
    start = text.find('{')
    end = text.rfind('}')
    if start == -1 or end == -1:
        return EMPTY
    text = text[start:end + 1]
    try:
        data = json.loads(text)
        # Ensure all top-level keys exist
        for key in EMPTY:
            if key not in data:
                data[key] = EMPTY[key]
        if not isinstance(data.get("contact"), dict):
            data["contact"] = CONTACT_DEFAULTS.copy()
        for field, default in CONTACT_DEFAULTS.items():
            if field not in data["contact"]:
                data["contact"][field] = default

        # ── Contact field rescue: fix all LLM inconsistencies ───────────────
        ct = data["contact"]

        # 1. Null/None/"null"/"undefined" → empty string for ALL contact fields
        for field in list(ct.keys()):
            v = ct[field]
            if v is None or str(v).strip() in ("null", "None", "undefined", "N/A", "n/a"):
                ct[field] = ""

        # 2. LLM put location/linkedin/github at top level instead of inside contact
        for top_key, contact_key in [
            ("location",  "location"),
            ("linkedin",  "linkedin"),
            ("github",    "github"),
            ("portfolio", "portfolio"),
            ("github_url","github"),
            ("linkedin_url","linkedin"),
            ("profile",   "linkedin"),
            ("website",   "portfolio"),
        ]:
            top_val = data.get(top_key, "") or ""
            if top_val and not ct.get(contact_key):
                ct[contact_key] = str(top_val).strip()

        # 3. LLM used alternate key names inside contact block
        ALT_KEYS = {
            "linkedin": ["linkedin_url", "linkedin_profile", "linkedin_profile_url",
                         "linkedIn", "linked_in", "profile_url", "profile"],
            "github":   ["github_url", "github_profile", "github_link",
                         "portfolio_url", "portfolio", "website", "repo"],
            "location": ["city", "address", "city_country", "current_location"],
        }
        for canonical, alts in ALT_KEYS.items():
            if not ct.get(canonical):
                for alt in alts:
                    v = ct.get(alt, "") or ""
                    if v and str(v).strip():
                        ct[canonical] = str(v).strip()
                        break

        # 4. Final null/empty cleanup — ensure every field is a string
        for field in CONTACT_DEFAULTS:
            if not isinstance(ct.get(field), str) or ct[field] is None:
                ct[field] = ""
        # Backfill missing experience fields
        _ZERO_EXP_PHRASES = re.compile(
            r'\b0\s*years?\s*(of\s*)?(experience|exp)?\b'
            r'|\bno\s+(work\s+)?experience\b'
            r'|\bN/?A\b'
            r'|\bnone\b',
            re.IGNORECASE,
        )
        for exp in data.get("experience", []):
            for f in ["role", "company", "duration", "description"]:
                if f not in exp:
                    exp[f] = ""
                # Scrub "0 years of experience" / "N/A" / "No experience" from any field
                if isinstance(exp[f], str) and _ZERO_EXP_PHRASES.search(exp[f]):
                    exp[f] = ""
            if "bullets" not in exp:
                exp["bullets"] = []
            # Scrub zero-exp phrases from bullets too
            exp["bullets"] = [
                b for b in exp["bullets"]
                if isinstance(b, str) and b.strip()
                and not _ZERO_EXP_PHRASES.search(b)
            ]
        # Drop experience entries that became entirely empty after scrubbing
        data["experience"] = [
            e for e in data.get("experience", [])
            if (e.get("role") and e["role"] not in ("", "[Not Provided]"))
            or (e.get("company") and e["company"] not in ("", "[Not Provided]"))
            or e.get("bullets")
        ]
        # Backfill missing project fields
        for proj in data.get("projects", []):
            for f in ["name", "duration", "tech_stack", "url", "description"]:
                if f not in proj:
                    proj[f] = ""
            if "bullets" not in proj:
                proj["bullets"] = []
        # Backfill missing education fields + normalize cgpa score format
        for edu in data.get("education", []):
            for f in ["degree", "institution", "year", "cgpa"]:
                if f not in edu:
                    edu[f] = ""
            # Normalize score format — fixes "CGPA: 7.0 GPA", "8.2 SGPA", "78%", etc.
            edu["cgpa"] = _normalize_cgpa(edu.get("cgpa", ""))
            if "bullets" not in edu:
                edu["bullets"] = []
        # Normalise additional — accept dicts, strings, or malformed objects
        raw_add = data.get("additional", [])
        norm_add = []
        for item in raw_add:
            if isinstance(item, dict):
                name = str(item.get("name", "") or "").strip()
                desc = str(item.get("description", "") or "").strip()
                dur  = str(item.get("duration", "") or "").strip()
                if not name and not desc:
                    continue
                norm_add.append({"name": name, "description": desc, "duration": dur})
            elif isinstance(item, str):
                s = item.strip()
                if not s or s in ("[Not Provided]",):
                    continue
                # If it looks like a leaked dict/JSON string, try to parse or salvage it
                if s.startswith("{") or ("name" in s and "description" in s):
                    salvaged = _salvage_additional_str(s)
                    if salvaged:
                        norm_add.append(salvaged)
                    # else discard entirely — do NOT render raw string
                else:
                    norm_add.append({"name": s, "description": "", "duration": ""})
        data["additional"] = norm_add

        # Normalise certifications — accept both flat strings and objects
        raw_certs = data.get("certifications", [])
        norm_certs = []
        for c in raw_certs:
            if isinstance(c, dict):
                norm_certs.append({
                    "name":     c.get("name", ""),
                    "issuer":   c.get("issuer", ""),
                    "duration": c.get("duration", ""),
                })
            elif isinstance(c, str) and c.strip():
                norm_certs.append({"name": c.strip(), "issuer": "", "duration": ""})
        data["certifications"] = norm_certs

        # ── Summary safety: collapse newlines, strip pronouns, cap overflow ──
        if isinstance(data.get("summary"), str):
            # Collapse newlines → single space (invalid in JSON strings)
            data["summary"] = re.sub(r'\s*\n\s*', ' ', data["summary"]).strip()
            # Strip leading punctuation artifacts — LLM sometimes starts with ": " or "- "
            data["summary"] = re.sub(r'^[:\-–—|•·]+\s*', '', data["summary"]).strip()
            # Strip leading banned phrases if LLM ignored the rule
            data["summary"] = re.sub(
                r'^(As a|I am|I have|I\'m)\s+', '', data["summary"],
                flags=re.IGNORECASE
            ).strip()
            # ── Overflow guard: if LLM dumped entire resume into summary ──
            # Cap at 3 sentences — anything beyond is a mis-extraction
            _summary_text = data["summary"]
            _sentences = re.split(r'(?<=[.!?])\s+', _summary_text.strip())
            if len(_sentences) > 3:
                # Keep only first 3 sentences
                data["summary"] = ' '.join(_sentences[:3]).strip()
            # Strip any job title suggestions block that leaked in
            for _leak_marker in ["### 🎯", "Suggested Job Titles", "linkedin.com/jobs/search"]:
                if _leak_marker in data["summary"]:
                    data["summary"] = data["summary"].split(_leak_marker)[0].strip()
                    # Re-cap to 3 sentences after stripping leak
                    _sentences = re.split(r'(?<=[.!?])\s+', data["summary"].strip())
                    if len(_sentences) > 3:
                        data["summary"] = ' '.join(_sentences[:3]).strip()
                    break

        return data
    except (json.JSONDecodeError, ValueError):
        return EMPTY


# ============================================================
# 📄 DOCX TEMPLATE GENERATORS — Three professional styles
# ============================================================

def _val(v) -> str:
    """Return value or 'Not Provided' placeholder — never empty, never None, never '[Not Provided]'."""
    if v is None:
        return "Not Provided"
    s = str(v).strip()
    if not s or s in ("[Not Provided]", "null", "None", "undefined"):
        return "Not Provided"
    return s


def _build_contact_header(doc, data: dict, name_size: int, name_color_rgb: tuple,
                           name_font: str, contact_font: str, contact_color_hex: str,
                           contact_size: int = 9, title_font: str = None,
                           title_size: int = 11, title_color_rgb: tuple = None,
                           separator: str = "  |  ",
                           label_color_hex: str = None,
                           accent_color_hex: str = None):
    """
    Builds the header block matching the exact sample template format:

      Line 1:  FULL NAME  (large bold centered)
      Line 2:  JOB TITLE  (smaller bold/normal centered uppercase)
      Line 3:  email  |  phone  |  LOCATION  |  linkedin_url  |  github_url
               (single pipe-separated line, centered, all fields always present)

    Every field always appears. Missing values show "Not Provided".
    No separate labeled rows. No dividers. Single clean contact line.
    """
    contact = data.get("contact", {})

    # ── Resolve contact color ────────────────────────────────────────────────
    cc = RGBColor(
        int(contact_color_hex[0:2], 16),
        int(contact_color_hex[2:4], 16),
        int(contact_color_hex[4:6], 16),
    )

    # ── ① Name ──────────────────────────────────────────────────────────────
    raw_name = contact.get("name", "") or ""
    name = raw_name if raw_name and raw_name not in ("", "[Not Provided]") else "Your Name"
    p_name = doc.add_paragraph()
    p_name.clear()
    r_name = p_name.add_run(name)
    r_name.bold = True
    r_name.font.size = Pt(name_size)
    r_name.font.name = name_font
    r_name.font.color.rgb = RGBColor(*name_color_rgb)
    p_name.alignment = 1
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(4)

    # ── ② Job Title — always shown, placeholder if missing ──────────────────
    raw_title = contact.get("title", "") or ""
    title_text = raw_title if raw_title and raw_title not in ("", "[Not Provided]") else "Job Title"
    p_title = doc.add_paragraph()
    p_title.clear()
    r_title = p_title.add_run(title_text.upper())
    r_title.font.size = Pt(title_size)
    r_title.font.name = title_font or name_font
    r_title.bold = True
    if title_color_rgb:
        r_title.font.color.rgb = RGBColor(*title_color_rgb)
    p_title.alignment = 1
    p_title.paragraph_format.space_after = Pt(4)

    # ── ③ Single pipe-separated contact line ────────────────────────────────
    # Build parts in order: email | phone | location | linkedin | github
    # Every field always included — "Not Provided" if missing
    def _clean(v):
        """Return value or 'Not Provided' — never None or empty."""
        if v is None:
            return "Not Provided"
        s = str(v).strip()
        return s if s and s not in ("[Not Provided]", "null", "None", "undefined") else "Not Provided"

    email_val    = _clean(contact.get("email", ""))
    phone_val    = _clean(contact.get("phone", ""))
    location_val = _clean(contact.get("location", ""))
    linkedin_val = _clean(contact.get("linkedin", ""))
    github_raw   = contact.get("github", "") or ""
    portfolio_raw= contact.get("portfolio", "") or ""
    github_val   = _clean(github_raw if github_raw and github_raw not in ("", "[Not Provided]") else portfolio_raw)

    _link_color = (
        int(contact_color_hex[0:2], 16),
        int(contact_color_hex[2:4], 16),
        int(contact_color_hex[4:6], 16),
    )

    p_contact = doc.add_paragraph()
    p_contact.clear()
    p_contact.alignment = 1
    p_contact.paragraph_format.space_before = Pt(2)
    p_contact.paragraph_format.space_after = Pt(6)

    def _plain_run(para, text):
        r = para.add_run(text)
        r.font.size = Pt(contact_size)
        r.font.name = contact_font
        r.font.color.rgb = cc
        return r

    SEP = "  |  "
    plain_parts = [email_val, phone_val, location_val]
    _plain_run(p_contact, SEP.join(plain_parts))

    # LinkedIn — clickable if it looks like a URL
    _plain_run(p_contact, SEP)
    if linkedin_val.startswith("http"):
        _add_hyperlink(p_contact, linkedin_val, linkedin_val,
                       font_name=contact_font, font_size=contact_size,
                       color_rgb=_link_color)
    else:
        _plain_run(p_contact, linkedin_val)

    # GitHub/Portfolio — clickable if it looks like a URL
    _plain_run(p_contact, SEP)
    if github_val.startswith("http"):
        _add_hyperlink(p_contact, github_val, github_val,
                       font_name=contact_font, font_size=contact_size,
                       color_rgb=_link_color)
    else:
        _plain_run(p_contact, github_val)

    # ── ④ Thin bottom rule below contact block ────────────────────────────
    # Signals end of header to ATS parsers and improves recruiter readability.
    from docx.oxml import OxmlElement as _OE
    from docx.oxml.ns import qn as _qn
    pPr = p_contact._p.get_or_add_pPr()
    pBdr = _OE('w:pBdr')
    btm = _OE('w:bottom')
    btm.set(_qn('w:val'), 'single')
    btm.set(_qn('w:sz'), '4')
    btm.set(_qn('w:space'), '1')
    # Use accent color if provided, else dark gray
    _border_col = accent_color_hex if accent_color_hex else "555555"
    btm.set(_qn('w:color'), _border_col)
    pBdr.append(btm)
    pPr.append(pBdr)


def _section_heading_bordered(doc, text: str, font_name: str,
                               font_size: int, bold: bool,
                               color_hex: str, border_color: str,
                               border_sz: str = "6",
                               space_before: float = 10, space_after: float = 4,
                               prefix: str = ""):
    """Universal bordered section heading."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    p.clear()
    label = f"{prefix}{text.upper()}" if prefix else text.upper()
    run = p.add_run(label)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), border_sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), border_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
    return p


def _add_hyperlink(paragraph, url: str, display_text: str, font_name: str = "Calibri",
                   font_size: int = 10, color_rgb: tuple = (0, 112, 192), underline: bool = True):
    """
    Add a real clickable hyperlink run to an existing paragraph in a python-docx document.
    Uses OOXML relationship injection — works in Word, LibreOffice, and Google Docs.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re

    # Sanitise URL — ensure it has a scheme
    url = url.strip()
    if url and not re.match(r"https?://", url, re.IGNORECASE):
        url = "https://" + url

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Font name
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    # Font size (half-points)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size * 2))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(font_size * 2))
    rPr.append(szCs)

    # Color
    color_el = OxmlElement("w:color")
    hex_color = "{:02X}{:02X}{:02X}".format(*color_rgb)
    color_el.set(qn("w:val"), hex_color)
    rPr.append(color_el)

    # Underline
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    # Style override to prevent default hyperlink style from overriding color
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.insert(0, rStyle)

    new_run.append(rPr)

    # Text
    t = OxmlElement("w:t")
    t.text = display_text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_bullet(doc, text: str, font_size: int = 10, font_name: str = "Arial",
                indent_left: int = 360, indent_hanging: int = 180,
                color_rgb: tuple = None):
    """
    Add a properly formatted ATS-compliant bullet point paragraph.
    Uses standard hanging indent matching Jobscan/Enhancv output.
    Bullet character is plain Unicode bullet (U+2022) — universally parsed by all ATS.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    # Strip any leading bullet/dash the LLM may have prepended to avoid double-bullets
    clean_text = text.strip()
    for prefix in ("\u2022", "-", "*", "\u25aa", "\u25cf"):
        if clean_text.startswith(prefix):
            clean_text = clean_text[len(prefix):].lstrip()
            break
    p = doc.add_paragraph(style="Normal")
    p.clear()
    run = p.add_run(f"\u2022  {clean_text}")
    run.font.size = Pt(font_size)
    run.font.name = font_name
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(indent_left))
    ind.set(qn('w:hanging'), str(indent_hanging))
    pPr.append(ind)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    _set_para_keep(p, keep_together=True, keep_with_next=False, widow_control=True)
    return p


def _set_para_keep(paragraph, keep_together: bool = True, keep_with_next: bool = False, widow_control: bool = True):
    """
    Apply paragraph-level page-break control to a python-docx paragraph.
    - keep_together:  prevents a paragraph from splitting across pages (good for bullets)
    - keep_with_next: keeps this paragraph on the same page as the next one (good for headings/role lines)
    - widow_control:  prevents single orphan/widow lines at top or bottom of page
    All three are set via the paragraph's XML pPr element.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    if keep_together:
        kT = OxmlElement('w:keepLines')
        pPr.append(kT)
    if keep_with_next:
        kN = OxmlElement('w:keepNext')
        pPr.append(kN)
    if widow_control:
        wC = OxmlElement('w:widowControl')
        wC.set(qn('w:val'), '1')
        pPr.append(wC)


def _add_role_line(doc, role: str, company: str, duration: str,
                   font_name: str, role_size: int = 11, meta_size: int = 9,
                   role_color: tuple = (0, 0, 0), company_color: tuple = (74, 74, 74),
                   duration_color: tuple = (128, 128, 128), separator: str = "  —  "):
    """Add role | company | duration header row for experience."""
    p = doc.add_paragraph()
    p.clear()
    if role:
        r1 = p.add_run(role)
        r1.bold = True
        r1.font.size = Pt(role_size)
        r1.font.name = font_name
        r1.font.color.rgb = RGBColor(*role_color)
    if company:
        r2 = p.add_run(f"{separator}{company}")
        r2.font.size = Pt(role_size - 1)
        r2.font.name = font_name
        r2.font.color.rgb = RGBColor(*company_color)
    if duration and duration not in ("", "[Not Provided]"):
        r3 = p.add_run(f"   [{duration}]")
        r3.italic = True
        r3.font.size = Pt(meta_size)
        r3.font.name = font_name
        r3.font.color.rgb = RGBColor(*duration_color)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
    return p


def _add_project_header(doc, name: str, duration: str, tech_stack: str, url: str,
                         font_name: str, name_size: int = 11,
                         name_color: tuple = (0, 0, 0),
                         meta_color: tuple = (74, 74, 74),
                         url_color: tuple = (30, 58, 95)):
    """Add project name | duration | tech stack | URL header row."""
    p = doc.add_paragraph()
    p.clear()
    if name:
        r1 = p.add_run(name)
        r1.bold = True
        r1.font.size = Pt(name_size)
        r1.font.name = font_name
        r1.font.color.rgb = RGBColor(*name_color)
    if duration and duration not in ("", "[Not Provided]"):
        rd = p.add_run(f"   [{duration}]")
        rd.italic = True
        rd.font.size = Pt(name_size - 2)
        rd.font.name = font_name
        rd.font.color.rgb = RGBColor(128, 128, 128)
    if tech_stack and tech_stack not in ("", "[Not Provided]"):
        r2 = p.add_run(f"  |  Tech: {tech_stack}")
        r2.font.size = Pt(name_size - 2)
        r2.font.name = font_name
        r2.font.color.rgb = RGBColor(*meta_color)
    if url and url not in ("", "[Not Provided]"):
        r3 = p.add_run(f"  |  {url}")
        r3.font.size = Pt(name_size - 2)
        r3.font.name = font_name
        r3.font.color.rgb = RGBColor(*url_color)
        r3.underline = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
    return p


def _add_education_row(doc, degree: str, institution: str, year: str,
                        edu_bullets: list, font_name: str,
                        degree_size: int = 10, meta_size: int = 9,
                        degree_color: tuple = (0, 0, 0),
                        inst_color: tuple = (74, 74, 74),
                        year_color: tuple = (128, 128, 128)):
    """Add degree | institution | year + optional bullets."""
    p = doc.add_paragraph()
    p.clear()
    if degree:
        r1 = p.add_run(degree)
        r1.bold = True
        r1.font.size = Pt(degree_size)
        r1.font.name = font_name
        r1.font.color.rgb = RGBColor(*degree_color)
    if institution and institution not in ("", "[Not Provided]"):
        r2 = p.add_run(f"  —  {institution}")
        r2.font.size = Pt(degree_size)
        r2.font.name = font_name
        r2.font.color.rgb = RGBColor(*inst_color)
    if year and year not in ("", "[Not Provided]"):
        r3 = p.add_run(f"  ({year})")
        r3.italic = True
        r3.font.size = Pt(meta_size)
        r3.font.name = font_name
        r3.font.color.rgb = RGBColor(*year_color)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
    for b in (edu_bullets or []):
        if b and b != "[Not Provided]":
            _add_bullet(doc, b, font_size=degree_size - 1, font_name=font_name)


def _render_additional(doc, data: dict, font_name: str, font_size: int,
                        heading_fn, bullet_fn,
                        name_color_rgb: tuple = (0,0,0),
                        desc_color_rgb: tuple = (80,80,80),
                        dur_color_rgb: tuple = (128,128,128)):
    """
    Render the Additional section from structured objects.
    Each item: bold name + optional [duration] on one line, description below.
    Handles dicts, flat strings, and completely skips raw leaked JSON strings.
    """
    raw = data.get("additional", [])
    # Final safety normalisation at render time
    items = []
    for item in raw:
        if isinstance(item, dict):
            name = str(item.get("name", "") or "").strip()
            desc = str(item.get("description", "") or "").strip()
            dur  = str(item.get("duration", "") or "").strip()
            # Skip if both name and desc are empty or placeholder
            if not name and not desc:
                continue
            if name in ("[Not Provided]", "") and desc in ("[Not Provided]", ""):
                continue
            items.append({"name": name, "description": desc, "duration": dur})
        elif isinstance(item, str):
            s = item.strip()
            if not s or s == "[Not Provided]":
                continue
            # Discard any raw dict/JSON leak silently
            if s.startswith("{") or ("'name'" in s and "'description'" in s) or s.startswith("["):
                continue
            items.append({"name": s, "description": "", "duration": ""})

    if not items:
        return

    heading_fn("Additional")

    for item in items:
        name = item.get("name", "")
        desc = item.get("description", "")
        dur  = item.get("duration", "")

        # Name line: bold name + italic [duration]
        p = doc.add_paragraph()
        p.clear()
        if name and name not in ("[Not Provided]", ""):
            r1 = p.add_run(name)
            r1.bold = True
            r1.font.size = Pt(font_size)
            r1.font.name = font_name
            r1.font.color.rgb = RGBColor(*name_color_rgb)
        if dur and dur not in ("[Not Provided]", ""):
            rd = p.add_run(f"   [{dur}]")
            rd.italic = True
            rd.font.size = Pt(font_size - 1)
            rd.font.name = font_name
            rd.font.color.rgb = RGBColor(*dur_color_rgb)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)

        # Description line (if present)
        if desc and desc not in ("[Not Provided]", ""):
            pd = doc.add_paragraph()
            pd.clear()
            rd2 = pd.add_run(desc)
            rd2.font.size = Pt(font_size - 1)
            rd2.font.name = font_name
            rd2.font.color.rgb = RGBColor(*desc_color_rgb)
            pd.paragraph_format.space_before = Pt(0)
            pd.paragraph_format.space_after = Pt(2)


# ─── MODERN TEMPLATE ──────────────────────────────────────────────────────────

def _chunk_skills(skills, size=7, sep="  |  "):
    """Split skills into rows of `size` joined by sep — prevents single-line overflow on long skill lists."""
    rows = [skills[i:i+size] for i in range(0, len(skills), size)]
    return ("\n" + " " * 12).join(sep.join(row) for row in rows)
def generate_modern_docx(data: dict) -> BytesIO:
    """
    Modern ATS-Optimized template — single-column, Calibri font, navy headings.
    Strictly follows ATS section ordering used by Workday, Greenhouse, and Lever:
      Header → Professional Summary → Skills → Work Experience →
      Projects → Education → Certifications → Languages → Interests → Additional

    All formatting decisions prioritize machine readability over visual design.
    No tables, no columns, no text boxes — pure linear paragraph flow for ATS parsers.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    NAVY = (0x1E, 0x3A, 0x5F)
    NAVY_HEX = "1E3A5F"
    FONT = "Calibri"
    BODY = 10

    # ── ATS RULE: Single-column header block — centered plain text, no tables ──
    _build_contact_header(
        doc, data,
        name_size=20, name_color_rgb=NAVY, name_font=FONT,
        contact_font=FONT, contact_color_hex="4A4A4A", contact_size=9,
        title_font=FONT, title_size=11, title_color_rgb=NAVY,
        accent_color_hex=NAVY_HEX,
    )

    def _heading(text):
        """
        ATS-standard section heading: ALL CAPS, bold, bottom-bordered.
        Border signals section boundary to ATS parsers without using tables.
        Matches heading labels expected by Workday/Greenhouse parsers.
        """
        _section_heading_bordered(doc, text, font_name=FONT, font_size=BODY,
                                   bold=True, color_hex=NAVY_HEX,
                                   border_color=NAVY_HEX, border_sz="6",
                                   space_before=10, space_after=4)

    def _body_para(text, italic=False, color_rgb=None):
        """Standard body paragraph — consistent 10pt Calibri, minimal spacing."""
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(BODY)
        run.font.name = FONT
        run.italic = italic
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        return p

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: PROFESSIONAL SUMMARY
    # ATS parsers (Workday, iCIMS) actively scan this section for role fit.
    # 2-3 sentences: Role identity + core competencies + value proposition.
    # ══════════════════════════════════════════════════════════════════════
    if data.get("summary") and data["summary"] not in ("", "[Not Provided]"):
        _heading("Professional Summary")
        _body_para(data["summary"])

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: CORE SKILLS
    # ATS keyword scanners parse Skills first for job description matching.
    # Categorized format (Technical / Professional) with pipe-separated values
    # is the industry standard used by Jobscan, Enhancv, and Greenhouse parsers.
    # ══════════════════════════════════════════════════════════════════════
    tech_skills = [s for s in data.get("skills", []) if s and s != "[Not Provided]"]
    soft_skills = [s for s in data.get("soft_skills", []) if s and s != "[Not Provided]"]
    if tech_skills or soft_skills:
        _heading("Core Skills")
        if tech_skills:
            p = doc.add_paragraph()
            p.clear()
            label_run = p.add_run("Technical:  ")
            label_run.bold = True
            label_run.font.size = Pt(BODY)
            label_run.font.name = FONT
            label_run.font.color.rgb = RGBColor(*NAVY)
            # Group into rows of max 7 skills for readability — ATS reads all as flat text
            skills_text = _chunk_skills(tech_skills, size=7, sep="  |  ")
            skills_run = p.add_run(skills_text)
            skills_run.font.size = Pt(BODY)
            skills_run.font.name = FONT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
        if soft_skills:
            p = doc.add_paragraph()
            p.clear()
            label_run = p.add_run("Professional:  ")
            label_run.bold = True
            label_run.font.size = Pt(BODY)
            label_run.font.name = FONT
            label_run.font.color.rgb = RGBColor(*NAVY)
            ss_run = p.add_run(_chunk_skills(soft_skills, size=7, sep="  |  "))
            ss_run.font.size = Pt(BODY)
            ss_run.font.name = FONT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: WORK EXPERIENCE
    # Industry-standard layout (Enhancv/Jobscan):
    #   Line 1: Job Title (bold, navy, larger)  |  Company Name
    #   Line 2: MMM YYYY – MMM YYYY  (italic, right-aligned)
    #   Line 3: 1-sentence role scope (italic gray) — optional
    #   Lines 4+: • Bullet points (action verb + task + tech + impact)
    # ATS parsers map Role Title and Company as separate parsed fields.
    # ══════════════════════════════════════════════════════════════════════
    valid_exp = [e for e in data.get("experience", [])
                 if (e.get("role") and e["role"] not in ("", "[Not Provided]"))
                 or (e.get("company") and e["company"] not in ("", "[Not Provided]"))]
    if valid_exp:
        _heading("Work Experience")
        for exp in valid_exp:
            _role     = exp.get("role", "")     if exp.get("role", "")     not in ("", "[Not Provided]") else ""
            _company  = exp.get("company", "")  if exp.get("company", "")  not in ("", "[Not Provided]") else ""
            _duration = exp.get("duration", "") if exp.get("duration", "") not in ("", "[Not Provided]") else ""

            # ── Row 1: Role Title (bold navy) — Company (regular gray) ──────
            p = doc.add_paragraph()
            p.clear()
            if _role:
                r_role = p.add_run(_role)
                r_role.bold = True
                r_role.font.size = Pt(BODY + 1)
                r_role.font.name = FONT
                r_role.font.color.rgb = RGBColor(*NAVY)
            if _company:
                sep = "  \u2014  " if _role else ""   # em-dash separator — ATS-safe
                r_co = p.add_run(f"{sep}{_company}")
                r_co.font.size = Pt(BODY)
                r_co.font.name = FONT
                r_co.font.color.rgb = RGBColor(74, 74, 74)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(0)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)

            # ── Row 2: Duration (italic, smaller, dark-gray) ─────────────────
            if _duration:
                p_dur = doc.add_paragraph()
                p_dur.clear()
                r_dur = p_dur.add_run(_duration)
                r_dur.italic = True
                r_dur.font.size = Pt(BODY - 1)
                r_dur.font.name = FONT
                r_dur.font.color.rgb = RGBColor(110, 110, 110)
                p_dur.paragraph_format.space_before = Pt(0)
                p_dur.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_dur, keep_together=True, keep_with_next=True, widow_control=True)

            # ── Row 3: Role scope summary (italic gray) ────────────────────
            if exp.get("description") and exp["description"] not in ("", "[Not Provided]"):
                _body_para(exp["description"], italic=True, color_rgb=(90, 90, 90))

            # ── Rows 4+: Achievement bullets ───────────────────────────────
            for b in exp.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: PROJECTS
    # Industry-standard project entry format (Jobscan/Enhancv):
    #   Line 1: Project Name (bold navy) | [Duration] (italic gray)
    #   Line 2: Tech Stack: ... (smaller, plain) | URL (plain text — no hyperlink)
    #   Line 3: 1-sentence project purpose (italic)
    #   Lines 4+: Achievement bullets
    # ATS parsers cannot follow hyperlinks — URLs must be plain text.
    # ══════════════════════════════════════════════════════════════════════
    valid_proj = [p for p in data.get("projects", [])
                  if p.get("name") and p["name"] not in ("", "[Not Provided]")]
    if valid_proj:
        _heading("Projects")
        for proj in valid_proj:
            # Line 1: Project Name + Duration
            p = doc.add_paragraph()
            p.clear()
            r_name = p.add_run(proj.get("name", ""))
            r_name.bold = True
            r_name.font.size = Pt(BODY + 1)
            r_name.font.name = FONT
            r_name.font.color.rgb = RGBColor(*NAVY)
            if proj.get("duration") and proj["duration"] not in ("", "[Not Provided]"):
                r_dur = p.add_run(f"   [{proj['duration']}]")
                r_dur.italic = True
                r_dur.font.size = Pt(BODY - 1)
                r_dur.font.name = FONT
                r_dur.font.color.rgb = RGBColor(110, 110, 110)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            # Line 2: Tech Stack (plain text) + URL (clickable hyperlink)
            has_tech = proj.get("tech_stack") and proj["tech_stack"] not in ("", "[Not Provided]")
            has_url  = proj.get("url") and proj["url"] not in ("", "[Not Provided]")
            if has_tech or has_url:
                p_meta = doc.add_paragraph()
                p_meta.clear()
                if has_tech:
                    r_tech = p_meta.add_run(f"Tech: {proj['tech_stack']}")
                    r_tech.font.size = Pt(BODY - 1)
                    r_tech.font.name = FONT
                    r_tech.font.color.rgb = RGBColor(74, 74, 74)
                if has_tech and has_url:
                    sep = p_meta.add_run("  |  ")
                    sep.font.size = Pt(BODY - 1)
                    sep.font.name = FONT
                    sep.font.color.rgb = RGBColor(74, 74, 74)
                if has_url:
                    _add_hyperlink(p_meta, proj["url"], proj["url"], font_name=FONT, font_size=BODY - 1, color_rgb=(0, 102, 204))
                p_meta.paragraph_format.space_before = Pt(0)
                p_meta.paragraph_format.space_after = Pt(2)
            if proj.get("description") and proj["description"] not in ("", "[Not Provided]"):
                _body_para(proj["description"], italic=True, color_rgb=(80, 80, 80))
            for b in proj.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: EDUCATION
    # Two-line entry: Degree + Institution on line 1, Year on line 2.
    # ATS parsers match: "Bachelor", "Master", "B.Tech", "MBA", "Ph.D".
    # ══════════════════════════════════════════════════════════════════════
    valid_edu = [e for e in data.get("education", [])
                 if e.get("degree") or e.get("institution")]
    if valid_edu:
        _heading("Education")
        for edu in valid_edu:
            p = doc.add_paragraph()
            p.clear()
            if edu.get("degree") and edu["degree"] not in ("", "[Not Provided]"):
                r_deg = p.add_run(edu["degree"])
                r_deg.bold = True
                r_deg.font.size = Pt(BODY + 1)
                r_deg.font.name = FONT
                r_deg.font.color.rgb = RGBColor(*NAVY)
            if edu.get("institution") and edu["institution"] not in ("", "[Not Provided]"):
                r_inst = p.add_run(f"  —  {edu['institution']}")
                r_inst.font.size = Pt(BODY)
                r_inst.font.name = FONT
                r_inst.font.color.rgb = RGBColor(74, 74, 74)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("year") and edu["year"] not in ("", "[Not Provided]"):
                p_yr = doc.add_paragraph()
                p_yr.clear()
                r_yr = p_yr.add_run(edu["year"])
                r_yr.italic = True
                r_yr.font.size = Pt(BODY - 1)
                r_yr.font.name = FONT
                r_yr.font.color.rgb = RGBColor(110, 110, 110)
                p_yr.paragraph_format.space_before = Pt(0)
                p_yr.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_yr, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("cgpa") and edu["cgpa"] not in ("", "[Not Provided]"):
                p_cgpa = doc.add_paragraph()
                p_cgpa.clear()
                _cgpa_val = edu['cgpa']
                # ── Smart grade label — delegates to _normalize_cgpa() ───────
                _cgpa_display = _normalize_cgpa(str(_cgpa_val).strip())
                r_cgpa = p_cgpa.add_run(_cgpa_display)
                r_cgpa.font.size = Pt(BODY - 1)
                r_cgpa.font.name = FONT
                r_cgpa.font.color.rgb = RGBColor(80, 80, 80)
                p_cgpa.paragraph_format.space_before = Pt(0)
                p_cgpa.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_cgpa, keep_together=True, keep_with_next=False, widow_control=True)
            for b in (edu.get("bullets") or []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY - 1, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: CERTIFICATIONS & LINKS
    # Each cert: Name (bold) | Issuer | Date — one bullet per credential.
    # Followed by LinkedIn / GitHub / Portfolio as plain-text URL bullets.
    # ATS systems (LinkedIn, Workday) match cert keywords directly.
    # ══════════════════════════════════════════════════════════════════════
    valid_certs = [c for c in data.get("certifications", [])
                   if isinstance(c, dict) and c.get("name") and c["name"] not in ("", "[Not Provided]")]
    contact = data.get("contact", {})
    has_links = any(contact.get(k, "") not in ("", "[Not Provided]", "Not Provided")
                    for k in ("linkedin", "github", "portfolio"))
    if valid_certs or has_links:
        _heading("Certifications & Links")
        for cert in valid_certs:
            parts = [cert["name"]]
            if cert.get("issuer") and cert["issuer"] not in ("", "[Not Provided]"):
                parts.append(cert["issuer"])
            if cert.get("duration") and cert["duration"] not in ("", "[Not Provided]"):
                parts.append(cert["duration"])
            _add_bullet(doc, "  |  ".join(parts), font_size=BODY, font_name=FONT)
        # Profile links as real clickable hyperlinks
        for label, key in [("LinkedIn", "linkedin"), ("GitHub", "github"), ("Portfolio", "portfolio")]:
            val = contact.get(key, "")
            if val and val not in ("", "[Not Provided]", "Not Provided"):
                p_link = doc.add_paragraph()
                p_link.paragraph_format.left_indent = Pt(0)
                p_link.paragraph_format.space_before = Pt(2)
                p_link.paragraph_format.space_after = Pt(2)
                label_run = p_link.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.size = Pt(BODY)
                label_run.font.name = FONT
                label_run.font.color.rgb = RGBColor(*NAVY)
                _add_hyperlink(p_link, val, val, font_name=FONT, font_size=BODY, color_rgb=(0, 102, 204))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7: LANGUAGES
    # ══════════════════════════════════════════════════════════════════════
    valid_lang = [l for l in data.get("languages", []) if l and l != "[Not Provided]"]
    if valid_lang:
        _heading("Languages")
        _body_para("  |  ".join(valid_lang))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8: INTERESTS
    # ══════════════════════════════════════════════════════════════════════
    valid_int = [i for i in data.get("interests", []) if i and i != "[Not Provided]"]
    if valid_int:
        _heading("Interests")
        _body_para("  |  ".join(valid_int))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9: ADDITIONAL (Awards, Training, Volunteering, Publications)
    # ══════════════════════════════════════════════════════════════════════
    _render_additional(doc, data, font_name=FONT, font_size=BODY,
                       heading_fn=_heading,
                       bullet_fn=lambda t: _add_bullet(doc, t, font_size=BODY, font_name=FONT),
                       name_color_rgb=NAVY, desc_color_rgb=(80, 80, 80))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── MINIMAL TEMPLATE ─────────────────────────────────────────────────────────
def generate_minimal_docx(data: dict) -> BytesIO:
    """
    Minimal ATS-Optimized template — pure black/white Arial, maximum machine readability.
    Highest ATS parse accuracy of all three templates.
    Follows identical section ordering to Modern template for ATS consistency:
      Header → Professional Summary → Skills → Work Experience →
      Projects → Education → Certifications → Languages → Interests → Additional

    No color, no decoration, no graphics — every byte serves ATS keyword matching.
    Preferred by Taleo, SmartRecruiters, and legacy HRIS systems.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(1.1)
        sec.right_margin = Inches(1.1)

    FONT = "Arial"
    BODY = 10
    BLACK_HEX = "000000"
    BLACK = (0, 0, 0)
    DARK_GRAY = (60, 60, 60)
    MID_GRAY = (100, 100, 100)

    # ── ATS RULE: Plain-text header — no color, no decoration ──
    _build_contact_header(
        doc, data,
        name_size=18, name_color_rgb=BLACK, name_font=FONT,
        contact_font=FONT, contact_color_hex="333333", contact_size=9,
        title_font=FONT, title_size=10, title_color_rgb=DARK_GRAY,
    )

    def _heading(text):
        """
        Pure black bold ALL-CAPS heading with bottom rule.
        Maximum compatibility with legacy ATS parsers that strip color.
        """
        _section_heading_bordered(doc, text, font_name=FONT, font_size=BODY,
                                   bold=True, color_hex=BLACK_HEX,
                                   border_color=BLACK_HEX, border_sz="4",
                                   space_before=10, space_after=3)

    def _body_para(text, italic=False):
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(BODY)
        run.font.name = FONT
        run.italic = italic
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        return p

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: PROFESSIONAL SUMMARY
    # ══════════════════════════════════════════════════════════════════════
    if data.get("summary") and data["summary"] not in ("", "[Not Provided]"):
        _heading("Professional Summary")
        _body_para(data["summary"])

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: CORE SKILLS
    # Comma-separated — maximum ATS keyword parser compatibility.
    # Taleo, SmartRecruiters, and legacy HRIS systems parse comma lists best.
    # Labeled "Technical" and "Professional" — matches Greenhouse/Lever field names.
    # ══════════════════════════════════════════════════════════════════════
    tech_skills = [s for s in data.get("skills", []) if s and s != "[Not Provided]"]
    soft_skills = [s for s in data.get("soft_skills", []) if s and s != "[Not Provided]"]
    if tech_skills or soft_skills:
        _heading("Core Skills")
        if tech_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Technical:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT
            skills_run = p.add_run(_chunk_skills(tech_skills, size=7, sep=", "))
            skills_run.font.size = Pt(BODY)
            skills_run.font.name = FONT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
        if soft_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Professional:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT
            ss_run = p.add_run(_chunk_skills(soft_skills, size=7, sep=", "))
            ss_run.font.size = Pt(BODY)
            ss_run.font.name = FONT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: WORK EXPERIENCE
    # Industry-standard two-line layout for maximum ATS field recognition:
    #   Line 1: Job Title (bold, black)  |  Duration (italic, gray)
    #   Line 2: Company Name (plain)
    #   Line 3: Role scope (italic, optional)
    #   Lines 4+: Bullet achievements
    # Plain black/white — preferred by Taleo and legacy HRIS parsers.
    # ══════════════════════════════════════════════════════════════════════
    valid_exp = [e for e in data.get("experience", [])
                 if (e.get("role") and e["role"] not in ("", "[Not Provided]"))
                 or (e.get("company") and e["company"] not in ("", "[Not Provided]"))]
    if valid_exp:
        _heading("Work Experience")
        for exp in valid_exp:
            _role     = exp.get("role", "")     if exp.get("role", "")     not in ("", "[Not Provided]") else ""
            _company  = exp.get("company", "")  if exp.get("company", "")  not in ("", "[Not Provided]") else ""
            _duration = exp.get("duration", "") if exp.get("duration", "") not in ("", "[Not Provided]") else ""

            # Line 1: Job Title + Duration
            p = doc.add_paragraph()
            p.clear()
            if _role:
                r1 = p.add_run(_role)
                r1.bold = True
                r1.font.size = Pt(BODY + 1)
                r1.font.name = FONT
            if _duration:
                r2 = p.add_run(f"  |  {_duration}")
                r2.italic = True
                r2.font.size = Pt(BODY - 1)
                r2.font.name = FONT
                r2.font.color.rgb = RGBColor(*MID_GRAY)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(1)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)

            # Line 2: Company Name
            if _company:
                p_co = doc.add_paragraph()
                p_co.clear()
                r_co = p_co.add_run(_company)
                r_co.font.size = Pt(BODY)
                r_co.font.name = FONT
                r_co.font.color.rgb = RGBColor(*DARK_GRAY)
                p_co.paragraph_format.space_before = Pt(0)
                p_co.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_co, keep_together=True, keep_with_next=True, widow_control=True)

            if exp.get("description") and exp["description"] not in ("", "[Not Provided]"):
                _body_para(exp["description"], italic=True)
            for b in exp.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: PROJECTS
    # Plain black project header: Name (bold) + [Duration] + tech + URL
    # Each on its own line for maximum ATS field parsing accuracy.
    # ══════════════════════════════════════════════════════════════════════
    valid_proj = [p for p in data.get("projects", [])
                  if p.get("name") and p["name"] not in ("", "[Not Provided]")]
    if valid_proj:
        _heading("Projects")
        for proj in valid_proj:
            # Line 1: Project Name + Duration
            p = doc.add_paragraph()
            p.clear()
            r1 = p.add_run(proj.get("name", ""))
            r1.bold = True
            r1.font.size = Pt(BODY + 1)
            r1.font.name = FONT
            if proj.get("duration") and proj["duration"] not in ("", "[Not Provided]"):
                rd = p.add_run(f"   [{proj['duration']}]")
                rd.italic = True
                rd.font.size = Pt(BODY - 1)
                rd.font.name = FONT
                rd.font.color.rgb = RGBColor(*MID_GRAY)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            # Line 2: Tech Stack (plain) + URL (clickable hyperlink)
            has_tech = proj.get("tech_stack") and proj["tech_stack"] not in ("", "[Not Provided]")
            has_url  = proj.get("url") and proj["url"] not in ("", "[Not Provided]")
            if has_tech or has_url:
                p_meta = doc.add_paragraph()
                p_meta.clear()
                if has_tech:
                    r_tech = p_meta.add_run(f"Tech: {proj['tech_stack']}")
                    r_tech.font.size = Pt(BODY - 1)
                    r_tech.font.name = FONT
                    r_tech.font.color.rgb = RGBColor(*DARK_GRAY)
                if has_tech and has_url:
                    sep = p_meta.add_run("  |  ")
                    sep.font.size = Pt(BODY - 1)
                    sep.font.name = FONT
                    sep.font.color.rgb = RGBColor(*DARK_GRAY)
                if has_url:
                    _add_hyperlink(p_meta, proj["url"], proj["url"], font_name=FONT, font_size=BODY - 1, color_rgb=(0, 0, 180))
                p_meta.paragraph_format.space_before = Pt(0)
                p_meta.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_meta, keep_together=True, keep_with_next=True, widow_control=True)
            if proj.get("description") and proj["description"] not in ("", "[Not Provided]"):
                _body_para(proj["description"], italic=True)
            for b in proj.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: EDUCATION
    # Degree (bold) — Institution — Year (italic)
    # ATS parsers match: "Bachelor", "Master", "B.Tech", "MBA", "Ph.D".
    # ══════════════════════════════════════════════════════════════════════
    valid_edu = [e for e in data.get("education", [])
                 if e.get("degree") or e.get("institution")]
    if valid_edu:
        _heading("Education")
        for edu in valid_edu:
            p = doc.add_paragraph()
            p.clear()
            if edu.get("degree") and edu["degree"] not in ("", "[Not Provided]"):
                r_deg = p.add_run(edu["degree"])
                r_deg.bold = True
                r_deg.font.size = Pt(BODY + 1)
                r_deg.font.name = FONT
            if edu.get("institution") and edu["institution"] not in ("", "[Not Provided]"):
                r_inst = p.add_run(f"  —  {edu['institution']}")
                r_inst.font.size = Pt(BODY)
                r_inst.font.name = FONT
                r_inst.font.color.rgb = RGBColor(*DARK_GRAY)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("year") and edu["year"] not in ("", "[Not Provided]"):
                p_yr = doc.add_paragraph()
                p_yr.clear()
                r_yr = p_yr.add_run(edu["year"])
                r_yr.italic = True
                r_yr.font.size = Pt(BODY - 1)
                r_yr.font.name = FONT
                r_yr.font.color.rgb = RGBColor(*MID_GRAY)
                p_yr.paragraph_format.space_before = Pt(0)
                p_yr.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_yr, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("cgpa") and edu["cgpa"] not in ("", "[Not Provided]"):
                p_cgpa = doc.add_paragraph()
                p_cgpa.clear()
                _cgpa_val = edu['cgpa']
                # ── Smart grade label — delegates to _normalize_cgpa() ───────
                _cgpa_display = _normalize_cgpa(str(_cgpa_val).strip())
                r_cgpa = p_cgpa.add_run(_cgpa_display)
                r_cgpa.font.size = Pt(BODY - 1)
                r_cgpa.font.name = FONT
                r_cgpa.font.color.rgb = RGBColor(80, 80, 80)
                p_cgpa.paragraph_format.space_before = Pt(0)
                p_cgpa.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_cgpa, keep_together=True, keep_with_next=False, widow_control=True)
            for b in (edu.get("bullets") or []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY - 1, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: CERTIFICATIONS & LINKS
    # Plain-text bullet per cert: Name | Issuer | Date
    # Profile URLs added as plain-text bullets — ATS-safe (no hyperlinks).
    # ══════════════════════════════════════════════════════════════════════
    valid_certs = [c for c in data.get("certifications", [])
                   if isinstance(c, dict) and c.get("name") and c["name"] not in ("", "[Not Provided]")]
    contact_min = data.get("contact", {})
    has_links_min = any(contact_min.get(k, "") not in ("", "[Not Provided]", "Not Provided")
                        for k in ("linkedin", "github", "portfolio"))
    if valid_certs or has_links_min:
        _heading("Certifications & Links")
        for cert in valid_certs:
            parts = [cert["name"]]
            if cert.get("issuer") and cert["issuer"] not in ("", "[Not Provided]"):
                parts.append(cert["issuer"])
            if cert.get("duration") and cert["duration"] not in ("", "[Not Provided]"):
                parts.append(cert["duration"])
            _add_bullet(doc, "  |  ".join(parts), font_size=BODY, font_name=FONT)
        for label, key in [("LinkedIn", "linkedin"), ("GitHub", "github"), ("Portfolio", "portfolio")]:
            val = contact_min.get(key, "")
            if val and val not in ("", "[Not Provided]", "Not Provided"):
                p_link = doc.add_paragraph()
                p_link.paragraph_format.space_before = Pt(2)
                p_link.paragraph_format.space_after = Pt(2)
                label_run = p_link.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.size = Pt(BODY)
                label_run.font.name = FONT
                label_run.font.color.rgb = RGBColor(*BLACK)
                _add_hyperlink(p_link, val, val, font_name=FONT, font_size=BODY, color_rgb=(0, 0, 180))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7: LANGUAGES
    # ══════════════════════════════════════════════════════════════════════
    valid_lang = [l for l in data.get("languages", []) if l and l != "[Not Provided]"]
    if valid_lang:
        _heading("Languages")
        _body_para(", ".join(valid_lang))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8: INTERESTS
    # ══════════════════════════════════════════════════════════════════════
    valid_int = [i for i in data.get("interests", []) if i and i != "[Not Provided]"]
    if valid_int:
        _heading("Interests")
        _body_para(", ".join(valid_int))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9: ADDITIONAL
    # ══════════════════════════════════════════════════════════════════════
    _render_additional(doc, data, font_name=FONT, font_size=BODY,
                       heading_fn=_heading,
                       bullet_fn=lambda t: _add_bullet(doc, t, font_size=BODY, font_name=FONT),
                       name_color_rgb=BLACK, desc_color_rgb=(80, 80, 80))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── CREATIVE TEMPLATE ────────────────────────────────────────────────────────
def generate_creative_docx(data: dict) -> BytesIO:
    """
    Executive ATS-Optimized template — teal/dark-navy Calibri, polished yet ATS-safe.
    Replaces design-heavy elements (decorative symbols, multi-font mixing, ◆/▌ glyphs)
    with ATS-safe equivalents that retain visual appeal without breaking parsers.
    Identical section ordering to Modern and Minimal templates for ATS consistency:
      Header → Professional Summary → Skills → Work Experience →
      Projects → Education → Certifications → Languages → Interests → Additional

    ATS COMPLIANCE NOTES:
    - All Unicode decorative characters (◆, ▌, @) removed from section headings/bullets.
    - Single font family (Calibri) used throughout body — multi-font mixing confuses some parsers.
    - Teal color is display-only; ATS parsers read plain text, not colors.
    - Georgia used ONLY for candidate name (header) — never in parseable body sections.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.95)
        sec.right_margin = Inches(0.95)

    TEAL_HEX = "0D7377"
    DARK_HEX = "14213D"
    TEAL = (0x0D, 0x73, 0x77)
    DARK = (0x14, 0x21, 0x3D)
    FONT_NAME = "Georgia"   # Name display only — ATS ignores header fonts
    FONT_BODY = "Calibri"   # Consistent body font — required for ATS parsing
    BODY = 10

    # ── Header: Name in Georgia (display), contact in Calibri (parseable) ──
    _build_contact_header(
        doc, data,
        name_size=22, name_color_rgb=DARK, name_font=FONT_NAME,
        contact_font=FONT_BODY, contact_color_hex="444444", contact_size=9,
        title_font=FONT_BODY, title_size=11, title_color_rgb=TEAL,
        accent_color_hex=TEAL_HEX,
    )

    def _heading(text):
        """
        ATS-safe teal heading — NO prefix symbols (◆/▌ break some ATS parsers).
        Teal color preserved for human readers; ATS reads underlying plain text.
        """
        _section_heading_bordered(doc, text, font_name=FONT_BODY, font_size=BODY + 1,
                                   bold=True, color_hex=TEAL_HEX,
                                   border_color=TEAL_HEX, border_sz="6",
                                   space_before=10, space_after=4)
        # NOTE: prefix="\u258c " removed — block character disrupts ATS text extraction

    def _body_para(text, italic=False, color_rgb=None):
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(BODY)
        run.font.name = FONT_BODY
        run.italic = italic
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        return p

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: PROFESSIONAL SUMMARY
    # "Profile" renamed to "Professional Summary" — standard ATS label.
    # Some older ATS systems (Taleo) fail to map "Profile" to summary field.
    # ══════════════════════════════════════════════════════════════════════
    if data.get("summary") and data["summary"] not in ("", "[Not Provided]"):
        _heading("Professional Summary")
        _body_para(data["summary"], italic=False)  # Not italic — ATS reads italic as emphasis, not content

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: CORE SKILLS
    # Teal labels + pipe-separated values — visually distinctive, ATS-safe.
    # Georgia used ONLY for name; Calibri throughout body for ATS compatibility.
    # ══════════════════════════════════════════════════════════════════════
    tech_skills = [s for s in data.get("skills", []) if s and s != "[Not Provided]"]
    soft_skills = [s for s in data.get("soft_skills", []) if s and s != "[Not Provided]"]
    if tech_skills or soft_skills:
        _heading("Core Skills")
        if tech_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Technical:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT_BODY
            lbl.font.color.rgb = RGBColor(*TEAL)
            skills_run = p.add_run(_chunk_skills(tech_skills, size=7, sep="  |  "))
            skills_run.font.size = Pt(BODY)
            skills_run.font.name = FONT_BODY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
        if soft_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Professional:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT_BODY
            lbl.font.color.rgb = RGBColor(*TEAL)
            ss_run = p.add_run(_chunk_skills(soft_skills, size=7, sep="  |  "))
            ss_run.font.size = Pt(BODY)
            ss_run.font.name = FONT_BODY
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: WORK EXPERIENCE
    # Two-line entry (Enhancv/Jobscan standard):
    #   Line 1: Job Title (bold dark-navy, 11pt)  |  Duration (italic gray)
    #   Line 2: Company Name (teal, 10pt)
    #   Line 3: Role scope (italic, optional)
    #   Lines 4+: Achievement bullets
    # ATS-safe: "@" separator replaced with " — " (em-dash).
    # ══════════════════════════════════════════════════════════════════════
    valid_exp = [e for e in data.get("experience", [])
                 if (e.get("role") and e["role"] not in ("", "[Not Provided]"))
                 or (e.get("company") and e["company"] not in ("", "[Not Provided]"))]
    if valid_exp:
        _heading("Work Experience")
        for exp in valid_exp:
            _role     = exp.get("role", "")     if exp.get("role", "")     not in ("", "[Not Provided]") else ""
            _company  = exp.get("company", "")  if exp.get("company", "")  not in ("", "[Not Provided]") else ""
            _duration = exp.get("duration", "") if exp.get("duration", "") not in ("", "[Not Provided]") else ""

            # Line 1: Job Title + Duration
            p = doc.add_paragraph()
            p.clear()
            if _role:
                r1 = p.add_run(_role)
                r1.bold = True
                r1.font.size = Pt(BODY + 1)
                r1.font.name = FONT_BODY
                r1.font.color.rgb = RGBColor(*DARK)
            if _duration:
                r3 = p.add_run(f"   {_duration}")
                r3.italic = True
                r3.font.size = Pt(BODY - 1)
                r3.font.name = FONT_BODY
                r3.font.color.rgb = RGBColor(110, 110, 110)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(1)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)

            # Line 2: Company Name (teal)
            if _company:
                p_co = doc.add_paragraph()
                p_co.clear()
                r_co = p_co.add_run(_company)
                r_co.font.size = Pt(BODY)
                r_co.font.name = FONT_BODY
                r_co.font.color.rgb = RGBColor(*TEAL)
                p_co.paragraph_format.space_before = Pt(0)
                p_co.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_co, keep_together=True, keep_with_next=True, widow_control=True)

            if exp.get("description") and exp["description"] not in ("", "[Not Provided]"):
                _body_para(exp["description"], italic=True, color_rgb=(80, 80, 80))
            for b in exp.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT_BODY)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: PROJECTS
    # Project Name (bold dark-navy, 11pt) + [Duration] (italic gray)
    # Tech Stack line (teal, 9pt) + URL as plain text (ATS-safe — no hyperlink)
    # ══════════════════════════════════════════════════════════════════════
    valid_proj = [pr for pr in data.get("projects", [])
                  if pr.get("name") and pr["name"] not in ("", "[Not Provided]")]
    if valid_proj:
        _heading("Projects")
        for proj in valid_proj:
            # Line 1: Project Name + Duration
            p = doc.add_paragraph()
            p.clear()
            r1 = p.add_run(proj.get("name", ""))
            r1.bold = True
            r1.font.size = Pt(BODY + 1)
            r1.font.name = FONT_BODY
            r1.font.color.rgb = RGBColor(*DARK)
            if proj.get("duration") and proj["duration"] not in ("", "[Not Provided]"):
                rd = p.add_run(f"   [{proj['duration']}]")
                rd.italic = True
                rd.font.size = Pt(BODY - 1)
                rd.font.name = FONT_BODY
                rd.font.color.rgb = RGBColor(110, 110, 110)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            # Line 2: Tech (teal plain text) + URL (clickable hyperlink)
            has_tech = proj.get("tech_stack") and proj["tech_stack"] not in ("", "[Not Provided]")
            has_url  = proj.get("url") and proj["url"] not in ("", "[Not Provided]")
            if has_tech or has_url:
                p_meta = doc.add_paragraph()
                p_meta.clear()
                if has_tech:
                    r_tech = p_meta.add_run(f"Tech: {proj['tech_stack']}")
                    r_tech.font.size = Pt(BODY - 1)
                    r_tech.font.name = FONT_BODY
                    r_tech.font.color.rgb = RGBColor(*TEAL)
                if has_tech and has_url:
                    sep = p_meta.add_run("  |  ")
                    sep.font.size = Pt(BODY - 1)
                    sep.font.name = FONT_BODY
                    sep.font.color.rgb = RGBColor(*TEAL)
                if has_url:
                    _add_hyperlink(p_meta, proj["url"], proj["url"], font_name=FONT_BODY, font_size=BODY - 1, color_rgb=(0, 128, 128))
                p_meta.paragraph_format.space_before = Pt(0)
                p_meta.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_meta, keep_together=True, keep_with_next=True, widow_control=True)
            if proj.get("description") and proj["description"] not in ("", "[Not Provided]"):
                _body_para(proj["description"], italic=True, color_rgb=(80, 80, 80))
            for b in proj.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT_BODY)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: EDUCATION
    # Degree (bold dark-navy) — Institution (teal) — Year (italic gray)
    # ══════════════════════════════════════════════════════════════════════
    valid_edu = [e for e in data.get("education", [])
                 if e.get("degree") or e.get("institution")]
    if valid_edu:
        _heading("Education")
        for edu in valid_edu:
            p = doc.add_paragraph()
            p.clear()
            if edu.get("degree") and edu["degree"] not in ("", "[Not Provided]"):
                r1 = p.add_run(edu["degree"])
                r1.bold = True
                r1.font.size = Pt(BODY + 1)
                r1.font.name = FONT_BODY
                r1.font.color.rgb = RGBColor(*DARK)
            if edu.get("institution") and edu["institution"] not in ("", "[Not Provided]"):
                r2 = p.add_run(f"  —  {edu['institution']}")
                r2.font.size = Pt(BODY)
                r2.font.name = FONT_BODY
                r2.font.color.rgb = RGBColor(*TEAL)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("year") and edu["year"] not in ("", "[Not Provided]"):
                p_yr = doc.add_paragraph()
                p_yr.clear()
                r3 = p_yr.add_run(edu["year"])
                r3.italic = True
                r3.font.size = Pt(BODY - 1)
                r3.font.name = FONT_BODY
                r3.font.color.rgb = RGBColor(110, 110, 110)
                p_yr.paragraph_format.space_before = Pt(0)
                p_yr.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_yr, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("cgpa") and edu["cgpa"] not in ("", "[Not Provided]"):
                p_cgpa = doc.add_paragraph()
                p_cgpa.clear()
                _cgpa_val = edu['cgpa']
                # ── Smart grade label — delegates to _normalize_cgpa() ───────
                _cgpa_display = _normalize_cgpa(str(_cgpa_val).strip())
                r_cgpa = p_cgpa.add_run(_cgpa_display)
                r_cgpa.font.size = Pt(BODY - 1)
                r_cgpa.font.name = FONT_BODY
                r_cgpa.font.color.rgb = RGBColor(80, 80, 80)
                p_cgpa.paragraph_format.space_before = Pt(0)
                p_cgpa.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_cgpa, keep_together=True, keep_with_next=False, widow_control=True)
            for b in (edu.get("bullets") or []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY - 1, font_name=FONT_BODY)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: CERTIFICATIONS & LINKS
    # Cert bullets: Name | Issuer | Date  (teal heading, plain bullets)
    # Profile URLs as plain-text bullets — no hyperlinks (ATS-safe).
    # ══════════════════════════════════════════════════════════════════════
    valid_certs = [c for c in data.get("certifications", [])
                   if isinstance(c, dict) and c.get("name") and c["name"] not in ("", "[Not Provided]")]
    contact_exec = data.get("contact", {})
    has_links_exec = any(contact_exec.get(k, "") not in ("", "[Not Provided]", "Not Provided")
                         for k in ("linkedin", "github", "portfolio"))
    if valid_certs or has_links_exec:
        _heading("Certifications & Links")
        for cert in valid_certs:
            parts = [cert["name"]]
            if cert.get("issuer") and cert["issuer"] not in ("", "[Not Provided]"):
                parts.append(cert["issuer"])
            if cert.get("duration") and cert["duration"] not in ("", "[Not Provided]"):
                parts.append(cert["duration"])
            _add_bullet(doc, "  |  ".join(parts), font_size=BODY, font_name=FONT_BODY)
        for label, key in [("LinkedIn", "linkedin"), ("GitHub", "github"), ("Portfolio", "portfolio")]:
            val = contact_exec.get(key, "")
            if val and val not in ("", "[Not Provided]", "Not Provided"):
                p_link = doc.add_paragraph()
                p_link.paragraph_format.space_before = Pt(2)
                p_link.paragraph_format.space_after = Pt(2)
                label_run = p_link.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.size = Pt(BODY)
                label_run.font.name = FONT_BODY
                label_run.font.color.rgb = RGBColor(*TEAL)
                _add_hyperlink(p_link, val, val, font_name=FONT_BODY, font_size=BODY, color_rgb=(0, 128, 128))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7: LANGUAGES
    # ══════════════════════════════════════════════════════════════════════
    valid_lang = [l for l in data.get("languages", []) if l and l != "[Not Provided]"]
    if valid_lang:
        _heading("Languages")
        _body_para("  |  ".join(valid_lang))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8: INTERESTS
    # ══════════════════════════════════════════════════════════════════════
    valid_int = [i for i in data.get("interests", []) if i and i != "[Not Provided]"]
    if valid_int:
        _heading("Interests")
        _body_para("  |  ".join(valid_int))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9: ADDITIONAL
    # ══════════════════════════════════════════════════════════════════════
    _render_additional(doc, data, font_name=FONT_BODY, font_size=BODY,
                       heading_fn=_heading,
                       bullet_fn=lambda t: _add_bullet(doc, t, font_size=BODY, font_name=FONT_BODY),
                       name_color_rgb=DARK, desc_color_rgb=(80, 80, 80))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ============================================================
# 🔍 RESUME ANALYSIS MODULE — Job Title Suggestions (UI only)
# ============================================================
# NOTE: rewrite_text_with_llm() above is the Analysis Module.
# It generates job title suggestions + rewritten text for DISPLAY ONLY.
# It is NEVER used for DOCX generation.
# DOCX generation uses optimize_resume_to_json() + extract_resume_json() exclusively.


def rewrite_and_highlight(text, replacement_mapping, user_location):
    highlighted_text = text
    masculine_count, feminine_count = 0, 0
    detected_masculine_words, detected_feminine_words = [], []
    matched_spans = []

    masculine_words = sorted(gender_words["masculine"], key=len, reverse=True)
    feminine_words = sorted(gender_words["feminine"], key=len, reverse=True)

    def span_overlaps(start, end):
        return any(s < end and e > start for s, e in matched_spans)

    # Highlight and count masculine words
    for word in masculine_words:
        pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
        for match in pattern.finditer(highlighted_text):
            start, end = match.span()
            if span_overlaps(start, end):
                continue

            word_match = match.group(0)
            colored = f"<span style='color:blue;'>{word_match}</span>"

            # Replace word in the highlighted text
            highlighted_text = highlighted_text[:start] + colored + highlighted_text[end:]
            shift = len(colored) - len(word_match)
            matched_spans = [(s if s < start else s + shift, e if s < start else e + shift) for s, e in matched_spans]
            matched_spans.append((start, start + len(colored)))

            masculine_count += 1

            # Get sentence context and highlight
            sentence_match = re.search(r'([^.]*?\b' + re.escape(word_match) + r'\b[^.]*\.)', text, re.IGNORECASE)
            if sentence_match:
                sentence = sentence_match.group(1).strip()
                colored_sentence = re.sub(
                    rf'\b({re.escape(word_match)})\b',
                    r"<span style='color:blue;'>\1</span>",
                    sentence,
                    flags=re.IGNORECASE
                )
                detected_masculine_words.append({
                    "word": word_match,
                    "sentence": colored_sentence
                })
            break  # Only one match per word

    # Highlight and count feminine words
    for word in feminine_words:
        pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
        for match in pattern.finditer(highlighted_text):
            start, end = match.span()
            if span_overlaps(start, end):
                continue

            word_match = match.group(0)
            colored = f"<span style='color:red;'>{word_match}</span>"

            # Replace word in the highlighted text
            highlighted_text = highlighted_text[:start] + colored + highlighted_text[end:]
            shift = len(colored) - len(word_match)
            matched_spans = [(s if s < start else s + shift, e if s < start else e + shift) for s, e in matched_spans]
            matched_spans.append((start, start + len(colored)))

            feminine_count += 1

            # Get sentence context and highlight
            sentence_match = re.search(r'([^.]*?\b' + re.escape(word_match) + r'\b[^.]*\.)', text, re.IGNORECASE)
            if sentence_match:
                sentence = sentence_match.group(1).strip()
                colored_sentence = re.sub(
                    rf'\b({re.escape(word_match)})\b',
                    r"<span style='color:red;'>\1</span>",
                    sentence,
                    flags=re.IGNORECASE
                )
                detected_feminine_words.append({
                    "word": word_match,
                    "sentence": colored_sentence
                })
            break  # Only one match per word

    # ⚡ Single LLM call — returns BOTH plain-text rewrite (for UI) AND JSON (for DOCX)
    # Replaces the old rewrite_text_with_llm call which discarded the JSON half.
    rewritten_text, json_str, rewrite_ok = rewrite_and_optimize_resume(
        text,
        replacement_mapping["masculine"] | replacement_mapping["feminine"],
        user_location
    )

    # Return rewrite_ok as 8th value so the caller can show a warning
    # when the LLM failed and rewritten_text is just the original unchanged text.
    return highlighted_text, rewritten_text, masculine_count, feminine_count, detected_masculine_words, detected_feminine_words, json_str, rewrite_ok

# ✅ Enhanced Grammar evaluation using LLM with suggestions
def get_grammar_score_with_llm(text, max_score=5):
    grammar_prompt = f"""
You are a senior HR language quality specialist and professional resume reviewer with 15+ years of experience evaluating resumes for Fortune 500 companies.

Analyze the following resume text across FIVE dimensions and provide an overall language quality score:

**EVALUATION DIMENSIONS:**
1. **Grammar & Mechanics** — Correct grammar, punctuation, subject-verb agreement, tense consistency
2. **Clarity & Conciseness** — Ideas expressed directly; no filler words or redundancy
3. **Professional Tone** — Appropriate formality, no informal slang or casual phrasing
4. **Action Verb Usage** — Starts bullet points with strong, quantifiable action verbs (e.g., "Led", "Engineered", "Reduced")
5. **ATS Language Alignment** — Industry-standard terminology, keyword density, no keyword stuffing

**SCORING SCALE (out of {max_score}):**
- {max_score}: Exceptional — Flawless grammar, powerful action verbs, crystal-clear and professional throughout
- {max_score-1}: Very Good — Minor stylistic issues; highly professional and readable
- {max_score-2}: Good — Some grammar or clarity issues but largely professional and effective
- {max_score-3}: Fair — Noticeable grammar, tone, or clarity problems that could affect readability
- {max_score-4}: Poor — Multiple errors affecting professional impression; needs significant editing
- 0-1: Very Poor — Significant language issues that would cause ATS rejection or recruiter dismissal

**IMPORTANT:** Be balanced — a technically competent resume with minor grammar issues should not be harshly penalized. Focus on overall professional impression.

Return EXACTLY in this format (no extra text):

Score: <number>
Feedback: <single sentence summarizing overall language quality and tone>
Suggestions:
- <Actionable suggestion 1 with example if helpful>
- <Actionable suggestion 2 with example if helpful>
- <Actionable suggestion 3 with example if helpful>
- <Actionable suggestion 4 with example if helpful>
- <Actionable suggestion 5 with example if helpful>

---
{text}
---
"""

    try:
        response = call_llm(grammar_prompt, session=st.session_state).strip()
    except Exception:
        response = ""

    # Guard: LLM error string or empty → return safe defaults
    _ERROR_PREFIXES = ("❌", "⚠️", "Error", "LLM unavailable", "No healthy", "rate limit", "quota")
    if not response or any(response.startswith(p) for p in _ERROR_PREFIXES):
        return max(0, min(max_score, max(3, max_score - 2))), "Language quality appears adequate for professional communication.", []

    score_match = re.search(r"Score:\s*(\d+)", response)
    feedback_match = re.search(r"Feedback:\s*(.+)", response)
    suggestions = re.findall(r"- (.+)", response)

    score = int(score_match.group(1)) if score_match else max(0, min(max_score, max(3, max_score - 2)))
    feedback = feedback_match.group(1).strip() if feedback_match else "Language quality appears adequate for professional communication."
    return score, feedback, suggestions

# ✅ Main ATS Evaluation Function
def ats_percentage_score(
    resume_text,
    job_description,
    job_title="Unknown",
    logic_profile_score=None,
    edu_weight=20,
    exp_weight=35,
    skills_weight=30,
    lang_weight=5,
    keyword_weight=10,
    format_data=None,       # pass pre-computed format check result
    resume_domain=None,     # FIX: accept pre-detected domain from main thread
    job_domain=None,        # FIX: accept pre-detected domain from main thread
    resume_depth=None,      # FIX: accept pre-detected depth — avoids reading stale "moderate" from session state
):
    import datetime

    _valid_domains = DOMAIN_VALID_LIST

    # FIX: use pre-detected value if passed in, only call LLM if not already set
    # When called from the parallel thread, resume_domain is set by the main thread
    # so no LLM call fires inside the thread (thread-safe).
    # ── BUG 3 FIX: Extract candidate title from resume for keyword title_overrides ──
    # The title_overrides block in detect_domain_from_title_and_description weights
    # title matches at 5× — but passing "" skips all of them. We extract the
    # most likely job title from the resume header/summary to use as the title arg.
    def _extract_resume_title(text: str) -> str:
        """Extract the most likely job title from resume text (first 800 chars)."""
        import re as _re_t
        header = text[:800].lower()
        # Common patterns: "Software Engineer | Python" or "Senior Data Analyst"
        # Check for known role keywords in the header
        _role_patterns = [
            r"(?:^|\n)([a-z][a-z/ |-]{4,40}(?:engineer|developer|analyst|scientist|"
            r"architect|manager|designer|specialist|consultant|lead|intern|fresher))",
        ]
        for pat in _role_patterns:
            m = _re_t.search(pat, header)
            if m:
                return m.group(1).strip()
        # Fallback: return first non-blank line after name (usually title/role line)
        lines = [l.strip() for l in text[:400].split("\n") if l.strip()]
        if len(lines) >= 2:
            candidate = lines[1]
            if len(candidate) < 60 and any(w in candidate.lower() for w in
               ["engineer","developer","analyst","designer","manager","scientist",
                "architect","intern","fresher","specialist","consultant"]):
                return candidate
        return ""

    _resume_title_hint = _extract_resume_title(resume_text)
    _resume_cache_key = f"resume_domain_{hash(resume_text[:500])}"
    if resume_domain is None and _resume_cache_key not in st.session_state:
        # ── Use shared prompt builder — single source of truth in db_manager ──
        # Pass extracted title hint into the prompt so LLM has the role context
        _resume_domain_prompt = build_resume_domain_prompt(
            resume_text,
            title_hint=_resume_title_hint
        )
        try:
            _raw = call_llm(_resume_domain_prompt, session=st.session_state).strip()

            # ── Parse domain + depth from LLM response ─────────────────────
            # LLMs sometimes add preamble, punctuation, or capitalization.
            # We strip all of those before validating.
            _domain_line = ""
            _depth_raw   = ""
            for _line in _raw.splitlines():
                _line = _line.strip()
                if _line.lower().startswith("domain:"):
                    _domain_line = _line.split(":", 1)[1].strip().rstrip(".")
                elif _line.lower().startswith("depth:"):
                    _depth_raw = _line.split(":", 1)[1].strip().lower().rstrip(".")

            # Normalise depth — strip whitespace, ignore case, strip trailing punct
            _depth_val = _depth_raw if _depth_raw in ("shallow", "moderate", "deep") else ""

            # ── Retry if depth is missing or unparseable ────────────────────
            if not _depth_val:
                import logging as _logging
                _logging.getLogger(__name__).warning(
                    "Depth line missing or unrecognised in LLM response. "
                    f"Raw depth token: {repr(_depth_raw)!r}. "
                    f"Full response (first 300 chars): {_raw[:300]!r}. "
                    "Retrying with stricter prompt."
                )
                _retry_prompt = (
                    f"{_resume_domain_prompt}\n\n"
                    "REMINDER: Your previous response was missing or had an invalid Depth line.\n"
                    "You MUST output exactly two lines and nothing else:\n"
                    "Domain: <domain>\n"
                    "Depth: shallow   (or moderate, or deep)\n"
                    "Do not write anything before or after these two lines."
                )
                try:
                    _raw2 = call_llm(_retry_prompt, session=st.session_state).strip()
                    for _line in _raw2.splitlines():
                        _line = _line.strip()
                        if _line.lower().startswith("domain:") and not _domain_line:
                            _domain_line = _line.split(":", 1)[1].strip().rstrip(".")
                        elif _line.lower().startswith("depth:"):
                            _depth_raw2 = _line.split(":", 1)[1].strip().lower().rstrip(".")
                            if _depth_raw2 in ("shallow", "moderate", "deep"):
                                _depth_val = _depth_raw2
                                break
                except Exception:
                    pass  # retry failed — fall through to keyword fallback below

            # ── Final safety net: keyword-based depth inference ─────────────
            # Only fires if BOTH the primary call and the retry failed to produce
            # a valid depth token. Uses simple resume-text signals so at least
            # shallow/deep extremes are caught correctly.
            if not _depth_val:
                _rt_lower = resume_text.lower()
                _has_fulltime = any(kw in _rt_lower for kw in [
                    "years of experience", "yrs of experience", "full-time", "full time",
                    "employed", "employment", "promoted", "production system",
                ])
                _has_quantified = bool(re.search(
                    r'\b(\d+[%x]|\d+\s*(users?|customers?|ms|seconds?|requests?|'
                    r'rpm|latency|revenue|million|billion|k\b))', _rt_lower
                ))
                _has_projects = len(re.findall(
                    r'\b(project|built|developed|implemented|designed)\b', _rt_lower
                )) >= 3
                _has_internship = bool(re.search(
                    r'\b(internship|intern\b|trainee)\b', _rt_lower
                ))
                _virtual_only = bool(re.search(
                    r'\b(aicte virtual|oasis infobyte|internshala|certificate program)\b', _rt_lower
                ))

                if _has_fulltime or _has_quantified:
                    _depth_val = "deep"
                elif (_has_internship and not _virtual_only) or _has_projects:
                    _depth_val = "moderate"
                else:
                    _depth_val = "shallow"

                import logging as _logging
                _logging.getLogger(__name__).warning(
                    f"Depth inferred via keyword fallback as '{_depth_val}' "
                    f"(fulltime={_has_fulltime}, quantified={_has_quantified}, "
                    f"projects={_has_projects}, internship={_has_internship})."
                )

            if _domain_line in _valid_domains:
                st.session_state[_resume_cache_key]            = _domain_line
                st.session_state[_resume_cache_key + "_depth"] = _depth_val
            else:
                _kw_fallback = (db_manager.detect_domain_with_confidence(_resume_title_hint, resume_text[:3000]).get("domain") or db_manager.detect_domain_from_title_and_description(_resume_title_hint, resume_text[:3000]))
                st.session_state[_resume_cache_key]            = _kw_fallback if _kw_fallback != "Unclassified" else "Software Engineering"
                st.session_state[_resume_cache_key + "_depth"] = _depth_val
        except Exception:
            # ── LLM call itself failed — infer depth from resume text ───────
            _rt_lower  = resume_text.lower()
            _has_ft    = any(kw in _rt_lower for kw in [
                "years of experience", "yrs of experience", "full-time", "full time",
                "employed", "employment", "promoted", "production system",
            ])
            _has_qt    = bool(re.search(
                r'\b(\d+[%x]|\d+\s*(users?|customers?|ms|seconds?|requests?|'
                r'rpm|latency|revenue|million|billion|k\b))', _rt_lower
            ))
            _has_pj    = len(re.findall(
                r'\b(project|built|developed|implemented|designed)\b', _rt_lower
            )) >= 3
            _has_int   = bool(re.search(r'\b(internship|intern\b|trainee)\b', _rt_lower))
            _virt_only = bool(re.search(
                r'\b(aicte virtual|oasis infobyte|internshala|certificate program)\b', _rt_lower
            ))
            if _has_ft or _has_qt:
                _inferred_depth = "deep"
            elif (_has_int and not _virt_only) or _has_pj:
                _inferred_depth = "moderate"
            else:
                _inferred_depth = "shallow"

            try:
                _kw_fallback = (db_manager.detect_domain_with_confidence(_resume_title_hint, resume_text[:3000]).get("domain") or db_manager.detect_domain_from_title_and_description(_resume_title_hint, resume_text[:3000]))
                st.session_state[_resume_cache_key]            = _kw_fallback if _kw_fallback != "Unclassified" else "Software Engineering"
                st.session_state[_resume_cache_key + "_depth"] = _inferred_depth
            except Exception:
                st.session_state[_resume_cache_key]            = "Software Engineering"
                st.session_state[_resume_cache_key + "_depth"] = _inferred_depth

    if resume_domain is None:
        resume_domain = st.session_state.get(_resume_cache_key, "Software Engineering")

    # ── Depth: use explicitly passed value first, then session state, never hardcode ──
    if resume_depth and resume_depth in ("shallow", "moderate", "deep"):
        _depth_str = resume_depth
    else:
        _depth_str = st.session_state.get(_resume_cache_key + "_depth")
        if _depth_str not in ("shallow", "moderate", "deep"):
            # Session state has nothing valid — infer from resume text directly
            _rt = resume_text.lower()
            _has_ft = any(kw in _rt for kw in [
                "years of experience", "yrs of experience", "full-time", "full time",
                "employed", "employment", "promoted", "production system",
            ])
            _has_qt = bool(re.search(
                r'\b(\d+[%x]|\d+\s*(users?|customers?|ms|seconds?|requests?|'
                r'rpm|latency|revenue|million|billion|k\b))', _rt
            ))
            _has_pj = len(re.findall(
                r'\b(project|built|developed|implemented|designed)\b', _rt
            )) >= 3
            _has_int = bool(re.search(r'\b(internship|intern\b|trainee)\b', _rt))
            _virt    = bool(re.search(
                r'\b(aicte virtual|oasis infobyte|internshala|certificate program)\b', _rt
            ))
            if _has_ft or _has_qt:
                _depth_str = "deep"
            elif (_has_int and not _virt) or _has_pj:
                _depth_str = "moderate"
            else:
                _depth_str = "shallow"
    _depth_score = {"shallow": 0.4, "moderate": 0.7, "deep": 1.0}.get(_depth_str, 0.7)

    # ── JOB DOMAIN: use pre-detected value if passed in, else detect here ──
    # If JD is non-English → skip LLM domain detection, use keyword fallback directly
    _jd_cache_key = f"jd_domain_{hash(job_description[:500])}"
    _jd_non_english = False
    if job_description:
        _jd_alpha = [c for c in job_description if c.isalpha()]
        if _jd_alpha:
            _jd_ascii_ratio = sum(1 for c in _jd_alpha if ord(c) < 128) / len(_jd_alpha)
            _jd_non_english = _jd_ascii_ratio < 0.70
    if _jd_non_english and job_domain is None:
        try:
            _jd_kw = (db_manager.detect_domain_with_confidence(job_title, job_description[:3000]).get("domain") or db_manager.detect_domain_from_title_and_description(job_title, job_description[:3000]))
            st.session_state[_jd_cache_key] = _jd_kw if _jd_kw != "Unclassified" else "Software Engineering"
        except Exception:
            st.session_state[_jd_cache_key] = "Software Engineering"
    if job_domain is None and _jd_cache_key not in st.session_state:
        _jd_domain_prompt = build_jd_domain_prompt(job_title, job_description)
        try:
            _raw_jd = call_llm(_jd_domain_prompt, session=st.session_state).strip()
            _jd_domain_line = ""
            for _line in _raw_jd.splitlines():
                _line = _line.strip()
                if _line.lower().startswith("domain:"):
                    _jd_domain_line = _line.split(":", 1)[1].strip()
                    break
            if _jd_domain_line in _valid_domains:
                st.session_state[_jd_cache_key] = _jd_domain_line
            else:
                _jd_kw = (db_manager.detect_domain_with_confidence(job_title, job_description[:3000]).get("domain") or db_manager.detect_domain_from_title_and_description(job_title, job_description[:3000]))
                st.session_state[_jd_cache_key] = _jd_kw if _jd_kw != "Unclassified" else "Software Engineering"
        except Exception:
            try:
                _jd_kw = (db_manager.detect_domain_with_confidence(job_title, job_description[:3000]).get("domain") or db_manager.detect_domain_from_title_and_description(job_title, job_description[:3000]))
                st.session_state[_jd_cache_key] = _jd_kw if _jd_kw != "Unclassified" else "Software Engineering"
            except Exception:
                st.session_state[_jd_cache_key] = "Software Engineering"

    # Use passed-in value if provided, otherwise use session state value
    if job_domain is None:
        job_domain = st.session_state.get(_jd_cache_key, "Software Engineering")

    similarity_score = get_domain_similarity(resume_domain, job_domain)

    # Grammar defaults — overwritten by values parsed from the ATS prompt response below
    grammar_score       = 0          # stays 0 if LLM section fails to parse — no phantom score
    grammar_feedback    = "Language quality appears adequate for professional communication."
    grammar_suggestions = []

    # ✅ Depth-weighted domain penalty
    # Same domain (similarity >= 0.90) → always zero penalty regardless of depth
    # Different domain → depth multiplier fires: shallow=0.4, moderate=0.7, deep=1.0
    MAX_DOMAIN_PENALTY = 15
    if similarity_score >= 0.90:
        domain_penalty       = 0
        effective_similarity = similarity_score
    else:
        effective_similarity = similarity_score * _depth_score
        domain_penalty = round((1 - effective_similarity) * MAX_DOMAIN_PENALTY)

    # ✅ Optional profile score note
    logic_score_note = (
        f"\n\nOptional Note: The system also calculated a logic-based profile score of {logic_profile_score}/100 "
        f"based on resume length, experience, and skills."
        if logic_profile_score else ""
    )

    # ✅ FIXED: Stable education scoring with 2025 cutoff
    current_year = datetime.datetime.now().year

    current_month = datetime.datetime.now().month

    # ✅ UPDATED: Stable education scoring with priority degrees minimum
    prompt = f"""
You are a senior ATS (Applicant Tracking System) Evaluator and Technical Recruiter with 15+ years of experience at top-tier tech firms.
Your evaluation must be rigorous, consistent, evidence-based, and match industry-standard hiring benchmarks.

You specialize in: AI/ML, Blockchain, Cloud Computing, Data Engineering, Software Development, DevOps, and Cybersecurity roles.

═══════════════════════════════════════════════════
🎯 EVALUATION PHILOSOPHY
═══════════════════════════════════════════════════
- Score based on EVIDENCE found in the resume — not assumptions
- Reward quantified achievements (numbers, percentages, scale)
- Credit projects, GitHub, hackathons, Kaggle, open-source contributions, certifications
- Penalize vague claims without evidence ("good communication skills")
- Recognize career stage: entry-level vs senior vs lead
- Prioritize recency: skills/experience from the last 3 years matter most
- Be encouraging but calibrated: do not inflate scores without evidence

═══════════════════════════════════════════════════
📐 SCORING FRAMEWORK
═══════════════════════════════════════════════════

**🎓 Education Score ({edu_weight} points max):**

PRIORITY RULE — Minimum {int(edu_weight * 0.75)} pts for these degrees (completed OR pursuing):
  • BSc/MSc Computer Science or Mathematics
  • MCA (Master of Computer Applications)  
  • BE/BTech Computer Science or IT
  • BCA + MCA combination

DATE PARSING (STRICT — Non-negotiable):
  • End year < 2025 → ✅ COMPLETED (hardcoded cutoff)
  • End year = 2025 → ✅ COMPLETED
  • End year > 2025 → 🔄 ONGOING
  • Keywords "pursuing", "in progress", "currently enrolled" → 🔄 ONGOING
  • Keywords "graduated", "completed", "finished" → ✅ COMPLETED
  • If end year < 2025, ALWAYS mark completed regardless of text

Scoring bands:
  • {int(edu_weight * 0.90)}–{edu_weight}: Outstanding — completed highly relevant degree + exceptional academic record
  • {int(edu_weight * 0.75)}–{int(edu_weight * 0.85)}: Excellent — priority degree (completed or ongoing), good standing
  • {int(edu_weight * 0.60)}–{int(edu_weight * 0.70)}: Very Good — related STEM/technical degree
  • {int(edu_weight * 0.45)}–{int(edu_weight * 0.55)}: Good — partially related degree with transferable foundation
  • {int(edu_weight * 0.30)}–{int(edu_weight * 0.40)}: Fair — unrelated degree with relevant self-learning evidence
  • {int(edu_weight * 0.15)}–{int(edu_weight * 0.25)}: Basic — minimal or no degree information
  • 0–{int(edu_weight * 0.10)}: Insufficient — no education details at all

**💼 Experience Score ({exp_weight} points max):**

Evaluate: years of relevant experience, role seniority, domain fit, impact, leadership, quantification.

  • {int(exp_weight * 0.91)}–{exp_weight}: Exceptional — exceeds requirements; strong leadership; quantified high-impact results
  • {int(exp_weight * 0.80)}–{int(exp_weight * 0.89)}: Excellent — meets/exceeds years; strong domain fit; clear achievements
  • {int(exp_weight * 0.69)}–{int(exp_weight * 0.77)}: Very Good — adequate years; good domain fit; solid responsibilities
  • {int(exp_weight * 0.57)}–{int(exp_weight * 0.66)}: Good — reasonable experience; relevant domain; some achievements
  • {int(exp_weight * 0.43)}–{int(exp_weight * 0.54)}: Fair — some gaps but shows clear potential and transferable skills
  • {int(exp_weight * 0.29)}–{int(exp_weight * 0.40)}: Basic — limited experience but relevant direction shown
  • {int(exp_weight * 0.14)}–{int(exp_weight * 0.26)}: Entry Level — minimal experience; strong potential only
  • 0–{int(exp_weight * 0.11)}: Insufficient — major gaps; no transferable evidence

NOTE: Internships, freelance projects, and open-source contributions count as valid experience.

**🛠️ Skills Score ({skills_weight} points max):**

Match each listed skill against job description requirements. Reward:
  • Hard skills: programming languages, frameworks, tools, platforms
  • Certifications: AWS, GCP, Azure, Kubernetes, Terraform, etc.
  • Emerging skills: LLMs, GenAI, Vector DBs, Web3, MLOps, DeFi, Smart Contracts

  • {int(skills_weight * 0.93)}–{skills_weight}: Outstanding — 90%+ required skills; expert proficiency; recent hands-on usage
  • {int(skills_weight * 0.80)}–{int(skills_weight * 0.90)}: Excellent — 80%+ required skills; advanced proficiency
  • {int(skills_weight * 0.67)}–{int(skills_weight * 0.77)}: Very Good — 70%+ required skills; competent usage
  • {int(skills_weight * 0.53)}–{int(skills_weight * 0.63)}: Good — 60%+ required skills; working knowledge
  • {int(skills_weight * 0.40)}–{int(skills_weight * 0.50)}: Fair — 50%+ skills OR strong foundational skills
  • {int(skills_weight * 0.27)}–{int(skills_weight * 0.37)}: Basic — 40%+ skills; clear learning trajectory
  • {int(skills_weight * 0.13)}–{int(skills_weight * 0.23)}: Limited — 30%+ skills; self-learning evident
  • 0–{int(skills_weight * 0.10)}: Insufficient — fewer than 30% required skills

**🔑 Keyword Score ({keyword_weight} points max):**

Systematically extract ALL critical terms from the job description:
technical tools, frameworks, methodologies, role titles, industry terms, certification names.
Compare against resume. Credit synonyms and equivalent terms.

  • {int(keyword_weight * 0.90)}–{keyword_weight}: Excellent — 85%+ critical terms; strong industry vocabulary
  • {int(keyword_weight * 0.80)}: Very Good — 75%+ critical terms
  • {int(keyword_weight * 0.60)}–{int(keyword_weight * 0.70)}: Good — 65%+ critical terms
  • {int(keyword_weight * 0.40)}–{int(keyword_weight * 0.50)}: Fair — 50%+ critical terms
  • {int(keyword_weight * 0.20)}–{int(keyword_weight * 0.30)}: Basic — 35%+ critical terms
  • {int(keyword_weight * 0.10)}: Limited — 20%+ critical terms
  • 0: Poor — fewer than 20% critical terms

═══════════════════════════════════════════════════
📋 REQUIRED OUTPUT FORMAT
═══════════════════════════════════════════════════

Follow this EXACT structure. Do not skip any section. Use the EXACT section markers shown (they are used by the parser — do not alter them):

[SEC:CANDIDATE_NAME]
<Copy the candidate's full name EXACTLY as it appears in the resume — character by character. Do NOT correct spelling, do NOT infer from context, do NOT paraphrase. Look at the very top of the resume (header/contact section). Output ONLY the name, nothing else. If you cannot find a name, write: Not Found>

[SEC:EDUCATION]
**Score:** <0–{edu_weight}> / {edu_weight}

**Scoring Rationale:**
- Degree Level & Relevance: <Does it qualify for minimum {int(edu_weight * 0.75)}-pt rule? Which degree?>
- Completion Status: <Apply strict 2025 cutoff rule; state year and final status>
- Academic Quality Indicators: <GPA, honors, relevant coursework if mentioned>
- **Score Justification:** <Explain exact score with evidence from resume>

[SEC:EXPERIENCE]
**Score:** <0–{exp_weight}> / {exp_weight}

**Experience Breakdown:**
- Total Years of Relevant Experience: <X years — include internships, freelance, open-source>
- Role Progression & Seniority: <Entry → Mid → Senior trajectory>
- Domain Alignment: <How well does background match job domain?>
- Quantified Achievements: <List metrics found: % improvement, $ savings, users served, etc.>
- Leadership & Ownership Evidence: <Managed teams? Led projects? Mentored?>
- Technology Currency: <Are skills/tools recent and relevant (last 3 years)?>
- **Score Justification:** <Explain score with specific resume evidence>

[SEC:SKILLS]
**Score:** <0–{skills_weight}> / {skills_weight}

**Skills Assessment:**
- Core Technical Skills Matched: <List matched skills with evidence>
- Emerging/Cutting-Edge Skills: <LLMs, GenAI, Web3, MLOps, Cloud, etc.>
- Certifications Detected: <List any certifications found>
- Soft Skills with Evidence: <Only count if backed by concrete examples>
- Proficiency Depth: <Surface knowledge vs. demonstrated project usage>

**Skills Gaps (Development Opportunities):**
- <Gap 1 — specific missing skill from job description>
- <Gap 2 — specific missing skill>
- <Gap 3 — specific missing skill>
- <Gap 4 — specific missing skill>
- <Gap 5 — specific missing skill>

**Score Justification:** <Explain with matched vs. required skills ratio>

[SEC:LANGUAGE]
**Score:** <evaluate and provide a score 0–{lang_weight}> / {lang_weight}
**Grammar & Professional Tone:** <single sentence summarising overall language quality>
**Suggestions:**
- <Actionable suggestion 1>
- <Actionable suggestion 2>
- <Actionable suggestion 3>
- <Actionable suggestion 4>
- <Actionable suggestion 5>
**Assessment:** <Specific feedback on action verb usage, clarity, tense consistency, and ATS language>

SCORING SCALE for language ({lang_weight} pts max):
- {lang_weight}: Exceptional — Flawless grammar, powerful action verbs, crystal-clear and professional throughout
- {lang_weight-1}: Very Good — Minor stylistic issues; highly professional and readable
- {lang_weight-2}: Good — Some grammar or clarity issues but largely professional
- {lang_weight-3}: Fair — Noticeable grammar or clarity problems
- 0-1: Poor — Significant language issues

[SEC:KEYWORD]
**Score:** <0–{keyword_weight}> / {keyword_weight}

**Keyword Assessment:**
- Industry Terminology Match: <Percentage and specific matches found>
- Role-Specific Keywords Present: <List matched keywords>
- Technical Vocabulary: <Tools, frameworks, platforms found in both>
- Keyword Density Quality: <Natural integration vs. stuffing>

**Keyword Enhancement Opportunities:**
- <Critical keyword 1 from job description — not in resume>
- <Critical keyword 2>
- <Critical keyword 3>
- <Critical keyword 4>
- <Critical keyword 5>
- <Critical keyword 6>
- <Critical keyword 7>
- <Critical keyword 8>

**Score Justification:** <Evidence-based explanation>

[SEC:FORMAT]
**Format Score:** {format_data.get("format_score", "N/A") if format_data else "N/A"} / 100  
**Format Grade:** {format_data.get("letter_grade", "N/A") if format_data else "N/A"} — {format_data.get("label", "") if format_data else ""}

⚠️ IMPORTANT: The Format Score and Format Grade above are SYSTEM-COMPUTED and LOCKED. Do NOT change these numbers. Only fill in the narrative fields below.

**Structural Assessment:**
- Section Completeness: <narrative only — do NOT include a score>
- Contact Block: <narrative only>
- Resume Length: {f"{format_data.get('word_count', 'N/A')} words — " + ("Optimal" if 300 <= (format_data.get('word_count') or 0) <= 1000 else "Too short" if (format_data.get('word_count') or 0) < 300 else "Too long") if format_data else "N/A"}
- Action Verb Strength: <narrative only>
- Quantification Quality: <narrative only>
- ATS Red Flags: <narrative only>

**Format Issues Detected:**
{chr(10).join(f"- {issue}" for issue in (format_data.get("issues", []) or ["No issues detected"])) if format_data else "- Format data not available"}

**Format Strengths:**
{chr(10).join(f"- {p}" for p in (format_data.get("passes", []) or ["No specific passes noted"])) if format_data else "- Format data not available"}

**Improvement Recommendations:**
- <Top format fix 1 — specific and actionable>
- <Top format fix 2>
- <Top format fix 3>

[SEC:FINAL]

**Overall Evaluation:**
<5–7 sentences covering: candidate's unique value proposition, strongest evidence-backed qualifications, key gaps, culture/team fit signals, and a clear hire/interview recommendation>

**Top 3 Strengths (with evidence):**
1. <Strength 1 — backed by resume evidence>
2. <Strength 2 — backed by resume evidence>
3. <Strength 3 — backed by resume evidence>

**Top 3 Development Areas:**
1. <Gap 1 framed as a growth opportunity>
2. <Gap 2 framed as a growth opportunity>
3. <Gap 3 framed as a growth opportunity>

**Hiring Recommendation:** <Strongly Recommend / Recommend / Recommend with Reservations / Do Not Recommend> — <2-sentence reasoning>

---

**EVALUATION CONTEXT:**
- Current Date: {datetime.datetime.now().strftime('%B %Y')} (Year: {current_year}, Month: {current_month})
- Resume Domain Detected: {resume_domain}
- Target Job Domain: {job_domain}
- Domain Similarity Score: {similarity_score:.2f}/1.0
- Domain Mismatch Penalty Applied: {domain_penalty}/{MAX_DOMAIN_PENALTY} pts

---

📄 **JOB DESCRIPTION:**
{job_description[:4000]}

📄 **RESUME TEXT:**
{resume_text[:5000]}

{logic_score_note}
"""
   
   
    try:
        ats_result = call_llm(prompt, session=st.session_state).strip()
    except Exception:
        ats_result = ""

    # Guard: LLM error string or empty → return a safe fallback ATS result
    _ERROR_PREFIXES = ("❌", "⚠️", "Error", "LLM unavailable", "No healthy", "rate limit", "quota")
    if not ats_result or any(ats_result.startswith(p) for p in _ERROR_PREFIXES):
        ats_result = (
            "**ATS Evaluation temporarily unavailable.**\n"
            "All API keys are currently exhausted or unavailable. "
            "Please try again in a few minutes."
        )

    # ── CRITICAL: Overwrite any LLM-modified Format Score/Grade lines ────
    # The LLM sometimes rewrites these despite instructions. Force the true
    # system-computed values back in so UI and narrative always match.
    _true_fmt_score = format_data.get("format_score", 75) if format_data else 75
    _true_fmt_grade = format_data.get("letter_grade", "N/A") if format_data else "N/A"
    _true_fmt_label = format_data.get("label", "") if format_data else ""

    ats_result = re.sub(
        r'\*\*Format Score:\*\*.*',
        f'**Format Score:** {_true_fmt_score} / 100',
        ats_result
    )
    ats_result = re.sub(
        r'\*\*Format Grade:\*\*.*',
        f'**Format Grade:** {_true_fmt_grade} — {_true_fmt_label}',
        ats_result
    )
    # ─────────────────────────────────────────────────────────────────────

    # ── GLOBAL LEAK GUARD: strip echoed prompt content from ats_result ───
    # Badly formatted resumes can cause the LLM to echo back the prompt
    # tail (EVALUATION CONTEXT / JOB DESCRIPTION / RESUME TEXT) anywhere
    # inside ats_result — which then renders visibly in the UI before the
    # ATS Score Breakdown chart. Strip at the first leaked marker found.
    _ATS_LEAK_MARKERS = [
        "**EVALUATION CONTEXT:**",
        "EVALUATION CONTEXT:",
        "📄 **JOB DESCRIPTION:**",
        "📄 **RESUME TEXT:**",
        "**JOB DESCRIPTION:**",
        "**RESUME TEXT:**",
        "JOB DESCRIPTION:",
        "RESUME TEXT:",
    ]
    _earliest_leak = len(ats_result)
    for _marker in _ATS_LEAK_MARKERS:
        _idx = ats_result.find(_marker)
        if _idx != -1 and _idx < _earliest_leak:
            _earliest_leak = _idx
    if _earliest_leak < len(ats_result):
        ats_result = ats_result[:_earliest_leak].strip()
    # ─────────────────────────────────────────────────────────────────────

    def extract_section(pattern, text, default="N/A"):
        match = re.search(pattern, text, re.DOTALL)
        return match.group(1).strip() if match else default

    def extract_score(pattern, text, default=0):
        match = re.search(pattern, text)
        return int(match.group(1)) if match else default

    # ── Section extraction using ASCII sentinel tags ──────────────────────
    # The prompt now uses [SEC:TAG] markers instead of emoji ### headers.
    # These are plain ASCII so the LLM can never accidentally alter them
    # (no emoji variation selectors, no rephrasing of heading words).
    # Pattern: \[SEC:TAG\][^\n]*\n  captures body until next sentinel or end.
    # Fallback: also accept old ### style headers so any cached legacy LLM
    # responses still parse correctly without re-running them.

    # ── PRE-CLEAN: normalise common LLM formatting deviations ────────────
    # Fix 1: LLM wraps tags in **bold** → strip the asterisks around tags
    #        e.g. **[SEC:LANGUAGE]** → [SEC:LANGUAGE]
    ats_result = re.sub(r'\*\*(\[SEC:[^\]]+\])\*\*', r'\1', ats_result)
    # Fix 2: LLM lowercases tags → uppercase them
    #        e.g. [sec:language] → [SEC:LANGUAGE]
    ats_result = re.sub(
        r'\[sec:([a-z_]+)\]',
        lambda m: '[SEC:' + m.group(1).upper() + ']',
        ats_result
    )
    # ─────────────────────────────────────────────────────────────────────

    # Known tag aliases — LLM sometimes shortens or renames tags
    _TAG_ALIASES = {
        "LANGUAGE":       ["LANG", "LANG_QUALITY", "LANGUAGE_QUALITY"],
        "KEYWORD":        ["KEYWORDS", "KW", "KEYWORD_ANALYSIS"],
        "CANDIDATE_NAME": ["NAME", "CANDIDATE"],
        "EXPERIENCE":     ["EXP", "WORK_EXPERIENCE"],
        "EDUCATION":      ["EDU"],
        "SKILLS":         ["SKILL"],
        "FORMAT":         ["FORMATTING"],
        "FINAL":          ["FINAL_THOUGHTS", "CONCLUSION"],
    }

    def _extract(tag, fallback_keyword, text, default="N/A"):
        """
        Primary:  [SEC:TAG] sentinel (case-insensitive after pre-clean).
        Aliases:  known shortened/alternate tag names the LLM might use.
        Fallback: ### line containing keyword (legacy responses).
        """
        # Primary — exact tag (already uppercased by pre-clean)
        m = re.search(
            r'\[SEC:' + re.escape(tag) + r'\][^\n]*\n(.*?)(?=\[SEC:|\Z)',
            text, re.DOTALL | re.IGNORECASE
        )
        if m:
            return m.group(1).strip()

        # Alias fallback — try known alternate tag names
        for alias in _TAG_ALIASES.get(tag, []):
            m = re.search(
                r'\[SEC:' + re.escape(alias) + r'\][^\n]*\n(.*?)(?=\[SEC:|\Z)',
                text, re.DOTALL | re.IGNORECASE
            )
            if m:
                return m.group(1).strip()

        # Legacy fallback — old ### markdown style headers
        m = re.search(
            r'###[^\n]*' + re.escape(fallback_keyword) + r'[^\n]*\n(.*?)(?=###|\Z)',
            text, re.DOTALL | re.IGNORECASE
        )
        return m.group(1).strip() if m else default

    _raw_name        = _extract("CANDIDATE_NAME", "Candidate Name", ats_result, "")
    edu_analysis     = _extract("EDUCATION",      "Education",      ats_result)
    exp_analysis     = _extract("EXPERIENCE",     "Experience",     ats_result)
    skills_analysis  = _extract("SKILLS",         "Skills",         ats_result)
    lang_analysis    = _extract("LANGUAGE",       "Language",       ats_result)
    keyword_analysis = _extract("KEYWORD",        "Keyword",        ats_result)
    format_analysis  = _extract("FORMAT",         "Format",         ats_result)
    final_thoughts   = _extract("FINAL",          "Final",          ats_result)

    # ── LEAK GUARD: strip prompt content if LLM echoed it back ──────────
    # Badly formatted resumes can cause the LLM to echo back parts of the
    # prompt (EVALUATION CONTEXT, JOB DESCRIPTION, RESUME TEXT) inside the
    # response sections. Truncate at the first leaked marker found.
    _LEAK_MARKERS = [
        "**EVALUATION CONTEXT:**",
        "EVALUATION CONTEXT:",
        "📄 **JOB DESCRIPTION:**",
        "📄 **RESUME TEXT:**",
        "**JOB DESCRIPTION:**",
        "**RESUME TEXT:**",
        "JOB DESCRIPTION:",
        "RESUME TEXT:",
    ]
    def _strip_leaked_prompt(text: str) -> str:
        if not text:
            return text
        earliest = len(text)
        for marker in _LEAK_MARKERS:
            idx = text.find(marker)
            if idx != -1 and idx < earliest:
                earliest = idx
        return text[:earliest].strip()

    final_thoughts   = _strip_leaked_prompt(final_thoughts)
    edu_analysis     = _strip_leaked_prompt(edu_analysis)
    exp_analysis     = _strip_leaked_prompt(exp_analysis)
    skills_analysis  = _strip_leaked_prompt(skills_analysis)
    lang_analysis    = _strip_leaked_prompt(lang_analysis)
    keyword_analysis = _strip_leaked_prompt(keyword_analysis)
    format_analysis  = _strip_leaked_prompt(format_analysis)
    # ─────────────────────────────────────────────────────────────────────

    candidate_name = re.sub(r"[*_`#\[\]<>]", "", _raw_name).strip()
    candidate_name = " ".join(candidate_name.split())
    _placeholder_values = {
        "not found", "n/a", "unknown", "none", "",
        "extract full name from resume header or contact section",
        "copy the candidate's full name exactly as it appears in the resume",
        "copy the candidates full name exactly as it appears in the resume",
        "name not found", "candidate name not found",
    }
    if candidate_name.lower() in _placeholder_values:
        candidate_name = "Not Found"

    # Extract scores with improved patterns (LLM now scores directly using sidebar weights)
    # NOTE: \s* between Score and : tolerates "Score :" (space before colon) LLM deviation
    edu_score     = extract_score(r"\*\*Score\s*:\*\*\s*(\d+)", edu_analysis)
    exp_score     = extract_score(r"\*\*Score\s*:\*\*\s*(\d+)", exp_analysis)
    skills_score  = extract_score(r"\*\*Score\s*:\*\*\s*(\d+)", skills_analysis)
    keyword_score = extract_score(r"\*\*Score\s*:\*\*\s*(\d+)", keyword_analysis)
    # ⚡ Parse grammar score + feedback from ATS result (no separate LLM call needed)
    _grammar_score_match    = re.search(r"\*\*Score\s*:\*\*\s*<evaluate.*?(\d+)>|Score.*?(\d+)\s*/\s*" + str(lang_weight), lang_analysis)
    _grammar_score_match2   = re.search(r"\*\*Score\s*:\*\*\s*(\d+)", lang_analysis)
    _grammar_feedback_match = re.search(r"\*\*Grammar & Professional Tone:\*\*\s*(.+)", lang_analysis)
    _grammar_sugg_raw       = re.findall(r"^- (.+)", lang_analysis, re.MULTILINE)

    if _grammar_score_match2:
        grammar_score = int(_grammar_score_match2.group(1))
    # else keep the safe default already set above

    if _grammar_feedback_match:
        grammar_feedback = _grammar_feedback_match.group(1).strip()

    if _grammar_sugg_raw:
        grammar_suggestions = _grammar_sugg_raw

    lang_score = grammar_score  # use value parsed from ATS result

    # ── Clamp every score: min floor + hard upper cap to its own weight ──
    # Upper cap prevents LLM hallucinating over-max scores (e.g. 25/20)
    # which would silently push content_score above 100.
    edu_score     = max(int(edu_weight * 0.15),     min(edu_score,     edu_weight))
    exp_score     = max(int(exp_weight * 0.15),     min(exp_score,     exp_weight))
    skills_score  = max(int(skills_weight * 0.15),  min(skills_score,  skills_weight))
    keyword_score = max(int(keyword_weight * 0.10), min(keyword_score, keyword_weight))
    lang_score    = max(0,                          min(lang_score,    lang_weight))

    # Extract missing items with better parsing - now called "opportunities"
    missing_keywords_section = extract_section(r"\*\*Keyword Enhancement Opportunities:\*\*(.*?)(?:\*\*|###|\Z)", keyword_analysis)
    missing_skills_section = extract_section(r"\*\*Skills Gaps \(Development Opportunities\):\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)

    # Fallback to old patterns if new ones don't match
    if not missing_keywords_section.strip():
        missing_keywords_section = extract_section(r"\*\*Missing Critical Keywords:\*\*(.*?)(?:\*\*|###|\Z)", keyword_analysis)
    if not missing_skills_section.strip():
        missing_skills_section = extract_section(r"\*\*Skills Gaps \(Opportunities for Growth\):\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)
    if not missing_skills_section.strip():
        missing_skills_section = extract_section(r"\*\*Missing Critical Skills:\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)
    
    # Improved extraction - handle multiple formats and get all items
    def extract_list_items(text):
        if not text.strip():
            return "None identified"
        
        # Find all bullet points with various formats
        items = []
        lines = text.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Remove various bullet point formats
            cleaned_line = re.sub(r'^[-•*]\s*', '', line)  # Remove -, •, * bullets
            cleaned_line = re.sub(r'^\d+\.\s*', '', cleaned_line)  # Remove numbered lists
            cleaned_line = cleaned_line.strip()
            
            if cleaned_line and len(cleaned_line) > 2:  # Avoid empty or very short items
                items.append(cleaned_line)
        
        return ', '.join(items) if items else "None identified"
    
    missing_keywords = extract_list_items(missing_keywords_section)
    missing_skills = extract_list_items(missing_skills_section)

    # ── Score assembly — fully deterministic integer arithmetic ──────────
    # Step 1: sum the five LLM-scored components (clamped to their individual weights)
    content_score = edu_score + exp_score + skills_score + lang_score + keyword_score

    # Normalise LLM components to 90-pt scale (format takes the remaining 10 pts)
    # This keeps total = 100 while giving format meaningful, visible weight.
    weight_total = edu_weight + exp_weight + skills_weight + lang_weight + keyword_weight
    # weight_total should always be > 0 here because the sidebar hard blocks evaluation
    # when total != 90 (min slider values are all >= 2). Guard kept as last-resort safety.
    if weight_total > 0:
        content_score = round(content_score / weight_total * 90)
    else:
        content_score = 0  # should never reach here after sidebar guard
    content_score = max(0, min(90, content_score))

    # Step 2: format score (0–100) contributes a fixed 10-pt component
    # Scaled proportionally: 100 format → 10 pts, 0 format → 0 pts
    fmt_score_raw = format_data.get("format_score", 75) if format_data else 75
    fmt_score_raw = max(0, min(100, int(fmt_score_raw)))
    FORMAT_WEIGHT = 10
    format_component = round(fmt_score_raw / 100 * FORMAT_WEIGHT)

    # Step 3: combine content + format → pre-penalty total (0–100)
    pre_penalty_score = content_score + format_component
    pre_penalty_score = max(0, min(100, pre_penalty_score))

    # Step 4: subtract domain mismatch penalty ONCE — straight subtraction
    total_score = pre_penalty_score - domain_penalty

    # Step 5: clamp final result 15–100
    total_score = max(15, min(100, total_score))

    # ✅ Industry-standard score labels with clear hiring signal
    formatted_score = (
        "Exceptional Match — Top 10% Candidate"    if total_score >= 85 else
        "Strong Match — Recommend for Interview"    if total_score >= 70 else
        "Good Potential — Competitive Candidate"    if total_score >= 55 else
        "Fair Match — Needs Resume Optimization"    if total_score >= 40 else
        "Developing — Significant Skill Gaps"       if total_score >= 25 else
        "Poor Match — Major Role Misalignment"
    )

    # ✅ Format suggestions nicely
    suggestions_html = ""
    if grammar_suggestions:
        suggestions_html = "<ul>" + "".join([f"<li>{s}</li>" for s in grammar_suggestions]) + "</ul>"

    # Convert LLM markdown to HTML so asterisks don't render literally
    # (lang_analysis is raw markdown; the appended content is already HTML)
    _lang_html = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', lang_analysis)   # **bold**
    _lang_html = re.sub(r'\*(.+?)\*',     r'<i>\1</i>', _lang_html)       # *italic*
    _lang_html = re.sub(r'^- ',           '• ',          _lang_html, flags=re.MULTILINE)  # bullet dash
    _lang_html = _lang_html.replace('\n', '<br>')

    updated_lang_analysis = (
        f"{_lang_html}"
        f"<br><b>LLM Feedback Summary:</b> {grammar_feedback}"
        f"<br><b>Improvement Suggestions:</b> {suggestions_html}"
    )

    # Enhanced final thoughts with domain analysis and industry benchmarks
    final_thoughts += f"""

**Score Breakdown:**
- Content Score: {content_score}/90
- Format Score: {format_component}/10 (raw: {fmt_score_raw}/100)
- Score before domain adjustment: {pre_penalty_score}/100
- Domain adjustment: {("-" + str(domain_penalty) + " pts (your field doesn\'t fully match the job)") if domain_penalty > 0 else "None — your field matches the job"}
- Final ATS Score: {total_score}/100
- Your field: {resume_domain}
- Job field: {job_domain}
- Experience level in your field: {_depth_str.capitalize()} {"(API/tool usage only)" if _depth_str == "shallow" else "(projects/internships)" if _depth_str == "moderate" else "(work experience/quantified results)"}

**Score Interpretation (Industry Benchmarks):**
- 85–100: Top 10% candidates — Strong interview recommendation
- 70–84: Above average — Likely to advance past ATS screening
- 55–69: Competitive — May advance with strong cover letter
- 40–54: Below average — Needs resume optimization before applying
- 25–39: Significant gaps — Upskilling recommended
- 0–24: Major misalignment — Not suitable for this specific role

**ATS Scoring Notes:**
- Scoring model: LLM components (90 pts) + Format (10 pts) − Domain penalty
- Format score is a real 10-pt component (not a delta) — poor formatting meaningfully lowers the score
- Domain penalty subtracted once as a flat deduction (max {MAX_DOMAIN_PENALTY} pts)
- Format checker v2: uses PDF block-coordinate multi-column detection, tiered deductions, bonus credits
- Transferable skills, projects, and open-source contributions were credited
- Career stage (entry/mid/senior) considered in experience scoring
"""

    return ats_result, {
        "Candidate Name": candidate_name,
        "Education Score": edu_score,
        "Experience Score": exp_score,
        "Skills Score": skills_score,
        "Language Score": lang_score,
        "Keyword Score": keyword_score,
        "Format Score": fmt_score_raw,
        "Format Grade": format_data.get("letter_grade", "N/A") if format_data else "N/A",
        "Format Label": format_data.get("label", "") if format_data else "",
        "Format Issues": format_data.get("issues", []) if format_data else [],
        "Format Passes": format_data.get("passes", []) if format_data else [],
        "ATS Match %": total_score,
        "Formatted Score": formatted_score,
        "Education Analysis": edu_analysis,
        "Experience Analysis": exp_analysis,
        "Skills Analysis": skills_analysis,
        "Language Analysis": updated_lang_analysis,
        "Keyword Analysis": keyword_analysis,
        "Format Analysis": format_analysis,
        "Final Thoughts": final_thoughts,
        "Missing Keywords": missing_keywords,
        "Missing Skills": missing_skills,
        "Resume Domain": resume_domain,
        "Job Domain": job_domain,
        "Resume Depth": _depth_str,
        "Domain Penalty": domain_penalty,
        "Domain Similarity Score": similarity_score
    }

# Setup Vector DB
def setup_vectorstore(documents):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    if DEVICE == "cuda":
        embeddings.model = embeddings.model.to(torch.device("cuda"))
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    doc_chunks = text_splitter.split_text("\n".join(documents))
    return FAISS.from_texts(doc_chunks, embeddings)

# Create Conversational Chain
def create_chain(vectorstore):
    # ✅ Use get_healthy_keys() so dead/quota keys are skipped (reads key_failures
    #    and key_usage from Supabase — same tables that call_llm() maintains).
    all_keys    = load_groq_api_keys()
    healthy     = get_healthy_keys(all_keys)
    if not healthy:
        raise ValueError("❌ No healthy Groq API keys available for chat chain.")
    # healthy list is already shuffled by get_healthy_keys — just take the first
    groq_api_key = healthy[0]
    # ✅ FIX: do NOT increment usage before the call — only after success

    # ✅ Create the ChatGroq object
    llm = ChatGroq(model="llama-3.3-70b-versatile", temperature=0, groq_api_key=groq_api_key)

    # ✅ Build the chain — report failures back so llm_manager skips this key next time
    try:
        chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=vectorstore.as_retriever(),
            return_source_documents=True
        )
        # Update in-memory usage instantly; flush to Supabase in background thread
        _mem_increment_usage(groq_api_key)
        _async_increment_usage(groq_api_key)
        _mem_clear_failure(groq_api_key)
        _async_clear_failure(groq_api_key)
        return chain
    except Exception as e:
        err_str = str(e).lower()
        if any(w in err_str for w in ["quota", "rate limit", "429", "too many requests"]):
            _mem_record_failure(groq_api_key, "quota")
            _async_mark_failure(groq_api_key, "quota")
        elif any(w in err_str for w in ["invalid api key", "unauthorized", "401", "403", "authentication"]):
            _mem_record_failure(groq_api_key, "error")
            _async_mark_failure(groq_api_key, "error")
        # transient errors (network blip, 500) — don't mark the key failed at all
        raise

# Chat history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
