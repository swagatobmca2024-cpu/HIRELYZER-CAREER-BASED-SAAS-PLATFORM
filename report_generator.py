import os
os.environ["STREAMLIT_WATCHDOG"] = "false"
import json
import random
import string
import re
import asyncio
import io
from concurrent.futures import ThreadPoolExecutor, as_completed
import urllib.parse
import base64
from io import BytesIO
from collections import Counter
from datetime import datetime
import time

import streamlit as st
import streamlit.components.v1 as components
from base64 import b64encode
import requests
import fitz
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import altair as alt
from PIL import Image
from pdf2image import convert_from_path
from dotenv import load_dotenv
from nltk.stem import WordNetLemmatizer
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from xhtml2pdf import pisa
from pydantic import BaseModel
from streamlit_pdf_viewer import pdf_viewer
import torch
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from llm_manager import (
    call_llm, load_groq_api_keys, get_healthy_keys, increment_key_usage,
    mark_key_failure, _mem_record_failure, _mem_clear_failure,
    _mem_increment_usage, _async_mark_failure, _async_increment_usage,
    _async_clear_failure,
)
from db_manager import (
    db_manager, insert_candidate, get_top_domains_by_score,
    get_database_stats, detect_domain_from_title_and_description,
    get_domain_similarity
)
from user_login import (
    create_user_table, add_user, complete_registration, verify_user,
    get_logins_today, get_total_registered_users, log_user_action,
    username_exists, email_exists, is_valid_email, get_all_user_logs,
    generate_otp, send_email_otp,
    get_user_by_email, update_password_by_email, is_strong_password,
    domain_has_mx_record, send_login_link, verify_login_token,
    cleanup_expired_login_tokens, check_and_gate_feature,
    record_feature_usage, get_usage_count_last_hour, check_brute_force,
)

# ── report_generator.py ─────────────────────────────────────────────────────

# ── Grade normaliser — single source of truth for ALL templates ─────────────
def _normalize_cgpa(raw: str) -> str:
    """
    Normalize academic score strings to a clean 'LABEL: value' format.
    Handles all input variants the LLM or raw PDF text might produce.
    Rules:
      - Never convert between score types (SGPA != CGPA != GPA != Percentage)
      - Never duplicate labels (e.g. 'CGPA: 7.0 GPA' -> 'CGPA: 7.0')
      - Numeric-only values -> assumed CGPA if <= 10, Percentage if > 10
      - Semester suffixes are preserved (e.g. 'SGPA: 7.4 (Semester 1)')
    """
    if not raw:
        return ""
    s = raw.strip()
    if not s:
        return ""

    # Pre-normalise: collapse spaces around slash e.g. "8.5 / 10" -> "8.5/10"
    s = re.sub(r'\s*/\s*', '/', s)

    # Pre-normalise: collapse "LABEL : value" (space before colon)
    s = re.sub(r'^(cgpa|sgpa|gpa|percentage)\s*:\s*', lambda m: m.group(1).upper() + ': ', s, flags=re.IGNORECASE)

    # Pre-normalise: collapse "LABEL - value" (dash format) → "LABEL: value"
    # e.g. "SGPA - 7.4 (1st Sem)" → "SGPA: 7.4 (1st Sem)"
    # e.g. "CGPA - 8.44" → "CGPA: 8.44"
    # e.g. "Percentage - 78.3%" → "Percentage: 78.3%"
    s = re.sub(r'^(cgpa|sgpa|gpa|percentage)\s*-\s*', lambda m: m.group(1).upper() + ': ', s, flags=re.IGNORECASE)

    # Already clean: CGPA/GPA/SGPA prefixes — strip only trailing duplicate labels
    for prefix in ("CGPA:", "GPA:", "SGPA:"):
        if s.upper().startswith(prefix.upper()):
            val = s[len(prefix):].strip()
            val = re.sub(r'\s+(cgpa|gpa|sgpa)\s*$', '', val, flags=re.IGNORECASE).strip()
            return f"{prefix} {val}"

    # Already clean: Percentage: prefix
    if s.upper().startswith("PERCENTAGE:"):
        val = s[len("Percentage:"):].strip()
        val = re.sub(r'\s+percent(?:age)?\s*$', '', val, flags=re.IGNORECASE).strip()
        if not val.endswith('%'):
            val = val.rstrip('%').strip() + '%'
        return f"Percentage: {val}"

    # Pattern: value then label + optional semester suffix e.g. "8.2 SGPA", "7.0 GPA"
    m = re.match(
        r'^(\d+(?:\.\d+)?(?:/\d+(?:\.\d+)?)?)\s*(cgpa|sgpa|gpa)((?:\s*[\(\[].*?[\)\]]|\s+Sem(?:ester)?\s*\d+)?)$',
        s, re.IGNORECASE
    )
    if m:
        val, label, suffix = m.group(1), m.group(2).upper(), m.group(3).strip()
        return f"{label}: {val}{(' ' + suffix) if suffix else ''}"

    # Pattern: label then value + optional semester suffix e.g. "CGPA 8.5", "SGPA 7.9"
    m = re.match(
        r'^(cgpa|sgpa|gpa)\s+(\d+(?:\.\d+)?(?:/\d+(?:\.\d+)?)?)((?:\s*[\(\[].*?[\)\]]|\s+Sem(?:ester)?\s*\d+)?)$',
        s, re.IGNORECASE
    )
    if m:
        label, val, suffix = m.group(1).upper(), m.group(2), m.group(3).strip()
        return f"{label}: {val}{(' ' + suffix) if suffix else ''}"

    # Pattern: percentage e.g. "78%", "78.3%", "87.4 percent", "83 %"
    m = re.match(r'^(\d+(?:\.\d+)?)\s*(%|percent(?:age)?)$', s, re.IGNORECASE)
    if m:
        return f"Percentage: {m.group(1)}%"

    # Pattern: x/y e.g. "8.5/10", "3.9/4.0", "85/100"
    m = re.match(r'^(\d+(?:\.\d+)?)/(\d+(?:\.\d+)?)$', s)
    if m:
        try:
            label = 'Percentage' if float(m.group(1)) > 10 else 'CGPA'
        except Exception:
            label = 'CGPA'
        return f"{label}: {m.group(1)}/{m.group(2)}"

    # Pattern: bare numeric (decimal or integer) e.g. "8.44", "7.0", "7", "85"
    m = re.match(r'^(\d+(?:\.\d+)?)$', s)
    if m:
        try:
            numeric = float(m.group(1))
            return f"{'Percentage' if numeric > 10 else 'CGPA'}: {m.group(1)}"
        except Exception:
            pass

    # Fallback: return as-is (unknown format, don't corrupt)
    return s


def html_to_pdf_bytes(html_string):
    styled_html = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: 400mm 297mm;  /* Original custom large page size */
                margin-top: 10mm;
                margin-bottom: 10mm;
                margin-left: 10mm;
                margin-right: 10mm;
            }}
            body {{
                font-size: 14pt;
                font-family: "Segoe UI", "Helvetica", sans-serif;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2, h3 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 10px;
            }}
            .box {{
                padding: 8px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;  /* More elegant than full border */
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.5em;
            }}
            li {{
                margin-bottom: 5px;
            }}
        </style>
    </head>
    <body>
        {html_string}
    </body>
    </html>
    """

    pdf_io = BytesIO()
    pisa.CreatePDF(styled_html, dest=pdf_io)
    pdf_io.seek(0)
    return pdf_io



def _val(v) -> str:
    """Return value or 'Not Provided' placeholder — never empty, never None, never '[Not Provided]'."""
    if v is None:
        return "Not Provided"
    s = str(v).strip()
    if not s or s in ("[Not Provided]", "null", "None", "undefined"):
        return "Not Provided"
    return s


def _build_contact_header(doc, data: dict, name_size: int, name_color_rgb: tuple,
                           name_font: str, contact_font: str, contact_color_hex: str,
                           contact_size: int = 9, title_font: str = None,
                           title_size: int = 11, title_color_rgb: tuple = None,
                           separator: str = "  |  ",
                           label_color_hex: str = None,
                           accent_color_hex: str = None):
    """
    Builds the header block matching the exact sample template format:

      Line 1:  FULL NAME  (large bold centered)
      Line 2:  JOB TITLE  (smaller bold/normal centered uppercase)
      Line 3:  email  |  phone  |  LOCATION  |  linkedin_url  |  github_url
               (single pipe-separated line, centered, all fields always present)

    Every field always appears. Missing values show "Not Provided".
    No separate labeled rows. No dividers. Single clean contact line.
    """
    contact = data.get("contact", {})

    # ── Resolve contact color ────────────────────────────────────────────────
    cc = RGBColor(
        int(contact_color_hex[0:2], 16),
        int(contact_color_hex[2:4], 16),
        int(contact_color_hex[4:6], 16),
    )

    # ── ① Name ──────────────────────────────────────────────────────────────
    raw_name = contact.get("name", "") or ""
    name = raw_name if raw_name and raw_name not in ("", "[Not Provided]") else "Your Name"
    p_name = doc.add_paragraph()
    p_name.clear()
    r_name = p_name.add_run(name)
    r_name.bold = True
    r_name.font.size = Pt(name_size)
    r_name.font.name = name_font
    r_name.font.color.rgb = RGBColor(*name_color_rgb)
    p_name.alignment = 1
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(4)

    # ── ② Job Title — always shown, placeholder if missing ──────────────────
    raw_title = contact.get("title", "") or ""
    title_text = raw_title if raw_title and raw_title not in ("", "[Not Provided]") else "Job Title"
    p_title = doc.add_paragraph()
    p_title.clear()
    r_title = p_title.add_run(title_text.upper())
    r_title.font.size = Pt(title_size)
    r_title.font.name = title_font or name_font
    r_title.bold = True
    if title_color_rgb:
        r_title.font.color.rgb = RGBColor(*title_color_rgb)
    p_title.alignment = 1
    p_title.paragraph_format.space_after = Pt(4)

    # ── ③ Single pipe-separated contact line ────────────────────────────────
    # Build parts in order: email | phone | location | linkedin | github
    # Every field always included — "Not Provided" if missing
    def _clean(v):
        """Return value or 'Not Provided' — never None or empty."""
        if v is None:
            return "Not Provided"
        s = str(v).strip()
        return s if s and s not in ("[Not Provided]", "null", "None", "undefined") else "Not Provided"

    email_val    = _clean(contact.get("email", ""))
    phone_val    = _clean(contact.get("phone", ""))
    location_val = _clean(contact.get("location", ""))
    linkedin_val = _clean(contact.get("linkedin", ""))
    github_raw   = contact.get("github", "") or ""
    portfolio_raw= contact.get("portfolio", "") or ""
    github_val   = _clean(github_raw if github_raw and github_raw not in ("", "[Not Provided]") else portfolio_raw)

    _link_color = (
        int(contact_color_hex[0:2], 16),
        int(contact_color_hex[2:4], 16),
        int(contact_color_hex[4:6], 16),
    )

    p_contact = doc.add_paragraph()
    p_contact.clear()
    p_contact.alignment = 1
    p_contact.paragraph_format.space_before = Pt(2)
    p_contact.paragraph_format.space_after = Pt(6)

    def _plain_run(para, text):
        r = para.add_run(text)
        r.font.size = Pt(contact_size)
        r.font.name = contact_font
        r.font.color.rgb = cc
        return r

    SEP = "  |  "
    plain_parts = [email_val, phone_val, location_val]
    _plain_run(p_contact, SEP.join(plain_parts))

    # LinkedIn — clickable if it looks like a URL
    _plain_run(p_contact, SEP)
    if linkedin_val.startswith("http"):
        _add_hyperlink(p_contact, linkedin_val, linkedin_val,
                       font_name=contact_font, font_size=contact_size,
                       color_rgb=_link_color)
    else:
        _plain_run(p_contact, linkedin_val)

    # GitHub/Portfolio — clickable if it looks like a URL
    _plain_run(p_contact, SEP)
    if github_val.startswith("http"):
        _add_hyperlink(p_contact, github_val, github_val,
                       font_name=contact_font, font_size=contact_size,
                       color_rgb=_link_color)
    else:
        _plain_run(p_contact, github_val)

    # ── ④ Thin bottom rule below contact block ────────────────────────────
    # Signals end of header to ATS parsers and improves recruiter readability.
    from docx.oxml import OxmlElement as _OE
    from docx.oxml.ns import qn as _qn
    pPr = p_contact._p.get_or_add_pPr()
    pBdr = _OE('w:pBdr')
    btm = _OE('w:bottom')
    btm.set(_qn('w:val'), 'single')
    btm.set(_qn('w:sz'), '4')
    btm.set(_qn('w:space'), '1')
    # Use accent color if provided, else dark gray
    _border_col = accent_color_hex if accent_color_hex else "555555"
    btm.set(_qn('w:color'), _border_col)
    pBdr.append(btm)
    pPr.append(pBdr)


def _section_heading_bordered(doc, text: str, font_name: str,
                               font_size: int, bold: bool,
                               color_hex: str, border_color: str,
                               border_sz: str = "6",
                               space_before: float = 10, space_after: float = 4,
                               prefix: str = ""):
    """Universal bordered section heading."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    p.clear()
    label = f"{prefix}{text.upper()}" if prefix else text.upper()
    run = p.add_run(label)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), border_sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), border_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
    return p


def _add_hyperlink(paragraph, url: str, display_text: str, font_name: str = "Calibri",
                   font_size: int = 10, color_rgb: tuple = (0, 112, 192), underline: bool = True):
    """
    Add a real clickable hyperlink run to an existing paragraph in a python-docx document.
    Uses OOXML relationship injection — works in Word, LibreOffice, and Google Docs.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re

    # Sanitise URL — ensure it has a scheme
    url = url.strip()
    if url and not re.match(r"https?://", url, re.IGNORECASE):
        url = "https://" + url

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Font name
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    # Font size (half-points)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size * 2))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(font_size * 2))
    rPr.append(szCs)

    # Color
    color_el = OxmlElement("w:color")
    hex_color = "{:02X}{:02X}{:02X}".format(*color_rgb)
    color_el.set(qn("w:val"), hex_color)
    rPr.append(color_el)

    # Underline
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    # Style override to prevent default hyperlink style from overriding color
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.insert(0, rStyle)

    new_run.append(rPr)

    # Text
    t = OxmlElement("w:t")
    t.text = display_text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_bullet(doc, text: str, font_size: int = 10, font_name: str = "Arial",
                indent_left: int = 360, indent_hanging: int = 180,
                color_rgb: tuple = None):
    """
    Add a properly formatted ATS-compliant bullet point paragraph.
    Uses standard hanging indent matching Jobscan/Enhancv output.
    Bullet character is plain Unicode bullet (U+2022) — universally parsed by all ATS.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    # Strip any leading bullet/dash the LLM may have prepended to avoid double-bullets
    clean_text = text.strip()
    for prefix in ("\u2022", "-", "*", "\u25aa", "\u25cf"):
        if clean_text.startswith(prefix):
            clean_text = clean_text[len(prefix):].lstrip()
            break
    p = doc.add_paragraph(style="Normal")
    p.clear()
    run = p.add_run(f"\u2022  {clean_text}")
    run.font.size = Pt(font_size)
    run.font.name = font_name
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(indent_left))
    ind.set(qn('w:hanging'), str(indent_hanging))
    pPr.append(ind)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    _set_para_keep(p, keep_together=True, keep_with_next=False, widow_control=True)
    return p


def _set_para_keep(paragraph, keep_together: bool = True, keep_with_next: bool = False, widow_control: bool = True):
    """
    Apply paragraph-level page-break control to a python-docx paragraph.
    - keep_together:  prevents a paragraph from splitting across pages (good for bullets)
    - keep_with_next: keeps this paragraph on the same page as the next one (good for headings/role lines)
    - widow_control:  prevents single orphan/widow lines at top or bottom of page
    All three are set via the paragraph's XML pPr element.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    if keep_together:
        kT = OxmlElement('w:keepLines')
        pPr.append(kT)
    if keep_with_next:
        kN = OxmlElement('w:keepNext')
        pPr.append(kN)
    if widow_control:
        wC = OxmlElement('w:widowControl')
        wC.set(qn('w:val'), '1')
        pPr.append(wC)


def _add_role_line(doc, role: str, company: str, duration: str,
                   font_name: str, role_size: int = 11, meta_size: int = 9,
                   role_color: tuple = (0, 0, 0), company_color: tuple = (74, 74, 74),
                   duration_color: tuple = (128, 128, 128), separator: str = "  —  "):
    """Add role | company | duration header row for experience."""
    p = doc.add_paragraph()
    p.clear()
    if role:
        r1 = p.add_run(role)
        r1.bold = True
        r1.font.size = Pt(role_size)
        r1.font.name = font_name
        r1.font.color.rgb = RGBColor(*role_color)
    if company:
        r2 = p.add_run(f"{separator}{company}")
        r2.font.size = Pt(role_size - 1)
        r2.font.name = font_name
        r2.font.color.rgb = RGBColor(*company_color)
    if duration and duration not in ("", "[Not Provided]"):
        r3 = p.add_run(f"   [{duration}]")
        r3.italic = True
        r3.font.size = Pt(meta_size)
        r3.font.name = font_name
        r3.font.color.rgb = RGBColor(*duration_color)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
    return p


def _add_project_header(doc, name: str, duration: str, tech_stack: str, url: str,
                         font_name: str, name_size: int = 11,
                         name_color: tuple = (0, 0, 0),
                         meta_color: tuple = (74, 74, 74),
                         url_color: tuple = (30, 58, 95)):
    """Add project name | duration | tech stack | URL header row."""
    p = doc.add_paragraph()
    p.clear()
    if name:
        r1 = p.add_run(name)
        r1.bold = True
        r1.font.size = Pt(name_size)
        r1.font.name = font_name
        r1.font.color.rgb = RGBColor(*name_color)
    if duration and duration not in ("", "[Not Provided]"):
        rd = p.add_run(f"   [{duration}]")
        rd.italic = True
        rd.font.size = Pt(name_size - 2)
        rd.font.name = font_name
        rd.font.color.rgb = RGBColor(128, 128, 128)
    if tech_stack and tech_stack not in ("", "[Not Provided]"):
        r2 = p.add_run(f"  |  Tech: {tech_stack}")
        r2.font.size = Pt(name_size - 2)
        r2.font.name = font_name
        r2.font.color.rgb = RGBColor(*meta_color)
    if url and url not in ("", "[Not Provided]"):
        r3 = p.add_run(f"  |  {url}")
        r3.font.size = Pt(name_size - 2)
        r3.font.name = font_name
        r3.font.color.rgb = RGBColor(*url_color)
        r3.underline = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
    return p


def _add_education_row(doc, degree: str, institution: str, year: str,
                        edu_bullets: list, font_name: str,
                        degree_size: int = 10, meta_size: int = 9,
                        degree_color: tuple = (0, 0, 0),
                        inst_color: tuple = (74, 74, 74),
                        year_color: tuple = (128, 128, 128)):
    """Add degree | institution | year + optional bullets."""
    p = doc.add_paragraph()
    p.clear()
    if degree:
        r1 = p.add_run(degree)
        r1.bold = True
        r1.font.size = Pt(degree_size)
        r1.font.name = font_name
        r1.font.color.rgb = RGBColor(*degree_color)
    if institution and institution not in ("", "[Not Provided]"):
        r2 = p.add_run(f"  —  {institution}")
        r2.font.size = Pt(degree_size)
        r2.font.name = font_name
        r2.font.color.rgb = RGBColor(*inst_color)
    if year and year not in ("", "[Not Provided]"):
        r3 = p.add_run(f"  ({year})")
        r3.italic = True
        r3.font.size = Pt(meta_size)
        r3.font.name = font_name
        r3.font.color.rgb = RGBColor(*year_color)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
    for b in (edu_bullets or []):
        if b and b != "[Not Provided]":
            _add_bullet(doc, b, font_size=degree_size - 1, font_name=font_name)


def _render_additional(doc, data: dict, font_name: str, font_size: int,
                        heading_fn, bullet_fn,
                        name_color_rgb: tuple = (0,0,0),
                        desc_color_rgb: tuple = (80,80,80),
                        dur_color_rgb: tuple = (128,128,128)):
    """
    Render the Additional section from structured objects.
    Each item: bold name + optional [duration] on one line, description below.
    Handles dicts, flat strings, and completely skips raw leaked JSON strings.
    """
    raw = data.get("additional", [])
    # Final safety normalisation at render time
    items = []
    for item in raw:
        if isinstance(item, dict):
            name = str(item.get("name", "") or "").strip()
            desc = str(item.get("description", "") or "").strip()
            dur  = str(item.get("duration", "") or "").strip()
            # Skip if both name and desc are empty or placeholder
            if not name and not desc:
                continue
            if name in ("[Not Provided]", "") and desc in ("[Not Provided]", ""):
                continue
            items.append({"name": name, "description": desc, "duration": dur})
        elif isinstance(item, str):
            s = item.strip()
            if not s or s == "[Not Provided]":
                continue
            # Discard any raw dict/JSON leak silently
            if s.startswith("{") or ("'name'" in s and "'description'" in s) or s.startswith("["):
                continue
            items.append({"name": s, "description": "", "duration": ""})

    if not items:
        return

    heading_fn("Additional")

    for item in items:
        name = item.get("name", "")
        desc = item.get("description", "")
        dur  = item.get("duration", "")

        # Name line: bold name + italic [duration]
        p = doc.add_paragraph()
        p.clear()
        if name and name not in ("[Not Provided]", ""):
            r1 = p.add_run(name)
            r1.bold = True
            r1.font.size = Pt(font_size)
            r1.font.name = font_name
            r1.font.color.rgb = RGBColor(*name_color_rgb)
        if dur and dur not in ("[Not Provided]", ""):
            rd = p.add_run(f"   [{dur}]")
            rd.italic = True
            rd.font.size = Pt(font_size - 1)
            rd.font.name = font_name
            rd.font.color.rgb = RGBColor(*dur_color_rgb)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)

        # Description line (if present)
        if desc and desc not in ("[Not Provided]", ""):
            pd = doc.add_paragraph()
            pd.clear()
            rd2 = pd.add_run(desc)
            rd2.font.size = Pt(font_size - 1)
            rd2.font.name = font_name
            rd2.font.color.rgb = RGBColor(*desc_color_rgb)
            pd.paragraph_format.space_before = Pt(0)
            pd.paragraph_format.space_after = Pt(2)


# ─── MODERN TEMPLATE ──────────────────────────────────────────────────────────
def generate_modern_docx(data: dict) -> BytesIO:
    """
    Modern ATS-Optimized template — single-column, Calibri font, navy headings.
    Strictly follows ATS section ordering used by Workday, Greenhouse, and Lever:
      Header → Professional Summary → Skills → Work Experience →
      Projects → Education → Certifications → Languages → Interests → Additional

    All formatting decisions prioritize machine readability over visual design.
    No tables, no columns, no text boxes — pure linear paragraph flow for ATS parsers.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    NAVY = (0x1E, 0x3A, 0x5F)
    NAVY_HEX = "1E3A5F"
    FONT = "Calibri"
    BODY = 10

    # ── ATS RULE: Single-column header block — centered plain text, no tables ──
    _build_contact_header(
        doc, data,
        name_size=20, name_color_rgb=NAVY, name_font=FONT,
        contact_font=FONT, contact_color_hex="4A4A4A", contact_size=9,
        title_font=FONT, title_size=11, title_color_rgb=NAVY,
        accent_color_hex=NAVY_HEX,
    )

    def _heading(text):
        """
        ATS-standard section heading: ALL CAPS, bold, bottom-bordered.
        Border signals section boundary to ATS parsers without using tables.
        Matches heading labels expected by Workday/Greenhouse parsers.
        """
        _section_heading_bordered(doc, text, font_name=FONT, font_size=BODY,
                                   bold=True, color_hex=NAVY_HEX,
                                   border_color=NAVY_HEX, border_sz="6",
                                   space_before=10, space_after=4)

    def _body_para(text, italic=False, color_rgb=None):
        """Standard body paragraph — consistent 10pt Calibri, minimal spacing."""
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(BODY)
        run.font.name = FONT
        run.italic = italic
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        return p

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: PROFESSIONAL SUMMARY
    # ATS parsers (Workday, iCIMS) actively scan this section for role fit.
    # 2-3 sentences: Role identity + core competencies + value proposition.
    # ══════════════════════════════════════════════════════════════════════
    if data.get("summary") and data["summary"] not in ("", "[Not Provided]"):
        _heading("Professional Summary")
        _body_para(data["summary"])

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: CORE SKILLS
    # ATS keyword scanners parse Skills first for job description matching.
    # Categorized format (Technical / Professional) with pipe-separated values
    # is the industry standard used by Jobscan, Enhancv, and Greenhouse parsers.
    # ══════════════════════════════════════════════════════════════════════
    tech_skills = [s for s in data.get("skills", []) if s and s != "[Not Provided]"]
    soft_skills = [s for s in data.get("soft_skills", []) if s and s != "[Not Provided]"]
    if tech_skills or soft_skills:
        _heading("Core Skills")
        if tech_skills:
            p = doc.add_paragraph()
            p.clear()
            label_run = p.add_run("Technical:  ")
            label_run.bold = True
            label_run.font.size = Pt(BODY)
            label_run.font.name = FONT
            label_run.font.color.rgb = RGBColor(*NAVY)
            # Group into rows of max 6 skills for readability — ATS reads all as flat text
            skills_text = "  |  ".join(tech_skills)
            skills_run = p.add_run(skills_text)
            skills_run.font.size = Pt(BODY)
            skills_run.font.name = FONT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
        if soft_skills:
            p = doc.add_paragraph()
            p.clear()
            label_run = p.add_run("Professional:  ")
            label_run.bold = True
            label_run.font.size = Pt(BODY)
            label_run.font.name = FONT
            label_run.font.color.rgb = RGBColor(*NAVY)
            ss_run = p.add_run("  |  ".join(soft_skills))
            ss_run.font.size = Pt(BODY)
            ss_run.font.name = FONT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: WORK EXPERIENCE
    # Industry-standard layout (Enhancv/Jobscan):
    #   Line 1: Job Title (bold, navy, larger)  |  Company Name
    #   Line 2: MMM YYYY – MMM YYYY  (italic, right-aligned)
    #   Line 3: 1-sentence role scope (italic gray) — optional
    #   Lines 4+: • Bullet points (action verb + task + tech + impact)
    # ATS parsers map Role Title and Company as separate parsed fields.
    # ══════════════════════════════════════════════════════════════════════
    valid_exp = [e for e in data.get("experience", [])
                 if (e.get("role") and e["role"] not in ("", "[Not Provided]"))
                 or (e.get("company") and e["company"] not in ("", "[Not Provided]"))]
    if valid_exp:
        _heading("Work Experience")
        for exp in valid_exp:
            _role     = exp.get("role", "")     if exp.get("role", "")     not in ("", "[Not Provided]") else ""
            _company  = exp.get("company", "")  if exp.get("company", "")  not in ("", "[Not Provided]") else ""
            _duration = exp.get("duration", "") if exp.get("duration", "") not in ("", "[Not Provided]") else ""

            # ── Row 1: Role Title (bold navy) — Company (regular gray) ──────
            p = doc.add_paragraph()
            p.clear()
            if _role:
                r_role = p.add_run(_role)
                r_role.bold = True
                r_role.font.size = Pt(BODY + 1)
                r_role.font.name = FONT
                r_role.font.color.rgb = RGBColor(*NAVY)
            if _company:
                sep = "  \u2014  " if _role else ""   # em-dash separator — ATS-safe
                r_co = p.add_run(f"{sep}{_company}")
                r_co.font.size = Pt(BODY)
                r_co.font.name = FONT
                r_co.font.color.rgb = RGBColor(74, 74, 74)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(0)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)

            # ── Row 2: Duration (italic, smaller, dark-gray) ─────────────────
            if _duration:
                p_dur = doc.add_paragraph()
                p_dur.clear()
                r_dur = p_dur.add_run(_duration)
                r_dur.italic = True
                r_dur.font.size = Pt(BODY - 1)
                r_dur.font.name = FONT
                r_dur.font.color.rgb = RGBColor(110, 110, 110)
                p_dur.paragraph_format.space_before = Pt(0)
                p_dur.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_dur, keep_together=True, keep_with_next=True, widow_control=True)

            # ── Row 3: Role scope summary (italic gray) ────────────────────
            if exp.get("description") and exp["description"] not in ("", "[Not Provided]"):
                _body_para(exp["description"], italic=True, color_rgb=(90, 90, 90))

            # ── Rows 4+: Achievement bullets ───────────────────────────────
            for b in exp.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: PROJECTS
    # Industry-standard project entry format (Jobscan/Enhancv):
    #   Line 1: Project Name (bold navy) | [Duration] (italic gray)
    #   Line 2: Tech Stack: ... (smaller, plain) | URL (plain text — no hyperlink)
    #   Line 3: 1-sentence project purpose (italic)
    #   Lines 4+: Achievement bullets
    # ATS parsers cannot follow hyperlinks — URLs must be plain text.
    # ══════════════════════════════════════════════════════════════════════
    valid_proj = [p for p in data.get("projects", [])
                  if p.get("name") and p["name"] not in ("", "[Not Provided]")]
    if valid_proj:
        _heading("Projects")
        for proj in valid_proj:
            # Line 1: Project Name + Duration
            p = doc.add_paragraph()
            p.clear()
            r_name = p.add_run(proj.get("name", ""))
            r_name.bold = True
            r_name.font.size = Pt(BODY + 1)
            r_name.font.name = FONT
            r_name.font.color.rgb = RGBColor(*NAVY)
            if proj.get("duration") and proj["duration"] not in ("", "[Not Provided]"):
                r_dur = p.add_run(f"   [{proj['duration']}]")
                r_dur.italic = True
                r_dur.font.size = Pt(BODY - 1)
                r_dur.font.name = FONT
                r_dur.font.color.rgb = RGBColor(110, 110, 110)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            # Line 2: Tech Stack (plain text) + URL (clickable hyperlink)
            has_tech = proj.get("tech_stack") and proj["tech_stack"] not in ("", "[Not Provided]")
            has_url  = proj.get("url") and proj["url"] not in ("", "[Not Provided]")
            if has_tech or has_url:
                p_meta = doc.add_paragraph()
                p_meta.clear()
                if has_tech:
                    r_tech = p_meta.add_run(f"Tech: {proj['tech_stack']}")
                    r_tech.font.size = Pt(BODY - 1)
                    r_tech.font.name = FONT
                    r_tech.font.color.rgb = RGBColor(74, 74, 74)
                if has_tech and has_url:
                    sep = p_meta.add_run("  |  ")
                    sep.font.size = Pt(BODY - 1)
                    sep.font.name = FONT
                    sep.font.color.rgb = RGBColor(74, 74, 74)
                if has_url:
                    _add_hyperlink(p_meta, proj["url"], proj["url"], font_name=FONT, font_size=BODY - 1, color_rgb=(0, 102, 204))
                p_meta.paragraph_format.space_before = Pt(0)
                p_meta.paragraph_format.space_after = Pt(2)
            if proj.get("description") and proj["description"] not in ("", "[Not Provided]"):
                _body_para(proj["description"], italic=True, color_rgb=(80, 80, 80))
            for b in proj.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: EDUCATION
    # Two-line entry: Degree + Institution on line 1, Year on line 2.
    # ATS parsers match: "Bachelor", "Master", "B.Tech", "MBA", "Ph.D".
    # ══════════════════════════════════════════════════════════════════════
    valid_edu = [e for e in data.get("education", [])
                 if e.get("degree") or e.get("institution")]
    if valid_edu:
        _heading("Education")
        for edu in valid_edu:
            p = doc.add_paragraph()
            p.clear()
            if edu.get("degree") and edu["degree"] not in ("", "[Not Provided]"):
                r_deg = p.add_run(edu["degree"])
                r_deg.bold = True
                r_deg.font.size = Pt(BODY + 1)
                r_deg.font.name = FONT
                r_deg.font.color.rgb = RGBColor(*NAVY)
            if edu.get("institution") and edu["institution"] not in ("", "[Not Provided]"):
                r_inst = p.add_run(f"  —  {edu['institution']}")
                r_inst.font.size = Pt(BODY)
                r_inst.font.name = FONT
                r_inst.font.color.rgb = RGBColor(74, 74, 74)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("year") and edu["year"] not in ("", "[Not Provided]"):
                p_yr = doc.add_paragraph()
                p_yr.clear()
                r_yr = p_yr.add_run(edu["year"])
                r_yr.italic = True
                r_yr.font.size = Pt(BODY - 1)
                r_yr.font.name = FONT
                r_yr.font.color.rgb = RGBColor(110, 110, 110)
                p_yr.paragraph_format.space_before = Pt(0)
                p_yr.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_yr, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("cgpa") and edu["cgpa"] not in ("", "[Not Provided]"):
                p_cgpa = doc.add_paragraph()
                p_cgpa.clear()
                _cgpa_val = edu['cgpa']
                # ── Smart grade label — delegates to _normalize_cgpa() ───────
                _cgpa_display = _normalize_cgpa(str(_cgpa_val).strip())
                r_cgpa = p_cgpa.add_run(_cgpa_display)
                r_cgpa.font.size = Pt(BODY - 1)
                r_cgpa.font.name = FONT
                r_cgpa.font.color.rgb = RGBColor(80, 80, 80)
                p_cgpa.paragraph_format.space_before = Pt(0)
                p_cgpa.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_cgpa, keep_together=True, keep_with_next=False, widow_control=True)
            for b in (edu.get("bullets") or []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY - 1, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: CERTIFICATIONS & LINKS
    # Each cert: Name (bold) | Issuer | Date — one bullet per credential.
    # Followed by LinkedIn / GitHub / Portfolio as plain-text URL bullets.
    # ATS systems (LinkedIn, Workday) match cert keywords directly.
    # ══════════════════════════════════════════════════════════════════════
    valid_certs = [c for c in data.get("certifications", [])
                   if isinstance(c, dict) and c.get("name") and c["name"] not in ("", "[Not Provided]")]
    contact = data.get("contact", {})
    has_links = any(contact.get(k, "") not in ("", "[Not Provided]", "Not Provided")
                    for k in ("linkedin", "github", "portfolio"))
    if valid_certs or has_links:
        _heading("Certifications & Links")
        for cert in valid_certs:
            parts = [cert["name"]]
            if cert.get("issuer") and cert["issuer"] not in ("", "[Not Provided]"):
                parts.append(cert["issuer"])
            if cert.get("duration") and cert["duration"] not in ("", "[Not Provided]"):
                parts.append(cert["duration"])
            _add_bullet(doc, "  |  ".join(parts), font_size=BODY, font_name=FONT)
        # Profile links as real clickable hyperlinks
        for label, key in [("LinkedIn", "linkedin"), ("GitHub", "github"), ("Portfolio", "portfolio")]:
            val = contact.get(key, "")
            if val and val not in ("", "[Not Provided]", "Not Provided"):
                p_link = doc.add_paragraph()
                p_link.paragraph_format.left_indent = Pt(0)
                p_link.paragraph_format.space_before = Pt(2)
                p_link.paragraph_format.space_after = Pt(2)
                label_run = p_link.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.size = Pt(BODY)
                label_run.font.name = FONT
                label_run.font.color.rgb = RGBColor(*NAVY)
                _add_hyperlink(p_link, val, val, font_name=FONT, font_size=BODY, color_rgb=(0, 102, 204))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7: LANGUAGES
    # ══════════════════════════════════════════════════════════════════════
    valid_lang = [l for l in data.get("languages", []) if l and l != "[Not Provided]"]
    if valid_lang:
        _heading("Languages")
        _body_para("  |  ".join(valid_lang))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8: INTERESTS
    # ══════════════════════════════════════════════════════════════════════
    valid_int = [i for i in data.get("interests", []) if i and i != "[Not Provided]"]
    if valid_int:
        _heading("Interests")
        _body_para("  |  ".join(valid_int))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9: ADDITIONAL (Awards, Training, Volunteering, Publications)
    # ══════════════════════════════════════════════════════════════════════
    _render_additional(doc, data, font_name=FONT, font_size=BODY,
                       heading_fn=_heading,
                       bullet_fn=lambda t: _add_bullet(doc, t, font_size=BODY, font_name=FONT),
                       name_color_rgb=NAVY, desc_color_rgb=(80, 80, 80))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── MINIMAL TEMPLATE ─────────────────────────────────────────────────────────
def generate_minimal_docx(data: dict) -> BytesIO:
    """
    Minimal ATS-Optimized template — pure black/white Arial, maximum machine readability.
    Highest ATS parse accuracy of all three templates.
    Follows identical section ordering to Modern template for ATS consistency:
      Header → Professional Summary → Skills → Work Experience →
      Projects → Education → Certifications → Languages → Interests → Additional

    No color, no decoration, no graphics — every byte serves ATS keyword matching.
    Preferred by Taleo, SmartRecruiters, and legacy HRIS systems.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(1.1)
        sec.right_margin = Inches(1.1)

    FONT = "Arial"
    BODY = 10
    BLACK_HEX = "000000"
    BLACK = (0, 0, 0)
    DARK_GRAY = (60, 60, 60)
    MID_GRAY = (100, 100, 100)

    # ── ATS RULE: Plain-text header — no color, no decoration ──
    _build_contact_header(
        doc, data,
        name_size=18, name_color_rgb=BLACK, name_font=FONT,
        contact_font=FONT, contact_color_hex="333333", contact_size=9,
        title_font=FONT, title_size=10, title_color_rgb=DARK_GRAY,
    )

    def _heading(text):
        """
        Pure black bold ALL-CAPS heading with bottom rule.
        Maximum compatibility with legacy ATS parsers that strip color.
        """
        _section_heading_bordered(doc, text, font_name=FONT, font_size=BODY,
                                   bold=True, color_hex=BLACK_HEX,
                                   border_color=BLACK_HEX, border_sz="4",
                                   space_before=10, space_after=3)

    def _body_para(text, italic=False):
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(BODY)
        run.font.name = FONT
        run.italic = italic
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        return p

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: PROFESSIONAL SUMMARY
    # ══════════════════════════════════════════════════════════════════════
    if data.get("summary") and data["summary"] not in ("", "[Not Provided]"):
        _heading("Professional Summary")
        _body_para(data["summary"])

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: CORE SKILLS
    # Comma-separated — maximum ATS keyword parser compatibility.
    # Taleo, SmartRecruiters, and legacy HRIS systems parse comma lists best.
    # Labeled "Technical" and "Professional" — matches Greenhouse/Lever field names.
    # ══════════════════════════════════════════════════════════════════════
    tech_skills = [s for s in data.get("skills", []) if s and s != "[Not Provided]"]
    soft_skills = [s for s in data.get("soft_skills", []) if s and s != "[Not Provided]"]
    if tech_skills or soft_skills:
        _heading("Core Skills")
        if tech_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Technical:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT
            skills_run = p.add_run(", ".join(tech_skills))
            skills_run.font.size = Pt(BODY)
            skills_run.font.name = FONT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
        if soft_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Professional:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT
            ss_run = p.add_run(", ".join(soft_skills))
            ss_run.font.size = Pt(BODY)
            ss_run.font.name = FONT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: WORK EXPERIENCE
    # Industry-standard two-line layout for maximum ATS field recognition:
    #   Line 1: Job Title (bold, black)  |  Duration (italic, gray)
    #   Line 2: Company Name (plain)
    #   Line 3: Role scope (italic, optional)
    #   Lines 4+: Bullet achievements
    # Plain black/white — preferred by Taleo and legacy HRIS parsers.
    # ══════════════════════════════════════════════════════════════════════
    valid_exp = [e for e in data.get("experience", [])
                 if (e.get("role") and e["role"] not in ("", "[Not Provided]"))
                 or (e.get("company") and e["company"] not in ("", "[Not Provided]"))]
    if valid_exp:
        _heading("Work Experience")
        for exp in valid_exp:
            _role     = exp.get("role", "")     if exp.get("role", "")     not in ("", "[Not Provided]") else ""
            _company  = exp.get("company", "")  if exp.get("company", "")  not in ("", "[Not Provided]") else ""
            _duration = exp.get("duration", "") if exp.get("duration", "") not in ("", "[Not Provided]") else ""

            # Line 1: Job Title + Duration
            p = doc.add_paragraph()
            p.clear()
            if _role:
                r1 = p.add_run(_role)
                r1.bold = True
                r1.font.size = Pt(BODY + 1)
                r1.font.name = FONT
            if _duration:
                r2 = p.add_run(f"  |  {_duration}")
                r2.italic = True
                r2.font.size = Pt(BODY - 1)
                r2.font.name = FONT
                r2.font.color.rgb = RGBColor(*MID_GRAY)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(1)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)

            # Line 2: Company Name
            if _company:
                p_co = doc.add_paragraph()
                p_co.clear()
                r_co = p_co.add_run(_company)
                r_co.font.size = Pt(BODY)
                r_co.font.name = FONT
                r_co.font.color.rgb = RGBColor(*DARK_GRAY)
                p_co.paragraph_format.space_before = Pt(0)
                p_co.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_co, keep_together=True, keep_with_next=True, widow_control=True)

            if exp.get("description") and exp["description"] not in ("", "[Not Provided]"):
                _body_para(exp["description"], italic=True)
            for b in exp.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: PROJECTS
    # Plain black project header: Name (bold) + [Duration] + tech + URL
    # Each on its own line for maximum ATS field parsing accuracy.
    # ══════════════════════════════════════════════════════════════════════
    valid_proj = [p for p in data.get("projects", [])
                  if p.get("name") and p["name"] not in ("", "[Not Provided]")]
    if valid_proj:
        _heading("Projects")
        for proj in valid_proj:
            # Line 1: Project Name + Duration
            p = doc.add_paragraph()
            p.clear()
            r1 = p.add_run(proj.get("name", ""))
            r1.bold = True
            r1.font.size = Pt(BODY + 1)
            r1.font.name = FONT
            if proj.get("duration") and proj["duration"] not in ("", "[Not Provided]"):
                rd = p.add_run(f"   [{proj['duration']}]")
                rd.italic = True
                rd.font.size = Pt(BODY - 1)
                rd.font.name = FONT
                rd.font.color.rgb = RGBColor(*MID_GRAY)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            # Line 2: Tech Stack (plain) + URL (clickable hyperlink)
            has_tech = proj.get("tech_stack") and proj["tech_stack"] not in ("", "[Not Provided]")
            has_url  = proj.get("url") and proj["url"] not in ("", "[Not Provided]")
            if has_tech or has_url:
                p_meta = doc.add_paragraph()
                p_meta.clear()
                if has_tech:
                    r_tech = p_meta.add_run(f"Tech: {proj['tech_stack']}")
                    r_tech.font.size = Pt(BODY - 1)
                    r_tech.font.name = FONT
                    r_tech.font.color.rgb = RGBColor(*DARK_GRAY)
                if has_tech and has_url:
                    sep = p_meta.add_run("  |  ")
                    sep.font.size = Pt(BODY - 1)
                    sep.font.name = FONT
                    sep.font.color.rgb = RGBColor(*DARK_GRAY)
                if has_url:
                    _add_hyperlink(p_meta, proj["url"], proj["url"], font_name=FONT, font_size=BODY - 1, color_rgb=(0, 0, 0))
                p_meta.paragraph_format.space_before = Pt(0)
                p_meta.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_meta, keep_together=True, keep_with_next=True, widow_control=True)
            if proj.get("description") and proj["description"] not in ("", "[Not Provided]"):
                _body_para(proj["description"], italic=True)
            for b in proj.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: EDUCATION
    # Degree (bold) — Institution — Year (italic)
    # ATS parsers match: "Bachelor", "Master", "B.Tech", "MBA", "Ph.D".
    # ══════════════════════════════════════════════════════════════════════
    valid_edu = [e for e in data.get("education", [])
                 if e.get("degree") or e.get("institution")]
    if valid_edu:
        _heading("Education")
        for edu in valid_edu:
            p = doc.add_paragraph()
            p.clear()
            if edu.get("degree") and edu["degree"] not in ("", "[Not Provided]"):
                r_deg = p.add_run(edu["degree"])
                r_deg.bold = True
                r_deg.font.size = Pt(BODY + 1)
                r_deg.font.name = FONT
            if edu.get("institution") and edu["institution"] not in ("", "[Not Provided]"):
                r_inst = p.add_run(f"  —  {edu['institution']}")
                r_inst.font.size = Pt(BODY)
                r_inst.font.name = FONT
                r_inst.font.color.rgb = RGBColor(*DARK_GRAY)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("year") and edu["year"] not in ("", "[Not Provided]"):
                p_yr = doc.add_paragraph()
                p_yr.clear()
                r_yr = p_yr.add_run(edu["year"])
                r_yr.italic = True
                r_yr.font.size = Pt(BODY - 1)
                r_yr.font.name = FONT
                r_yr.font.color.rgb = RGBColor(*MID_GRAY)
                p_yr.paragraph_format.space_before = Pt(0)
                p_yr.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_yr, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("cgpa") and edu["cgpa"] not in ("", "[Not Provided]"):
                p_cgpa = doc.add_paragraph()
                p_cgpa.clear()
                _cgpa_val = edu['cgpa']
                # ── Smart grade label — delegates to _normalize_cgpa() ───────
                _cgpa_display = _normalize_cgpa(str(_cgpa_val).strip())
                r_cgpa = p_cgpa.add_run(_cgpa_display)
                r_cgpa.font.size = Pt(BODY - 1)
                r_cgpa.font.name = FONT
                r_cgpa.font.color.rgb = RGBColor(80, 80, 80)
                p_cgpa.paragraph_format.space_before = Pt(0)
                p_cgpa.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_cgpa, keep_together=True, keep_with_next=False, widow_control=True)
            for b in (edu.get("bullets") or []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY - 1, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: CERTIFICATIONS & LINKS
    # Plain-text bullet per cert: Name | Issuer | Date
    # Profile URLs added as plain-text bullets — ATS-safe (no hyperlinks).
    # ══════════════════════════════════════════════════════════════════════
    valid_certs = [c for c in data.get("certifications", [])
                   if isinstance(c, dict) and c.get("name") and c["name"] not in ("", "[Not Provided]")]
    contact_min = data.get("contact", {})
    has_links_min = any(contact_min.get(k, "") not in ("", "[Not Provided]", "Not Provided")
                        for k in ("linkedin", "github", "portfolio"))
    if valid_certs or has_links_min:
        _heading("Certifications & Links")
        for cert in valid_certs:
            parts = [cert["name"]]
            if cert.get("issuer") and cert["issuer"] not in ("", "[Not Provided]"):
                parts.append(cert["issuer"])
            if cert.get("duration") and cert["duration"] not in ("", "[Not Provided]"):
                parts.append(cert["duration"])
            _add_bullet(doc, "  |  ".join(parts), font_size=BODY, font_name=FONT)
        for label, key in [("LinkedIn", "linkedin"), ("GitHub", "github"), ("Portfolio", "portfolio")]:
            val = contact_min.get(key, "")
            if val and val not in ("", "[Not Provided]", "Not Provided"):
                p_link = doc.add_paragraph()
                p_link.paragraph_format.space_before = Pt(2)
                p_link.paragraph_format.space_after = Pt(2)
                label_run = p_link.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.size = Pt(BODY)
                label_run.font.name = FONT
                _add_hyperlink(p_link, val, val, font_name=FONT, font_size=BODY, color_rgb=(0, 0, 0))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7: LANGUAGES
    # ══════════════════════════════════════════════════════════════════════
    valid_lang = [l for l in data.get("languages", []) if l and l != "[Not Provided]"]
    if valid_lang:
        _heading("Languages")
        _body_para(", ".join(valid_lang))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8: INTERESTS
    # ══════════════════════════════════════════════════════════════════════
    valid_int = [i for i in data.get("interests", []) if i and i != "[Not Provided]"]
    if valid_int:
        _heading("Interests")
        _body_para(", ".join(valid_int))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9: ADDITIONAL
    # ══════════════════════════════════════════════════════════════════════
    _render_additional(doc, data, font_name=FONT, font_size=BODY,
                       heading_fn=_heading,
                       bullet_fn=lambda t: _add_bullet(doc, t, font_size=BODY, font_name=FONT),
                       name_color_rgb=BLACK, desc_color_rgb=(80, 80, 80))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── CREATIVE TEMPLATE ────────────────────────────────────────────────────────
def generate_creative_docx(data: dict) -> BytesIO:
    """
    Executive ATS-Optimized template — teal/dark-navy Calibri, polished yet ATS-safe.
    Replaces design-heavy elements (decorative symbols, multi-font mixing, ◆/▌ glyphs)
    with ATS-safe equivalents that retain visual appeal without breaking parsers.
    Identical section ordering to Modern and Minimal templates for ATS consistency:
      Header → Professional Summary → Skills → Work Experience →
      Projects → Education → Certifications → Languages → Interests → Additional

    ATS COMPLIANCE NOTES:
    - All Unicode decorative characters (◆, ▌, @) removed from section headings/bullets.
    - Single font family (Calibri) used throughout body — multi-font mixing confuses some parsers.
    - Teal color is display-only; ATS parsers read plain text, not colors.
    - Georgia used ONLY for candidate name (header) — never in parseable body sections.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.95)
        sec.right_margin = Inches(0.95)

    TEAL_HEX = "0D7377"
    DARK_HEX = "14213D"
    TEAL = (0x0D, 0x73, 0x77)
    DARK = (0x14, 0x21, 0x3D)
    FONT_NAME = "Georgia"   # Name display only — ATS ignores header fonts
    FONT_BODY = "Calibri"   # Consistent body font — required for ATS parsing
    BODY = 10

    # ── Header: Name in Georgia (display), contact in Calibri (parseable) ──
    _build_contact_header(
        doc, data,
        name_size=22, name_color_rgb=DARK, name_font=FONT_NAME,
        contact_font=FONT_BODY, contact_color_hex="444444", contact_size=9,
        title_font=FONT_BODY, title_size=11, title_color_rgb=TEAL,
        accent_color_hex=TEAL_HEX,
    )

    def _heading(text):
        """
        ATS-safe teal heading — NO prefix symbols (◆/▌ break some ATS parsers).
        Teal color preserved for human readers; ATS reads underlying plain text.
        """
        _section_heading_bordered(doc, text, font_name=FONT_BODY, font_size=BODY + 1,
                                   bold=True, color_hex=TEAL_HEX,
                                   border_color=TEAL_HEX, border_sz="6",
                                   space_before=10, space_after=4)
        # NOTE: prefix="\u258c " removed — block character disrupts ATS text extraction

    def _body_para(text, italic=False, color_rgb=None):
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(BODY)
        run.font.name = FONT_BODY
        run.italic = italic
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        return p

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: PROFESSIONAL SUMMARY
    # "Profile" renamed to "Professional Summary" — standard ATS label.
    # Some older ATS systems (Taleo) fail to map "Profile" to summary field.
    # ══════════════════════════════════════════════════════════════════════
    if data.get("summary") and data["summary"] not in ("", "[Not Provided]"):
        _heading("Professional Summary")
        _body_para(data["summary"], italic=False)  # Not italic — ATS reads italic as emphasis, not content

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: CORE SKILLS
    # Teal labels + pipe-separated values — visually distinctive, ATS-safe.
    # Georgia used ONLY for name; Calibri throughout body for ATS compatibility.
    # ══════════════════════════════════════════════════════════════════════
    tech_skills = [s for s in data.get("skills", []) if s and s != "[Not Provided]"]
    soft_skills = [s for s in data.get("soft_skills", []) if s and s != "[Not Provided]"]
    if tech_skills or soft_skills:
        _heading("Core Skills")
        if tech_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Technical:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT_BODY
            lbl.font.color.rgb = RGBColor(*TEAL)
            skills_run = p.add_run("  |  ".join(tech_skills))
            skills_run.font.size = Pt(BODY)
            skills_run.font.name = FONT_BODY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
        if soft_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Professional:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT_BODY
            lbl.font.color.rgb = RGBColor(*TEAL)
            ss_run = p.add_run("  |  ".join(soft_skills))
            ss_run.font.size = Pt(BODY)
            ss_run.font.name = FONT_BODY
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: WORK EXPERIENCE
    # Two-line entry (Enhancv/Jobscan standard):
    #   Line 1: Job Title (bold dark-navy, 11pt)  |  Duration (italic gray)
    #   Line 2: Company Name (teal, 10pt)
    #   Line 3: Role scope (italic, optional)
    #   Lines 4+: Achievement bullets
    # ATS-safe: "@" separator replaced with " — " (em-dash).
    # ══════════════════════════════════════════════════════════════════════
    valid_exp = [e for e in data.get("experience", [])
                 if (e.get("role") and e["role"] not in ("", "[Not Provided]"))
                 or (e.get("company") and e["company"] not in ("", "[Not Provided]"))]
    if valid_exp:
        _heading("Work Experience")
        for exp in valid_exp:
            _role     = exp.get("role", "")     if exp.get("role", "")     not in ("", "[Not Provided]") else ""
            _company  = exp.get("company", "")  if exp.get("company", "")  not in ("", "[Not Provided]") else ""
            _duration = exp.get("duration", "") if exp.get("duration", "") not in ("", "[Not Provided]") else ""

            # Line 1: Job Title + Duration
            p = doc.add_paragraph()
            p.clear()
            if _role:
                r1 = p.add_run(_role)
                r1.bold = True
                r1.font.size = Pt(BODY + 1)
                r1.font.name = FONT_BODY
                r1.font.color.rgb = RGBColor(*DARK)
            if _duration:
                r3 = p.add_run(f"   {_duration}")
                r3.italic = True
                r3.font.size = Pt(BODY - 1)
                r3.font.name = FONT_BODY
                r3.font.color.rgb = RGBColor(110, 110, 110)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(1)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)

            # Line 2: Company Name (teal)
            if _company:
                p_co = doc.add_paragraph()
                p_co.clear()
                r_co = p_co.add_run(_company)
                r_co.font.size = Pt(BODY)
                r_co.font.name = FONT_BODY
                r_co.font.color.rgb = RGBColor(*TEAL)
                p_co.paragraph_format.space_before = Pt(0)
                p_co.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_co, keep_together=True, keep_with_next=True, widow_control=True)

            if exp.get("description") and exp["description"] not in ("", "[Not Provided]"):
                _body_para(exp["description"], italic=True, color_rgb=(80, 80, 80))
            for b in exp.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT_BODY)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: PROJECTS
    # Project Name (bold dark-navy, 11pt) + [Duration] (italic gray)
    # Tech Stack line (teal, 9pt) + URL as plain text (ATS-safe — no hyperlink)
    # ══════════════════════════════════════════════════════════════════════
    valid_proj = [pr for pr in data.get("projects", [])
                  if pr.get("name") and pr["name"] not in ("", "[Not Provided]")]
    if valid_proj:
        _heading("Projects")
        for proj in valid_proj:
            # Line 1: Project Name + Duration
            p = doc.add_paragraph()
            p.clear()
            r1 = p.add_run(proj.get("name", ""))
            r1.bold = True
            r1.font.size = Pt(BODY + 1)
            r1.font.name = FONT_BODY
            r1.font.color.rgb = RGBColor(*DARK)
            if proj.get("duration") and proj["duration"] not in ("", "[Not Provided]"):
                rd = p.add_run(f"   [{proj['duration']}]")
                rd.italic = True
                rd.font.size = Pt(BODY - 1)
                rd.font.name = FONT_BODY
                rd.font.color.rgb = RGBColor(110, 110, 110)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            # Line 2: Tech (teal plain text) + URL (clickable hyperlink)
            has_tech = proj.get("tech_stack") and proj["tech_stack"] not in ("", "[Not Provided]")
            has_url  = proj.get("url") and proj["url"] not in ("", "[Not Provided]")
            if has_tech or has_url:
                p_meta = doc.add_paragraph()
                p_meta.clear()
                if has_tech:
                    r_tech = p_meta.add_run(f"Tech: {proj['tech_stack']}")
                    r_tech.font.size = Pt(BODY - 1)
                    r_tech.font.name = FONT_BODY
                    r_tech.font.color.rgb = RGBColor(*TEAL)
                if has_tech and has_url:
                    sep = p_meta.add_run("  |  ")
                    sep.font.size = Pt(BODY - 1)
                    sep.font.name = FONT_BODY
                    sep.font.color.rgb = RGBColor(*TEAL)
                if has_url:
                    _add_hyperlink(p_meta, proj["url"], proj["url"], font_name=FONT_BODY, font_size=BODY - 1, color_rgb=(0, 128, 128))
                p_meta.paragraph_format.space_before = Pt(0)
                p_meta.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_meta, keep_together=True, keep_with_next=True, widow_control=True)
            if proj.get("description") and proj["description"] not in ("", "[Not Provided]"):
                _body_para(proj["description"], italic=True, color_rgb=(80, 80, 80))
            for b in proj.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT_BODY)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: EDUCATION
    # Degree (bold dark-navy) — Institution (teal) — Year (italic gray)
    # ══════════════════════════════════════════════════════════════════════
    valid_edu = [e for e in data.get("education", [])
                 if e.get("degree") or e.get("institution")]
    if valid_edu:
        _heading("Education")
        for edu in valid_edu:
            p = doc.add_paragraph()
            p.clear()
            if edu.get("degree") and edu["degree"] not in ("", "[Not Provided]"):
                r1 = p.add_run(edu["degree"])
                r1.bold = True
                r1.font.size = Pt(BODY + 1)
                r1.font.name = FONT_BODY
                r1.font.color.rgb = RGBColor(*DARK)
            if edu.get("institution") and edu["institution"] not in ("", "[Not Provided]"):
                r2 = p.add_run(f"  —  {edu['institution']}")
                r2.font.size = Pt(BODY)
                r2.font.name = FONT_BODY
                r2.font.color.rgb = RGBColor(*TEAL)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            _set_para_keep(p, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("year") and edu["year"] not in ("", "[Not Provided]"):
                p_yr = doc.add_paragraph()
                p_yr.clear()
                r3 = p_yr.add_run(edu["year"])
                r3.italic = True
                r3.font.size = Pt(BODY - 1)
                r3.font.name = FONT_BODY
                r3.font.color.rgb = RGBColor(110, 110, 110)
                p_yr.paragraph_format.space_before = Pt(0)
                p_yr.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_yr, keep_together=True, keep_with_next=True, widow_control=True)
            if edu.get("cgpa") and edu["cgpa"] not in ("", "[Not Provided]"):
                p_cgpa = doc.add_paragraph()
                p_cgpa.clear()
                _cgpa_val = edu['cgpa']
                # ── Smart grade label — delegates to _normalize_cgpa() ───────
                _cgpa_display = _normalize_cgpa(str(_cgpa_val).strip())
                r_cgpa = p_cgpa.add_run(_cgpa_display)
                r_cgpa.font.size = Pt(BODY - 1)
                r_cgpa.font.name = FONT_BODY
                r_cgpa.font.color.rgb = RGBColor(80, 80, 80)
                p_cgpa.paragraph_format.space_before = Pt(0)
                p_cgpa.paragraph_format.space_after = Pt(2)
                _set_para_keep(p_cgpa, keep_together=True, keep_with_next=False, widow_control=True)
            for b in (edu.get("bullets") or []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY - 1, font_name=FONT_BODY)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: CERTIFICATIONS & LINKS
    # Cert bullets: Name | Issuer | Date  (teal heading, plain bullets)
    # Profile URLs as plain-text bullets — no hyperlinks (ATS-safe).
    # ══════════════════════════════════════════════════════════════════════
    valid_certs = [c for c in data.get("certifications", [])
                   if isinstance(c, dict) and c.get("name") and c["name"] not in ("", "[Not Provided]")]
    contact_exec = data.get("contact", {})
    has_links_exec = any(contact_exec.get(k, "") not in ("", "[Not Provided]", "Not Provided")
                         for k in ("linkedin", "github", "portfolio"))
    if valid_certs or has_links_exec:
        _heading("Certifications & Links")
        for cert in valid_certs:
            parts = [cert["name"]]
            if cert.get("issuer") and cert["issuer"] not in ("", "[Not Provided]"):
                parts.append(cert["issuer"])
            if cert.get("duration") and cert["duration"] not in ("", "[Not Provided]"):
                parts.append(cert["duration"])
            _add_bullet(doc, "  |  ".join(parts), font_size=BODY, font_name=FONT_BODY)
        for label, key in [("LinkedIn", "linkedin"), ("GitHub", "github"), ("Portfolio", "portfolio")]:
            val = contact_exec.get(key, "")
            if val and val not in ("", "[Not Provided]", "Not Provided"):
                p_link = doc.add_paragraph()
                p_link.paragraph_format.space_before = Pt(2)
                p_link.paragraph_format.space_after = Pt(2)
                label_run = p_link.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.size = Pt(BODY)
                label_run.font.name = FONT_BODY
                label_run.font.color.rgb = RGBColor(*TEAL)
                _add_hyperlink(p_link, val, val, font_name=FONT_BODY, font_size=BODY, color_rgb=(0, 128, 128))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7: LANGUAGES
    # ══════════════════════════════════════════════════════════════════════
    valid_lang = [l for l in data.get("languages", []) if l and l != "[Not Provided]"]
    if valid_lang:
        _heading("Languages")
        _body_para("  |  ".join(valid_lang))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8: INTERESTS
    # ══════════════════════════════════════════════════════════════════════
    valid_int = [i for i in data.get("interests", []) if i and i != "[Not Provided]"]
    if valid_int:
        _heading("Interests")
        _body_para("  |  ".join(valid_int))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9: ADDITIONAL
    # ══════════════════════════════════════════════════════════════════════
    _render_additional(doc, data, font_name=FONT_BODY, font_size=BODY,
                       heading_fn=_heading,
                       bullet_fn=lambda t: _add_bullet(doc, t, font_size=BODY, font_name=FONT_BODY),
                       name_color_rgb=DARK, desc_color_rgb=(80, 80, 80))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ============================================================
# 🔍 RESUME ANALYSIS MODULE — Job Title Suggestions (UI only)
# ============================================================
# NOTE: rewrite_text_with_llm() above is the Analysis Module.
# It generates job title suggestions + rewritten text for DISPLAY ONLY.
# It is NEVER used for DOCX generation.
# DOCX generation uses optimize_resume_to_json() + extract_resume_json() exclusively.




def generate_resume_report_html(resume, user_location=""):
    candidate_name = resume.get('Candidate Name', 'Not Found')
    resume_name = resume.get('Resume Name', 'Unknown')
    _raw_rewritten = resume.get('Rewritten Text', '')
    if "### 🎯 Suggested Job Titles" in _raw_rewritten:
        _resume_part, _jobs_part = _raw_rewritten.split("### 🎯 Suggested Job Titles", 1)
    else:
        _resume_part, _jobs_part = _raw_rewritten, ""
    rewritten_text = _resume_part.replace("\n", "<br/>")
    _job_titles_html = ""
    _location_param = urllib.parse.quote(user_location) if user_location else "India"
    _location_raw   = user_location.strip() if user_location and user_location.strip() else "India"
    if _jobs_part:
        _job_titles_html = "<div class='section-title'>Suggested Job Titles</div><div class='box'><ul>"
        for _line in _jobs_part.split('\n'):
            _m = re.match(r'^\d+\.\s+\*\*(.+?)\*\*\s*[—-]?\s*(.*)', _line.strip())
            if _m:
                _title = _m.group(1).strip()
                _desc = re.sub(r'https?://\S+', '', _m.group(2)).strip().rstrip('.')
                # LinkedIn only reliably applies `keywords` for logged-out clicks —
                # `location` alone gets silently dropped, so fold it into keywords too.
                _encoded = urllib.parse.quote(f"{_title},{_location_raw}")
                _url = f"https://www.linkedin.com/jobs/search/?keywords={_encoded}&location={_location_param}"
                _job_titles_html += f'<li><b><a href="{_url}">{_title}</a></b>{(" — " + _desc) if _desc else ""}</li>'
        _job_titles_html += "</ul></div>"

    masculine_words_list = resume.get("Detected Masculine Words", [])
    masculine_words = "".join(
        f"<b>{item.get('word','')}</b>: {item.get('sentence','')}<br/>"
        for item in masculine_words_list
    ) if masculine_words_list else "<i>None detected.</i>"

    feminine_words_list = resume.get("Detected Feminine Words", [])
    feminine_words = "".join(
        f"<b>{item.get('word','')}</b>: {item.get('sentence','')}<br/>"
        for item in feminine_words_list
    ) if feminine_words_list else "<i>None detected.</i>"

    ats_report_html = resume.get("ATS Report", "").replace("\n", "<br/>")

    def style_analysis(analysis, fallback="N/A"):
        if not analysis or analysis == "N/A":
            return f"<p><i>{fallback}</i></p>"

        if "**Score:**" in analysis:
            parts = analysis.split("**Score:**")
            rest = parts[1].split("**", 1)
            score_text = rest[0].strip()
            remaining = rest[1].strip() if len(rest) > 1 else ""
            return f"<p><b>Score:</b> {score_text}</p><p>{remaining}</p>"
        else:
            return f"<p>{analysis}</p>"

    edu_analysis = style_analysis(resume.get("Education Analysis", "").replace("\n", "<br/>"))
    exp_analysis = style_analysis(resume.get("Experience Analysis", "").replace("\n", "<br/>"))
    skills_analysis = style_analysis(resume.get("Skills Analysis", "").replace("\n", "<br/>"))
    keyword_analysis = style_analysis(resume.get("Keyword Analysis", "").replace("\n", "<br/>"))
    final_thoughts = resume.get("Final Thoughts", "N/A").replace("\n", "<br/>")

    lang_analysis_raw = resume.get("Language Analysis", "").replace("\n", "<br/>")
    lang_analysis = f"<div>{lang_analysis_raw}</div>" if lang_analysis_raw else "<p><i>No language analysis available.</i></p>"

    ats_match = resume.get('ATS Match %', 'N/A')
    edu_score = resume.get('Education Score', 'N/A')
    exp_score = resume.get('Experience Score', 'N/A')
    skills_score = resume.get('Skills Score', 'N/A')
    lang_score = resume.get('Language Score', 'N/A')
    keyword_score = resume.get('Keyword Score', 'N/A')
    format_score = resume.get('Format Score', 'N/A')
    format_grade = resume.get('Format Grade', 'N/A')
    format_label = resume.get('Format Label', '')
    masculine_count = len(masculine_words_list)
    feminine_count = len(feminine_words_list)
    bias_score = resume.get('Bias Score (0 = Fair, 1 = Biased)', 'N/A')

    return f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Helvetica, sans-serif;
                font-size: 12pt;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.4em;
            }}
            li {{
                margin-bottom: 5px;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 12px;
                border-left: 4px solid #666;
            }}
            .box {{
                padding: 10px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;
            }}
        </style>
    </head>
    <body>

    <h1>Resume Analysis Report</h1>

    <h2>Candidate: {candidate_name}</h2>
    <p><b>Resume File:</b> {resume_name}</p>

    <h2>ATS Evaluation</h2>
    <table>
        <tr><td><b>Overall ATS Match</b></td><td>{ats_match}%</td></tr>
        <tr><td><b>Education Score</b></td><td>{edu_score}</td></tr>
        <tr><td><b>Experience Score</b></td><td>{exp_score}</td></tr>
        <tr><td><b>Skills Score</b></td><td>{skills_score}</td></tr>
        <tr><td><b>Language Score</b></td><td>{lang_score}</td></tr>
        <tr><td><b>Keyword Score</b></td><td>{keyword_score}</td></tr>
        <tr><td><b>Format Score</b></td><td>{format_score}/100 — {format_grade} ({format_label})</td></tr>
    </table>

    <div class="section-title">ATS Report</div>
    <div class="box">{ats_report_html}</div>

    <div class="section-title">Education Analysis</div>
    <div class="box">{edu_analysis}</div>

    <div class="section-title">Experience Analysis</div>
    <div class="box">{exp_analysis}</div>

    <div class="section-title">Skills Analysis</div>
    <div class="box">{skills_analysis}</div>

    <div class="section-title">Language Analysis</div>
    <div class="box">{lang_analysis}</div>

    <div class="section-title">Keyword Analysis</div>
    <div class="box">{keyword_analysis}</div>

    <div class="section-title">Final Thoughts</div>
    <div class="box">{final_thoughts}</div>

    <h2>Gender Bias Analysis</h2>
    <table>
        <tr><td><b>Masculine Words</b></td><td>{masculine_count}</td></tr>
        <tr><td><b>Feminine Words</b></td><td>{feminine_count}</td></tr>
        <tr><td><b>Bias Score (0 = Fair, 1 = Biased)</b></td><td>{bias_score}</td></tr>
    </table>

    <div class="section-title">Masculine Words Detected</div>
    <div class="box">{masculine_words}</div>

    <div class="section-title">Feminine Words Detected</div>
    <div class="box">{feminine_words}</div>

    <h2>Rewritten Bias-Free Resume</h2>
    <div class="box">{rewritten_text}</div>

    {_job_titles_html}

    </body>
    </html>
    """
