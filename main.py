import os
os.environ["STREAMLIT_WATCHDOG"] = "false"
import json
import random
import string
import re
import asyncio
import io
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from job_scam_detector import render_job_scam_detector_tab
import urllib.parse
import base64
from io import BytesIO
from collections import Counter
from datetime import datetime
import time

import streamlit as st
import streamlit.components.v1 as components
from base64 import b64encode
import requests
import fitz
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import altair as alt
from PIL import Image
from pdf2image import convert_from_path
from dotenv import load_dotenv
from nltk.stem import WordNetLemmatizer
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from xhtml2pdf import pisa
from pydantic import BaseModel
from streamlit_pdf_viewer import pdf_viewer
import torch
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from llm_manager import (
    call_llm, load_groq_api_keys, get_healthy_keys, increment_key_usage,
    mark_key_failure, _mem_record_failure, _mem_clear_failure,
    _mem_increment_usage, _async_mark_failure, _async_increment_usage,
    _async_clear_failure,
)
from db_manager import (
    db_manager, insert_candidate, get_top_domains_by_score,
    get_database_stats, detect_domain_from_title_and_description,
    get_domain_similarity
)
from user_login import (
    create_user_table, add_user, complete_registration, verify_user,
    get_logins_today, get_total_registered_users, log_user_action,
    username_exists, email_exists, is_valid_email, get_all_user_logs,
    generate_otp, send_email_otp,
    get_user_by_email, update_password_by_email, is_strong_password,
    domain_has_mx_record, send_login_link, verify_login_token,
    cleanup_expired_login_tokens, check_and_gate_feature,
    record_feature_usage, get_usage_count_last_hour, check_brute_force,
    get_user_email_by_username, send_analysis_email,
)

from resume_processor import (
    get_easyocr_reader, ensure_nltk, generate_docx,
    _extract_page_text_smart, _classify_pdf, _render_scanned_rejection_card,
    extract_text_from_pdf, extract_text_from_images, safe_extract_text,
    _detect_multicolumn_pdf, check_resume_format, _SCANNED_SENTINEL, _NON_ENGLISH_SENTINEL,
    _render_long_resume_warning, _LONG_RESUME_THRESHOLD,
)
from resume_engine import (
    gender_words, detect_bias, replacement_mapping,
    rewrite_and_optimize_resume, rewrite_text_with_llm,
    optimize_resume_to_json, _salvage_additional_str, extract_resume_json,
    rewrite_and_highlight, get_grammar_score_with_llm,
    ats_percentage_score, setup_vectorstore, create_chain,
)
from report_generator import (
    html_to_pdf_bytes, _val, _build_contact_header,
    _section_heading_bordered, _add_hyperlink, _add_bullet,
    _set_para_keep, _add_role_line, _add_project_header,
    _add_education_row, _render_additional,
    generate_modern_docx, generate_minimal_docx, generate_creative_docx,
    generate_resume_report_html,
)

# ── TAB_1_RESUME.py — Main UI Entrypoint ─────────────────────────────────────
@st.cache_data(ttl=60)
def _cached_hero_stats():
    return (
        get_total_registered_users(),
        get_logins_today(),
        get_database_stats(),
    )

@st.cache_data(ttl=30)   # admin panel metrics — slightly fresher
def _cached_admin_metrics():
    return (
        get_total_registered_users(),
        get_logins_today(),
        get_all_user_logs(),
    )

def generate_cover_letter_from_resume_builder():
    name = st.session_state.get("name", "")
    job_title = st.session_state.get("job_title", "")
    summary = st.session_state.get("summary", "")
    skills = st.session_state.get("skills", "")
    location = st.session_state.get("location", "")
    today_date = datetime.today().strftime("%B %d, %Y")

    # ✅ Input boxes for contact info
    company = st.text_input("🏢 Target Company", placeholder="e.g., Google")
    linkedin = st.text_input("🔗 LinkedIn URL", placeholder="e.g., https://linkedin.com/in/username")
    email = st.text_input("📧 Email", placeholder="e.g., you@example.com")
    mobile = st.text_input("📞 Mobile Number", placeholder="e.g., +91 9876543210")

    # ✅ Button to prevent relooping
    if st.button("✉️ Generate Cover Letter"):
        # ✅ Validate input before generating
        if not all([name, job_title, summary, skills, company, linkedin, email, mobile]):
            st.warning("⚠️ Please fill in all fields including LinkedIn, email, and mobile.")
            return

        prompt = f"""
You are a world-class executive cover letter writer with 20+ years of experience helping candidates land roles at top-tier companies.

Write a compelling, personalized cover letter using the candidate information below.

COVER LETTER STRUCTURE:
1. **Header** — Date, Hiring Manager, Company Name
2. **Opening Paragraph** — Hook with specific value proposition; mention the role and why this company specifically
3. **Core Value Paragraph** — Top 2 achievements from their background (quantified if possible); connect directly to company's likely needs
4. **Skills-Fit Paragraph** — Bridge candidate skills to the role requirements; show cultural awareness
5. **Closing Paragraph** — Confident call to action; express enthusiasm; professional closing

TONE: Professional, confident, specific — NOT generic. Avoid clichés like "I am passionate about..." or "I believe I would be a great fit."
INCLUDE the contact block at the very top: Name, LinkedIn, Email, Phone.
ENSURE company name appears only once (in the header or salutation).
LENGTH: 3 short-to-medium paragraphs. Maximum 350 words.

### CANDIDATE DETAILS:
- Full Name: {name}
- Job Title Applying For: {job_title}
- Professional Summary: {summary}
- Key Skills: {skills}
- Location: {location}
- Date: {today_date}

### COMPANY DETAILS:
- Target Company: {company}
- Candidate LinkedIn: {linkedin}
- Candidate Email: {email}
- Candidate Phone: {mobile}

### INSTRUCTIONS:
- Return PLAIN TEXT ONLY — no HTML, no markdown, no asterisks
- Do NOT mention company name more than once
- Make it feel tailored — not templated
- End with: "Sincerely," followed by the candidate's full name
"""

        # ✅ Call LLM
        with st.spinner("✉️ Generating cover letter..."):
            try:
                cover_letter = call_llm(prompt, session=st.session_state).strip()
            except Exception as e:
                st.error(f"❌ Failed to generate cover letter: {e}")
                return

        # ✅ Store plain text
        st.session_state["cover_letter"] = cover_letter

        # ✅ Build HTML wrapper for preview (safe)
        cover_letter_html = f"""
        <div style="font-family: Georgia, serif; font-size: 13pt; line-height: 1.6; 
                    color: #000; background: #fff; padding: 25px; 
                    border-radius: 8px; box-shadow: 0px 2px 6px rgba(0,0,0,0.1); 
                    max-width: 800px; margin: auto;">
            <div style="text-align:center; margin-bottom:15px;">
                <div style="font-size:18pt; font-weight:bold; color:#003366;">{name}</div>
                <div style="font-size:14pt; color:#555;">{job_title}</div>
                <div style="font-size:10pt; margin-top:5px;">
                    <a href="{linkedin}" style="color:#003366;">{linkedin}</a><br/>
                    📧 {email} | 📞 {mobile}
                </div>
            </div>
            <hr/>
            <pre style="white-space: pre-wrap; font-family: Georgia, serif; font-size: 12pt; color:#000;">
{cover_letter}
            </pre>
        </div>
        """

        st.session_state["cover_letter_html"] = cover_letter_html

        # ✅ Show nicely in Streamlit
        st.markdown(cover_letter_html, unsafe_allow_html=True)

# ------------------- Initialize -------------------
# ✅ Initialize database in persistent storage
create_user_table()
cleanup_expired_login_tokens()

# ── Magic link token check — runs on every page load ──────────────────────────
# When user clicks the login link, ?login_token=<uuid> is in the URL.
# We handle it here BEFORE rendering any UI so the session is set instantly.
if not st.session_state.get("authenticated"):
    _qp = st.query_params
    _token = _qp.get("login_token", "")
    if _token:
        _ok, _result = verify_login_token(_token)
        if _ok:
            log_user_action(st.session_state.username, "login")
            # Clear the token from the URL so a refresh doesn't re-trigger
            st.query_params.clear()
            st.rerun()
        else:
            # Show the error on the login page — don't block the app
            st.session_state["_token_error"] = _result
            st.query_params.clear()
# ──────────────────────────────────────────────────────────────────────────────

# ------------------- Tab-Specific Notification System -------------------
if "login_notification" not in st.session_state:
    st.session_state.login_notification = {"type": None, "text": None, "expires": 0.0}
if "register_notification" not in st.session_state:
    st.session_state.register_notification = {"type": None, "text": None, "expires": 0.0}

def notify(tab, msg_type, text, duration=3.0):
    """Show auto-disappearing message for specific tab (login/register)."""
    notification_key = f"{tab}_notification"
    st.session_state[notification_key] = {
        "type": msg_type,
        "text": text,
        "expires": time.time() + duration,
    }

def render_notification(tab):
    """Render notification in a fixed-height slot — button position never shifts."""
    notification_key = f"{tab}_notification"
    notif = st.session_state[notification_key]

    # Map type to inline style colours (avoids Streamlit's full-height alert boxes)
    _styles = {
        "success": ("rgba(52,211,153,0.13)", "rgba(52,211,153,0.28)", "#6ee7b7"),
        "error":   ("rgba(251,113,133,0.13)", "rgba(251,113,133,0.28)", "#fca5a5"),
        "warning": ("rgba(251,191,36,0.13)",  "rgba(251,191,36,0.28)",  "#fde68a"),
        "info":    ("rgba(56,189,248,0.13)",  "rgba(56,189,248,0.28)",  "#7dd3fc"),
    }

    # Always emit a min-height wrapper so nothing below shifts on empty state
    if notif["type"] and time.time() < notif["expires"]:
        bg, border, color = _styles.get(notif["type"], _styles["info"])
        st.markdown(
            f"""<div style='min-height:48px; display:flex; align-items:center;'>
                <div style='width:100%; padding:8px 14px; border-radius:8px;
                            background:{bg}; border:1px solid {border};
                            color:{color}; font-size:0.85rem; font-weight:500;
                            font-family:-apple-system,sans-serif; line-height:1.4;
                            white-space:normal; word-wrap:break-word; overflow:visible;'>
                    {notif["text"]}
                </div>
            </div>""",
            unsafe_allow_html=True
        )
    else:
        # Reserved space — invisible, same height
        st.markdown("<div style='height:48px;'></div>", unsafe_allow_html=True)


def display_timer(remaining_seconds, expired=False, key_suffix=""):
    """
    Display a server-synced timer with glassmorphism styling.
    Server-side validation ensures OTP expiry is accurately enforced.

    Args:
        remaining_seconds: Time remaining in seconds (server-calculated)
        expired: Whether the timer has expired
        key_suffix: Unique suffix for the timer component
    """
    minutes = remaining_seconds // 60
    seconds = remaining_seconds % 60

    if expired or remaining_seconds <= 0:
        st.markdown(
            "<div class='timer-display timer-expired' style=\"background:linear-gradient(135deg,rgba(255,99,71,0.18) 0%,rgba(255,99,71,0.08) 100%);backdrop-filter:blur(15px);-webkit-backdrop-filter:blur(15px);border:2px solid rgba(255,99,71,0.4);border-radius:14px;padding:16px 24px;margin:20px 0;text-align:center;box-shadow:0 4px 20px rgba(255,99,71,0.15),inset 0 1px 0 rgba(255,255,255,0.1);\">"
            "<span class='timer-text' style=\"color:#FF6347;font-size:1.15em;font-weight:bold;font-family:-apple-system,sans-serif;text-shadow:0 0 18px rgba(255,99,71,0.5);\">OTP Expired</span>"
            "</div>",
            unsafe_allow_html=True
        )
    else:
        # Client-side countdown for UX, but server validates on action
        st.components.v1.html(f"""
        <div class='timer-display' id='timer-{key_suffix}' style="
            background: linear-gradient(135deg, rgba(255, 215, 0, 0.18) 0%, rgba(255, 165, 0, 0.08) 100%);
            backdrop-filter: blur(15px);
            -webkit-backdrop-filter: blur(15px);
            border: 2px solid rgba(255, 215, 0, 0.4);
            border-radius: 14px;
            padding: 16px 24px;
            margin: 20px 0;
            text-align: center;
            box-shadow: 0 4px 20px rgba(255, 215, 0, 0.15), inset 0 1px 0 rgba(255, 255, 255, 0.1);
        ">
            <span class='timer-text' style="
                color: #FFD700;
                font-size: 1.15em;
                font-weight: bold;
                font-family: 'Orbitron', sans-serif;
                text-shadow: 0 0 18px rgba(255, 215, 0, 0.5);
            ">⏱️ Time Remaining: <span id='countdown-{key_suffix}'>{minutes:02d}:{seconds:02d}</span></span>
        </div>
        <script>
        (function() {{
            let remaining = {remaining_seconds};
            const countdownEl = document.getElementById('countdown-{key_suffix}');
            const timerEl = document.getElementById('timer-{key_suffix}');

            const interval = setInterval(() => {{
                remaining--;
                if (remaining <= 0) {{
                    clearInterval(interval);
                    if (timerEl) {{
                        timerEl.style.background = 'linear-gradient(135deg, rgba(255, 99, 71, 0.18) 0%, rgba(255, 99, 71, 0.08) 100%)';
                        timerEl.style.border = '2px solid rgba(255, 99, 71, 0.4)';
                        timerEl.innerHTML = "<span style='color: #FF6347; font-size: 1.15em; font-weight: bold; font-family: Orbitron, sans-serif; text-shadow: 0 0 18px rgba(255, 99, 71, 0.5);'>⏱️ OTP Expired</span>";
                    }}
                }} else {{
                    const mins = Math.floor(remaining / 60);
                    const secs = remaining % 60;
                    if (countdownEl) {{
                        countdownEl.textContent = `${{mins.toString().padStart(2, '0')}}:${{secs.toString().padStart(2, '0')}}`;
                    }}
                }}
            }}, 1000);
        }})();
        </script>
        """, height=80)


def mask_email(email: str) -> str:
    """Mask email for display: show first 2 chars of local part, mask the rest, keep domain.
    e.g. brandontiger231@gmail.com → br***@gmail.com
    """
    if not email or '@' not in email:
        return email
    local, domain = email.split('@', 1)
    if len(local) <= 2:
        masked_local = local[0] + '***'
    else:
        masked_local = local[:2] + '***'
    return f"{masked_local}@{domain}"


# ------------------- Initialize Session State -------------------
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "username" not in st.session_state:
    st.session_state.username = None
if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

# Forgot password session states
if "reset_stage" not in st.session_state:
    st.session_state.reset_stage = "none"
if "reset_email" not in st.session_state:
    st.session_state.reset_email = ""
if "reset_otp" not in st.session_state:
    st.session_state.reset_otp = ""
if "reset_otp_time" not in st.session_state:
    st.session_state.reset_otp_time = 0

# Validation message state for register form (populated by on_change callbacks)
# _email_msg, _user_msg, _pass_msg are initialised inside the register form block

# ------------------- CSS Styling -------------------
st.markdown("""
<style>
/* ═══════════════════════════════════════════════════════════════
   HIRELYZER — Premium Apple-Style Dark Theme
   Font Stack: SF Pro Display → Segoe UI → Roboto → sans-serif
   Design Language: Glassmorphism · Soft gradients · Refined motion
   ═══════════════════════════════════════════════════════════════ */

@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=DM+Sans:wght@300;400;500;600;700&family=Orbitron:wght@600&display=swap');

:root {
    --bg-primary:       #080c12;
    --bg-secondary:     #0e1420;
    --bg-tertiary:      #141c2b;
    --surface-01:       rgba(255,255,255,0.04);
    --surface-02:       rgba(255,255,255,0.07);
    --surface-hover:    rgba(255,255,255,0.10);
    --border-subtle:    rgba(255,255,255,0.07);
    --border-accent:    rgba(99,179,237,0.30);
    --accent-blue:      #4fa3e3;
    --accent-cyan:      #38bdf8;
    --accent-violet:    #818cf8;
    --accent-emerald:   #34d399;
    --accent-amber:     #fbbf24;
    --accent-rose:      #fb7185;
    --text-primary:     #f0f4f8;
    --text-secondary:   #94a3b8;
    --text-muted:       #4a5568;
    --radius-sm:        8px;
    --radius-md:        14px;
    --radius-lg:        20px;
    --radius-xl:        28px;
    --shadow-glow-blue: 0 0 30px rgba(79,163,227,0.15);
    --shadow-card:      0 8px 40px rgba(0,0,0,0.45), 0 1px 0 rgba(255,255,255,0.06) inset;
    --font-sans:        -apple-system, BlinkMacSystemFont, "SF Pro Display", "DM Sans", "Segoe UI", Roboto, sans-serif;
    --transition-fast:  0.18s cubic-bezier(0.4,0,0.2,1);
    --transition-base:  0.28s cubic-bezier(0.4,0,0.2,1);
    --transition-slow:  0.45s cubic-bezier(0.4,0,0.2,1);
}

/* ── Base Reset ── */
html, body, [class*="css"], .stApp {
    font-family: var(--font-sans) !important;
    background-color: var(--bg-primary) !important;
    color: var(--text-primary) !important;
    -webkit-font-smoothing: antialiased;
    -moz-osx-font-smoothing: grayscale;
    scroll-behavior: smooth;
}

/* ── Force unified background — eliminates the horizontal seam glitch ── */
.stApp > header,
.stApp [data-testid="stAppViewContainer"],
.stApp [data-testid="stAppViewBlockContainer"],
.stApp [data-testid="block-container"],
.main,
.main > div,
section[data-testid="stMain"],
section[data-testid="stMain"] > div {
    background-color: var(--bg-primary) !important;
    background: var(--bg-primary) !important;
}

/* ── Streamlit top toolbar / header bar — the actual seam source ── */
header[data-testid="stHeader"],
header[data-testid="stHeader"] > div,
header[data-testid="stHeader"] > div > div,
.stApp header,
div[data-testid="stToolbar"],
div[data-testid="stStatusWidget"] {
    background-color: var(--bg-primary) !important;
    background: var(--bg-primary) !important;
    border-bottom: none !important;
    box-shadow: none !important;
}

/* ── Remove the decorative top colour bar Streamlit injects ── */
div[data-testid="stDecoration"],
#stDecoration {
    background: var(--bg-primary) !important;
    background-image: none !important;
    display: none !important;
}

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: var(--bg-secondary); }
::-webkit-scrollbar-thumb { background: rgba(79,163,227,0.35); border-radius: 99px; }
::-webkit-scrollbar-thumb:hover { background: rgba(79,163,227,0.6); }

/* ── Main container ── */
.main .block-container {
    padding: 1.5rem 2rem 3rem !important;
    max-width: 1280px;
}

/* ══════════════════════════════════════
   FADE ANIMATIONS
   ══════════════════════════════════════ */
@keyframes fadein  { from { opacity: 0; transform: translateY(6px); } to { opacity: 1; transform: translateY(0); } }
@keyframes fadeout { from { opacity: 1; } to { opacity: 0; } }
@keyframes fadeSlideUp { from { opacity: 0; transform: translateY(16px); } to { opacity: 1; transform: translateY(0); } }
@keyframes pulseGlow {
    0%, 100% { box-shadow: var(--shadow-card); }
    50%       { box-shadow: var(--shadow-card), var(--shadow-glow-blue); }
}
@keyframes shimmerSlide {
    0%   { transform: translateX(-100%) skewX(-12deg); }
    100% { transform: translateX(220%) skewX(-12deg); }
}
@keyframes glassShimmer {
    0%   { transform: translateX(-100%) skewX(-15deg); }
    100% { transform: translateX(200%) skewX(-15deg); }
}
@keyframes slideIn {
    0%   { transform: translateX(-50px); opacity: 0; }
    100% { transform: translateX(0); opacity: 1; }
}
@keyframes floatUp {
    0%, 100% { transform: translateY(0); }
    50%       { transform: translateY(-6px); }
}

/* ── Animated cards removed — replaced by premium hero section ── */

/* ══════════════════════════════════════
   STREAMLIT ALERT TOASTS
   ══════════════════════════════════════ */
div.stAlert {
    border-radius: var(--radius-md) !important;
    padding: 12px 18px !important;
    animation: fadein 0.3s ease, fadeout 0.3s 2.7s ease;
    backdrop-filter: blur(20px);
    border: 1px solid var(--border-subtle);
    font-size: 0.875rem;
    font-family: var(--font-sans) !important;
}

/* ══════════════════════════════════════
   GLASSMORPHISM CARD — base class
   ══════════════════════════════════════ */
.glass-card {
    background: var(--surface-01);
    backdrop-filter: blur(24px) saturate(180%);
    -webkit-backdrop-filter: blur(24px) saturate(180%);
    border: 1px solid var(--border-subtle);
    border-radius: var(--radius-lg);
    box-shadow: var(--shadow-card);
    transition: transform var(--transition-base), box-shadow var(--transition-base);
    position: relative;
    overflow: hidden;
}
.glass-card::before {
    content: '';
    position: absolute;
    inset: 0;
    background: linear-gradient(135deg, rgba(255,255,255,0.05) 0%, transparent 60%);
    pointer-events: none;
    border-radius: inherit;
}
.glass-card:hover {
    transform: translateY(-4px);
    box-shadow: var(--shadow-card), 0 0 50px rgba(79,163,227,0.10);
    border-color: var(--border-accent);
}

/* ══════════════════════════════════════
   LOGIN / AUTH CARD
   ══════════════════════════════════════ */
.login-card {
    background: linear-gradient(160deg,
        rgba(14,20,32,0.95) 0%,
        rgba(8,12,18,0.98) 100%);
    backdrop-filter: blur(32px) saturate(160%);
    -webkit-backdrop-filter: blur(32px) saturate(160%);
    border: 1px solid rgba(99,179,237,0.18);
    border-radius: var(--radius-xl);
    padding: 28px 32px 36px;
    box-shadow: var(--shadow-card), 0 0 60px rgba(79,163,227,0.07);
    transition: all var(--transition-slow);
    position: relative;
    overflow: hidden;
    animation: fadeSlideUp 0.7s cubic-bezier(0.22,1,0.36,1) forwards;
}
.login-card::after {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 60%; height: 100%;
    background: linear-gradient(90deg, transparent, rgba(79,163,227,0.06), transparent);
    animation: shimmerSlide 3.5s ease-in-out infinite;
}
.login-card:hover {
    border-color: rgba(99,179,237,0.32);
    box-shadow: var(--shadow-card), 0 0 80px rgba(79,163,227,0.12);
}

/* ══════════════════════════════════════
   TEXT INPUTS & TEXTAREAS
   ══════════════════════════════════════ */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stSelectbox > div > div > div {
    background: rgba(255,255,255,0.04) !important;
    color: var(--text-primary) !important;
    border: 1px solid rgba(255,255,255,0.10) !important;
    border-radius: var(--radius-sm) !important;
    font-family: var(--font-sans) !important;
    font-size: 0.9rem !important;
    padding: 10px 14px !important;
    transition: border-color var(--transition-fast), box-shadow var(--transition-fast) !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: rgba(79,163,227,0.5) !important;
    box-shadow: 0 0 0 3px rgba(79,163,227,0.12) !important;
    outline: none !important;
}
.stTextInput > div > div > input:hover,
.stTextArea > div > div > textarea:hover {
    border-color: rgba(79,163,227,0.28) !important;
}
.stTextInput > label,
.stTextArea > label,
.stSelectbox > label,
.stSlider > label,
.stFileUploader > label {
    color: var(--text-secondary) !important;
    font-size: 0.8rem !important;
    font-weight: 500 !important;
    letter-spacing: 0.03em !important;
    text-transform: uppercase !important;
    font-family: var(--font-sans) !important;
}

/* ══════════════════════════════════════
   BUTTONS — Primary style
   ══════════════════════════════════════ */
.stButton > button {
    background: linear-gradient(135deg,
        rgba(56,189,248,0.18) 0%,
        rgba(79,163,227,0.12) 100%) !important;
    color: var(--accent-cyan) !important;
    border: 1px solid rgba(56,189,248,0.3) !important;
    border-radius: var(--radius-sm) !important;
    font-family: var(--font-sans) !important;
    font-weight: 600 !important;
    font-size: 0.875rem !important;
    letter-spacing: 0.02em !important;
    padding: 10px 22px !important;
    backdrop-filter: blur(12px) !important;
    box-shadow: 0 2px 12px rgba(56,189,248,0.08), inset 0 1px 0 rgba(255,255,255,0.08) !important;
    transition: all var(--transition-fast) !important;
}
.stButton > button:hover {
    background: linear-gradient(135deg,
        rgba(56,189,248,0.28) 0%,
        rgba(79,163,227,0.22) 100%) !important;
    border-color: rgba(56,189,248,0.55) !important;
    box-shadow: 0 4px 20px rgba(56,189,248,0.18), inset 0 1px 0 rgba(255,255,255,0.12) !important;
    transform: translateY(-2px) !important;
    color: #e0f6ff !important;
}
.stButton > button:active {
    transform: translateY(0px) !important;
    box-shadow: 0 1px 6px rgba(56,189,248,0.10) !important;
}

/* ══════════════════════════════════════
   DOWNLOAD BUTTONS
   ══════════════════════════════════════ */
.stDownloadButton > button {
    background: linear-gradient(135deg,
        rgba(52,211,153,0.16) 0%,
        rgba(52,211,153,0.08) 100%) !important;
    color: var(--accent-emerald) !important;
    border: 1px solid rgba(52,211,153,0.28) !important;
    border-radius: var(--radius-sm) !important;
    font-family: var(--font-sans) !important;
    font-weight: 600 !important;
    font-size: 0.875rem !important;
    transition: all var(--transition-fast) !important;
}
.stDownloadButton > button:hover {
    background: linear-gradient(135deg,
        rgba(52,211,153,0.26) 0%,
        rgba(52,211,153,0.16) 100%) !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 4px 18px rgba(52,211,153,0.15) !important;
}

/* ══════════════════════════════════════
   METRICS
   ══════════════════════════════════════ */
div[data-testid="metric-container"] {
    background: var(--surface-01) !important;
    border: 1px solid var(--border-subtle) !important;
    border-radius: var(--radius-md) !important;
    padding: 18px 20px !important;
    backdrop-filter: blur(16px) !important;
    transition: all var(--transition-base) !important;
    animation: fadeSlideUp 0.5s ease forwards;
}
div[data-testid="metric-container"]:hover {
    border-color: var(--border-accent) !important;
    background: var(--surface-02) !important;
    transform: translateY(-3px) !important;
    box-shadow: var(--shadow-glow-blue) !important;
}
div[data-testid="metric-container"] label {
    color: var(--text-secondary) !important;
    font-size: 0.75rem !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.06em !important;
    font-family: var(--font-sans) !important;
}
div[data-testid="metric-container"] [data-testid="stMetricValue"] {
    color: var(--text-primary) !important;
    font-size: 1.75rem !important;
    font-weight: 700 !important;
    letter-spacing: -0.02em !important;
    line-height: 1.2 !important;
    font-family: var(--font-sans) !important;
}

/* ══════════════════════════════════════
   TABS
   ══════════════════════════════════════ */
.stTabs [data-baseweb="tab-list"] {
    gap: 4px !important;
    background: rgba(255,255,255,0.03) !important;
    padding: 5px !important;
    border-radius: var(--radius-md) !important;
    border: 1px solid var(--border-subtle) !important;
    backdrop-filter: blur(16px) !important;
}
.stTabs [data-baseweb="tab"] {
    border-radius: var(--radius-sm) !important;
    color: var(--text-secondary) !important;
    font-family: var(--font-sans) !important;
    font-weight: 500 !important;
    font-size: 0.875rem !important;
    padding: 9px 18px !important;
    transition: all var(--transition-fast) !important;
    border: none !important;
    background: transparent !important;
}
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg,
        rgba(56,189,248,0.18) 0%,
        rgba(79,163,227,0.12) 100%) !important;
    color: var(--accent-cyan) !important;
    font-weight: 600 !important;
    border: 1px solid rgba(56,189,248,0.25) !important;
    box-shadow: 0 2px 10px rgba(56,189,248,0.12) !important;
}
.stTabs [data-baseweb="tab-highlight"] { display: none !important; }
.stTabs [data-baseweb="tab-panel"] {
    background: transparent !important;
    padding: 20px 0 !important;
}

/* ══════════════════════════════════════
   DATAFRAME / TABLE
   ══════════════════════════════════════ */
.dataframe, .stDataFrame {
    border-radius: var(--radius-md) !important;
    overflow: hidden !important;
    border: 1px solid var(--border-subtle) !important;
}
.stDataFrame [data-testid="stDataFrameResizable"] {
    background: var(--bg-secondary) !important;
}

/* ══════════════════════════════════════
   SIDEBAR
   ══════════════════════════════════════ */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg,
        rgba(10,14,22,0.98) 0%,
        rgba(8,12,18,1) 100%) !important;
    border-right: 1px solid var(--border-subtle) !important;
}
section[data-testid="stSidebar"] .block-container {
    padding: 1.5rem 1rem !important;
}
section[data-testid="stSidebar"] label {
    color: var(--text-secondary) !important;
    font-size: 0.8rem !important;
    text-transform: uppercase !important;
    letter-spacing: 0.05em !important;
    font-weight: 600 !important;
}
section[data-testid="stSidebar"] .stTextInput > div > div > input,
section[data-testid="stSidebar"] .stTextArea > div > div > textarea {
    background: rgba(255,255,255,0.03) !important;
    border-color: rgba(255,255,255,0.08) !important;
}

/* ══════════════════════════════════════
   EXPANDER
   ══════════════════════════════════════ */
.streamlit-expanderHeader {
    background: var(--surface-01) !important;
    border: 1px solid var(--border-subtle) !important;
    border-radius: var(--radius-md) !important;
    color: var(--text-primary) !important;
    font-family: var(--font-sans) !important;
    font-weight: 500 !important;
    transition: all var(--transition-fast) !important;
}
.streamlit-expanderHeader:hover {
    background: var(--surface-hover) !important;
    border-color: var(--border-accent) !important;
}
.streamlit-expanderContent {
    background: rgba(255,255,255,0.02) !important;
    border: 1px solid var(--border-subtle) !important;
    border-top: none !important;
    border-radius: 0 0 var(--radius-md) var(--radius-md) !important;
}

/* ══════════════════════════════════════
   SLIDERS
   ══════════════════════════════════════ */
.stSlider [data-baseweb="slider"] [role="slider"] {
    background: var(--accent-cyan) !important;
    border: 2px solid var(--bg-primary) !important;
    box-shadow: 0 0 0 3px rgba(56,189,248,0.3) !important;
}
.stSlider [data-baseweb="slider"] [data-testid="stTickBar"] > div {
    background: rgba(56,189,248,0.6) !important;
}

/* ══════════════════════════════════════
   FILE UPLOADER
   ══════════════════════════════════════ */
.stFileUploader > div {
    background: var(--surface-01) !important;
    border: 1.5px dashed rgba(79,163,227,0.3) !important;
    border-radius: var(--radius-lg) !important;
    transition: all var(--transition-base) !important;
}
.stFileUploader > div:hover {
    border-color: rgba(79,163,227,0.6) !important;
    background: rgba(79,163,227,0.04) !important;
    box-shadow: 0 0 40px rgba(79,163,227,0.07) !important;
}

/* ══════════════════════════════════════
   DIVIDER
   ══════════════════════════════════════ */
hr {
    border: none !important;
    border-top: 1px solid var(--border-subtle) !important;
    margin: 28px 0 !important;
}

/* ══════════════════════════════════════
   HEADINGS
   ══════════════════════════════════════ */
h1, h2, h3, h4, h5, h6,
.stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {
    font-family: var(--font-sans) !important;
    color: var(--text-primary) !important;
    letter-spacing: -0.02em !important;
    font-weight: 700 !important;
}
h1, .stMarkdown h1 { font-size: 2rem !important; }
h2, .stMarkdown h2 { font-size: 1.4rem !important; }
h3, .stMarkdown h3 {
    font-size: 1.1rem !important;
    color: var(--text-secondary) !important;
    font-weight: 600 !important;
}

/* ══════════════════════════════════════
   FEATURE CARDS (sidebar pre-login)
   ══════════════════════════════════════ */
.feature-card {
    background: var(--surface-01);
    backdrop-filter: blur(20px);
    -webkit-backdrop-filter: blur(20px);
    border: 1px solid var(--border-subtle);
    border-radius: var(--radius-md);
    padding: 18px 16px;
    margin-bottom: 12px;
    transition: all var(--transition-base);
    position: relative;
    overflow: hidden;
    animation: fadeSlideUp 0.5s ease forwards;
}
.feature-card::before {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 50%; height: 100%;
    background: linear-gradient(90deg, transparent, rgba(79,163,227,0.05), transparent);
    transition: left 0.6s ease;
}
.feature-card:hover { transform: translateY(-4px); border-color: var(--border-accent); }
.feature-card:hover::before { left: 150%; }
.feature-card h3 {
    color: var(--accent-cyan) !important;
    font-size: 0.9rem !important;
    font-weight: 600 !important;
    margin: 8px 0 4px !important;
}
.feature-card p {
    color: var(--text-secondary) !important;
    font-size: 0.78rem !important;
    line-height: 1.5 !important;
    margin: 0 !important;
}

/* ══════════════════════════════════════
   SLIDE MESSAGES (inline notifications)
   ══════════════════════════════════════ */
.slide-message {
    position: relative;
    overflow: hidden;
    margin: 12px 0;
    padding: 12px 18px;
    border-radius: var(--radius-md);
    font-weight: 500;
    font-size: 0.875rem;
    display: flex;
    align-items: center;
    justify-content: flex-start;
    gap: 10px;
    animation: fadein 0.4s cubic-bezier(0.34,1.56,0.64,1) forwards;
    backdrop-filter: blur(16px);
    -webkit-backdrop-filter: blur(16px);
    box-shadow: 0 4px 16px rgba(0,0,0,0.15), inset 0 1px 0 rgba(255,255,255,0.08);
    width: 100%;
    box-sizing: border-box;
    font-family: var(--font-sans) !important;
    transition: all var(--transition-fast);
    min-height: 46px;
}
.slide-message:hover { transform: translateY(-2px); }
.slide-message svg { display: inline-block !important; flex-shrink: 0; vertical-align: middle; min-width: 16px; min-height: 16px; }
.slide-message-text { flex: 1; position: relative; z-index: 2; word-wrap: break-word; }
.success-msg {
    background: linear-gradient(135deg, rgba(52,211,153,0.15) 0%, rgba(52,211,153,0.05) 100%);
    border: 1px solid rgba(52,211,153,0.30);
    color: #6ee7b7;
}
.error-msg {
    background: linear-gradient(135deg, rgba(251,113,133,0.15) 0%, rgba(251,113,133,0.05) 100%);
    border: 1px solid rgba(251,113,133,0.30);
    color: #fca5a5;
}
.info-msg {
    background: linear-gradient(135deg, rgba(56,189,248,0.15) 0%, rgba(56,189,248,0.05) 100%);
    border: 1px solid rgba(56,189,248,0.30);
    color: #7dd3fc;
}
.warn-msg {
    background: linear-gradient(135deg, rgba(251,191,36,0.15) 0%, rgba(251,191,36,0.05) 100%);
    border: 1px solid rgba(251,191,36,0.30);
    color: #fde68a;
}

/* ══════════════════════════════════════
   COUNTER GRID (landing page stats)
   ══════════════════════════════════════ */
.counter-grid {
    display: grid;
    grid-template-columns: repeat(2, 1fr);
    gap: 16px;
    padding: 24px 0;
    max-width: 520px;
    margin: 0 auto;
}
.counter-box {
    background: var(--surface-01);
    backdrop-filter: blur(20px);
    -webkit-backdrop-filter: blur(20px);
    border: 1px solid var(--border-subtle);
    border-radius: var(--radius-md);
    padding: 22px 18px;
    display: flex;
    flex-direction: column;
    justify-content: center;
    align-items: center;
    position: relative;
    overflow: hidden;
    transition: all var(--transition-base);
    animation: floatUp 4s ease-in-out infinite;
}
.counter-box::before {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 50%; height: 100%;
    background: linear-gradient(90deg, transparent, rgba(56,189,248,0.06), transparent);
    animation: shimmerSlide 3s infinite;
}
.counter-box:hover {
    transform: translateY(-6px) scale(1.02);
    border-color: var(--border-accent);
    box-shadow: 0 12px 40px rgba(56,189,248,0.10);
}
.counter-box:nth-child(1) { animation-delay: 0s; }
.counter-box:nth-child(2) { animation-delay: 0.6s; }
.counter-box:nth-child(3) { animation-delay: 1.2s; }
.counter-box:nth-child(4) { animation-delay: 1.8s; }
.counter-number {
    font-size: 2rem;
    font-weight: 700;
    color: var(--accent-cyan);
    letter-spacing: -0.03em;
    line-height: 1;
    position: relative;
    z-index: 2;
    font-family: var(--font-sans);
}
.counter-label {
    margin-top: 6px;
    font-size: 0.78rem;
    color: var(--text-secondary);
    font-weight: 500;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    position: relative;
    z-index: 2;
    font-family: var(--font-sans);
}

/* ══════════════════════════════════════
   ATS SECTION CARDS (analysis results)
   ══════════════════════════════════════ */
.ats-section-header {
    background: linear-gradient(135deg,
        rgba(56,189,248,0.14) 0%,
        rgba(79,163,227,0.08) 100%);
    border: 1px solid rgba(56,189,248,0.22);
    border-radius: var(--radius-sm) var(--radius-sm) 0 0;
    padding: 12px 18px;
    font-family: var(--font-sans) !important;
    font-weight: 700;
    font-size: 0.875rem;
    color: var(--accent-cyan);
    letter-spacing: 0.02em;
    text-transform: uppercase;
}
.ats-section-body {
    background: rgba(255,255,255,0.025);
    border: 1px solid rgba(255,255,255,0.06);
    border-top: none;
    border-radius: 0 0 var(--radius-sm) var(--radius-sm);
    padding: 16px 18px;
    color: var(--text-secondary);
    font-family: var(--font-sans) !important;
    font-size: 0.875rem;
    line-height: 1.65;
    margin-bottom: 14px;
}
.score-badge {
    display: inline-flex;
    align-items: center;
    background: linear-gradient(135deg, rgba(56,189,248,0.18) 0%, rgba(56,189,248,0.08) 100%);
    border: 1px solid rgba(56,189,248,0.30);
    border-radius: 99px;
    padding: 4px 14px;
    font-size: 0.78rem;
    font-weight: 700;
    color: var(--accent-cyan);
    letter-spacing: 0.04em;
    font-family: var(--font-sans);
    margin-bottom: 10px;
}

/* ══════════════════════════════════════
   WELCOME BANNER (post-login)
   ══════════════════════════════════════ */
.welcome-banner {
    background: linear-gradient(135deg,
        rgba(14,20,32,0.9) 0%,
        rgba(10,16,26,0.95) 100%);
    border: 1px solid rgba(56,189,248,0.15);
    border-radius: var(--radius-lg);
    padding: 24px 32px;
    margin-bottom: 28px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    animation: fadeSlideUp 0.6s ease forwards;
    position: relative;
    overflow: hidden;
}
.welcome-banner::before {
    content: '';
    position: absolute;
    top: -50%; left: -20%;
    width: 60%; height: 200%;
    background: radial-gradient(ellipse, rgba(56,189,248,0.05) 0%, transparent 70%);
    pointer-events: none;
}
.welcome-title {
    font-family: var(--font-sans) !important;
    font-size: 1.45rem !important;
    font-weight: 700 !important;
    letter-spacing: -0.025em !important;
    color: var(--text-primary) !important;
    line-height: 1.3 !important;
}
.welcome-subtitle {
    font-size: 0.85rem;
    color: var(--text-secondary);
    margin-top: 4px;
    font-family: var(--font-sans);
}
.welcome-username {
    color: var(--accent-cyan);
    font-weight: 700;
}

/* ══════════════════════════════════════
   ADMIN DASHBOARD
   ══════════════════════════════════════ */
.admin-header {
    background: linear-gradient(135deg,
        rgba(129,140,248,0.12) 0%,
        rgba(99,102,241,0.06) 100%);
    border: 1px solid rgba(129,140,248,0.20);
    border-radius: var(--radius-md);
    padding: 16px 24px;
    margin-bottom: 24px;
}
.admin-header h2 {
    color: var(--accent-violet) !important;
    margin: 0 !important;
    font-size: 1.2rem !important;
}

/* ══════════════════════════════════════
   SECTION DIVIDER with label
   ══════════════════════════════════════ */
.section-label {
    font-family: var(--font-sans) !important;
    font-size: 0.72rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.12em !important;
    text-transform: uppercase !important;
    color: var(--text-muted) !important;
    margin-bottom: 12px !important;
    padding-bottom: 8px !important;
    border-bottom: 1px solid var(--border-subtle) !important;
}

/* ══════════════════════════════════════
   SPINNER
   ══════════════════════════════════════ */
.stSpinner > div {
    border-top-color: var(--accent-cyan) !important;
}

/* ══════════════════════════════════════
   TIMER DISPLAY
   ══════════════════════════════════════ */
.timer-display {
    background: linear-gradient(135deg,
        rgba(251,191,36,0.12) 0%,
        rgba(251,191,36,0.05) 100%);
    backdrop-filter: blur(16px);
    -webkit-backdrop-filter: blur(16px);
    border: 1px solid rgba(251,191,36,0.28);
    border-radius: var(--radius-md);
    padding: 16px 24px;
    margin: 18px 0;
    text-align: center;
    box-shadow: 0 4px 20px rgba(251,191,36,0.08), inset 0 1px 0 rgba(255,255,255,0.06);
    transition: all var(--transition-base);
    position: relative;
    overflow: hidden;
}
.timer-display::before {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 50%; height: 100%;
    background: linear-gradient(90deg, transparent, rgba(251,191,36,0.08), transparent);
    animation: shimmerSlide 3s infinite;
}
.timer-display:hover { transform: translateY(-2px); box-shadow: 0 8px 28px rgba(251,191,36,0.12); }
.timer-text {
    color: var(--accent-amber);
    font-size: 1rem;
    font-weight: 700;
    font-family: var(--font-sans);
    letter-spacing: 0.04em;
    position: relative;
    z-index: 2;
}
.timer-expired { 
    background: linear-gradient(135deg, rgba(251,113,133,0.12) 0%, rgba(251,113,133,0.05) 100%);
    border-color: rgba(251,113,133,0.28);
}
.timer-expired .timer-text { color: var(--accent-rose); }

/* ══════════════════════════════════════
   BANNER / MARQUEE (tab1 dashboard)
   ══════════════════════════════════════ */
.banner-container {
    width: 100%;
    height: 60px;
    background: linear-gradient(90deg,
        rgba(8,12,18,1) 0%,
        rgba(14,20,32,0.9) 50%,
        rgba(8,12,18,1) 100%);
    border: 1px solid var(--border-subtle);
    overflow: hidden;
    display: flex;
    align-items: center;
    position: relative;
    margin-bottom: 24px;
    border-radius: var(--radius-md);
    backdrop-filter: blur(20px);
}
.pulse-bar {
    position: absolute;
    display: flex;
    align-items: center;
    font-size: 0.9rem;
    font-weight: 600;
    font-family: var(--font-sans);
    color: var(--accent-cyan);
    white-space: nowrap;
    letter-spacing: 0.04em;
    animation: glideIn 14s linear infinite;
}
.pulse-bar .bar {
    width: 3px;
    height: 18px;
    margin-right: 12px;
    background: var(--accent-cyan);
    border-radius: 2px;
    box-shadow: 0 0 8px var(--accent-cyan);
    animation: pulse 1s ease-in-out infinite;
}
@keyframes glideIn {
    0%   { left: -40%; opacity: 0; }
    8%   { opacity: 1; }
    92%  { opacity: 1; }
    100% { left: 110%; opacity: 0; }
}
@keyframes pulse {
    0%, 100% { height: 14px; background: var(--accent-cyan); }
    50%       { height: 22px; background: var(--accent-violet); }
}

/* ══════════════════════════════════════
   HEADER BOX (dashboard title area)
   ══════════════════════════════════════ */
.header {
    font-size: 1.5rem;
    font-weight: 700;
    text-align: center;
    letter-spacing: -0.02em;
    padding: 20px 28px;
    color: var(--text-primary);
    position: relative;
    overflow: hidden;
    border-radius: var(--radius-md);
    background: linear-gradient(135deg,
        rgba(14,20,32,0.8) 0%,
        rgba(10,16,26,0.9) 100%);
    border: 1px solid rgba(56,189,248,0.18);
    box-shadow: var(--shadow-card);
    font-family: var(--font-sans);
}
.header span { color: var(--accent-cyan); }
.header::before {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 50%; height: 100%;
    background: linear-gradient(90deg, transparent, rgba(56,189,248,0.05), transparent);
    transition: left 0.7s ease;
}
.header:hover::before { left: 150%; }

/* ══════════════════════════════════════
   ANALYSIS RESULT CARDS
   ══════════════════════════════════════ */
.result-card {
    background: var(--surface-01);
    border: 1px solid var(--border-subtle);
    border-radius: var(--radius-md);
    padding: 20px;
    margin-bottom: 14px;
    transition: all var(--transition-base);
    animation: fadeSlideUp 0.5s ease forwards;
    position: relative;
    overflow: hidden;
}
.result-card::before {
    content: '';
    position: absolute;
    left: 0; top: 0; bottom: 0;
    width: 3px;
    background: linear-gradient(180deg, var(--accent-cyan) 0%, var(--accent-violet) 100%);
    border-radius: 0 0 0 var(--radius-md);
}
.result-card:hover {
    border-color: var(--border-accent);
    background: var(--surface-02);
    transform: translateX(3px);
}

/* ══════════════════════════════════════
   DARK STREAMLIT OVERRIDES
   ══════════════════════════════════════ */
.stMarkdown p, .stText {
    color: var(--text-secondary) !important;
    font-family: var(--font-sans) !important;
    font-size: 0.9rem !important;
    line-height: 1.65 !important;
}
.stMarkdown strong { color: var(--text-primary) !important; }
.stMarkdown code {
    background: rgba(255,255,255,0.06) !important;
    color: var(--accent-cyan) !important;
    border-radius: 4px !important;
    padding: 2px 6px !important;
    font-size: 0.83rem !important;
}
.stInfo, .stSuccess, .stWarning, .stError {
    border-radius: var(--radius-md) !important;
    font-family: var(--font-sans) !important;
    font-size: 0.875rem !important;
}

/* Caption */
.stCaption {
    color: var(--text-muted) !important;
    font-size: 0.78rem !important;
    font-family: var(--font-sans) !important;
}

/* ══════════════════════════════════════
   FILE UPLOADER INNER LABEL
   ══════════════════════════════════════ */
.stFileUploader [data-testid="stFileUploaderDropzone"] {
    background: rgba(56,189,248,0.02) !important;
}
/* Only style the direct instruction text — not internal layout spans */
.stFileUploader [data-testid="stFileUploaderDropzone"] > span {
    color: var(--text-secondary) !important;
    font-family: var(--font-sans) !important;
}
/* ══════════════════════════════════════
   FILE UPLOADER — CSS ISOLATION RESET
   Prevents shimmer pseudo-elements,
   position/overflow/animation overrides
   from leaking into the uploader widget.
   ══════════════════════════════════════ */

/* 1. Hard-reset every descendant's layout model */
[data-testid="stFileUploader"] *,
[data-testid="stFileUploaderDropzone"] * {
    position: static !important;
    overflow: visible !important;
    animation: none !important;
    transition: none !important;
    transform: none !important;
    backdrop-filter: none !important;
    -webkit-backdrop-filter: none !important;
}

/* 2. Kill ALL pseudo-elements inside the uploader */
[data-testid="stFileUploader"] *::before,
[data-testid="stFileUploader"] *::after,
[data-testid="stFileUploaderDropzone"] *::before,
[data-testid="stFileUploaderDropzone"] *::after {
    content: none !important;
    display: none !important;
    position: static !important;
    animation: none !important;
}

/* 3. Restore the dropzone container itself
      (needs overflow:hidden for border clipping, relative for internal layout) */
[data-testid="stFileUploaderDropzone"] {
    overflow: hidden !important;
    position: relative !important;
    background: rgba(56,189,248,0.02) !important;
}

/* 4. Restore flex row on the dropzone's direct children
      (Streamlit renders: icon | instruction text | Browse button) */
[data-testid="stFileUploaderDropzone"] > div {
    display: flex !important;
    flex-direction: row !important;
    align-items: center !important;
    position: relative !important;
    overflow: visible !important;
}

/* 5. Let the uploaded file chip and its delete button render normally */
[data-testid="stFileUploaderDeleteBtn"] button,
[data-testid="stFileUploaderFile"],
[data-testid="stFileUploaderFile"] * {
    position: relative !important;
    overflow: visible !important;
    animation: none !important;
    transform: none !important;
}
</style>
""", unsafe_allow_html=True)
# 🔹 VIDEO BACKGROUND & GLOW TEXT

# ------------------- BEFORE LOGIN -------------------
if not st.session_state.authenticated:
    

    # -------- Sidebar --------
    with st.sidebar:
        st.markdown(
            "<div style=\"padding:16px 4px 20px;border-bottom:1px solid rgba(255,255,255,0.07);margin-bottom:16px;\">"
            "<div style=\"font-family:-apple-system,BlinkMacSystemFont,sans-serif;font-size:1.1rem;font-weight:700;letter-spacing:-0.02em;color:#f0f4f8;line-height:1.2;\">HIRELYZER</div>"
            "<div style=\"font-size:0.72rem;font-weight:600;letter-spacing:0.1em;text-transform:uppercase;color:#38bdf8;margin-top:3px;\">AI Resume Intelligence</div>"
            "</div>"
            "<p style=\"color:#64748b;font-size:0.8rem;line-height:1.55;font-family:-apple-system,BlinkMacSystemFont,sans-serif;margin-bottom:16px;\">Transform your career with AI-powered resume analysis, job matching, and smart insights.</p>",
            unsafe_allow_html=True
        )

        features = [
            ("https://img.icons8.com/fluency/48/resume.png", "Resume Analyzer", "Get feedback, scores, and tips powered by AI along with the biased words detection and rewriting the resume in an inclusive way."),
            ("https://img.icons8.com/fluency/48/resume-website.png", "Resume Builder", "Build modern, eye-catching resumes easily."),
            ("https://img.icons8.com/fluency/48/job.png", "Job Search", "Find tailored job matches."),
            ("https://img.icons8.com/fluency/48/classroom.png", "Course Suggestions", "Get upskilling recommendations based on your goals."),
            ("https://img.icons8.com/fluency/48/combo-chart.png", "Interactive Dashboard", "Visualize trends, scores, and analytics."),
        ]

        for icon, title, desc in features:
            st.markdown(f"""
            <div class="feature-card">
                <img src="{icon}" width="40"/>
                <h3>{title}</h3>
                <p>{desc}</p>
            </div>
            """, unsafe_allow_html=True)

    # -------- Premium Hero Section --------
    # Fetch live stats for subtle ribbon (cached — no Supabase hit on every rerun)
    total_users, active_logins, stats = _cached_hero_stats()
    resumes_uploaded = stats.get("total_candidates", 0)
    active_domains = stats.get("unique_domains", 0)

    # ── Hero HTML (no script — Streamlit strips <script> from st.markdown) ──
    st.markdown(f"""
    <style>
    .hero-section {{
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 36px 24px 14px;
        position: relative;
        overflow: hidden;
    }}
    .hero-section::before {{
        content: '';
        position: absolute;
        top: -60px; left: 50%;
        transform: translateX(-50%);
        width: 520px; height: 320px;
        background: radial-gradient(ellipse, rgba(56,189,248,0.08) 0%, transparent 70%);
        pointer-events: none;
        z-index: 0;
    }}
    .hero-brand {{
        position: relative;
        z-index: 2;
        text-align: center;
        margin-bottom: 18px;
        animation: fadeSlideUp 0.7s cubic-bezier(0.22,1,0.36,1) both;
    }}
    .hero-wordmark {{
        font-size: 3rem;
        font-weight: 800;
        letter-spacing: -0.04em;
        line-height: 1;
        color: #f0f4f8;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", "DM Sans", sans-serif;
    }}
    .hero-wordmark span {{
        background: linear-gradient(135deg, #38bdf8 0%, #818cf8 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }}
    .hero-tagline {{
        margin-top: 8px;
        font-size: 0.8rem;
        color: #334155;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
        letter-spacing: 0.04em;
        text-transform: uppercase;
        font-weight: 500;
    }}
    .hero-pills-container {{
        display: inline-flex;
        gap: 6px;
        flex-wrap: wrap;
        justify-content: center;
        align-items: center;
        background: rgba(255,255,255,0.03);
        border: 1px solid rgba(255,255,255,0.08);
        border-radius: 99px;
        padding: 6px 10px;
        margin-bottom: 32px;
        position: relative;
        z-index: 2;
        backdrop-filter: blur(16px);
        -webkit-backdrop-filter: blur(16px);
        animation: fadeSlideUp 1s cubic-bezier(0.22,1,0.36,1) 0.25s both;
    }}
    .hero-pill {{
        display: inline-flex;
        align-items: center;
        gap: 5px;
        padding: 4px 11px;
        border-radius: 99px;
        background: transparent;
        font-size: 0.75rem;
        font-weight: 500;
        color: #64748b;
        font-family: -apple-system, BlinkMacSystemFont, sans-serif;
        letter-spacing: 0.01em;
        transition: color 0.2s ease;
    }}
    .hero-pill:hover {{
        color: #38bdf8;
    }}
    .hero-pill-dot {{
        width: 5px; height: 5px;
        border-radius: 50%;
        background: #38bdf8;
        opacity: 0.55;
        flex-shrink: 0;
    }}
    .hero-pill-sep {{
        width: 1px; height: 12px;
        background: rgba(255,255,255,0.10);
        flex-shrink: 0;
        align-self: center;
    }}
    .hero-stat-ribbon {{
        display: inline-flex;
        gap: 0;
        justify-content: center;
        align-items: stretch;
        position: relative;
        z-index: 2;
        margin-bottom: 8px;
        background: rgba(255,255,255,0.025);
        border: 1px solid rgba(255,255,255,0.07);
        border-radius: 16px;
        padding: 0;
        overflow: hidden;
        animation: fadeSlideUp 1.1s cubic-bezier(0.22,1,0.36,1) 0.35s both;
    }}
    .hero-stat-item {{
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 16px 28px;
        gap: 2px;
        transition: background 0.2s ease;
    }}
    .hero-stat-item:hover {{
        background: rgba(56,189,248,0.05);
    }}
    .hero-stat-item:not(:last-child) {{
        border-right: 1px solid rgba(255,255,255,0.07);
    }}
    .hero-stat-num {{
        font-size: 1.4rem;
        font-weight: 700;
        color: #f0f4f8;
        letter-spacing: -0.03em;
        line-height: 1;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
    }}
    .hero-stat-icon {{
        display: flex;
        align-items: center;
        justify-content: center;
        width: 28px; height: 28px;
        border-radius: 8px;
        background: rgba(56,189,248,0.08);
        border: 1px solid rgba(56,189,248,0.15);
        margin-bottom: 4px;
        flex-shrink: 0;
    }}
    .hero-stat-lbl {{
        font-size: 0.65rem;
        color: #4a5568;
        text-transform: uppercase;
        letter-spacing: 0.09em;
        font-weight: 600;
        font-family: -apple-system, BlinkMacSystemFont, sans-serif;
    }}
    </style>

    <div class="hero-section">
        <div class="hero-brand">
            <div class="hero-wordmark">HIRE<span>LYZER</span></div>
            <div class="hero-tagline">AI-Powered Resume Intelligence Platform</div>
        </div>
        <div class="hero-pills-container">
            <div class="hero-pill"><span class="hero-pill-dot"></span>Bias Detection</div>
            <div class="hero-pill-sep"></div>
            <div class="hero-pill"><span class="hero-pill-dot"></span>ATS Scoring</div>
            <div class="hero-pill-sep"></div>
            <div class="hero-pill"><span class="hero-pill-dot"></span>Resume Builder</div>
            <div class="hero-pill-sep"></div>
            <div class="hero-pill"><span class="hero-pill-dot"></span>Job Matching</div>
            <div class="hero-pill-sep"></div>
            <div class="hero-pill"><span class="hero-pill-dot"></span>AI Coach</div>
        </div>
        <div class="hero-stat-ribbon">
            <div class="hero-stat-item">
                <div class="hero-stat-icon">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <circle cx="12" cy="7" r="4" stroke="#38bdf8" stroke-width="1.8" fill="rgba(56,189,248,0.12)"/>
                        <path d="M4 20c0-4 3.582-7 8-7s8 3 8 7" stroke="#38bdf8" stroke-width="1.8" stroke-linecap="round"/>
                    </svg>
                </div>
                <div class="hero-stat-num">{total_users}</div>
                <div class="hero-stat-lbl">Users</div>
            </div>
            <div class="hero-stat-item">
                <div class="hero-stat-icon">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <rect x="4" y="2" width="16" height="20" rx="2" stroke="#38bdf8" stroke-width="1.8" fill="rgba(56,189,248,0.08)"/>
                        <path d="M8 7h8M8 11h8M8 15h5" stroke="#38bdf8" stroke-width="1.5" stroke-linecap="round"/>
                    </svg>
                </div>
                <div class="hero-stat-num">{resumes_uploaded}</div>
                <div class="hero-stat-lbl">Resumes Analysed</div>
            </div>
            <div class="hero-stat-item">
                <div class="hero-stat-icon">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <circle cx="12" cy="12" r="9" stroke="#38bdf8" stroke-width="1.8" fill="rgba(56,189,248,0.08)"/>
                        <path d="M12 3v18M3 12h18" stroke="#38bdf8" stroke-width="1.5" stroke-linecap="round"/>
                    </svg>
                </div>
                <div class="hero-stat-num">{active_domains}</div>
                <div class="hero-stat-lbl">Domains</div>
            </div>
            <div class="hero-stat-item">
                <div class="hero-stat-icon">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <path d="M13 2L4.5 13.5H12L11 22L19.5 10.5H12L13 2Z" stroke="#38bdf8" stroke-width="1.8" stroke-linejoin="round" fill="rgba(56,189,248,0.12)"/>
                    </svg>
                </div>
                <div class="hero-stat-num">{active_logins}</div>
                <div class="hero-stat-lbl">Active Today</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Typewriter animation via components.html (scripts work here) ──
    st.components.v1.html("""
    <style>
    * { box-sizing: border-box; }
    body { margin: 0; padding: 0; background: transparent; overflow: hidden; }
    .tw-wrap {
        text-align: center;
        padding: 4px 0 8px;
        width: 100%;
    }
    .tw-text {
        font-size: 1rem;
        font-weight: 600;
        color: #38bdf8;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
        letter-spacing: -0.01em;
        border-right: 2px solid #38bdf8;
        white-space: nowrap;
        display: inline-block;
        max-width: 100%;
        animation: twBlink 1s ease-in-out infinite;
        background: transparent;
        vertical-align: middle;
    }
    @keyframes twBlink {
        0%, 100% { border-color: #38bdf8; }
        50%       { border-color: transparent; }
    }
    </style>
    <div class="tw-wrap"><span class="tw-text" id="tw">&nbsp;</span></div>
    <script>
    (function() {
        var phrases = [
            "Analyse resumes with zero bias.",
            "Score smarter. Hire better.",
            "AI that reads between the lines.",
            "Ethical hiring starts here.",
            "10x faster resume screening."
        ];
        var idx = 0, charIdx = 0, deleting = false;
        var el = document.getElementById('tw');
        function tick() {
            var phrase = phrases[idx];
            if (!deleting) {
                el.textContent = phrase.slice(0, ++charIdx);
                if (charIdx === phrase.length) { deleting = true; setTimeout(tick, 1800); return; }
            } else {
                el.textContent = phrase.slice(0, --charIdx);
                if (charIdx === 0) { deleting = false; idx = (idx + 1) % phrases.length; }
            }
            setTimeout(tick, deleting ? 38 : 62);
        }
        setTimeout(tick, 600);
    })();
    </script>
    """, height=44, scrolling=False)

if not st.session_state.get("authenticated", False):

    # -------- Login/Register Layout --------
    left, center, right = st.columns([1, 2, 1])

    # Shared spinner label style — used by all async buttons in this section
    st.markdown("""
    <style>
    .hly-spinner-wrap {
        display: flex; align-items: center; gap: 10px;
        padding: 10px 0 4px 0;
        font-family: -apple-system, BlinkMacSystemFont, sans-serif;
        font-size: 0.88rem; font-weight: 500; color: #94a3b8;
    }
    .hly-spinner-wrap svg { flex-shrink: 0; }
    </style>
    """, unsafe_allow_html=True)

    with center:
        st.markdown(
            """<div class='login-card'>
            <div style='text-align:center; margin-bottom:12px;'>
                <div style='display:inline-flex; align-items:center; justify-content:center; width:40px; height:40px; border-radius:10px; background:rgba(79,140,255,0.12); border:1px solid rgba(79,140,255,0.22); margin-bottom:12px;'>
                    <svg width="18" height="18" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <rect x="5" y="11" width="14" height="10" rx="2" stroke="#4f8cff" stroke-width="1.5" fill="rgba(79,140,255,0.12)"/>
                        <path d="M8 11V7a4 4 0 0 1 8 0v4" stroke="#4f8cff" stroke-width="1.5" stroke-linecap="round"/>
                        <circle cx="12" cy="16" r="1.2" fill="#4f8cff"/>
                    </svg>
                </div>
            </div>
            <h2 id='auth-heading' style='text-align:center; font-family:-apple-system,BlinkMacSystemFont,"SF Pro Display","Segoe UI",Roboto,sans-serif; font-size:1.3rem; font-weight:700; letter-spacing:-0.025em; color:#e6edf3; margin:0 0 18px 0;'>Sign in to <span style='color:#4f8cff;'>HIRELYZER</span></h2>
            <script>
            (function() {
                function updateHeading() {
                    var tabs = window.parent.document.querySelectorAll('[data-baseweb="tab"]');
                    var heading = window.parent.document.getElementById('auth-heading');
                    if (!heading || tabs.length < 2) return;
                    var activeTab = window.parent.document.querySelector('[data-baseweb="tab"][aria-selected="true"]');
                    if (activeTab) {
                        var label = activeTab.textContent.trim().toLowerCase();
                        if (label === 'register') {
                            heading.innerHTML = 'Register to <span style="color:#4f8cff;">HIRELYZER</span>';
                        } else {
                            heading.innerHTML = 'Sign in to <span style="color:#4f8cff;">HIRELYZER</span>';
                        }
                    }
                    tabs.forEach(function(tab) {
                        tab.addEventListener('click', function() {
                            setTimeout(updateHeading, 80);
                        });
                    });
                }
                setTimeout(updateHeading, 400);
                var observer = new MutationObserver(function() { updateHeading(); });
                setTimeout(function() {
                    var tabBar = window.parent.document.querySelector('[data-baseweb="tab-list"]');
                    if (tabBar) observer.observe(tabBar, { attributes: true, subtree: true, attributeFilter: ['aria-selected'] });
                }, 600);
            })();
            </script>""",
            unsafe_allow_html=True,
        )

        login_tab, register_tab = st.tabs(["Login", "Register"])

        # ---------------- LOGIN TAB ----------------
        with login_tab:
            # Show login or forgot password flow based on reset_stage
            if st.session_state.reset_stage == "none":

                # ── Show token error if magic link was invalid ──
                if st.session_state.get("_token_error"):
                    st.error(st.session_state.pop("_token_error"))

                # ── Pending magic link state ──
                if st.session_state.get("_magic_link_pending"):
                    _masked = st.session_state.get("_magic_link_email", "your registered email")
                    st.markdown(f"""
                    <div style='text-align:center; padding:28px 12px;'>
                        <div style='display:flex; align-items:center; justify-content:center; margin-bottom:14px;'>
                            <div style='display:inline-flex; align-items:center; justify-content:center;
                                        width:52px; height:52px; border-radius:14px;
                                        background:rgba(56,189,248,0.10); border:1px solid rgba(56,189,248,0.22);'>
                                <svg width="26" height="26" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                    <path d="M20 4H4a2 2 0 0 0-2 2v12a2 2 0 0 0 2 2h16a2 2 0 0 0 2-2V6a2 2 0 0 0-2-2z"
                                          stroke="#38bdf8" stroke-width="1.6" fill="rgba(56,189,248,0.08)"/>
                                    <path d="M2 8l10 7 10-7" stroke="#38bdf8" stroke-width="1.6" stroke-linecap="round"/>
                                </svg>
                            </div>
                        </div>
                        <div style='color:#e6edf3; font-size:1.05rem; font-weight:600; margin-bottom:8px;'>Check your inbox!</div>
                        <div style='color:#8b949e; font-size:0.88rem; line-height:1.6;'>
                            A login link has been sent to<br>
                            <strong style='color:#38bdf8; font-family:monospace;'>{_masked}</strong><br><br>
                            Click the link in the email to sign in.<br>
                            <span style='font-size:0.8rem;'>Link expires in 10 minutes.</span>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    render_notification("login")
                    if st.button("Back / Try Again", key="magic_link_back_btn", use_container_width=True):
                        st.session_state.pop("_magic_link_pending", None)
                        st.session_state.pop("_magic_link_email", None)
                        st.rerun()

                else:
                    # Normal Login UI
                    st.markdown("""<h3 style='color:#9aa4af; text-align:center; font-family:-apple-system,BlinkMacSystemFont,"SF Pro Display","Segoe UI",Roboto,sans-serif; font-size:0.82rem; font-weight:500; letter-spacing:0.06em; text-transform:uppercase; margin-bottom:24px;'>Welcome Back</h3>""", unsafe_allow_html=True)

                    user = st.text_input("Username or Email", key="login_user")
                    pwd = st.text_input("Password", type="password", key="login_pass")

                    # Render notification area (reserves space)
                    render_notification("login")

                    if st.button("Sign In", key="login_btn", use_container_width=True):
                        if user.strip() and pwd.strip():
                            # Detect admin: either typed "admin" as username,
                            # or typed an email that belongs to the admin account.
                            _input = user.strip()
                            _is_admin = _input.lower() == "admin"
                            if not _is_admin and '@' in _input:
                                # Check if this email maps to the admin username
                                _resolved = get_user_by_email(_input.lower())
                                if _resolved and _resolved.lower() == "admin":
                                    _is_admin = True

                            if _is_admin:
                                # ── Admin: magic link flow ──
                                status, message, _uname = send_login_link(_input, pwd.strip())
                                if status == "link_sent":
                                    st.session_state["_magic_link_pending"] = True
                                    st.session_state["_magic_link_email"] = "your registered email"
                                    notify("login", "success", "Login link sent to admin email! Click it to sign in.")
                                    st.rerun()
                                elif status == "bad_creds":
                                    notify("login", "error", message)
                                    st.rerun()
                                else:
                                    notify("login", "error", message)
                                    st.rerun()
                            else:
                                # ── Regular users: direct login ──
                                _allowed, _lock_msg = check_brute_force(_input)
                                if not _allowed:
                                    notify("login", "error", _lock_msg)
                                    st.rerun()
                                else:
                                    st.markdown("""
                                    <div class="hly-spinner-wrap">
                                        <svg width="16" height="16" viewBox="0 0 24 24" fill="none"
                                             xmlns="http://www.w3.org/2000/svg">
                                            <rect x="5" y="11" width="14" height="10" rx="2"
                                                  stroke="#4f8cff" stroke-width="1.6" fill="rgba(79,140,255,0.10)"/>
                                            <path d="M8 11V7a4 4 0 0 1 8 0v4"
                                                  stroke="#4f8cff" stroke-width="1.6" stroke-linecap="round"/>
                                            <circle cx="12" cy="16" r="1.2" fill="#4f8cff"/>
                                        </svg>
                                        Signing you in...
                                    </div>
                                    """, unsafe_allow_html=True)
                                    with st.spinner(""):
                                        success, saved_key = verify_user(_input, pwd.strip())
                                        if success:
                                            st.session_state.authenticated = True
                                            log_user_action(st.session_state.username, "login")
                                            notify("login", "success", "Login successful!")
                                            time.sleep(1.5)
                                    if success:
                                        st.rerun()
                                    else:
                                        notify("login", "error", "Invalid credentials. Please try again.")
                                        st.rerun()
                        else:
                            notify("login", "warning", "Please enter your username/email and password.")
                            st.rerun()

                    st.markdown("<br>", unsafe_allow_html=True)

                    # Forgot Password Link
                    if st.button("Forgot Password?", key="forgot_pw_link"):
                        st.session_state.reset_stage = "request_email"
                        st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 1: Request Email
            # ============================================================
            elif st.session_state.reset_stage == "request_email":
                st.markdown("""<h3 style='color:#9aa4af; text-align:center; font-family:-apple-system,BlinkMacSystemFont,"SF Pro Display","Segoe UI",Roboto,sans-serif; font-size:0.82rem; font-weight:500; letter-spacing:0.06em; text-transform:uppercase; margin-bottom:16px;'>Reset Password</h3>""", unsafe_allow_html=True)
                st.markdown("<p style='color:#c9d1d9; text-align:center;'>Enter your registered email to receive an OTP</p>", unsafe_allow_html=True)

                email_input = st.text_input("Email Address", key="reset_email_input")

                # Render notification area (reserves space)
                render_notification("login")

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("Send OTP", key="send_otp_btn", use_container_width=True):
                        if email_input.strip():
                            _email_exists = get_user_by_email(email_input.strip())
                            if _email_exists:
                                st.markdown("""
                                <div class="hly-spinner-wrap">
                                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                        <path d="M20 4H4a2 2 0 0 0-2 2v12a2 2 0 0 0 2 2h16a2 2 0 0 0 2-2V6a2 2 0 0 0-2-2z"
                                              stroke="#38bdf8" stroke-width="1.6" fill="rgba(56,189,248,0.10)"/>
                                        <path d="M2 8l10 7 10-7" stroke="#38bdf8" stroke-width="1.6" stroke-linecap="round"/>
                                    </svg>
                                    Sending OTP to your email...
                                </div>""", unsafe_allow_html=True)
                                with st.spinner(""):
                                    otp = generate_otp()
                                    success = send_email_otp(email_input.strip(), otp)
                                    if success:
                                        st.session_state.reset_email = email_input.strip()
                                        st.session_state.reset_otp = otp
                                        st.session_state.reset_otp_time = time.time()
                                        st.session_state.reset_stage = "verify_otp"
                                        notify("login", "success", "OTP sent successfully to your email!")
                                        time.sleep(0.5)
                                if success:
                                    st.rerun()
                                else:
                                    notify("login", "error", "Failed to send OTP. Please try again.")
                                    st.rerun()
                            else:
                                notify("login", "error", "Email not found. Please register first.")
                                st.rerun()
                        else:
                            notify("login", "warning", "Please enter your email address.")
                            st.rerun()

                with col2:
                    if st.button("Back to Login", key="back_to_login_1", use_container_width=True):
                        st.session_state.reset_stage = "none"
                        st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 2: Verify OTP
            # ============================================================
            elif st.session_state.reset_stage == "verify_otp":
                st.markdown("""<h3 style='color:#e6edf3; text-align:center; font-family:-apple-system,sans-serif; font-size:1.05rem; font-weight:600;'>
                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" style="vertical-align:-3px; margin-right:6px;" xmlns="http://www.w3.org/2000/svg"><path d="M4 4h16v16H4z" rx="2" stroke="#38bdf8" stroke-width="1.5" fill="none"/><path d="M4 9h16" stroke="#38bdf8" stroke-width="1.5"/><path d="M8 4v5" stroke="#38bdf8" stroke-width="1.5" stroke-linecap="round"/><path d="M16 4v5" stroke="#38bdf8" stroke-width="1.5" stroke-linecap="round"/></svg>
                    Verify OTP</h3>""", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#c9d1d9; text-align:center;'>Enter the 6-digit OTP sent to <strong>{mask_email(st.session_state.reset_email)}</strong></p>", unsafe_allow_html=True)

                # Calculate elapsed and remaining time (server-side)
                elapsed_time = time.time() - st.session_state.reset_otp_time
                remaining_time = max(0, int(180 - elapsed_time))

                # Display timer
                display_timer(remaining_time, expired=(remaining_time == 0), key_suffix="forgot_pw")

                # Check if OTP expired (3 minutes)
                if remaining_time == 0:
                    # OTP Expired - Show resend option
                    render_notification("login")
                    notify("login", "error", "OTP expired. Please request a new one.")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("Resend OTP", key="resend_otp_btn", use_container_width=True):
                            st.markdown("""
                            <div class="hly-spinner-wrap">
                                <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                    <path d="M20 4H4a2 2 0 0 0-2 2v12a2 2 0 0 0 2 2h16a2 2 0 0 0 2-2V6a2 2 0 0 0-2-2z"
                                          stroke="#38bdf8" stroke-width="1.6" fill="rgba(56,189,248,0.10)"/>
                                    <path d="M2 8l10 7 10-7" stroke="#38bdf8" stroke-width="1.6" stroke-linecap="round"/>
                                </svg>
                                Sending new OTP...
                            </div>""", unsafe_allow_html=True)
                            with st.spinner(""):
                                otp = generate_otp()
                                success = send_email_otp(st.session_state.reset_email, otp)
                                if success:
                                    st.session_state.reset_otp = otp
                                    st.session_state.reset_otp_time = time.time()
                                    notify("login", "info", "New OTP sent!")
                                    time.sleep(0.5)
                            if success:
                                st.rerun()
                            else:
                                notify("login", "error", "Failed to send OTP. Please try again.")
                                st.rerun()

                    with col2:
                        if st.button("Back to Login", key="back_to_login_expired", use_container_width=True):
                            st.session_state.reset_stage = "none"
                            st.rerun()
                else:
                    # OTP still valid - Show verification form
                    otp_input = st.text_input("Enter 6-Digit OTP", key="otp_input", max_chars=6)

                    # Render notification area (reserves space)
                    render_notification("login")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("Verify OTP", key="verify_otp_btn", use_container_width=True):
                            current_elapsed = time.time() - st.session_state.reset_otp_time
                            if current_elapsed >= 180:
                                st.markdown("""
                                <div class="hly-spinner-wrap">
                                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                        <circle cx="12" cy="12" r="9" stroke="#fca5a5" stroke-width="1.6"/>
                                        <line x1="12" y1="8" x2="12" y2="12" stroke="#fca5a5" stroke-width="2" stroke-linecap="round"/>
                                        <circle cx="12" cy="16" r="1" fill="#fca5a5"/>
                                    </svg>
                                    Checking OTP...
                                </div>""", unsafe_allow_html=True)
                                with st.spinner(""):
                                    notify("login", "error", "OTP has expired. Please request a new one.")
                                    time.sleep(0.6)
                                st.rerun()
                            elif otp_input.strip() == st.session_state.reset_otp:
                                st.markdown("""
                                <div class="hly-spinner-wrap">
                                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                        <path d="M9 12l2 2 4-4" stroke="#34d399" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round"/>
                                        <circle cx="12" cy="12" r="9" stroke="#34d399" stroke-width="1.6"/>
                                    </svg>
                                    Verifying OTP...
                                </div>""", unsafe_allow_html=True)
                                with st.spinner(""):
                                    st.session_state.reset_stage = "reset_password"
                                    notify("login", "success", "OTP verified successfully!")
                                    time.sleep(0.8)
                                st.rerun()
                            else:
                                st.markdown("""
                                <div class="hly-spinner-wrap">
                                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                        <circle cx="12" cy="12" r="9" stroke="#fca5a5" stroke-width="1.6"/>
                                        <line x1="15" y1="9" x2="9" y2="15" stroke="#fca5a5" stroke-width="2" stroke-linecap="round"/>
                                        <line x1="9" y1="9" x2="15" y2="15" stroke="#fca5a5" stroke-width="2" stroke-linecap="round"/>
                                    </svg>
                                    Checking OTP...
                                </div>""", unsafe_allow_html=True)
                                with st.spinner(""):
                                    notify("login", "error", "Invalid OTP. Please try again.")
                                    time.sleep(0.6)
                                st.rerun()

                    with col2:
                        if st.button("Back to Login", key="back_to_login_2", use_container_width=True):
                            st.session_state.reset_stage = "none"
                            st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 3: Reset Password
            # ============================================================
            elif st.session_state.reset_stage == "reset_password":
                st.markdown("""<h3 style='color:#e6edf3; text-align:center; font-family:-apple-system,sans-serif; font-size:1.05rem; font-weight:600;'>
                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" style="vertical-align:-3px; margin-right:6px;" xmlns="http://www.w3.org/2000/svg"><path d="M12 3a4 4 0 0 1 4 4v1H8V7a4 4 0 0 1 4-4z" stroke="#38bdf8" stroke-width="1.5" fill="none"/><rect x="5" y="11" width="14" height="10" rx="2" stroke="#38bdf8" stroke-width="1.5" fill="rgba(56,189,248,0.08)"/><circle cx="12" cy="16" r="1.5" fill="#38bdf8"/></svg>
                    Set New Password</h3>""", unsafe_allow_html=True)
                st.markdown("<p style='color:#c9d1d9; text-align:center;'>Enter your new password</p>", unsafe_allow_html=True)

                new_password = st.text_input("New Password", type="password", key="new_password_input")
                confirm_password = st.text_input("Confirm Password", type="password", key="confirm_password_input")

                st.caption("Password must be at least 8 characters, include uppercase, lowercase, number, and special character.")

                # Render notification area (reserves space)
                render_notification("login")

                if st.button("Reset Password", key="reset_password_btn", use_container_width=True):
                    if new_password.strip() and confirm_password.strip():
                        if new_password == confirm_password:
                            st.markdown("""
                            <div class="hly-spinner-wrap">
                                <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                    <rect x="5" y="11" width="14" height="10" rx="2"
                                          stroke="#4f8cff" stroke-width="1.6" fill="rgba(79,140,255,0.10)"/>
                                    <path d="M8 11V7a4 4 0 0 1 8 0v4"
                                          stroke="#4f8cff" stroke-width="1.6" stroke-linecap="round"/>
                                    <circle cx="12" cy="16" r="1.2" fill="#4f8cff"/>
                                </svg>
                                Resetting your password...
                            </div>""", unsafe_allow_html=True)
                            with st.spinner(""):
                                success = update_password_by_email(st.session_state.reset_email, new_password)
                                if success:
                                    log_user_action(st.session_state.reset_email, "password_reset")
                                    st.session_state.reset_stage = "none"
                                    st.session_state.reset_email = ""
                                    st.session_state.reset_otp = ""
                                    st.session_state.reset_otp_time = 0
                                    notify("login", "success", "Password reset successful! Please log in again.")
                                    time.sleep(1)
                            if success:
                                st.rerun()
                            else:
                                notify("login", "error", "Failed to reset password. Please try again.")
                                st.rerun()
                        else:
                            notify("login", "error", "Passwords do not match.")
                            st.rerun()
                    else:
                        notify("login", "warning", "Please fill in both password fields.")
                        st.rerun()

                if st.button("Back to Login", key="back_to_login_3"):
                    st.session_state.reset_stage = "none"
                    st.rerun()

        # ---------------- REGISTER TAB ----------------
        with register_tab:
            # Check if OTP was sent and pending verification
            if 'pending_registration' in st.session_state:
                st.markdown("""<h3 style='color:#e6edf3; text-align:center; font-family:-apple-system,sans-serif; font-size:1.05rem; font-weight:600;'>
                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" style="vertical-align:-3px; margin-right:6px;" xmlns="http://www.w3.org/2000/svg"><path d="M20 4H4a2 2 0 0 0-2 2v12a2 2 0 0 0 2 2h16a2 2 0 0 0 2-2V6a2 2 0 0 0-2-2z" stroke="#38bdf8" stroke-width="1.5" fill="none"/><path d="M2 8l10 7 10-7" stroke="#38bdf8" stroke-width="1.5"/></svg>
                    Verify Your Email</h3>""", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#c9d1d9; text-align:center;'>Enter the 6-digit OTP sent to <strong>{mask_email(st.session_state.pending_registration['email'])}</strong></p>", unsafe_allow_html=True)

                # Calculate remaining time
                from datetime import datetime
                elapsed = (datetime.now(st.session_state.pending_registration['timestamp'].tzinfo) - st.session_state.pending_registration['timestamp']).total_seconds()
                remaining = max(0, 180 - int(elapsed))

                # Display timer
                display_timer(remaining, expired=(remaining == 0), key_suffix="register")

                if remaining == 0:
                    # OTP Expired
                    render_notification("register")
                    notify("register", "error", "OTP expired. Please request a new one.")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("Resend OTP", key="reg_resend_expired_btn", use_container_width=True):
                            pending = st.session_state.pending_registration
                            st.markdown("""
                            <div class="hly-spinner-wrap">
                                <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                    <path d="M20 4H4a2 2 0 0 0-2 2v12a2 2 0 0 0 2 2h16a2 2 0 0 0 2-2V6a2 2 0 0 0-2-2z"
                                          stroke="#38bdf8" stroke-width="1.6" fill="rgba(56,189,248,0.10)"/>
                                    <path d="M2 8l10 7 10-7" stroke="#38bdf8" stroke-width="1.6" stroke-linecap="round"/>
                                </svg>
                                Sending new OTP...
                            </div>""", unsafe_allow_html=True)
                            with st.spinner(""):
                                success, message = add_user(pending['username'], pending['password'], pending['email'])
                                if success:
                                    notify("register", "success", "New OTP sent!")
                                    time.sleep(0.5)
                            if success:
                                st.rerun()
                            else:
                                notify("register", "error", message)
                                st.rerun()
                    with col2:
                        if st.button("Start Over", key="reg_start_over_btn", use_container_width=True):
                            del st.session_state.pending_registration
                            st.rerun()
                else:
                    # OTP still valid
                    otp_input = st.text_input("Enter 6-Digit OTP", key="reg_otp_input", max_chars=6)

                    # Render notification area (reserves space)
                    render_notification("register")

                    col1, col2, col3 = st.columns(3)
                    with col1:
                        if st.button("Verify", key="verify_reg_otp_btn", use_container_width=True):
                            cached_username = st.session_state.pending_registration['username']
                            current_elapsed = (datetime.now(st.session_state.pending_registration['timestamp'].tzinfo) - st.session_state.pending_registration['timestamp']).total_seconds()
                            if current_elapsed >= 180:
                                st.markdown("""
                                <div class="hly-spinner-wrap">
                                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                        <circle cx="12" cy="12" r="9" stroke="#fca5a5" stroke-width="1.6"/>
                                        <line x1="12" y1="8" x2="12" y2="12" stroke="#fca5a5" stroke-width="2" stroke-linecap="round"/>
                                        <circle cx="12" cy="16" r="1" fill="#fca5a5"/>
                                    </svg>
                                    Checking OTP...
                                </div>""", unsafe_allow_html=True)
                                with st.spinner(""):
                                    notify("register", "error", "OTP has expired. Please request a new one.")
                                    time.sleep(0.6)
                                st.rerun()
                            else:
                                st.markdown("""
                                <div class="hly-spinner-wrap">
                                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                        <path d="M9 12l2 2 4-4" stroke="#34d399" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round"/>
                                        <circle cx="12" cy="12" r="9" stroke="#34d399" stroke-width="1.6"/>
                                    </svg>
                                    Verifying OTP...
                                </div>""", unsafe_allow_html=True)
                                with st.spinner(""):
                                    success, message = complete_registration(otp_input.strip())
                                    if success:
                                        log_user_action(cached_username, "register")
                                        notify("register", "success", message)
                                        time.sleep(0.5)
                                    else:
                                        notify("register", "error", message)
                                        time.sleep(0.6)
                                st.rerun()

                    with col2:
                        if st.button("Resend", key="resend_reg_otp_btn", use_container_width=True):
                            pending = st.session_state.pending_registration
                            st.markdown("""
                            <div class="hly-spinner-wrap">
                                <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                                    <path d="M20 4H4a2 2 0 0 0-2 2v12a2 2 0 0 0 2 2h16a2 2 0 0 0 2-2V6a2 2 0 0 0-2-2z"
                                          stroke="#38bdf8" stroke-width="1.6" fill="rgba(56,189,248,0.10)"/>
                                    <path d="M2 8l10 7 10-7" stroke="#38bdf8" stroke-width="1.6" stroke-linecap="round"/>
                                </svg>
                                Sending new OTP...
                            </div>""", unsafe_allow_html=True)
                            with st.spinner(""):
                                success, message = add_user(pending['username'], pending['password'], pending['email'])
                                if success:
                                    notify("register", "info", "New OTP sent successfully!")
                                    time.sleep(0.5)
                            if success:
                                st.rerun()
                            else:
                                notify("register", "error", message)
                                st.rerun()

                    with col3:
                        if st.button("Back", key="back_to_reg_btn", use_container_width=True):
                            del st.session_state.pending_registration
                            st.rerun()

            else:
                # Normal registration form
                st.markdown("""
                <h3 style='text-align:center; color:#e6edf3; font-family:-apple-system,BlinkMacSystemFont,sans-serif;
                            font-size:1.05rem; font-weight:600; display:flex; align-items:center;
                            justify-content:center; gap:8px; margin-bottom:16px;'>
                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg"
                         style="vertical-align:-2px; flex-shrink:0;">
                        <path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"
                              stroke="#38bdf8" stroke-width="1.6" fill="rgba(56,189,248,0.08)"/>
                        <polyline points="14 2 14 8 20 8" stroke="#38bdf8" stroke-width="1.6" stroke-linejoin="round"/>
                        <line x1="12" y1="12" x2="12" y2="18" stroke="#38bdf8" stroke-width="1.6" stroke-linecap="round"/>
                        <line x1="9" y1="15" x2="15" y2="15" stroke="#38bdf8" stroke-width="1.6" stroke-linecap="round"/>
                    </svg>
                    Register New User
                </h3>""", unsafe_allow_html=True)

                # ── CSS: fixed-height validation slot — zero layout shift ──
                st.markdown("""
                <style>
                /* Outer wrapper: zero document-flow height — no spacing contribution */
                .val-slot {
                    height: 0;
                    overflow: visible;
                    position: relative;
                    margin: 0;
                    padding: 0;
                    line-height: 0;
                }
                /* The badge floats above the next field via negative top offset */
                .val-badge {
                    position: absolute;
                    top: -26px;
                    left: 0;
                    right: 0;
                    display: flex;
                    align-items: center;
                    gap: 6px;
                    padding: 3px 10px;
                    border-radius: 5px;
                    font-size: 0.75rem;
                    font-weight: 500;
                    font-family: var(--font-sans), -apple-system, sans-serif;
                    line-height: 1.3;
                    opacity: 0;
                    transform: translateY(2px);
                    transition: opacity 0.18s ease, transform 0.18s ease;
                    pointer-events: none;
                    white-space: nowrap;
                    overflow: hidden;
                    text-overflow: ellipsis;
                    z-index: 10;
                }
                .val-badge.val-visible {
                    opacity: 1;
                    transform: translateY(0);
                    pointer-events: auto;
                }
                .val-badge.val-success {
                    background: rgba(52,211,153,0.12);
                    border: 1px solid rgba(52,211,153,0.25);
                    color: #6ee7b7;
                }
                .val-badge.val-error {
                    background: rgba(251,113,133,0.12);
                    border: 1px solid rgba(251,113,133,0.25);
                    color: #fca5a5;
                }
                .val-badge.val-warn {
                    background: rgba(251,191,36,0.12);
                    border: 1px solid rgba(251,191,36,0.25);
                    color: #fde68a;
                }
                @keyframes _val_autofade {
                    0%   { opacity: 1; }
                    65%  { opacity: 1; }
                    100% { opacity: 0; }
                }
                .val-badge.val-autofade {
                    animation: _val_autofade 3.5s ease forwards;
                }
                </style>
                """, unsafe_allow_html=True)

                # ── on_change callbacks — DB hits ONLY when field value changes ──
                def _validate_email():
                    val = st.session_state.get("reg_email", "").strip()
                    if not val:
                        st.session_state._email_msg = ("", "")
                        return
                    if not is_valid_email(val):
                        st.session_state._email_msg = ("warn", "Invalid email format.")
                    elif email_exists(val):
                        st.session_state._email_msg = ("error", "Email already registered.")
                    else:
                        st.session_state._email_msg = ("success", "Email is available.")

                def _validate_username():
                    val = st.session_state.get("reg_user", "").strip()
                    if not val:
                        st.session_state._user_msg = ("", "")
                        return
                    if username_exists(val):
                        st.session_state._user_msg = ("error", "Username already exists.")
                    else:
                        st.session_state._user_msg = ("success", "Username is available.")

                def _validate_password():
                    val = st.session_state.get("reg_pass", "")
                    if not val:
                        st.session_state._pass_msg = ("", "")
                        return
                    if not is_strong_password(val):
                        st.session_state._pass_msg = ("warn", "Password must be at least 8 characters and strong.")
                    else:
                        st.session_state._pass_msg = ("success", "Strong password.")

                # Initialise message state once
                if "_email_msg" not in st.session_state:
                    st.session_state._email_msg = ("", "")
                if "_user_msg" not in st.session_state:
                    st.session_state._user_msg = ("", "")
                if "_pass_msg" not in st.session_state:
                    st.session_state._pass_msg = ("", "")

                def _render_val_msg(state_key):
                    """Render a compact validation badge with SVG icon in a fixed-height slot — zero layout shift."""
                    kind, text = st.session_state.get(state_key, ("", ""))
                    if not kind or not text:
                        st.markdown(
                            '<div class="val-slot"><div class="val-badge"></div></div>',
                            unsafe_allow_html=True
                        )
                        return
                    type_class = {"warn": "val-warn", "error": "val-error", "success": "val-success"}.get(kind, "val-warn")
                    fade_class = " val-autofade" if kind == "success" else ""
                    # SVG icon per type
                    _icons = {
                        "success": '<svg width="11" height="11" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg"><path d="M9 12l2 2 4-4" stroke="#34d399" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/><circle cx="12" cy="12" r="9" stroke="#34d399" stroke-width="1.8"/></svg>',
                        "error":   '<svg width="11" height="11" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg"><circle cx="12" cy="12" r="9" stroke="#fca5a5" stroke-width="1.8"/><line x1="15" y1="9" x2="9" y2="15" stroke="#fca5a5" stroke-width="2" stroke-linecap="round"/><line x1="9" y1="9" x2="15" y2="15" stroke="#fca5a5" stroke-width="2" stroke-linecap="round"/></svg>',
                        "warn":    '<svg width="11" height="11" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg"><path d="M12 9v4M12 17h.01" stroke="#fde68a" stroke-width="2" stroke-linecap="round"/><path d="M10.29 3.86L1.82 18a2 2 0 0 0 1.71 3h16.94a2 2 0 0 0 1.71-3L13.71 3.86a2 2 0 0 0-3.42 0z" stroke="#fde68a" stroke-width="1.8" stroke-linejoin="round"/></svg>',
                    }
                    icon_svg = _icons.get(kind, "")
                    st.markdown(
                        f'<div class="val-slot"><div class="val-badge {type_class} val-visible{fade_class}">{icon_svg} {text}</div></div>',
                        unsafe_allow_html=True
                    )

                # ── Inputs wired to on_change — NO inline DB calls ──
                new_email = st.text_input(
                    "Email", key="reg_email",
                    placeholder="your@email.com",
                    on_change=_validate_email
                )
                _render_val_msg("_email_msg")

                new_user = st.text_input(
                    "Username", key="reg_user",
                    on_change=_validate_username
                )
                _render_val_msg("_user_msg")

                new_pass = st.text_input(
                    "Password", type="password", key="reg_pass",
                    on_change=_validate_password
                )
                st.caption("Password must be at least 8 characters, include uppercase, lowercase, number, and special character.")
                _render_val_msg("_pass_msg")

                # Render notification area (reserves space)
                render_notification("register")

                st.markdown("""
                <div style="display:flex; align-items:center; gap:8px; margin-bottom:4px; color:#94a3b8;
                            font-size:0.78rem; font-family:-apple-system,BlinkMacSystemFont,sans-serif;">
                    <svg width="13" height="13" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <path d="M20 4H4a2 2 0 0 0-2 2v12a2 2 0 0 0 2 2h16a2 2 0 0 0 2-2V6a2 2 0 0 0-2-2z"
                              stroke="#38bdf8" stroke-width="1.6" fill="rgba(56,189,248,0.08)"/>
                        <path d="M2 8l10 7 10-7" stroke="#38bdf8" stroke-width="1.6" stroke-linecap="round"/>
                    </svg>
                    A verification code will be sent to your email
                </div>""", unsafe_allow_html=True)
                if st.button("Register & Send OTP", key="register_btn", use_container_width=True):
                    if new_email.strip() and new_user.strip() and new_pass.strip():
                        # Validate before attempting registration
                        if not is_valid_email(new_email.strip()):
                            notify("register", "warning", "Invalid email format.")
                            st.rerun()
                        elif email_exists(new_email.strip()):
                            notify("register", "error", "Email already registered.")
                            st.rerun()
                        elif username_exists(new_user.strip()):
                            notify("register", "error", "Username already exists.")
                            st.rerun()
                        else:
                            st.markdown("""
                            <div class="hly-spinner-wrap">
                                <svg width="16" height="16" viewBox="0 0 24 24" fill="none"
                                     xmlns="http://www.w3.org/2000/svg">
                                    <path d="M20 4H4a2 2 0 0 0-2 2v12a2 2 0 0 0 2 2h16a2 2 0 0 0 2-2V6a2 2 0 0 0-2-2z"
                                          stroke="#38bdf8" stroke-width="1.6" fill="rgba(56,189,248,0.10)"/>
                                    <path d="M2 8l10 7 10-7"
                                          stroke="#38bdf8" stroke-width="1.6" stroke-linecap="round"/>
                                </svg>
                                Sending verification OTP to your email...
                            </div>
                            """, unsafe_allow_html=True)
                            with st.spinner(""):
                                success, message = add_user(new_user.strip(), new_pass.strip(), new_email.strip())
                                if success:
                                    notify("register", "success", message)
                                    time.sleep(0.8)
                            if success:
                                st.rerun()
                            else:
                                notify("register", "error", message)
                                st.rerun()
                    else:
                        notify("register", "warning", "Please fill in all fields (email, username, and password).")
                        st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)

    st.stop()

# ------------------- AFTER LOGIN -------------------
if st.session_state.get("authenticated"):
    st.markdown(
        f'<div class="welcome-banner">'
        f'<div>'
        f'<div class="welcome-title">Welcome back, <span class="welcome-username">{st.session_state.username}</span> 👋</div>'
        f'<div class="welcome-subtitle">HIRELYZER — AI-Powered Resume Intelligence Platform</div>'
        f'</div>'
        f'<div style="display:flex;align-items:center;gap:8px;">'
        f'<div style="background:linear-gradient(135deg,rgba(52,211,153,0.15) 0%,rgba(52,211,153,0.06) 100%);border:1px solid rgba(52,211,153,0.25);border-radius:99px;padding:5px 14px;font-size:0.75rem;font-weight:600;color:#6ee7b7;letter-spacing:0.04em;text-transform:uppercase;font-family:-apple-system,sans-serif;">&#9679; Live</div>'
        f'</div>'
        f'</div>',
        unsafe_allow_html=True
    )

    # 🔓 LOGOUT BUTTON
    if st.button("🚪 Logout"):
        log_user_action(st.session_state.get("username", "unknown"), "logout")

        # ✅ Clear all session keys safely
        for key in list(st.session_state.keys()):
            del st.session_state[key]

        st.success("✅ Logged out successfully.")
        st.rerun()  # Force rerun to prevent stale UI


if st.session_state.username == "admin":
    _adm_hdr_col, _adm_btn_col = st.columns([6, 1])
    with _adm_hdr_col:
        st.markdown("""
    <div class="admin-header">
        <h2>
        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="#4f8cff" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round" style="vertical-align:-3px;margin-right:8px;"><path d="M12 22s8-4 8-10V5l-8-3-8 3v7c0 6 8 10 8 10z"/></svg>
        Admin Control Panel</h2>
    </div>
    """, unsafe_allow_html=True)
    with _adm_btn_col:
        st.markdown("<div style='padding-top:18px;'>", unsafe_allow_html=True)
        if st.button("↻ Refresh", key="admin_refresh_btn", help="Force fetch latest data from database", use_container_width=True):
            _cached_admin_metrics.clear()
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    # Metrics row — cached, no Supabase hit on every rerun
    _reg_users, _logins_today, _logs = _cached_admin_metrics()

    # ── Build analytics dataframe from logs ──────────────────────────────────
    import pandas as pd
    from collections import Counter

    if _logs:
        _df_logs = pd.DataFrame(_logs, columns=["Username", "Action", "Timestamp"])
        _df_logs["Timestamp"] = pd.to_datetime(_df_logs["Timestamp"], errors="coerce")
        _df_logs = _df_logs.dropna(subset=["Timestamp"])
        _df_logs["Date"]     = _df_logs["Timestamp"].dt.date
        _df_logs["Hour"]     = _df_logs["Timestamp"].dt.hour
        _df_logs["DayName"]  = _df_logs["Timestamp"].dt.day_name()
        _login_df  = _df_logs[_df_logs["Action"] == "login"]
        _logout_df = _df_logs[_df_logs["Action"] == "logout"]
    else:
        _df_logs = pd.DataFrame(columns=["Username","Action","Timestamp","Date","Hour","DayName"])
        _login_df  = _df_logs.copy()
        _logout_df = _df_logs.copy()

    # ── KPI Row ───────────────────────────────────────────────────────────────
    _total_logins   = len(_login_df)
    _unique_users   = _login_df["Username"].nunique() if not _login_df.empty else 0
    _total_actions  = len(_df_logs)
    _peak_hour      = int(_login_df["Hour"].mode()[0]) if not _login_df.empty else 0

    _peak_full = f"{_peak_hour:02d}:00 – {(_peak_hour+1)%24:02d}:00"
    st.markdown(f"""
    <style>
    .kpi-grid {{
        display: grid;
        grid-template-columns: repeat(4, 1fr);
        gap: 14px;
        margin-bottom: 8px;
    }}
    .kpi-card {{
        background: #0d1117;
        border: 1px solid rgba(56,189,248,0.18);
        border-radius: 12px;
        padding: 18px 20px 14px 20px;
        position: relative;
        cursor: default;
        transition: border-color 0.25s, box-shadow 0.25s, transform 0.2s;
        overflow: visible;
    }}
    .kpi-card:hover {{
        border-color: rgba(56,189,248,0.55);
        box-shadow: 0 0 18px rgba(56,189,248,0.15);
        transform: translateY(-2px);
    }}
    .kpi-label {{
        font-size: 0.78rem;
        color: #8b949e;
        font-weight: 500;
        letter-spacing: 0.02em;
        margin-bottom: 6px;
    }}
    .kpi-value {{
        font-size: 1.95rem;
        font-weight: 700;
        color: #e6edf3;
        line-height: 1.15;
        white-space: nowrap;
        overflow: hidden;
        text-overflow: ellipsis;
    }}
    .kpi-sub {{
        font-size: 0.7rem;
        color: #8b949e;
        margin-top: 8px;
        display: flex;
        align-items: center;
        gap: 4px;
    }}
    /* Tooltip */
    .kpi-card .kpi-tooltip {{
        visibility: hidden;
        opacity: 0;
        background: #161b22;
        border: 1px solid rgba(56,189,248,0.4);
        color: #e6edf3;
        font-size: 0.8rem;
        padding: 7px 13px;
        border-radius: 8px;
        white-space: nowrap;
        position: absolute;
        bottom: calc(100% + 8px);
        left: 50%;
        transform: translateX(-50%);
        z-index: 9999;
        pointer-events: none;
        transition: opacity 0.2s;
        box-shadow: 0 4px 16px rgba(0,0,0,0.5);
    }}
    .kpi-card .kpi-tooltip::after {{
        content: "";
        position: absolute;
        top: 100%;
        left: 50%;
        transform: translateX(-50%);
        border: 6px solid transparent;
        border-top-color: rgba(56,189,248,0.4);
    }}
    .kpi-card:hover .kpi-tooltip {{
        visibility: visible;
        opacity: 1;
    }}
    </style>
    <div class="kpi-grid">
      <div class="kpi-card">
        <div class="kpi-tooltip">Total registered accounts: {_reg_users}</div>
        <div class="kpi-label">Registered Users</div>
        <div class="kpi-value">{_reg_users}</div>
        <div class="kpi-sub">
          <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#8b949e" stroke-width="2"><path d="M20 21v-2a4 4 0 0 0-4-4H8a4 4 0 0 0-4 4v2"/><circle cx="12" cy="7" r="4"/></svg>
          Total accounts
        </div>
      </div>
      <div class="kpi-card">
        <div class="kpi-tooltip">Logins since midnight IST: {_logins_today}</div>
        <div class="kpi-label">Logins Today (IST)</div>
        <div class="kpi-value">{_logins_today}</div>
        <div class="kpi-sub">
          <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#8b949e" stroke-width="2"><rect x="3" y="4" width="18" height="18" rx="2"/><line x1="16" y1="2" x2="16" y2="6"/><line x1="8" y1="2" x2="8" y2="6"/><line x1="3" y1="10" x2="21" y2="10"/></svg>
          Since midnight IST
        </div>
      </div>
      <div class="kpi-card">
        <div class="kpi-tooltip">All-time cumulative logins: {_total_logins}</div>
        <div class="kpi-label">Total Logins (All Time)</div>
        <div class="kpi-value">{_total_logins}</div>
        <div class="kpi-sub">
          <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#8b949e" stroke-width="2"><path d="M15 3h4a2 2 0 0 1 2 2v14a2 2 0 0 1-2 2h-4"/><polyline points="10 17 15 12 10 7"/><line x1="15" y1="12" x2="3" y2="12"/></svg>
          Cumulative logins
        </div>
      </div>
      <div class="kpi-card">
        <div class="kpi-tooltip">Busiest login window: {_peak_full} IST</div>
        <div class="kpi-label">Peak Hour</div>
        <div class="kpi-value" style="font-size:1.55rem;">{_peak_full}</div>
        <div class="kpi-sub">
          <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#8b949e" stroke-width="2"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg>
          Busiest login hour
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Charts — helper to render chart OR table with a toggle button ─────────
    def _chart_panel(key, title_html, chart_fn, table_df):
        """Render a labelled chart with a chart/table toggle. chart_fn() returns an Altair chart."""
        _tkey = f"_tbl_view_{key}"
        if _tkey not in st.session_state:
            st.session_state[_tkey] = False
        _hdr, _btn = st.columns([6, 1])
        with _hdr:
            st.markdown(title_html, unsafe_allow_html=True)
        with _btn:
            _label = "📊" if st.session_state[_tkey] else "📋"
            _tip   = "Back to Chart" if st.session_state[_tkey] else "View as Table"
            if st.button(_label, key=f"_toggle_{key}", use_container_width=True, help=_tip):
                st.session_state[_tkey] = not st.session_state[_tkey]
                st.rerun()
        if st.session_state[_tkey]:
            st.dataframe(table_df.reset_index(drop=True), use_container_width=True, height=260)
        else:
            st.altair_chart(chart_fn(), use_container_width=True)

    # ── Charts Row 1: Logins per Day + Logins by Day of Week ─────────────────
    if not _login_df.empty:
        ch_col1, ch_col2 = st.columns(2)

        with ch_col1:
            _daily = _login_df.groupby("Date").size().reset_index(name="Logins")
            _daily["Date"] = _daily["Date"].astype(str)
            _max_day = _daily.loc[_daily["Logins"].idxmax(), "Date"]
            _min_day = _daily.loc[_daily["Logins"].idxmin(), "Date"]
            _daily["Color"] = _daily["Date"].apply(
                lambda d: "#34d399" if d == _max_day else ("#f87171" if d == _min_day else "#38bdf8")
            )
            _daily_w = max(260, len(_daily) * 28)
            def _make_daily_chart():
                return (
                    alt.Chart(_daily)
                    .mark_bar(cornerRadiusTopLeft=4, cornerRadiusTopRight=4)
                    .encode(
                        x=alt.X("Date:O", axis=alt.Axis(labelAngle=-45, labelColor="#8b949e", titleColor="#8b949e", labelFontSize=10), title="Date"),
                        y=alt.Y("Logins:Q", axis=alt.Axis(labelColor="#8b949e", titleColor="#8b949e"), title="Login Count"),
                        color=alt.Color("Color:N", scale=None, legend=None),
                        tooltip=[alt.Tooltip("Date:O", title="Date"), alt.Tooltip("Logins:Q", title="Logins")]
                    )
                    .properties(height=260, width=_daily_w, background="transparent")
                    .configure_view(strokeWidth=0)
                    .configure_axis(labelLimit=120)
                )
            _chart_panel(
                "daily",
                """<p class='section-label'><svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align:-2px;margin-right:5px;"><polyline points="22 12 18 12 15 21 9 3 6 12 2 12"/></svg>Daily Login Trend</p>""",
                _make_daily_chart,
                _daily[["Date", "Logins"]],
            )
            st.markdown(f'<p style="font-size:0.75rem;color:#8b949e;display:flex;gap:14px;"><span style="color:#34d399;display:flex;align-items:center;gap:4px;"><svg width="10" height="10" viewBox="0 0 24 24" fill="#34d399"><circle cx="12" cy="12" r="10"/></svg> Peak: <b style="color:#e6edf3;">{_max_day}</b></span><span style="color:#f87171;display:flex;align-items:center;gap:4px;"><svg width="10" height="10" viewBox="0 0 24 24" fill="#f87171"><circle cx="12" cy="12" r="10"/></svg> Lowest: <b style="color:#e6edf3;">{_min_day}</b></span></p>', unsafe_allow_html=True)

        with ch_col2:
            _dow_order = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
            _dow = _login_df.groupby("DayName").size().reset_index(name="Logins")
            _dow["DayName"] = pd.Categorical(_dow["DayName"], categories=_dow_order, ordered=True)
            _dow = _dow.sort_values("DayName")
            def _make_dow_chart():
                return (
                    alt.Chart(_dow)
                    .mark_bar(cornerRadiusTopLeft=4, cornerRadiusTopRight=4, color="#818cf8")
                    .encode(
                        x=alt.X("DayName:O", sort=_dow_order, axis=alt.Axis(labelAngle=-25, labelColor="#8b949e", titleColor="#8b949e"), title="Day"),
                        y=alt.Y("Logins:Q", axis=alt.Axis(labelColor="#8b949e", titleColor="#8b949e"), title="Login Count"),
                        tooltip=[alt.Tooltip("DayName:O", title="Day"), alt.Tooltip("Logins:Q", title="Logins")]
                    )
                    .properties(height=260, background="transparent")
                    .configure_view(strokeWidth=0)
                )
            _chart_panel(
                "dow",
                """<p class='section-label'><svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align:-2px;margin-right:5px;"><rect x="3" y="4" width="18" height="18" rx="2"/><line x1="16" y1="2" x2="16" y2="6"/><line x1="8" y1="2" x2="8" y2="6"/><line x1="3" y1="10" x2="21" y2="10"/></svg>Logins by Day of Week</p>""",
                _make_dow_chart,
                _dow[["DayName", "Logins"]],
            )

        # ── Charts Row 2: Peak Hour + Top Users ───────────────────────────────
        ch_col3, ch_col4 = st.columns(2)

        with ch_col3:
            _hourly = _login_df.groupby("Hour").size().reset_index(name="Logins")
            _all_hours = pd.DataFrame({"Hour": range(24)})
            _hourly = _all_hours.merge(_hourly, on="Hour", how="left").fillna(0)
            _hourly["Logins"] = _hourly["Logins"].astype(int)
            _hourly["HourLabel"] = _hourly["Hour"].apply(lambda h: f"{h:02d}:00-{(h+1)%24:02d}:00")
            _peak_label = f"{_peak_hour:02d}:00 – {(_peak_hour+1)%24:02d}:00"
            def _make_hour_chart():
                return (
                    alt.Chart(_hourly)
                    .mark_bar(cornerRadiusTopLeft=3, cornerRadiusTopRight=3)
                    .encode(
                        x=alt.X("HourLabel:O", sort=None, axis=alt.Axis(labelAngle=-45, labelColor="#8b949e", titleColor="#8b949e", labelFontSize=9), title="Hour (IST)"),
                        y=alt.Y("Logins:Q", axis=alt.Axis(labelColor="#8b949e", titleColor="#8b949e"), title="Logins"),
                        color=alt.Color("Logins:Q", scale=alt.Scale(scheme="blues"), legend=None),
                        tooltip=[alt.Tooltip("HourLabel:O", title="Hour"), alt.Tooltip("Logins:Q", title="Logins")]
                    )
                    .properties(height=260, background="transparent")
                    .configure_view(strokeWidth=0)
                )
            _chart_panel(
                "hour",
                """<p class='section-label'><svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align:-2px;margin-right:5px;"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg>Peak Login Hours</p>""",
                _make_hour_chart,
                _hourly[["HourLabel", "Logins"]].rename(columns={"HourLabel": "Hour (IST)"}),
            )
            st.markdown(f'<p style="font-size:0.75rem;color:#8b949e;display:flex;align-items:center;gap:5px;"><svg width="11" height="11" viewBox="0 0 24 24" fill="none" stroke="#fbbf24" stroke-width="2.5"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg> Busiest: <b style="color:#e6edf3;">{_peak_label} IST</b></p>', unsafe_allow_html=True)

        with ch_col4:
            _top_users = _login_df.groupby("Username").size().reset_index(name="Logins").sort_values("Logins", ascending=False).head(10)
            _users_h = max(260, len(_top_users) * 32)
            def _make_users_chart():
                return (
                    alt.Chart(_top_users)
                    .mark_bar(cornerRadiusTopLeft=4, cornerRadiusTopRight=4, color="#f472b6")
                    .encode(
                        x=alt.X("Logins:Q", axis=alt.Axis(labelColor="#8b949e", titleColor="#8b949e"), title="Login Count"),
                        y=alt.Y("Username:N", sort="-x", axis=alt.Axis(labelColor="#8b949e", titleColor="#8b949e", labelLimit=150), title="User"),
                        tooltip=[alt.Tooltip("Username:N", title="User"), alt.Tooltip("Logins:Q", title="Logins")]
                    )
                    .properties(height=_users_h, background="transparent")
                    .configure_view(strokeWidth=0)
                )
            _chart_panel(
                "users",
                """<p class='section-label'><svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align:-2px;margin-right:5px;"><path d="M17 21v-2a4 4 0 0 0-4-4H5a4 4 0 0 0-4 4v2"/><circle cx="9" cy="7" r="4"/><path d="M23 21v-2a4 4 0 0 0-3-3.87"/><path d="M16 3.13a4 4 0 0 1 0 7.75"/></svg>Most Active Users</p>""",
                _make_users_chart,
                _top_users[["Username", "Logins"]],
            )

        # ── Charts Row 3: Action Breakdown + Login vs Logout ──────────────────
        ch_col5, ch_col6 = st.columns(2)

        with ch_col5:
            _actions = _df_logs.groupby("Action").size().reset_index(name="Count")
            _action_colors = {"login": "#34d399", "logout": "#f87171", "register": "#38bdf8", "password_reset": "#fbbf24"}
            _actions["Color"] = _actions["Action"].map(lambda a: _action_colors.get(a, "#a78bfa"))
            def _make_actions_chart():
                return (
                    alt.Chart(_actions)
                    .mark_arc(innerRadius=55, outerRadius=100)
                    .encode(
                        theta=alt.Theta("Count:Q"),
                        color=alt.Color("Color:N", scale=None, legend=alt.Legend(labelColor="#8b949e", titleColor="#8b949e")),
                        tooltip=[alt.Tooltip("Action:N", title="Action"), alt.Tooltip("Count:Q", title="Count")]
                    )
                    .properties(height=260, background="transparent")
                    .configure_view(strokeWidth=0)
                )
            _chart_panel(
                "actions",
                """<p class='section-label'><svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align:-2px;margin-right:5px;"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>Action Breakdown</p>""",
                _make_actions_chart,
                _actions[["Action", "Count"]],
            )
            # Legend below donut
            _total_actions_count = _actions["Count"].sum()
            # Largest remainder method — guarantees percentages always sum to exactly 100%
            _exact_pcts = [100 * r["Count"] / _total_actions_count for _, r in _actions.iterrows()]
            _floored_pcts = [int(x) for x in _exact_pcts]
            _remainders = sorted(enumerate(_exact_pcts), key=lambda x: x[1] - int(x[1]), reverse=True)
            for i in range(100 - sum(_floored_pcts)):
                _floored_pcts[_remainders[i][0]] += 1
            _legend_items = "".join([
                f'<div style="display:flex;align-items:center;justify-content:space-between;'
                f'padding:5px 10px;border-radius:6px;background:rgba(255,255,255,0.03);margin-bottom:4px;">'
                f'<span style="display:flex;align-items:center;gap:7px;">'
                f'<svg width="10" height="10" viewBox="0 0 10 10"><circle cx="5" cy="5" r="5" fill="{_action_colors.get(row["Action"], "#a78bfa")}"/></svg>'
                f'<span style="color:#c9d1d9;font-size:0.8rem;text-transform:capitalize;">{row["Action"].replace("_"," ")}</span>'
                f'</span>'
                f'<span style="color:#8b949e;font-size:0.8rem;">'
                f'<b style="color:#e6edf3;">{row["Count"]}</b>'
                f'&nbsp;<span style="font-size:0.72rem;">({_floored_pcts[idx]}%)</span>'
                f'</span>'
                f'</div>'
                for idx, (_, row) in enumerate(_actions.iterrows())
            ])
            st.markdown(
                f'<div style="margin-top:6px;">{_legend_items}</div>',
                unsafe_allow_html=True,
            )

        with ch_col6:
            _ll = _df_logs[_df_logs["Action"].isin(["login","logout"])].copy()
            _ll["Date"] = _ll["Date"].astype(str)
            _ll_grouped = _ll.groupby(["Date","Action"]).size().reset_index(name="Count")
            _ll_dates = _ll_grouped["Date"].nunique()
            _ll_w = max(260, _ll_dates * 28)
            def _make_ll_chart():
                return (
                    alt.Chart(_ll_grouped)
                    .mark_line(point=True, strokeWidth=2)
                    .encode(
                        x=alt.X("Date:O", axis=alt.Axis(labelAngle=-45, labelColor="#8b949e", titleColor="#8b949e", labelFontSize=10), title="Date"),
                        y=alt.Y("Count:Q", axis=alt.Axis(labelColor="#8b949e", titleColor="#8b949e"), title="Count"),
                        color=alt.Color("Action:N", scale=alt.Scale(domain=["login","logout"], range=["#34d399","#f87171"]),
                                        legend=alt.Legend(labelColor="#c9d1d9", titleColor="#8b949e")),
                        tooltip=[alt.Tooltip("Date:O"), alt.Tooltip("Action:N"), alt.Tooltip("Count:Q")]
                    )
                    .properties(height=260, width=_ll_w, background="transparent")
                    .configure_view(strokeWidth=0)
                    .configure_axis(labelLimit=120)
                )
            _chart_panel(
                "ll",
                """<p class='section-label'><svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align:-2px;margin-right:5px;"><line x1="18" y1="20" x2="18" y2="10"/><line x1="12" y1="20" x2="12" y2="4"/><line x1="6" y1="20" x2="6" y2="14"/></svg>Login vs Logout Over Time</p>""",
                _make_ll_chart,
                _ll_grouped,
            )

    else:
        st.markdown("""
        <div style="display:flex;align-items:center;gap:10px;padding:16px;background:rgba(56,189,248,0.06);border:1px solid rgba(56,189,248,0.15);border-radius:8px;color:#8b949e;font-size:0.88rem;">
            <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>
            Not enough data yet to render charts. Check back after more users log in.
        </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Activity Log with filters ─────────────────────────────────────────────
    st.markdown("""<p class='section-label'>
        <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align:-2px;margin-right:5px;"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/><polyline points="10 9 9 9 8 9"/></svg>
        Activity Log</p>""", unsafe_allow_html=True)

    if _logs:
        _f1, _f2, _f3 = st.columns([2, 2, 1])
        with _f1:
            _all_users = ["All"] + sorted(_df_logs["Username"].unique().tolist())
            _filter_user = st.selectbox("Filter by User", _all_users, key="admin_filter_user")
        with _f2:
            _all_actions = ["All"] + sorted(_df_logs["Action"].unique().tolist())
            _filter_action = st.selectbox("Filter by Action", _all_actions, key="admin_filter_action")
        with _f3:
            _filter_rows = st.selectbox("Show rows", [25, 50, 100, 200, "All"], key="admin_filter_rows")

        _filtered = _df_logs.copy()
        if _filter_user != "All":
            _filtered = _filtered[_filtered["Username"] == _filter_user]
        if _filter_action != "All":
            _filtered = _filtered[_filtered["Action"] == _filter_action]
        _filtered = _filtered.sort_values("Timestamp", ascending=False)
        if _filter_rows != "All":
            _filtered = _filtered.head(int(_filter_rows))

        st.dataframe(
            _filtered[["Username", "Action", "Timestamp"]].reset_index(drop=True),
            use_container_width=True,
            height=320
        )
        st.markdown(f'<p style="font-size:0.75rem;color:#8b949e;display:flex;align-items:center;gap:5px;margin-top:4px;"><svg width="11" height="11" viewBox="0 0 24 24" fill="none" stroke="#8b949e" stroke-width="2"><circle cx="11" cy="11" r="8"/><line x1="21" y1="21" x2="16.65" y2="16.65"/></svg> Showing <b style="color:#c9d1d9;">{len(_filtered)}</b> of <b style="color:#c9d1d9;">{len(_df_logs)}</b> total records</p>', unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="display:flex;align-items:center;gap:10px;padding:16px;background:rgba(56,189,248,0.06);border:1px solid rgba(56,189,248,0.15);border-radius:8px;color:#8b949e;font-size:0.88rem;">
            <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>
            No logs found yet.
        </div>""", unsafe_allow_html=True)

    st.divider()
    st.markdown("""
    <div style="display:flex;align-items:center;gap:8px;padding:12px 16px;background:rgba(56,189,248,0.06);border:1px solid rgba(56,189,248,0.15);border-radius:8px;color:#8b949e;font-size:0.82rem;">
        <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>
        Data is stored in Supabase PostgreSQL. Use the Admin DB View tab to export records as CSV.
    </div>""", unsafe_allow_html=True)
# Always-visible tabs
tab_labels = [
    "📊 Dashboard",
    "🧾 Resume Builder",
    "💼 Job Search",
    "📚 Course Recommendation",
	"🛡️ Scam Detector"
]

# Add Admin tab only for admin user
if st.session_state.username == "admin":
    tab_labels.append("📁 Admin DB View")

# Create tabs dynamically
tabs = st.tabs(tab_labels)

# Unpack first five (always exist)
tab1, tab2, tab3, tab4, tab_scam = tabs[:5]

# Handle optional admin tab (index shifts to 5 now)
tab5 = tabs[5] if len(tabs) > 5 else None
with tab1:
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Orbitron', sans-serif;
        background-color: #0b0c10;
        color: #c5c6c7;
        scroll-behavior: smooth;
    }

    /* ---------- SCROLLBAR ---------- */
    ::-webkit-scrollbar { width: 8px; }
    ::-webkit-scrollbar-track { background: #1f2833; }
    ::-webkit-scrollbar-thumb { background: #00ffff; border-radius: 4px; }

    /* ---------- BANNER ---------- */
    .banner-container {
        width: 100%;
        height: 80px;
        background: linear-gradient(90deg, #000428, #004e92);
        border-bottom: 2px solid cyan;
        overflow: hidden;
        display: flex;
        align-items: center;
        justify-content: flex-start;
        position: relative;
        margin-bottom: 20px;
        border-radius: 12px;
        backdrop-filter: blur(14px);
    }
    .pulse-bar {
        position: absolute;
        display: flex;
        align-items: center;
        font-size: 22px;
        font-weight: bold;
        color: #00ffff;
        white-space: nowrap;
        animation: glideIn 12s linear infinite;
        text-shadow: 0 0 10px #00ffff;
    }
    .pulse-bar .bar {
        width: 10px;
        height: 30px;
        margin-right: 10px;
        background: #00ffff;
        box-shadow: 0 0 8px cyan;
        animation: pulse 1s ease-in-out infinite;
    }
    @keyframes glideIn {
        0% { left: -50%; opacity: 0; }
        10% { opacity: 1; }
        90% { opacity: 1; }
        100% { left: 110%; opacity: 0; }
    }
    @keyframes pulse {
        0%, 100% { height: 20px; background-color: #00ffff; }
        50% { height: 40px; background-color: #ff00ff; }
    }

    /* ---------- HEADER ---------- */
    .header {
        font-size: 28px;
        font-weight: bold;
        text-align: center;
        text-transform: uppercase;
        letter-spacing: 2px;
        padding: 20px 30px;  /* ✅ More spacing inside the bar */
        color: #00ffff;
        text-shadow: 0px 0px 10px #00ffff;
        position: relative;
        overflow: hidden;
        border-radius: 14px;
        background: rgba(10,20,40,0.35);
        backdrop-filter: blur(14px);
        border: 1px solid rgba(0,200,255,0.5);
        box-shadow: 0 0 12px rgba(0,200,255,0.25);
    }
    .header::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.18) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .header:hover::before { left: 100%; top: 100%; }

    /* ---------- SHIMMER (COMMON) ---------- */
    .shimmer::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .shimmer:hover::before { left: 100%; top: 100%; }

    /* ---------- FILE UPLOADER ---------- */
    .stFileUploader > div > div {
        border: 1px solid rgba(0,200,255,0.5);
        border-radius: 14px;
        background: rgba(10,20,40,0.35);
        backdrop-filter: blur(14px);
        color: #cce6ff;
        box-shadow: 0 0 12px rgba(0,200,255,0.3);
        position: relative;
        overflow: hidden;
    }
    .stFileUploader > div > div::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stFileUploader > div > div:hover::before { left: 100%; top: 100%; }

    /* ---------- BUTTONS ---------- */
    .stButton > button {
        position: relative;
        overflow: hidden;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        color: #e6f7ff;
        border-radius: 14px;
        padding: 10px 20px;
        font-size: 16px;
        font-weight: 500;
        text-transform: uppercase;
        backdrop-filter: blur(16px);
        box-shadow: 0 0 12px rgba(0,200,255,0.35),
                    inset 0 0 20px rgba(0,200,255,0.05);
        transition: all 0.3s ease-in-out;
    }
    .stButton > button::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stButton > button:hover::before { left: 100%; top: 100%; }

    /* ---------- INPUTS ---------- */
    .stTextInput > div > input,
    .stTextArea > div > textarea {
        position: relative;
        overflow: hidden;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        border-radius: 14px;
        color: #e6f7ff;
        padding: 10px;
        backdrop-filter: blur(16px);
        box-shadow: 0 0 12px rgba(0,200,255,0.3),
                    inset 0 0 15px rgba(0,200,255,0.05);
        transition: all 0.3s ease-in-out;
    }

    /* ---------- CHAT MESSAGES ---------- */
    .stChatMessage {
        position: relative;
        overflow: hidden;
        font-size: 18px;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.5);
        border-radius: 14px;
        padding: 14px;
        color: #e6f7ff;
        text-shadow: 0 0 6px rgba(0,200,255,0.7);
        box-shadow: 0 0 12px rgba(0,200,255,0.3),
                    inset 0 0 15px rgba(0,200,255,0.05);
    }
    .stChatMessage::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stChatMessage:hover::before { left: 100%; top: 100%; }

    /* ---------- METRICS ---------- */
    .stMetric {
        position: relative;
        overflow: hidden;
        background-color: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        border-radius: 14px;
        padding: 15px;
        box-shadow: 0 0 12px rgba(0,200,255,0.35),
                    inset 0 0 20px rgba(0,200,255,0.05);
        text-align: center;
    }
    .stMetric::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stMetric:hover::before { left: 100%; top: 100%; }

    /* ---------- MOBILE ---------- */
    @media (max-width: 768px) {
        .pulse-bar { font-size: 16px; }
        .header { font-size: 20px; }
    }
    </style>

    <!-- Banner -->
    <div class="banner-container">
        <div class="pulse-bar">
            <div class="bar"></div>
            <div>HIRELYZER - Elevate Your Resume Analysis</div>
        </div>
    </div>

    <!-- Header -->
    <div class="header">💼 HIRELYZER - AI BASED ETHICAL RESUME ANALYZER</div>
    """, unsafe_allow_html=True)

# Load environment variables
load_dotenv()

# Detect Device
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
torch.backends.cudnn.benchmark = True
working_dir = os.path.dirname(os.path.abspath(__file__))

# ------------------- Lazy Initialization -------------------
# (get_easyocr_reader and ensure_nltk are in resume_processor.py)

# ---------------- Sidebar Layout with Inline Images ----------------
st.sidebar.markdown(
    "<p style='font-size:0.72rem;font-weight:700;letter-spacing:0.10em;text-transform:uppercase;color:#4a5568;border-bottom:1px solid rgba(255,255,255,0.06);padding-bottom:8px;margin-bottom:12px;font-family:-apple-system,BlinkMacSystemFont,sans-serif;'>Job Configuration</p>",
    unsafe_allow_html=True
)

# ── Job Title options — one clean title per VALID_DOMAIN in db_manager ──────
_JOB_TITLE_OPTIONS = [
    "— Select Job Title —",
    # ── Core Software ──
    "Software Engineer",           # Software Engineering
    "Backend Developer",           # Backend Development
    "Frontend Developer",          # Frontend Development
    "Full Stack Developer",        # Full Stack Development
    "Mobile Developer",            # Mobile Development
    # ── Data & AI ──
    "Data Scientist",              # Data Science
    "ML Engineer",                 # AI/Machine Learning
    # ── Infrastructure ──
    "DevOps Engineer",             # DevOps/Infrastructure
    "Cloud Engineer",              # Cloud Engineering
    "Site Reliability Engineer",   # Site Reliability Engineering
    "Database Administrator",      # Database Management
    "Network Engineer",            # Networking
    "Embedded Systems Engineer",   # Embedded Systems
    "IoT Engineer",                # IoT Development
    # ── Design & Quality ──
    "UI/UX Designer",              # UI/UX Design
    "QA Engineer",                 # Quality Assurance
    # ── Architecture & Management ──
    "Solution Architect",          # System Architecture
    "Product Manager",             # Product Management
    "Project Manager",             # Project Management
    "Business Analyst",            # Business Analysis
    "Agile Coach / Scrum Master",  # Agile Coaching
    # ── Specialised Tech ──
    "Cybersecurity Analyst",       # Cybersecurity
    "Blockchain Developer",        # Blockchain Development
    "Game Developer",              # Game Development
    "AR/VR Developer",             # AR/VR Development
    # ── Non-Tech Domains ──
    "Digital Marketing Specialist",# Digital Marketing
    "Technical Writer",            # Technical Writing
    "Technical Sales Engineer",    # Technical Sales
    "E-commerce Specialist",       # E-commerce
    "Fintech Developer",           # Fintech
    "Healthcare Tech Specialist",  # Healthcare Tech
    "EdTech Specialist",           # EdTech
    # ── Other ──
    "Other (type below)",
]

# ── India top hiring cities — ordered by tech job market size ───────────────
_LOCATION_OPTIONS = [
    "— Select Location —",
    "Bangalore, India",
    "Hyderabad, India",
    "Mumbai, India",
    "Pune, India",
    "Chennai, India",
    "Delhi, India",
    "Noida, India",
    "Gurgaon, India",
    "Kolkata, India",
    "Ahmedabad, India",
    "Coimbatore, India",
    "Indore, India",
    "Jaipur, India",
    "Kochi, India",
    "Bhubaneswar, India",
    "Chandigarh, India",
    "Nagpur, India",
    "Thiruvananthapuram, India",
    "Remote — India",
    "Other (type below)",
]

# ── Standard JD templates — one per job title (auto-fill on selection) ──────
_JD_TEMPLATES = {
    "Software Engineer": """\
Position: Software Engineer
Location: [City, India] | Full-Time

About the Role:
We are looking for a passionate and skilled Software Engineer to join our growing engineering team. You will design, develop, and maintain high-quality software solutions that power our core products.

Key Responsibilities:
- Design, develop, test, and deploy scalable software applications
- Write clean, maintainable, and well-documented code
- Collaborate with cross-functional teams including product managers and designers
- Participate in code reviews and contribute to engineering best practices
- Identify and resolve performance bottlenecks and software defects
- Contribute to architectural decisions and technical roadmap discussions
- Mentor junior engineers and share knowledge across the team
- Integrate third-party APIs and services as needed
- Ensure software security and data protection standards are met
- Continuously improve development processes and tooling

Required Skills & Qualifications:
- B.E./B.Tech/M.Tech in Computer Science or related field
- 2–5 years of software development experience
- Proficiency in one or more languages: Python, Java, C++, Go, or JavaScript
- Strong understanding of data structures, algorithms, and design patterns
- Experience with version control (Git), CI/CD pipelines, and Agile methodologies
- Familiarity with relational and NoSQL databases
- Strong problem-solving and analytical skills

Good to Have:
- Experience with cloud platforms (AWS, GCP, Azure)
- Knowledge of containerization (Docker, Kubernetes)
- Open-source contributions
""",

    "Backend Developer": """\
Position: Backend Developer
Location: [City, India] | Full-Time

About the Role:
We are seeking an experienced Backend Developer to build and maintain robust, scalable server-side systems. You will work closely with frontend teams and DevOps to deliver reliable APIs and microservices.

Key Responsibilities:
- Design and develop high-performance RESTful and GraphQL APIs
- Build and maintain microservices architectures
- Optimize database queries and data models for performance and scalability
- Implement authentication, authorization, and security best practices
- Collaborate with frontend developers to define API contracts
- Write unit and integration tests to ensure code quality
- Monitor system performance and resolve production issues proactively
- Participate in system design discussions and technical planning
- Document APIs using tools like Swagger/OpenAPI
- Conduct code reviews and enforce coding standards

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 2–5 years of backend development experience
- Proficiency in Python, Node.js, Java, or Go
- Strong experience with REST API design and development
- Hands-on with databases: PostgreSQL, MySQL, MongoDB, or Redis
- Experience with message queues (RabbitMQ, Kafka)
- Solid understanding of HTTP, caching, and distributed systems

Good to Have:
- Experience with AWS Lambda, EC2, or GCP services
- Knowledge of gRPC and Protobuf
- Familiarity with containerization (Docker, Kubernetes)
""",

    "Frontend Developer": """\
Position: Frontend Developer
Location: [City, India] | Full-Time

About the Role:
We are looking for a creative and detail-oriented Frontend Developer to craft exceptional user interfaces. You will translate UI/UX designs into pixel-perfect, performant, and accessible web experiences.

Key Responsibilities:
- Build responsive, cross-browser-compatible web applications using modern frameworks
- Translate Figma/Adobe XD designs into high-quality HTML, CSS, and JavaScript
- Optimize frontend performance for speed, accessibility, and SEO
- Integrate RESTful APIs and manage application state effectively
- Write unit and end-to-end tests for frontend components
- Collaborate closely with designers, backend developers, and product managers
- Maintain and improve existing frontend codebases
- Implement animations, transitions, and interactive UI components
- Participate in design reviews and provide technical feasibility feedback
- Follow accessibility (WCAG) and web standards best practices

Required Skills & Qualifications:
- B.E./B.Tech or equivalent in Computer Science or related field
- 2–5 years of frontend development experience
- Proficiency in React.js, Vue.js, or Angular
- Strong command of HTML5, CSS3, and modern JavaScript (ES6+)
- Experience with state management (Redux, Zustand, Vuex)
- Familiarity with build tools: Webpack, Vite, or Parcel
- Understanding of browser rendering, performance optimization, and caching

Good to Have:
- Experience with TypeScript
- Knowledge of Next.js or Nuxt.js
- Familiarity with testing frameworks (Jest, Cypress, Playwright)
""",

    "Full Stack Developer": """\
Position: Full Stack Developer
Location: [City, India] | Full-Time

About the Role:
We are hiring a Full Stack Developer who is comfortable working across the entire web application stack. You will own features end-to-end — from database design to pixel-perfect UI delivery.

Key Responsibilities:
- Develop and maintain both frontend and backend components of web applications
- Design and implement RESTful or GraphQL APIs consumed by web and mobile clients
- Build responsive UIs using modern JavaScript frameworks (React, Vue, Angular)
- Design and manage relational and NoSQL databases
- Write automated tests (unit, integration, E2E) across the stack
- Deploy and manage applications on cloud infrastructure
- Collaborate with product managers, designers, and QA engineers
- Conduct code reviews and mentor junior developers
- Identify and resolve performance, security, and scalability issues
- Contribute to CI/CD pipeline setup and maintenance

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 3–6 years of full stack development experience
- Proficiency in React.js/Vue.js (frontend) and Node.js/Python/Java (backend)
- Strong knowledge of SQL and NoSQL databases
- Experience with Git, Docker, and cloud deployment (AWS/GCP/Azure)
- Understanding of Agile/Scrum methodologies

Good to Have:
- Experience with TypeScript and serverless architectures
- Knowledge of microservices and event-driven systems
- Mobile development experience (React Native / Flutter)
""",

    "Mobile Developer": """\
Position: Mobile Developer
Location: [City, India] | Full-Time

About the Role:
We are looking for a talented Mobile Developer to build high-quality iOS and/or Android applications. You will collaborate with designers and backend engineers to deliver seamless mobile experiences to millions of users.

Key Responsibilities:
- Design and develop mobile applications for Android and/or iOS platforms
- Write clean, maintainable, and testable code following platform best practices
- Integrate RESTful APIs, push notifications, and third-party SDKs
- Optimize app performance, memory usage, and battery consumption
- Implement complex UI/UX designs with smooth animations
- Debug and resolve production issues reported by users
- Publish and maintain apps on Google Play Store and Apple App Store
- Collaborate with backend teams to define and consume APIs
- Write unit and UI tests using platform-specific testing frameworks
- Stay updated with the latest mobile platform changes and best practices

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 2–5 years of mobile app development experience
- Proficiency in Kotlin/Java (Android) or Swift/Objective-C (iOS), or Flutter/React Native
- Experience with REST API integration and async programming
- Familiarity with mobile CI/CD tools (Fastlane, Bitrise, Firebase App Distribution)
- Understanding of MVVM, MVP, or Clean Architecture patterns

Good to Have:
- Cross-platform experience with Flutter or React Native
- Experience with Firebase, Google Maps SDK, or payment gateway SDKs
- Published apps on Play Store or App Store
""",

    "Data Scientist": """\
Position: Data Scientist
Location: [City, India] | Full-Time

About the Role:
We are seeking an analytical and curious Data Scientist to turn raw data into actionable insights and predictive models. You will work with large, complex datasets and collaborate with engineering and business teams to drive data-driven decision-making.

Key Responsibilities:
- Collect, clean, and preprocess large structured and unstructured datasets
- Develop and validate statistical models and machine learning algorithms
- Build predictive and prescriptive models to solve business problems
- Perform exploratory data analysis (EDA) and communicate findings clearly
- Design and run A/B experiments to measure product impact
- Deploy models into production in collaboration with ML engineers
- Create dashboards and reports using BI tools (Tableau, Power BI, Metabase)
- Work closely with product managers and business stakeholders
- Document methodology, results, and model assumptions rigorously
- Monitor model performance and retrain as data distributions shift

Required Skills & Qualifications:
- B.Tech/M.Tech/M.Sc in Computer Science, Statistics, Mathematics, or related field
- 2–5 years of data science experience
- Proficiency in Python (Pandas, NumPy, Scikit-learn, Matplotlib, Seaborn)
- Strong understanding of statistical concepts: regression, classification, clustering, hypothesis testing
- Experience with SQL for data extraction and transformation
- Familiarity with ML frameworks: XGBoost, LightGBM, TensorFlow, or PyTorch

Good to Have:
- Experience with cloud ML platforms (AWS SageMaker, GCP Vertex AI)
- Knowledge of NLP, time-series forecasting, or recommendation systems
- Familiarity with MLflow or DVC for experiment tracking
""",

    "ML Engineer": """\
Position: ML Engineer (AI/Machine Learning)
Location: [City, India] | Full-Time

About the Role:
We are hiring an ML Engineer to bridge the gap between data science and production engineering. You will build, deploy, and scale machine learning systems that power intelligent features across our platform.

Key Responsibilities:
- Design, train, evaluate, and deploy machine learning models at scale
- Build ML pipelines for data ingestion, feature engineering, model training, and inference
- Optimize models for latency, throughput, and resource efficiency
- Implement MLOps practices: experiment tracking, model versioning, monitoring, and retraining
- Collaborate with data scientists to productionize research models
- Build and maintain feature stores and data pipelines
- Develop APIs and microservices to serve ML model predictions
- Monitor model drift and implement automated retraining workflows
- Work with GPU/TPU infrastructure for large-scale model training
- Research and apply state-of-the-art ML techniques to business problems

Required Skills & Qualifications:
- B.Tech/M.Tech in Computer Science, AI, or related field
- 3–6 years of ML engineering or related experience
- Strong proficiency in Python and ML frameworks (TensorFlow, PyTorch, Scikit-learn)
- Experience with ML pipeline orchestration (Airflow, Kubeflow, MLflow)
- Solid software engineering skills: APIs, testing, CI/CD
- Experience with cloud platforms for ML (AWS SageMaker, GCP Vertex AI, Azure ML)

Good to Have:
- Experience with LLMs, RAG pipelines, or generative AI
- Knowledge of ONNX, TensorRT, or model quantization techniques
- Familiarity with Spark or Dask for distributed data processing
""",

    "DevOps Engineer": """\
Position: DevOps Engineer
Location: [City, India] | Full-Time

About the Role:
We are looking for a skilled DevOps Engineer to streamline our development lifecycle, automate infrastructure, and ensure the reliability and scalability of our production systems.

Key Responsibilities:
- Design, implement, and maintain CI/CD pipelines for automated build, test, and deployment
- Manage and provision cloud infrastructure using IaC tools (Terraform, CloudFormation)
- Monitor system health, performance, and availability using observability tools
- Implement and maintain containerized deployments using Docker and Kubernetes
- Manage cloud environments on AWS, GCP, or Azure
- Collaborate with development teams to improve deployment frequency and reliability
- Implement security best practices across infrastructure and pipelines
- Automate repetitive operational tasks using scripting (Bash, Python)
- Manage secrets, certificates, and access control policies
- Respond to and resolve production incidents, conducting root cause analysis

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 3–6 years of DevOps or infrastructure engineering experience
- Strong experience with CI/CD tools: Jenkins, GitHub Actions, GitLab CI, or CircleCI
- Proficiency with Docker, Kubernetes, and Helm
- Experience with cloud platforms (AWS, GCP, or Azure)
- Scripting skills in Bash and Python
- Knowledge of monitoring tools: Prometheus, Grafana, ELK Stack, or Datadog

Good to Have:
- Experience with service mesh (Istio, Linkerd)
- Knowledge of GitOps workflows (ArgoCD, Flux)
- Cloud certifications (AWS Solutions Architect, CKA, etc.)
""",

    "Cloud Engineer": """\
Position: Cloud Engineer
Location: [City, India] | Full-Time

About the Role:
We are seeking an experienced Cloud Engineer to design, build, and manage our cloud infrastructure. You will ensure our systems are secure, scalable, cost-optimized, and highly available across multi-cloud environments.

Key Responsibilities:
- Architect and implement scalable, resilient cloud infrastructure on AWS, GCP, or Azure
- Automate infrastructure provisioning using Terraform, Pulumi, or CloudFormation
- Manage cloud networking: VPCs, subnets, load balancers, DNS, and security groups
- Implement cloud security best practices including IAM, encryption, and compliance controls
- Optimize cloud costs through right-sizing, reserved instances, and spot usage
- Set up and manage monitoring, alerting, and logging pipelines
- Support cloud migration projects from on-premises to cloud environments
- Collaborate with DevOps, security, and application teams
- Develop and maintain disaster recovery and business continuity plans
- Evaluate and adopt new cloud services and technologies

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 3–6 years of cloud engineering experience
- Deep expertise in at least one major cloud platform (AWS, GCP, Azure)
- Strong experience with Infrastructure as Code (Terraform, Ansible)
- Knowledge of cloud networking, security, and identity management
- Familiarity with containerization (Docker, Kubernetes, EKS/GKE/AKS)

Good to Have:
- Cloud certifications (AWS Solutions Architect Pro, GCP Professional, Azure Expert)
- Experience with multi-cloud and hybrid cloud architectures
- Knowledge of FinOps practices and cost optimization tools
""",

    "Site Reliability Engineer": """\
Position: Site Reliability Engineer (SRE)
Location: [City, India] | Full-Time

About the Role:
We are hiring a Site Reliability Engineer to ensure the reliability, scalability, and performance of our production systems. You will apply software engineering principles to infrastructure and operations problems.

Key Responsibilities:
- Define and maintain SLOs, SLIs, and error budgets for critical services
- Build and maintain observability systems: metrics, logs, traces (Prometheus, Grafana, Jaeger)
- Lead incident response, post-mortems, and systematic reliability improvements
- Automate toil and manual operational work through engineering solutions
- Collaborate with development teams to improve software reliability and deployability
- Capacity planning and performance testing for high-traffic systems
- Design fault-tolerant, self-healing distributed systems
- Implement and improve on-call rotations and escalation policies
- Drive chaos engineering practices to proactively uncover weaknesses
- Review system designs for reliability, scalability, and maintainability

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 4–7 years of SRE, DevOps, or systems engineering experience
- Strong software engineering skills (Python, Go, or Java)
- Deep experience with Linux systems administration
- Proficiency with Kubernetes, Docker, and cloud platforms (AWS/GCP/Azure)
- Experience with distributed systems concepts and microservices architectures

Good to Have:
- Experience with chaos engineering tools (Chaos Monkey, Gremlin)
- Knowledge of service meshes (Istio, Linkerd)
- Contributions to open-source reliability or observability tooling
""",

    "Database Administrator": """\
Position: Database Administrator (DBA)
Location: [City, India] | Full-Time

About the Role:
We are looking for an experienced Database Administrator to manage, optimize, and secure our databases. You will ensure high availability, performance, and integrity of mission-critical data systems.

Key Responsibilities:
- Install, configure, and maintain relational and NoSQL database systems
- Monitor database performance and proactively tune queries and indexes
- Design and implement backup, recovery, and disaster recovery strategies
- Manage database security: user roles, permissions, encryption, and auditing
- Plan and execute database migrations and schema changes with zero downtime
- Capacity planning and storage management for growing data volumes
- Set up and maintain database replication and high-availability clusters
- Collaborate with developers to design optimal data models and queries
- Troubleshoot and resolve database performance and connectivity issues
- Document database architecture, procedures, and runbooks

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 3–6 years of DBA experience
- Strong expertise in PostgreSQL, MySQL, or Oracle
- Experience with NoSQL databases: MongoDB, Cassandra, or DynamoDB
- Proficiency in SQL query optimization and execution plan analysis
- Knowledge of replication, clustering, and failover strategies
- Familiarity with backup tools and point-in-time recovery

Good to Have:
- Experience with cloud-managed databases (RDS, Cloud SQL, Cosmos DB)
- Knowledge of database monitoring tools (pgBadger, Percona Monitoring)
- Experience with data warehousing (Redshift, BigQuery, Snowflake)
""",

    "Network Engineer": """\
Position: Network Engineer
Location: [City, India] | Full-Time

About the Role:
We are seeking a skilled Network Engineer to design, implement, and maintain our network infrastructure. You will ensure high availability, security, and performance of our on-premises and cloud network environments.

Key Responsibilities:
- Design and implement LAN, WAN, SD-WAN, and cloud networking solutions
- Configure and manage routers, switches, firewalls, and load balancers
- Monitor network performance and proactively resolve latency and connectivity issues
- Implement network security policies, VPNs, and zero-trust architectures
- Manage DNS, DHCP, and IP address management (IPAM)
- Plan and execute network upgrades, migrations, and expansions
- Collaborate with cloud teams on VPC design, peering, and direct connect setups
- Respond to network incidents and conduct root cause analysis
- Develop and maintain network documentation and topology diagrams
- Evaluate and recommend new networking technologies and vendors

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science, Electronics, or equivalent
- 3–6 years of network engineering experience
- Strong knowledge of TCP/IP, BGP, OSPF, MPLS, and VLANs
- Experience with Cisco, Juniper, Palo Alto, or Fortinet devices
- Familiarity with cloud networking (AWS VPC, Azure VNet, GCP VPC)
- Understanding of network security: firewalls, IDS/IPS, and DDoS mitigation

Good to Have:
- Certifications: CCNA, CCNP, CCIE, or equivalent
- Experience with SD-WAN solutions (Cisco Viptela, VMware VeloCloud)
- Knowledge of network automation using Python or Ansible
""",

    "Embedded Systems Engineer": """\
Position: Embedded Systems Engineer
Location: [City, India] | Full-Time

About the Role:
We are hiring an Embedded Systems Engineer to develop firmware and software for embedded hardware platforms. You will work closely with hardware engineers to bring intelligent devices to life.

Key Responsibilities:
- Develop, test, and debug firmware for microcontrollers and embedded processors
- Write low-level C/C++ code for bare-metal and RTOS-based systems
- Develop and integrate device drivers for peripherals (UART, SPI, I2C, CAN, USB)
- Collaborate with hardware engineers on board bring-up and hardware validation
- Optimize firmware for real-time performance, power consumption, and memory usage
- Implement communication protocols for IoT and industrial systems
- Write automated test frameworks and hardware-in-the-loop (HIL) test cases
- Review hardware schematics and PCB designs from a firmware perspective
- Maintain and improve existing embedded codebases
- Document firmware architecture, APIs, and porting guides

Required Skills & Qualifications:
- B.E./B.Tech in Electronics, Computer Science, or Electrical Engineering
- 3–6 years of embedded systems development experience
- Strong proficiency in C and C++ for embedded targets
- Experience with MCU families: STM32, ESP32, NXP, TI, or AVR
- Knowledge of RTOS: FreeRTOS, Zephyr, or ThreadX
- Hands-on with JTAG/SWD debugging using tools like OpenOCD or Segger J-Link
- Understanding of communication protocols: UART, SPI, I2C, CAN, USB

Good to Have:
- Experience with Linux-based embedded systems (Yocto, Buildroot)
- Knowledge of functional safety standards (IEC 61508, ISO 26262)
- Familiarity with wireless protocols (BLE, Zigbee, LoRa, Wi-Fi)
""",

    "IoT Engineer": """\
Position: IoT Engineer
Location: [City, India] | Full-Time

About the Role:
We are looking for an innovative IoT Engineer to design and build end-to-end Internet of Things solutions. You will work across firmware, connectivity, cloud integration, and data pipelines to deliver smart, connected products.

Key Responsibilities:
- Design and develop IoT solutions spanning devices, gateways, and cloud backends
- Develop firmware for IoT devices using C/C++ or MicroPython on embedded platforms
- Implement wireless communication protocols: MQTT, CoAP, HTTP, BLE, Zigbee, LoRa
- Integrate IoT devices with cloud IoT platforms (AWS IoT Core, Google Cloud IoT, Azure IoT Hub)
- Build data ingestion pipelines for real-time telemetry and event processing
- Develop device management features: OTA updates, remote config, diagnostics
- Ensure device and data security through encryption, secure boot, and certificate management
- Collaborate with hardware, cloud, and product teams for end-to-end integration
- Perform field testing, debugging, and reliability validation of IoT deployments
- Document hardware interfaces, APIs, and integration guides

Required Skills & Qualifications:
- B.E./B.Tech in Electronics, Computer Science, or related field
- 3–6 years of IoT development experience
- Proficiency in C/C++ (firmware) and Python (backend/scripting)
- Experience with MQTT, REST APIs, and IoT communication protocols
- Knowledge of cloud IoT platforms and serverless computing
- Familiarity with edge computing concepts and gateways

Good to Have:
- Experience with TinyML or edge AI inference on microcontrollers
- Knowledge of industrial IoT protocols (Modbus, OPC-UA)
- Hands-on with Raspberry Pi, ESP32, or similar platforms
""",

    "UI/UX Designer": """\
Position: UI/UX Designer
Location: [City, India] | Full-Time

About the Role:
We are seeking a talented UI/UX Designer to create intuitive and visually compelling digital experiences. You will work closely with product managers and engineers to design products that delight users and achieve business goals.

Key Responsibilities:
- Lead end-to-end design processes: research, ideation, wireframing, prototyping, and delivery
- Conduct user research, usability testing, and competitive analysis to inform design decisions
- Create user personas, journey maps, information architecture, and user flows
- Design high-fidelity mockups and interactive prototypes using Figma or Adobe XD
- Develop and maintain a consistent design system and component library
- Collaborate closely with frontend developers to ensure pixel-perfect implementation
- Gather and incorporate feedback from stakeholders and end-users iteratively
- Ensure designs meet accessibility standards (WCAG 2.1)
- Present design concepts and rationale to cross-functional teams and leadership
- Stay current with design trends, tools, and emerging interaction patterns

Required Skills & Qualifications:
- Bachelor's degree in Design, HCI, or equivalent field
- 2–5 years of UI/UX design experience for web and/or mobile products
- Proficiency in Figma, Sketch, or Adobe XD
- Strong portfolio demonstrating end-to-end design process and shipped products
- Solid understanding of design principles: typography, color, layout, and accessibility
- Experience conducting user interviews and usability tests

Good to Have:
- Experience with motion design and micro-interactions (After Effects, Principle)
- Knowledge of HTML/CSS for better developer collaboration
- Familiarity with product analytics tools (Mixpanel, Hotjar, FullStory)
""",

    "QA Engineer": """\
Position: QA Engineer
Location: [City, India] | Full-Time

About the Role:
We are looking for a detail-oriented QA Engineer to ensure the quality of our software products. You will design and execute test strategies, build automation frameworks, and act as the last line of defence before features reach our users.

Key Responsibilities:
- Design, develop, and maintain automated test suites for web, mobile, and API testing
- Define and execute test plans, test cases, and regression test suites
- Perform functional, integration, performance, and exploratory testing
- Work closely with developers during sprint planning and development to shift quality left
- Identify, document, and track bugs through the full defect lifecycle
- Set up and maintain CI/CD integrated test pipelines
- Perform API testing using tools like Postman, REST-Assured, or Karate
- Conduct load and stress testing using JMeter or Locust
- Review requirements and user stories to identify gaps and edge cases
- Mentor junior QA engineers and promote quality culture across the team

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 2–5 years of QA engineering experience
- Proficiency in test automation frameworks: Selenium, Playwright, Cypress, or Appium
- Strong knowledge of testing methodologies and SDLC
- Experience with API testing and performance testing tools
- Scripting ability in Python, Java, or JavaScript for test automation
- Familiarity with defect tracking tools (JIRA, Bugzilla)

Good to Have:
- ISTQB certification
- Experience with BDD frameworks (Cucumber, Behave)
- Knowledge of mobile testing for iOS and Android
""",

    "Solution Architect": """\
Position: Solution Architect
Location: [City, India] | Full-Time

About the Role:
We are hiring a Solution Architect to lead the design of scalable, secure, and cost-effective technology solutions. You will bridge business requirements with technical implementation, guiding engineering teams and influencing technology strategy.

Key Responsibilities:
- Design end-to-end technical architectures for complex enterprise systems
- Evaluate and select appropriate technologies, frameworks, and platforms
- Create architecture blueprints, technical specifications, and design documents
- Collaborate with product, engineering, and business stakeholders to align solutions with goals
- Lead technical proof-of-concept projects to validate architectural decisions
- Define integration patterns, API contracts, and data flow architectures
- Ensure solutions meet non-functional requirements: scalability, security, availability, and performance
- Provide technical leadership and mentorship to development teams
- Conduct architecture reviews and code quality assessments
- Stay current with emerging technologies and industry best practices

Required Skills & Qualifications:
- B.E./B.Tech/M.Tech in Computer Science or equivalent
- 8+ years of software development with 3+ years in architecture or senior technical roles
- Deep expertise in distributed systems, microservices, and cloud-native architectures
- Strong knowledge of cloud platforms (AWS, GCP, Azure) and enterprise integration patterns
- Experience designing APIs (REST, GraphQL, gRPC) and event-driven systems
- Strong communication skills to present technical concepts to non-technical audiences

Good to Have:
- Cloud architect certifications (AWS Solutions Architect Professional, GCP Professional Architect)
- Experience with TOGAF or other enterprise architecture frameworks
- Background in pre-sales and RFP/RFI responses
""",

    "Product Manager": """\
Position: Product Manager
Location: [City, India] | Full-Time

About the Role:
We are looking for a strategic and execution-focused Product Manager to own the product roadmap and drive the delivery of features that create value for our users and business.

Key Responsibilities:
- Define product vision, strategy, and roadmap aligned with company goals
- Gather and prioritize requirements from customers, stakeholders, and data insights
- Write detailed product requirement documents (PRDs), user stories, and acceptance criteria
- Collaborate with engineering, design, data, and business teams throughout the product lifecycle
- Define and track key product metrics and KPIs to measure success
- Conduct user research, interviews, and usability sessions to deeply understand user needs
- Prioritize backlog and manage sprint planning with engineering teams
- Communicate product updates, decisions, and trade-offs to leadership and stakeholders
- Analyse competitor products and market trends to identify opportunities
- Drive go-to-market strategy for new feature launches

Required Skills & Qualifications:
- Bachelor's/MBA degree in Business, Computer Science, or related field
- 3–6 years of product management experience in a tech company
- Strong analytical skills with experience using data tools (SQL, Amplitude, Mixpanel, GA)
- Experience with Agile/Scrum methodologies and project management tools (JIRA, Confluence)
- Excellent communication and stakeholder management skills
- Ability to make data-driven decisions and manage ambiguity

Good to Have:
- Technical background or experience working closely with engineering teams
- Experience with A/B experimentation and growth product management
- Knowledge of B2B SaaS, fintech, or consumer internet domains
""",

    "Project Manager": """\
Position: Project Manager
Location: [City, India] | Full-Time

About the Role:
We are seeking an experienced Project Manager to plan, execute, and deliver technology projects on time and within budget. You will coordinate across multiple teams and ensure seamless project delivery.

Key Responsibilities:
- Plan, schedule, and manage end-to-end delivery of technology projects
- Define project scope, goals, deliverables, timelines, and resource requirements
- Create and maintain detailed project plans using project management tools
- Facilitate daily stand-ups, sprint reviews, retrospectives, and stakeholder meetings
- Identify, assess, and mitigate project risks and dependencies proactively
- Manage project budgets and resource allocation
- Track and report project status, milestones, and KPIs to leadership
- Coordinate across engineering, QA, design, business, and vendor teams
- Manage change requests and scope creep effectively
- Ensure projects meet quality, compliance, and delivery standards

Required Skills & Qualifications:
- Bachelor's degree in Business, Computer Science, or equivalent
- 4–8 years of project management experience in IT or software delivery
- PMP, PRINCE2, or equivalent certification preferred
- Proficiency with project management tools: MS Project, JIRA, Asana, or Monday.com
- Strong knowledge of Agile, Scrum, and Waterfall methodologies
- Excellent stakeholder management and communication skills

Good to Have:
- Experience managing cross-border or distributed teams
- Technical background in software development or DevOps
- Familiarity with program-level planning and portfolio management
""",

    "Business Analyst": """\
Position: Business Analyst
Location: [City, India] | Full-Time

About the Role:
We are looking for an insightful Business Analyst to bridge the gap between business needs and technology solutions. You will gather requirements, analyse data, and ensure technology projects deliver real business value.

Key Responsibilities:
- Elicit, analyse, and document business requirements from stakeholders
- Translate business needs into functional specifications, use cases, and user stories
- Conduct gap analysis between current and desired business processes
- Facilitate workshops and stakeholder meetings to gather and validate requirements
- Create process flow diagrams, data flow diagrams, and wireframes
- Collaborate with development teams to ensure requirements are correctly understood and implemented
- Define and execute user acceptance testing (UAT) plans
- Analyse business data to identify trends, patterns, and improvement opportunities
- Track and manage requirement changes throughout the project lifecycle
- Prepare business case documents and ROI analysis for new initiatives

Required Skills & Qualifications:
- Bachelor's degree in Business, Computer Science, or related field
- 3–5 years of business analysis experience in IT projects
- Proficiency with tools: JIRA, Confluence, Visio, Lucidchart, or similar
- Strong SQL skills for data analysis and reporting
- Experience with Agile and Waterfall project methodologies
- Excellent communication, documentation, and facilitation skills

Good to Have:
- CBAP certification or equivalent
- Experience in BFSI, e-commerce, or healthcare domains
- Familiarity with BI tools (Tableau, Power BI, Looker)
""",

    "Agile Coach / Scrum Master": """\
Position: Agile Coach / Scrum Master
Location: [City, India] | Full-Time

About the Role:
We are hiring an experienced Agile Coach / Scrum Master to embed agile culture and practices across our engineering and product teams. You will coach teams to self-organize, continuously improve, and deliver value iteratively.

Key Responsibilities:
- Facilitate Scrum ceremonies: sprint planning, daily stand-ups, sprint reviews, and retrospectives
- Coach teams on Agile principles, Scrum framework, Kanban, and SAFe practices
- Remove impediments and shield teams from external disruptions to maintain flow
- Track and report team velocity, sprint burn-down, and delivery metrics
- Work with product managers to maintain a healthy and prioritised product backlog
- Identify and help resolve team dysfunctions and improve collaboration
- Drive continuous improvement through retrospectives and process experiments
- Support organizational agile transformation initiatives
- Mentor and coach other Scrum Masters and team leads
- Facilitate cross-team dependency management and PI planning (SAFe context)

Required Skills & Qualifications:
- Bachelor's degree in any field; technical background preferred
- 4–8 years of experience as a Scrum Master or Agile Coach
- CSM, PSM, or SAFe certification required; CSP or ICP-ACC preferred
- Deep knowledge of Agile frameworks: Scrum, Kanban, XP, and SAFe
- Experience with Agile tools: JIRA, Azure DevOps, or Rally
- Strong facilitation, coaching, and conflict resolution skills

Good to Have:
- Experience with large-scale agile transformations
- Technical background in software development or QA
- Knowledge of DevOps and Lean practices
""",

    "Cybersecurity Analyst": """\
Position: Cybersecurity Analyst
Location: [City, India] | Full-Time

About the Role:
We are seeking a vigilant and skilled Cybersecurity Analyst to protect our systems, networks, and data from cyber threats. You will monitor, detect, investigate, and respond to security incidents while strengthening our overall security posture.

Key Responsibilities:
- Monitor security events and alerts using SIEM tools (Splunk, IBM QRadar, Microsoft Sentinel)
- Investigate and respond to security incidents, conducting thorough forensic analysis
- Perform vulnerability assessments and penetration testing on systems and applications
- Manage and maintain security tools: firewalls, IDS/IPS, EDR, DLP, and WAF
- Conduct threat intelligence analysis and proactive threat hunting
- Implement and enforce security policies, standards, and compliance requirements
- Perform security audits and risk assessments across infrastructure and applications
- Support security awareness training and phishing simulation programs
- Collaborate with IT and development teams to implement security best practices
- Prepare incident reports and security metrics for management

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science, Information Security, or equivalent
- 2–5 years of cybersecurity experience
- Strong knowledge of networking, TCP/IP, firewalls, and intrusion detection
- Experience with SIEM platforms, log analysis, and threat detection
- Familiarity with security frameworks: NIST CSF, ISO 27001, OWASP
- Knowledge of regulatory compliance: GDPR, PCI-DSS, HIPAA, or IT Act

Good to Have:
- Certifications: CEH, CompTIA Security+, CISSP, or OSCP
- Experience with cloud security (AWS Security Hub, Azure Defender, GCP Security Command Center)
- Knowledge of malware analysis and reverse engineering
""",

    "Blockchain Developer": """\
Position: Blockchain Developer
Location: [City, India] | Full-Time

About the Role:
We are looking for a skilled Blockchain Developer to design and build decentralized applications and smart contracts. You will work on cutting-edge blockchain infrastructure powering fintech, supply chain, or Web3 products.

Key Responsibilities:
- Design, develop, and deploy smart contracts on Ethereum, Polygon, Solana, or Hyperledger
- Build decentralized applications (dApps) with Web3 frontend integration
- Develop and audit smart contracts for correctness, gas optimization, and security
- Integrate blockchain backends with traditional web APIs and data systems
- Implement token standards: ERC-20, ERC-721, ERC-1155
- Build and manage nodes, validators, and private blockchain networks
- Implement cryptographic protocols, consensus mechanisms, and wallet integrations
- Conduct smart contract security audits and vulnerability assessments
- Collaborate with product and frontend teams on blockchain-powered features
- Stay current with emerging blockchain protocols, DeFi, and NFT ecosystems

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 2–5 years of blockchain development experience
- Strong proficiency in Solidity, Rust, or Go for smart contract development
- Experience with Web3.js, Ethers.js, or Hardhat/Truffle development environments
- Solid understanding of blockchain fundamentals: cryptography, consensus, and distributed ledgers
- Experience with IPFS, The Graph, or Chainlink for decentralized infrastructure

Good to Have:
- Experience with Layer 2 solutions (Polygon, Optimism, Arbitrum)
- Knowledge of DeFi protocols (Uniswap, Aave, Compound)
- Smart contract audit experience or familiarity with tools like Slither/MythX
""",

    "Game Developer": """\
Position: Game Developer
Location: [City, India] | Full-Time

About the Role:
We are seeking a creative and technically strong Game Developer to design and build engaging gaming experiences. You will work on gameplay mechanics, graphics, physics, and performance to deliver polished games across platforms.

Key Responsibilities:
- Design and implement gameplay mechanics, systems, and features
- Develop high-quality game code using Unity (C#) or Unreal Engine (C++)
- Implement physics simulations, AI behaviours, and animation systems
- Optimize game performance for target platforms: mobile, PC, or console
- Build multiplayer networking systems and real-time game state synchronization
- Integrate monetization systems: in-app purchases, ads, and reward systems
- Collaborate with artists, designers, and sound engineers to integrate assets
- Debug and profile games to resolve performance, memory, and stability issues
- Implement analytics and crash reporting for live game monitoring
- Contribute to game design discussions and player experience improvements

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent; game development diploma a plus
- 2–5 years of professional game development experience
- Proficiency in Unity (C#) or Unreal Engine (C++)
- Strong understanding of 3D math, physics, and rendering pipelines
- Experience shipping at least one commercial game on mobile, PC, or console
- Knowledge of platform-specific SDKs (iOS, Android, Steam, PlayStation)

Good to Have:
- Experience with multiplayer game development (Photon, Mirror, or custom solutions)
- Knowledge of shader programming (HLSL, GLSL)
- Familiarity with procedural generation or ML-based game AI
""",

    "AR/VR Developer": """\
Position: AR/VR Developer
Location: [City, India] | Full-Time

About the Role:
We are hiring an AR/VR Developer to build immersive augmented and virtual reality experiences. You will push the boundaries of spatial computing to create next-generation applications for training, entertainment, and enterprise use cases.

Key Responsibilities:
- Develop immersive AR/VR applications using Unity (C#) or Unreal Engine (C++)
- Implement spatial interactions, hand tracking, gaze input, and 6DoF controller support
- Build AR experiences using ARKit (iOS), ARCore (Android), or HoloLens SDK
- Optimize rendering pipelines for high frame rates and low latency on XR hardware
- Implement 3D spatial audio and haptic feedback systems
- Integrate real-world sensor data (LiDAR, depth cameras) into AR applications
- Collaborate with designers and 3D artists to create compelling spatial UX
- Test and debug across various XR devices: Meta Quest, HoloLens, Magic Leap, iOS/Android
- Research and apply the latest spatial computing technologies and standards (OpenXR)
- Document technical architecture and integration guides

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 2–5 years of AR/VR development experience
- Proficiency in Unity or Unreal Engine for XR development
- Experience with AR frameworks: ARKit, ARCore, Vuforia, or Windows Mixed Reality
- Strong 3D math skills: vectors, quaternions, transformation matrices
- Understanding of XR performance constraints and optimization strategies

Good to Have:
- Experience with WebXR for browser-based AR/VR
- Knowledge of spatial mapping and SLAM algorithms
- Familiarity with OpenXR standard and multi-device deployment
""",

    "Digital Marketing Specialist": """\
Position: Digital Marketing Specialist
Location: [City, India] | Full-Time

About the Role:
We are looking for a data-driven Digital Marketing Specialist to plan, execute, and optimize our online marketing campaigns. You will drive user acquisition, engagement, and retention across multiple digital channels.

Key Responsibilities:
- Plan and execute multi-channel digital marketing campaigns: SEO, SEM, social media, email, and content
- Manage and optimize paid advertising campaigns on Google Ads, Meta Ads, LinkedIn Ads, and programmatic platforms
- Conduct keyword research, on-page and off-page SEO optimization
- Develop content calendars and oversee content creation for blogs, social media, and email newsletters
- Analyse campaign performance using Google Analytics 4, Search Console, and marketing dashboards
- Run A/B tests on ad creatives, landing pages, and email subject lines
- Manage marketing automation workflows using HubSpot, Mailchimp, or similar tools
- Collaborate with the product and sales teams on lead generation and conversion optimization
- Track and report marketing KPIs: CAC, ROAS, CTR, conversion rates, and MQL pipeline
- Monitor competitor digital presence and identify market opportunities

Required Skills & Qualifications:
- Bachelor's degree in Marketing, Business, or Communications
- 2–5 years of digital marketing experience
- Proficiency in Google Ads, Meta Ads Manager, and LinkedIn Campaign Manager
- Strong analytical skills with Google Analytics, Data Studio, or similar BI tools
- Knowledge of SEO best practices, tools (Ahrefs, SEMrush, Moz), and CRO techniques
- Experience with email marketing and marketing automation platforms

Good to Have:
- Google Ads and Meta Blueprint certifications
- Experience with video marketing (YouTube, Instagram Reels)
- Knowledge of growth hacking and performance marketing for SaaS or e-commerce
""",

    "Technical Writer": """\
Position: Technical Writer
Location: [City, India] | Full-Time

About the Role:
We are seeking an experienced Technical Writer to produce clear, accurate, and user-friendly technical documentation. You will transform complex technical information into content that developers, end-users, and stakeholders can easily understand.

Key Responsibilities:
- Create and maintain developer documentation: API references, SDKs, integration guides, and tutorials
- Write end-user documentation: user manuals, help articles, FAQs, and release notes
- Collaborate with engineers, product managers, and designers to gather technical information
- Develop and maintain a documentation style guide and information architecture
- Structure documentation sites using tools like Confluence, Notion, GitBook, or Docusaurus
- Review and edit technical content written by engineers for clarity and accuracy
- Gather user feedback on documentation quality and address gaps
- Produce video scripts and UI tooltips as part of in-product help
- Keep documentation up-to-date with product releases and API changes
- Ensure consistency of terminology, tone, and style across all documentation

Required Skills & Qualifications:
- Bachelor's degree in English, Journalism, Computer Science, or equivalent
- 2–5 years of technical writing experience for software or hardware products
- Ability to understand and document APIs, code samples, and CLI tools
- Proficiency with documentation tools: Confluence, Notion, ReadMe, or Docusaurus
- Strong written communication and editing skills
- Basic familiarity with markup languages: Markdown, reStructuredText, or XML/DITA

Good to Have:
- Experience writing API documentation using OpenAPI/Swagger
- Familiarity with developer tooling: Git, VS Code, terminal/CLI
- Knowledge of documentation-as-code workflows
""",

    "Technical Sales Engineer": """\
Position: Technical Sales Engineer
Location: [City, India] | Full-Time

About the Role:
We are looking for a Technical Sales Engineer to bridge the gap between our product capabilities and customer needs. You will support the sales team by providing technical expertise, conducting product demos, and ensuring customers achieve success with our solutions.

Key Responsibilities:
- Serve as the primary technical resource for the sales team during pre-sales engagements
- Conduct product demonstrations, proof-of-concept deployments, and technical evaluations
- Understand customer technical requirements and map them to our product capabilities
- Respond to RFPs, RFIs, and security questionnaires with accurate technical content
- Build and present custom technical proposals and solution architectures
- Collaborate with product and engineering teams to convey customer feedback and feature requests
- Support customer onboarding and post-sale technical transitions
- Develop technical sales collateral: battle cards, solution briefs, and integration guides
- Attend industry conferences, webinars, and customer events as a technical spokesperson
- Build long-term trusted advisor relationships with technical stakeholders at customer organizations

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 3–6 years of experience in technical sales, solutions engineering, or pre-sales
- Strong technical aptitude: APIs, cloud platforms, SaaS architectures, and integrations
- Excellent presentation, communication, and stakeholder management skills
- Experience with CRM tools (Salesforce, HubSpot) and demo environments
- Ability to explain complex technical concepts to both technical and non-technical audiences

Good to Have:
- Background in software engineering or DevOps
- Industry certifications (AWS, Azure, GCP, or product-specific)
- Experience in B2B SaaS, cloud infrastructure, or cybersecurity sales
""",

    "E-commerce Specialist": """\
Position: E-commerce Specialist
Location: [City, India] | Full-Time

About the Role:
We are seeking an E-commerce Specialist to manage and grow our online retail operations. You will own the end-to-end product listing, catalogue management, campaign execution, and performance optimization across e-commerce platforms.

Key Responsibilities:
- Manage product listings, catalogue, pricing, and inventory across e-commerce platforms (Amazon, Flipkart, Meesho, own website)
- Optimize product pages for SEO, conversion, and discoverability
- Plan and execute promotional campaigns, flash sales, and festive offers
- Analyse sales data and customer behaviour to identify growth opportunities
- Coordinate with supply chain and logistics teams for fulfilment and returns management
- Manage seller accounts, ratings, and brand storefronts on marketplace platforms
- Run and optimize performance marketing campaigns on Amazon Ads, Google Shopping
- Monitor competitor pricing, assortment, and promotional strategies
- Collaborate with design and content teams for creative assets and copy
- Prepare regular performance reports and present insights to leadership

Required Skills & Qualifications:
- Bachelor's degree in Business, Marketing, or equivalent
- 2–5 years of e-commerce or marketplace management experience
- Hands-on experience with Amazon Seller Central, Flipkart Seller Hub, or Shopify
- Strong analytical skills with Excel/Google Sheets and data dashboards
- Understanding of e-commerce SEO, listing optimization, and A+ content
- Knowledge of digital marketing: paid ads, email marketing, and social commerce

Good to Have:
- Experience with D2C brand building and own-website e-commerce
- Familiarity with ERP or inventory management software
- Knowledge of GST compliance for e-commerce in India
""",

    "Fintech Developer": """\
Position: Fintech Developer
Location: [City, India] | Full-Time

About the Role:
We are looking for an experienced Fintech Developer to build secure, reliable, and scalable financial technology solutions. You will develop features spanning payments, lending, wealth management, or banking platforms.

Key Responsibilities:
- Design and develop core fintech platform features: payments, wallets, lending, or investment modules
- Integrate with payment gateways (Razorpay, PayU, Stripe, PayPal) and banking APIs (UPI, NEFT, IMPS)
- Implement RBI-compliant KYC/AML workflows and regulatory reporting features
- Build secure financial transaction processing systems with ACID guarantees
- Develop APIs for internal microservices and third-party integrations (NBFC, bank, insurance partners)
- Implement double-entry accounting systems and reconciliation engines
- Ensure PCI-DSS compliance and end-to-end encryption for sensitive financial data
- Collaborate with compliance, risk, and product teams on regulatory and feature requirements
- Write comprehensive unit, integration, and stress tests for financial logic
- Monitor and troubleshoot production financial systems with zero-tolerance for downtime

Required Skills & Qualifications:
- B.E./B.Tech in Computer Science or equivalent
- 3–6 years of software development experience with 2+ years in fintech
- Strong backend proficiency: Java, Python, Go, or Node.js
- Experience with payment gateway and banking API integrations
- Knowledge of financial regulations: RBI guidelines, PCI-DSS, SEBI compliance
- Deep understanding of database transactions, consistency, and fault tolerance

Good to Have:
- Experience with core banking systems or lending management platforms
- Knowledge of blockchain for financial applications
- Familiarity with ISO 8583 or FIX protocol for financial messaging
""",

    "Healthcare Tech Specialist": """\
Position: Healthcare Tech Specialist
Location: [City, India] | Full-Time

About the Role:
We are seeking a Healthcare Tech Specialist to build and implement technology solutions for healthcare providers, payers, or digital health platforms. You will combine technical expertise with domain knowledge to improve patient outcomes and healthcare operations.

Key Responsibilities:
- Develop and maintain healthcare software solutions: EMR/EHR integrations, telehealth platforms, or health analytics
- Implement HL7 FHIR, HL7 v2, and DICOM standards for healthcare data interoperability
- Build secure patient data management systems compliant with HIPAA and India's DPDP Act
- Integrate with ABDM (Ayushman Bharat Digital Mission) health stack and NHA APIs
- Develop clinical decision support and care coordination features
- Build data pipelines for population health analytics and clinical reporting
- Collaborate with clinical staff, doctors, and healthcare administrators to define requirements
- Ensure system availability, data integrity, and security for critical healthcare workflows
- Implement telemedicine, scheduling, and billing modules
- Conduct compliance reviews and prepare documentation for health IT audits

Required Skills & Qualifications:
- B.E./B.Tech/B.Sc in Computer Science, Biomedical Engineering, or equivalent
- 3–6 years of healthcare IT experience
- Knowledge of HL7 FHIR, ABDM standards, and clinical terminologies (ICD-10, SNOMED-CT, LOINC)
- Strong backend development skills (Python, Java, or Node.js)
- Experience with healthcare compliance frameworks (HIPAA, ISO 27799)
- Familiarity with cloud-based healthcare platforms (AWS HealthLake, Google Healthcare API)

Good to Have:
- Experience with AI/ML applications in clinical decision support or medical imaging
- Knowledge of DICOM for medical imaging
- Background in pharmacy management or lab information systems
""",

    "EdTech Specialist": """\
Position: EdTech Specialist
Location: [City, India] | Full-Time

About the Role:
We are looking for an EdTech Specialist to build and enhance digital learning platforms that transform education. You will combine instructional design expertise with technical skills to deliver engaging, personalized learning experiences at scale.

Key Responsibilities:
- Design, develop, and maintain Learning Management Systems (LMS) and digital learning platforms
- Build adaptive learning pathways using learner data and AI-driven personalization
- Develop interactive content modules, assessments, and gamification features
- Integrate video streaming, live class, and collaborative learning tools
- Implement SCORM, xAPI, and LTI standards for content interoperability
- Build learner analytics dashboards to track engagement, progress, and outcomes
- Collaborate with educators, instructional designers, and content teams
- Develop mobile learning applications for Android and iOS
- Ensure platform accessibility (WCAG) for diverse learner populations
- Optimize platform performance for low-bandwidth environments in Tier 2/3 India

Required Skills & Qualifications:
- B.E./B.Tech or Bachelor's in Education Technology, Computer Science, or related field
- 3–6 years of EdTech product or platform development experience
- Proficiency in full stack web development (React.js, Node.js/Python)
- Experience with LMS platforms (Moodle, Canvas, or custom-built)
- Knowledge of e-learning standards: SCORM, xAPI (Tin Can), LTI
- Familiarity with video streaming technologies (HLS, WebRTC, Zoom/Agora SDK)

Good to Have:
- Experience with AI tutoring, NLP-based assessments, or adaptive learning algorithms
- Knowledge of content authoring tools (Articulate Storyline, Adobe Captivate)
- Background in K-12, higher education, or corporate training domains
""",
}

# ── Auto-fill JD from template when job title changes ───────────────────────
def _get_jd_default() -> str:
    """Return the JD template for the currently selected job title, or empty string."""
    _sel = st.session_state.get("jt_select", "— Select Job Title —")
    if _sel in ("— Select Job Title —", "Other (type below)", ""):
        return ""
    return _JD_TEMPLATES.get(_sel, "")

def _on_jt_change() -> None:
    """Called when job title selectbox changes — push template into JD field."""
    _new_title = st.session_state.get("jt_select", "")
    if _new_title not in ("— Select Job Title —", "Other (type below)", ""):
        st.session_state["jd_textarea"] = _JD_TEMPLATES.get(_new_title, "")
    else:
        st.session_state["jd_textarea"] = ""

# ---------------- Job Information Dropdown ----------------
with st.sidebar.expander("![Job](https://img.icons8.com/ios-filled/20/briefcase.png) Enter Job Details", expanded=False):

    # ── Job Title ────────────────────────────────────────────────────────────
    _jt_choice = st.selectbox(
        "![Job](https://img.icons8.com/ios-filled/20/briefcase.png) Job Title",
        _JOB_TITLE_OPTIONS,
        key="jt_select",
        on_change=_on_jt_change,
    )
    if _jt_choice == "Other (type below)":
        job_title = st.text_input(
            "Enter Job Title",
            placeholder="e.g. Prompt Engineer",
            key="jt_other_input",
        )
    elif _jt_choice == "— Select Job Title —":
        job_title = ""
    else:
        job_title = _jt_choice

    # ── Location ─────────────────────────────────────────────────────────────
    _loc_choice = st.selectbox(
        "![Location](https://img.icons8.com/ios-filled/20/marker.png) Preferred Job Location",
        _LOCATION_OPTIONS,
        key="loc_select",
    )
    if _loc_choice == "Other (type below)":
        user_location = st.text_input(
            "Enter Location",
            placeholder="e.g. Mysore, India",
            key="loc_other_input",
        )
    elif _loc_choice == "— Select Location —":
        user_location = ""
    else:
        user_location = _loc_choice

    # ── Job Description — pre-filled from template, fully editable ───────────
    if "jd_textarea" not in st.session_state:
        st.session_state["jd_textarea"] = ""

    job_description = st.text_area(
        "![Description](https://img.icons8.com/ios-filled/20/document.png) Paste Job Description",
        height=200,
        key="jd_textarea",
        placeholder="Select a Job Title above to auto-fill a standard JD, or paste your own here.",
    )

    # ── Resume Analyzer quota badge — unchanged below this line ──────────────

    # ── Resume Analyzer quota badge ───────────────────────────────────────────
    _ra_username = st.session_state.get("username")
    if _ra_username:
        _ra_used = get_usage_count_last_hour(_ra_username, "resume_analyzer")
        _ra_remaining = max(0, 2 - _ra_used)

        # colours
        _ra_accent   = "#34d399" if _ra_remaining > 0 else "#fb7185"
        _ra_bg       = "rgba(52,211,153,0.08)" if _ra_remaining > 0 else "rgba(251,113,133,0.08)"
        _ra_border   = "rgba(52,211,153,0.25)" if _ra_remaining > 0 else "rgba(251,113,133,0.25)"
        _ra_dot_col  = _ra_accent

        # animated pulse dot (green = ok, red = exhausted)
        _ra_dot = (
            f'<span style="display:inline-block;width:7px;height:7px;border-radius:50%;'
            f'background:{_ra_dot_col};box-shadow:0 0 0 0 {_ra_dot_col};'
            f'animation:raPulse 2s infinite;flex-shrink:0;"></span>'
        )

        # clock icon
        _ra_clock = (
            '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" '
            f'stroke="{_ra_accent}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" '
            'style="flex-shrink:0;">'
            '<circle cx="12" cy="12" r="10"/>'
            '<polyline points="12 6 12 12 16 14"/>'
            '</svg>'
        )

        # pill counter badge
        _ra_pill = (
            f'<span style="display:inline-flex;align-items:center;justify-content:center;'
            f'background:{_ra_accent};color:#0f172a;font-weight:700;font-size:0.72rem;'
            f'border-radius:999px;padding:1px 8px;min-width:32px;letter-spacing:0.01em;">'
            f'{_ra_remaining}/2</span>'
        )

        st.markdown(
            f"""
            <style>
            @keyframes raPulse {{
                0%   {{ box-shadow: 0 0 0 0 {_ra_dot_col}66; }}
                70%  {{ box-shadow: 0 0 0 5px {_ra_dot_col}00; }}
                100% {{ box-shadow: 0 0 0 0 {_ra_dot_col}00; }}
            }}
            </style>
            <div style="display:flex;align-items:center;gap:8px;
                        margin-top:10px;padding:8px 12px;
                        background:{_ra_bg};
                        border:1px solid {_ra_border};
                        border-radius:10px;
                        font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;">
                {_ra_dot}
                {_ra_clock}
                <span style="font-size:0.78rem;color:#94a3b8;flex:1;">
                    Analyses remaining this hour
                </span>
                {_ra_pill}
            </div>
            """,
            unsafe_allow_html=True
        )

    if job_description.strip() == "":
        st.warning("Please enter a job description to evaluate the resumes.")

# ── ATS Scoring Weights — Career Presets + Fine-tune Sliders ─────────────────

# SVG icon templates — {color} replaced at render time
_SVG_FRESHER = (
    '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" '
    'stroke="{color}" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">'
    '<path d="M22 10v6M2 10l10-5 10 5-10 5z"/>'
    '<path d="M6 12v5c3 3 9 3 12 0v-5"/>'
    '</svg>'
)
_SVG_MIDLEVEL = (
    '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" '
    'stroke="{color}" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">'
    '<rect x="2" y="7" width="20" height="14" rx="2"/>'
    '<path d="M16 7V5a2 2 0 00-2-2h-4a2 2 0 00-2 2v2"/>'
    '</svg>'
)
_SVG_EXPERIENCED = (
    '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" '
    'stroke="{color}" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">'
    '<path d="M20 21v-2a4 4 0 00-4-4H8a4 4 0 00-4 4v2"/>'
    '<circle cx="12" cy="7" r="4"/>'
    '</svg>'
)

# Rich preset definitions — color, glow, hint, svg all bundled per tier
_CAREER_PRESETS = {
    "fresher": {
        "label": "Fresher", "sublabel": "0–1 yr",
        "edu": 30, "exp": 15, "skills": 30, "lang": 5, "kw": 10,
        "hint": "Education and skills dominate — no work history to evaluate.",
        "color":  "#34d399",
        "bg":     "rgba(52,211,153,0.11)",
        "border": "rgba(52,211,153,0.30)",
        "glow":   "rgba(52,211,153,0.18)",
        "svg_tpl": _SVG_FRESHER,
    },
    "midlevel": {
        "label": "Mid-level", "sublabel": "2–5 yrs",
        "edu": 20, "exp": 30, "skills": 25, "lang": 5, "kw": 10,
        "hint": "Balanced — experience begins to outweigh education.",
        "color":  "#38bdf8",
        "bg":     "rgba(56,189,248,0.11)",
        "border": "rgba(56,189,248,0.30)",
        "glow":   "rgba(56,189,248,0.18)",
        "svg_tpl": _SVG_MIDLEVEL,
    },
    "expert": {
        "label": "Expert", "sublabel": "5+ yrs",
        "edu": 10, "exp": 40, "skills": 25, "lang": 5, "kw": 10,
        "hint": "Experience is the dominant signal at senior level.",
        "color":  "#818cf8",
        "bg":     "rgba(129,140,248,0.11)",
        "border": "rgba(129,140,248,0.30)",
        "glow":   "rgba(129,140,248,0.18)",
        "svg_tpl": _SVG_EXPERIENCED,
    },
}

# ── on_change callback — writes preset values into slider session-state keys.
# Sliders read from session_state → callback writes to session_state →
# only sidebar widgets re-render, NO full page rerun, NO flicker.
def _apply_career_preset():
    mode = st.session_state.get("career_mode_radio", "fresher")
    p = _CAREER_PRESETS.get(mode, _CAREER_PRESETS["fresher"])
    st.session_state["sl_edu"]    = p["edu"]
    st.session_state["sl_exp"]    = p["exp"]
    st.session_state["sl_skills"] = p["skills"]
    st.session_state["sl_lang"]   = p["lang"]
    st.session_state["sl_kw"]     = p["kw"]

# Session-state init — OUTSIDE expander so values survive collapse
if "career_mode_radio" not in st.session_state:
    st.session_state["career_mode_radio"] = "fresher"
if "sl_edu" not in st.session_state:
    _apply_career_preset()

# ---------------- Advanced Weights Dropdown ----------------
with st.sidebar.expander("![Settings](https://img.icons8.com/ios-filled/20/settings.png) Customize ATS Scoring Weights", expanded=False):

    # ── Description ──────────────────────────────────────────────────────────
    st.markdown(
        "<div style='font-size:0.72rem;color:#64748b;margin-bottom:14px;"
        "font-family:-apple-system,sans-serif;line-height:1.55;'>"
        "Format quality is scored automatically (10 pts fixed). "
        "Adjust the remaining <b style='color:#94a3b8;'>90 pts</b> below."
        "</div>",
        unsafe_allow_html=True,
    )

    # ── Slider CSS — red track, cyan thumb, bold red value label, uppercase labels
    st.markdown("""
    <style>
    div[data-testid="stSidebar"] .stSlider > label {
        font-size: 0.68rem !important;
        font-weight: 700 !important;
        letter-spacing: 0.10em !important;
        text-transform: uppercase !important;
        color: #64748b !important;
        font-family: -apple-system, BlinkMacSystemFont, sans-serif !important;
        margin-bottom: 2px !important;
    }
    div[data-testid="stSidebar"] .stSlider [data-testid="stSliderTrackFill"] {
        background: linear-gradient(90deg, #dc2626, #ef4444) !important;
    }
    div[data-testid="stSidebar"] .stSlider [role="slider"] {
        background: #38bdf8 !important;
        width: 17px !important;
        height: 17px !important;
        border: 2px solid #0f172a !important;
        box-shadow: 0 0 0 3px rgba(56,189,248,0.35), 0 2px 6px rgba(0,0,0,0.5) !important;
    }
    div[data-testid="stSidebar"] .stSlider [data-baseweb="tooltip"] div {
        color: #f87171 !important;
        font-weight: 700 !important;
        font-size: 0.78rem !important;
        font-family: -apple-system, BlinkMacSystemFont, sans-serif !important;
        background: transparent !important;
        box-shadow: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # ── Hidden radio — actual state driver; styled cards below are the visual UI
    _cur_level = st.session_state.get("career_mode_radio", "fresher")
    selected_mode = st.radio(
        "Career level",
        options=list(_CAREER_PRESETS.keys()),
        index=list(_CAREER_PRESETS.keys()).index(_cur_level),
        key="career_mode_radio",
        on_change=_apply_career_preset,
        label_visibility="collapsed",
        horizontal=True,
    )

    # ── 3 styled preset cards ─────────────────────────────────────────────────
    _pcols = st.columns(3)
    for _col, (_pkey, _pd) in zip(_pcols, _CAREER_PRESETS.items()):
        _is_active   = (selected_mode == _pkey)
        _icon_color  = _pd["color"] if _is_active else "#4a5568"
        _label_color = _pd["color"] if _is_active else "#4a5568"
        _sub_color   = _pd["color"] if _is_active else "#2d3748"
        _border      = f"2px solid {_pd['border']}" if _is_active else "1px solid rgba(255,255,255,0.07)"
        _bg          = _pd["bg"] if _is_active else "rgba(255,255,255,0.025)"
        _shadow      = f"0 0 12px {_pd['glow']}" if _is_active else "none"
        _icon_svg    = _pd["svg_tpl"].replace("{color}", _icon_color)
        with _col:
            st.markdown(
                f"""<div style="
                    border:{_border};background:{_bg};box-shadow:{_shadow};
                    border-radius:10px;padding:12px 4px 10px;text-align:center;
                    font-family:-apple-system,BlinkMacSystemFont,sans-serif;
                    user-select:none;">
                    <div style="display:flex;justify-content:center;margin-bottom:6px;">
                        {_icon_svg}
                    </div>
                    <div style="font-size:0.65rem;font-weight:700;
                                color:{_label_color};letter-spacing:0.01em;line-height:1.2;
                                white-space:nowrap;overflow:hidden;text-overflow:ellipsis;
                                width:100%;">
                        {_pd['label']}
                    </div>
                    <div style="font-size:0.60rem;color:{_sub_color};
                                margin-top:3px;font-weight:500;white-space:nowrap;">
                        {_pd['sublabel']}
                    </div>
                </div>""",
                unsafe_allow_html=True,
            )

    # ── Hint card — icon + color matches active tier ──────────────────────────
    _ap = _CAREER_PRESETS[selected_mode]
    _hint_icon = _ap["svg_tpl"].replace("{color}", _ap["color"])
    st.markdown(
        f"""<div style="
            margin-top:14px;padding:10px 12px;
            background:{_ap['bg']};border:1px solid {_ap['border']};
            border-radius:9px;box-shadow:0 0 14px {_ap['glow']};
            font-family:-apple-system,BlinkMacSystemFont,sans-serif;
            font-size:0.75rem;color:{_ap['color']};line-height:1.5;
            display:flex;align-items:flex-start;gap:9px;">
            <div style="flex-shrink:0;margin-top:1px;">{_hint_icon}</div>
            <span style="font-weight:500;">{_ap['hint']}</span>
        </div>""",
        unsafe_allow_html=True,
    )

    # ── FINE-TUNE label ───────────────────────────────────────────────────────
    st.markdown(
        "<div style='margin:16px 0 4px;font-size:0.63rem;font-weight:700;"
        "letter-spacing:0.13em;text-transform:uppercase;color:#334155;"
        "font-family:-apple-system,sans-serif;'>Fine-Tune</div>",
        unsafe_allow_html=True,
    )

    # ── Sliders — keyed to session_state; callback writes here, no full rerun ─
    edu_weight = st.slider(
        "Education", 5, 40, key="sl_edu",
        help="Weight given to academic qualifications and degrees.",
    )
    exp_weight = st.slider(
        "Experience", 5, 45, key="sl_exp",
        help="Weight given to work history, roles, and tenure.",
    )
    skills_weight = st.slider(
        "Skills", 5, 40, key="sl_skills",
        help="Weight given to technical and domain skill matches.",
    )
    lang_weight = st.slider(
        "Language", 2, 10, key="sl_lang",
        help="Weight given to grammar quality and language clarity.",
    )
    keyword_weight = st.slider(
        "Keywords", 3, 20, key="sl_kw",
        help="Weight given to job-description keyword alignment.",
    )

    total_weight  = edu_weight + exp_weight + skills_weight + lang_weight + keyword_weight
    weights_valid = (total_weight == 90)

    # ── Validation badge ──────────────────────────────────────────────────────
    if not weights_valid:
        _remaining = 90 - total_weight
        _direction = f"remove {abs(_remaining)}" if _remaining < 0 else f"add {_remaining}"
        st.markdown(
            f"""<div style="margin-top:12px;display:flex;align-items:center;gap:8px;
                border:1px solid rgba(251,113,133,0.3);
                background:linear-gradient(135deg,rgba(251,113,133,0.12) 0%,rgba(251,113,133,0.05) 100%);
                padding:10px 13px;border-radius:10px;">
                <svg xmlns="http://www.w3.org/2000/svg" width="15" height="15" viewBox="0 0 24 24"
                     fill="none" stroke="#fb7185" stroke-width="2" stroke-linecap="round"
                     stroke-linejoin="round" style="flex-shrink:0;">
                    <circle cx="12" cy="12" r="10"/>
                    <line x1="12" y1="8" x2="12" y2="12"/>
                    <line x1="12" y1="16" x2="12.01" y2="16"/>
                </svg>
                <span style="color:#fca5a5;font-weight:600;font-size:0.76rem;
                             font-family:-apple-system,sans-serif;">
                    Total = {total_weight} / 90 &mdash; {_direction} pts to balance.
                </span>
            </div>""",
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            """<div style="margin-top:12px;display:flex;align-items:center;gap:8px;
                border:1px solid rgba(52,211,153,0.28);
                background:linear-gradient(135deg,rgba(52,211,153,0.12) 0%,rgba(52,211,153,0.05) 100%);
                padding:10px 13px;border-radius:10px;">
                <svg xmlns="http://www.w3.org/2000/svg" width="15" height="15" viewBox="0 0 24 24"
                     fill="none" stroke="#34d399" stroke-width="2.5" stroke-linecap="round"
                     stroke-linejoin="round" style="flex-shrink:0;">
                    <polyline points="20 6 9 17 4 12"/>
                </svg>
                <span style="color:#6ee7b7;font-weight:600;font-size:0.76rem;
                             font-family:-apple-system,sans-serif;">
                    Weights balanced &middot; Content = 90 pts &middot; Format = 10 pts &middot; Total = 100
                </span>
            </div>""",
            unsafe_allow_html=True,
        )

with tab1:
    # Slide message styles already defined in global CSS — no extra block needed

    uploaded_files = st.file_uploader(
        "📄 Upload PDF Resumes",
        type=["pdf"],
        accept_multiple_files=True,
        help="Upload one or more resumes in PDF format (max 5MB each)."
    )

    # ── 5 MB hard cap ────────────────────────────────────────────────────────
    _MAX_FILE_MB  = 5
    _MAX_FILE_BYTES = _MAX_FILE_MB * 1024 * 1024

    if uploaded_files:
        for uploaded_file in uploaded_files:
            with st.container():

                # ── Size gate — reject before any processing ─────────────────
                file_bytes = uploaded_file.size  # Streamlit exposes .size directly
                file_mb    = round(file_bytes / (1024 * 1024), 2)

                if file_bytes > _MAX_FILE_BYTES:
                    _svg_oversized = f'<svg width="28" height="28" viewBox="0 0 24 24" fill="none" stroke="#fb7185" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="12" y1="11" x2="12" y2="17"/><line x1="9" y1="14" x2="15" y2="14"/></svg>'
                    _svg_dot       = '<svg width="9" height="9" viewBox="0 0 9 9" fill="none" xmlns="http://www.w3.org/2000/svg" style="flex-shrink:0;margin-top:1px;"><circle cx="4.5" cy="4.5" r="4.5" fill="#fb7185"/></svg>'
                    _svg_info      = '<svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>'
                    _svg_compress  = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><path d="M21 15v4a2 2 0 01-2 2H5a2 2 0 01-2-2v-4"/><polyline points="17 8 12 3 7 8"/><line x1="12" y1="3" x2="12" y2="15"/></svg>'
                    _svg_export    = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><rect x="4" y="2" width="16" height="20" rx="2"/><polyline points="14 2 14 8 20 8"/><line x1="8" y1="13" x2="16" y2="13"/><line x1="8" y1="17" x2="13" y2="17"/></svg>'
                    _svg_img       = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><rect x="3" y="3" width="18" height="18" rx="2"/><circle cx="8.5" cy="8.5" r="1.5"/><polyline points="21 15 16 10 5 21"/></svg>'

                    fix_items_size = [
                        (_svg_compress, "Open your resume in <strong>Microsoft Word or Google Docs</strong> and re-export as PDF"),
                        (_svg_export,   "Use <strong>File &rarr; Export &rarr; Reduce File Size</strong> or <strong>Save as PDF (Optimised)</strong>"),
                        (_svg_img,      "Remove embedded high-resolution photos or images from the resume before saving"),
                    ]
                    fix_html_size = "".join(
                        f"<li style='display:flex;align-items:flex-start;gap:8px;margin-bottom:8px;list-style:none;'>{icon}<span>{text}</span></li>"
                        for icon, text in fix_items_size
                    )

                    size_card = (
                        '<div style="background:linear-gradient(135deg,rgba(251,113,133,0.15) 0%,rgba(0,0,0,0) 100%);border:1px solid rgba(251,113,133,0.35);border-radius:16px;padding:22px 24px;margin:14px 0;backdrop-filter:blur(20px);-webkit-backdrop-filter:blur(20px);box-shadow:0 8px 32px rgba(0,0,0,0.3),inset 0 1px 0 rgba(255,255,255,0.06);font-family:-apple-system,BlinkMacSystemFont,sans-serif;position:relative;overflow:hidden;">'
                        '<div style="position:absolute;top:0;left:0;right:0;height:3px;background:linear-gradient(90deg,transparent,#fb7185,transparent);opacity:0.6;"></div>'
                        f'<div style="display:flex;align-items:flex-start;gap:14px;margin-bottom:16px;">'
                        f'<div style="width:44px;height:44px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:10px;display:flex;align-items:center;justify-content:center;flex-shrink:0;">{_svg_oversized}</div>'
                        f'<div style="flex:1;">'
                        f'<div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;color:#fb7185;margin-bottom:4px;">{_svg_dot} File Too Large — Exceeds 5 MB Limit</div>'
                        f'<div style="font-size:1rem;font-weight:600;color:#f0f4f8;word-break:break-all;">{uploaded_file.name}</div>'
                        f'</div></div>'
                        f'<div style="display:flex;gap:10px;flex-wrap:wrap;margin-bottom:16px;">'
                        f'<div style="display:flex;align-items:center;gap:6px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:8px;padding:6px 12px;font-size:0.78rem;color:#94a3b8;"><svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;"><ellipse cx="12" cy="5" rx="9" ry="3"/><path d="M21 12c0 1.66-4 3-9 3s-9-1.34-9-3"/><path d="M3 5v14c0 1.66 4 3 9 3s9-1.34 9-3V5"/></svg> {file_mb} MB uploaded</div>'
                        f'<div style="display:flex;align-items:center;gap:6px;background:rgba(251,113,133,0.10);border:1px solid rgba(251,113,133,0.25);border-radius:8px;padding:6px 12px;font-size:0.78rem;color:#fca5a5;"><svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg> Limit: {_MAX_FILE_MB} MB</div>'
                        f'</div>'
                        f'<div style="background:rgba(56,189,248,0.07);border:1px solid rgba(56,189,248,0.18);border-radius:10px;padding:12px 16px;">'
                        f'<div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;font-weight:700;letter-spacing:0.07em;text-transform:uppercase;color:#38bdf8;margin-bottom:10px;">{_svg_info} How to Reduce File Size</div>'
                        f'<ul style="margin:0;padding:0;color:#7dd3fc;font-size:0.82rem;line-height:1.8;">{fix_html_size}</ul>'
                        f'</div>'
                        f'</div>'
                    )
                    st.markdown(size_card, unsafe_allow_html=True)
                    continue  # skip all further processing for this file

                # ── Normal flow — file is within size limit ───────────────────
                st.subheader(f"📄 Original Resume Preview: {uploaded_file.name}")

                try:
                    # ✅ Show PDF preview safely
                    pdf_viewer(
                        uploaded_file.read(),
                        key=f"pdf_viewer_{uploaded_file.name}"
                    )

                    # Reset pointer so file can be read again later
                    uploaded_file.seek(0)
                    # NOTE: Text extraction + scanned rejection card is handled
                    # exclusively in the processing loop below to avoid double rendering.

                except Exception as e:
                    st.markdown(
                        f'<div class="slide-message error-msg">'
                        f'<svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;vertical-align:middle;"><circle cx="12" cy="12" r="10"/><line x1="15" y1="9" x2="9" y2="15"/><line x1="9" y1="9" x2="15" y2="15"/></svg>'
                        f' Could not display or process <b>{uploaded_file.name}</b>: {e}</div>',
                        unsafe_allow_html=True
                    )

# ✅ Initialize state
# Initialize session state
if "resume_data" not in st.session_state:
    st.session_state.resume_data = []

if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

resume_data = st.session_state.resume_data

# ✏️ Resume Evaluation Logic
if uploaded_files and job_description and not weights_valid:
    st.warning(
        f"⚠️ Resume analysis is blocked — your scoring weights add up to **{total_weight}/90**. "
        f"Please adjust the sliders in the sidebar until the total equals exactly **90**.",
        icon=None
    )

if uploaded_files and job_description and weights_valid:
    # ── Usage gate: only check when there are NEW files not yet processed ─────
    _gate_username = st.session_state.get("username")
    _new_files = [
        f for f in uploaded_files
        if f.name not in st.session_state.get("processed_files", set())
        and f.size <= _MAX_FILE_BYTES
    ]
    if _gate_username and _new_files:
        _gate_allowed, _gate_msg = check_and_gate_feature(_gate_username, "resume_analyzer")
        if not _gate_allowed:
            st.markdown(_gate_msg, unsafe_allow_html=True)
            _gate_used = get_usage_count_last_hour(_gate_username, "resume_analyzer")
            _gate_svg_clock = (
                '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" '
                'stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" '
                'style="display:inline-block;vertical-align:middle;margin-right:5px;">'
                '<circle cx="12" cy="12" r="10"/>'
                '<polyline points="12 6 12 12 16 14"/>'
                '</svg>'
            )
            st.markdown(
                f'<div style="display:flex;align-items:center;font-size:0.88rem;color:#7dd3fc;'
                f'background:rgba(56,189,248,0.08);border:1px solid rgba(56,189,248,0.2);'
                f'border-radius:8px;padding:10px 14px;margin-top:8px;font-family:-apple-system,sans-serif;">'
                f'{_gate_svg_clock} You have used <b style="margin:0 3px;">{_gate_used}/2</b> resume analyses this hour. Resets on a rolling 60-minute window.</div>',
                unsafe_allow_html=True
            )
            st.stop()
    # ─────────────────────────────────────────────────────────────────────────

    all_text = []

    for uploaded_file in uploaded_files:
        if uploaded_file.name in st.session_state.processed_files:
            continue

        # ── Skip files that exceeded the 5 MB size cap ───────────────────────
        if uploaded_file.size > _MAX_FILE_BYTES:
            continue

        # ✅ Improved optimized scanner animation with better performance
        scanner_placeholder = st.empty()

        # ✅ IMPROVED: More efficient CSS animations with GPU acceleration
        OPTIMIZED_SCANNER_HTML = f"""
        <style>
        .scanner-overlay {{
            position: fixed;
            top: 0; left: 0;
            width: 100vw; height: 100vh;
            background: linear-gradient(135deg, #0b0c10 0%, #1a1c29 100%);
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            z-index: 9999;
            will-change: transform, opacity;
        }}
        
        .scanner-doc {{
            width: 280px;
            height: 340px;
            background: linear-gradient(145deg, #f8f9fa, #e9ecef);
            border-radius: 16px;
            position: relative;
            overflow: hidden;
            box-shadow: 0 20px 40px rgba(0, 191, 255, 0.3);
            transform: translateZ(0);
            will-change: transform;
            animation: docFloat 3s ease-in-out infinite alternate;
        }}
        
        @keyframes docFloat {{
            0% {{ transform: translateY(0px) scale(1); }}
            100% {{ transform: translateY(-8px) scale(1.02); }}
        }}
        
        .doc-header {{
            padding: 20px;
            text-align: center;
            border-bottom: 2px solid #e9ecef;
        }}
        
        .doc-avatar {{
            width: 50px;
            height: 50px;
            background: linear-gradient(135deg, #667eea, #764ba2);
            border-radius: 50%;
            margin: 0 auto 10px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 20px;
            color: white;
        }}
        
        .doc-title {{
            font-size: 16px;
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 5px;
            font-family: 'Segoe UI', sans-serif;
        }}
        
        .doc-content {{
            padding: 15px;
            font-size: 12px;
            color: #6c757d;
            line-height: 1.4;
        }}
        
        .scan-line {{
            position: absolute;
            top: 0; left: 0;
            width: 100%; height: 4px;
            background: linear-gradient(90deg, transparent, rgba(0,191,255,0.8), transparent);
            animation: scanMove 2.5s ease-in-out infinite;
            box-shadow: 0 0 20px rgba(0,191,255,0.6);
            transform: translateZ(0);
            will-change: transform;
        }}
        
        @keyframes scanMove {{
            0% {{ top: 0; opacity: 1; }}
            50% {{ opacity: 0.8; }}
            100% {{ top: 340px; opacity: 1; }}
        }}
        
        .scanner-text {{
            margin-top: 30px;
            font-family: 'Orbitron', 'Segoe UI', sans-serif;
            font-weight: 600;
            font-size: 18px;
            color: #00bfff;
            text-shadow: 0 0 10px rgba(0,191,255,0.5);
            animation: textPulse 2s ease-in-out infinite;
        }}
        
        @keyframes textPulse {{
            0%, 100% {{ opacity: 1; transform: scale(1); }}
            50% {{ opacity: 0.8; transform: scale(1.05); }}
        }}
        
        .progress-bar {{
            width: 200px;
            height: 4px;
            background: rgba(255,255,255,0.2);
            border-radius: 2px;
            margin-top: 20px;
            overflow: hidden;
        }}
        
        .progress-fill {{
            height: 100%;
            background: linear-gradient(90deg, #00bfff, #1e90ff);
            border-radius: 2px;
            animation: progressFill 3s ease-in-out infinite;
            transform: translateX(-100%);
        }}
        
        @keyframes progressFill {{
            0% {{ transform: translateX(-100%); }}
            100% {{ transform: translateX(0); }}
        }}
        
        /* Mobile optimizations */
        @media (max-width: 768px) {{
            .scanner-doc {{ width: 240px; height: 300px; }}
            .scanner-text {{ font-size: 16px; }}
        }}
        </style>
        
        <div class="scanner-overlay">
            <div class="scanner-doc">
                <div class="scan-line"></div>
                <div class="doc-header">
                    <div class="doc-avatar">👤</div>
                    <div class="doc-title">{job_title}</div>
                </div>
                <div class="doc-content">
                    • Analyzing candidate profile...<br>
                    • Extracting key skills...<br>
                    • Matching with job requirements...<br>
                    • Calculating ATS compatibility...<br>
                    • Checking for bias patterns...
                </div>
            </div>
            <div class="scanner-text">Scanning Resume...</div>
            <div class="progress-bar">
                <div class="progress-fill"></div>
            </div>
        </div>
        """
        
        scanner_placeholder.markdown(OPTIMIZED_SCANNER_HTML, unsafe_allow_html=True)

        # ✅ Save uploaded file
        file_path = os.path.join(working_dir, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # ✅ Reduced delay for better UX
        time.sleep(4)

        # ✅ Extract text from PDF (scanned files return _SCANNED_SENTINEL)
        uploaded_file.seek(0)
        full_text = safe_extract_text(uploaded_file, container=tab1)
        if full_text is None or full_text in (_SCANNED_SENTINEL, _NON_ENGLISH_SENTINEL):
            # Rejection card already rendered by safe_extract_text for scanned/non-English files.
            # Plain None means unreadable for another reason — warning already shown.
            scanner_placeholder.empty()
            continue

        all_text.append(full_text)

        # ── Long resume warning — styled card, non-blocking ──────────────────
        # Content beyond 8000 chars is silently truncated in the LLM prompt.
        # _render_long_resume_warning is imported from resume_processor and
        # matches the app's glassmorphism dark theme exactly.
        if len(full_text) > _LONG_RESUME_THRESHOLD:
            with tab1:
                _render_long_resume_warning(uploaded_file.name, len(full_text), container=tab1)

        # ✅ Bias detection
        bias_score, masc_count, fem_count, detected_masc, detected_fem = detect_bias(full_text)

        # ✅ Format check (industry standard — no LLM call, run before parallel block)
        try:
            doc_check = fitz.open(file_path)
            num_pages = doc_check.page_count
            doc_check.close()
        except Exception:
            num_pages = 1
        format_data = check_resume_format(full_text, num_pages, pdf_path=file_path, session=st.session_state)

        # FIX: detect domains on the main thread BEFORE spawning parallel threads.
        # This prevents both threads from simultaneously reading/writing st.session_state
        # and firing duplicate LLM calls for domain detection.
        _pre_valid_domains = [
            "Data Science", "AI/Machine Learning", "UI/UX Design", "Mobile Development",
            "Frontend Development", "Backend Development", "Full Stack Development", "Cybersecurity",
            "Cloud Engineering", "DevOps/Infrastructure", "Quality Assurance", "Game Development",
            "Blockchain Development", "Embedded Systems", "System Architecture", "Database Management",
            "Networking", "Site Reliability Engineering", "Product Management", "Project Management",
            "Business Analysis", "Technical Writing", "Digital Marketing", "E-commerce", "Fintech",
            "Healthcare Tech", "EdTech", "IoT Development", "AR/VR Development", "Technical Sales",
            "Agile Coaching", "Software Engineering"
        ]
        _pre_domain_list = ", ".join(_pre_valid_domains)

        _pre_resume_cache_key = f"resume_domain_{hash(full_text[:500])}"
        if _pre_resume_cache_key not in st.session_state:
            _pre_resume_prompt = f"""You are a senior technical recruiter with 15+ years of experience classifying candidate profiles across all levels — freshers, students, mid-level, and senior professionals.

Your ONLY job: identify the candidate's PRIMARY professional domain from their resume text below.

════════════════════════════════════════════════════════
STEP 1 — DETERMINE CANDIDATE LEVEL FIRST
════════════════════════════════════════════════════════

Classify the candidate into one of these levels before picking a domain:

LEVEL A — Pure Fresher / Student with NO specialization evidence:
  • Still studying OR just graduated
  • No internship OR only 1 internship with no described work
  • Projects listed as names only (no descriptions, no tech stack mentioned)
  • Skills are only basic CS fundamentals (Java, C, C++, Python, HTML, SQL alone)
  → DEFAULT to "Software Engineering" immediately. Do not over-classify.
  → EXAMPLES: Only Java+MySQL+DBMS listed, no projects described → "Software Engineering"

LEVEL B — Fresher / Student WITH specialization evidence:
  • Still studying OR recently graduated BUT has AT LEAST ONE of:
    - 1 internship where the domain is clearly described (e.g. "frontend web development internship")
    - 1 project with a description mentioning domain-specific technologies
    - Skills showing a clear technology stack (not just basics)
  → DO classify into a specific domain based on the strongest evidence
  → EXAMPLES:
    - HTML+CSS+JS+React + frontend internship described → "Frontend Development"
    - Django/Laravel + MySQL + web project described, NO frontend tech mentioned → "Backend Development"
    - Django/Laravel + MySQL + HTML+CSS+JS + web project described → "Full Stack Development"
    - Android/Flutter + built a mobile app described → "Mobile Development"
    - TensorFlow/PyTorch + ML project described → "AI/Machine Learning"

LEVEL C — Experienced Professional (1+ years full-time work):
  → ALWAYS classify into a specific domain — never default to "Software Engineering" unless truly mixed
  → Use job titles + tech stack + years of experience as primary signals

════════════════════════════════════════════════════════
STEP 2 — DOMAIN CLASSIFICATION RULES (for Level B and C)
════════════════════════════════════════════════════════

RULE A — DO NOT over-classify from basic skills alone (applies to ALL levels):
  ✗ Java + MySQL + DBMS alone → NOT "Backend Development"
  ✗ HTML + CSS alone → NOT "Frontend Development"
  ✗ Python alone → NOT "AI/Machine Learning" or "Data Science"
  ✗ SQL alone → NOT "Database Management" or "Data Science"
  ✗ C / C++ alone → NOT "Embedded Systems" or "Software Engineering" specialist
  ✓ Basic CS languages without frameworks + no described projects → "Software Engineering"

RULE B — WHAT COUNTS AS TRUE DOMAIN EVIDENCE:
  → Frontend Development:
     MUST have: HTML+CSS+JS PLUS at least one of (React/Vue/Angular/Bootstrap/jQuery)
     AND: at least 1 described project OR internship explicitly about frontend/web UI
     
  → Backend Development:
     MUST have: A backend framework (Django/Flask/Spring Boot/Laravel/Express/Node.js/FastAPI)
     AND: database integration (MySQL/PostgreSQL/MongoDB) in a described project
     NOT just: Java + SQL listed in skills with no project context
     ⚠ "website" in project name does NOT mean Full Stack. Django + database + no frontend = Backend.
     
  → Full Stack Development:
     MUST have: frontend technologies (HTML+CSS+JS or React/Vue/Angular/Bootstrap/jQuery)
     AND: backend framework + database — ALL THREE explicitly present
     AND: at least 1 project or internship that uses both frontend and backend
     SELF-IDENTIFICATION counts: if summary says "full stack" or "front-end and back-end" → Full Stack
     ⚠ "website" + backend framework alone is NOT Full Stack — frontend tech must be explicitly named.
     
  → Mobile Development:
     MUST have: Android/iOS/Flutter/React Native/Kotlin/Swift
     AND: at least 1 described mobile app project
     
  → Data Science:
     MUST have: pandas/numpy/matplotlib/seaborn/tableau/power bi
     AND: actual data analysis or visualization project described
     NOT just: SQL or Excel listed in skills
     
  → AI/Machine Learning:
     MUST have: TensorFlow/PyTorch/scikit-learn/Keras/HuggingFace/LLM/NLP/Computer Vision
     AND: model training or ML pipeline described in a project
     
  → Cybersecurity:
     MUST have: security tools (Kali/Burp Suite/Wireshark/Metasploit) OR security concepts (pentesting/OWASP/CTF)
     AND: security internship or project described
     NOTE: A cybersecurity VIRTUAL internship with no tools described = weak signal, check other evidence too
     
  → DevOps/Infrastructure:
     MUST have: Docker/Kubernetes/CI-CD/Jenkins/Terraform/Ansible
     AND: deployment or infrastructure project described
     
  → Cloud Engineering:
     MUST have: AWS/Azure/GCP services (not just "cloud" mentioned)
     AND: cloud deployment or architecture in a project
     
  → UI/UX Design:
     MUST have: Figma/Adobe XD/Sketch/InVision
     AND: wireframes/prototypes/user research described
     
  → Database Management:
     MUST have: DBA role OR database optimization/administration as PRIMARY focus
     NOT just: SQL listed as one of many skills
     
  → Product Management:
     MUST have: product ownership, roadmaps, PRDs, stakeholder management
     NOT just: Agile/Scrum keywords
     
  → Project Management:
     MUST have: managing teams, project delivery, PMP/Prince2 or equivalent experience
     
  → Business Analysis:
     MUST have: requirements gathering, process mapping, business case writing
     
  → Digital Marketing:
     MUST have: SEO/SEM/campaigns/social media marketing with actual results
     
  → Blockchain Development:
     MUST have: Solidity/Web3/Smart Contracts/Ethereum/DeFi in described projects
     
  → Game Development:
     MUST have: Unity/Unreal Engine/game mechanics in described projects
     
  → Embedded Systems:
     MUST have: microcontroller/RTOS/firmware/hardware programming described
     
  → IoT Development:
     MUST have: IoT devices/sensors/protocols (MQTT/CoAP) + hardware integration
     
  → AR/VR Development:
     MUST have: ARKit/ARCore/Unity3D/Unreal/Oculus in described projects

RULE C — MIXED SIGNALS → pick the DOMINANT domain:
  • Count: technologies + described projects + internship titles per domain
  • The domain with the most evidence wins
  • If frontend has 3 signals and cybersecurity has 1 virtual internship → Frontend wins
  • If truly equal across 2 domains → "Full Stack Development" if they're frontend+backend, else "Software Engineering"
  ⚠ INTERNSHIP TITLE CONFLICT RULE (critical for Level B):
    If internship title suggests Domain A BUT skills + projects have 3+ strong signals for Domain B
    AND Domain B is more specific than Domain A → Domain B wins over the internship title.
    EXAMPLE: "Full Stack Developer Intern" + LangChain/LLaMA/RAG/FAISS/LLMs in skills+projects
             → "AI/Machine Learning" wins, NOT "Full Stack Development"
    EXAMPLE: "Full Stack Developer Intern" + only HTML/CSS/React/Node projects, no AI tools
             → "Full Stack Development" wins correctly
    EXAMPLE: "Android Developer Intern" + Flutter/Kotlin projects → "Mobile Development" wins correctly

RULE D — RESEARCH / ACADEMIC profiles:
  • Research intern at university/NIT/IIT/ISRO/DRDO etc. → classify by research TOPIC
  • AI/accessibility/NLP research → "AI/Machine Learning"
  • Security research → "Cybersecurity"  
  • Hardware/systems research → "Embedded Systems" or "Software Engineering"
  • Generic CS research → "Software Engineering"

RULE E — CAREER SWITCHERS:
  • If candidate has old domain (e.g. mechanical engineer) but new projects/courses in tech → classify by new tech domain
  • Recent certifications + projects in new domain outweigh old job titles

RULE F — JOB TITLE as strong signal (Level C only):
  • ONLY applies to Level C (1+ years full-time work experience)
  • For Level C: explicit job title is the STRONGEST single signal
  • "Backend Developer" → "Backend Development", "Data Analyst" → "Data Science"
  ⚠ For Level B (freshers/students): internship title is ONE signal among many.
    It can be OVERRIDDEN if skills + projects show 3+ strong signals for a different domain.
    Do NOT blindly use internship title for Level B — apply Rule C conflict check first.

════════════════════════════════════════════════════════
STEP 3 — FINAL CHECK BEFORE ANSWERING
════════════════════════════════════════════════════════

Ask yourself:
1. What is the candidate's LEVEL? (A / B / C)
2. If Level A → return "Software Engineering"
3. If Level B or C → what domain has the MOST evidence (technologies + described projects + internship/job titles)?
4. Does that domain meet the TRUE EVIDENCE bar from Rule B?
5. If Full Stack → are frontend tech + backend framework + database ALL explicitly mentioned? If frontend is missing → Backend, not Full Stack.
6. If Level B → did I check Rule C conflict? Does the internship title conflict with skills+projects?
   If yes → let skills+projects override the internship title.
7. If Level C → is there a job title confirming the domain (Rule F)?
8. If yes → return that domain. If no → return "Software Engineering"

════════════════════════════════════════════════════════
Resume Text:
{full_text[:2500]}
════════════════════════════════════════════════════════

Return ONLY one domain from this list, nothing else:
{_pre_domain_list}
"""
            try:
                _r = call_llm(_pre_resume_prompt, session=st.session_state).strip()
                if _r in _pre_valid_domains:
                    st.session_state[_pre_resume_cache_key] = _r
                else:
                    # LLM returned invalid domain — fall back to keyword detection
                    _kw = db_manager.detect_domain_from_title_and_description("", full_text[:3000])
                    st.session_state[_pre_resume_cache_key] = _kw if _kw != "Unclassified" else "Software Engineering"
            except Exception:
                # LLM failed entirely — fall back to keyword detection
                try:
                    _kw = db_manager.detect_domain_from_title_and_description("", full_text[:3000])
                    st.session_state[_pre_resume_cache_key] = _kw if _kw != "Unclassified" else "Software Engineering"
                except Exception:
                    st.session_state[_pre_resume_cache_key] = "Software Engineering"
        _pre_resume_domain = st.session_state[_pre_resume_cache_key]

        _pre_jd_cache_key = f"jd_domain_{hash(job_description[:500])}"
        if _pre_jd_cache_key not in st.session_state:
            _pre_jd_prompt = f"""You are an expert technical recruiter with 15+ years of experience classifying job descriptions across all industries and levels.

Your ONLY job: identify the PRIMARY professional domain this job description is hiring for.

════════════════════════════════════════════════════════
STEP 1 — READ THE JOB TITLE FIRST (strongest signal)
════════════════════════════════════════════════════════

Job Title: {job_title}

If the job title EXPLICITLY names a domain (e.g. "Backend Developer", "Data Scientist", "DevOps Engineer", "UX Designer"), use that domain directly — do not over-analyse the description.

Title override examples:
  "Backend Developer" → "Backend Development"
  "Data Analyst" → "Data Science"
  "ML Engineer" / "AI Engineer" → "AI/Machine Learning"
  "DevOps Engineer" / "Platform Engineer" → "DevOps/Infrastructure"
  "Cloud Architect" / "Cloud Engineer" → "Cloud Engineering"
  "QA Engineer" / "SDET" / "Test Engineer" → "Quality Assurance"
  "Mobile Developer" / "Android" / "iOS" / "Flutter" → "Mobile Development"
  "Full Stack Developer" → "Full Stack Development"
  "Frontend Developer" / "Front End" → "Frontend Development"
  "UX Designer" / "UI Designer" / "Product Designer" → "UI/UX Design"
  "Security Engineer" / "Security Analyst" / "Penetration Tester" → "Cybersecurity"
  "SRE" / "Site Reliability Engineer" → "Site Reliability Engineering"
  "Blockchain Developer" / "Web3 Developer" → "Blockchain Development"
  "Game Developer" / "Game Engineer" → "Game Development"
  "Embedded Engineer" / "Firmware Engineer" → "Embedded Systems"
  "IoT Engineer" → "IoT Development"
  "Network Engineer" / "Network Admin" → "Networking"
  "Database Administrator" / "DBA" → "Database Management"
  "Product Manager" → "Product Management"
  "Project Manager" / "Program Manager" → "Project Management"
  "Business Analyst" → "Business Analysis"
  "Scrum Master" / "Agile Coach" → "Agile Coaching"
  "Technical Writer" → "Technical Writing"
  "Sales Engineer" / "Pre-Sales" → "Technical Sales"
  "Solution Architect" / "Enterprise Architect" → "System Architecture"

════════════════════════════════════════════════════════
STEP 2 — IF TITLE IS AMBIGUOUS, ANALYSE THE JD BELOW
════════════════════════════════════════════════════════

Job Description:
{job_description[:2000]}

Classification rules:
  • Backend: Node.js/Django/Spring Boot/FastAPI + database + API work
  • Frontend: React/Vue/Angular/HTML+CSS+JS + UI work
  • Full Stack: Both frontend AND backend tech explicitly required
  • Data Science: SQL/Python analytics + pandas/numpy/Tableau/Power BI + analysis work
  • AI/ML: TensorFlow/PyTorch/scikit-learn/LLM/NLP/model training required
  • DevOps: Docker/Kubernetes/CI-CD/Terraform/Jenkins required
  • Cloud: AWS/Azure/GCP services explicitly required (not just "cloud" mentioned)
  • Cybersecurity: pentesting/OWASP/SIEM/SOC/security tools required
  • Mobile: Android/iOS/Flutter/React Native explicitly required
  • UI/UX: Figma/wireframes/prototyping/user research required
  • Product Management: roadmap/PRD/stakeholder management (not just Agile)
  • Project Management: team delivery/PMP/programme management
  • Business Analysis: requirements/BRD/process mapping as primary duty
  • Quality Assurance: test automation/test planning as primary duty
  • Fintech: payment/banking/trading/KYC/AML systems
  • Healthcare Tech: EHR/EMR/HIPAA/clinical systems
  • EdTech: LMS/e-learning/educational platform
  • Game Development: Unity/Unreal/game mechanics explicitly required
  • Blockchain: Solidity/Web3/smart contracts explicitly required
  • Embedded: firmware/RTOS/microcontroller/hardware explicitly required

════════════════════════════════════════════════════════
STEP 3 — FINAL CHECK
════════════════════════════════════════════════════════

1. Did the job title directly name a domain? → Use that.
2. If not, which domain has the MOST required skills/responsibilities in the JD?
3. If truly unclear → "Software Engineering"

Return ONLY one domain from this list, nothing else:
{_pre_domain_list}
"""
            try:
                _j = call_llm(_pre_jd_prompt, session=st.session_state).strip()
                if _j in _pre_valid_domains:
                    st.session_state[_pre_jd_cache_key] = _j
                else:
                    # LLM returned invalid — fall back to keyword detection
                    _jd_kw = db_manager.detect_domain_from_title_and_description(job_title, job_description[:3000])
                    st.session_state[_pre_jd_cache_key] = _jd_kw if _jd_kw != "Unclassified" else "Software Engineering"
            except Exception:
                # LLM failed — fall back to keyword detection
                try:
                    _jd_kw = db_manager.detect_domain_from_title_and_description(job_title, job_description[:3000])
                    st.session_state[_pre_jd_cache_key] = _jd_kw if _jd_kw != "Unclassified" else "Software Engineering"
                except Exception:
                    st.session_state[_pre_jd_cache_key] = "Software Engineering"
        _pre_job_domain = st.session_state[_pre_jd_cache_key]

        # ⚡ PARALLEL: rewrite + ATS run simultaneously using threads.
        # Both are network-bound (Groq API) so they benefit from parallelism
        # without needing async — ThreadPoolExecutor handles it safely.
        # Domains pre-detected above on main thread — no LLM calls fire inside threads.
        def _task_rewrite():
            return rewrite_and_highlight(full_text, replacement_mapping, user_location)

        def _task_ats():
            return ats_percentage_score(
                resume_text=full_text,
                job_description=job_description,
                logic_profile_score=None,
                edu_weight=edu_weight,
                exp_weight=exp_weight,
                skills_weight=skills_weight,
                lang_weight=lang_weight,
                keyword_weight=keyword_weight,
                format_data=format_data,
                resume_domain=_pre_resume_domain,   # FIX: pre-detected, no thread LLM call
                job_domain=_pre_job_domain,         # FIX: pre-detected, no thread LLM call
            )

        with st.spinner("✍️ Rewriting resume & running ATS evaluation in parallel..."):
            # ── Key-aware parallel/sequential decision ────────────────────────
            # When only 1–2 healthy keys remain, running both tasks in parallel
            # bursts the same key simultaneously and triggers rate limiting.
            # Switch to sequential mode with a short gap to let the TPM window recover.
            try:
                from llm_manager import load_groq_api_keys, get_healthy_keys as _ghk
                _n_healthy = len(_ghk(load_groq_api_keys()))
            except Exception:
                _n_healthy = 99  # assume enough keys if check fails

            if _n_healthy >= 3:
                # Enough keys: run truly in parallel (original behaviour)
                with ThreadPoolExecutor(max_workers=2) as _executor:
                    _future_rewrite = _executor.submit(_task_rewrite)
                    _future_ats     = _executor.submit(_task_ats)

                    try:
                        highlighted_text, rewritten_text, _, _, _, _, json_str, rewrite_ok = _future_rewrite.result()
                    except Exception:
                        highlighted_text = full_text
                        rewritten_text   = full_text
                        json_str         = ""
                        rewrite_ok       = False

                    ats_result, ats_scores = _future_ats.result()
            else:
                # Low on keys: run sequentially with a gap to protect the TPM window
                try:
                    highlighted_text, rewritten_text, _, _, _, _, json_str, rewrite_ok = _task_rewrite()
                except Exception:
                    highlighted_text = full_text
                    rewritten_text   = full_text
                    json_str         = ""
                    rewrite_ok       = False

                time.sleep(4)   # let per-minute token window partially recover

                ats_result, ats_scores = _task_ats()

        # ── Rewrite failure warning — shown once, immediately after processing ──
        # rewrite_ok=False means all API keys were exhausted and the LLM never ran.
        # The user would otherwise see their original resume silently presented as
        # "optimized" with no explanation — this makes the failure explicit.
        if not rewrite_ok:
            with tab1:
                st.warning(
                    "⚠️ **Resume rewrite unavailable** — API keys are currently exhausted. "
                    "The original resume is shown below. "
                    "Download buttons will produce a document based on the original text. "
                    "Please try again in a few minutes.",
                    icon=None,
                )

        # ✅ Resume Optimization Module — reuse JSON already produced above (0 extra LLM calls)
        try:
            optimized_resume_data = extract_resume_json(json_str)
        except Exception:
            optimized_resume_data = extract_resume_json("")  # falls back to empty skeleton

        # ✅ Extract structured ATS values
        candidate_name = ats_scores.get("Candidate Name", "Not Found")

        # ── Candidate name resolution (LLM-first, filename as validated fallback) ──
        #
        # Design: neither source is blindly trusted. Both are run through the
        # same _looks_like_person_name() validator before use. LLM always wins
        # when valid; filename is only used when the LLM result fails validation.
        #
        # Fixes:
        #   1. Job-title filenames (e.g. "mobile_app_developer_resume.pdf") no
        #      longer override a correct LLM name.
        #   2. Low character-overlap no longer triggers a blind filename override.
        #   3. Single-word extractions from either source are rejected.

        # Comprehensive set of words that appear in job-title filenames but are
        # never part of a person's name. Extend as needed.
        _NAME_STOP_WORDS: set[str] = {
            # document meta
            "resume", "cv", "curriculum", "vitae", "updated", "final",
            "new", "latest", "copy", "draft", "version", "doc",
            "v1", "v2", "v3", "v4", "v5",
            "2022", "2023", "2024", "2025", "2026",
            # seniority / role qualifiers
            "senior", "junior", "lead", "principal", "staff", "associate",
            "intern", "entry", "mid", "level",
            # job-title nouns
            "developer", "engineer", "designer", "manager", "analyst",
            "consultant", "architect", "director", "officer", "specialist",
            "coordinator", "executive", "recruiter", "advisor", "strategist",
            "scientist", "researcher", "administrator", "technician",
            # tech-domain prefixes that appear in filenames
            "mobile", "app", "web", "data", "software", "frontend", "backend",
            "fullstack", "full", "stack", "cloud", "devops", "qa", "product",
            "project", "platform", "site", "ui", "ux", "ml", "ai", "it",
            "cyber", "security", "network", "systems", "database", "infra",
        }

        def _looks_like_person_name(name: str) -> bool:
            """Return True only if *name* plausibly looks like a human full name.

            Rules (all must pass):
              • 2–4 tokens (first + last, optionally middle / suffix)
              • Every token is letters-only, length 2-25
              • No token is a known job-title / document-meta stop word
              • Not a bare placeholder string
            """
            _placeholder_values = {
                "not found", "n/a", "unknown", "none", "", "name not found",
                "candidate name not found",
                "extract full name from resume header or contact section",
                "copy the candidate's full name exactly as it appears in the resume",
                "copy the candidates full name exactly as it appears in the resume",
            }
            if name.lower().strip() in _placeholder_values:
                return False
            tokens = name.strip().split()
            if not (2 <= len(tokens) <= 4):
                return False
            for tok in tokens:
                tok_l = tok.lower()
                if tok_l in _NAME_STOP_WORDS:
                    return False
                if not re.match(r"^[a-zA-Z]{2,25}$", tok):
                    return False
            return True

        def _name_from_filename(fname: str) -> str:
            """Extract a candidate name from the filename, or return '' if none found.

            Stops at the first stop-word or digit token; requires ≥ 2 name tokens
            so that single-word job titles (e.g. 'Engineer.pdf') are rejected.
            """
            base = os.path.splitext(fname)[0]
            base = re.sub(r"[\(\)\[\]_\-\.]", " ", base)
            base = re.sub(r"\s+", " ", base).strip()
            parts: list[str] = []
            for word in base.split():
                if word.lower() in _NAME_STOP_WORDS or word.isdigit():
                    break
                if re.match(r"^[A-Za-z]{2,25}$", word):
                    parts.append(word.title())
            return " ".join(parts) if len(parts) >= 2 else ""

        # ── Resolution logic ──────────────────────────────────────────────────
        _llm_valid      = _looks_like_person_name(candidate_name)
        _filename_name  = _name_from_filename(uploaded_file.name)
        _filename_valid = _looks_like_person_name(_filename_name)

        if _llm_valid:
            pass                            # LLM passed validation — keep it
        elif _filename_valid:
            candidate_name = _filename_name  # LLM failed, filename looks real
        else:
            candidate_name = "Not Found"    # both unreliable
        # ─────────────────────────────────────────────────────────────────────

        ats_score = ats_scores.get("ATS Match %", 0)
        edu_score = ats_scores.get("Education Score", 0)
        exp_score = ats_scores.get("Experience Score", 0)
        skills_score = ats_scores.get("Skills Score", 0)
        lang_score = ats_scores.get("Language Score", 0)
        keyword_score = ats_scores.get("Keyword Score", 0)
        fmt_score = ats_scores.get("Format Score", format_data.get("format_score", 0))
        formatted_score = ats_scores.get("Formatted Score", "N/A")
        fit_summary = ats_scores.get("Final Thoughts", "N/A")
        language_analysis_full = ats_scores.get("Language Analysis", "N/A")

        missing_keywords_raw = ats_scores.get("Missing Keywords", "N/A")
        missing_skills_raw = ats_scores.get("Missing Skills", "N/A")
        missing_keywords = [kw.strip() for kw in missing_keywords_raw.split(",") if kw.strip()] if missing_keywords_raw != "N/A" else []
        missing_skills = [sk.strip() for sk in missing_skills_raw.split(",") if sk.strip()] if missing_skills_raw != "N/A" else []

        bias_flag = "High Bias" if bias_score > 0.6 else "Fair"
        ats_flag  = "Low ATS"   if ats_score < 50   else "Good ATS"

        # Reuse domain already detected inside ats_percentage_score — no extra LLM call
        domain = ats_scores.get("Resume Domain", "Unknown")

        # ✅ Store everything in session state
        st.session_state.resume_data.append({
            "Resume Name": uploaded_file.name,
            "Candidate Name": candidate_name,
            "ATS Report": ats_result,
            "ATS Match %": ats_score,
            "Formatted Score": formatted_score,
            "Education Score": edu_score,
            "Experience Score": exp_score,
            "Skills Score": skills_score,
            "Language Score": lang_score,
            "Keyword Score": keyword_score,
            "Format Score": ats_scores.get("Format Score", 0),
            "Format Grade": ats_scores.get("Format Grade", "N/A"),
            "Format Label": ats_scores.get("Format Label", ""),
            "Format Issues": ats_scores.get("Format Issues", []),
            "Format Passes": ats_scores.get("Format Passes", []),
            "Education Analysis": ats_scores.get("Education Analysis", ""),
            "Experience Analysis": ats_scores.get("Experience Analysis", ""),
            "Skills Analysis": ats_scores.get("Skills Analysis", ""),
            "Language Analysis": language_analysis_full,
            "Keyword Analysis": ats_scores.get("Keyword Analysis", ""),
            "Format Analysis": ats_scores.get("Format Analysis", ""),
            "Final Thoughts": fit_summary,
            "Missing Keywords": missing_keywords,
            "Missing Skills": missing_skills,
            "Bias Score (0 = Fair, 1 = Biased)": bias_score,
            "Bias Status": bias_flag,
            "Masculine Words": masc_count,
            "Feminine Words": fem_count,
            "Detected Masculine Words": detected_masc,
            "Detected Feminine Words": detected_fem,
            "Text Preview": full_text[:300] + "...",
            "Highlighted Text": highlighted_text,
            "Rewritten Text": rewritten_text,
            "Optimized Resume Data": optimized_resume_data,
            "Domain": domain,
            "Domain Penalty": ats_scores.get("Domain Penalty", 0),
            "Domain Similarity Score": ats_scores.get("Domain Similarity Score", 1.0),
            "Resume Domain": ats_scores.get("Resume Domain", domain),
            "Job Domain": ats_scores.get("Job Domain", "Unknown"),
        })

        insert_candidate(
            (
                uploaded_file.name,
                candidate_name,
                ats_score,
                edu_score,
                exp_score,
                skills_score,
                lang_score,
                keyword_score,
                bias_score,
                fmt_score,   # ← format_score now saved to DB
            ),
            job_title=job_title,
            job_description=job_description,
            resume_domain=domain,   # ← pass pre-detected resume domain, never re-detected
        )

        st.session_state.processed_files.add(uploaded_file.name)

        # ── Record usage after successful analysis ────────────────────────────
        _rec_username = st.session_state.get("username")
        if _rec_username:
            record_feature_usage(_rec_username, "resume_analyzer")
        # ─────────────────────────────────────────────────────────────────────

        # ── Silently email the report + optimised resume to the user ─────────
        # Runs in a daemon thread so the UI is never blocked.
        # Generates the PDF and Modern DOCX from data already in memory.
        try:
            _email_username  = st.session_state.get("username", "")
            _email_to        = get_user_email_by_username(_email_username) if _email_username else ""
            _email_candidate = candidate_name
            _email_resume_fn = uploaded_file.name

            if _email_to:
                # Build both attachments now (in the main thread, data is in scope)
                _email_html_report = generate_resume_report_html(
                    st.session_state.resume_data[-1],
                    user_location=user_location,
                )
                _email_pdf_bytes  = html_to_pdf_bytes(_email_html_report)
                _email_docx_bytes = generate_modern_docx(optimized_resume_data)

                threading.Thread(
                    target=send_analysis_email,
                    args=(
                        _email_to,
                        _email_candidate,
                        _email_pdf_bytes,
                        _email_docx_bytes,
                        _email_resume_fn,
                    ),
                    daemon=True,
                ).start()
        except Exception:
            pass  # Silent — never surface email errors to the user
        # ─────────────────────────────────────────────────────────────────────

        # ✅ IMPROVED: Smoother success animation with better transitions
        SUCCESS_HTML = """
        <style>
        .success-overlay {
            position: fixed;
            top: 0; left: 0;
            width: 100vw; height: 100vh;
            background: linear-gradient(135deg, #0b0c10 0%, #1a1c29 100%);
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            z-index: 9999;
            animation: fadeIn 0.5s ease-out;
        }
        
        @keyframes fadeIn {
            0% { opacity: 0; }
            100% { opacity: 1; }
        }
        
        .success-circle {
            width: 140px;
            height: 140px;
            border: 3px solid #00bfff;
            border-radius: 50%;
            position: relative;
            display: flex;
            align-items: center;
            justify-content: center;
            background: radial-gradient(circle, rgba(0,191,255,0.1) 0%, rgba(0,191,255,0.05) 50%, transparent 100%);
            animation: successPulse 2s ease-in-out infinite;
        }
        
        @keyframes successPulse {
            0%, 100% { 
                transform: scale(1);
                box-shadow: 0 0 20px rgba(0,191,255,0.3);
            }
            50% { 
                transform: scale(1.05);
                box-shadow: 0 0 30px rgba(0,191,255,0.6);
            }
        }
        
        .success-checkmark {
            font-size: 48px;
            color: #00ff7f;
            animation: checkmarkPop 0.8s ease-out;
        }
        
        @keyframes checkmarkPop {
            0% { transform: scale(0) rotate(-45deg); opacity: 0; }
            50% { transform: scale(1.2) rotate(-10deg); opacity: 0.8; }
            100% { transform: scale(1) rotate(0deg); opacity: 1; }
        }
        
        .success-text {
            margin-top: 25px;
            font-family: 'Orbitron', 'Segoe UI', sans-serif;
            font-size: 20px;
            font-weight: 600;
            color: #00bfff;
            text-shadow: 0 0 10px rgba(0,191,255,0.5);
            animation: textSlideUp 0.8s ease-out 0.3s both;
        }
        
        @keyframes textSlideUp {
            0% { transform: translateY(20px); opacity: 0; }
            100% { transform: translateY(0); opacity: 1; }
        }
        
        .success-subtitle {
            margin-top: 10px;
            font-size: 14px;
            color: #8e9aaf;
            animation: textSlideUp 0.8s ease-out 0.5s both;
        }
        </style>
        
        <div class="success-overlay">
            <div class="success-circle">
                <div class="success-checkmark">✓</div>
            </div>
            <div class="success-text">Scan Complete!</div>
            <div class="success-subtitle">Resume analysis ready</div>
        </div>
        """
        
        # Clear scanner and show success animation
        scanner_placeholder.empty()
        success_placeholder = st.empty()
        success_placeholder.markdown(SUCCESS_HTML, unsafe_allow_html=True)

        # ⏳ Shorter delay for better UX, then clear and rerun
        time.sleep(3)
        success_placeholder.empty()
        st.rerun()

    # ✅ Optional vectorstore setup
    if all_text:
        st.session_state.vectorstore = setup_vectorstore(all_text)
        st.session_state.chain = create_chain(st.session_state.vectorstore)

# 🔄 Developer Reset Button
with tab1:
    if st.button("🔄 Refresh view"):
        st.session_state.processed_files.clear()
        st.session_state.resume_data.clear()

        # Temporary placeholder for sliding success message
        msg_placeholder = st.empty()
        msg_placeholder.markdown("""
        <div class='slide-message success-msg'>
            <svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;vertical-align:middle;"><polyline points="20 6 9 17 4 12"/></svg>
            Cleared uploaded resume history. You can re-upload now.
        </div>
        """, unsafe_allow_html=True)

        # Wait 3 seconds then clear message
        time.sleep(3)
        msg_placeholder.empty()

# === TAB 1: Dashboard ===
with tab1:
    resume_data = st.session_state.get("resume_data", [])

    if resume_data:
        # ✅ Calculate total counts safely
        total_masc = sum(len(r.get("Detected Masculine Words", [])) for r in resume_data)
        total_fem = sum(len(r.get("Detected Feminine Words", [])) for r in resume_data)
        avg_bias = round(np.mean([r.get("Bias Score (0 = Fair, 1 = Biased)", 0) for r in resume_data]), 2)
        total_resumes = len(resume_data)

        st.markdown("<p class='section-label'>Session Summary</p>", unsafe_allow_html=True)
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Resumes Uploaded", total_resumes)
        with col2:
            st.metric("Avg. Bias Score", avg_bias)
        with col3:
            st.metric("Total Masculine Words", total_masc)
        with col4:
            st.metric("Total Feminine Words", total_fem)

        st.markdown("<p class='section-label'>Resumes Overview</p>", unsafe_allow_html=True)
        df = pd.DataFrame(resume_data)

        # ✅ Add calculated count columns safely
        df["Masculine Words Count"] = df["Detected Masculine Words"].apply(lambda x: len(x) if isinstance(x, list) else 0)
        df["Feminine Words Count"] = df["Detected Feminine Words"].apply(lambda x: len(x) if isinstance(x, list) else 0)

        overview_cols = [
            "Resume Name", "Candidate Name", "ATS Match %", "Education Score",
            "Experience Score", "Skills Score", "Language Score", "Keyword Score",
            "Format Score",
            "Bias Score (0 = Fair, 1 = Biased)", "Masculine Words Count", "Feminine Words Count"
        ]

        st.dataframe(df[overview_cols], use_container_width=True)

        st.markdown("<p class='section-label'>Visual Analysis</p>", unsafe_allow_html=True)
        chart_tab1, chart_tab2 = st.tabs(["Bias Score Chart", "Gender-Coded Words"])
        with chart_tab1:
            st.subheader("Bias Score Comparison Across Resumes")
            bias_chart_df = df[["Resume Name", "Bias Score (0 = Fair, 1 = Biased)"]].copy()
            bias_chart_df.columns = ["Resume", "Bias Score"]
            bias_altair = alt.Chart(bias_chart_df).mark_bar(
                cornerRadiusTopLeft=4,
                cornerRadiusTopRight=4,
                color="#4f8cff"
            ).encode(
                x=alt.X("Resume:N", sort=None, axis=alt.Axis(labelAngle=-35, labelFontSize=11, titleFontSize=12)),
                y=alt.Y("Bias Score:Q", scale=alt.Scale(domain=[0, 1]), axis=alt.Axis(titleFontSize=12)),
                tooltip=["Resume", alt.Tooltip("Bias Score:Q", format=".2f")]
            ).properties(height=260).configure_view(strokeWidth=0).configure_axis(
                grid=False, domainColor="#2d3748"
            )
            st.altair_chart(bias_altair, use_container_width=True)
        with chart_tab2:
            st.subheader("Masculine vs Feminine Word Usage")
            gender_df = pd.DataFrame({
                "Resume": list(df["Resume Name"]) * 2,
                "Type": ["Masculine"] * len(df) + ["Feminine"] * len(df),
                "Count": list(df["Masculine Words Count"]) + list(df["Feminine Words Count"])
            })
            color_scale = alt.Scale(domain=["Masculine", "Feminine"], range=["#4f8cff", "#fb7185"])
            gender_altair = alt.Chart(gender_df).mark_bar(cornerRadiusTopLeft=3, cornerRadiusTopRight=3).encode(
                x=alt.X("Resume:N", sort=None, axis=alt.Axis(labelAngle=-35, labelFontSize=11, titleFontSize=12)),
                y=alt.Y("Count:Q", axis=alt.Axis(titleFontSize=12)),
                color=alt.Color("Type:N", scale=color_scale, legend=alt.Legend(orient="top", titleFontSize=11)),
                xOffset="Type:N",
                tooltip=["Resume", "Type", "Count"]
            ).properties(height=260).configure_view(strokeWidth=0).configure_axis(
                grid=False, domainColor="#2d3748"
            )
            st.altair_chart(gender_altair, use_container_width=True)

        st.markdown("<p class='section-label'>Detailed Resume Reports</p>", unsafe_allow_html=True)
        for resume in resume_data:
            candidate_name = resume.get("Candidate Name", "Not Found")
            resume_name = resume.get("Resume Name", "Unknown")
            missing_keywords = resume.get("Missing Keywords", [])
            missing_skills = resume.get("Missing Skills", [])

            with st.expander(f"{resume_name} | {candidate_name}"):
                st.markdown(
                    f'<div style="background:linear-gradient(135deg,rgba(56,189,248,0.10) 0%,rgba(79,163,227,0.05) 100%);border:1px solid rgba(56,189,248,0.18);border-radius:14px;padding:18px 22px;margin-bottom:20px;">'
                    f'<div style="font-family:-apple-system,BlinkMacSystemFont,sans-serif;font-size:1rem;font-weight:700;color:#f0f4f8;letter-spacing:-0.01em;">ATS Evaluation — <span style="color:#38bdf8;">{candidate_name}</span></div>'
                    f'<div style="font-size:0.75rem;color:#64748b;margin-top:4px;font-family:-apple-system,sans-serif;text-transform:uppercase;letter-spacing:0.05em;">Resume Intelligence Report</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

                # ── SVG icon helper ──────────────────────────────────────────────
                SVG_ICONS = {
                    "overall": '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg>',
                    "grade":   '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M22 11.08V12a10 10 0 1 1-5.93-9.14"/><polyline points="22 4 12 14.01 9 11.01"/></svg>',
                    "edu":     '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M22 10v6M2 10l10-5 10 5-10 5z"/><path d="M6 12v5c3 3 9 3 12 0v-5"/></svg>',
                    "exp":     '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="2" y="7" width="20" height="14" rx="2"/><path d="M16 7V5a2 2 0 0 0-2-2h-4a2 2 0 0 0-2 2v2"/></svg>',
                    "skills":  '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="3"/><path d="M12 1v4M12 19v4M4.22 4.22l2.83 2.83M16.95 16.95l2.83 2.83M1 12h4M19 12h4M4.22 19.78l2.83-2.83M16.95 7.05l2.83-2.83"/></svg>',
                    "lang":    '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M21 15a2 2 0 0 1-2 2H7l-4 4V5a2 2 0 0 1 2-2h14a2 2 0 0 1 2 2z"/></svg>',
                    "keyword": '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="11" cy="11" r="8"/><line x1="21" y1="21" x2="16.65" y2="16.65"/></svg>',
                    "format":  '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="9 11 12 14 22 4"/><path d="M21 12v7a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11"/></svg>',
                    "pass":    '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#34d399" stroke-width="2.5"><polyline points="20 6 9 17 4 12"/></svg>',
                    "fail":    '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#f87171" stroke-width="2.5"><line x1="18" y1="6" x2="6" y2="18"/><line x1="6" y1="6" x2="18" y2="18"/></svg>',
                }

                def svg_ats_card(svg_key, label, value, tooltip=None):
                    tooltip_attr = f'title="{tooltip}"' if tooltip else ""
                    return (
                        f'<div style="background:rgba(15,23,42,0.85);border:1px solid rgba(56,189,248,0.25);border-radius:12px;padding:14px 16px;margin-bottom:8px;height:86px;display:flex;flex-direction:column;justify-content:center;overflow:hidden;box-sizing:border-box;">'
                        f'<div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;color:#94a3b8;">'
                        f'<span style="color:#38bdf8;flex-shrink:0;">{SVG_ICONS.get(svg_key,"")}</span>'
                        f'<span style="white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{label}</span>'
                        f'</div>'
                        f'<div {tooltip_attr} style="font-size:1.35rem;font-weight:700;color:#f0f4f8;margin-top:6px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{value}</div>'
                        f'</div>'
                    )

                # ── Overall Score Gauge (SVG) ────────────────────────────────────
                overall_pct = resume.get("ATS Match %", 0)
                fmt_score   = resume.get("Format Score", 0)
                fmt_grade   = resume.get("Format Grade", "N/A")
                fmt_label   = resume.get("Format Label", "")

                # Gauge colour
                if overall_pct >= 75:
                    gauge_color = "#22c55e"
                elif overall_pct >= 55:
                    gauge_color = "#f59e0b"
                else:
                    gauge_color = "#ef4444"

                # SVG arc gauge
                radius = 70
                cx, cy = 90, 90
                circumference = 3.14159 * radius  # half-circle arc = π*r
                arc_offset = circumference * (1 - overall_pct / 100)

                gauge_html = (
                    f'<div style="display:flex;align-items:center;gap:32px;padding:20px 24px;background:rgba(15,23,42,0.9);border:1px solid rgba(56,189,248,0.2);border-radius:16px;margin-bottom:20px;flex-wrap:wrap;">'
                    f'<div style="flex-shrink:0;text-align:center;">'
                    f'<svg width="180" height="100" viewBox="0 0 180 100">'
                    f'<path d="M 20 90 A 70 70 0 0 1 160 90" fill="none" stroke="rgba(255,255,255,0.08)" stroke-width="12" stroke-linecap="round"/>'
                    f'<path d="M 20 90 A 70 70 0 0 1 160 90" fill="none" stroke="{gauge_color}" stroke-width="12" stroke-linecap="round" stroke-dasharray="{circumference}" stroke-dashoffset="{arc_offset}" style="transition:stroke-dashoffset 0.8s ease;"/>'
                    f'<text x="90" y="80" text-anchor="middle" font-size="28" font-weight="700" fill="{gauge_color}" font-family="-apple-system,sans-serif">{overall_pct}</text>'
                    f'<text x="90" y="98" text-anchor="middle" font-size="11" fill="#64748b" font-family="-apple-system,sans-serif">/ 100</text>'
                    f'</svg>'
                    f'<div style="font-size:0.75rem;color:#64748b;margin-top:2px;font-family:-apple-system,sans-serif;letter-spacing:0.04em;text-transform:uppercase;">Overall ATS Score</div>'
                    f'</div>'
                    f'<div style="flex:1;min-width:200px;">'
                    f'<div style="font-size:1.1rem;font-weight:700;color:#f0f4f8;font-family:-apple-system,sans-serif;">{resume.get("Formatted Score","N/A")}</div>'
                    f'<div style="margin-top:12px;display:flex;align-items:center;gap:10px;">'
                    f'<span style="color:#38bdf8;">{SVG_ICONS["format"]}</span>'
                    f'<span style="font-size:0.82rem;color:#94a3b8;">Format Score:</span>'
                    f'<span style="font-size:0.95rem;font-weight:700;color:#f0f4f8;">{fmt_score}/100</span>'
                    f'<span style="background:rgba(56,189,248,0.12);border:1px solid rgba(56,189,248,0.25);border-radius:6px;padding:2px 8px;font-size:0.75rem;font-weight:700;color:#38bdf8;">{fmt_grade}</span>'
                    f'</div>'
                    f'<div style="margin-top:6px;font-size:0.78rem;color:#64748b;">{fmt_label}</div>'
                    f'</div>'
                    f'</div>'
                )
                st.markdown(gauge_html, unsafe_allow_html=True)

                # ── Score cards row 1 ──────────────────────────────────────────
                formatted_val = resume.get("Formatted Score", "N/A")
                score_col1, score_col2, score_col3 = st.columns(3)
                with score_col1:
                    st.markdown(svg_ats_card("overall", "Overall ATS Match", f"{resume.get('ATS Match %', 'N/A')}%"), unsafe_allow_html=True)
                with score_col2:
                    st.markdown(svg_ats_card("grade", "Hire Signal", formatted_val, tooltip=formatted_val), unsafe_allow_html=True)
                with score_col3:
                    st.markdown(svg_ats_card("lang", "Language Quality", f"{resume.get('Language Score', 'N/A')} / {lang_weight}"), unsafe_allow_html=True)

                # ── Score cards row 2 ──────────────────────────────────────────
                col_a, col_b, col_c, col_d = st.columns(4)
                with col_a:
                    st.markdown(svg_ats_card("edu", "Education", f"{resume.get('Education Score', 'N/A')} / {edu_weight}"), unsafe_allow_html=True)
                with col_b:
                    st.markdown(svg_ats_card("exp", "Experience", f"{resume.get('Experience Score', 'N/A')} / {exp_weight}"), unsafe_allow_html=True)
                with col_c:
                    st.markdown(svg_ats_card("skills", "Skills", f"{resume.get('Skills Score', 'N/A')} / {skills_weight}"), unsafe_allow_html=True)
                with col_d:
                    st.markdown(svg_ats_card("keyword", "Keywords", f"{resume.get('Keyword Score', 'N/A')} / {keyword_weight}"), unsafe_allow_html=True)

                # ── Score cards row 3: bias + domain status ────────────────────
                SVG_BIAS  = '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>'
                SVG_DOM   = '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="10"/><line x1="2" y1="12" x2="22" y2="12"/><path d="M12 2a15.3 15.3 0 0 1 4 10 15.3 15.3 0 0 1-4 10 15.3 15.3 0 0 1-4-10 15.3 15.3 0 0 1 4-10z"/></svg>'

                bias_raw   = resume.get("Bias Score (0 = Fair, 1 = Biased)", 0)
                bias_pct   = round(bias_raw * 100)
                bias_label = "High Bias" if bias_raw > 0.6 else ("Moderate" if bias_raw > 0.3 else "Fair")
                bias_color = "#ef4444" if bias_raw > 0.6 else ("#f59e0b" if bias_raw > 0.3 else "#22c55e")

                dom_penalty = resume.get("Domain Penalty", 0)
                dom_penalty = dom_penalty if isinstance(dom_penalty, (int, float)) else 0
                dom_sim     = resume.get("Domain Similarity Score", 1.0)
                dom_sim     = dom_sim if isinstance(dom_sim, (int, float)) else 1.0
                dom_pct     = round(dom_sim * 100)
                dom_label   = resume.get("Resume Domain", resume.get("Domain", "Unknown"))

                r3c1, r3c2, r3c3, r3c4 = st.columns(4)
                with r3c1:
                    st.markdown(f"""
                    <div style="background:rgba(15,23,42,0.85);border:1px solid rgba(56,189,248,0.25);
                                border-radius:12px;padding:14px 16px;margin-bottom:8px;height:86px;
                                display:flex;flex-direction:column;justify-content:center;overflow:hidden;">
                        <div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;color:#94a3b8;">
                            <span style="color:{bias_color};flex-shrink:0;">{SVG_BIAS}</span>
                            <span>Bias Status</span>
                        </div>
                        <div style="font-size:1.1rem;font-weight:700;color:{bias_color};margin-top:6px;">
                            {bias_label} <span style="font-size:0.8rem;color:#64748b;">({bias_pct}%)</span>
                        </div>
                    </div>""", unsafe_allow_html=True)
                with r3c2:
                    st.markdown(f"""
                    <div style="background:rgba(15,23,42,0.85);border:1px solid rgba(56,189,248,0.25);
                                border-radius:12px;padding:14px 16px;margin-bottom:8px;height:86px;
                                display:flex;flex-direction:column;justify-content:center;overflow:hidden;">
                        <div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;color:#94a3b8;">
                            <span style="color:#38bdf8;flex-shrink:0;">{SVG_DOM}</span>
                            <span>Domain Match</span>
                        </div>
                        <div style="font-size:1.1rem;font-weight:700;color:#f0f4f8;margin-top:6px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">
                            {dom_pct}% <span style="font-size:0.75rem;color:#64748b;">(-{dom_penalty} pts)</span>
                        </div>
                    </div>""", unsafe_allow_html=True)
                with r3c3:
                    st.markdown(svg_ats_card("format", "Format Score", f"{resume.get('Format Score', 'N/A')}/100 · {resume.get('Format Grade','N/A')}"), unsafe_allow_html=True)
                with r3c4:
                    masc_c = len(resume.get("Detected Masculine Words", []))
                    fem_c  = len(resume.get("Detected Feminine Words", []))
                    SVG_WORDS = '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M21 15a2 2 0 0 1-2 2H7l-4 4V5a2 2 0 0 1 2-2h14a2 2 0 0 1 2 2z"/></svg>'
                    st.markdown(f"""
                    <div style="background:rgba(15,23,42,0.85);border:1px solid rgba(56,189,248,0.25);
                                border-radius:12px;padding:14px 16px;margin-bottom:8px;height:86px;
                                display:flex;flex-direction:column;justify-content:center;overflow:hidden;">
                        <div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;color:#94a3b8;">
                            <span style="color:#38bdf8;flex-shrink:0;">{SVG_WORDS}</span>
                            <span>Gender Words</span>
                        </div>
                        <div style="font-size:0.95rem;font-weight:700;color:#f0f4f8;margin-top:6px;">
                            <span style="color:#60a5fa;">{masc_c} M</span>
                            <span style="color:#64748b;margin:0 4px;">/</span>
                            <span style="color:#f87171;">{fem_c} F</span>
                        </div>
                    </div>""", unsafe_allow_html=True)

                # ── Format Checker Panel ───────────────────────────────────────
                fmt_issues = resume.get("Format Issues", [])
                fmt_passes = resume.get("Format Passes", [])
                st.markdown("""
                <div style="margin:16px 0 6px;font-size:0.72rem;font-weight:700;color:#64748b;
                            letter-spacing:0.08em;text-transform:uppercase;font-family:-apple-system,sans-serif;">
                    Format &amp; ATS Compatibility Check
                </div>""", unsafe_allow_html=True)
                
                issues_html = "".join(
                    f"<div style='display:flex;align-items:flex-start;gap:6px;margin-bottom:5px;font-size:0.8rem;color:#fca5a5;'>{SVG_ICONS['fail']}<span>{iss}</span></div>"
                    for iss in fmt_issues
                ) if fmt_issues else "<div style='font-size:0.8rem;color:#94a3b8;'>No critical issues detected.</div>"
                passes_html = "".join(
                    f"<div style='display:flex;align-items:flex-start;gap:6px;margin-bottom:5px;font-size:0.8rem;color:#6ee7b7;'>{SVG_ICONS['pass']}<span>{p}</span></div>"
                    for p in fmt_passes
                ) if fmt_passes else ""

                fmt_col1, fmt_col2 = st.columns(2)
                with fmt_col1:
                    st.markdown(f"""
                    <div style="background:rgba(239,68,68,0.06);border:1px solid rgba(239,68,68,0.2);
                                border-radius:10px;padding:12px 14px;">
                        <div style="font-size:0.72rem;font-weight:700;color:#f87171;text-transform:uppercase;
                                    letter-spacing:0.06em;margin-bottom:8px;display:flex;align-items:center;gap:6px;">
                            {SVG_ICONS['fail']} Issues ({len(fmt_issues)})
                        </div>
                        {issues_html}
                    </div>""", unsafe_allow_html=True)
                with fmt_col2:
                    st.markdown(f"""
                    <div style="background:rgba(52,211,153,0.06);border:1px solid rgba(52,211,153,0.2);
                                border-radius:10px;padding:12px 14px;">
                        <div style="font-size:0.72rem;font-weight:700;color:#34d399;text-transform:uppercase;
                                    letter-spacing:0.06em;margin-bottom:8px;display:flex;align-items:center;gap:6px;">
                            {SVG_ICONS['pass']} Passed ({len(fmt_passes)})
                        </div>
                        {passes_html}
                    </div>""", unsafe_allow_html=True)

                # Fit summary
                st.markdown("""
                <div style="margin:18px 0 6px;font-size:0.72rem;font-weight:700;color:#64748b;
                            letter-spacing:0.08em;text-transform:uppercase;font-family:-apple-system,sans-serif;">
                    Fit Summary
                </div>""", unsafe_allow_html=True)
                st.write(resume.get('Final Thoughts', 'N/A'))

                # ATS Report
                if resume.get("ATS Report"):
                    st.markdown("<p class='section-label'>ATS Evaluation Report</p>", unsafe_allow_html=True)
                    st.markdown(resume["ATS Report"], unsafe_allow_html=True)

                # ATS Chart
                st.markdown("<p class='section-label'>ATS Score Breakdown</p>", unsafe_allow_html=True)
                # Normalize each component score to 0–100 scale for fair visual comparison
                def _pct(score, weight):
                    return round(score / weight * 100) if weight > 0 else 0
                ats_df = pd.DataFrame({
                    'Component': ['Education', 'Experience', 'Skills', 'Language', 'Keywords', 'Format'],
                    'Score': [
                        _pct(resume.get("Education Score", 0), edu_weight),
                        _pct(resume.get("Experience Score", 0), exp_weight),
                        _pct(resume.get("Skills Score", 0), skills_weight),
                        _pct(resume.get("Language Score", 0), lang_weight) if lang_weight > 0 else 0,
                        _pct(resume.get("Keyword Score", 0), keyword_weight),
                        resume.get("Format Score", 0),  # Already on 0–100 scale
                    ]
                })
                ats_chart = alt.Chart(ats_df).mark_bar().encode(
                    x=alt.X('Component', sort=None),
                    y=alt.Y('Score', scale=alt.Scale(domain=[0, 100]), title='Score (% of weight)'),
                    color='Component',
                    tooltip=['Component', 'Score']
                ).properties(
                    title="ATS Evaluation Breakdown (All scores normalized to 0–100%)",
                    width=600,
                    height=300
                )
                st.altair_chart(ats_chart, use_container_width=True)

                st.markdown("<p class='section-label'>Detailed ATS Section Analyses</p>", unsafe_allow_html=True)

                # ── Fixed section order — always rendered in this sequence ──────────
                # Markdown is converted to HTML so **bold**, bullet lists etc. render
                # correctly instead of showing raw asterisks.
                def _md_to_html(text: str) -> str:
                    """Minimal markdown → HTML converter for ATS section bodies."""
                    import html as _html
                    lines = text.split("\n")
                    out, in_ul = [], False
                    for line in lines:
                        # Bold  **text**
                        line = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
                        # Italic *text*
                        line = re.sub(r'\*(.+?)\*', r'<i>\1</i>', line)
                        # Bullet point lines  - item  or  • item
                        if re.match(r'^\s*[-•]\s+', line):
                            if not in_ul:
                                out.append("<ul style='margin:6px 0 6px 16px;padding:0;'>")
                                in_ul = True
                            content = re.sub(r'^\s*[-•]\s+', '', line)
                            out.append(f"<li style='margin-bottom:3px;'>{content}</li>")
                        else:
                            if in_ul:
                                out.append("</ul>")
                                in_ul = False
                            if line.strip():
                                out.append(f"<p style='margin:4px 0;'>{line}</p>")
                    if in_ul:
                        out.append("</ul>")
                    return "".join(out)

                _SECTION_ORDER = [
                    ("Education Analysis",        "Education Analysis"),
                    ("Experience Analysis",        "Experience Analysis"),
                    ("Skills Analysis",            "Skills Analysis"),
                    ("Language Quality",           "Language Analysis"),
                    ("Keyword Analysis",           "Keyword Analysis"),
                    ("Format & ATS Compatibility", "Format Analysis"),
                    ("Final Assessment",           "Final Thoughts"),
                ]

                for section_title, key in _SECTION_ORDER:
                    analysis_content = resume.get(key, "") or ""
                    # Guard: treat empty / whitespace-only / literal "N/A" the same way
                    _is_empty = not analysis_content.strip() or analysis_content.strip().upper() == "N/A"

                    if _is_empty:
                        body_html = (
                            "<div style='color:#94a3b8;font-size:0.85rem;font-style:italic;padding:4px 0;'>"
                            "Analysis not available for this resume — the LLM response did not include this section."
                            "</div>"
                        )
                    elif "**Score:**" in analysis_content:
                        parts    = analysis_content.split("**Score:**", 1)
                        after    = parts[1]
                        # Score value ends at next ** or newline
                        score_end = re.search(r'[\n*]', after)
                        score_text = after[:score_end.start()].strip() if score_end else after[:20].strip()
                        remaining  = after[score_end.start():].strip() if score_end else ""
                        score_html = (
                            f"<span class='score-badge' style='display:inline-block;margin-bottom:8px;'>"
                            f"Score: {score_text}</span>"
                        )
                        body_html = f"{score_html}<div style='margin-top:6px;'>{_md_to_html(remaining)}</div>"
                    else:
                        body_html = f"<div>{_md_to_html(analysis_content)}</div>"

                    st.markdown(f"""
<div class="ats-section-header">{section_title}</div>
<div class="ats-section-body">{body_html}</div>
""", unsafe_allow_html=True)

                st.divider()

                detail_tab1, detail_tab2 = st.tabs(["Bias Analysis", "Rewritten Resume"])

                with detail_tab1:
                    st.markdown("""
                    <div style="display:flex;align-items:center;gap:8px;margin:12px 0 6px;">
                        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2"><circle cx="11" cy="11" r="8"/><line x1="21" y1="21" x2="16.65" y2="16.65"/></svg>
                        <span class='section-label' style="margin:0;">Bias-Highlighted Original Text</span>
                    </div>""", unsafe_allow_html=True)
                    st.markdown(resume["Highlighted Text"], unsafe_allow_html=True)

                    st.markdown("""
                    <div style="display:flex;align-items:center;gap:8px;margin:14px 0 6px;">
                        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2"><line x1="8" y1="6" x2="21" y2="6"/><line x1="8" y1="12" x2="21" y2="12"/><line x1="8" y1="18" x2="21" y2="18"/><line x1="3" y1="6" x2="3.01" y2="6"/><line x1="3" y1="12" x2="3.01" y2="12"/><line x1="3" y1="18" x2="3.01" y2="18"/></svg>
                        <span class='section-label' style="margin:0;">Gender-Coded Word Counts</span>
                    </div>""", unsafe_allow_html=True)
                    bias_col1, bias_col2 = st.columns(2)

                    with bias_col1:
                        st.metric("Masculine Words", len(resume["Detected Masculine Words"]))
                        if resume["Detected Masculine Words"]:
                            st.markdown("<p class='section-label'>Masculine Words with Context</p>", unsafe_allow_html=True)
                            for item in resume["Detected Masculine Words"]:
                                word = item['word']
                                sentence = item['sentence']
                                st.markdown(f"""<div style='margin-bottom:6px;font-size:0.85rem;'>
                                    <span style='color:#60a5fa;font-weight:600;'>{word}</span>: {sentence}</div>""",
                                    unsafe_allow_html=True)
                        else:
                            st.info("No masculine words detected.")

                    with bias_col2:
                        st.metric("Feminine Words", len(resume["Detected Feminine Words"]))
                        if resume["Detected Feminine Words"]:
                            st.markdown("<p class='section-label'>Feminine Words with Context</p>", unsafe_allow_html=True)
                            for item in resume["Detected Feminine Words"]:
                                word = item['word']
                                sentence = item['sentence']
                                st.markdown(f"""<div style='margin-bottom:6px;font-size:0.85rem;'>
                                    <span style='color:#f87171;font-weight:600;'>{word}</span>: {sentence}</div>""",
                                    unsafe_allow_html=True)
                        else:
                            st.info("No feminine words detected.")

                with detail_tab2:
                    st.markdown("""
                    <div style="display:flex;align-items:center;gap:8px;margin:12px 0 6px;">
                        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="#34d399" stroke-width="2"><polyline points="9 11 12 14 22 4"/><path d="M21 12v7a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11"/></svg>
                        <span class='section-label' style="margin:0;">Bias-Free Rewritten Resume</span>
                    </div>""", unsafe_allow_html=True)

                    # ── Job Title Suggestions (Analysis Module — displayed here, NOT in DOCX) ──
                    rewritten_raw = resume.get("Rewritten Text", "")

                    # Flexible split — handles every LLM variant of the job title header:
                    #   "### 🎯 Suggested Job Titles (Based on Resume)"  — standard
                    #   "### Suggested Job Titles"                        — no emoji
                    #   "## 🎯 Suggested Job Titles"                      — two hashes
                    #   "**Suggested Job Titles (Based on Resume)**"       — bold
                    #   "Suggested Job Titles (Based on Resume)"           — plain text
                    _jt_pattern = re.compile(
                        r'(?:'
                        r'(?:#{1,4})\s*'         # 1–4 # chars (optional heading)
                        r')?'
                        r'(?:🎯\s*)?'             # optional 🎯 emoji
                        r'\*{0,2}'                # optional ** bold open
                        r'Suggested Job Titles'   # the fixed phrase
                        r'.*',                    # anything after (Based on Resume), etc.
                        re.IGNORECASE
                    )
                    _jt_split = _jt_pattern.split(rewritten_raw, maxsplit=1)
                    # Guard: only treat as a real split if the remainder contains a numbered list item
                    # Prevents false splits when "Suggested Job Titles" appears in the resume body text
                    if len(_jt_split) == 2 and re.search(r'^\s*\d+\.', _jt_split[1], re.MULTILINE):
                        resume_text_display     = _jt_split[0].strip()
                        job_suggestions_display = "### 🎯 Suggested Job Titles" + _jt_split[1]
                    else:
                        resume_text_display     = rewritten_raw
                        job_suggestions_display = ""

                    st.write(resume_text_display)

                    if job_suggestions_display:
                        st.markdown("""
                        <div style="margin:18px 0 8px;font-size:0.72rem;font-weight:700;color:#64748b;
                                    letter-spacing:0.08em;text-transform:uppercase;font-family:-apple-system,sans-serif;">
                            Job Title Suggestions (for reference only — not included in resume files)
                        </div>""", unsafe_allow_html=True)

                        # Build LinkedIn search location param — fallback to "India" if blank
                        _loc_param = urllib.parse.quote(user_location.strip()) if user_location and user_location.strip() else "India"

                        def _strip_urls(s):
                            """Remove ALL URLs and link emoji from a string."""
                            s = re.sub(r'https?://\S+', '', s)
                            s = re.sub(r'🔗', '', s)
                            return s.strip()

                        # Pre-process: strip all URLs and 🔗 from the entire block first
                        # so no URL fragment ever reaches the line-level patterns.
                        _clean_block = re.sub(r'https?://\S+', '', job_suggestions_display)
                        _clean_block = re.sub(r'🔗', '', _clean_block)
                        lines = _clean_block.split('\n')

                        # Separator group: covers —  –  -  :  →  and variants with spaces
                        # Made OPTIONAL (?) so bold-titled lines with no separator also match
                        _SEP     = r'(?:[\s]*[—–\-:→][\s]*)?' # optional separator
                        _SEP_REQ = r'[\s]*[—–\-:→][\s]*'      # required separator (for non-bold patterns)

                        def _clean_title(t):
                            """Strip ** bold markers, URLs, 🔗, and leading/trailing punctuation."""
                            t = re.sub(r'\*\*(.+?)\*\*', r'\1', t)  # remove ** wrappers
                            t = re.sub(r'\*\*', '', t)               # remove any stray **
                            t = re.sub(r'https?://\S+', '', t)
                            t = re.sub(r'🔗', '', t)
                            return t.strip()

                        items_html = ""
                        for line in lines:
                            line = line.strip()
                            if not line:
                                continue

                            title = ""
                            desc  = ""

                            # Pattern A: "1. **Title** — desc"  OR  "1. **Title**"  (numbered + bold, sep optional)
                            _ma = re.match(r'^\d+\.\s+\*\*(.+?)\*\*' + _SEP + r'(.*)', line)
                            if _ma:
                                title = _clean_title(_ma.group(1))
                                desc  = _strip_urls(_ma.group(2).strip())

                            # Pattern B: "**Title** — desc"  OR  "**Title**"  (bold only, sep optional)
                            if not title:
                                _mb = re.match(r'^\*\*(.+?)\*\*' + _SEP + r'(.*)', line)
                                if _mb:
                                    title = _clean_title(_mb.group(1))
                                    desc  = _strip_urls(_mb.group(2).strip())

                            # Pattern C: "1. Title — description"  (numbered, no bold, separator required)
                            if not title:
                                _mc = re.match(r'^\d+\.\s+(.+?)' + _SEP_REQ + r'(.*)', line)
                                if _mc:
                                    title = _clean_title(_mc.group(1))
                                    desc  = _strip_urls(_mc.group(2).strip())

                            # Pattern D: "1. Title"  (numbered, no separator, no bold — title only line)
                            if not title:
                                _md = re.match(r'^\d+\.\s+(.+)', line)
                                if _md:
                                    title = _clean_title(_md.group(1))
                                    desc  = ""

                            # Pattern E: "Title — description"  (no number, no bold, separator required)
                            if not title:
                                _me = re.match(r'^([^*\d].+?)' + _SEP_REQ + r'(.*)', line)
                                if _me:
                                    _candidate = _clean_title(_me.group(1))
                                    # Guard: max 6 words, not a header/decoration line
                                    if _candidate and len(_candidate.split()) <= 6 and not _candidate.startswith('#'):
                                        title = _candidate
                                        desc  = _strip_urls(_me.group(2).strip())

                            # Skip if nothing matched, title looks like a URL, or is a decoration line
                            if not title or title.startswith('http') or title.startswith('#'):
                                continue
                            # Skip subtitle lines like "(Based on Resume)"
                            if re.match(r'^\(.*\)$', title.strip()):
                                continue

                            desc = desc.rstrip('.')

                            encoded      = urllib.parse.quote(title)
                            linkedin_url = f"https://www.linkedin.com/jobs/search/?keywords={encoded}&location={_loc_param}"
                            link_icon = (
                                '<a href="' + linkedin_url + '" target="_blank" style="text-decoration:none;margin-left:6px;">'
                                '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" '
                                'stroke="#38bdf8" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" '
                                'style="display:inline-block;vertical-align:middle;">'
                                '<path d="M18 13v6a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V8a2 2 0 0 1 2-2h6"/>'
                                '<polyline points="15 3 21 3 21 9"/>'
                                '<line x1="10" y1="14" x2="21" y2="3"/>'
                                '</svg></a>'
                            )
                            items_html += (
                                f'<div style="margin-bottom:10px;font-size:0.88rem;color:#c9d1d9;'
                                f'display:flex;align-items:center;flex-wrap:nowrap;gap:0 4px;overflow:hidden;">'
                                f'<span style="white-space:nowrap;flex-shrink:0;">'
                                f'<b style="color:#e6edf3;">{title}</b>{link_icon}'
                                f'</span>'
                                f'{("<span style=\"color:#94a3b8;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;\"> — " + desc + "</span>") if desc else ""}'
                                f'</div>'
                            )

                        if items_html:
                            st.markdown("### 🎯 Suggested Job Titles (Based on Resume)")
                            st.markdown(f'<div style="margin-top:4px;">{items_html}</div>', unsafe_allow_html=True)
                        else:
                            # Hard fallback: LLM returned something we couldn't parse at all.
                            # Strip headers/decorations and render whatever text we got — with links.
                            _fallback_lines = [
                                l.strip() for l in _clean_block.split('\n')
                                if l.strip() and not l.strip().startswith('#') and not re.match(r'^[═─=\-]{3,}', l.strip())
                            ]
                            _fallback_html = ""
                            for _fl in _fallback_lines:
                                _ft = re.sub(r'\*\*(.+?)\*\*', r'\1', _fl)  # strip bold markers
                                _ft = _ft.strip()
                                if not _ft:
                                    continue
                                _fe = urllib.parse.quote(_ft[:60])
                                _furl = f"https://www.linkedin.com/jobs/search/?keywords={_fe}&location={_loc_param}"
                                _ficon = (
                                    '<a href="' + _furl + '" target="_blank" style="text-decoration:none;margin-left:6px;">'
                                    '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" '
                                    'stroke="#38bdf8" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" '
                                    'style="display:inline-block;vertical-align:middle;">'
                                    '<path d="M18 13v6a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V8a2 2 0 0 1 2-2h6"/>'
                                    '<polyline points="15 3 21 3 21 9"/>'
                                    '<line x1="10" y1="14" x2="21" y2="3"/>'
                                    '</svg></a>'
                                )
                                _fallback_html += f'<div style="margin-bottom:8px;font-size:0.88rem;color:#c9d1d9;">{_ft}{_ficon}</div>'
                            if _fallback_html:
                                st.markdown("### 🎯 Suggested Job Titles (Based on Resume)")
                                st.markdown(f'<div style="margin-top:4px;">{_fallback_html}</div>', unsafe_allow_html=True)

                    # ── 3-Template DOCX Download Buttons (Optimization Module — JSON data only) ──
                    st.markdown("""
                    <div style="margin:20px 0 10px;font-size:0.72rem;font-weight:700;color:#64748b;
                                letter-spacing:0.08em;text-transform:uppercase;font-family:-apple-system,sans-serif;">
                        Download Optimized Resume — Choose Template
                    </div>""", unsafe_allow_html=True)

                    optimized_data = resume.get("Optimized Resume Data", {})
                    base_name = resume['Resume Name'].split('.')[0]

                    dl_col1, dl_col2, dl_col3 = st.columns(3)

                    with dl_col1:
                        try:
                            modern_buf = generate_modern_docx(optimized_data)
                            st.download_button(
                                label="⬇ Modern (ATS)",
                                data=modern_buf,
                                file_name=f"{base_name}_modern_ats.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key=f"dl_modern_{resume['Resume Name']}",
                                help="Navy headings · Calibri · Labeled Skills block · ATS section order · Workday/Greenhouse optimized"
                            )
                        except Exception as e:
                            st.error(f"Modern template error: {e}")

                    with dl_col2:
                        try:
                            minimal_buf = generate_minimal_docx(optimized_data)
                            st.download_button(
                                label="⬇ Minimal (ATS)",
                                data=minimal_buf,
                                file_name=f"{base_name}_minimal_ats.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key=f"dl_minimal_{resume['Resume Name']}",
                                help="Pure black/white Arial · Maximum parse accuracy · Taleo/iCIMS/SmartRecruiters compatible"
                            )
                        except Exception as e:
                            st.error(f"Minimal template error: {e}")

                    with dl_col3:
                        try:
                            creative_buf = generate_creative_docx(optimized_data)
                            st.download_button(
                                label="⬇ Executive (ATS)",
                                data=creative_buf,
                                file_name=f"{base_name}_executive_ats.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key=f"dl_creative_{resume['Resume Name']}",
                                help="Teal/navy accents · ATS-safe symbols · Consistent Calibri body · Standard section labels"
                            )
                        except Exception as e:
                            st.error(f"Executive template error: {e}")

                    html_report = generate_resume_report_html(resume, user_location=user_location)
                    pdf_file = html_to_pdf_bytes(html_report)
                    st.download_button(
                        label="Download Full Analysis Report (.pdf)",
                        data=pdf_file,
                        file_name=f"{base_name}_report.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key=f"download_pdf_{resume['Resume Name']}"
                    )

    elif not uploaded_files:
        st.warning("⚠️ Please upload resumes to view dashboard analytics.")

def _sanitize_html_for_pdf(html_string):
    """
    Strip / replace CSS properties that xhtml2pdf (pisa) does not support.

    xhtml2pdf uses a very limited CSS 2.1 subset and crashes with a
    CSSParseError on any property it cannot parse — including all of the
    modern layout primitives used by the resume templates for the browser
    preview:  display:flex, flex-wrap, gap, align-items, justify-content,
    flex-direction, flex-shrink, object-fit, object-position, letter-spacing,
    text-transform (partial), border-radius on some elements, box-shadow,
    background-clip, linear-gradient(), rgba() in some positions, etc.

    Strategy:
      1. Use regex to find every inline style="..." attribute.
      2. Within each style block, strip individual property:value pairs that
         are known to crash the parser, while keeping PDF-safe ones intact.
      3. Replace display:flex → display:block  so containers still render.
      4. Replace gap:... → margin-bottom on the container (best-effort).

    This sanitisation is applied ONLY for PDF export — the browser preview
    still uses the full modern CSS.
    """
    import re as _re

    # Properties whose entire declaration should be dropped for PDF safety.
    # Each entry is a regex that matches  "property-name:value"  (no semicolon).
    STRIP_PROPS = [
        r'flex-wrap\s*:[^;]*',
        r'flex-direction\s*:[^;]*',
        r'flex-shrink\s*:[^;]*',
        r'flex-grow\s*:[^;]*',
        r'flex\s*:[^;]*',           # shorthand — must come after flex-*
        r'align-items\s*:[^;]*',
        r'align-self\s*:[^;]*',
        r'justify-content\s*:[^;]*',
        r'justify-self\s*:[^;]*',
        r'gap\s*:[^;]*',
        r'row-gap\s*:[^;]*',
        r'column-gap\s*:[^;]*',
        r'grid[^:]*:[^;]*',         # any grid-* property
        r'object-fit\s*:[^;]*',
        r'object-position\s*:[^;]*',
        r'box-shadow\s*:[^;]*',
        r'text-shadow\s*:[^;]*',
        r'background-clip\s*:[^;]*',
        r'-webkit-[^:]*:[^;]*',     # any vendor prefix
        r'-moz-[^:]*:[^;]*',
        r'-ms-[^:]*:[^;]*',
        r'transition\s*:[^;]*',
        r'transform\s*:[^;]*',
        r'animation[^:]*:[^;]*',
        r'will-change\s*:[^;]*',
        r'pointer-events\s*:[^;]*',
        r'resize\s*:[^;]*',
        r'cursor\s*:[^;]*',
        r'overflow-x\s*:[^;]*',
        r'overflow-y\s*:[^;]*',
        r'overflow\s*:\s*(?!hidden)[^;]*',   # keep overflow:hidden, drop others
        r'white-space\s*:[^;]*',
        r'word-break\s*:[^;]*',
        r'overflow-wrap\s*:[^;]*',
        r'text-overflow\s*:[^;]*',
    ]

    # background: linear-gradient(...) crashes pisa.
    # Extract the FIRST colour argument and use it as a solid background so
    # sidebar columns keep their intended colour instead of turning grey.
    GRADIENT_RE = _re.compile(
        r'background\s*:\s*linear-gradient\(([^)]*)\)\s*(?:;|$)', _re.IGNORECASE
    )
    # Matches the first hex/rgb/named colour inside the gradient args
    FIRST_COLOR_RE = _re.compile(
        r'(#[0-9a-fA-F]{3,8}|rgba?\([^)]+\))'
    )

    # display:flex → display:block  (pisa only knows block/inline/table/none)
    FLEX_DISPLAY_RE = _re.compile(
        r'display\s*:\s*flex\b', _re.IGNORECASE
    )
    INLINE_FLEX_RE = _re.compile(
        r'display\s*:\s*inline-flex\b', _re.IGNORECASE
    )

    # Viewport / dynamic units that xhtml2pdf cannot resolve → blank space
    # Strip the whole property when its value uses vh/vw/vmin/vmax
    VIEWPORT_UNIT_RE = _re.compile(
        r'[\w-]+\s*:[^;]*\d+v[hwmin][^;]*(?:;|$)', _re.IGNORECASE
    )

    # Compile all strip patterns once
    strip_re_list = [_re.compile(p, _re.IGNORECASE) for p in STRIP_PROPS]

    def _gradient_to_solid(m):
        """Replace linear-gradient(...) with its first colour as a solid bg."""
        args = m.group(1)
        colour_m = FIRST_COLOR_RE.search(args)
        colour = colour_m.group(0) if colour_m else '#f5f5f5'
        return f'background:{colour};'

    def _clean_style(style_value):
        """Clean a single style="..." value string."""
        # Replace gradients with extracted first colour (preserves sidebar colour)
        style_value = GRADIENT_RE.sub(_gradient_to_solid, style_value)
        # Strip any property whose value contains a viewport unit (vh/vw/etc.)
        # e.g. min-height:100vh causes blank pages; height:100vh same issue
        style_value = VIEWPORT_UNIT_RE.sub('', style_value)
        # Replace flex displays
        style_value = FLEX_DISPLAY_RE.sub('display:block', style_value)
        style_value = INLINE_FLEX_RE.sub('display:inline-block', style_value)
        # Strip unsupported properties
        for pat in strip_re_list:
            style_value = pat.sub('', style_value)
        # Clean up leftover semicolons / whitespace e.g. ";;  ;" → ";"
        style_value = _re.sub(r'\s*;\s*;+', ';', style_value)
        style_value = _re.sub(r'^\s*;+', '', style_value)
        style_value = style_value.strip().strip(';')
        return style_value

    # Match every  style="..."  or  style='...'  attribute in the HTML
    def _replace_style_attr(m):
        quote = m.group(1)           # ' or "
        style_content = m.group(2)
        cleaned = _clean_style(style_content)
        if not cleaned.strip():
            return ''                 # remove empty style attributes entirely
        return f'style={quote}{cleaned}{quote}'

    sanitized = _re.sub(
        r'''style=(['"])(.*?)\1''',
        _replace_style_attr,
        html_string,
        flags=_re.DOTALL | _re.IGNORECASE,
    )

    # Also strip any <style> blocks that contain flex/grid — pisa parses
    # embedded <style> tags through its CSS engine and will also crash there.
    def _clean_style_tag(m):
        css_text = m.group(1)
        css_text = GRADIENT_RE.sub('background:#f5f5f5;', css_text)
        css_text = FLEX_DISPLAY_RE.sub('display:block', css_text)
        css_text = INLINE_FLEX_RE.sub('display:inline-block', css_text)
        for pat in strip_re_list:
            css_text = pat.sub('', css_text)
        return f'<style>{css_text}</style>'

    sanitized = _re.sub(
        r'<style[^>]*>(.*?)</style>',
        _clean_style_tag,
        sanitized,
        flags=_re.DOTALL | _re.IGNORECASE,
    )

    return sanitized


def html_to_pdf_bytes(html_string):
    # NOTE: Do NOT use f-string here. Template HTML contains CSS variables like
    # {C_PRIMARY} which Python re-evaluates as f-string placeholders → CSSParseError crash.
    # Use plain string + .replace() to safely inject html_string.

    # Sanitise modern CSS that xhtml2pdf cannot parse before wrapping.
    safe_html = _sanitize_html_for_pdf(html_string)

    wrapper = """
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {
                size: A4 portrait;
                margin: 0;
            }
            html, body {
                margin: 0;
                padding: 0;
                font-size: 12pt;
                font-family: "Segoe UI", "Helvetica", sans-serif;
                line-height: 1.5;
                color: #000;
                background: #fff;
            }
            table {
                width: 100%;
                border-collapse: collapse;
            }
        </style>
    </head>
    <body>
        __HTML_CONTENT__
    </body>
    </html>
    """
    styled_html = wrapper.replace("__HTML_CONTENT__", safe_html)

    pdf_io = BytesIO()
    pisa.CreatePDF(styled_html, dest=pdf_io)
    pdf_io.seek(0)
    return pdf_io
# ══════════════════════════════════════════════════════════════════════════════
# IMPORTS — modular resume & cover letter engines
# ══════════════════════════════════════════════════════════════════════════════
from resume_builder import (
    render_template_default, render_template_modern, render_template_sidebar,
    render_template_classic, render_template_executive, render_template_timeline,
    render_template_corporate, render_template_creative_green,
    render_template_terracotta, render_template_navy_prestige,
    render_template_slate_gray, render_template_teal_impact,
    render_template_burgundy_classic, render_template_indigo_tech,
    render_template_forest_green,
    render_template_pure_white, render_template_midnight_black,
    render_template_soft_lavender, render_template_warm_sand,
    render_template_ice_blue,
    RESUME_TEMPLATES, render_resume,
    _fmt_desc, _cert_name_html,
)
from cover_letter import (
    render_cover_letter_professional, render_cover_letter_modern,
    render_cover_letter_creative, render_cover_letter_executive,
    render_cover_letter_entry_level, render_cover_letter_ats,
    COVER_LETTER_TEMPLATES, render_cover_letter,
    generate_cover_letter_from_resume_builder,
)

from collections import Counter as _Counter

# ── Filler phrase / word constants ────────────────────────────────────────────
_FILLER_PHRASES = {
    "lorem ipsum", "test test", "sample text", "placeholder",
    "your text here", "enter here", "tbd", "todo",
    "fill in", "coming soon", "to be added", "add here",
    "description here", "write here", "dummy text",
    "hello world", "foo bar", "asdf", "qwerty",
}
_FILLER_WORDS = {"placeholder", "tbd", "todo", "dummy", "example", "test", "asdf", "qwerty"}


def detect_garbage_text(text: str) -> bool:
    """
    Returns True when the text is meaningless / garbage — should earn 0 XP.

    Detection layers (cheapest first):
      1. Empty / whitespace-only
      2. Pure numeric string  ("123456789", "+91 98765")
      3. Known filler/placeholder phrases (multi-word substring match)
      4. All-same character repeated  ("aaaaaaa", "--------")
      5. Keyboard-mash: any token > 18 chars with < 12% vowels
      6. Single long token (> 20 chars) with < 15% vowels
      7. Repeated-token spam: dominant word > 55% of all tokens (≥ 3 tokens)
      8. Very low unique-word ratio: < 35% unique among 4+ word inputs
      9. Pathological token length: avg < 2.0 chars with > 60% short tokens
    """
    t = str(text).strip()
    if not t:
        return True

    # 1. Pure numeric (phone-safe: strip spaces, +, -, .)
    stripped_num = t.replace(" ", "").replace("-", "").replace("+", "").replace(".", "")
    if stripped_num.isdigit() and len(stripped_num) > 5:
        return True

    t_lower = t.lower()

    # 2. Known filler phrases (substring match)
    for phrase in _FILLER_PHRASES:
        if phrase in t_lower:
            return True

    # 3. "n/a" exact or sole token
    if t_lower.strip() in {"n/a", "na", "n.a.", "n.a"}:
        return True

    # 4. Single-word filler: whole-token match only
    t_tokens_lower = [tok.lower() for tok in t.split() if tok]
    if len(t_tokens_lower) <= 2 and any(w in _FILLER_WORDS for w in t_tokens_lower):
        return True

    # 5. All-same character ("aaaaaaa", "-------")
    stripped_chars = t.replace(" ", "")
    if len(stripped_chars) >= 3 and len(set(stripped_chars.lower())) == 1:
        return True

    tokens = [tok for tok in t.replace(",", " ").replace(";", " ").split() if tok]
    if not tokens:
        return True

    # 5b. Per-token all-same-character check: if majority of tokens are
    #     single-repeated-char words (e.g. "aaaaaa bbbb cccc"), reject.
    #     Threshold: ≥ 60% of tokens are all-same-char and len ≥ 3.
    _same_char_count = sum(
        1 for tok in tokens
        if len(tok) >= 3 and len(set(tok.lower())) == 1
    )
    if len(tokens) >= 2 and (_same_char_count / len(tokens)) >= 0.60:
        return True

    # 6. Any token > 18 chars with < 12% vowels → keyboard mash
    for tok in tokens:
        if len(tok) > 18:
            vowels = sum(1 for c in tok.lower() if c in "aeiou")
            if vowels / len(tok) < 0.12:
                return True

    # 7. Single mega-token keyboard mash (> 20 chars, < 15% vowels)
    if len(tokens) == 1 and len(tokens[0]) > 20:
        vowels = sum(1 for c in tokens[0].lower() if c in "aeiou")
        if vowels / len(tokens[0]) < 0.15:
            return True

    # 8. Repeated-token spam: one word > 55% of all tokens
    lowered = [tok.lower() for tok in tokens]
    if len(tokens) >= 3:
        most_common_count = _Counter(lowered).most_common(1)[0][1]
        if most_common_count / len(tokens) > 0.55:
            return True

    # 9. Low unique-word ratio: for 4+ word inputs, require ≥ 35% unique
    if len(tokens) >= 4:
        if calculate_unique_word_ratio(t) < 0.35:
            return True

    # 10. Pathological token length pattern
    avg_len = sum(len(tok) for tok in tokens) / len(tokens)
    short_ratio = sum(1 for tok in tokens if len(tok) <= 2) / len(tokens)
    if avg_len < 2.0 and short_ratio > 0.60:
        return True

    return False


def is_meaningful_text(text: str, min_words: int = 5) -> bool:
    """
    Returns True when text passes quality gates and meets the minimum word count.
    Inverse of detect_garbage_text with an additional word-count floor.

    Parameters
    ----------
    text      : the raw string to validate
    min_words : minimum number of space-separated tokens required (default 5)

    Usage
    -----
    Suitable for experience / project / education description fields.
    Short-form fields (skills, cert name) should use detect_garbage_text directly.
    """
    t = str(text).strip()
    if detect_garbage_text(t):
        return False
    tokens = [tok for tok in t.split() if tok]
    return len(tokens) >= min_words


def detect_repeated_words(text: str, threshold: float = 0.55) -> bool:
    """
    Returns True when a single word dominates more than `threshold` fraction
    of all tokens — indicating repetitive / spam input.

    Parameters
    ----------
    text      : raw input string
    threshold : fraction above which dominant word is considered spam (default 0.55)
    """
    tokens = [tok.lower() for tok in str(text).split() if tok]
    if len(tokens) < 3:
        return False
    most_common_count = _Counter(tokens).most_common(1)[0][1]
    return (most_common_count / len(tokens)) > threshold


def calculate_unique_word_ratio(text: str) -> float:
    """
    Returns the ratio of unique lowercase words to total words (0.0 – 1.0).
    A ratio of 1.0 means every word is unique; 0.0 means only one distinct word.
    Returns 0.0 for empty / single-token input.

    Example
    -------
    "hello world hello" → 2 unique / 3 total → 0.667
    """
    tokens = [tok.lower() for tok in str(text).split() if tok]
    if not tokens:
        return 0.0
    return round(len(set(tokens)) / len(tokens), 4)


# ── Internal quality scorer (used by all section scorers) ─────────────────────

def _text_quality_score(text: str, min_words: int = 3) -> float:
    """
    Returns a quality multiplier 0.0–1.0 for free-text fields.
    Garbage inputs always score 0.0.
    Rewards: sufficient word count + unique-word diversity.
    """
    t = str(text).strip()
    if not t:
        return 0.0
    if detect_garbage_text(t):
        return 0.0

    tokens = [tok for tok in t.split() if tok]
    word_count = len(tokens)

    if word_count < min_words:
        return 0.0

    unique_ratio = calculate_unique_word_ratio(t)
    diversity_mult = 1.0
    if word_count >= 6 and unique_ratio < 0.50:
        diversity_mult = max(0.30, min(1.0, unique_ratio * 1.5))

    if word_count < 5:
        base = 0.25
    elif word_count < 15:
        base = 0.55
    elif word_count < 30:
        base = 0.80
    elif word_count < 50:
        base = 0.92
    else:
        base = 1.0

    return round(min(base * diversity_mult, 1.0), 3)


def _desc_score(text: str) -> float:
    """
    Score a long-form description field (experience / project / education / cert).
    Combines word-count quality with character-length signal.
    Garbage → 0.0 always.
    """
    t = str(text).strip()
    if not t or detect_garbage_text(t):
        return 0.0

    char_len = len(t)
    word_score = _text_quality_score(t, min_words=4)

    if char_len < 30:
        char_tier = 0.20
    elif char_len < 80:
        char_tier = 0.55
    elif char_len < 180:
        char_tier = 0.82
    else:
        char_tier = 1.0

    return round(min((word_score * 0.60) + (char_tier * 0.40), 1.0), 3)


def _count_valid_tokens(raw_str: str) -> int:
    """Return count of unique, non-empty, non-garbage comma-separated tokens."""
    if not raw_str:
        return 0
    seen = set()
    count = 0
    for tok in raw_str.split(","):
        tok = tok.strip()
        if not tok:
            continue
        tok_norm = tok.lower()
        if tok_norm in seen:
            continue
        seen.add(tok_norm)
        if not detect_garbage_text(tok):
            count += 1
    return count


# ── Industry-standard XP weights ──────────────────────────────────────────────
# Inspired by LinkedIn Profile Strength, Indeed Resume Score, and Jobscan:
#   Experience    → 30 XP  (highest — core employability signal)
#   Projects      → 22 XP  (second — demonstrates practical skills)
#   Skills & More → 16 XP  (medium — keyword matching & breadth)
#   Education     → 14 XP  (medium — credential verification)
#   Summary       →  7 XP  (differentiator — first impression text)
#   Certificates  →  5 XP  (bonus — validated expertise)
#   Personal Info →  4 XP  (baseline — completeness check)
#   Contact       →  2 XP  (small — reachability signal)
#                  ─────
#   TOTAL MAX     → 100 XP
XP_WEIGHTS = {
    "Experience":    30,
    "Projects":      22,
    "Skills & More": 16,
    "Education":     14,
    "Summary":        7,
    "Certificates":   5,
    "Personal Info":  4,
    "Contact":        2,
}
XP_TOTAL_MAX = sum(XP_WEIGHTS.values())  # 100


def score_experience_section(experience_entries: list) -> float:
    """
    Score the Experience section. Returns float 0.0–1.0.

    Scoring per entry:
      Required (40%): job title + company — both validated (non-garbage)
      Quality  (60%): duration 20% + rich description 80%
    Cap: 1 entry → max 0.80; 2+ entries can reach 1.0.
    """
    if not experience_entries:
        return 0.0
    n = len(experience_entries)
    total = 0.0
    for e in experience_entries:
        title   = str(e.get("title",       "")).strip()
        company = str(e.get("company",     "")).strip()
        dur     = str(e.get("duration",    "")).strip()
        desc    = str(e.get("description", "")).strip()
        if not title and not company:
            continue
        title_ok   = bool(title)   and not detect_garbage_text(title)
        company_ok = bool(company) and not detect_garbage_text(company)
        req   = (float(title_ok) + float(company_ok)) / 2
        qual  = (float(bool(dur)) * 0.20) + (_desc_score(desc) * 0.80)
        total += (req * 0.40) + (qual * 0.60)
    avg = total / n
    if n == 1:
        avg = min(avg, 0.80)
    return round(min(avg, 1.0), 3)


def score_education_section(education_entries: list) -> float:
    """
    Score the Education section. Returns float 0.0–1.0.

    Scoring per entry:
      Required (40%): institution + degree — both validated
      Quality  (60%): year 20% + academic details 80%
    Cap: 1 entry → max 0.85; 2+ can reach 1.0.
    """
    if not education_entries:
        return 0.0
    n = len(education_entries)
    total = 0.0
    for e in education_entries:
        inst   = str(e.get("institution", "")).strip()
        degree = str(e.get("degree",      "")).strip()
        if isinstance(degree, list):
            degree = ", ".join(degree)
        year   = str(e.get("year",    "")).strip()
        det    = str(e.get("details", "")).strip()
        if not inst and not degree:
            continue
        inst_ok   = bool(inst)   and not detect_garbage_text(inst)
        degree_ok = bool(degree) and not detect_garbage_text(degree)
        req  = (float(inst_ok) + float(degree_ok)) / 2
        qual = (float(bool(year)) * 0.20) + (_desc_score(det) * 0.80)
        total += (req * 0.40) + (qual * 0.60)
    avg = total / n
    if n == 1:
        avg = min(avg, 0.85)
    return round(min(avg, 1.0), 3)


def score_project_section(project_entries: list, project_links: list = None) -> float:
    """
    Score the Projects section. Returns float 0.0–1.0.

    Scoring per entry:
      Required (35%): title (validated) + tech stack (≥ 1 valid token)
      Quality  (65%): duration 15% + rich description 85%
    Cap: 1 project → max 0.75; 2 → max 0.90; 3+ → 1.0.
    """
    if not project_entries:
        return 0.0
    n = len(project_entries)
    total = 0.0
    for e in project_entries:
        title = str(e.get("title",       "")).strip()
        tech  = str(e.get("tech",        "")).strip()
        dur   = str(e.get("duration",    "")).strip()
        desc  = str(e.get("description", "")).strip()
        if not title:
            continue
        title_ok  = not detect_garbage_text(title)
        tech_cnt  = _count_valid_tokens(tech)
        tech_sc   = min(tech_cnt / 2.0, 1.0)
        req  = (float(title_ok) + tech_sc) / 2
        qual = (float(bool(dur)) * 0.15) + (_desc_score(desc) * 0.85)
        total += (req * 0.35) + (qual * 0.65)
    avg = total / n
    if n == 1:
        avg = min(avg, 0.75)
    elif n == 2:
        avg = min(avg, 0.90)
    return round(min(avg, 1.0), 3)


def score_certificate_section(certificate_entries: list) -> float:
    """
    Score the Certificates section. Returns float 0.0–1.0.

    Scoring per entry:
      Required (30%): certificate name (validated — non-garbage)
      Quality  (70%): link 20% + duration 20% + description 60%
    1 complete cert CAN reach 1.0 (certs are optional, 1 complete = excellent).
    """
    if not certificate_entries:
        return 0.0
    n = len(certificate_entries)
    total = 0.0
    for e in certificate_entries:
        name = str(e.get("name",        "")).strip()
        link = str(e.get("link",        "")).strip()
        dur  = str(e.get("duration",    "")).strip()
        desc = str(e.get("description", "")).strip()
        if not name:
            continue
        name_ok = not detect_garbage_text(name)
        req  = float(name_ok)
        qual = (float(bool(link)) * 0.20) + (float(bool(dur)) * 0.20) + (_desc_score(desc) * 0.60)
        total += (req * 0.30) + (qual * 0.70)
    return round(min(total / n, 1.0), 3)


def score_skills_section(
    skills: str,
    soft_skills: str = "",
    interests: str = "",
    languages: str = "",
) -> float:
    """
    Score the Skills & More composite section. Returns float 0.0–1.0.

    Sub-weights: Skills 50% | Soft Skills 20% | Interests 20% | Languages 10%
    Each sub-score uses _count_valid_tokens for garbage rejection.
    Skills: 0 → 0.0, 1 → 0.20, 2 → 0.45, 3 → 0.65, 4 → 0.82, 5+ → 1.0
    Soft/Interests: 0 → 0.0, 1 → 0.35, 2 → 0.70, 3+ → 1.0
    Languages: 0 → 0.0, 1 → 0.50, 2+ → 1.0
    """
    skill_count = _count_valid_tokens(skills)
    if skill_count == 0:    sub_skills = 0.0
    elif skill_count == 1:  sub_skills = 0.20
    elif skill_count == 2:  sub_skills = 0.45
    elif skill_count == 3:  sub_skills = 0.65
    elif skill_count == 4:  sub_skills = 0.82
    else:                   sub_skills = 1.0

    interest_count = _count_valid_tokens(interests)
    sub_interests = 0.0 if interest_count == 0 else (0.35 if interest_count == 1 else (0.70 if interest_count == 2 else 1.0))

    soft_count = _count_valid_tokens(soft_skills)
    sub_soft = 0.0 if soft_count == 0 else (0.35 if soft_count == 1 else (0.70 if soft_count == 2 else 1.0))

    lang_count = _count_valid_tokens(languages)
    sub_lang = 0.0 if lang_count == 0 else (0.50 if lang_count == 1 else 1.0)

    return round(
        (sub_skills * 0.50) + (sub_interests * 0.20) + (sub_soft * 0.20) + (sub_lang * 0.10),
        3
    )


def calculate_resume_xp(section_fills: dict, weights: dict = None) -> tuple:
    """
    Compute total resume XP from section fill scores and weights.
    XP is capped at XP_TOTAL_MAX (100).

    Parameters
    ----------
    section_fills : dict[section_name → float 0.0–1.0]
    weights       : dict[section_name → int] — defaults to module XP_WEIGHTS

    Returns
    -------
    (raw_xp: float, xp: int, pct: int)

    Usage
    -----
    xp_score = calculate_resume_xp(st.session_state)  # legacy call
    raw, xp, pct = calculate_resume_xp(fills, XP_WEIGHTS)
    """
    if weights is None:
        weights = XP_WEIGHTS
    raw_xp = sum(section_fills.get(k, 0.0) * weights.get(k, 0) for k in weights)
    raw_xp = min(raw_xp, XP_TOTAL_MAX)
    xp_int = int(round(raw_xp))
    pct    = int(round((raw_xp / XP_TOTAL_MAX) * 100))
    return raw_xp, xp_int, pct


# ── Public aliases used in external imports / tests ───────────────────────────
score_experience_section.__module__ = __name__
score_project_section.__module__    = __name__
score_education_section.__module__  = __name__
score_certificate_section.__module__ = __name__
score_skills_section.__module__     = __name__

# ══════════════════════════════════════════════════════════════════════════════
# END OF MODULE-LEVEL XP ENGINE
# ══════════════════════════════════════════════════════════════════════════════

# Import necessary modules first
import streamlit as st
import time

# Tab setup (assuming this is within a tab2 context)
with tab2:
    st.session_state.active_tab = "Resume Builder"

    # ---------- Title with Blue Glassmorphism + Shine ----------
    st.markdown("""
    <style>
    .glass-title {
        background: rgba(10, 20, 40, 0.5);
        border-radius: 20px;
        padding: 20px;
        backdrop-filter: blur(14px);
        box-shadow: 0 8px 32px rgba(0, 200, 255, 0.25);
        border: 1px solid rgba(0, 200, 255, 0.3);
        text-align: center;
        position: relative;
        overflow: hidden;
    }
    .glass-title h2 {
        color: #4da6ff;
        margin: 0;
        text-shadow: 0 0 12px rgba(0,200,255,0.7);
        font-weight: 600;
    }
    .glass-title::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.18) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .glass-title:hover::before {
        left: 100%;
        top: 100%;
    }
    </style>

    <div class="glass-title">
        <h2>🧾 Advanced Resume Builder</h2>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<hr style='border-top: 2px solid rgba(0,200,255,0.4);'>", unsafe_allow_html=True)

    # ---------- Anti-Flicker / Smooth Rerun CSS ----------
    st.markdown("""
        <style>
        /* Prevent white flash and blinking on Streamlit reruns */
        [data-testid="stAppViewContainer"],
        [data-testid="stVerticalBlock"],
        [data-testid="stForm"],
        [data-testid="stSidebar"],
        section.main > div {
            transition: opacity 0.15s ease-in-out !important;
        }

        /* Suppress the brief layout jump when widgets remount */
        iframe, [data-testid="stIFrame"] {
            transition: none !important;
        }

        /* Prevent scrollbar flicker */
        html {
            overflow-y: scroll !important;
            scrollbar-gutter: stable !important;
        }

        /* Smooth button press — no jump */
        button[kind="formSubmit"],
        button[kind="secondary"],
        .stButton > button {
            transition: background-color 0.2s ease, box-shadow 0.2s ease, transform 0.1s ease !important;
        }
        .stButton > button:active {
            transform: scale(0.98) !important;
        }

        /* Prevent input field flicker on focus/blur */
        input, textarea, select {
            transition: border-color 0.15s ease, box-shadow 0.15s ease !important;
        }

        /* Prevent expander flicker */
        details summary {
            transition: background 0.2s ease !important;
        }

        /* Prevent layout shift during widget remounts */
        [data-testid="stVerticalBlock"] > div {
            min-height: 0 !important;
        }

        /* Smooth fade-in for newly rendered blocks */
        @keyframes fadeInBlock {
            from { opacity: 0.6; }
            to   { opacity: 1; }
        }
        [data-testid="stVerticalBlock"] {
            animation: fadeInBlock 0.12s ease-out !important;
        }
        </style>

        <script>
        (function() {
            var lastScrollY = 0;
            var ticking = false;
            var observer = new MutationObserver(function() {
                if (!ticking) {
                    requestAnimationFrame(function() {
                        if (Math.abs(window.scrollY - lastScrollY) > 200) {
                            window.scrollTo({ top: lastScrollY, behavior: 'instant' });
                        }
                        ticking = false;
                    });
                    ticking = true;
                }
            });
            observer.observe(document.body, { childList: true, subtree: false });
            window.addEventListener('scroll', function() {
                lastScrollY = window.scrollY;
            }, { passive: true });
        })();
        </script>
    """, unsafe_allow_html=True)

    # ---------- Global Styles (Glassmorphism + Glow + Shine) ----------
    st.markdown("""
        <style>
        /* File uploader */
        .uploadedFile { 
            background: rgba(10, 20, 40, 0.6) !important;
            border: 1px solid rgba(0,200,255,0.5) !important;
            border-radius: 14px !important;
            color: #cce6ff !important;
            box-shadow: 0 0 12px rgba(0,200,255,0.3) !important;
        }

        /* Sidebar expander style */
        .streamlit-expanderHeader {
            background: rgba(10, 20, 40, 0.45);
            border-radius: 12px;
            color: #4da6ff !important;
            font-weight: bold;
            backdrop-filter: blur(12px);
            box-shadow: 0 4px 12px rgba(0,200,255,0.25);
            transition: all 0.3s ease-in-out;
        }
        .streamlit-expanderHeader:hover {
            background: rgba(0, 200, 255, 0.12);
            box-shadow: 0 0 16px rgba(0,200,255,0.4);
        }
        .streamlit-expanderContent {
            background: rgba(10, 20, 40, 0.45);
            border-radius: 10px;
            padding: 8px;
            color: #e6f7ff;
        }

        /* Selectbox */
        div[data-baseweb="select"] {
            background: rgba(10, 20, 40, 0.35);
            border: 1px solid rgba(0, 200, 255, 0.6);
            border-radius: 12px;
            color: #e6f7ff;
            backdrop-filter: blur(14px);
            box-shadow: 0 0 10px rgba(0,200,255,0.3);
        }

        /* Buttons with Shine Effect */
        div.stButton > button {
            position: relative;
            background: rgba(10, 20, 40, 0.35);
            border: 1px solid rgba(0, 200, 255, 0.6);
            color: #e6f7ff;
            border-radius: 14px;
            padding: 10px 20px;
            font-size: 15px;
            font-weight: 500;
            backdrop-filter: blur(16px);
            box-shadow: 0 0 12px rgba(0, 200, 255, 0.35),
                        inset 0 0 20px rgba(0, 200, 255, 0.05);
            overflow: hidden;
            transition: all 0.3s ease-in-out;
        }
        div.stButton > button::before {
            content: "";
            position: absolute;
            top: -50%;
            left: -50%;
            width: 200%;
            height: 200%;
            background: linear-gradient(
                120deg,
                rgba(255,255,255,0.15) 0%,
                rgba(255,255,255,0.05) 40%,
                transparent 60%
            );
            transform: rotate(25deg);
            transition: all 0.6s;
        }
        div.stButton > button:hover::before {
            left: 100%;
            top: 100%;
        }
        div.stButton > button:hover {
            background: rgba(0, 200, 255, 0.12);
            box-shadow: 0 0 20px rgba(0, 200, 255, 0.65),
                        inset 0 0 25px rgba(0, 200, 255, 0.15);
            transform: translateY(-2px);
        }
        div.stButton > button:active {
            transform: scale(0.95);
            box-shadow: 0 0 10px rgba(0, 200, 255, 0.45);
        }
        </style>
    """, unsafe_allow_html=True)

    # 🎨 Template Selection — visual card grid
    st.markdown("""
    <style>
    div[data-testid="stHorizontalBlock"] .tpl-card-wrap { padding: 4px; }
    </style>
    """, unsafe_allow_html=True)

    TEMPLATE_META = [
        ("Default (Professional)",        "#2f4f6f", "#e8f0fe"),
        ("Modern Minimal",                 "#0d9488", "#f0fdfa"),
        ("Elegant Sidebar",                "#7c3aed", "#f5f3ff"),
        ("Classic Clean (Single Column)",  "#374151", "#f9fafb"),
        ("Executive (Single Column)",      "#1e3a5f", "#eff6ff"),
        ("Timeline (Single Column)",       "#b45309", "#fffbeb"),
        ("Corporate Blue (Two Column)",    "#1d4ed8", "#eff6ff"),
        ("Creative Green (Two Column)",    "#166534", "#f0fdf4"),
        ("Warm Terracotta (Two Column)",   "#c2410c", "#fff7ed"),
        ("Navy Prestige (Two Column)",     "#1e3a5f", "#f0f4ff"),
        ("Slate Gray (Single Column)",     "#475569", "#f8fafc"),
        ("Teal Impact (Two Column)",       "#0f766e", "#f0fdfa"),
        ("Burgundy Classic (Single Column)","#881337","#fff1f2"),
        ("Indigo Tech (Two Column)",       "#4338ca", "#eef2ff"),
        ("Forest Green (Single Column)",   "#14532d", "#f0fdf4"),
        # ── 6 new premium templates ──────────────────────────────────────────
        ("Pure White (Single Column)",     "linear-gradient(to right,#ffffff 70%,#111111 100%)", "#ffffff"),
        ("Midnight Black (Single Column)", "#f59e0b", "#111827"),
        ("Soft Lavender (Single Column)",  "#6366f1", "#f5f3ff"),
        ("Warm Sand (Single Column)",      "#b45309", "#fdf8f0"),
        ("Ice Blue (Single Column)",       "#0369a1", "#eff9ff"),
    ]
    TEMPLATE_NAMES = [t[0] for t in TEMPLATE_META]

    if "selected_template_name" not in st.session_state:
        st.session_state["selected_template_name"] = TEMPLATE_NAMES[0]

    st.markdown("<div style='margin:18px 0 8px;font-size:14px;font-weight:600;color:#93c5fd;'>🎨 Choose Resume Template</div>", unsafe_allow_html=True)

    # Show 5 cards per row
    _tpl_rows = [TEMPLATE_META[i:i+5] for i in range(0, len(TEMPLATE_META), 5)]
    for _row in _tpl_rows:
        _cols = st.columns(len(_row))
        for _ci, (_tname, _color, _light) in enumerate(_row):
            with _cols[_ci]:
                _is_sel = st.session_state["selected_template_name"] == _tname
                _border = "2px solid #4da6ff" if _is_sel else "1px solid rgba(0,180,255,0.15)"
                _glow   = "box-shadow: 0 0 12px rgba(77,166,255,0.45);" if _is_sel else ""
                _is_light_swatch = _color.startswith("linear-gradient") or _color in ("#fdf8f0", "#eff9ff", "#fff0f6", "#f5f3ff", "#f8fafc", "#fffbeb")
                _swatch_extra = "border:1px solid rgba(255,255,255,0.25);" if _is_light_swatch else ""
                st.markdown(
                    f"<div style='background:rgba(13,20,40,0.6);border:{_border};border-radius:10px;"
                    f"padding:8px 6px 6px;text-align:center;{_glow}'>"
                    f"<div style='height:28px;border-radius:6px;background:{_color};margin-bottom:6px;{_swatch_extra}'></div>"
                    f"<div style='font-size:9.5px;color:{'#93c5fd' if _is_sel else '#6b7280'};font-weight:{'700' if _is_sel else '500'};line-height:1.3;'>{_tname}</div>"
                    f"</div>",
                    unsafe_allow_html=True,
                )
                if st.button("✓" if _is_sel else "Select", key=f"tpl_btn_{_tname}", use_container_width=True):
                    if st.session_state["selected_template_name"] != _tname:
                        st.session_state["selected_template_name"] = _tname
                        st.rerun()

    selected_template = st.session_state["selected_template_name"]

    # 📸 Upload profile photo
    st.markdown("<div style='margin:18px 0 6px;font-size:14px;font-weight:600;color:#93c5fd;'>📸 Profile Photo</div>", unsafe_allow_html=True)
    _img_col1, _img_col2 = st.columns([3, 1])
    with _img_col1:
        uploaded_image = st.file_uploader("Upload a Profile Image (PNG/JPG, square preferred)", type=["png", "jpg", "jpeg"], key="profile_img_upload", label_visibility="collapsed")
        # ── FIX: Encode and store as soon as a new file is uploaded, then rerun
        # so the preview column (already rendered above) refreshes immediately.
        if uploaded_image is not None:
            import base64 as _base64
            _new_encoded = _base64.b64encode(uploaded_image.read()).decode()
            if _new_encoded != st.session_state.get("encoded_profile_image"):
                st.session_state["encoded_profile_image"] = _new_encoded
                st.rerun()
        # ── FIX: "Remove Photo" button clears session state so image disappears.
        if st.session_state.get("encoded_profile_image"):
            if st.button("🗑️ Remove Photo", key="remove_profile_photo"):
                st.session_state.pop("encoded_profile_image", None)
                st.rerun()
    with _img_col2:
        if st.session_state.get("encoded_profile_image"):
            st.markdown(
                f"<img src='data:image/png;base64,{st.session_state['encoded_profile_image']}' "
                f"class='photo-preview' />",
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                "<div style='width:72px;height:72px;border-radius:50%;background:#1e2535;"
                "border:2px dashed #374151;display:flex;align-items:center;justify-content:center;"
                "font-size:22px;margin:4px auto;'>👤</div>",
                unsafe_allow_html=True,
            )
    profile_img_html = ""

    if st.session_state.get("encoded_profile_image"):
        encoded_image = st.session_state["encoded_profile_image"]
        profile_img_html = f"""
        <div style="display: flex; justify-content: flex-end; margin-top: 20px;">
            <img src="data:image/png;base64,{encoded_image}" alt="Profile Photo"
                 style="
                    width: 140px;
                    height: 140px;
                    border-radius: 50%;
                    object-fit: cover;
                    object-position: center;
                    border: 4px solid rgba(255,255,255,0.6);
                    box-shadow:
                        0 0 0 3px #4da6ff,
                        0 8px 25px rgba(77, 166, 255, 0.3),
                        0 4px 15px rgba(0, 0, 0, 0.15);
                    transition: transform 0.3s ease-in-out;
                "
                onmouseover="this.style.transform='scale(1.07)'"
                onmouseout="this.style.transform='scale(1)'"
             />
        </div>
        """
    else:
        st.markdown("<div style='font-size:12px;color:#4b5563;margin-top:4px;'>📸 Upload a clear, front-facing photo (square or portrait preferred)</div>", unsafe_allow_html=True)

    # ---------------- Session State Defaults ----------------
    fields = ["name", "email", "phone", "linkedin", "location", "portfolio", "summary",
              "skills", "languages", "interests", "Softskills", "job_title"]
    for f in fields:
        st.session_state.setdefault(f, "")

    st.session_state.setdefault("experience_entries", [{"title": "", "company": "", "duration": "", "description": ""}])
    st.session_state.setdefault("education_entries", [{"degree": "", "institution": "", "year": "", "details": ""}])
    st.session_state.setdefault("project_entries", [{"title": "", "tech": "", "duration": "", "description": ""}])
    st.session_state.setdefault("project_links", [])
    st.session_state.setdefault("certificate_links", [{"name": "", "link": "", "duration": "", "description": ""}])
    st.session_state.setdefault("form_key_counter", 0)

    # ─────────────────────────────────────────────────────────────────────────
    # GAMIFIED SIDEBAR
    # ─────────────────────────────────────────────────────────────────────────
    def render_gamified_sidebar(ss, fk):
        """
        Renders a fully gamified sidebar with:
        - XP counter + rank badge
        - Master progress bar
        - Streak dot row
        - Per-section rows with inline SVG icons, mini-bars, and check circles
        - Stats footer (Done / XP / % Complete)
        - Section add/delete controls (Experience, Education, Projects, Certificates)
        All icons are inline SVG — no emojis.
        """

        # ══════════════════════════════════════════════════════════════════════
        # PRE-SCORING SYNC: push live widget values into entry dicts BEFORE
        # scoring runs. This eliminates the double-press/stale-score bug that
        # occurs because the sidebar renders BEFORE the form widgets, so the
        # entry dicts still hold the previous run's committed values.
        # By reading directly from session_state widget keys here, we always
        # score the text the user has typed RIGHT NOW.
        # ══════════════════════════════════════════════════════════════════════
        def _sync_entries():
            # Sync simple scalar fields
            for widget_key, ss_key in [
                (f"name_input_{fk}",    "name"),
                (f"email_input_{fk}",   "email"),
                (f"phone_input_{fk}",   "phone"),
                (f"loc_input_{fk}",     "location"),
                (f"job_input_{fk}",     "job_title"),
                (f"summary_input_{fk}", "summary"),
                (f"skills_input_{fk}",  "skills"),
                (f"lang_input_{fk}",    "languages"),
                (f"int_input_{fk}",     "interests"),
                (f"soft_input_{fk}",    "Softskills"),
                (f"ln_input_{fk}",      "linkedin"),
                (f"phone_input_{fk}",   "phone"),
            ]:
                if widget_key in ss:
                    ss[ss_key] = ss[widget_key]

            # Sync experience entries
            entries = ss.get("experience_entries", [])
            n = len(entries)
            for i, e in enumerate(entries):
                for widget_suffix, dict_key in [
                    (f"title_{i}_{n}_{fk}",       "title"),
                    (f"company_{i}_{n}_{fk}",     "company"),
                    (f"duration_{i}_{n}_{fk}",    "duration"),
                    (f"description_{i}_{n}_{fk}", "description"),
                ]:
                    if widget_suffix in ss:
                        e[dict_key] = ss[widget_suffix]

            # Sync education entries
            entries = ss.get("education_entries", [])
            n = len(entries)
            for i, e in enumerate(entries):
                for widget_suffix, dict_key in [
                    (f"degree_{i}_{n}_{fk}",      "degree"),
                    (f"institution_{i}_{n}_{fk}", "institution"),
                    (f"edu_year_{i}_{n}_{fk}",    "year"),
                    (f"edu_details_{i}_{n}_{fk}", "details"),
                ]:
                    if widget_suffix in ss:
                        e[dict_key] = ss[widget_suffix]

            # Sync project entries
            entries = ss.get("project_entries", [])
            n = len(entries)
            for i, e in enumerate(entries):
                for widget_suffix, dict_key in [
                    (f"proj_title_{i}_{n}_{fk}",    "title"),
                    (f"proj_tech_{i}_{n}_{fk}",     "tech"),
                    (f"proj_duration_{i}_{n}_{fk}", "duration"),
                    (f"proj_desc_{i}_{n}_{fk}",     "description"),
                ]:
                    if widget_suffix in ss:
                        e[dict_key] = ss[widget_suffix]

            # Sync certificate entries
            entries = ss.get("certificate_links", [])
            n = len(entries)
            for i, e in enumerate(entries):
                for widget_suffix, dict_key in [
                    (f"cert_name_{i}_{n}_{fk}",        "name"),
                    (f"cert_link_{i}_{n}_{fk}",        "link"),
                    (f"cert_duration_{i}_{n}_{fk}",    "duration"),
                    (f"cert_description_{i}_{n}_{fk}", "description"),
                ]:
                    if widget_suffix in ss:
                        e[dict_key] = ss[widget_suffix]

        _sync_entries()

        # ── SVG icon library ──────────────────────────────────────────────────
        SVG = {
            "personal": '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><circle cx="8" cy="5" r="2.5"/><path d="M2.5 13.5c0-3 2.5-4.5 5.5-4.5s5.5 1.5 5.5 4.5"/></svg>',
            "summary":  '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><rect x="2" y="2" width="12" height="12" rx="2"/><line x1="5" y1="6" x2="11" y2="6"/><line x1="5" y1="9" x2="9" y2="9"/></svg>',
            "exp":      '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><rect x="2" y="6" width="12" height="8" rx="1.5"/><path d="M5 6V4.5A2.5 2.5 0 0 1 11 4.5V6"/><line x1="8" y1="9" x2="8" y2="11"/></svg>',
            "edu":      '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><path d="M8 2L14 5.5 8 9 2 5.5Z"/><path d="M4.5 7.5V11.5c0 0 1.5 1.5 3.5 1.5s3.5-1.5 3.5-1.5V7.5"/><line x1="14" y1="5.5" x2="14" y2="9"/></svg>',
            "projects": '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><polyline points="2,12 6,7 9,10 12,5 14,7"/><circle cx="14" cy="4" r="1.2" fill="currentColor" stroke="none"/></svg>',
            "skills":   '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><path d="M10.5 2.5l1 1-7 7-1-1z"/><path d="M12 4l1.5-1.5-1-1L11 3"/><path d="M3 11l-0.5 2 2-0.5"/></svg>',
            "certs":    '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><circle cx="8" cy="7" r="3.5"/><path d="M5.5 10L4 14l4-2 4 2-1.5-4"/></svg>',
            "contact":  '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><path d="M2 3h12v9a1 1 0 01-1 1H3a1 1 0 01-1-1V3z"/><polyline points="2,3 8,8.5 14,3"/></svg>',
            "add":      '<svg width="12" height="12" viewBox="0 0 12 12" fill="none" stroke="currentColor" stroke-width="1.6" stroke-linecap="round"><line x1="6" y1="2" x2="6" y2="10"/><line x1="2" y1="6" x2="10" y2="6"/></svg>',
            "remove":   '<svg width="12" height="12" viewBox="0 0 12 12" fill="none" stroke="currentColor" stroke-width="1.6" stroke-linecap="round"><line x1="2" y1="6" x2="10" y2="6"/></svg>',
        }

        # ══════════════════════════════════════════════════════════════════════
        # PARTIAL-FILL SCORING
        # Each section returns a float 0.0–1.0 based on how many of its
        # required fields the user has actually filled.
        # The mini-bar width, XP, and master bar all reflect partial progress.
        # A section is "complete" (check circle on) only at 1.0.
        # ══════════════════════════════════════════════════════════════════════

        def _wv(widget_key, fallback_key=""):
            """Live widget value → stored session value → empty string."""
            v = ss.get(widget_key, "")
            if not v and fallback_key:
                v = ss.get(fallback_key, "")
            return str(v).strip()

        def _filled(*values):
            """Count how many of the given values are non-empty."""
            return sum(1 for v in values if str(v).strip())

        # ── Personal Info: 5 fields, each worth 0.2 ───────────────────────────
        pi_name     = _wv(f"name_input_{fk}",  "name")
        pi_email    = _wv(f"email_input_{fk}",  "email")
        pi_phone    = _wv(f"phone_input_{fk}",  "phone")
        pi_location = _wv(f"loc_input_{fk}",    "location")
        pi_jobtitle = _wv(f"job_input_{fk}",    "job_title")
        _fill_personal = round(_filled(pi_name, pi_email, pi_phone, pi_location, pi_jobtitle) / 5, 2)

        # ── Summary: quality-gated length scoring ─────────────────────────────
        # Garbage text earns 0 regardless of length. Valid text is scored by
        # character length tiers, rewarding rich professional summaries.
        _summary_text = _wv(f"summary_input_{fk}", "summary")
        _summary_len  = len(_summary_text)
        if _summary_len == 0 or detect_garbage_text(_summary_text):
            _fill_summary = 0.0
        elif _summary_len < 40:
            _fill_summary = 0.25   # started but too short
        elif _summary_len < 100:
            _fill_summary = 0.60   # decent but brief
        elif _summary_len < 200:
            _fill_summary = 0.85   # good
        else:
            _fill_summary = 1.0    # full — rich summary

        # ══════════════════════════════════════════════════════════════════════
        # TEXT QUALITY VALIDATION ENGINE — delegates to module-level functions
        # All logic lives at module scope for reuse; these are local aliases.
        # ══════════════════════════════════════════════════════════════════════
        _is_low_quality_text  = detect_garbage_text      # module-level
        _is_gibberish         = detect_garbage_text      # backwards-compat alias

        # ══════════════════════════════════════════════════════════════════════
        # LIVE VALUE READER
        # ══════════════════════════════════════════════════════════════════════

        def _get_val(ss, widget_key, entry, stored_key, fk):
            """
            Read the most up-to-date value for a field.
            Priority:
              1. Live Streamlit widget value (ss[widget_key])
              2. Stored entry dict value (entry[stored_key])
            """
            live = ss.get(widget_key, "")
            if live:
                return str(live).strip()
            stored = entry.get(stored_key, "")
            return str(stored).strip()

        # XP_WEIGHTS and XP_TOTAL_MAX are defined at module level — use them directly.

        # ── Skills & More — delegates to module-level score_skills_section ──────
        _skills_raw    = _wv(f"skills_input_{fk}", "skills")
        _interests_raw = _wv(f"int_input_{fk}",    "interests")
        _soft_raw      = _wv(f"soft_input_{fk}",   "Softskills")
        _lang_raw      = _wv(f"lang_input_{fk}",   "languages")
        # Keep individual counts for feedback tips
        _skill_count   = _count_valid_tokens(_skills_raw)
        _soft_count    = _count_valid_tokens(_soft_raw)
        _fill_skills   = score_skills_section(
            skills=_skills_raw,
            soft_skills=_soft_raw,
            interests=_interests_raw,
            languages=_lang_raw,
        )

        # ── Contact: phone + linkedin, each worth 0.5 ─────────────────────────
        pi_phone2   = _wv(f"phone_input_{fk}", "phone")
        pi_linkedin = _wv(f"ln_input_{fk}",    "linkedin")
        _fill_contact = round(_filled(pi_phone2, pi_linkedin) / 2, 3)

        # ── Experience scoring — delegates to module-level score_experience_section ──
        def _score_experience():
            entries = ss.get("experience_entries", [])
            if not entries:
                return 0.0
            n = len(entries)
            # Merge live widget values into a temporary list for scoring
            merged = []
            for i, e in enumerate(entries):
                merged.append({
                    "title":       _get_val(ss, f"title_{i}_{n}_{fk}",       e, "title",       fk),
                    "company":     _get_val(ss, f"company_{i}_{n}_{fk}",     e, "company",     fk),
                    "duration":    _get_val(ss, f"duration_{i}_{n}_{fk}",    e, "duration",    fk),
                    "description": _get_val(ss, f"description_{i}_{n}_{fk}", e, "description", fk),
                })
            return score_experience_section(merged)  # module-level

        _fill_exp = _score_experience()

        # ── Education scoring — delegates to module-level score_education_section ──
        def _score_education():
            entries = ss.get("education_entries", [])
            if not entries:
                return 0.0
            n = len(entries)
            merged = []
            for i, e in enumerate(entries):
                merged.append({
                    "institution": _get_val(ss, f"institution_{i}_{n}_{fk}", e, "institution", fk),
                    "degree":      _get_val(ss, f"degree_{i}_{n}_{fk}",      e, "degree",      fk),
                    "year":        _get_val(ss, f"edu_year_{i}_{n}_{fk}",    e, "year",        fk),
                    "details":     _get_val(ss, f"edu_details_{i}_{n}_{fk}", e, "details",     fk),
                })
            return score_education_section(merged)  # module-level

        _fill_edu = _score_education()

        # ── Projects scoring — delegates to module-level score_project_section ──
        def _score_projects():
            entries = ss.get("project_entries", [])
            if not entries:
                return 0.0
            n = len(entries)
            merged = []
            for i, e in enumerate(entries):
                merged.append({
                    "title":       _get_val(ss, f"proj_title_{i}_{n}_{fk}",    e, "title",       fk),
                    "tech":        _get_val(ss, f"proj_tech_{i}_{n}_{fk}",     e, "tech",        fk),
                    "duration":    _get_val(ss, f"proj_duration_{i}_{n}_{fk}", e, "duration",    fk),
                    "description": _get_val(ss, f"proj_desc_{i}_{n}_{fk}",     e, "description", fk),
                })
            return score_project_section(merged)  # module-level

        _fill_proj = _score_projects()

        # ── Certificates scoring — delegates to module-level score_certificate_section ──
        def _score_certificates():
            entries = ss.get("certificate_links", [])
            if not entries:
                return 0.0
            n = len(entries)
            merged = []
            for i, e in enumerate(entries):
                merged.append({
                    "name":        _get_val(ss, f"cert_name_{i}_{n}_{fk}",        e, "name",        fk),
                    "link":        _get_val(ss, f"cert_link_{i}_{n}_{fk}",        e, "link",        fk),
                    "duration":    _get_val(ss, f"cert_duration_{i}_{n}_{fk}",    e, "duration",    fk),
                    "description": _get_val(ss, f"cert_description_{i}_{n}_{fk}", e, "description", fk),
                })
            return score_certificate_section(merged)  # module-level

        _fill_cert = _score_certificates()

        # ── Aggregate XP via module-level calculate_resume_xp() — idempotent ──
        # Order matches XP_WEIGHTS for consistent rendering.
        SECTIONS = {
            "Experience":     _fill_exp,
            "Projects":       _fill_proj,
            "Skills & More":  _fill_skills,
            "Education":      _fill_edu,
            "Summary":        _fill_summary,
            "Certificates":   _fill_cert,
            "Personal Info":  _fill_personal,
            "Contact":        _fill_contact,
        }
        ICON_KEYS = ["exp", "projects", "skills", "edu", "summary", "certs", "personal", "contact"]

        # ── Done thresholds — matched to new scoring formulas ─────────────────
        DONE_THRESHOLD = {
            "Personal Info":  1.0,   # all 5 fields filled
            "Summary":        0.85,  # rich summary with good word quality
            "Experience":     0.72,  # role + company + quality description
            "Education":      0.68,  # institution + degree + year + details
            "Projects":       0.65,  # title + techs + quality description
            "Skills & More":  0.57,  # 3+ tech skills + soft skills + interests
            "Certificates":   0.62,  # name + link + duration + description
            "Contact":        1.0,   # both phone AND linkedin
        }

        # ── Aggregate XP via calculate_resume_xp() — idempotent, no accumulation ──
        _raw_xp, xp, pct = calculate_resume_xp(SECTIONS, XP_WEIGHTS)
        max_xp    = XP_TOTAL_MAX   # always 100
        total     = len(SECTIONS)
        fully_done = sum(
            1 for (k, v) in SECTIONS.items()
            if v >= DONE_THRESHOLD.get(k, 1.0)
        )

        # Store computed XP in session state for external access
        ss["resume_xp"]  = xp
        ss["resume_pct"] = pct

        if   pct == 0:    rank, rank_color, rank_bg, rank_border = "Unranked",   "#6b7280", "#1e2535", "#374151"
        elif pct <= 20:   rank, rank_color, rank_bg, rank_border = "Beginner",   "#d97706", "#2a1f12", "#92400e"
        elif pct <= 40:   rank, rank_color, rank_bg, rank_border = "Builder",    "#94a3b8", "#1a2133", "#475569"
        elif pct <= 60:   rank, rank_color, rank_bg, rank_border = "Proficient", "#22d3ee", "#0c2233", "#0e4f60"
        elif pct <= 80:   rank, rank_color, rank_bg, rank_border = "Advanced",   "#f59e0b", "#2a2410", "#92700e"
        elif pct < 100:   rank, rank_color, rank_bg, rank_border = "Expert",     "#a78bfa", "#1a1a2e", "#6d28d9"
        else:             rank, rank_color, rank_bg, rank_border = "Pro Resume", "#34d399", "#0a2318", "#065f46"

        # ── helper: section row HTML — now takes fill float 0.0–1.0 ──────────
        def _section_row(label, icon_key, fill):
            done      = fill >= DONE_THRESHOLD.get(label, 1.0)
            partial   = 0.0 < fill < 1.0
            bar_pct   = f"{int(fill * 100)}%"
            # Per-section XP earned vs max for this section
            _sec_max  = XP_WEIGHTS.get(label, 0)
            _sec_earned = int(round(fill * _sec_max))
            # colour ramp: empty=dark, partial=amber, done=blue
            if done:
                icon_bg, icon_col, name_col = "#1d3a6e", "#93c5fd", "#93c5fd"
                bar_col = "#3b82f6"
                row_bg  = "#131c33"
                row_bdr = "#1d4ed8"
            elif partial:
                icon_bg, icon_col, name_col = "#2a1f12", "#f59e0b", "#d4a017"
                bar_col = "#f59e0b"
                row_bg  = "#1a1600"
                row_bdr = "#78450a"
            else:
                icon_bg, icon_col, name_col = "#1e2535", "#6b7280", "#6b7280"
                bar_col = "#374151"
                row_bg  = "#161b27"
                row_bdr = "#1e2535"
            chk_bg  = "#2563eb"    if done    else "transparent"
            chk_bdr = "#2563eb"    if done    else ("#78450a" if partial else "#374151")
            chk_op  = "1"          if done    else "0"
            # XP badge: shows "earned/max XP" — always visible, colour-coded
            if done:
                xp_badge_col = "#93c5fd"
            elif partial:
                xp_badge_col = "#f59e0b"
            else:
                xp_badge_col = "#4b5563"
            xp_badge = (
                "<div style='font-size:9px;color:" + xp_badge_col + ";font-weight:600;"
                "white-space:nowrap;'>"
                + str(_sec_earned) + "<span style='opacity:0.55;font-weight:400;'>/"
                + str(_sec_max) + "</span></div>"
            )
            return (
                "<div style='display:flex;align-items:center;gap:10px;padding:9px 10px;"
                "border-radius:9px;background:" + row_bg + ";border:0.5px solid " + row_bdr + ";"
                "margin-bottom:7px;'>"
                "<div style='width:28px;height:28px;border-radius:7px;background:" + icon_bg + ";"
                "display:flex;align-items:center;justify-content:center;flex-shrink:0;color:" + icon_col + ";'>"
                + SVG[icon_key] +
                "</div>"
                "<div style='flex:1;min-width:0;'>"
                "<div style='display:flex;justify-content:space-between;align-items:center;'>"
                "<div style='font-size:12px;font-weight:500;color:" + name_col + ";'>" + label + "</div>"
                + xp_badge +
                "</div>"
                "<div style='height:3px;background:#1e2535;border-radius:3px;margin-top:4px;overflow:hidden;'>"
                "<div style='height:100%;width:" + bar_pct + ";background:" + bar_col + ";border-radius:3px;'></div>"
                "</div>"
                "</div>"
                "<div style='width:18px;height:18px;border-radius:50%;background:" + chk_bg + ";"
                "border:1.5px solid " + chk_bdr + ";display:flex;align-items:center;"
                "justify-content:center;flex-shrink:0;'>"
                "<svg width='9' height='9' viewBox='0 0 10 10' fill='none' stroke='#fff' "
                "stroke-width='1.8' stroke-linecap='round' stroke-linejoin='round' "
                "style='opacity:" + chk_op + ";'>"
                "<polyline points='2,5 4.5,7.5 8.5,2.5'/>"
                "</svg>"
                "</div>"
                "</div>"
            )

        # ── streak dots — one dot per section, lit when section is 100% ─────────
        dots_html = "".join(
            f"<div style='flex:1;height:4px;border-radius:3px;"
            f"background:{'#3b82f6' if i < fully_done else '#1e2535'};'></div>"
            for i in range(total)
        )

        # ── render into sidebar ────────────────────────────────────────────────
        if st.session_state.get("username") != "admin":
            with st.sidebar:

                # ── XP header ─────────────────────────────────────────────────────
                st.markdown(f"""
<div style='margin-bottom:6px;display:flex;justify-content:space-between;align-items:center;'>
  <span style='font-size:10px;letter-spacing:1.2px;text-transform:uppercase;
               color:#6b7280;font-weight:500;'>Resume XP</span>
  <span style='font-size:11px;color:#9ca3af;font-weight:500;'>{xp} / {max_xp} XP</span>
</div>
<div style='margin-bottom:8px;'>
  <span style='font-size:11px;font-weight:500;padding:2px 10px;border-radius:20px;
               background:{rank_bg};border:0.5px solid {rank_border};color:{rank_color};'>{rank}</span>
</div>
<div style='width:100%;height:7px;background:#1e2535;border-radius:6px;overflow:hidden;margin-bottom:10px;'>
  <div style='height:100%;width:{pct}%;background:#3b82f6;border-radius:6px;'></div>
</div>
<div style='display:flex;gap:4px;margin-bottom:18px;'>{dots_html}</div>
<div style='font-size:10px;letter-spacing:1.4px;text-transform:uppercase;
            color:#4b5563;margin-bottom:10px;font-weight:500;'>Sections</div>
""", unsafe_allow_html=True)

                # ── section rows ──────────────────────────────────────────────────
                for (label, fill), icon_key in zip(SECTIONS.items(), ICON_KEYS):
                    st.markdown(_section_row(label, icon_key, fill), unsafe_allow_html=True)

                # ── UX Feedback Panel ─────────────────────────────────────────────
                # Generate actionable, ranked feedback tips based on current state.
                # Tips are prioritised by XP weight (highest-impact sections first).
                _feedback_tips = []

                # Experience feedback
                if _fill_exp == 0.0:
                    _feedback_tips.append("Add work experience to strengthen your resume (+30 XP potential)")
                elif _fill_exp < DONE_THRESHOLD["Experience"]:
                    _exp_entries = ss.get("experience_entries", [])
                    _has_desc = any(
                        len(str(e.get("description", "")).strip()) > 40
                        for e in _exp_entries
                    )
                    if not _has_desc:
                        _feedback_tips.append("Add detailed descriptions to your work experience to unlock more XP")
                    else:
                        _feedback_tips.append("Enrich your experience descriptions with specific achievements and metrics")

                # Projects feedback
                if _fill_proj == 0.0:
                    _feedback_tips.append("Add a project to gain up to 22 XP — projects are highly valued")
                elif _fill_proj < DONE_THRESHOLD["Projects"]:
                    _proj_entries = ss.get("project_entries", [])
                    _has_tech = any(
                        _count_valid_tokens(e.get("tech", "")) >= 1
                        for e in _proj_entries if e.get("title")
                    )
                    if not _has_tech:
                        _feedback_tips.append("List the technologies used in your projects to improve your score")
                    else:
                        _feedback_tips.append("Improve your project descriptions — explain the problem solved and your impact")
                elif _fill_proj < 0.90:
                    n_projs = len([e for e in ss.get("project_entries", []) if e.get("title")])
                    if n_projs < 2:
                        _feedback_tips.append("Add another project to increase your Projects score further")

                # Skills feedback
                if _fill_skills < DONE_THRESHOLD["Skills & More"]:
                    if _skill_count < 3:
                        _feedback_tips.append(f"Add more technical skills — you have {_skill_count}, aim for 5+ (comma-separated)")
                    if _soft_count < 2:
                        _feedback_tips.append("Add 2–3 soft skills (e.g. Leadership, Communication) to boost your score")

                # Education feedback
                if _fill_edu == 0.0:
                    _feedback_tips.append("Add your education details to build a complete resume")
                elif _fill_edu < DONE_THRESHOLD["Education"]:
                    _feedback_tips.append("Add graduation year and academic details/achievements to complete education")

                # Summary feedback
                if _fill_summary == 0.0:
                    _feedback_tips.append("Write a professional summary — it's your first impression on recruiters")
                elif _fill_summary < DONE_THRESHOLD["Summary"]:
                    _feedback_tips.append("Expand your summary with more specific skills, experience, and career goals")

                # Certificate feedback
                if _fill_cert == 0.0 and pct >= 40:
                    _feedback_tips.append("Add a certification to differentiate yourself from other candidates")
                elif 0.0 < _fill_cert < DONE_THRESHOLD["Certificates"]:
                    _feedback_tips.append("Add a verification link and description to your certificates for full credit")

                # Contact feedback
                if _fill_contact < 1.0:
                    if not pi_phone2:
                        _feedback_tips.append("Add your phone number to make your resume complete")
                    if not pi_linkedin:
                        _feedback_tips.append("Add your LinkedIn profile URL — recruiters always check it")

                # Show up to 3 tips (highest priority = highest XP weight = listed first)
                if _feedback_tips:
                    _tips_to_show = _feedback_tips[:3]
                    _tip_html_items = "".join(
                        f"<div style='display:flex;align-items:flex-start;gap:8px;margin-bottom:8px;'>"
                        f"<span style='color:#f59e0b;font-size:12px;flex-shrink:0;margin-top:1px;'>&#9654;</span>"
                        f"<span style='font-size:11px;color:#cbd5e1;line-height:1.5;'>{tip}</span>"
                        f"</div>"
                        for tip in _tips_to_show
                    )
                    st.markdown(
                        f"<div style='margin:12px 0 6px;padding:10px 12px;background:#111827;"
                        f"border-radius:8px;border:0.5px solid #374151;'>"
                        f"<div style='font-size:9px;letter-spacing:1.2px;text-transform:uppercase;"
                        f"color:#6b7280;font-weight:600;margin-bottom:8px;'>Tips to Boost XP</div>"
                        f"{_tip_html_items}"
                        f"</div>",
                        unsafe_allow_html=True
                    )

                # ── divider + stats footer ─────────────────────────────────────────
                st.markdown(f"""
<hr style='border:none;border-top:0.5px solid #1e2535;margin:14px 0;'>
<div style='display:flex;justify-content:space-between;text-align:center;margin-bottom:18px;'>
  <div>
    <span style='font-size:16px;font-weight:500;color:#e2e8f0;display:block;'>{fully_done}</span>
    <span style='font-size:10px;color:#4b5563;letter-spacing:0.8px;text-transform:uppercase;'>Done</span>
  </div>
  <div>
    <span style='font-size:16px;font-weight:500;color:#e2e8f0;display:block;'>{xp}<span style='font-size:10px;color:#4b5563;'>/{max_xp}</span></span>
    <span style='font-size:10px;color:#4b5563;letter-spacing:0.8px;text-transform:uppercase;'>XP</span>
  </div>
  <div>
    <span style='font-size:16px;font-weight:500;color:#e2e8f0;display:block;'>{pct}%</span>
    <span style='font-size:10px;color:#4b5563;letter-spacing:0.8px;text-transform:uppercase;'>Complete</span>
  </div>
</div>
<hr style='border:none;border-top:0.5px solid #1e2535;margin:0 0 14px;'>
<div style='font-size:10px;letter-spacing:1.4px;text-transform:uppercase;
                color:#4b5563;margin-bottom:10px;font-weight:500;'>Manage Sections</div>
""", unsafe_allow_html=True)

                # ── section add/delete controls ────────────────────────────────────
                if "edit_mode" not in ss:
                    ss.edit_mode = "Add"

                mode = st.selectbox(
                    "Mode",
                    ["Add", "Delete"],
                    index=0,
                    key="mode_dropdown",
                    label_visibility="collapsed",
                )
                ss.edit_mode = mode

                st.markdown("<div style='height:8px;'></div>", unsafe_allow_html=True)

                # Experience
                with st.expander("Experience", expanded=False):
                    if st.button(("Add" if mode == "Add" else "Delete") + " Experience", key="exp_btn", use_container_width=True):
                        if mode == "Add":
                            ss.experience_entries.append({"title": "", "company": "", "duration": "", "description": ""})
                        elif mode == "Delete" and len(ss.experience_entries) > 1:
                            ss.experience_entries.pop()

                # Education
                with st.expander("Education", expanded=False):
                    if st.button(("Add" if mode == "Add" else "Delete") + " Education", key="edu_btn", use_container_width=True):
                        if mode == "Add":
                            ss.education_entries.append({"degree": "", "institution": "", "year": "", "details": ""})
                        elif mode == "Delete" and len(ss.education_entries) > 1:
                            ss.education_entries.pop()

                # Projects
                with st.expander("Projects", expanded=False):
                    if st.button(("Add" if mode == "Add" else "Delete") + " Project", key="proj_btn", use_container_width=True):
                        if mode == "Add":
                            ss.project_entries.append({"title": "", "tech": "", "duration": "", "description": ""})
                        elif mode == "Delete" and len(ss.project_entries) > 1:
                            ss.project_entries.pop()

                # Certificates
                with st.expander("Certificates", expanded=False):
                    if st.button(("Add" if mode == "Add" else "Delete") + " Certificate", key="cert_btn", use_container_width=True):
                        if mode == "Add":
                            ss.certificate_links.append({"name": "", "link": "", "duration": "", "description": ""})
                        elif mode == "Delete" and len(ss.certificate_links) > 1:
                            ss.certificate_links.pop()

    # ── call gamified sidebar AFTER fk is known so widget keys resolve correctly ──
    fk = st.session_state["form_key_counter"]
    render_gamified_sidebar(st.session_state, fk)
    mode = st.session_state.get("edit_mode", "Add")

    # ── Shared section-header style injected once ────────────────────────────
    st.markdown("""
    <style>
    .sec-header {
        display: flex; align-items: center; gap: 10px;
        background: linear-gradient(90deg, rgba(0,180,255,0.10) 0%, rgba(0,180,255,0.03) 100%);
        border-left: 3px solid #4da6ff;
        border-radius: 0 10px 10px 0;
        padding: 9px 14px; margin: 18px 0 10px;
    }
    .sec-header .sec-icon { font-size: 18px; line-height: 1; }
    .sec-header .sec-title {
        font-size: 15px; font-weight: 700;
        color: #93c5fd; letter-spacing: 0.4px; margin: 0;
    }
    .sec-header .sec-badge {
        margin-left: auto; font-size: 10px; font-weight: 600;
        color: #4b5563; background: #1e2535;
        padding: 2px 8px; border-radius: 20px;
        border: 0.5px solid #374151;
    }
    .entry-card {
        background: rgba(13,20,40,0.55);
        border: 0.5px solid rgba(0,180,255,0.18);
        border-radius: 12px; padding: 14px 16px; margin-bottom: 12px;
    }
    .entry-card-label {
        font-size: 12px; font-weight: 600; color: #4da6ff;
        margin-bottom: 10px; letter-spacing: 0.3px;
    }
    .field-hint { font-size: 11px; color: #4b5563; margin-top: -8px; margin-bottom: 6px; }
    .tag-chip {
        display: inline-block; background: rgba(0,180,255,0.12);
        border: 0.5px solid rgba(0,180,255,0.35); color: #93c5fd;
        font-size: 12px; font-weight: 500;
        padding: 3px 10px; border-radius: 20px; margin: 3px 3px 3px 0;
    }
    .confirm-warn {
        background: rgba(239,68,68,0.10); border: 0.5px solid rgba(239,68,68,0.35);
        border-radius: 10px; padding: 10px 14px; margin-top: 6px;
        font-size: 13px; color: #fca5a5;
    }
    .photo-preview {
        width: 90px; height: 90px; border-radius: 50%; object-fit: cover;
        border: 3px solid #4da6ff;
        box-shadow: 0 0 14px rgba(77,166,255,0.4);
        display: block; margin: 8px auto 0;
    }
    @media (max-width: 768px) {
        [data-testid="column"] { min-width: 100% !important; }
    }
    </style>
    """, unsafe_allow_html=True)

    def _sec_hdr(icon, title, badge=None):
        badge_html = f"<span class='sec-badge'>{badge}</span>" if badge else ""
        st.markdown(
            f"<div class='sec-header'>"
            f"<span class='sec-icon'>{icon}</span>"
            f"<span class='sec-title'>{title}</span>"
            f"{badge_html}</div>",
            unsafe_allow_html=True,
        )

    def _hint(text):
        st.markdown(f"<div class='field-hint'>💡 {text}</div>", unsafe_allow_html=True)

    def _tag_chips(raw, label):
        items = [s.strip() for s in raw.split(",") if s.strip()]
        if not items:
            return
        chips = "".join(f"<span class='tag-chip'>{t}</span>" for t in items)
        st.markdown(
            f"<div style='margin-bottom:4px;font-size:11px;color:#6b7280;'>{label}</div>"
            f"<div style='margin-bottom:10px;'>{chips}</div>",
            unsafe_allow_html=True,
        )

    # ---------------- Resume Form ----------------
    with st.form(f"resume_form_{fk}", clear_on_submit=False):
        _sec_hdr("👤", "Personal Information")
        col1, col2 = st.columns(2)
        with col1:
            # FIX: Do NOT assign back to session_state inside the form.
            # Streamlit batches form widgets — writing to session_state here
            # triggers an immediate rerun on every keystroke, defeating the form.
            # Instead, just render the widget with `value=` for pre-fill.
            # Values are committed to session_state only when submitted=True below.
            st.text_input("👤 Full Name", value=st.session_state.name, placeholder="e.g., Arjun Sharma", key=f"name_input_{fk}")
            st.text_input("📞 Phone Number", value=st.session_state.phone, placeholder="e.g., +91 98765 43210", key=f"phone_input_{fk}")
            st.text_input("📍 Location", value=st.session_state.location, placeholder="e.g., Kolkata, West Bengal", key=f"loc_input_{fk}")
        with col2:
            st.text_input("📧 Email", value=st.session_state.email, placeholder="e.g., arjun@gmail.com", key=f"email_input_{fk}")
            st.text_input("🔗 LinkedIn", value=st.session_state.linkedin, placeholder="e.g., linkedin.com/in/arjun", key=f"ln_input_{fk}")
            st.text_input("🌐 Portfolio", value=st.session_state.portfolio, placeholder="e.g., arjun.dev or github.com/arjun", key=f"port_input_{fk}")
            st.text_input("💼 Job Title / Target Role", value=st.session_state.job_title, placeholder="e.g., Full Stack Developer", key=f"job_input_{fk}")

        _sec_hdr("📝", "Professional Summary")
        st.text_area(
            "Summary",
            value=st.session_state.summary,
            placeholder="Write 3–5 sentences about your career goals, key strengths, and what makes you stand out. E.g., 'Results-driven software engineer with 3+ years building scalable web apps...'",
            height=120,
            key=f"summary_input_{fk}",
        )
        _hint("Aim for 80–200 characters. Recruiters read this first — make it count.")

        _sec_hdr("🛠️", "Skills, Languages, Interests & Soft Skills")
        st.text_area(
            "Technical Skills (comma-separated)",
            value=st.session_state.skills,
            placeholder="e.g., Python, React, Node.js, PostgreSQL, Docker, AWS",
            height=70,
            key=f"skills_input_{fk}",
        )
        _hint("List 5+ skills for best score. Separate each with a comma.")
        # FIX: _tag_chips now reads from the widget key directly (live value),
        # not from session_state.skills which lags by one submit cycle.
        _tag_chips(st.session_state.get(f"skills_input_{fk}", st.session_state.skills), "Preview:")

        st.text_area(
            "Languages (comma-separated)",
            value=st.session_state.languages,
            placeholder="e.g., English, Bengali, Hindi",
            height=60,
            key=f"lang_input_{fk}",
        )
        _tag_chips(st.session_state.get(f"lang_input_{fk}", st.session_state.languages), "Preview:")
        st.text_area(
            "Interests / Hobbies (comma-separated)",
            value=st.session_state.interests,
            placeholder="e.g., Open Source, Machine Learning, Chess, Blogging",
            height=60,
            key=f"int_input_{fk}",
        )
        _tag_chips(st.session_state.get(f"int_input_{fk}", st.session_state.interests), "Preview:")
        st.text_area(
            "Soft Skills (comma-separated)",
            value=st.session_state.Softskills,
            placeholder="e.g., Leadership, Communication, Problem Solving, Teamwork",
            height=60,
            key=f"soft_input_{fk}",
        )
        _tag_chips(st.session_state.get(f"soft_input_{fk}", st.session_state.Softskills), "Preview:")

        _sec_hdr("🧱", "Work Experience", badge=f"{len(st.session_state.experience_entries)} entr{'y' if len(st.session_state.experience_entries)==1 else 'ies'}")
        for idx, exp in enumerate(st.session_state.experience_entries):
            _entry_label = exp.get("title", "") or f"Experience #{idx+1}"
            _entry_company = exp.get("company", "")
            _display = f"{_entry_label} @ {_entry_company}" if _entry_company else _entry_label
            with st.expander(f"🏢 {_display}", expanded=True):
                st.markdown(f"<div class='entry-card-label'>Entry #{idx+1}</div>", unsafe_allow_html=True)
                # FIX: Do not assign back to exp dict here — that mutates session_state
                # inside the form, causing a rerun on every keystroke.
                # Widget keys are unique and Streamlit persists their values automatically.
                # _sync_entries() in the sidebar and the submit handler below read them.
                st.text_input("Job Title", value=exp.get("title", ""), placeholder="e.g., Software Engineer", key=f"title_{idx}_{len(st.session_state.experience_entries)}_{fk}")
                st.text_input("Company", value=exp.get("company", ""), placeholder="e.g., Infosys, TCS, Google", key=f"company_{idx}_{len(st.session_state.experience_entries)}_{fk}")
                st.text_input("Duration", value=exp.get("duration", ""), placeholder="e.g., Jun 2022 – Present", key=f"duration_{idx}_{len(st.session_state.experience_entries)}_{fk}")
                st.text_area("Description", value=exp.get("description", ""), placeholder="• Developed REST APIs using Node.js that reduced response time by 35%\n• Led a team of 4 engineers to deliver the project 2 weeks ahead of schedule", height=100, key=f"description_{idx}_{len(st.session_state.experience_entries)}_{fk}")
                _hint("Use bullet points starting with action verbs. Include metrics where possible.")

        _sec_hdr("🎓", "Education", badge=f"{len(st.session_state.education_entries)} entr{'y' if len(st.session_state.education_entries)==1 else 'ies'}")
        for idx, edu in enumerate(st.session_state.education_entries):
            _edu_label = edu.get("degree", "") or f"Education #{idx+1}"
            _edu_inst = edu.get("institution", "")
            _edu_display = f"{_edu_label} — {_edu_inst}" if _edu_inst else _edu_label
            with st.expander(f"🏫 {_edu_display}", expanded=True):
                st.markdown(f"<div class='entry-card-label'>Entry #{idx+1}</div>", unsafe_allow_html=True)
                st.text_input("Degree / Qualification", value=edu.get("degree", ""), placeholder="e.g., B.Tech in Computer Science", key=f"degree_{idx}_{len(st.session_state.education_entries)}_{fk}")
                st.text_input("Institution", value=edu.get("institution", ""), placeholder="e.g., Jadavpur University", key=f"institution_{idx}_{len(st.session_state.education_entries)}_{fk}")
                st.text_input("Year / Duration", value=edu.get("year", ""), placeholder="e.g., 2019 – 2023", key=f"edu_year_{idx}_{len(st.session_state.education_entries)}_{fk}")
                st.text_area("Academic Details", value=edu.get("details", ""), placeholder="e.g., CGPA: 8.7/10 | Relevant: Data Structures, OS, DBMS | Dean's List 2022", height=80, key=f"edu_details_{idx}_{len(st.session_state.education_entries)}_{fk}")

        _sec_hdr("🚀", "Projects", badge=f"{len(st.session_state.project_entries)} entr{'y' if len(st.session_state.project_entries)==1 else 'ies'}")
        for idx, proj in enumerate(st.session_state.project_entries):
            _proj_label = proj.get("title", "") or f"Project #{idx+1}"
            with st.expander(f"📌 {_proj_label}", expanded=True):
                st.markdown(f"<div class='entry-card-label'>Project #{idx+1}</div>", unsafe_allow_html=True)
                st.text_input("Project Title", value=proj.get("title", ""), placeholder="e.g., AI Resume Builder", key=f"proj_title_{idx}_{len(st.session_state.project_entries)}_{fk}")
                st.text_input("Tech Stack", value=proj.get("tech", ""), placeholder="e.g., Python, Streamlit, OpenAI API, PostgreSQL", key=f"proj_tech_{idx}_{len(st.session_state.project_entries)}_{fk}")
                st.text_input("Duration", value=proj.get("duration", ""), placeholder="e.g., Jan 2024 – Mar 2024  (or  2 months)", key=f"proj_duration_{idx}_{len(st.session_state.project_entries)}_{fk}")
                st.text_area("Description", value=proj.get("description", ""), placeholder="• Built a full-stack resume builder with AI-powered cover letter generation\n• Reduced resume creation time by 70% compared to manual methods", height=100, key=f"proj_desc_{idx}_{len(st.session_state.project_entries)}_{fk}")
                _hint("Describe the problem solved, your role, and the impact or outcome.")

        _sec_hdr("🔗", "Project Links")
        st.text_area(
            "Enter one project link per line:",
            value="\n".join(st.session_state.project_links),
            placeholder="https://github.com/yourname/project1\nhttps://yourproject.netlify.app",
            height=80,
            key=f"proj_links_input_{fk}",
        )
        # FIX: Do NOT write session_state.project_links here — that's a mutation
        # inside the form. It will be committed on submit below.

        _sec_hdr("🏅", "Certificates", badge=f"{len(st.session_state.certificate_links)} entr{'y' if len(st.session_state.certificate_links)==1 else 'ies'}")
        for idx, cert in enumerate(st.session_state.certificate_links):
            _cert_label = cert.get("name", "") or f"Certificate #{idx+1}"
            with st.expander(f"🎖️ {_cert_label}", expanded=True):
                st.markdown(f"<div class='entry-card-label'>Certificate #{idx+1}</div>", unsafe_allow_html=True)
                st.text_input("Certificate Name", value=cert.get("name", ""), placeholder="e.g., AWS Certified Solutions Architect", key=f"cert_name_{idx}_{len(st.session_state.certificate_links)}_{fk}")
                st.text_input("Verification Link", value=cert.get("link", ""), placeholder="e.g., https://credly.com/badges/...", key=f"cert_link_{idx}_{len(st.session_state.certificate_links)}_{fk}")
                st.text_input("Issued Date", value=cert.get("duration", ""), placeholder="e.g., March 2024", key=f"cert_duration_{idx}_{len(st.session_state.certificate_links)}_{fk}")
                st.text_area("Description", value=cert.get("description", ""), placeholder="e.g., Demonstrates expertise in designing distributed systems on AWS. Covers EC2, S3, RDS, and networking.", height=80, key=f"cert_description_{idx}_{len(st.session_state.certificate_links)}_{fk}")

        st.markdown("<div style='height:16px;'></div>", unsafe_allow_html=True)
        btn_col1, btn_col2 = st.columns([1, 1])
        with btn_col1:
            submitted = st.form_submit_button(
                "📑 Generate Resume",
                use_container_width=True,
                type="primary",
            )
        with btn_col2:
            clear_clicked = st.form_submit_button(
                "🗑️ Clear All",
                use_container_width=True,
            )

        if submitted:
            st.session_state["_resume_generated_msg"] = True
            st.session_state["_resume_generating"] = True
            # ── Commit all form widget values to session_state on submit ──────
            # This is the ONLY place we write widget values back — not during typing.
            ss = st.session_state
            ss.name      = ss.get(f"name_input_{fk}",    ss.name)
            ss.email     = ss.get(f"email_input_{fk}",   ss.email)
            ss.phone     = ss.get(f"phone_input_{fk}",   ss.phone)
            ss.location  = ss.get(f"loc_input_{fk}",     ss.location)
            ss.linkedin  = ss.get(f"ln_input_{fk}",      ss.linkedin)
            ss.portfolio = ss.get(f"port_input_{fk}",    ss.portfolio)
            ss.job_title = ss.get(f"job_input_{fk}",     ss.job_title)
            ss.summary   = ss.get(f"summary_input_{fk}", ss.summary)
            ss.skills    = ss.get(f"skills_input_{fk}",  ss.skills)
            ss.languages = ss.get(f"lang_input_{fk}",    ss.languages)
            ss.interests = ss.get(f"int_input_{fk}",     ss.interests)
            ss.Softskills = ss.get(f"soft_input_{fk}",   ss.Softskills)
            # Project links
            _pl_raw = ss.get(f"proj_links_input_{fk}", "")
            ss.project_links = [lnk.strip() for lnk in _pl_raw.splitlines() if lnk.strip()]
            # Sync experience entry dicts
            _n_exp = len(ss.experience_entries)
            for _i, _e in enumerate(ss.experience_entries):
                _e["title"]       = ss.get(f"title_{_i}_{_n_exp}_{fk}",       _e.get("title", ""))
                _e["company"]     = ss.get(f"company_{_i}_{_n_exp}_{fk}",     _e.get("company", ""))
                _e["duration"]    = ss.get(f"duration_{_i}_{_n_exp}_{fk}",    _e.get("duration", ""))
                _e["description"] = ss.get(f"description_{_i}_{_n_exp}_{fk}", _e.get("description", ""))
            # Sync education entry dicts
            _n_edu = len(ss.education_entries)
            for _i, _e in enumerate(ss.education_entries):
                _e["degree"]      = ss.get(f"degree_{_i}_{_n_edu}_{fk}",      _e.get("degree", ""))
                _e["institution"] = ss.get(f"institution_{_i}_{_n_edu}_{fk}", _e.get("institution", ""))
                _e["year"]        = ss.get(f"edu_year_{_i}_{_n_edu}_{fk}",    _e.get("year", ""))
                _e["details"]     = ss.get(f"edu_details_{_i}_{_n_edu}_{fk}", _e.get("details", ""))
            # Sync project entry dicts
            _n_proj = len(ss.project_entries)
            for _i, _e in enumerate(ss.project_entries):
                _e["title"]       = ss.get(f"proj_title_{_i}_{_n_proj}_{fk}",    _e.get("title", ""))
                _e["tech"]        = ss.get(f"proj_tech_{_i}_{_n_proj}_{fk}",     _e.get("tech", ""))
                _e["duration"]    = ss.get(f"proj_duration_{_i}_{_n_proj}_{fk}", _e.get("duration", ""))
                _e["description"] = ss.get(f"proj_desc_{_i}_{_n_proj}_{fk}",     _e.get("description", ""))
            # Sync certificate entry dicts
            _n_cert = len(ss.certificate_links)
            for _i, _e in enumerate(ss.certificate_links):
                _e["name"]        = ss.get(f"cert_name_{_i}_{_n_cert}_{fk}",        _e.get("name", ""))
                _e["link"]        = ss.get(f"cert_link_{_i}_{_n_cert}_{fk}",        _e.get("link", ""))
                _e["duration"]    = ss.get(f"cert_duration_{_i}_{_n_cert}_{fk}",    _e.get("duration", ""))
                _e["description"] = ss.get(f"cert_description_{_i}_{_n_cert}_{fk}", _e.get("description", ""))

        if clear_clicked:
            st.session_state["_confirm_clear"] = True

    # ── Clear confirmation (outside form so it can render fresh buttons) ──────
    if st.session_state.get("_confirm_clear"):
        st.markdown(
            "<div class='confirm-warn'>⚠️ <strong>This will erase all entered data.</strong> "
            "This cannot be undone.</div>",
            unsafe_allow_html=True,
        )
        cc1, cc2 = st.columns([1, 1])
        with cc1:
            if st.button("✅ Yes, Clear", key="confirm_clear_yes", use_container_width=True):
                _new_counter = st.session_state.get("form_key_counter", 0) + 1
                resume_fields = ["name", "email", "phone", "linkedin", "location",
                                 "portfolio", "summary", "skills", "languages",
                                 "interests", "Softskills", "job_title"]
                for _f in resume_fields:
                    st.session_state[_f] = ""
                st.session_state["experience_entries"] = [{"title": "", "company": "", "duration": "", "description": ""}]
                st.session_state["education_entries"] = [{"degree": "", "institution": "", "year": "", "details": ""}]
                st.session_state["project_entries"] = [{"title": "", "tech": "", "duration": "", "description": ""}]
                st.session_state["project_links"] = []
                st.session_state["certificate_links"] = [{"name": "", "link": "", "duration": "", "description": ""}]
                for _key in ["generated_html", "ai_output", "cover_letter",
                             "cover_letter_html", "encoded_profile_image"]:
                    st.session_state.pop(_key, None)
                st.session_state["form_key_counter"] = _new_counter
                st.session_state.pop("_confirm_clear", None)
                st.rerun()
        with cc2:
            if st.button("❌ Cancel", key="confirm_clear_no", use_container_width=True):
                st.session_state.pop("_confirm_clear", None)
                st.rerun()

    st.markdown("""
    <style>
        .heading-large {
            font-size: 36px;
            font-weight: bold;
            color: #336699;
        }
        .subheading-large {
            font-size: 30px;
            font-weight: bold;
            color: #336699;
        }
        .tab-section {
            margin-top: 20px;
        }
    </style>
    """, unsafe_allow_html=True)

    # ── Loading button pulse animation CSS ────────────────────────────────────
    st.markdown("""
    <style>
    /* Pulse animation for buttons during processing */
    @keyframes btn-pulse {
        0%   { box-shadow: 0 0 0 0 rgba(0,200,255,0.55); }
        70%  { box-shadow: 0 0 0 10px rgba(0,200,255,0); }
        100% { box-shadow: 0 0 0 0 rgba(0,200,255,0); }
    }
    /* Spinner overlay for the stSpinner */
    [data-testid="stSpinner"] > div {
        background: rgba(10, 20, 40, 0.75) !important;
        border: 1px solid rgba(0,200,255,0.4) !important;
        border-radius: 12px !important;
        padding: 14px 20px !important;
        backdrop-filter: blur(10px) !important;
        color: #93c5fd !important;
        font-size: 14px !important;
        font-weight: 500 !important;
    }
    [data-testid="stSpinner"] svg {
        color: #4da6ff !important;
        stroke: #4da6ff !important;
    }
    /* Download button loading feel */
    [data-testid="stDownloadButton"] > button {
        position: relative;
        overflow: hidden;
    }
    [data-testid="stDownloadButton"] > button:active::after {
        content: "";
        position: absolute;
        inset: 0;
        background: rgba(0,200,255,0.18);
        animation: btn-pulse 0.6s ease-out;
    }
    /* Form submit button active state — pulse */
    button[kind="formSubmit"]:active,
    button[data-testid="baseButton-primary"]:active {
        animation: btn-pulse 0.5s ease-out !important;
    }
    /* Disabled state for buttons during loading */
    .stButton > button:disabled,
    button[disabled] {
        opacity: 0.55 !important;
        cursor: not-allowed !important;
        animation: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # --- Visual Resume Preview Section (only shown after form is submitted) ---
    if st.session_state.get("_resume_generated_msg"):
        st.success("✅ Resume Generated Successfully! Scroll down to preview or download.")
        st.session_state["_resume_generated_msg"] = False  # show only once per submit

    if "generated_html" in st.session_state:
        st.markdown("## 🧾 <span style='color:#336699;'>Resume Preview</span>", unsafe_allow_html=True)
        st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

        left, right = st.columns([1, 2])

        with left:
            st.markdown(f"""
                <h2 style='color:#2f2f2f;margin-bottom:0;'>{st.session_state['name']}</h2>
                <h4 style='margin-top:5px;color:#444;'>{st.session_state['job_title']}</h4>
                <p style='font-size:14px;'>
                📍 {st.session_state['location']}<br>
                📞 {st.session_state['phone']}<br>
                📧 <a href="mailto:{st.session_state['email']}">{st.session_state['email']}</a><br>
                🔗 <a href="{st.session_state['linkedin']}" target="_blank">LinkedIn</a><br>
                🌐 <a href="{st.session_state['portfolio']}" target="_blank">Portfolio</a>
                </p>
            """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Skills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            _skill_items = [s.strip() for s in st.session_state["skills"].split(",") if s.strip()]
            if _skill_items:
                _chips = "".join(f"<span class='tag-chip'>{s}</span>" for s in _skill_items)
                st.markdown(f"<div style='margin-bottom:8px;'>{_chips}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Languages</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            _lang_items = [l.strip() for l in st.session_state["languages"].split(",") if l.strip()]
            if _lang_items:
                _lang_chips = "".join(f"<span class='tag-chip'>{l}</span>" for l in _lang_items)
                st.markdown(f"<div style='margin-bottom:8px;'>{_lang_chips}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Interests</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            _int_items = [i.strip() for i in st.session_state["interests"].split(",") if i.strip()]
            if _int_items:
                _int_chips = "".join(f"<span class='tag-chip'>{t}</span>" for t in _int_items)
                st.markdown(f"<div style='margin-bottom:8px;'>{_int_chips}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Soft Skills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            _soft_items = [i.strip() for i in st.session_state["Softskills"].split(",") if i.strip()]
            if _soft_items:
                _soft_chips = "".join(f"<span class='tag-chip'>{t}</span>" for t in _soft_items)
                st.markdown(f"<div style='margin-bottom:8px;'>{_soft_chips}</div>", unsafe_allow_html=True)

        with right:
            st.markdown("<h4 style='color:#336699;'>Summary</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            summary_text = st.session_state["summary"].replace("\n", "<br>")
            st.markdown(f"<p style='font-size:17px;'>{summary_text}</p>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Experience</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for exp in st.session_state.experience_entries:
                if exp["company"] or exp["title"]:
                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px; border-radius:8px;'>
                        <div style='display:flex; justify-content:space-between;'>
                            <b>🏢 {exp['company']}</b><span style='color:gray;'>📆 {exp['duration']}</span>
                        </div>
                        <div style='font-size:14px;'>💼 <i>{exp['title']}</i></div>
                        <div style='font-size:17px;'>📝 {exp['description']}</div>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>🎓 Education</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for edu in st.session_state.education_entries:
                if edu["institution"] or edu["degree"]:
                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px 15px; border-radius:8px;'>
                        <div style='display:flex; justify-content:space-between; font-size:16px; font-weight:bold;'>
                            <span>🏫 {edu['institution']}</span>
                            <span style='color:gray;'>📅 {edu['year']}</span>
                        </div>
                        <div style='font-size:14px;'>🎓 <i>{edu['degree']}</i></div>
                        <div style='font-size:14px;'>📄 {edu['details']}</div>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Projects</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for proj in st.session_state.project_entries:
                if proj.get("title"):
                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px;'>
                        <strong style='font-size:16px;'>{proj['title']}</strong><br>
                        <span style='font-size:14px;'>🛠️ <strong>Tech Stack:</strong> {proj['tech']}</span><br>
                        <span style='font-size:14px;'>⏳ <strong>Duration:</strong> {proj['duration']}</span><br>
                        <span style='font-size:17px;'>📝 <strong>Description:</strong> {proj['description']}</span>
                    </div>
                    """, unsafe_allow_html=True)

            if st.session_state.project_links:
                st.markdown("<h4 style='color:#336699;'>Project Links</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for i, link in enumerate(st.session_state.project_links):
                    st.markdown(f"[🔗 Project {i+1}]({link})", unsafe_allow_html=True)

            if st.session_state.certificate_links:
                st.markdown("<h4 style='color:#336699;'>Certificates</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for cert in st.session_state.certificate_links:
                    if cert["name"]:
                        name_html = (f"<a href=\"{cert['link']}\" target=\"_blank\"><b>\U0001f4c4 {cert['name']}</b></a>"
                                     if cert.get("link") else f"<b>\U0001f4c4 {cert['name']}</b>")
                        st.markdown(f"""
                        <div style='display:flex; justify-content:space-between;'>
                            {name_html}
                            <span style='color:gray;'>{cert['duration']}</span>
                        </div>
                        <div style='margin-bottom:10px; font-size:14px;'>{cert['description']}</div>
                        """, unsafe_allow_html=True)

import re

with tab2:
    st.markdown("## ✨ <span style='color:#336699;'>Enhanced AI Resume Preview</span>", unsafe_allow_html=True)
    st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

    col1, spacer, col2 = st.columns([1, 0.2, 1])

    with col1:
        if st.button("🔁 Clear Preview"):
            st.session_state.pop("ai_output", None)
            st.toast("🗑️ Preview cleared!")

    with col2:
        if st.button("🚀 Generate AI Resume Preview"):

            # ── Job Title Spell-Correction ─────────────────────────────────────────
            # Fuzzy-matches the user's input against the canonical role list.
            # If a close match is found (≥ 70% similarity) it silently corrects the
            # title in session_state before anything else runs.
            import difflib as _difflib

            _CANONICAL_JOB_TITLES = [
                # ── Software Engineering ──────────────────────────────────────
                "Software Engineer", "Senior Software Engineer", "Staff Software Engineer",
                "Frontend Developer", "Backend Developer", "Full Stack Developer",
                "React Developer", "Angular Developer", "Vue.js Developer",
                "Node.js Developer", "Python Developer", "Java Developer",
                "Go Developer", "Rust Developer", "C++ Developer",
                "PHP Developer", "Ruby on Rails Developer", ".NET Developer",
                "iOS Developer", "Android Developer", "Mobile App Developer",
                "Flutter Developer", "React Native Developer",
                "Embedded Systems Engineer", "Firmware Engineer",
                # ── Data & AI ─────────────────────────────────────────────────
                "Data Scientist", "Senior Data Scientist",
                "Data Analyst", "Senior Data Analyst",
                "Data Engineer", "Senior Data Engineer",
                "Machine Learning Engineer", "ML Engineer",
                "AI Engineer", "Generative AI Engineer",
                "LLM Engineer", "Prompt Engineer",
                "Computer Vision Engineer", "NLP Engineer",
                "Business Intelligence Analyst", "BI Developer",
                "Analytics Engineer", "Quantitative Analyst",
                # ── Infrastructure & Cloud ────────────────────────────────────
                "DevOps Engineer", "Senior DevOps Engineer",
                "Cloud Engineer", "AWS Engineer", "Azure Engineer", "GCP Engineer",
                "Site Reliability Engineer", "Platform Engineer",
                "Infrastructure Engineer", "Systems Administrator",
                "Network Engineer", "Network Administrator",
                "Database Administrator", "Database Engineer",
                # ── Security ──────────────────────────────────────────────────
                "Cybersecurity Analyst", "Information Security Analyst",
                "Security Engineer", "Penetration Tester",
                "SOC Analyst", "Cloud Security Engineer",
                # ── QA & Testing ──────────────────────────────────────────────
                "QA Engineer", "QA Analyst", "SDET",
                "Automation Test Engineer", "Performance Test Engineer",
                # ── Architecture & Leadership ─────────────────────────────────
                "Solutions Architect", "Cloud Architect", "Enterprise Architect",
                "Technical Lead", "Engineering Manager",
                "Chief Technology Officer", "VP of Engineering",
                # ── Product & Design ──────────────────────────────────────────
                "Product Manager", "Senior Product Manager",
                "Product Owner", "Technical Product Manager",
                "UI/UX Designer", "UX Designer", "UI Designer",
                "Product Designer", "Graphic Designer",
                "Visual Designer", "Motion Designer",
                "Interaction Designer", "Design Lead",
                # ── Project & Delivery ────────────────────────────────────────
                "Project Manager", "Senior Project Manager",
                "Scrum Master", "Agile Coach",
                "Program Manager", "Delivery Manager",
                "IT Project Manager",
                # ── Business & Analysis ───────────────────────────────────────
                "Business Analyst", "Senior Business Analyst",
                "Systems Analyst", "Functional Consultant",
                "ERP Consultant", "Salesforce Developer",
                "Salesforce Administrator",
                # ── Marketing & Growth ────────────────────────────────────────
                "Digital Marketing Specialist", "Digital Marketing Manager",
                "SEO Specialist", "SEM Specialist",
                "Content Strategist", "Content Writer",
                "Social Media Manager", "Growth Hacker",
                "Performance Marketing Manager", "Email Marketing Specialist",
                "Brand Manager", "Marketing Analyst",
                "E-commerce Specialist", "E-commerce Manager",
                # ── Finance & Accounting ──────────────────────────────────────
                "Financial Analyst", "Senior Financial Analyst",
                "Chartered Accountant", "Cost Accountant",
                "Investment Analyst", "Equity Research Analyst",
                "Risk Analyst", "Credit Analyst",
                "Fintech Developer", "Quantitative Developer",
                "Accounts Manager", "Tax Consultant",
                "Audit Manager", "CFO",
                # ── HR & People ───────────────────────────────────────────────
                "HR Manager", "HR Business Partner",
                "Talent Acquisition Specialist", "Recruiter",
                "Technical Recruiter", "HR Analyst",
                "Learning and Development Manager", "Compensation Analyst",
                # ── Sales & CRM ───────────────────────────────────────────────
                "Sales Manager", "Account Executive",
                "Business Development Manager", "Sales Engineer",
                "Technical Sales Engineer", "Pre-Sales Consultant",
                "Customer Success Manager",
                # ── Operations & Supply Chain ─────────────────────────────────
                "Operations Manager", "Supply Chain Analyst",
                "Logistics Manager", "Procurement Manager",
                # ── Niche Tech ────────────────────────────────────────────────
                "Game Developer", "Blockchain Developer",
                "AR/VR Developer", "IoT Engineer",
                "EdTech Developer", "HealthTech Developer",
                "Healthcare Software Engineer",
                # ── Writing & Documentation ───────────────────────────────────
                "Technical Writer", "API Documentation Specialist",
                # ── Support & Admin ───────────────────────────────────────────
                "IT Support Engineer", "Help Desk Analyst",
                "Systems Engineer",
            ]

            def _normalize_job_title(raw: str) -> tuple[str, bool]:
                """
                Returns (corrected_title, was_corrected).
                Tries an exact case-insensitive match first, then fuzzy.
                Threshold: 0.70 similarity — catches typos but won't misfire on
                completely unrelated inputs (e.g. 'Chef').
                """
                if not raw or not raw.strip():
                    return raw, False
                raw_stripped = raw.strip()
                raw_lower = raw_stripped.lower()

                # 1. Exact case-insensitive match → silently fix capitalisation only
                #    was_corrected = False here because only casing changed, not spelling
                for title in _CANONICAL_JOB_TITLES:
                    if title.lower() == raw_lower:
                        return title, False  # no toast — user spelled it right

                # 2. Fuzzy match against canonical list (case-insensitive compare)
                lower_map = {t.lower(): t for t in _CANONICAL_JOB_TITLES}
                matches = _difflib.get_close_matches(
                    raw_lower,
                    lower_map.keys(),
                    n=1,
                    cutoff=0.70,
                )
                if matches:
                    corrected = lower_map[matches[0]]
                    return corrected, True

                # 3. No confident match — return original unchanged
                return raw_stripped, False

            _raw_title = st.session_state.get("job_title", "").strip()
            _corrected_title, _was_corrected = _normalize_job_title(_raw_title)
            if _was_corrected:
                st.session_state["job_title"] = _corrected_title
                st.toast(f"✏️ Job title corrected: \"{_raw_title}\" → \"{_corrected_title}\"", icon="✅")
            # ── End Job Title Spell-Correction ────────────────────────────────────

            # ── Helper: detect if a field has real user-entered content ──
            def _has_real_content(value, min_len=4):
                if not value:
                    return False
                v = str(value).strip().lower()
                placeholders = {"placeholder", "sample", "n/a", "na", "none", "tbd", ""}
                return len(v) >= min_len and v not in placeholders

            # ── Normalize experience entries + detect if user provided real data ──
            experience_entries = st.session_state.get('experience_entries', [])
            normalized_experience_entries = []
            user_has_real_experience = False

            for entry in experience_entries:
                if isinstance(entry, dict):
                    title    = entry.get("title", "").strip()
                    company  = entry.get("company", "").strip()
                    duration = entry.get("duration", "").strip()
                    desc     = entry.get("description", "").strip()
                    if _has_real_content(company) or _has_real_content(desc):
                        user_has_real_experience = True
                    parts = []
                    if title:    parts.append(f"Role: {title}")
                    if company:  parts.append(f"Company: {company}")
                    if duration: parts.append(f"Duration: {duration}")
                    if desc:     parts.append(f"Description: {desc}")
                    formatted = "\n".join(parts)
                else:
                    formatted = entry.strip()
                    if _has_real_content(formatted, min_len=10):
                        user_has_real_experience = True
                if formatted:
                    normalized_experience_entries.append(formatted)

            # ── Normalize project entries + detect if user provided real data ──
            project_entries = st.session_state.get('project_entries', [])
            normalized_project_entries = []
            user_has_real_projects = False

            for entry in project_entries:
                if isinstance(entry, dict):
                    title    = entry.get("title", "").strip()
                    tech     = entry.get("tech", "").strip()
                    duration = entry.get("duration", "").strip()
                    desc     = entry.get("description", "").strip()
                    if _has_real_content(title) or _has_real_content(desc):
                        user_has_real_projects = True
                    parts = []
                    if title:    parts.append(f"Title: {title}")
                    if tech:     parts.append(f"Tech Stack: {tech}")
                    if duration: parts.append(f"Duration: {duration}")
                    if desc:     parts.append(f"Description: {desc}")
                    formatted = "\n".join(parts)
                else:
                    formatted = entry.strip()
                    if _has_real_content(formatted, min_len=10):
                        user_has_real_projects = True
                if formatted:
                    normalized_project_entries.append(formatted)

            # ── Build experience prompt section based on whether user has real data ──
            if user_has_real_experience:
                experience_instruction = f"""2. EXPERIENCE (USER HAS PROVIDED REAL DATA — LIGHT POLISH ONLY):
               The user has already entered their experience. Your ONLY job is to lightly polish the bullet descriptions.
               STRICT RULES — violating any of these is an error:
               - DO NOT change, rename, or reword any company name. Copy it exactly as given.
               - DO NOT change any role/job title. Copy it exactly as given.
               - DO NOT change any date or duration. Copy it exactly as given.
               - DO NOT add new entries that the user did not provide.
               - DO NOT restructure or reorder entries.
               - ONLY rewrite existing description bullets to be cleaner, more ATS-friendly, and results-oriented.
               - If a description is already strong, make minimal or no edits.
               Present as A., B., C. using the user's original data below:
               {normalized_experience_entries}"""
            else:
                experience_instruction = f"""2. EXPERIENCE (NO USER DATA PROVIDED — GENERATE REALISTIC DUMMY DATA):
               The user's target job title is: "{st.session_state['job_title']}".
               Generate 2–3 realistic experience entries showing a natural career progression toward this SPECIFIC role.

               COMPANY NAMING RULES:
               - Choose company names that are REALISTIC and DOMAIN-APPROPRIATE for "{st.session_state['job_title']}".
               - For tech/software roles: use companies like Google, Microsoft, Amazon, Flipkart, Razorpay, Zomato, Swiggy, PhonePe, Freshworks, Zoho, Paytm, Byju's, Ola, or similar product/tech companies.
               - For finance/banking roles: use HDFC Bank, ICICI Bank, Axis Bank, Kotak, JPMorgan, Goldman Sachs, KPMG, Deloitte, EY, or similar.
               - For data/analytics roles: use Mu Sigma, Fractal Analytics, ThoughtWorks, Tiger Analytics, or product companies with data teams.
               - For marketing/design/HR roles: use relevant advertising agencies, startups, or consumer brands.
               - NEVER use the same company list for every job title. Pick what makes sense for THIS domain.
               - Each entry MUST have a DIFFERENT company name.

               DATE RULES:
               - Dates must show logical progression (most recent first, oldest last).
               - NEVER repeat the same date range across entries.
               - Use realistic tenure lengths (1–3 years per role).

               CONTENT RULES:
               - Role titles must be SPECIFIC to "{st.session_state['job_title']}" — not generic.
               - Each entry must have 3–4 achievement bullets with measurable metrics relevant to this role.
               - Tools, technologies, and responsibilities must match what someone in "{st.session_state['job_title']}" actually does.
               Present as A., B., C. format."""

            # ── Build projects prompt section based on whether user has real data ──
            if user_has_real_projects:
                projects_instruction = f"""3. PROJECTS (USER HAS PROVIDED REAL DATA — LIGHT POLISH ONLY):
               The user has already entered their projects. Your ONLY job is to lightly improve the descriptions.
               STRICT RULES — violating any of these is an error:
               - DO NOT change, rename, or reword any project title. Copy it exactly as given.
               - DO NOT change any tech stack. Copy it exactly as given.
               - DO NOT change any duration/date. Copy it exactly as given.
               - DO NOT add new projects that the user did not provide.
               - DO NOT restructure or reorder entries.
               - ONLY rewrite existing description bullets to be more technical, impactful, and measurable.
               - If a description is already strong, make minimal or no edits.
               Present as A., B., C. using the user's original data below:
               {normalized_project_entries}"""
            else:
                projects_instruction = f"""3. PROJECTS (NO USER DATA PROVIDED — GENERATE REALISTIC DUMMY DATA):
               The user's target job title is: "{st.session_state['job_title']}".
               Generate 2–3 realistic, industry-standard projects that someone in THIS SPECIFIC ROLE would actually build.

               PROJECT NAMING RULES:
               - Project names MUST be derived directly from the domain of "{st.session_state['job_title']}".
               - Do NOT use any generic or pre-existing example names. Think from scratch for this role.
               - The name should sound like a real internal tool, product feature, or system — not a textbook exercise.
               - For example: a "Digital Marketing Manager" would NOT build a Kubernetes migration; they'd build a "Campaign Attribution Analytics Dashboard" or "SEO Content Performance Tracker".
               - A "Chartered Accountant" would NOT build a fraud detection pipeline; they'd build a "GST Reconciliation Automation Tool" or "Audit Trail Ledger System".
               - Generate names that ONLY make sense for "{st.session_state['job_title']}" — if the name could apply to a different role, discard it and try again.

               TECH STACK RULES:
               - Use ONLY tools, frameworks, and technologies that are standard for "{st.session_state['job_title']}".
               - Do NOT use backend/data engineering stacks (Kafka, Spark, Flink) for non-technical roles.
               - Do NOT use frontend stacks (React, TypeScript) for data or infrastructure roles.
               - Match the stack precisely to what this role uses day-to-day.

               DATE RULES:
               - All project dates must be DIFFERENT from each other.
               - Ordered most recent first. Use realistic 3–5 month project durations.

               CONTENT RULES:
               - Each project must have 3–5 strong technical/functional bullets with measurable outcomes.
               - Bullets must use vocabulary and actions that belong to "{st.session_state['job_title']}" — not generic software engineering language unless the role demands it.
               - NEVER use names like "Sample Project", "My Project", "Project 1", or any name from unrelated domains.
               Present as A., B., C. format."""

            # ── Build certificates prompt section ────────────────────────────────
            _cert_entries = st.session_state.get("certificate_links", [])
            _user_has_real_certs = any(
                _has_real_content(c.get("name", "")) for c in _cert_entries
            )
            if _user_has_real_certs:
                # Build structured representation passing ALL fields to the AI
                _cert_lines = []
                for c in _cert_entries:
                    _cname = c.get("name", "").strip()
                    _clink = c.get("link", "").strip()
                    _cdur  = c.get("duration", "").strip()
                    _cdesc = c.get("description", "").strip()
                    if not _cname:
                        continue
                    parts = [f"Name: {_cname}"]
                    if _clink:  parts.append(f"Link: {_clink}")
                    if _cdur:   parts.append(f"Date: {_cdur}")
                    if _cdesc:  parts.append(f"Description: {_cdesc}")
                    _cert_lines.append("\n".join(parts))
                _cert_data_str = "\n\n".join(_cert_lines)
                certificates_instruction = f"""8. CERTIFICATES (USER HAS PROVIDED REAL DATA — PRESERVE EVERYTHING):
               STRICT RULES — violating any of these is an error:
               - DO NOT change the certificate name. Copy it exactly as given.
               - DO NOT change, alter, or reformat the issued date in ANY way. Copy it character-for-character.
               - DO NOT change the verification link. Copy it exactly as given.
               - ONLY lightly polish the description if one was provided; otherwise leave it as-is.
               - DO NOT invent or add certificates the user did not provide.
               Output each certificate on one line as: [Name] - [Issuer] ([Date])
               where [Date] is EXACTLY the date the user entered — no reformatting, no substitution.
               User's certificate data:
               {_cert_data_str}"""
            else:
                certificates_instruction = f"""8. CERTIFICATES (NO USER DATA — GENERATE):
               Generate 3 realistic, industry-recognized certifications for {st.session_state['job_title']} with provider name."""

            enhance_prompt = f"""
            You are a professional Resume Optimization Specialist with deep expertise in ATS systems,
            industry hiring standards, and professional resume writing.
            Target role: "{st.session_state['job_title']}"

            ⚠️ CRITICAL DOMAIN RULE — READ BEFORE GENERATING ANYTHING:
            Every single piece of content you generate (projects, experience, skills, summary, certificates)
            MUST be tailored specifically and exclusively to the role: "{st.session_state['job_title']}".
            - A "Digital Marketing Manager" should NEVER have Java/Kafka/Kubernetes projects.
            - A "Java Backend Developer" should NEVER have Excel/VLOOKUP or GST reconciliation content.
            - A "Chartered Accountant" should NEVER have React/TypeScript or ML pipeline content.
            Before writing any section, ask yourself: "Would someone hiring a {st.session_state['job_title']} 
            care about this?" If no → discard and generate something domain-appropriate.
            The output must look like it was written BY a {st.session_state['job_title']} FOR a {st.session_state['job_title']} role.

            LANGUAGE & TONE:
            - Neutral, professional, ATS-optimized tone throughout.
            - NO first-person (I, me, my, we, our). NO gendered pronouns.
            - NO marketing terms (rockstar, guru, ninja).
            - Concise, quantifiable, outcome-focused language only.
            - Each section must use a DISTINCT verb set — no verb/phrase/action repeated across sections.

            SECTION LANGUAGE RULES:
            - SUMMARY: Third-person PRESENT tense. Strategic positioning bullets only.
            - EXPERIENCE: PAST tense. Ownership, delivery, accountability language.
            - PROJECTS: PAST tense. Use vocabulary natural to the {st.session_state['job_title']} domain.
            - SKILLS / SOFTSKILLS: Nouns only. Comma-separated list.
            - INTERESTS: Professional domain-engagement language.

            SECTION INSTRUCTIONS:

            1. SUMMARY:
               {"Enhance the provided summary." if _has_real_content(st.session_state.get('summary','')) else f"Generate a 3-4 bullet professional summary for a {st.session_state['job_title']}."}
               Write 3-4 bullets defining professional identity, specialization, and measurable strengths.

            {experience_instruction}

            {projects_instruction}

            4. SKILLS:
               {"Enhance and expand the provided skills list to be more specific and ATS-optimized for this role." if _has_real_content(st.session_state.get('skills','')) else f"Generate 6-8 current, highly specific technical/functional skills that are EXCLUSIVELY relevant to a {st.session_state['job_title']}. Do NOT list generic skills that apply to every role."}
               List only - no sentences.

            5. SOFTSKILLS:
               {"Enhance the provided soft skills to align with this role's demands." if _has_real_content(st.session_state.get('Softskills','')) else f"Generate 5-6 soft skills that are most valued specifically for a {st.session_state['job_title']} — not a generic list."}
               List only - no sentences.

            6. LANGUAGES:
               {"Use provided languages." if _has_real_content(st.session_state.get('languages','')) else "Generate 2-3 relevant languages (include English)."}

            7. INTERESTS:
               {"Enhance provided interests." if _has_real_content(st.session_state.get('interests','')) else f"Generate 3-5 professional interests aligned with {st.session_state['job_title']}."}

            {certificates_instruction}

            OUTPUT FORMAT (FOLLOW EXACTLY):

            Summary:
            * [bullet]
            * [bullet]
            * [bullet]

            Experience:
            A. [Company Name] ([Start Month Year - End Month Year or Present])
               * [Role Title]
               * [Achievement with metric]
               * [Ownership/delivery bullet]
               * [Impact/improvement bullet]

            B. [Company Name] ([Start Month Year - End Month Year])
               * [Role Title]
               * [Achievement]
               * [Responsibility]

            Projects:
            A. [Unique Realistic Project Name]
               * Tech Stack: [tools]
               * Duration: [Start Month Year - End Month Year]
               * Description:
                 - [bullet]
                 - [bullet]
                 - [bullet]

            B. [Unique Realistic Project Name]
               * Tech Stack: [tools]
               * Duration: [Start Month Year - End Month Year]
               * Description:
                 - [bullet]
                 - [bullet]
                 - [bullet]

            Skills:
            [skill1], [skill2], [skill3], [skill4], [skill5], [skill6]

            SoftSkills:
            [soft1], [soft2], [soft3], [soft4], [soft5]

            Languages:
            [lang1], [lang2]

            Interests:
            [interest1], [interest2], [interest3]

            Certificates:
            [Certificate] - [Provider] ([Year/Level])
            [Certificate] - [Provider] ([Year/Level])

            SOURCE DATA TO ENHANCE (where provided):
            Summary: {st.session_state.get('summary', '')}
            Skills: {st.session_state.get('skills', '')}
            SoftSkills: {st.session_state.get('Softskills', '')}
            Languages: {st.session_state.get('languages', '')}
            Interests: {st.session_state.get('interests', '')}
            Certificates: {[{"name": c.get("name",""), "date": c.get("duration",""), "link": c.get("link","")} for c in st.session_state.get('certificate_links', []) if c.get('name')]}

            CRITICAL RULES:
            - Output ONLY the formatted resume content. No explanations, no preamble.
            - NEVER use "Sample Project", "Previous Company", "Placeholder", or any generic names.
            - ALL dates must be DIFFERENT across experience entries AND across project entries.
            - Experience dates must show logical career progression (most recent first, oldest last).
            - Project dates must all be different and logically ordered (most recent first).
            - If user provided real experience/project data, PRESERVE every company name, project title, tech stack, and date EXACTLY as written. Only polish the description bullets.
            - If user provided real certificate data, PRESERVE the certificate name, issued date, and link EXACTLY as written. Never substitute or reformat the date.
            - UNIQUENESS RULE: Every generation must produce fresh, original content. Never repeat the same project names, company names, or bullet phrasing across different runs. Treat each generation as a brand-new resume for a brand-new person.
            - DOMAIN LOCK: Every project name, tech stack, skill, and certificate must be something a real "{st.session_state['job_title']}" would have. Cross-domain content is forbidden.
            """





            import uuid as _uuid
            import datetime as _datetime
            _unique_seed = _uuid.uuid4().hex[:8]
            _timestamp = _datetime.datetime.now().strftime("%H%M%S")
            enhance_prompt += f"\n[Generation ID: {_unique_seed}-{_timestamp} — produce content unique to this exact run]"

            with st.spinner("🧠 Thinking..."):
                ai_output = call_llm(enhance_prompt, session=st.session_state)
                st.session_state["ai_output"] = ai_output

    # ------------------------- PARSE + RENDER -------------------------
    if "ai_output" in st.session_state:
        ai_output = st.session_state["ai_output"]

        def extract_section(label, output, default=""):
            pattern = rf"(?m)^{re.escape(label)}:\s*\n?(.*?)(?=\n[A-Za-z][A-Za-z\s]*:\s*\n?|\Z)"
            match = re.search(pattern, output, re.DOTALL)
            return match.group(1).strip() if match else default

        summary_enhanced  = extract_section("Summary",      ai_output, st.session_state["summary"])
        experience_raw    = extract_section("Experience",   ai_output)
        projects_raw      = extract_section("Projects",     ai_output)
        skills_list       = extract_section("Skills",       ai_output, st.session_state["skills"])
        softskills_list   = extract_section("SoftSkills",   ai_output, st.session_state["Softskills"])
        languages_list    = extract_section("Languages",    ai_output, st.session_state["languages"])
        interests_list    = extract_section("Interests",    ai_output, st.session_state["interests"])
        certificates_list = extract_section("Certificates", ai_output)

        experience_blocks = [b.strip() for b in re.split(r"\n(?=[A-Z]\. )", experience_raw.strip()) if b.strip()]
        projects_blocks   = [b.strip() for b in re.split(r"\n(?=[A-Z]\. )", projects_raw.strip())   if b.strip()]

        left, right = st.columns([1, 2])

        with left:
            st.markdown(
                f"<h2 style='color:#2f2f2f;margin-bottom:0;'>{st.session_state['name']}</h2>"
                f"<h4 style='margin-top:5px;color:#444;'>{st.session_state['job_title']}</h4>"
                f"<p style='font-size:14px;'>"
                f"📍 {st.session_state['location']}<br>"
                f"📞 {st.session_state['phone']}<br>"
                f"📧 <a href='mailto:{st.session_state['email']}'>{st.session_state['email']}</a><br>"
                f"🔗 <a href='{st.session_state['linkedin']}' target='_blank'>LinkedIn</a><br>"
                f"🌐 <a href='{st.session_state['portfolio']}' target='_blank'>Portfolio</a></p>",
                unsafe_allow_html=True
            )

            def render_bullet_section(title, items):
                st.markdown(f"<h4 style='color:#336699;'>{title}</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for item in [i.strip() for i in items.split(",") if i.strip()]:
                    st.markdown(f"<div style='margin-left:10px;'>• {item}</div>", unsafe_allow_html=True)

            render_bullet_section("Skills",      skills_list)
            render_bullet_section("Languages",   languages_list)
            render_bullet_section("Interests",   interests_list)
            render_bullet_section("Soft Skills", softskills_list)

        with right:
            formatted_summary = summary_enhanced.replace("\n• ", "<br>• ").replace("\n* ", "<br>• ").replace("\n", "<br>")
            st.markdown("<h4 style='color:#336699;'>Summary</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            st.markdown(f"<p style='font-size:17px;'>{formatted_summary}</p>", unsafe_allow_html=True)

            if experience_blocks:
                st.markdown("<h4 style='color:#336699;'>Experience</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for idx, exp_block in enumerate(experience_blocks):
                    lines_exp = [l for l in exp_block.strip().split("\n") if l.strip()]
                    if not lines_exp:
                        continue
                    heading = lines_exp[0]
                    m = re.match(r"[A-Z]\.\s*(.+?)\s*\((.+?)\)\s*$", heading)
                    if m:
                        company, duration = m.group(1).strip(), m.group(2).strip()
                    else:
                        m2 = re.match(r"[A-Z]\.\s*(.+?)\s*[\u2014\u2013-]+\s*(.+)$", heading)
                        if m2:
                            company, duration = m2.group(1).strip(), m2.group(2).strip()
                        else:
                            company  = re.sub(r"^[A-Z]\.\s*", "", heading).strip()
                            duration = ""
                    role = ""
                    bullet_lines = []
                    for line in lines_exp[1:]:
                        stripped = line.strip().lstrip("\u2022\u00b7*-\u2013\u2014 ").strip()
                        if not role and len(stripped) <= 60 and not re.search(r"\d+%|\d+ [a-z]", stripped):
                            role = stripped
                        else:
                            bullet_lines.append(line.strip())
                    if not role:
                        ss_entries = st.session_state.get("experience_entries", [])
                        role = ss_entries[idx].get("title", "") if idx < len(ss_entries) else ""
                    formatted_exp = "".join(
                        f"<div style='margin-left:12px;margin-bottom:4px;'>• {l.lstrip('\u2022\u00b7*-\u2013\u2014 ').strip()}</div>"
                        for l in bullet_lines if l.strip()
                    )
                    st.markdown(
                        f"<div style='margin-bottom:15px;padding:10px;border-radius:8px;border-left:3px solid #336699;'>"
                        f"<div style='display:flex;justify-content:space-between;flex-wrap:wrap;gap:4px;'>"
                        f"<b style='font-size:15px;'>🏢 {company}</b>"
                        f"<span style='color:gray;font-size:13px;'>📆 {duration}</span></div>"
                        f"<div style='font-size:14px;margin-top:3px;'>💼 <i>{role}</i></div>"
                        f"<div style='font-size:14px;margin-top:6px;'>{formatted_exp}</div></div>",
                        unsafe_allow_html=True
                    )

            edu_to_show = [e for e in st.session_state.education_entries if e.get("institution") or e.get("degree")]
            if edu_to_show:
                st.markdown("<h4 style='color:#336699;'>🎓 Education</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for edu in edu_to_show:
                    degree_val = edu.get("degree", "")
                    if isinstance(degree_val, list):
                        degree_val = ", ".join(degree_val)
                    st.markdown(
                        f"<div style='margin-bottom:15px;padding:10px 15px;border-radius:8px;border-left:3px solid #336699;'>"
                        f"<div style='display:flex;justify-content:space-between;font-size:16px;font-weight:bold;flex-wrap:wrap;gap:4px;'>"
                        f"<span>🏫 {edu.get('institution','')}</span>"
                        f"<span style='color:gray;font-size:13px;'>📅 {edu.get('year','')}</span></div>"
                        f"<div style='font-size:14px;'>🎓 <i>{degree_val}</i></div>"
                        f"<div style='font-size:14px;color:#555;'>📄 {edu.get('details','')}</div></div>",
                        unsafe_allow_html=True
                    )

            if projects_blocks:
                st.markdown("<h4 style='color:#336699;'>Projects</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for idx, proj_block in enumerate(projects_blocks):
                    plines = proj_block.strip().split("\n")
                    label  = chr(65 + idx)
                    ai_title = re.sub(r"^[A-Z]\.\s*", "", plines[0]).strip() if plines else ""
                    ai_tech = ai_duration = ""
                    desc_lines = []
                    in_desc = False
                    for line in plines[1:]:
                        stripped = line.strip()
                        tl = stripped.lstrip("\u2022\u00b7*-\u2013\u2014 ").strip()
                        if re.match(r"Tech\s*Stack\s*:", tl, re.I):
                            ai_tech = re.sub(r"(?i)^Tech\s*Stack\s*:\s*", "", tl).strip()
                        elif re.match(r"Duration\s*:", tl, re.I):
                            ai_duration = re.sub(r"(?i)^Duration\s*:\s*", "", tl).strip()
                        elif re.match(r"Description\s*:", tl, re.I):
                            in_desc = True
                            remainder = re.sub(r"(?i)^Description\s*:\s*", "", tl).strip()
                            if remainder:
                                desc_lines.append(remainder)
                        elif in_desc and stripped:
                            desc_lines.append(stripped)
                    ss_proj = st.session_state.project_entries[idx] if idx < len(st.session_state.project_entries) else {}
                    final_title    = ai_title    or ss_proj.get("title", "")
                    final_tech     = ai_tech     or ss_proj.get("tech", "")
                    final_duration = ai_duration or ss_proj.get("duration", "")
                    formatted_proj = "".join(
                        f"<div style='margin-left:12px;margin-bottom:4px;'>• {dl.lstrip('\u2022\u00b7*-\u2013\u2014 ').strip()}</div>"
                        for dl in desc_lines if dl.strip()
                    )
                    st.markdown(
                        f"<div style='margin-bottom:15px;padding:10px;border-radius:8px;border-left:3px solid #336699;'>"
                        f"<strong style='font-size:16px;'>📌 <span style='color:#444;'>{label}. </span>{final_title}</strong><br>"
                        f"<span style='font-size:13px;color:#555;'>🛠️ <strong>Tech Stack:</strong> {final_tech}</span><br>"
                        f"<span style='font-size:13px;color:#555;'>⏳ <strong>Duration:</strong> {final_duration}</span><br>"
                        f"<div style='font-size:14px;margin-top:6px;'>{formatted_proj}</div></div>",
                        unsafe_allow_html=True
                    )

            # ── Certificate rendering — strip AI-invented "Unknown" placeholders ──
            # Build a name→date lookup from session state so we can fall back to
            # the user's real data if the AI mangles or omits it.
            _ss_cert_lookup = {
                c.get("name", "").strip(): c.get("duration", "").strip()
                for c in st.session_state.get("certificate_links", [])
                if c.get("name", "").strip()
            }

            _cert_lines_raw = [c.strip() for c in certificates_list.split("\n") if c.strip()] if certificates_list else []

            # Fall back to session state when AI produced nothing or only garbage
            if not _cert_lines_raw and _ss_cert_lookup:
                _cert_lines_raw = [
                    f"{name} ({date})" if date else name
                    for name, date in _ss_cert_lookup.items()
                ]

            if _cert_lines_raw:
                st.markdown("<h4 style='color:#336699;'>📜 Certificates</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for _cert_line in _cert_lines_raw:
                    # Remove any " - Unknown" or "Unknown - " the AI injected
                    _cleaned = re.sub(r'\s*-\s*Unknown\b', '', _cert_line, flags=re.IGNORECASE).strip()
                    _cleaned = re.sub(r'\bUnknown\s*-\s*', '', _cleaned, flags=re.IGNORECASE).strip()
                    _cleaned = re.sub(r'\bUnknown\b', '', _cleaned, flags=re.IGNORECASE).strip(" -–—").strip()

                    # If AI dropped the date, re-inject it from session state
                    for _ss_name, _ss_date in _ss_cert_lookup.items():
                        if _ss_name.lower() in _cleaned.lower() and _ss_date and _ss_date not in _cleaned:
                            _cleaned = f"{_cleaned} ({_ss_date})"
                            break

                    if _cleaned:
                        st.markdown(f"<div style='margin-left:10px;margin-bottom:4px;'>• {_cleaned}</div>", unsafe_allow_html=True)

            if st.session_state.project_links:
                st.markdown("<h4 style='color:#336699;'>Project Links</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for i, link in enumerate(st.session_state.project_links):
                    st.markdown(f"[🔗 Project {i+1}]({link})", unsafe_allow_html=True)

    # Generate HTML content based on selected template — only on submit, stored in session_state
    if submitted:
        with st.spinner("⚙️ Generating your resume... please wait"):
            # Render selected resume template via the registry dispatcher (resume_builder.py)
            html_content = render_resume(selected_template, st.session_state, profile_img_html)

            # Store the generated content and invalidate cached PDF so it's recomputed fresh
            # NOTE: Use direct assignment instead of .pop() — .pop() on an existing key
            # triggers an extra Streamlit rerun which causes visible page blinking.
            st.session_state["generated_html"] = html_content
            st.session_state["pdf_resume_bytes"] = None   # invalidate cache without extra rerun
            st.session_state["show_template_preview"] = False
        st.session_state.pop("_resume_generating", None)

with tab2:
    # ==========================
    # 📥 Resume Download Header
    # ==========================
    if "generated_html" in st.session_state:
        st.markdown(
            """
            <div style='text-align: center; margin-top: 20px; margin-bottom: 30px;'>
                <h2 style='color: #2f4f6f; font-family: Arial, sans-serif; font-size: 24px;'>
                    📥 Download Your Resume
                </h2>
                <p style="color:#555; font-size:14px;">
                    Choose your preferred format below
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )

        # Cache PDF bytes in session_state to avoid expensive recomputation on every rerun
        if not st.session_state.get("pdf_resume_bytes"):
            st.session_state["pdf_resume_bytes"] = html_to_pdf_bytes(
                st.session_state["generated_html"]
            ).read()

        col1, spacer, col2 = st.columns([1, 0.15, 0.85])

        # HTML Resume Download Button
        with col1:
            html_bytes = st.session_state["generated_html"].encode("utf-8")
            html_file = BytesIO(html_bytes)

            st.download_button(
                label="⬇️ Download as Template",
                data=html_file,
                file_name=f"{st.session_state['name'].replace(' ', '_')}_Resume.html",
                mime="text/html",
                key="download_resume_html"
            )

        # Preview Template Button — smart toggle: spinner only when opening, instant when closing
        with col2:
            is_previewing = st.session_state.get("show_template_preview", False)
            if st.button("👁️ Preview Template", key="preview_template_btn"):
                if not is_previewing:
                    # Opening — show spinner since we're loading the iframe
                    with st.spinner("Loading template preview..."):
                        time.sleep(2)
                        st.session_state["show_template_preview"] = True
                else:
                    # Closing — instant, no spinner
                    st.session_state["show_template_preview"] = False

        # Show/hide the template preview iframe
        if st.session_state.get("show_template_preview", False):
            import streamlit.components.v1 as components
            st.markdown(
                "<p style='color:#555; font-size:13px; margin-top:8px;'>"
                "📄 Template Preview (scroll to explore):</p>",
                unsafe_allow_html=True,
            )
            components.html(
                st.session_state["generated_html"],
                height=600,
                scrolling=True,
            )

        # PDF Resume Download Button — use cached bytes
        pdf_resume_bytes = BytesIO(st.session_state["pdf_resume_bytes"])
        
        # ✅ Extra Help Note
        st.markdown("""
        ✅ After downloading your HTML resume, you can 
        <a href="https://www.sejda.com/html-to-pdf" target="_blank" style="color:#2f4f6f; text-decoration:none;">
        convert it to PDF using Sejda's free online tool</a>.
        """, unsafe_allow_html=True)

        # ==========================
        # 📩 Cover Letter Expander
        # ==========================
        with st.expander("📩 Generate Cover Letter from This Resume"):
            generate_cover_letter_from_resume_builder()

        # ==========================
        # ✉️ Generated Cover Letter Downloads (NO PREVIEW HERE)
        # ==========================
        if "cover_letter" in st.session_state:
            st.markdown(
                """
                <div style="margin-top: 30px; margin-bottom: 20px;">
                    <h3 style="color: #003366;">✉️ Generated Cover Letter</h3>
                    <p style="color:#555; font-size:14px;">
                        You can download your generated cover letter in multiple formats.
                    </p>
                </div>
                """,
                unsafe_allow_html=True
            )

            # ✅ Use already-rendered HTML from session (don't show again)
            styled_cover_letter = st.session_state.get("cover_letter_html", "")

            # ✅ Generate PDF from styled HTML
            pdf_file = html_to_pdf_bytes(styled_cover_letter)

            # ✅ DOCX Generator (preserves line breaks)
            def create_docx_from_text(text, filename="cover_letter.docx"):
                from docx import Document
                bio = BytesIO()
                doc = Document()
                doc.add_heading("Cover Letter", 0)

                for line in text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)
                    else:
                        doc.add_paragraph("")  # preserve empty lines

                doc.save(bio)
                bio.seek(0)
                return bio

            # ==========================
            # 📥 Cover Letter Download Buttons
            # ==========================
            st.markdown("""
            <div style="margin-top: 25px; margin-bottom: 15px;">
                <strong>⬇️ Download Your Cover Letter:</strong>
            </div>
            """, unsafe_allow_html=True)

            col1,col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📥 Download Cover Letter (.docx)",
                    data=create_docx_from_text(st.session_state["cover_letter"]),
                    file_name=f"{st.session_state['name'].replace(' ', '_')}_Cover_Letter.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_coverletter_docx"
                )
            
            with col2:
                st.download_button(
                    label="📥 Download Cover Letter (Template)",
                    data=styled_cover_letter.encode("utf-8"),
                    file_name=f"{st.session_state['name'].replace(' ', '_')}_Cover_Letter.html",
                    mime="text/html",
                    key="download_coverletter_html"
                )

            # ✅ Helper note
            st.markdown("""
            ✅ If the HTML cover letter doesn't display properly, you can 
            <a href="https://www.sejda.com/html-to-pdf" target="_blank" style="color:#2f4f6f; text-decoration:none;">
            convert it to PDF using Sejda's free online tool</a>.
            """, unsafe_allow_html=True)
import streamlit as st


# ── Sub-module imports ────────────────────────────────────────
from tab3_backend import init_job_search_db
from tab3_ui import (
    _inject_tab3_css,
    _job_search_interactive,
    _analytics_dashboard,
    render_featured_companies,
    render_market_trends,
    render_salary_insights,
)

# ── One-time DB initialisation ────────────────────────────────
# The flag lives in st.session_state so it is skipped on all
# subsequent reruns within the same browser session.
# The @st.cache_resource connection already ensures the
# psycopg2 connect() call itself is not repeated either.
if not st.session_state.get("_db_initialized"):
    init_job_search_db()
    st.session_state["_db_initialized"] = True

# ── Tab 3 rendering ───────────────────────────────────────────
with tab3:
    _inject_tab3_css()
    _job_search_interactive()
    _analytics_dashboard()          # ← separate fragment: updates instantly after search

    # ---------- Featured Companies ----------
    render_featured_companies()

    # ---------- Market Insights ----------
    render_market_trends()

    # ---------- Salary Insights ----------
    render_salary_insights()

def evaluate_interview_answer(answer: str, question: str = None):
    """
    Uses an LLM to strictly evaluate an interview answer.
    Returns (score out of 5, feedback string).
    """
    from llm_manager import call_llm
    import re
    import streamlit as st

    # Empty check
    if not answer.strip():
        return 0, "⚠️ No answer provided."

    # 🔹 LLM Prompt (STRICTER)
    prompt = f"""
    You are an expert technical interview evaluator.

    ### Task:
    Evaluate the candidate's answer to the question below.
    Be STRICT. Only give high scores if the answer is technically correct, relevant, and detailed.

    ### Question:
    {question if question else "N/A"}

    ### Candidate Answer:
    {answer}

    ### Strict Scoring Rubric:
    - 5 = Exceptional: Fully correct, highly relevant, clear, detailed, technically accurate.
    - 4 = Good: Mostly correct and relevant, but missing some depth/clarity.
    - 3 = Average: Partially correct OR generic, but somewhat relevant.
    - 2 = Weak: Mostly irrelevant, shallow, or major gaps in correctness.
    - 1 = Poor: Completely irrelevant, incoherent, or very wrong.
    - 0 = No answer / total nonsense.

    ### Output Format:
    Score: <number between 0 and 5>
    Feedback: <constructive feedback in 1–2 sentences>
    """

    try:
        # Call LLM
        response = call_llm(prompt, session=st.session_state).strip()

        # Extract Score
        score_match = re.search(r"Score:\s*(\d+)", response)
        score = int(score_match.group(1)) if score_match else 1  # stricter fallback

        # Extract Feedback
        feedback_match = re.search(r"Feedback:\s*(.+)", response)
        feedback = feedback_match.group(1).strip() if feedback_match else "Answer was unclear or irrelevant."

        # ✅ Keep score in 0–5 range
        score = max(0, min(score, 5))

    except Exception as e:
        score = 1
        feedback = f"⚠️ Evaluation fallback due to error: {e}"

    return score, feedback


def format_score(score) -> str:
    """
    Uniform score formatter for all UI display.
    Always returns a 2-decimal-place string (e.g. 6.47, 6.50, 6.00).
    Returns 'N/A' for None / NaN values.
    Raw database values are never modified — formatting is display-layer only.
    """
    import math
    if score is None:
        return "N/A"
    try:
        val = float(score)
        if math.isnan(val):
            return "N/A"
        return f"{val:.2f}"
    except (TypeError, ValueError):
        return "N/A"


def evaluate_interview_answer_for_scores(answer: str, question: str, difficulty: str, role: str = "", domain: str = ""):
    """
    UPGRADED: Intelligent evaluation with chain-of-thought reasoning and structured feedback.
    Uses JSON-based parsing for robustness and provides detailed, actionable feedback.

    Returns dict with keys: knowledge, communication, relevance, feedback (list), followup

    Features:
    - Chain-of-thought evaluation: extracts key concepts, identifies strengths/gaps
    - Structured feedback: detailed paragraph with specific, actionable insights
    - Difficulty calibration: Easy (encouraging), Medium (balanced), Hard (strict)
    - JSON-based parsing for reliability
    """
    from llm_manager import call_llm
    import json
    import streamlit as st

    # Empty check or junk answers
    if not answer.strip() or answer == "⚠️ No Answer" or len(answer.strip()) < 3:
        return {
            "knowledge": 0,
            "communication": 0,
            "relevance": 0,
            "feedback": "No answer provided. Try using the STAR method: Situation, Task, Action, Result. Provide specific examples from your experience to demonstrate your understanding and capabilities.",
            "followup": ""
        }

    # Check for obvious junk answers (single character, just symbols, etc.)
    if len(answer.strip()) == 1 or not any(c.isalnum() for c in answer):
        return {
            "knowledge": 0,
            "communication": 0,
            "relevance": 0,
            "feedback": "Answer appears incomplete or invalid. Please provide a meaningful response with technical details and structure your answer clearly with concrete examples from your experience.",
            "followup": ""
        }

    # STRICTER JUNK FILTERING: Check word count and meaningful tokens
    words = answer.strip().split()
    meaningful_words = [w for w in words if len(w) > 2 and any(c.isalpha() for c in w)]

    if len(words) < 5 or len(meaningful_words) < 2:
        return {
            "knowledge": 0,
            "communication": 0,
            "relevance": 0,
            "feedback": "Answer too short or lacks substance. Provide a detailed response with at least 3-4 sentences and include specific examples or technical details to demonstrate your understanding.",
            "followup": ""
        }

    # Difficulty-based evaluation guidance
    # TEXT-INTERVIEW-OPTIMISED difficulty guidance.
    # Hard is senior-level but scoped to ONE focused challenge — not whiteboard mega-design.
    difficulty_guidance = {
        "Easy": {
            "tone": "encouraging and patient",
            "expectations": (
                "Concept clarity and accurate definitions. The candidate should explain WHAT something is, "
                "WHY it exists, and give a simple real-world example. No implementation depth required."
            ),
            "scoring": (
                "5-10 for clear definitions with a correct example. "
                "3-4 for partially correct or vague answers that show some understanding. "
                "0-2 for wrong definitions or no answer."
            ),
            "feedback_style": (
                "Encouraging. Praise correct parts. Point out the one or two missing elements "
                "with a simple suggestion. Keep feedback under 3 paragraphs."
            ),
            "answer_scope": "3-5 structured paragraphs",
        },
        "Medium": {
            "tone": "balanced and scenario-focused",
            "expectations": (
                "Scenario reasoning with ONE practical constraint. The candidate should describe their "
                "approach, make ONE explicit decision or tradeoff, and briefly justify it with a real example. "
                "No multi-layer system design required."
            ),
            "scoring": (
                "7-10 for answers that frame the scenario, make a clear decision, and justify it with reasoning. "
                "4-6 for answers that address the scenario but miss the decision logic or give only conceptual responses. "
                "0-3 for answers that treat it like an Easy question (pure definition) or are off-topic."
            ),
            "feedback_style": (
                "Constructive. Acknowledge the scenario framing they used, then identify the ONE key "
                "reasoning step they missed. Give a concrete direction for improvement in 4-5 paragraphs."
            ),
            "answer_scope": "5-6 structured paragraphs",
        },
        "Hard": {
            "tone": "precise and technically demanding",
            "expectations": (
                "Deep technical reasoning on ONE focused challenge — either a tradeoff analysis, a failure handling "
                "scenario, or an optimisation under constraint. NOT a full system design. "
                "The candidate should: state their reasoning framework, analyse the core challenge, "
                "explain their decision with specific technical justification, and mention one edge case or risk. "
                "Answerable in 6-8 paragraphs — no whiteboard required."
            ),
            "scoring": (
                "8-10 for answers that isolate the core challenge, reason through it with technical specifics, "
                "make a justified decision, and acknowledge a risk or edge case. "
                "5-7 for answers that address the challenge but stay too high-level or skip justification. "
                "0-4 for vague, off-topic, or purely conceptual answers."
            ),
            "feedback_style": (
                "Precise and senior-level. Identify exactly WHERE the reasoning stopped — was it before the tradeoff, "
                "before the edge case, or before quantification? Give one concrete example of what a strong answer "
                "would have added. 5-6 focused paragraphs."
            ),
            "answer_scope": "6-8 structured paragraphs",
        },
    }

    guidance = difficulty_guidance.get(difficulty, difficulty_guidance["Medium"])

    # Build context for relevance checking
    context_info = f" for {role} in {domain}" if role and domain else ""

    # UPGRADED CHAIN-OF-THOUGHT EVALUATION PROMPT
    prompt = f"""You are an expert technical interviewer evaluating a candidate's answer{context_info}.

QUESTION: {question}
CANDIDATE'S ANSWER: {answer}
DIFFICULTY LEVEL: {difficulty}

EVALUATION APPROACH — {difficulty.upper()} MODE ({guidance['tone']}):
What to expect: {guidance['expectations']}
Scoring guide: {guidance['scoring']}
Feedback style: {guidance['feedback_style']}
Expected answer scope: {guidance['answer_scope']}

STEP-BY-STEP EVALUATION PROCESS:

STEP 1 — IDENTIFY THE QUESTION'S CORE CHALLENGE:
State in one sentence what this question is actually testing (concept recall / scenario reasoning / focused technical analysis).
List 3-5 key concepts or reasoning moves a strong answer must include.

STEP 2 — ANALYSE THE CANDIDATE'S ANSWER:
✅ WHAT THEY GOT RIGHT: Which key concepts did they cover? What reasoning was correct or well-expressed?
⚠️ WHAT IS MISSING OR WEAK: Which expected concepts or reasoning steps are absent, shallow, or wrong?
🔴 SCOPE CHECK: Did the answer stay within the question's scope, or did it over-engineer / under-explain?

STEP 3 — SCORE ON 3 DIMENSIONS (1-10 each):
- Knowledge: Correctness and depth of technical content for THIS difficulty tier.
- Communication: Clarity, logical structure, and how easy it is to follow the reasoning.
- Relevance: How directly the answer addresses the specific question asked — not adjacent topics.

STEP 4 — WRITE FEEDBACK ({guidance['answer_scope']} equivalent):
Write {{"Easy": "2-3", "Medium": "3-4", "Hard": "4-5"}}.get(difficulty, "3-4") flowing paragraphs that:
1. Start with what the candidate did well (be specific — quote or paraphrase their answer)
2. Identify the ONE or TWO most important gaps for this difficulty level
3. Give a concrete, actionable suggestion — what would a stronger answer have included?
4. For Hard: note whether the answer stayed text-answerable and focused, or drifted into vague system design

Do NOT write bullet points. Write as a knowledgeable interviewer giving verbal feedback.

{"STEP 5 — FOLLOW-UP: Generate ONE tightly scoped follow-up question. It must: (a) directly reference something in their answer, (b) probe ONE specific gap identified above, (c) be answerable in 4-6 paragraphs of text — not a whiteboard exercise. Choose from: Metric Justification, Tradeoff Challenge, Edge Case Scenario, Failure Handling, Constraint Injection, or Depth Probe." if difficulty == "Hard" else ""}

OUTPUT FORMAT (strict JSON):
{{
  "key_concepts": ["concept1", "concept2", "concept3"],
  "strengths": ["strength1", "strength2"],
  "gaps": ["gap1", "gap2"],
  "knowledge": <number 1-10>,
  "communication": <number 1-10>,
  "relevance": <number 1-10>,
  "feedback": "Detailed, comprehensive feedback in 2-4 flowing paragraphs. Be specific about what the candidate did well, what they missed, and how they can improve. Reference actual content from their answer. Make it constructive, actionable, and personalized."{',\n  "followup": "One probing follow-up question"' if difficulty == "Hard" else ''}
}}

IMPORTANT RULES:
- If answer is off-topic or from wrong domain, set relevance to 0-2
- If answer is junk/minimal, set all scores to 0-2
- Feedback must be specific to THIS answer, not generic templates
- Reference actual content from the candidate's answer in feedback
- Each feedback point should feel personalized and human

Provide ONLY the JSON output, no additional text."""

    try:
        response = call_llm(prompt, session=st.session_state).strip()

        # Clean response - remove markdown code blocks if present
        if response.startswith("```"):
            response = response.split("```")[1]
            if response.startswith("json"):
                response = response[4:]
            response = response.strip()

        # Parse JSON response
        result = json.loads(response)

        # Extract and validate scores
        knowledge = int(result.get("knowledge", 1))
        communication = int(result.get("communication", 1))
        relevance = int(result.get("relevance", 1))

        # Clamp scores to 0-10 range
        knowledge = max(0, min(10, knowledge))
        communication = max(0, min(10, communication))
        relevance = max(0, min(10, relevance))

        # Extract feedback (should be a detailed string, not a list)
        feedback = result.get("feedback", "")

        # If feedback comes as a list (fallback), join it into paragraphs
        if isinstance(feedback, list):
            feedback = "\n\n".join(feedback)

        # Ensure we have substantial feedback
        if not feedback or len(feedback.strip()) < 50:
            feedback = "Your answer shows some understanding, but could benefit from more technical depth and specific examples. Consider structuring your response more clearly and providing concrete details from your experience. Focus on addressing all aspects of the question comprehensively."

        # ── POST-PROCESSING CALIBRATION ──────────────────────────────────────────
        # Safeguard 1: keyword overlap check — clamp knowledge/relevance if answer
        # contains almost none of the question's key terms
        key_terms = [kw.lower() for kw in question.split() if len(kw) > 3]
        match_count = sum(1 for t in key_terms if t in answer.lower())
        if key_terms and match_count < max(2, len(key_terms) // 10):
            knowledge = min(knowledge, 3)
            relevance = min(relevance, 3)

        # Safeguard 2: off-topic or low-relevance signal from LLM
        if "off-topic" in feedback.lower() or relevance < 4:
            knowledge = min(knowledge, 4)
            relevance = min(relevance, 4)

        # Safeguard 3: if average would be >4 but relevance is very low, enforce ceiling
        raw_avg = (knowledge + communication + relevance) / 3
        if relevance <= 2 and raw_avg > 4.0:
            # Wrong/irrelevant answers: all scores clamped to ≤ 2
            knowledge = min(knowledge, 2)
            communication = min(communication, 2)
            relevance = min(relevance, 2)
        elif relevance < 4 and raw_avg > 5.0:
            # Generic/partially relevant: clamp average ≤ 5
            excess = raw_avg - 5.0
            knowledge = max(0, min(knowledge, knowledge - int(excess * 1.5 + 0.5)))
            relevance = max(0, min(relevance, relevance - 1))
        # ─────────────────────────────────────────────────────────────────────────

        # Extract follow-up question
        followup = result.get("followup", "") if difficulty == "Hard" else ""

        return {
            "knowledge": knowledge,
            "communication": communication,
            "relevance": relevance,
            "feedback": feedback,  # Now a string, not a list
            "followup": followup
        }

    except json.JSONDecodeError as e:
        # Fallback: try to extract scores from non-JSON response
        import re
        try:
            knowledge = int(re.search(r'"?knowledge"?\s*:\s*(\d+)', response, re.IGNORECASE).group(1))
            communication = int(re.search(r'"?communication"?\s*:\s*(\d+)', response, re.IGNORECASE).group(1))
            relevance = int(re.search(r'"?relevance"?\s*:\s*(\d+)', response, re.IGNORECASE).group(1))

            # Extract feedback (try both string and array format)
            feedback_match = re.search(r'"feedback"\s*:\s*"([^"]+)"', response, re.DOTALL)
            if feedback_match:
                feedback = feedback_match.group(1)
            else:
                # Fallback: try array format and join
                feedback_array_match = re.search(r'"feedback"\s*:\s*\[(.*?)\]', response, re.DOTALL)
                if feedback_array_match:
                    feedback_text = feedback_array_match.group(1)
                    feedback_items = [f.strip(' "\'') for f in re.findall(r'"([^"]+)"', feedback_text)]
                    feedback = "\n\n".join(feedback_items) if feedback_items else "Answer evaluated but formatting unclear. Provide more structured responses with clear examples and explanations."
                else:
                    feedback = "Answer evaluated but formatting unclear. Provide more structured responses with clear examples and explanations."

            # Apply same post-processing calibration to fallback path
            knowledge = max(0, min(10, knowledge))
            communication = max(0, min(10, communication))
            relevance = max(0, min(10, relevance))

            key_terms_fb = [kw.lower() for kw in question.split() if len(kw) > 3]
            match_count_fb = sum(1 for t in key_terms_fb if t in answer.lower())
            if key_terms_fb and match_count_fb < max(2, len(key_terms_fb) // 10):
                knowledge = min(knowledge, 3)
                relevance = min(relevance, 3)
            if "off-topic" in feedback.lower() if isinstance(feedback, str) else False or relevance < 4:
                knowledge = min(knowledge, 4)
                relevance = min(relevance, 4)
            raw_avg_fb = (knowledge + communication + relevance) / 3
            if relevance <= 2 and raw_avg_fb > 4.0:
                knowledge = min(knowledge, 2)
                communication = min(communication, 2)
                relevance = min(relevance, 2)
            elif relevance < 4 and raw_avg_fb > 5.0:
                excess_fb = raw_avg_fb - 5.0
                knowledge = max(0, min(knowledge, knowledge - int(excess_fb * 1.5 + 0.5)))
                relevance = max(0, min(relevance, relevance - 1))

            return {
                "knowledge": knowledge,
                "communication": communication,
                "relevance": relevance,
                "feedback": feedback if isinstance(feedback, str) else "\n\n".join(feedback[:5]),
                "followup": ""
            }
        except:
            pass

    except Exception as e:
        pass

    # Final fallback based on difficulty
    fallback_scores = {"Easy": 3, "Medium": 2, "Hard": 1}
    fallback_score = fallback_scores.get(difficulty, 2)

    return {
        "knowledge": fallback_score,
        "communication": fallback_score,
        "relevance": fallback_score,
        "feedback": "Unable to evaluate properly. Please provide a clear, structured answer. Use the STAR method for behavioral questions and include technical details and examples for technical questions.",
        "followup": ""
    }


def get_ist_time():
    """Get current time in IST timezone"""
    try:
        from datetime import datetime
        import pytz
        ist = pytz.timezone('Asia/Kolkata')
        return datetime.now(ist).strftime('%Y-%m-%d %H:%M:%S')
    except:
        from datetime import datetime
        return datetime.now().strftime('%Y-%m-%d %H:%M:%S')


def log_user_action(username: str, action: str):
    """Log user actions - placeholder for compatibility"""
    pass


def create_interview_database():
    """Create interview_results table if not exists, safely migrate new columns"""
    try:
        conn = _get_live_conn()
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS interview_results (
                id SERIAL PRIMARY KEY,
                username TEXT NOT NULL,
                role TEXT,
                domain TEXT,
                avg_score REAL,
                total_questions INTEGER,
                completed_on TEXT,
                feedback_summary TEXT
            )
        """)
        conn.commit()

        # Safe migration: add new columns only if they don't exist
        cursor.execute("""
            SELECT column_name FROM information_schema.columns
            WHERE table_name = 'interview_results'
        """)
        existing_columns = [row[0] for row in cursor.fetchall()]

        migrations = [
            ("knowledge_avg", "REAL"),
            ("communication_avg", "REAL"),
            ("relevance_avg", "REAL"),
            ("difficulty", "TEXT"),
            ("duration_seconds", "INTEGER"),
            ("interview_mode", "TEXT"),
            ("created_timestamp", "TIMESTAMP DEFAULT CURRENT_TIMESTAMP"),
            ("weighted_score", "REAL"),
            ("raw_avg_score", "REAL"),
            ("follow_up_count", "INTEGER DEFAULT 0"),
            ("depth_score", "REAL"),
            ("behavior_class", "TEXT"),
        ]

        for col_name, col_type in migrations:
            if col_name not in existing_columns:
                try:
                    cursor.execute(f"ALTER TABLE interview_results ADD COLUMN {col_name} {col_type}")
                    conn.commit()
                except Exception:
                    conn.rollback()

        # Also ensure interview_questions table exists
        create_interview_questions_table()
    except Exception as e:
        import streamlit as st
        st.error(f"Database error: {e}")


def create_interview_questions_table():
    """
    Create interview_questions table for storing every question and answer with full context.
    This is the SINGLE SOURCE OF TRUTH for PDF generation.
    """
    try:
        conn = _get_live_conn()
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS interview_questions (
                id SERIAL PRIMARY KEY,
                interview_id TEXT NOT NULL,
                question_text TEXT NOT NULL,
                answer_text TEXT,
                difficulty TEXT,
                is_follow_up INTEGER DEFAULT 0,
                parent_question_id INTEGER,
                timestamp TEXT NOT NULL,
                score_breakdown TEXT,
                question_order INTEGER DEFAULT 0
            )
        """)
        conn.commit()
    except Exception as e:
        import streamlit as st
        st.error(f"Failed to create interview_questions table: {e}")


def save_interview_question(interview_id: str, question_text: str, answer_text: str = None,
                             difficulty: str = "Medium", is_follow_up: bool = False,
                             parent_question_id: int = None, score_breakdown: dict = None,
                             question_order: int = 0) -> int:
    """
    Save a single question (and optionally its answer) to the interview_questions table.
    Returns the row id of the inserted record, or -1 on failure.
    This must be called immediately when a question is answered.
    """
    import json
    try:
        conn = _get_live_conn()
        cursor = conn.cursor()
        score_json = json.dumps(score_breakdown) if score_breakdown else None
        timestamp = get_ist_time()
        cursor.execute("""
            INSERT INTO interview_questions
                (interview_id, question_text, answer_text, difficulty, is_follow_up,
                 parent_question_id, timestamp, score_breakdown, question_order)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
        """, (interview_id, question_text, answer_text,
              difficulty, 1 if is_follow_up else 0,
              parent_question_id, timestamp, score_json, question_order))
        conn.commit()
        row_id = cursor.fetchone()[0]
        return row_id
    except Exception as e:
        import streamlit as st
        st.error(f"Failed to save interview question: {e}")
        return -1


def get_interview_questions_from_db(interview_id: str) -> list:
    """
    Fetch all questions for an interview from DB, ordered by timestamp then question_order.
    Returns list of dicts with keys: id, question_text, answer_text, difficulty,
    is_follow_up, parent_question_id, timestamp, score_breakdown, question_order.
    """
    import json
    try:
        conn = _get_live_conn()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id, question_text, answer_text, difficulty, is_follow_up,
                   parent_question_id, timestamp, score_breakdown, question_order
            FROM interview_questions
            WHERE interview_id = %s
            ORDER BY question_order ASC, timestamp ASC
        """, (interview_id,))
        rows = cursor.fetchall()

        result = []
        for row in rows:
            score = None
            if row[7]:
                try:
                    score = json.loads(row[7])
                except Exception:
                    score = None
            result.append({
                "id": row[0],
                "question_text": row[1],
                "answer_text": row[2] or "",
                "difficulty": row[3],
                "is_follow_up": bool(row[4]),
                "parent_question_id": row[5],
                "timestamp": row[6],
                "score_breakdown": score,
                "question_order": row[8],
            })
        return result
    except Exception as e:
        import streamlit as st
        st.error(f"Failed to fetch interview questions: {e}")
        return []


def save_interview_result(username: str, role: str, domain: str, avg_score: float, total_questions: int, feedback_summary: str,
                          knowledge_avg: float = None, communication_avg: float = None, relevance_avg: float = None,
                          difficulty: str = None, duration_seconds: int = None, interview_mode: str = None,
                          weighted_score: float = None, raw_avg_score: float = None,
                          follow_up_count: int = 0, depth_score: float = None, behavior_class: str = None):
    """Save interview result to database with extended columns"""
    try:
        conn = _get_live_conn()
        cursor = conn.cursor()
        completed_on = get_ist_time()
        cursor.execute("""
            INSERT INTO interview_results (username, role, domain, avg_score, total_questions, completed_on, feedback_summary,
                                          knowledge_avg, communication_avg, relevance_avg, difficulty, duration_seconds, interview_mode, created_timestamp,
                                          weighted_score, raw_avg_score, follow_up_count, depth_score, behavior_class)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, (NOW() AT TIME ZONE 'Asia/Kolkata'), %s, %s, %s, %s, %s)
        """, (username, role, domain, avg_score, total_questions, completed_on, feedback_summary,
              knowledge_avg, communication_avg, relevance_avg, difficulty, duration_seconds, interview_mode,
              weighted_score, raw_avg_score, follow_up_count, depth_score, behavior_class))
        conn.commit()
        # Invalidate dashboard data cache so next visit shows fresh results
        _dirty_key = f"_dashboard_dirty_{username}"
        import streamlit as _st_cache
        if hasattr(_st_cache, 'session_state'):
            _st_cache.session_state[_dirty_key] = True
        return True
    except Exception as e:
        import streamlit as st
        st.error(f"Failed to save interview result: {e}")
        return False


def format_feedback_text(feedback):
    """
    Format feedback text into bullet points for clean display
    """
    import re
    import html
    sentences = re.split(r'(?<=\.)\s+', feedback.strip())
    sentences = [s.strip() for s in sentences if len(s.strip()) > 0]
    formatted = "<b>💡 Improvement Tips:</b><br><ul style='margin-top:5px;'>"
    for s in sentences:
        # Escape HTML special characters to display tags like <header>, <section>, etc.
        safe_sentence = html.escape(s)
        formatted += f"<li>{safe_sentence}</li>"
    formatted += "</ul>"
    return formatted


def generate_interview_pdf_report(username, role, domain, completed_on, questions, answers, scores, feedbacks, overall_avg, badge, difficulty="Medium", interview_id=None):
    """
    Generate PDF report for interview using xhtml2pdf.

    ARCHITECTURE FIX: When interview_id is provided, fetches ALL Q&A data exclusively
    from the interview_questions DB table (the single source of truth).
    Never regenerates follow-up questions. Preserves original order via timestamp/question_order.
    Falls back to passed-in arrays only when interview_id is unavailable (legacy).
    """
    try:
        from xhtml2pdf import pisa
        from io import BytesIO

        # ── SINGLE SOURCE OF TRUTH: fetch from DB when interview_id is available ──
        db_rows = []
        if interview_id:
            db_rows = get_interview_questions_from_db(interview_id)

        # Build XHTML content
        xhtml = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1 {{ color: #00c3ff; text-align: center; }}
                h2 {{ color: #0099cc; margin-top: 20px; }}
                .header {{ background: #f0f0f0; padding: 15px; border-radius: 8px; margin-bottom: 20px; }}
                .question-block {{ margin: 20px 0; padding: 15px; border: 1px solid #ddd; border-radius: 8px; page-break-inside: avoid; }}
                .followup-block {{ margin: 10px 0 20px 30px; padding: 15px; border: 1px solid #ffc107; border-radius: 8px; background: #fffdf0; page-break-inside: avoid; }}
                .score {{ font-weight: bold; color: #00c3ff; }}
                .feedback {{ color: #666; margin-top: 10px; padding: 10px; background: #f9f9f9; border-left: 3px solid #00c3ff; }}
                .feedback ul {{ margin: 5px 0 0 0; padding-left: 20px; }}
                .feedback li {{ margin: 8px 0; line-height: 1.5; }}
                .summary {{ background: #fffacd; padding: 15px; border-radius: 8px; margin: 20px 0; }}
                .answer-text {{ white-space: pre-wrap; word-wrap: break-word; margin: 10px 0; }}
                .followup-label {{ color: #b8860b; font-weight: bold; font-size: 13px; margin-bottom: 6px; }}
            </style>
        </head>
        <body>
            <h1>Interview Practice Report</h1>
            <div class="header">
                <p><strong>Candidate:</strong> {username}</p>
                <p><strong>Role:</strong> {role}</p>
                <p><strong>Domain:</strong> {domain}</p>
                <p><strong>Date:</strong> {completed_on}</p>
            </div>
            <div class="summary">
                <h2>Overall Performance</h2>
                <p class="score">Average Score: {overall_avg:.2f}/10</p>
                <p><strong>Badge Earned:</strong> {badge}</p>
            </div>
            <h2>Detailed Q&amp;A Review</h2>
        """

        if db_rows:
            # ── DB-backed path: use ONLY stored data, never regenerate ──
            # Separate main questions and follow-ups
            main_questions = [r for r in db_rows if not r["is_follow_up"]]
            followup_map = {}  # parent_question_id -> list of follow-up rows
            for r in db_rows:
                if r["is_follow_up"] and r["parent_question_id"] is not None:
                    followup_map.setdefault(r["parent_question_id"], []).append(r)

            import re
            for idx, row in enumerate(main_questions, 1):
                q_escaped = row["question_text"].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                a_escaped = row["answer_text"].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                score_dict = row["score_breakdown"] or {}
                if isinstance(score_dict, dict) and score_dict:
                    avg_q_score = (score_dict.get('knowledge', 5) + score_dict.get('communication', 5) + score_dict.get('relevance', 5)) / 3
                else:
                    avg_q_score = 5.0
                    score_dict = {'knowledge': 5, 'communication': 5, 'relevance': 5}

                feedback_raw = score_dict.get("feedback", "") if isinstance(score_dict, dict) else ""
                if isinstance(feedback_raw, list):
                    feedback_raw = "\n\n".join(feedback_raw)
                sentences = re.split(r'(?<=\.)\s+', str(feedback_raw).strip())
                sentences = [s.strip() for s in sentences if len(s.strip()) > 0]
                bullet_feedback = "<b>💡 Improvement Tips:</b><ul>"
                for sent in sentences:
                    sent_esc = sent.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    bullet_feedback += f"<li>{sent_esc}</li>"
                bullet_feedback += "</ul>"

                xhtml += f"""
            <div class="question-block">
                <h3>Question {idx}</h3>
                <p><strong>Q:</strong> {q_escaped}</p>
                <div class="answer-text"><strong>Your Answer:</strong><br/>{a_escaped}</div>
                <p class="score">Knowledge: {score_dict.get('knowledge', 0)}/10 | Communication: {score_dict.get('communication', 0)}/10 | Relevance: {score_dict.get('relevance', 0)}/10</p>
                <p class="score">Question Score: {avg_q_score:.2f}/10</p>
                <div class="feedback">{bullet_feedback}</div>
            </div>
                """

                # Nest follow-up questions under this main question
                for fu in followup_map.get(row["id"], []):
                    fu_q_esc = fu["question_text"].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    fu_a_esc = fu["answer_text"].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    fu_score = fu["score_breakdown"] or {}
                    if isinstance(fu_score, dict) and fu_score:
                        fu_avg = (fu_score.get('knowledge', 5) + fu_score.get('communication', 5) + fu_score.get('relevance', 5)) / 3
                    else:
                        fu_avg = 5.0
                        fu_score = {'knowledge': 5, 'communication': 5, 'relevance': 5}

                    fu_feedback_raw = fu_score.get("feedback", "") if isinstance(fu_score, dict) else ""
                    if isinstance(fu_feedback_raw, list):
                        fu_feedback_raw = "\n\n".join(fu_feedback_raw)
                    fu_sentences = re.split(r'(?<=\.)\s+', str(fu_feedback_raw).strip())
                    fu_sentences = [s.strip() for s in fu_sentences if len(s.strip()) > 0]
                    fu_bullets = "<b>💡 Improvement Tips:</b><ul>"
                    for s in fu_sentences:
                        s_esc = s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        fu_bullets += f"<li>{s_esc}</li>"
                    fu_bullets += "</ul>"

                    xhtml += f"""
            <div class="followup-block">
                <div class="followup-label">↳ Follow-Up Question (Hard Mode)</div>
                <p><strong>Q:</strong> {fu_q_esc}</p>
                <div class="answer-text"><strong>Your Answer:</strong><br/>{fu_a_esc}</div>
                <p class="score">Knowledge: {fu_score.get('knowledge', 0)}/10 | Communication: {fu_score.get('communication', 0)}/10 | Relevance: {fu_score.get('relevance', 0)}/10</p>
                <p class="score">Follow-Up Score: {fu_avg:.2f}/10</p>
                <div class="feedback">{fu_bullets}</div>
            </div>
                    """
        else:
            # ── Legacy fallback: use passed-in arrays (no interview_id) ──
            import re
            for i, (q, a, score_dict, f) in enumerate(zip(questions, answers, scores, feedbacks), 1):
                if isinstance(score_dict, dict):
                    avg_q_score = (score_dict.get('knowledge', 5) + score_dict.get('communication', 5) + score_dict.get('relevance', 5)) / 3
                else:
                    avg_q_score = 5.0
                    score_dict = {'knowledge': 5, 'communication': 5, 'relevance': 5}

                q_escaped = q.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                a_escaped = a.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                if isinstance(f, list):
                    f_text = "\n\n".join(f)
                else:
                    f_text = str(f)

                sentences = re.split(r'(?<=\.)\s+', f_text.strip())
                sentences = [sent.strip() for sent in sentences if len(sent.strip()) > 0]
                bullet_feedback = "<b>💡 Improvement Tips:</b><ul>"
                for sent in sentences:
                    sent_escaped = sent.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    bullet_feedback += f"<li>{sent_escaped}</li>"
                bullet_feedback += "</ul>"

                followup_text = ""
                if difficulty == "Hard" and isinstance(score_dict, dict) and score_dict.get('followup'):
                    followup_escaped = score_dict['followup'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    followup_text = f"""<div style="margin-top: 10px; padding: 10px; background: #fff3cd; border-radius: 5px;">
                        <strong>Follow-up Question (for Hard interviews):</strong><br/>
                        {followup_escaped}
                    </div>"""

                xhtml += f"""
            <div class="question-block">
                <h3>Question {i}</h3>
                <p><strong>Q:</strong> {q_escaped}</p>
                <div class="answer-text"><strong>Your Answer:</strong><br/>{a_escaped}</div>
                <p class="score">Knowledge: {score_dict.get('knowledge', 0)}/10 | Communication: {score_dict.get('communication', 0)}/10 | Relevance: {score_dict.get('relevance', 0)}/10</p>
                <p class="score">Question Score: {avg_q_score:.2f}/10</p>
                <div class="feedback">{bullet_feedback}</div>
                {followup_text}
            </div>
                """

        xhtml += """
        </body>
        </html>
        """

        # Convert to PDF
        pdf_out = BytesIO()
        pisa_status = pisa.CreatePDF(xhtml, dest=pdf_out)
        pdf_out.seek(0)

        if pisa_status.err:
            return None

        return pdf_out.getvalue()

    except Exception as e:
        import streamlit as st
        st.error(f"PDF generation failed: {e}")
        return None



import streamlit as st
import plotly.graph_objects as go
from courses import COURSES_BY_CATEGORY, RESUME_VIDEOS, INTERVIEW_VIDEOS, get_courses_for_role
from llm_manager import call_llm
import time
import threading
import json


# =============================================================================
# BROWSER-SIDE JS TIMER  (final concurrency fix)
# =============================================================================
# ROOT CAUSE of slow/jammed timers under concurrent load:
#   st.fragment(run_every=1) wakes up a SERVER thread every second PER USER.
#   With 2 users → 2 threads fighting every second → GIL contention → timer
#   ticks arrive late → display freezes / lags.
#
# SOLUTION: Move the countdown 100% into the BROWSER with JavaScript.
#   - JS setInterval runs in the user's own browser tab — zero server load.
#   - 10 users running interviews = 0 extra server threads for timer ticks.
#   - The server is only contacted ONCE per question: when time expires,
#     the JS clicks a hidden Streamlit button to trigger auto-submit.
#   - question_timer_start (time.time()) is still stored in session_state
#     so the server always knows the true elapsed time for scoring/duration.
# =============================================================================

def _render_js_timer(remaining_seconds: float, total_seconds: int, submitted: bool, q_idx: int):
    """
    Render a pure-JS countdown timer inside an st.components.v1.html block.

    The countdown runs entirely in the browser — no server round-trip per tick.
    When the countdown hits zero the JS clicks the hidden Streamlit button
    whose key is  '__timer_expired_btn_{q_idx}'  which triggers a normal
    Streamlit interaction → sets session_state flags → auto-submits answer.

    Parameters
    ----------
    remaining_seconds : float   Seconds left as calculated by the server on this render.
    total_seconds     : int     Total seconds for this question (for the progress bar).
    submitted         : bool    If True, show "Answer Submitted" banner instead.
    q_idx             : int     Current question index (used to key the hidden button).
    """
    if submitted:
        st.markdown("""
        <div style="background:linear-gradient(135deg,rgba(52,211,153,0.10),rgba(52,211,153,0.05));
                    border:1px solid rgba(52,211,153,0.30);border-radius:12px;
                    padding:14px;text-align:center;">
          <div style="font-size:1.2rem;font-weight:700;color:#34d399;">
            ✅ Answer Submitted
          </div>
        </div>
        """, unsafe_allow_html=True)
        return

    # Clamp so we never pass a negative value into JS
    remaining_seconds = max(0.0, remaining_seconds)

    st.components.v1.html(f"""
    <style>
      @keyframes t4pulse {{
        0%,100% {{ box-shadow: 0 0 0 0 rgba(244,67,54,0.0); }}
        50%      {{ box-shadow: 0 0 0 6px rgba(244,67,54,0.18); }}
      }}
      #timer-wrap {{
        font-family: -apple-system, BlinkMacSystemFont, 'SF Pro Display', sans-serif;
        border-radius: 12px;
        padding: 14px;
        text-align: center;
        transition: background 0.5s, border-color 0.5s;
      }}
      #timer-wrap.normal {{
        background: linear-gradient(135deg,rgba(251,191,36,0.08),rgba(251,191,36,0.04));
        border: 1px solid rgba(251,191,36,0.25);
      }}
      #timer-wrap.urgent {{
        background: linear-gradient(135deg,rgba(244,67,54,0.12),rgba(244,67,54,0.06));
        border: 1px solid rgba(244,67,54,0.45);
        animation: t4pulse 1s ease-in-out infinite;
      }}
      #timer-text {{
        font-size: 1.4rem;
        font-weight: 700;
        letter-spacing: -0.01em;
        transition: color 0.5s;
      }}
      #timer-track {{
        width: 100%; height: 4px;
        background: rgba(255,255,255,0.08);
        border-radius: 99px;
        margin-top: 10px;
        overflow: hidden;
      }}
      #timer-bar {{
        height: 100%;
        border-radius: 99px;
        transition: width 0.95s linear, background 0.5s;
      }}
    </style>

    <div id="timer-wrap" class="normal">
      <div id="timer-text">⏰ Time Remaining: <span id="timer-display">--:--</span></div>
      <div id="timer-track"><div id="timer-bar"></div></div>
    </div>

    <script>
    (function() {{
      // Server told us exactly how many seconds remain at render time.
      // Use performance.now() for drift-free sub-millisecond accuracy in browser.
      var remaining  = {remaining_seconds:.3f};
      var total      = {total_seconds};
      var startedAt  = performance.now();
      var expired    = false;

      var wrap    = document.getElementById('timer-wrap');
      var display = document.getElementById('timer-display');
      var bar     = document.getElementById('timer-bar');
      var text    = document.getElementById('timer-text');

      function fmt(secs) {{
        var m = Math.floor(secs / 60);
        var s = Math.floor(secs % 60);
        return (m < 10 ? '0' : '') + m + ':' + (s < 10 ? '0' : '') + s;
      }}

      function tick() {{
        var elapsed = (performance.now() - startedAt) / 1000;
        var left    = Math.max(0, remaining - elapsed);
        var pct     = total > 0 ? (1 - left / total) * 100 : 100;
        var urgent  = left <= 30;

        display.textContent = fmt(left);
        bar.style.width     = pct.toFixed(1) + '%';
        bar.style.background= urgent ? '#ef4444' : '#f59e0b';
        text.style.color    = urgent ? '#f87171' : '#fbbf24';
        wrap.className      = urgent ? 'urgent' : 'normal';

        if (left <= 0 && !expired) {{
          expired = true;
          // Display-only: freeze at 00:00.
          // Auto-submit is handled entirely server-side via a background thread.
          display.textContent = '00:00';
          bar.style.width = '100%';
          bar.style.background = '#ef4444';
          text.style.color = '#f87171';
          wrap.className = 'urgent';
        }}
      }}

      // Run immediately then every 250 ms for smooth display without hammering.
      tick();
      setInterval(tick, 250);
    }})();
    </script>
    """, height=80)  # compact fixed height — no scrollbar


# =============================================================================
# SUPABASE POSTGRESQL — SINGLE CACHED CONNECTION  (anti-flicker fix)
# =============================================================================
# @st.cache_resource creates the psycopg2 connection ONCE per server process
# and reuses it on every Streamlit rerun.  Previously, sqlite3.connect() was
# called inside every function → new I/O on every widget interaction → page
# scrolled back to top and flickered.  A single cached connection eliminates
# that entirely.
#
# We also add a lightweight ping so that if the connection goes idle and the
# server closes it, we transparently reconnect without the user seeing an error.

@st.cache_resource
def get_progress_db():
    """Return a single cached psycopg2 connection to Supabase PostgreSQL."""
    import psycopg2
    return psycopg2.connect(
        host=st.secrets["SUPABASE_HOST"],
        database=st.secrets["SUPABASE_DB"],
        user=st.secrets["SUPABASE_USER"],
        password=st.secrets["SUPABASE_PASSWORD"],
        port=st.secrets["SUPABASE_PORT"]
    )


def _get_live_conn():
    """
    Return the cached connection.  If the connection has gone idle/closed,
    clear the cache so get_progress_db() reconnects on next call.
    This prevents 'connection already closed' errors without any visible flicker.
    """
    import psycopg2
    conn = get_progress_db()
    try:
        # Lightweight ping — no round-trip if connection is healthy
        conn.cursor().execute("SELECT 1")
    except Exception:
        # Connection is dead — clear cache and reconnect
        get_progress_db.clear()
        conn = get_progress_db()
    return conn


# =============================================================================
# ONE-TIME DB INIT GUARD  (anti-flicker fix)
# =============================================================================
# create_interview_database() was called on EVERY Streamlit rerun (every widget
# interaction).  Each call ran DDL queries against Supabase — completely
# unnecessary after the first run.  We guard it with a session_state flag so
# the DDL only runs once per browser session.

def _ensure_db_initialized():
    """Run create_interview_database() at most once per browser session."""
    if not st.session_state.get("_db_initialized", False):
        create_interview_database()
        st.session_state["_db_initialized"] = True


import json
import time
import re
import streamlit as st

# =============================================================================
# ARCHITECTURAL FIX 1: DOMAIN AUTHORITY LAYER
# =============================================================================
# Problem: Resume context dominates LLM prompts, causing Full Stack resumes to
# produce Full Stack questions even when "Data Analyst" is selected.
# Solution: Strip and suppress resume content that contradicts the selected domain,
# then inject domain-specific mandatory keywords into every question generation prompt.

DOMAIN_AUTHORITY_CONFIG = {
    "Data Science & Analytics": {
        "aliases": ["data analyst", "data science", "analytics", "business intelligence", "bi", "ml", "machine learning"],
        "mandatory_topics": [
            # Core Python data stack
            "pandas", "NumPy", "SciPy", "Matplotlib", "Seaborn", "Plotly",
            # SQL & querying
            "SQL", "window functions", "CTEs", "query optimization", "joins", "GROUP BY aggregations",
            # Statistical foundations
            "descriptive statistics", "inferential statistics", "hypothesis testing", "p-values",
            "confidence intervals", "A/B testing", "statistical significance", "effect size",
            "probability distributions", "Bayesian reasoning",
            # EDA & data quality
            "exploratory data analysis", "data cleaning", "missing value imputation",
            "outlier detection", "data profiling", "feature distributions",
            # Classical ML
            "linear regression", "logistic regression", "decision trees", "random forests",
            "gradient boosting", "XGBoost", "LightGBM", "k-means clustering", "PCA",
            "bias-variance tradeoff", "cross-validation", "regularization (L1/L2)",
            # Model evaluation
            "precision", "recall", "F1 score", "ROC-AUC", "confusion matrix",
            "RMSE", "MAE", "R-squared", "lift curves",
            # BI & visualization
            "Tableau", "Power BI", "Looker", "dashboard design", "data storytelling",
            "KPI definition", "metric frameworks", "cohort analysis",
            # Data pipelines
            "ETL pipelines", "data warehousing", "OLAP vs OLTP", "star schema",
            "dbt", "Apache Airflow", "data lineage",
            # Advanced / modern
            "time series forecasting", "ARIMA", "Prophet", "feature engineering",
            "dimensionality reduction", "class imbalance handling", "SMOTE",
        ],
        "forbidden_resume_keywords": [
            "react", "angular", "vue", "next.js", "nuxt", "svelte",
            "node.js", "express", "fastify", "django", "flask", "spring boot",
            "frontend", "css", "html", "tailwind", "bootstrap", "figma",
            "mobile app", "swift", "kotlin", "flutter", "react native",
            "graphql", "REST API design", "OAuth", "JWT authentication",
        ],
        "context_override": (
            "This is a Data Science & Analytics interview. Focus EXCLUSIVELY on data analysis, "
            "statistical reasoning, SQL, Python data libraries (pandas/numpy/scikit-learn), "
            "model evaluation, data visualization, A/B testing, EDA, and business intelligence tools. "
            "Probe for depth on statistical rigour, evaluation metric selection, and data storytelling — "
            "not just tool familiarity."
        ),
    },

    "Full Stack Development": {
        "aliases": ["full stack", "fullstack", "web developer", "mern", "mean", "full-stack engineer"],
        "mandatory_topics": [
            # Frontend
            "React", "Vue.js", "Angular", "Next.js", "TypeScript", "JavaScript (ES6+)",
            "state management", "Redux", "Zustand", "Pinia", "component lifecycle",
            "virtual DOM", "server-side rendering", "static site generation",
            "CSS-in-JS", "Tailwind CSS", "responsive design", "accessibility (WCAG)",
            "browser performance", "lazy loading", "code splitting", "web vitals",
            # Backend
            "Node.js", "Express", "NestJS", "Django", "FastAPI", "Spring Boot",
            "REST API design", "GraphQL", "API versioning", "rate limiting",
            "middleware patterns", "input validation", "error handling",
            # Databases
            "PostgreSQL", "MySQL", "MongoDB", "Redis", "database indexing",
            "query optimization", "ORM (Prisma/Sequelize/SQLAlchemy)", "N+1 problem",
            "transactions and ACID", "database migrations",
            # Auth & security
            "JWT", "OAuth 2.0", "session management", "CORS", "CSRF", "XSS", "SQL injection",
            "HTTPS", "secrets management", "RBAC",
            # DevOps basics
            "Docker", "CI/CD", "environment variables", "12-factor app",
            "Nginx", "reverse proxy", "load balancing basics",
            # Testing
            "unit testing", "integration testing", "Jest", "React Testing Library",
            "end-to-end testing", "Cypress", "Playwright", "test coverage",
            # Architecture
            "monolith vs microservices", "BFF pattern", "caching strategies",
            "WebSockets", "real-time updates", "message queues basics",
        ],
        "forbidden_resume_keywords": [
            "tensorflow", "pytorch", "sklearn", "scikit-learn", "pandas",
            "regression model", "clustering", "NLP pipeline", "deep learning model",
            "Kubernetes operator", "Terraform modules", "Ansible playbooks",
            "pen testing", "SIEM", "SOC analyst",
        ],
        "context_override": (
            "This is a Full Stack Development interview. Cover the entire request lifecycle: "
            "browser → frontend framework → API layer → database → response. "
            "Focus on React/Next.js or Vue, Node.js or Django backends, PostgreSQL/MongoDB, "
            "authentication patterns, caching, testing strategies, and deployment pipelines. "
            "Probe for real decisions: state management choices, API contract design, N+1 fixes, "
            "and security hardening — not just stack enumeration."
        ),
    },

    "Backend Development": {
        "aliases": ["backend", "server-side", "api developer", "java developer", "python developer",
                    "golang developer", "backend engineer", "software engineer backend"],
        "mandatory_topics": [
            # API design
            "REST API design principles", "GraphQL", "gRPC", "API versioning",
            "idempotency", "pagination", "rate limiting", "API gateway",
            "OpenAPI / Swagger documentation",
            # Databases
            "PostgreSQL", "MySQL", "database indexing", "query execution plans",
            "EXPLAIN ANALYZE", "N+1 problem", "connection pooling",
            "transactions", "ACID properties", "isolation levels", "deadlocks",
            "database migrations", "schema design", "normalization",
            "Redis", "caching patterns (cache-aside, write-through)", "cache invalidation",
            # System design
            "horizontal vs vertical scaling", "load balancing", "reverse proxy",
            "microservices", "service mesh", "event-driven architecture",
            "message queues (RabbitMQ, Kafka, SQS)", "pub/sub patterns",
            "circuit breaker", "bulkhead pattern", "retry with backoff",
            "CAP theorem", "eventual consistency", "distributed transactions",
            # Auth & security
            "JWT", "OAuth 2.0", "OpenID Connect", "RBAC", "ABAC",
            "password hashing (bcrypt/argon2)", "secrets management",
            "input validation", "SQL injection", "OWASP Top 10 for APIs",
            # Performance
            "async programming", "concurrency models", "thread pools",
            "profiling", "bottleneck identification", "background jobs",
            "batch processing", "streaming",
            # Testing
            "unit testing", "integration testing", "contract testing",
            "mocking dependencies", "test isolation", "CI test pipelines",
            # Languages & runtimes
            "Python (asyncio/FastAPI/Django)", "Node.js (event loop)",
            "Java (Spring Boot, JVM tuning)", "Go (goroutines, channels)",
        ],
        "forbidden_resume_keywords": [
            "react", "css", "html", "angular", "vue", "next.js", "tailwind",
            "figma", "photoshop", "sketch", "frontend", "ui design",
            "pandas", "sklearn", "tensorflow", "pytorch",
            "mobile app", "swift", "kotlin", "flutter",
        ],
        "context_override": (
            "This is a Backend Development interview. Focus on server-side engineering: "
            "REST/GraphQL/gRPC API design, relational and NoSQL databases (including indexing, "
            "transactions, and query optimisation), caching strategies, message queues, "
            "distributed system patterns, authentication/authorisation, and backend testing. "
            "Push candidates to explain WHY they made design decisions, not just WHAT they used."
        ),
    },

    "Frontend Development": {
        "aliases": ["frontend", "ui developer", "react developer", "angular developer",
                    "vue developer", "frontend engineer", "web ui engineer"],
        "mandatory_topics": [
            # Core JavaScript
            "JavaScript (ES6+)", "TypeScript", "closures", "event loop", "async/await",
            "Promises", "prototypal inheritance", "hoisting", "debounce/throttle",
            "Web APIs (Fetch, localStorage, IntersectionObserver)",
            # Frameworks
            "React (hooks, context, reconciliation)", "Vue.js (Options API vs Composition API)",
            "Angular (dependency injection, change detection)",
            "Next.js (SSR vs SSG vs ISR)", "Nuxt.js",
            # State management
            "Redux (flux architecture)", "Redux Toolkit", "Zustand", "Recoil",
            "React Query / TanStack Query", "SWR", "Pinia",
            # CSS & styling
            "CSS specificity", "flexbox", "CSS Grid", "responsive design",
            "CSS-in-JS (styled-components, Emotion)", "Tailwind CSS",
            "CSS custom properties", "animations and transitions",
            # Performance
            "Core Web Vitals (LCP, CLS, FID/INP)", "code splitting",
            "lazy loading", "tree shaking", "bundle analysis (Webpack, Vite)",
            "image optimisation", "caching strategies (HTTP cache, service workers)",
            "virtual DOM and reconciliation", "memoization (useMemo, useCallback, React.memo)",
            # Testing
            "Jest", "React Testing Library", "Vitest",
            "end-to-end testing (Cypress, Playwright)",
            "snapshot testing", "accessibility testing",
            # Accessibility
            "WCAG 2.1 guidelines", "ARIA attributes", "keyboard navigation",
            "screen reader compatibility", "semantic HTML",
            # Architecture
            "micro-frontends", "module federation", "design systems",
            "component composition patterns", "render props vs HOC vs hooks",
            "Storybook", "monorepo (Nx, Turborepo)",
            # Security
            "XSS prevention", "CSP headers", "CSRF in SPAs", "sanitisation",
        ],
        "forbidden_resume_keywords": [
            "kubernetes", "docker-compose", "terraform", "ansible",
            "CI/CD pipeline design", "microservices orchestration", "kafka",
            "grpc", "tensorflow", "pytorch", "pandas", "sklearn",
            "SQL joins", "database schema design", "backend API design",
            "pen testing", "SIEM", "firewall rules",
        ],
        "context_override": (
            "This is a Frontend Development interview. Focus on deep JavaScript/TypeScript knowledge, "
            "React or Vue framework internals (reconciliation, hooks, reactivity), state management "
            "tradeoffs, Core Web Vitals and performance optimisation, CSS layout and architecture, "
            "accessibility standards, testing strategies, and client-side security. "
            "Go beyond tool lists — probe for understanding of browser behaviour, "
            "render performance, and component design decisions."
        ),
    },

    "Machine Learning & AI": {
        "aliases": ["machine learning", "ml engineer", "ai engineer", "deep learning",
                    "nlp engineer", "computer vision engineer", "mlops engineer", "ai researcher"],
        "mandatory_topics": [
            # Foundations
            "supervised vs unsupervised vs reinforcement learning",
            "bias-variance tradeoff", "overfitting", "underfitting", "regularisation (L1/L2/dropout)",
            "cross-validation (k-fold, stratified)", "train/val/test split strategy",
            "data leakage", "target encoding pitfalls",
            # Feature engineering
            "feature selection (mutual information, SHAP)", "feature scaling (standardisation, normalisation)",
            "handling missing values", "categorical encoding (one-hot, ordinal, target)",
            "dimensionality reduction (PCA, t-SNE, UMAP)", "class imbalance (SMOTE, cost-sensitive learning)",
            # Classical ML
            "linear/logistic regression", "decision trees and ensemble methods",
            "random forests", "gradient boosting (XGBoost, LightGBM, CatBoost)",
            "SVMs", "k-means clustering", "DBSCAN",
            # Deep learning
            "neural network architecture", "backpropagation", "gradient descent variants (Adam, SGD)",
            "batch normalisation", "dropout", "learning rate scheduling",
            "CNNs (convolution, pooling)", "RNNs / LSTMs", "attention mechanism",
            "Transformer architecture", "transfer learning", "fine-tuning",
            "BERT, GPT, and LLM fundamentals",
            # Evaluation metrics
            "precision, recall, F1", "ROC-AUC, PR-AUC", "NDCG", "MAP",
            "RMSE, MAE, MAPE", "calibration", "offline vs online evaluation",
            # MLOps
            "ML pipelines (Kubeflow, MLflow, SageMaker Pipelines)",
            "experiment tracking (MLflow, Weights & Biases)",
            "model versioning", "model registry", "feature stores",
            "data drift detection", "model monitoring in production",
            "A/B testing for ML models", "shadow deployment", "canary deployment",
            # Deployment
            "model serialisation (ONNX, pickle, TorchScript)",
            "serving (TensorFlow Serving, Triton, FastAPI)", "batch vs real-time inference",
            "latency-throughput tradeoffs", "model quantisation", "pruning", "distillation",
            # LLM-specific
            "prompt engineering", "RAG (retrieval-augmented generation)",
            "vector databases (Pinecone, Weaviate, FAISS)", "embedding models",
            "LLM fine-tuning (LoRA, PEFT)", "hallucination mitigation",
            "LLM evaluation (BERTScore, RAGAS)", "token context windows",
        ],
        "forbidden_resume_keywords": [
            "react", "angular", "vue", "next.js", "tailwind",
            "node.js", "express", "spring boot", "django REST framework",
            "frontend", "css", "html", "figma", "mobile app",
            "swift", "kotlin", "flutter", "react native",
        ],
        "context_override": (
            "This is a Machine Learning & AI interview. Cover the full ML lifecycle: "
            "problem framing, data preparation, feature engineering, model selection and evaluation, "
            "production deployment, and monitoring. For senior roles, probe on MLOps maturity, "
            "LLM engineering, and system design for ML (feature stores, serving infrastructure). "
            "Demand technical depth — evaluation metric justification, tradeoff reasoning, "
            "and real failure scenarios, not just algorithm definitions."
        ),
    },

    "DevOps & Cloud": {
        "aliases": ["devops", "cloud engineer", "platform engineer", "sre", "site reliability",
                    "infrastructure engineer", "cloud architect", "devsecops"],
        "mandatory_topics": [
            # CI/CD
            "CI/CD pipeline design (GitHub Actions, GitLab CI, Jenkins, CircleCI)",
            "pipeline stages (build, test, security scan, deploy)",
            "artifact management", "deployment strategies (blue-green, canary, rolling)",
            "feature flags", "rollback strategies", "trunk-based development",
            # Containers & orchestration
            "Docker (images, layers, multi-stage builds, registry)",
            "Kubernetes (pods, deployments, services, ingress, namespaces)",
            "Kubernetes resource requests/limits", "HPA and VPA",
            "Helm charts", "operators", "StatefulSets", "PersistentVolumes",
            "service mesh (Istio, Linkerd)", "container security scanning",
            # Infrastructure as Code
            "Terraform (state management, modules, workspaces)",
            "Ansible", "Pulumi", "CloudFormation / CDK",
            "GitOps (ArgoCD, Flux)", "drift detection",
            # Cloud platforms
            "AWS (EC2, ECS/EKS, Lambda, S3, RDS, CloudFront, IAM, VPC, Route53)",
            "Azure (AKS, App Service, Azure Functions, Blob Storage, AAD)",
            "GCP (GKE, Cloud Run, Cloud Functions, BigQuery, IAM, VPC)",
            "multi-cloud and hybrid cloud patterns",
            "cloud cost optimisation (reserved instances, spot/preemptible, rightsizing)",
            # Observability
            "metrics (Prometheus, CloudWatch, Datadog)", "logging (ELK stack, Loki, CloudWatch Logs)",
            "tracing (Jaeger, Zipkin, AWS X-Ray, OpenTelemetry)",
            "alerting (PagerDuty, OpsGenie)", "SLOs, SLAs, SLIs, error budgets",
            "on-call practices", "incident management (runbooks, postmortems)",
            # Networking
            "DNS", "load balancers (ALB, NLB, HAProxy)", "CDN",
            "VPC design (subnets, NACLs, security groups)", "VPN and Direct Connect",
            "service discovery", "network policies in Kubernetes",
            # Security & compliance
            "IAM least privilege", "secrets management (Vault, AWS Secrets Manager)",
            "SAST/DAST in pipelines", "image vulnerability scanning (Trivy, Snyk)",
            "CIS benchmarks", "SOC2/PCI compliance automation",
            # Reliability
            "chaos engineering", "fault injection", "disaster recovery",
            "RTO and RPO", "backup strategies", "multi-region failover",
        ],
        "forbidden_resume_keywords": [
            "react", "angular", "vue", "next.js", "tailwind", "figma",
            "pandas", "sklearn", "tensorflow", "pytorch",
            "mobile app", "swift", "kotlin", "flutter",
            "SEO optimisation", "UI component library",
        ],
        "context_override": (
            "This is a DevOps & Cloud interview. Cover CI/CD pipeline design, container orchestration "
            "(Kubernetes internals and production operations), Infrastructure as Code (Terraform/GitOps), "
            "cloud platform services (AWS/GCP/Azure), observability (metrics/logs/traces/SLOs), "
            "networking, and cloud security. Probe for real operational experience: "
            "incident response war stories, cost optimisation decisions, and reliability engineering "
            "— not just tool enumerations."
        ),
    },

    "Cybersecurity": {
        "aliases": ["cybersecurity", "security engineer", "pen tester", "information security",
                    "appsec", "application security", "cloud security", "devsecops engineer",
                    "soc analyst", "threat intelligence"],
        "mandatory_topics": [
            # Fundamentals
            "CIA triad (confidentiality, integrity, availability)",
            "defence in depth", "principle of least privilege", "zero trust architecture",
            "threat modelling (STRIDE, PASTA, attack trees)",
            "risk assessment and risk scoring (CVSS)", "security controls taxonomy",
            # Application security
            "OWASP Top 10 (SQLi, XSS, SSRF, IDOR, broken auth, etc.)",
            "input validation and output encoding", "parameterised queries",
            "authentication flows (OAuth 2.0, OIDC, SAML)", "JWT security pitfalls",
            "session management", "CSRF", "clickjacking", "security headers (CSP, HSTS)",
            "API security (rate limiting, auth, mass assignment)",
            "SAST and DAST", "software composition analysis (SCA)", "secret scanning",
            # Network security
            "TCP/IP fundamentals", "TLS/SSL (handshake, certificate chains, HSTS)",
            "firewalls and WAFs", "IDS/IPS", "VPN and Zero Trust Network Access",
            "DNS security (DNSSEC, DNS poisoning)", "DDoS mitigation",
            "network segmentation and micro-segmentation",
            # Penetration testing
            "OWASP Testing Guide", "recon and OSINT", "exploitation frameworks (Metasploit)",
            "web app pen testing (Burp Suite)", "privilege escalation techniques",
            "post-exploitation and lateral movement", "reporting and severity classification",
            # Cloud security
            "IAM misconfiguration", "S3 bucket exposure", "AWS security best practices",
            "cloud security posture management (CSPM)", "container security",
            "secrets management (Vault, AWS Secrets Manager)", "CWPP",
            # Incident response
            "incident response lifecycle (preparation, detection, containment, eradication, recovery)",
            "digital forensics basics", "log analysis and SIEM (Splunk, Microsoft Sentinel, Chronicle)",
            "threat hunting", "IoCs and IoAs", "MITRE ATT&CK framework",
            # Cryptography
            "symmetric vs asymmetric encryption", "AES, RSA, ECC",
            "hashing (SHA-256, bcrypt, argon2)", "PKI and certificate management",
            "key management", "TLS configuration best practices",
            # Compliance & governance
            "GDPR", "SOC 2", "ISO 27001", "PCI-DSS", "HIPAA",
            "security policies and standards", "vulnerability management lifecycle",
        ],
        "forbidden_resume_keywords": [
            "react", "angular", "vue", "pandas", "sklearn",
            "mobile app", "ui design", "figma", "photoshop",
            "CSS animations", "frontend state management",
        ],
        "context_override": (
            "This is a Cybersecurity interview. Cover application security (OWASP Top 10, "
            "auth flows, API security), network security, penetration testing methodology, "
            "cloud security, incident response, and cryptography fundamentals. "
            "Probe for offensive AND defensive mindset: threat modelling, exploit chaining, "
            "detection engineering, and security architecture decisions — not just tool familiarity."
        ),
    },

    "UI/UX Design": {
        "aliases": ["ui designer", "ux designer", "product designer", "interaction designer",
                    "ux researcher", "design lead", "experience designer"],
        "mandatory_topics": [
            # Research methods
            "user interviews", "contextual inquiry", "diary studies",
            "surveys (quantitative vs qualitative)", "usability testing (moderated/unmoderated)",
            "card sorting", "tree testing", "A/B testing for UX",
            "affinity mapping", "persona development", "jobs-to-be-done framework",
            # Information architecture & flows
            "information architecture", "site maps", "user flows", "task flows",
            "mental models", "navigation patterns", "progressive disclosure",
            # Wireframing & prototyping
            "low-fidelity wireframes", "high-fidelity mockups",
            "interactive prototyping (Figma, Axure, ProtoPie)",
            "design handoff (Figma Dev Mode, Zeplin)",
            "micro-interactions and animation principles",
            # Design systems
            "design tokens", "component libraries", "atomic design",
            "responsive and adaptive design", "platform guidelines (HIG, Material Design)",
            "version control for design (Figma branching, Abstract)",
            # Visual design
            "typography hierarchy", "colour theory and accessible colour contrast (WCAG AA/AAA)",
            "gestalt principles", "visual hierarchy and layout grids",
            "icon design", "illustration style consistency",
            # Accessibility
            "WCAG 2.1 / 2.2 guidelines", "ARIA roles and labels",
            "keyboard navigation design", "screen reader compatibility",
            "colour blindness considerations", "inclusive design principles",
            # Metrics & measurement
            "usability metrics (task completion, error rate, time-on-task)",
            "NPS and CSAT", "System Usability Scale (SUS)",
            "funnel analysis", "heatmaps and session recordings",
            "design iteration cycles", "OKRs tied to UX outcomes",
            # Collaboration
            "design critique facilitation", "stakeholder alignment",
            "design sprints", "cross-functional collaboration with engineering",
            "documenting design decisions and rationale",
        ],
        "forbidden_resume_keywords": [
            "tensorflow", "docker", "kubernetes", "SQL queries",
            "backend API design", "microservices", "CI/CD pipelines",
            "server infrastructure", "network security", "penetration testing",
        ],
        "context_override": (
            "This is a UI/UX Design interview. Cover the full design process from research "
            "(user interviews, usability testing) through information architecture, wireframing, "
            "high-fidelity prototyping, design systems, and accessibility. "
            "Probe for research rigour, design decision justification, stakeholder communication, "
            "and how the candidate measures design impact — not just tool proficiency."
        ),
    },

    "Project Management": {
        "aliases": ["project manager", "product manager", "scrum master", "agile coach",
                    "program manager", "technical program manager", "delivery manager"],
        "mandatory_topics": [
            # Methodologies
            "Agile (Scrum, Kanban, SAFe, LeSS)", "Waterfall and hybrid approaches",
            "sprint planning", "backlog refinement", "sprint retrospectives",
            "definition of done vs definition of ready", "velocity and story points",
            "epic, story, task hierarchy", "release planning",
            # Product management
            "product vision and strategy", "product roadmapping (now/next/later, theme-based)",
            "OKRs and KPI definition", "prioritisation frameworks (RICE, MoSCoW, Kano, WSJF)",
            "product discovery", "opportunity sizing", "market research",
            "customer journey mapping", "user story writing",
            "go-to-market planning", "launch checklists",
            # Stakeholder management
            "stakeholder mapping", "RACI matrix", "executive communication",
            "managing up vs managing down", "conflict resolution",
            "requirements gathering and sign-off", "change management",
            # Risk & delivery
            "risk identification and RAID log", "risk mitigation strategies",
            "dependency mapping", "critical path analysis",
            "scope creep management", "escalation paths",
            "delivery metrics (cycle time, lead time, throughput)",
            # Technical program management
            "technical debt management", "cross-team dependency management",
            "technical roadmap alignment", "engineering capacity planning",
            "incident retrospectives", "architecture decision records (ADRs)",
            # Data & metrics
            "funnel metrics", "retention metrics", "activation, engagement, churn",
            "hypothesis-driven development", "experiment design",
            "dashboard creation", "reporting to leadership",
            # Tools
            "Jira", "Linear", "Confluence", "Notion", "Asana",
            "Miro / FigJam for workshops", "ProductBoard", "Amplitude / Mixpanel",
        ],
        "forbidden_resume_keywords": [
            "react", "tensorflow", "docker", "SQL joins", "API development",
            "frontend CSS", "kubernetes", "penetration testing",
            "SIEM tools", "malware analysis",
        ],
        "context_override": (
            "This is a Project/Product Management interview. Cover planning methodologies (Agile/Scrum), "
            "prioritisation frameworks (RICE, MoSCoW, WSJF), stakeholder communication, risk management, "
            "product discovery and roadmapping, delivery metrics, and data-driven decision making. "
            "For PM roles, probe for customer empathy and impact measurement. "
            "For TPM roles, probe for technical dependency management and engineering collaboration. "
            "Demand specific examples — STAR method, not generic process descriptions."
        ),
    },
}

# Generic fallback for domains not explicitly configured
_DEFAULT_DOMAIN_CONFIG = {
    "mandatory_topics": [],
    "forbidden_resume_keywords": [],
    "context_override": "",
}

# =============================================================================
# TASK 1: DOMAIN → ROLE MAPPING
# Used to validate and synchronise selected_domain and target_role in session_state.
# Always pass these together: st.session_state.selected_domain + st.session_state.target_role
# =============================================================================
DOMAIN_ROLES = {
    "Software Development and Engineering": [
        "Frontend Developer",
        "Backend Developer",
        "Full Stack Developer",
        "Mobile App Developer",
        "Game Developer",
    ],
    "Data Science and Analytics": [
        "Data Scientist",
        "Data Analyst",
        "Machine Learning Engineer",
    ],
    "Cloud Computing and DevOps": [
        "Cloud Architect",
        "DevOps Engineer",
        "Site Reliability Engineer",
    ],
    "Cybersecurity": [
        "Security Analyst",
        "Penetration Tester",
    ],
    "UI/UX Design": [
        "UI Designer",
        "UX Designer",
    ],
    "Project Management": [
        "Project Manager",
        "Product Manager",
    ],
}


def get_valid_role_for_domain(domain: str, current_role: str = None) -> str:
    """
    Returns a valid role for the given domain.
    If current_role is already valid for the domain, returns it unchanged.
    Otherwise returns the first valid role for the domain.
    Always call this when domain changes to keep selected_domain and target_role in sync.
    """
    roles = DOMAIN_ROLES.get(domain, [])
    if not roles:
        return current_role or ""
    if current_role and current_role in roles:
        return current_role
    return roles[0]


# =============================================================================
# TASK 2: EXPANDED KEY_TOPICS_IN_SCOPE
# Comprehensive interview topic bank per domain and role.
# Structure: { "Domain": { "Role": ["topic1", "topic2", ...] } }
# Compatible with existing generate_key_topics() and get_domain_config() logic.
# =============================================================================
KEY_TOPICS_IN_SCOPE = {
    "Software Development and Engineering": {
        "Frontend Developer": [
            "HTML5 semantics and accessibility",
            "CSS3 layouts: Flexbox and Grid",
            "JavaScript ES6+ features (arrow functions, destructuring, spread/rest)",
            "TypeScript fundamentals and type safety",
            "DOM manipulation and event handling",
            "Responsive and adaptive design",
            "React hooks (useState, useEffect, useContext, useMemo, useCallback)",
            "React component lifecycle and reconciliation",
            "State management (Redux, Zustand, Recoil)",
            "React Query / TanStack Query for async state",
            "Next.js: SSR vs SSG vs ISR",
            "API integration with Fetch and Axios",
            "Frontend performance optimisation (Core Web Vitals, LCP, CLS, INP)",
            "Code splitting and lazy loading",
            "Bundle analysis and tree shaking (Webpack, Vite)",
            "CSS-in-JS (styled-components, Emotion)",
            "Tailwind CSS utility-first approach",
            "Accessibility (WCAG 2.1, ARIA attributes, keyboard navigation)",
            "Cross-browser compatibility and polyfills",
            "Unit testing with Jest and React Testing Library",
            "End-to-end testing with Cypress or Playwright",
            "Web security basics (XSS, CSRF, CSP headers)",
            "Browser storage (localStorage, sessionStorage, IndexedDB)",
            "Micro-frontend architecture and module federation",
            "Design systems and component libraries",
            "Storybook for component documentation",
            "Progressive Web Apps (PWA) and service workers",
            "WebSockets and real-time UI updates",
            "Internationalisation (i18n) and localisation",
            "Monorepo tooling (Nx, Turborepo)",
        ],
        "Backend Developer": [
            "REST API design principles (CRUD, statelessness, versioning)",
            "GraphQL schema design, resolvers, and N+1 problem",
            "gRPC and Protocol Buffers",
            "Node.js event loop and async programming",
            "Python (FastAPI, Django, Flask) backend patterns",
            "Java Spring Boot and dependency injection",
            "Go goroutines and channels",
            "PostgreSQL indexing, query optimisation, EXPLAIN ANALYZE",
            "MySQL transactions and ACID properties",
            "MongoDB schema design and aggregation pipelines",
            "Redis caching strategies (cache-aside, write-through, cache invalidation)",
            "Database connection pooling and N+1 prevention",
            "Microservices architecture and inter-service communication",
            "Message queues (RabbitMQ, Kafka, AWS SQS)",
            "Event-driven architecture and pub/sub patterns",
            "Circuit breaker and bulkhead patterns",
            "CAP theorem and eventual consistency",
            "Horizontal vs vertical scaling",
            "Load balancing and reverse proxy (Nginx)",
            "Authentication (JWT, OAuth 2.0, OpenID Connect)",
            "Role-based access control (RBAC) and ABAC",
            "Input validation and OWASP Top 10 for APIs",
            "Secrets management (Vault, AWS Secrets Manager)",
            "Background jobs and task queues (Celery, BullMQ)",
            "Unit testing and integration testing with mocking",
            "Contract testing (Pact)",
            "API rate limiting and throttling",
            "Idempotency in distributed systems",
            "Database migrations and schema evolution",
            "OpenAPI / Swagger documentation",
        ],
        "Full Stack Developer": [
            "Full request lifecycle: browser → frontend → API → database → response",
            "React or Vue.js frontend framework internals",
            "Next.js or Nuxt.js full-stack routing and data fetching",
            "Node.js or Django/FastAPI backend API design",
            "PostgreSQL and MongoDB data modelling",
            "REST API and GraphQL API integration",
            "JWT and session-based authentication",
            "OAuth 2.0 social login flows",
            "CORS configuration and security headers",
            "State management (Redux, Zustand, Pinia)",
            "Frontend performance (Core Web Vitals, lazy loading)",
            "Backend performance (connection pooling, caching, async workers)",
            "Docker containerisation for local development",
            "CI/CD pipelines (GitHub Actions, GitLab CI)",
            "Deployment to Vercel, Railway, AWS, or GCP",
            "Environment variables and 12-factor app principles",
            "WebSockets for real-time features",
            "Serverless functions and edge computing",
            "Unit, integration, and end-to-end testing across the stack",
            "Database ORM usage (Prisma, Sequelize, SQLAlchemy)",
            "API error handling and validation",
            "Monorepo structure for shared frontend/backend code",
            "TypeScript across frontend and backend",
            "BFF (Backend For Frontend) pattern",
            "Feature flags and A/B testing",
            "Accessibility (WCAG) and SEO basics",
            "Web security: XSS, CSRF, SQL injection prevention",
            "Logging and error monitoring (Sentry, Datadog)",
            "Agile development and sprint workflows",
            "Code review practices and pull request workflows",
        ],
        "Mobile App Developer": [
            "React Native architecture (bridge, JSI, Hermes engine)",
            "Flutter widgets and Dart language fundamentals",
            "Swift UIKit and SwiftUI for iOS",
            "Kotlin Jetpack Compose for Android",
            "Mobile navigation patterns (stack, tab, drawer)",
            "State management in mobile (Redux, MobX, Provider, Riverpod)",
            "Offline-first architecture and local data persistence (SQLite, Realm, MMKV)",
            "Mobile networking (REST APIs, GraphQL, WebSockets)",
            "Push notifications (FCM, APNs)",
            "Deep linking and universal links",
            "Authentication in mobile (biometrics, OAuth, Keychain/Keystore)",
            "Mobile performance optimisation (FlatList, image caching, lazy rendering)",
            "Mobile accessibility (VoiceOver, TalkBack)",
            "App store submission and code signing (iOS and Android)",
            "Over-the-air updates (CodePush, EAS Update)",
            "Mobile testing (unit, integration, E2E with Detox or Maestro)",
            "Mobile CI/CD (Fastlane, Bitrise, EAS Build)",
            "Native modules and bridging to platform APIs",
            "Responsive layouts for different screen sizes and orientations",
            "Background processing and app lifecycle management",
            "Camera, GPS, and sensor integration",
            "Mobile security (certificate pinning, secure storage, jailbreak detection)",
            "Analytics and crash reporting (Firebase, Mixpanel, Crashlytics)",
            "Mobile design patterns (MVVM, Clean Architecture)",
            "App size optimisation and ProGuard/R8 for Android",
        ],
        "Game Developer": [
            "Unity engine architecture (GameObjects, Components, Scenes)",
            "Unreal Engine Blueprints and C++ gameplay programming",
            "Game loop design (update, render, fixed update)",
            "Physics engine integration (Rigidbody, colliders, raycasting)",
            "2D sprite animation and skeletal animation systems",
            "3D mesh, material, and shader fundamentals",
            "Shader programming (HLSL, GLSL, ShaderLab)",
            "Lighting models (PBR, baked vs real-time lighting)",
            "Camera systems (third-person, first-person, cinematic)",
            "AI for games (finite state machines, behaviour trees, NavMesh pathfinding)",
            "Multiplayer networking (authoritative server, client-side prediction, lag compensation)",
            "Game data persistence (save systems, PlayerPrefs, serialisation)",
            "Audio integration (FMOD, Wwise, Unity Audio)",
            "UI/UX for games (HUD design, menu flows, accessibility)",
            "Memory management and garbage collection in game engines",
            "Performance profiling (frame time, draw calls, batching)",
            "Level of Detail (LOD) and occlusion culling",
            "Procedural content generation",
            "Particle systems and VFX",
            "Game design patterns (object pooling, observer, command)",
            "Input systems (keyboard, controller, touch)",
            "Asset pipeline and version control with large binary files (Git LFS)",
            "Platform publishing (Steam, App Stores, console certification)",
            "Monetisation mechanics (IAP, ads, battle pass)",
            "Testing in games (unit tests for game logic, QA playtesting)",
        ],
    },

    "Data Science and Analytics": {
        "Data Scientist": [
            "Exploratory data analysis (EDA) and data profiling",
            "Data cleaning and missing value imputation",
            "Feature engineering and feature selection",
            "Pandas, NumPy, and SciPy for data manipulation",
            "Linear regression and logistic regression",
            "Decision trees, random forests, and gradient boosting (XGBoost, LightGBM)",
            "Bias-variance tradeoff and regularisation (L1/L2, ElasticNet)",
            "Cross-validation (k-fold, stratified, time-series split)",
            "Model evaluation metrics (precision, recall, F1, ROC-AUC, RMSE, MAE)",
            "Hypothesis testing (t-test, chi-squared, ANOVA)",
            "A/B testing design and statistical significance",
            "Bayesian reasoning and confidence intervals",
            "SQL for analytical queries (window functions, CTEs, GROUP BY)",
            "Time series forecasting (ARIMA, Prophet, SARIMA)",
            "Clustering algorithms (k-means, DBSCAN, hierarchical)",
            "Dimensionality reduction (PCA, t-SNE, UMAP)",
            "Class imbalance handling (SMOTE, cost-sensitive learning)",
            "Natural language processing basics (TF-IDF, word embeddings)",
            "Data visualisation with Matplotlib, Seaborn, and Plotly",
            "Dashboard design and data storytelling",
            "Business communication of technical findings",
            "Experiment design and causal inference",
            "Data ethics and bias detection in models",
            "ETL pipelines and data warehouse basics",
            "Cloud-based data platforms (BigQuery, Redshift, Snowflake)",
            "Jupyter notebooks and reproducible research practices",
            "Model interpretability (SHAP, LIME, feature importance)",
            "Deployment of models as APIs (Flask, FastAPI)",
            "Version control for data science (DVC, MLflow)",
            "Regulatory considerations (GDPR, data privacy in ML)",
        ],
        "Data Analyst": [
            "SQL proficiency (joins, subqueries, window functions, CTEs)",
            "Query optimisation and execution plans",
            "Descriptive statistics (mean, median, mode, variance, standard deviation)",
            "Data cleaning: handling nulls, duplicates, and outliers",
            "Pandas for tabular data manipulation",
            "Exploratory data analysis workflows",
            "Data visualisation principles (chart selection, colour, labelling)",
            "Tableau dashboard design and calculated fields",
            "Power BI reports, DAX formulas, and data modelling",
            "Looker and LookML basics",
            "Google Sheets and Excel advanced formulas (VLOOKUP, pivot tables)",
            "KPI definition and metric frameworks",
            "Cohort analysis and retention metrics",
            "Funnel analysis and conversion rate optimisation",
            "A/B test result interpretation",
            "Statistical significance and practical significance",
            "Business intelligence and self-serve analytics culture",
            "Data storytelling and executive presentations",
            "ETL processes and data pipeline basics",
            "Data warehouse concepts (star schema, OLAP, OLTP)",
            "dbt for data transformation",
            "Stakeholder requirements gathering for analytical projects",
            "Data governance and data quality standards",
            "Python scripting for data tasks (pandas, matplotlib)",
            "Working with APIs to pull data",
            "Time series trend analysis",
            "Segmentation and customer analytics",
            "Hypothesis formation and testing for business decisions",
            "Reporting automation and scheduled reports",
            "Data documentation and lineage tracking",
        ],
        "Machine Learning Engineer": [
            "Supervised vs unsupervised vs reinforcement learning paradigms",
            "Feature engineering, scaling, and encoding",
            "Scikit-learn pipelines for reproducible ML workflows",
            "Gradient boosting: XGBoost, LightGBM, CatBoost",
            "Deep learning with PyTorch and TensorFlow",
            "CNN architecture (convolution, pooling, stride)",
            "RNN, LSTM, and GRU for sequential data",
            "Transformer architecture and self-attention mechanism",
            "Transfer learning and fine-tuning pre-trained models",
            "BERT, GPT, and LLM fundamentals",
            "Prompt engineering and RAG (retrieval-augmented generation)",
            "Vector databases (Pinecone, Weaviate, FAISS, pgvector)",
            "MLOps: experiment tracking with MLflow and Weights & Biases",
            "ML pipelines: Kubeflow, SageMaker Pipelines, Vertex AI",
            "Feature stores (Feast, Tecton)",
            "Model versioning and model registry",
            "Data drift and concept drift detection",
            "Model monitoring in production (performance, latency, quality)",
            "A/B testing and shadow deployment for ML models",
            "Model serialisation (ONNX, TorchScript, pickle)",
            "Model serving (TensorFlow Serving, Triton, FastAPI inference)",
            "Batch vs real-time inference tradeoffs",
            "Model quantisation, pruning, and distillation",
            "GPU utilisation and distributed training (DDP, FSDP)",
            "Hyperparameter tuning (Optuna, Ray Tune, Hyperopt)",
            "LLM fine-tuning (LoRA, QLoRA, PEFT)",
            "Responsible AI: fairness, explainability, and bias auditing",
            "Data labelling workflows and annotation quality",
            "Cloud ML platforms (AWS SageMaker, GCP Vertex AI, Azure ML)",
            "Cost optimisation for ML workloads",
        ],
    },

    "Cloud Computing and DevOps": {
        "Cloud Architect": [
            "Cloud service models: IaaS, PaaS, SaaS, FaaS",
            "AWS core services: EC2, S3, RDS, Lambda, VPC, IAM, CloudFront",
            "Azure core services: AKS, App Service, Azure Functions, Blob Storage, AAD",
            "GCP core services: GKE, Cloud Run, BigQuery, Cloud Functions, IAM",
            "Multi-cloud and hybrid cloud architectures",
            "Cloud networking: VPC design, subnets, security groups, NACLs, peering",
            "DNS management and Route 53 / Cloud DNS",
            "Content delivery networks (CDN) and edge caching",
            "Load balancing: ALB, NLB, Global Load Balancer",
            "Auto-scaling groups and elasticity patterns",
            "Serverless architecture design and cold start optimisation",
            "Containerisation with Docker and orchestration with Kubernetes (EKS, GKE, AKS)",
            "Infrastructure as Code: Terraform state management, modules, workspaces",
            "AWS CloudFormation and CDK",
            "GitOps workflows (ArgoCD, Flux)",
            "Cloud cost optimisation (reserved instances, spot/preemptible, rightsizing)",
            "Cloud storage options (object, block, file) and their tradeoffs",
            "Database cloud services: RDS, DynamoDB, Firestore, Cloud Spanner",
            "Data streaming and analytics (Kinesis, Pub/Sub, Dataflow, EMR)",
            "Cloud security: IAM least privilege, KMS, Secrets Manager, CSPM",
            "Disaster recovery: RTO, RPO, multi-region failover, backup strategies",
            "Well-Architected Framework (reliability, security, performance, cost, ops)",
            "Cloud migration strategies (6 Rs: rehost, replatform, refactor, etc.)",
            "Cloud compliance and governance (SOC 2, ISO 27001, PCI-DSS)",
            "Landing zone design and account/organisation structure",
        ],
        "DevOps Engineer": [
            "CI/CD pipeline design (GitHub Actions, GitLab CI, Jenkins, CircleCI)",
            "Pipeline stages: build, test, security scan, deploy",
            "Deployment strategies: blue-green, canary, rolling, feature flags",
            "Rollback strategies and release management",
            "Trunk-based development and branching strategies",
            "Docker: image layers, multi-stage builds, registry management",
            "Kubernetes: pods, deployments, services, ingress, ConfigMaps, Secrets",
            "Kubernetes resource requests/limits, HPA, and VPA",
            "Helm charts for Kubernetes packaging",
            "Service mesh (Istio, Linkerd) and mTLS",
            "Terraform for infrastructure provisioning",
            "Ansible for configuration management",
            "GitOps with ArgoCD or Flux",
            "AWS/GCP/Azure DevOps toolchain integration",
            "Observability: Prometheus metrics, Grafana dashboards",
            "Centralized logging: ELK stack, Loki, CloudWatch Logs",
            "Distributed tracing: Jaeger, Zipkin, OpenTelemetry",
            "Alerting and on-call: PagerDuty, OpsGenie, SLOs, SLAs, error budgets",
            "Secrets management: HashiCorp Vault, AWS Secrets Manager",
            "SAST/DAST and vulnerability scanning in pipelines (Trivy, Snyk, Sonar)",
            "Container security and CIS benchmarks",
            "Network policies and Kubernetes security contexts",
            "Chaos engineering and fault injection (Chaos Monkey, LitmusChaos)",
            "Incident management: runbooks, postmortems, blameless culture",
            "Developer experience: internal developer platforms, golden paths",
            "Platform engineering and self-service infrastructure",
        ],
        "Site Reliability Engineer": [
            "SLO, SLA, SLI definitions and error budget calculation",
            "Toil reduction and automation philosophy",
            "Incident management: response, escalation, and postmortems",
            "On-call runbook creation and maintenance",
            "Capacity planning and demand forecasting",
            "Observability stack: metrics (Prometheus), logs (ELK/Loki), traces (Jaeger)",
            "Alerting strategy: avoiding alert fatigue, symptom-based alerts",
            "Distributed systems failure modes (network partition, split brain, cascading failures)",
            "Chaos engineering and game days",
            "Kubernetes cluster operations and node management",
            "Service mesh for traffic management and canary analysis",
            "Blue-green and canary deployments with automated rollback",
            "Database reliability: replication, failover, backup/restore testing",
            "Load testing and performance benchmarking (k6, Locust, JMeter)",
            "Cache reliability and Redis/Memcached HA patterns",
            "CDN configuration and origin shield strategies",
            "DDoS mitigation and rate limiting at the infrastructure level",
            "Secret rotation and certificate lifecycle management",
            "Cloud cost attribution and chargeback models",
            "Multi-region active-active and active-passive architectures",
            "Terraform and GitOps for infrastructure changes",
            "Change management and progressive delivery",
            "Engineering reliability into system design (redundancy, bulkheads)",
            "Communication of reliability metrics to product and leadership",
            "Developer enablement and reliability culture building",
        ],
    },

    "Cybersecurity": {
        "Security Analyst": [
            "CIA triad: confidentiality, integrity, availability",
            "Defence in depth and zero trust architecture principles",
            "OWASP Top 10 vulnerabilities and mitigations",
            "Threat modelling (STRIDE, PASTA, attack trees)",
            "Risk assessment and CVSS scoring",
            "SIEM tools: Splunk, Microsoft Sentinel, Chronicle",
            "Log analysis: identifying IOCs and IOAs",
            "MITRE ATT&CK framework: tactics, techniques, procedures",
            "Threat intelligence: feeds, indicators, and enrichment",
            "Incident response lifecycle: preparation, detection, containment, eradication, recovery",
            "Digital forensics basics: disk imaging, memory forensics, chain of custody",
            "Network traffic analysis (Wireshark, Zeek, Suricata)",
            "Endpoint detection and response (EDR) tools (CrowdStrike, SentinelOne)",
            "Vulnerability management lifecycle: scan, prioritise, remediate, verify",
            "Patch management processes",
            "Authentication security: MFA, OAuth 2.0, SAML, JWT pitfalls",
            "Cloud security: IAM misconfiguration, S3 exposure, CSPM tools",
            "Container security scanning (Trivy, Snyk, Aqua Security)",
            "Secrets management (Vault, AWS Secrets Manager)",
            "Security awareness and phishing simulation",
            "Compliance frameworks: GDPR, SOC 2, ISO 27001, PCI-DSS",
            "Security policies, procedures, and governance",
            "SOAR platforms and playbook automation",
            "Cryptography basics: AES, RSA, TLS/SSL, PKI",
            "Secure SDLC integration and DevSecOps practices",
        ],
        "Penetration Tester": [
            "Penetration testing methodology (planning, recon, scanning, exploitation, reporting)",
            "OWASP Testing Guide for web applications",
            "Recon and OSINT techniques (Shodan, theHarvester, Maltego)",
            "Network scanning with Nmap (service enumeration, OS detection)",
            "Vulnerability scanning (Nessus, OpenVAS)",
            "Web application testing with Burp Suite (intercept, scanner, intruder)",
            "SQL injection: manual and automated exploitation",
            "Cross-site scripting (XSS): reflected, stored, DOM-based",
            "SSRF, IDOR, and business logic flaws",
            "Authentication bypass and session management attacks",
            "Exploitation frameworks: Metasploit, ExploitDB",
            "Privilege escalation techniques: Linux (SUID, cron, sudo) and Windows (token impersonation, registry)",
            "Post-exploitation: persistence, lateral movement, pivoting",
            "Active Directory attacks: Kerberoasting, Pass-the-Hash, BloodHound",
            "API penetration testing (REST, GraphQL, broken object level auth)",
            "Cloud penetration testing (AWS, GCP, Azure misconfigurations)",
            "Mobile app penetration testing (SSL pinning bypass, insecure storage)",
            "Wireless security testing (WPA2 cracking, rogue AP)",
            "Social engineering and phishing campaign design",
            "Buffer overflow and memory corruption basics",
            "Reporting: severity classification, CVSS scoring, executive summary writing",
            "Red team vs blue team vs purple team concepts",
            "Bug bounty platforms and responsible disclosure",
            "Legal and ethical considerations in penetration testing",
            "Defensive mindset: hardening recommendations after assessment",
        ],
    },

    "UI/UX Design": {
        "UI Designer": [
            "Visual design principles: hierarchy, contrast, alignment, proximity, repetition",
            "Typography: typeface selection, scale, line height, readability",
            "Colour theory: hue, saturation, value, accessible contrast (WCAG AA/AAA)",
            "Grid systems and layout composition",
            "Iconography: consistency, clarity, and scalability",
            "Illustration style definition and usage guidelines",
            "Figma: components, variants, auto-layout, design tokens",
            "Design system creation: component libraries and token architecture",
            "Atomic design methodology (atoms, molecules, organisms)",
            "Responsive and adaptive design: breakpoints, fluid grids",
            "Platform-specific guidelines: iOS HIG and Material Design 3",
            "Interactive prototyping in Figma (smart animate, overlays)",
            "Micro-interactions and animation principles (easing, duration, feedback)",
            "Design handoff: Figma Dev Mode, Zeplin, annotated specs",
            "Version control for design (Figma branching, Abstract)",
            "Dark mode and theme support in design systems",
            "Accessibility in UI: touch target size, colour contrast, motion sensitivity",
            "ARIA roles and semantic HTML from a design perspective",
            "High-fidelity mockup production workflow",
            "Design critique facilitation and feedback incorporation",
            "Cross-functional collaboration with developers and product managers",
            "Motion design and after-effects basics for UI transitions",
            "Brand consistency across digital products",
            "Measuring design quality: visual QA and pixel-perfection standards",
            "Design documentation and decision rationale recording",
        ],
        "UX Designer": [
            "User research methods: interviews, contextual inquiry, diary studies, surveys",
            "Research synthesis: affinity mapping, thematic analysis, insight generation",
            "Persona development and jobs-to-be-done framework",
            "Customer journey mapping and service blueprinting",
            "Information architecture: site maps, card sorting, tree testing",
            "User flows, task flows, and decision trees",
            "Mental models and progressive disclosure",
            "Low-fidelity wireframing and rapid sketching",
            "High-fidelity mockups and interactive prototyping (Figma, ProtoPie, Axure)",
            "Usability testing: moderated and unmoderated, think-aloud protocol",
            "A/B testing for UX decisions",
            "Heuristic evaluation (Nielsen's 10 heuristics)",
            "Accessibility: WCAG 2.1/2.2, inclusive design, keyboard navigation",
            "System Usability Scale (SUS) and usability metrics",
            "NPS, CSAT, task completion rate, and error rate measurement",
            "Funnel analysis and heatmaps (Hotjar, FullStory)",
            "Design sprints and rapid ideation workshops",
            "Stakeholder alignment and design strategy communication",
            "OKR-tied UX outcomes and design impact measurement",
            "Design systems usage from a UX consistency perspective",
            "Responsive design considerations for UX flows",
            "Mobile UX patterns (bottom navigation, gestures, thumb zones)",
            "Ethical design: dark patterns, consent, and privacy by design",
            "Cross-cultural UX and localisation considerations",
            "Continuous discovery and product-led growth thinking",
        ],
    },

    "Project Management": {
        "Project Manager": [
            "Agile frameworks: Scrum, Kanban, SAFe, LeSS",
            "Waterfall and hybrid project management approaches",
            "Sprint planning, backlog refinement, and sprint retrospectives",
            "Definition of Done vs Definition of Ready",
            "Velocity, story points, and capacity planning",
            "Epic, story, and task hierarchy",
            "Release planning and roadmap management",
            "Risk identification and RAID log management",
            "Risk mitigation strategies and contingency planning",
            "Dependency mapping and critical path analysis",
            "Scope creep management and change control",
            "Stakeholder mapping and stakeholder communication plans",
            "RACI matrix and accountability frameworks",
            "Managing up: executive communication and status reporting",
            "Conflict resolution and team facilitation",
            "Requirements gathering and sign-off process",
            "Delivery metrics: cycle time, lead time, throughput, burndown",
            "Budget management and earned value analysis",
            "Cross-team dependency management",
            "Technical debt management and negotiation with engineering",
            "Incident retrospectives and process improvement",
            "Tools: Jira, Confluence, Linear, Asana, Notion",
            "Miro / FigJam for workshops and planning sessions",
            "Escalation paths and decision-making frameworks",
            "Agile transformation and coaching teams",
        ],
        "Product Manager": [
            "Product vision and strategy definition",
            "Product roadmapping: now/next/later, theme-based roadmaps",
            "OKR definition and product KPI frameworks",
            "Prioritisation frameworks: RICE, MoSCoW, Kano, WSJF",
            "Product discovery: user interviews, problem validation, opportunity sizing",
            "Market research and competitive analysis",
            "Customer journey mapping and pain point identification",
            "User story writing (as a user, I want... so that...)",
            "Acceptance criteria and definition of done",
            "Go-to-market planning and launch checklists",
            "Hypothesis-driven development and experiment design",
            "A/B testing and feature flag rollouts",
            "Funnel metrics: activation, engagement, retention, churn",
            "NPS, CSAT, and product analytics (Amplitude, Mixpanel)",
            "Dashboard creation and data-driven decision making",
            "Working with engineering: technical feasibility and effort estimation",
            "Working with design: UX research integration and design feedback",
            "Stakeholder management and executive buy-in",
            "Product-led growth (PLG) principles",
            "Platform vs product thinking",
            "Agile ceremonies as product manager (sprint planning, review, retro)",
            "Business model understanding: revenue, margins, unit economics",
            "Regulatory and compliance considerations for product decisions",
            "Technical literacy: APIs, databases, system architecture basics",
            "ProductBoard and roadmap communication tools",
        ],
    },
}


def get_domain_config(domain: str) -> dict:
    """Return domain config by exact name or alias match."""
    if domain in DOMAIN_AUTHORITY_CONFIG:
        return DOMAIN_AUTHORITY_CONFIG[domain]
    domain_lower = domain.lower()
    for key, cfg in DOMAIN_AUTHORITY_CONFIG.items():
        if any(alias in domain_lower for alias in cfg.get("aliases", [])):
            return cfg
    return _DEFAULT_DOMAIN_CONFIG


def generate_key_topics(resume_context: dict, domain_config: dict, selected_role: str) -> list:
    """
    Generate a ranked, deduplicated list of up to 10 key topics in scope for the
    interview, combining three sources in priority order:

        1. Resume technologies  (most specific to the candidate)
        2. Resume skills        (secondary candidate signal)
        3. Domain mandatory_topics (domain-level fundamentals, fill remaining slots)

    Parameters
    ----------
    resume_context : dict
        Parsed resume data with keys: "technologies", "skills", "projects", "experience".
        May be None or empty — function degrades gracefully to domain topics only.
    domain_config : dict
        Output of get_domain_config() for the selected domain.
    selected_role : str
        The sub-role chosen by the user (e.g. "Frontend Developer", "ML Engineer").
        Used as a tiebreaker to surface more relevant domain topics when the resume
        is sparse.

    Returns
    -------
    list[str]
        Between 1 and 10 topic strings, never empty.
    """
    MAX_TOPICS = 10

    # ── Normalise inputs ──────────────────────────────────────────────────────
    rc = resume_context or {}
    domain_mandatory = domain_config.get("mandatory_topics", [])
    forbidden = [kw.lower() for kw in domain_config.get("forbidden_resume_keywords", [])]

    def _clean(items):
        """Deduplicate, strip empties, remove items that are forbidden in this domain."""
        seen = set()
        out = []
        for item in items:
            norm = item.strip()
            if not norm:
                continue
            # Drop resume items that conflict with the selected domain
            if any(f in norm.lower() for f in forbidden):
                continue
            key = norm.lower()
            if key not in seen:
                seen.add(key)
                out.append(norm)
        return out

    # ── Source 1: resume technologies ────────────────────────────────────────
    resume_techs = _clean(rc.get("technologies", []))

    # ── Source 2: resume skills ───────────────────────────────────────────────
    # Exclude anything already captured from technologies
    tech_keys = {t.lower() for t in resume_techs}
    resume_skills = _clean(
        [s for s in rc.get("skills", []) if s.strip().lower() not in tech_keys]
    )

    # ── Source 3: domain mandatory topics ────────────────────────────────────
    # Prioritise mandatory topics whose text overlaps with the selected role name,
    # so a "Frontend Developer" sees React/CSS before obscure backend topics.
    role_lower = selected_role.lower()
    role_keywords = set(role_lower.replace("-", " ").split())

    def _role_relevance(topic: str) -> int:
        """Higher = more relevant to the selected role."""
        t_lower = topic.lower()
        return sum(1 for kw in role_keywords if kw in t_lower)

    already_seen = {t.lower() for t in resume_techs + resume_skills}
    domain_topics_filtered = [
        t for t in domain_mandatory if t.strip().lower() not in already_seen
    ]
    domain_topics_sorted = sorted(
        domain_topics_filtered,
        key=_role_relevance,
        reverse=True
    )

    # ── Merge in priority order ───────────────────────────────────────────────
    combined = resume_techs[:MAX_TOPICS]
    remaining = MAX_TOPICS - len(combined)

    if remaining > 0:
        combined += resume_skills[:remaining]
        remaining = MAX_TOPICS - len(combined)

    if remaining > 0:
        combined += domain_topics_sorted[:remaining]

    # ── Guarantee non-empty list ──────────────────────────────────────────────
    if not combined:
        combined = domain_mandatory[:MAX_TOPICS] or [selected_role + " fundamentals"]

    return combined[:MAX_TOPICS]


def filter_resume_for_domain(resume_context: dict, selected_domain: str) -> dict:
    """
    DOMAIN AUTHORITY LAYER — Core Function.

    Strips resume skills/technologies that are IRRELEVANT to the selected domain
    and flags that domain override is active. This prevents a Full Stack resume
    from contaminating a Data Analyst interview prompt.

    Returns a modified resume_context dict safe to pass to question generators.
    """
    cfg = get_domain_config(selected_domain)
    forbidden = [kw.lower() for kw in cfg.get("forbidden_resume_keywords", [])]

    if not forbidden:
        # No filtering needed for this domain
        return resume_context

    def clean_list(items: list) -> list:
        cleaned = []
        for item in items:
            item_lower = item.lower()
            if not any(f in item_lower for f in forbidden):
                cleaned.append(item)
        return cleaned

    filtered = {
        "skills": clean_list(resume_context.get("skills", [])),
        "technologies": clean_list(resume_context.get("technologies", [])),
        # Keep projects/experience but append a domain caveat so LLM understands the interview scope
        "projects": resume_context.get("projects", []),
        "experience": resume_context.get("experience", []),
        "_domain_override": True,
        "_domain_name": selected_domain,
    }

    # If filtering removed everything, add a note so LLM doesn't get empty context
    if not filtered["skills"]:
        filtered["skills"] = [f"Candidate background may differ from {selected_domain} domain"]
    if not filtered["technologies"]:
        filtered["technologies"] = [f"Domain: {selected_domain}"]

    return filtered


def build_domain_authority_block(selected_domain: str, selected_role: str) -> str:
    """
    Returns a strong domain-authority instruction block to prepend to ALL
    question-generation prompts. Forces LLM to stay domain-aligned regardless
    of resume content.
    """
    cfg = get_domain_config(selected_domain)
    override = cfg.get("context_override", "")
    mandatory = cfg.get("mandatory_topics", [])

    block = f"""
⚠️ DOMAIN AUTHORITY OVERRIDE — HIGHEST PRIORITY ⚠️
The candidate has SELECTED to be interviewed as: {selected_role} in {selected_domain}.
Even if the resume shows different experience, ALL questions MUST be about {selected_domain}.
{override}

MANDATORY TOPIC POOL (draw from these for every question):
{', '.join(mandatory) if mandatory else selected_domain + ' core concepts'}

STRICT RULE: Do NOT ask about technologies or concepts outside {selected_domain}.
If resume content conflicts with the selected domain, IGNORE the resume content.
"""
    return block.strip()


# =============================================================================
# ARCHITECTURAL FIX 2: STRUCTURED DIFFICULTY ENFORCER
# =============================================================================
# Problem: Easy/Medium/Hard produce stylistically different questions but not
# structurally different ones. They all look similar in depth.
# Solution: Define a strict question-type contract per difficulty level, enforced
# at the prompt level with explicit templates and forbidden patterns.

# =============================================================================
# TEXT-INTERVIEW-OPTIMISED DIFFICULTY CONTRACTS
# =============================================================================
# Core design principle: every question must be answerable in 3-8 structured
# paragraphs of text. Hard is senior-level but scoped to ONE focused challenge.
# No whiteboard mega-design prompts. No combining scaling + tradeoff + failure
# + architecture in a single question. One core challenge per question.
# =============================================================================

DIFFICULTY_CONTRACTS = {
    "Easy": {
        # ── What it tests ──────────────────────────────────────────────────────
        "label": "Concept Clarity",
        "description": (
            "Concept clarity questions only. The candidate explains WHAT something is, "
            "WHY it exists, and gives a grounded real-world example. "
            "No architecture, no scaling, no tradeoffs, no production scenarios."
        ),
        "answer_scope": "3-5 paragraphs",

        # ── Question structural templates ──────────────────────────────────────
        # LLM must pick ONE of these patterns and fill in the domain-specific topic.
        # Templates are written at the question-generation level — they set structure,
        # not just tone.
        "question_templates": [
            "What is [concept] and why does it exist? Give a concrete example of where you would use it.",
            "Explain the difference between [concept A] and [concept B]. When would you choose one over the other?",
            "Walk me through how [concept] works at a high level. What problem does it solve?",
            "What are the core properties or guarantees of [concept]? Why do those properties matter in practice?",
            "Describe a situation where [concept] would be the right tool and one where it would be the wrong choice.",
        ],

        # ── Hard constraints — these must NEVER appear in Easy questions ────────
        "forbidden_patterns": [
            "design a system", "design an architecture", "at scale", "1 million",
            "production outage", "failure scenario", "optimize for latency",
            "handle 10x", "migrate from", "distributed", "multi-region",
            "zero downtime", "fault tolerance", "SLA", "tradeoff between",
            "compare and contrast in a production context",
        ],

        # ── Scoring calibration ────────────────────────────────────────────────
        "scoring_note": (
            "Award full marks for a clear, correct definition with one real-world example. "
            "Do NOT penalise for missing implementation detail — that belongs to Medium/Hard. "
            "Deduct marks for wrong definitions or examples that show misunderstanding."
        ),
        "followup_allowed": False,
        "cognitive_load": "LOW",
        "cognitive_load_detail": "Definition → Example → One simple comparison. No design decisions.",

        # ── Fallback questions (used when LLM fails) ──────────────────────────
        "fallback_questions": [
            "What is [topic] and why is it important in {domain}?",
            "Explain the difference between two core concepts in {domain} that are often confused.",
            "Give an example of when you would use [concept] in a real project.",
        ],
    },

    "Medium": {
        # ── What it tests ──────────────────────────────────────────────────────
        "label": "Scenario Reasoning",
        "description": (
            "Scenario-based questions with ONE practical constraint or decision point. "
            "The candidate must describe their approach, make one explicit decision or tradeoff, "
            "and briefly justify it. No multi-layer system design. No combined scaling + failure + architecture."
        ),
        "answer_scope": "5-6 paragraphs",

        # ── Question structural templates ──────────────────────────────────────
        "question_templates": [
            "You're implementing [feature/component] for a small production service. "
            "Walk through your approach and explain the ONE key decision you'd make and why.",

            "Your team is choosing between [option A] and [option B] for [use case]. "
            "What factors would you evaluate, and which would you recommend for this context?",

            "You've been asked to add [capability] to an existing codebase without breaking current behaviour. "
            "Describe your implementation strategy and one challenge you'd anticipate.",

            "A junior engineer on your team is confused about when to use [concept]. "
            "How would you explain it, and what example would you use to make it concrete?",

            "You notice [specific problem/smell] in a codebase you've just joined. "
            "What's your diagnosis, and what's the first concrete step you'd take to address it?",
        ],

        # ── Hard constraints ───────────────────────────────────────────────────
        "forbidden_patterns": [
            "design a system from scratch", "handle 1 million concurrent users",
            "multi-region active-active", "full microservices migration",
            "design the entire architecture", "production outage at peak traffic",
            "compare all possible approaches", "list every tradeoff",
        ],

        # ── Scoring calibration ────────────────────────────────────────────────
        "scoring_note": (
            "Award high marks for scenario framing + one clear decision + concrete justification. "
            "Penalise answers that stay purely definitional (no scenario engagement) or that jump to "
            "full system design without scoping to the constraint given. "
            "A good Medium answer reads like a thoughtful Slack message from a mid-level engineer."
        ),
        "followup_allowed": True,
        "cognitive_load": "MEDIUM",
        "cognitive_load_detail": "Scenario → ONE decision → Justification → One risk or alternative considered.",

        "fallback_questions": [
            "Describe a specific implementation challenge you faced with {domain} and how you resolved it.",
            "How would you approach adding [feature] to an existing {role} project without breaking existing behaviour?",
            "A teammate proposes using [technology] for [use case]. What questions would you ask before agreeing?",
        ],
    },

    "Hard": {
        # ── What it tests ──────────────────────────────────────────────────────
        "label": "Focused Technical Depth",
        "description": (
            "Deep technical reasoning on ONE focused high-impact challenge. "
            "Pick EXACTLY ONE of: tradeoff analysis, failure handling, or optimisation under constraint. "
            "Do NOT combine all three in one question. "
            "The candidate must reason at a senior level but the question must remain "
            "text-answerable in 6-8 paragraphs — no whiteboard, no full system diagram."
        ),
        "answer_scope": "6-8 paragraphs",

        # ── Question structural templates (one per challenge type) ─────────────
        # Each template is intentionally NARROW — one challenge, one decision axis.
        "question_templates": [
            # Tradeoff Analysis (one axis only)
            "You need to choose between [approach A] and [approach B] for [specific use case]. "
            "Both are technically valid. Walk through your decision framework: what data or signals "
            "would drive your choice, and what would you accept as a known limitation of your decision?",

            # Failure Handling (one failure mode only)
            "Your [component/service] starts returning elevated error rates under normal load — "
            "no obvious upstream failures. Walk through your diagnostic process step by step: "
            "what would you check first, what signals would you look for, and how would you isolate the cause?",

            # Optimisation under a single constraint
            "You're asked to reduce [specific metric: latency / cost / memory] for [component] "
            "by 40% without changing its external interface. What's your investigation process, "
            "what are the two or three highest-leverage changes you'd consider, and what would you measure to validate success?",

            # Edge case / correctness challenge
            "Describe a non-obvious edge case or failure mode in [concept/system component] "
            "that a developer might miss during implementation. How would you detect it, "
            "handle it gracefully, and prevent it from recurring?",

            # Depth probe on a specific technical decision
            "You've used [technology/pattern] in production. What is the single most important "
            "limitation or risk of that choice that most engineers underestimate? "
            "How did you mitigate it, or how would you mitigate it if given the chance?",
        ],

        # ── Hard constraints — what must NEVER appear in Hard questions ─────────
        "forbidden_patterns": [
            # These turn a Hard question into an unanswerable whiteboard session
            "design a complete system", "design the entire architecture",
            "design and implement X from scratch handling Y million users",
            "walk through every layer of the stack",
            "describe all possible failure modes", "list every tradeoff",
            "design for global scale with multi-region active-active",
            "design + implement + monitor + scale + secure",
            # Scope combiners — pick ONE axis, not all
            "tradeoff AND failure AND scaling AND security",
            "compare all alternatives AND handle failures AND optimize",
        ],

        # ── Scoring calibration ────────────────────────────────────────────────
        "scoring_note": (
            "Award 8-10 for answers that: isolate the core challenge clearly, reason with technical specifics "
            "(not just buzzwords), make a justified decision, and acknowledge one concrete risk or edge case. "
            "Award 5-7 for correct but high-level answers that skip quantification or justification. "
            "Award 0-4 for answers that treat the question like an Easy/Medium or give purely theoretical responses. "
            "A strong Hard answer reads like a well-structured Slack thread from a senior engineer explaining "
            "a decision to their team — not a conference talk or architecture document."
        ),
        "followup_allowed": True,
        "cognitive_load": "HIGH",
        "cognitive_load_detail": (
            "ONE challenge axis → Reasoning framework → Technical specifics → Justified decision → "
            "ONE edge case or risk. Answerable in text. No diagram needed."
        ),

        "fallback_questions": [
            "Describe the most counterintuitive technical tradeoff you've encountered in {domain}. "
            "What made it hard, and how did you ultimately decide?",
            "Walk through how you would diagnose an unexpected performance regression "
            "in a {domain} system you own. What would you check and in what order?",
            "What is the single most dangerous assumption developers make when using "
            "[core technology in {domain}], and how would you guard against it?",
        ],
    },
}


def get_difficulty_instruction_block(difficulty: str) -> str:
    """
    Returns a structured, text-interview-optimised difficulty instruction block.

    This block is injected into EVERY question-generation prompt. It enforces:
    - The correct question TYPE (not just tone)
    - Structural templates the LLM must follow
    - Hard forbidden patterns that prevent mega system-design questions
    - Scope reminders so Hard questions stay text-answerable (6-8 paragraphs)
    """
    contract = DIFFICULTY_CONTRACTS.get(difficulty, DIFFICULTY_CONTRACTS["Medium"])

    templates = "\n".join(f"  TEMPLATE {i+1}: {t}" for i, t in enumerate(contract["question_templates"]))
    forbidden = contract.get("forbidden_patterns", [])
    forbidden_str = (
        "\n⛔ THESE PATTERNS ARE FORBIDDEN — never generate questions that contain:\n"
        + "\n".join(f"  - {p}" for p in forbidden)
        if forbidden else ""
    )

    scope_reminder = ""
    if difficulty == "Hard":
        scope_reminder = """
⚠️  TEXT-INTERVIEW SCOPE RULE FOR HARD:
The question must be answerable in 6-8 paragraphs of text.
Pick EXACTLY ONE challenge axis: tradeoff OR failure handling OR optimisation OR edge case.
Do NOT combine multiple axes (e.g. "design + scale + handle failures + secure" = INVALID).
A well-formed Hard question targets one decision, one failure mode, or one constraint.
"""

    block = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
DIFFICULTY CONTRACT: {difficulty.upper()} — {contract["label"]}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
What this level tests: {contract["description"]}

Cognitive load: {contract["cognitive_load"]} — {contract["cognitive_load_detail"]}
Expected answer scope: {contract["answer_scope"]}

STRUCTURAL TEMPLATES — choose the ONE template that best fits the topic:
{templates}
{forbidden_str}
{scope_reminder}
Scoring context: {contract["scoring_note"]}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"""
    return block.strip()


# =============================================================================
# ARCHITECTURAL FIX 3: SMART ESCALATION ENGINE
# =============================================================================
# Problem: Escalation layer doesn't increase dynamically, weakness detection is
# keyword-based, and Hard mode doesn't increase cognitive pressure per follow-up.
# Solution: Layer-based escalation map with LLM-scored weakness analysis that
# selects strategy from a deterministic mapping of score + answer quality signals.

ESCALATION_LAYER_MAP = {
    1: {
        "name": "Clarification",
        "instruction": "Ask the candidate to clarify or expand on a specific part of their answer that was vague or ambiguous.",
        "trigger": "Used when the answer lacks depth or contains unclear statements.",
        "cognitive_pressure": "LOW",
    },
    2: {
        "name": "Metrics",
        "instruction": "Ask the candidate to justify their answer with specific numbers, benchmarks, or measurable outcomes. Push for concrete data.",
        "trigger": "Used when the answer is conceptually correct but lacks evidence or quantification.",
        "cognitive_pressure": "MEDIUM",
    },
    3: {
        "name": "Tradeoff",
        "instruction": "Challenge the candidate with a direct tradeoff: their approach vs. an alternative. Ask them to defend their choice with clear pros/cons.",
        "trigger": "Used when no tradeoffs were mentioned or the answer seems too one-sided.",
        "cognitive_pressure": "MEDIUM-HIGH",
    },
    4: {
        "name": "Scalability",
        "instruction": "Inject a scale constraint (10x traffic, 100x data volume, global users) and ask how their approach holds up. Push for architectural changes.",
        "trigger": "Used after tradeoffs are discussed or to pressure-test their design thinking.",
        "cognitive_pressure": "HIGH",
    },
    5: {
        "name": "Failure Simulation",
        "instruction": "Simulate a production failure related to their approach. Describe a realistic incident and ask them to diagnose, mitigate, and prevent it.",
        "trigger": "Maximum pressure — only at layer 5. Tests crisis thinking and system ownership.",
        "cognitive_pressure": "MAXIMUM",
    },
}

# Strategy → Layer mapping (which layer best fits each strategy)
STRATEGY_TO_LAYER = {
    "Clarification": 1,
    "Metric Justification": 2,
    "Tradeoff Challenge": 3,
    "Alternative Design Comparison": 3,
    "Scalability Challenge": 4,
    "Constraint Injection": 4,
    "Failure Simulation": 5,
    "Security Consideration": 5,
    "Depth Probe": 1,
    "Edge Case Scenario": 3,
    "Architecture Breakdown": 4,
}


def analyze_answer_weaknesses_smart(answer_text: str, scoring: dict, escalation_layer: int = 1) -> dict:
    """
    UPGRADED Weakness Analyzer — replaces keyword-matching with score-driven
    multi-signal strategy selection.

    Signal priority (in order):
    1. Score deltas between knowledge/communication/relevance
    2. Answer length and structural quality signals
    3. Current escalation layer (forces progression through layers 1→5)
    4. Keyword presence as secondary signals (not sole determinant)

    Returns:
        weaknesses (list): detected weakness signals
        strategy (str): selected follow-up strategy
        next_layer (int): escalation layer for this follow-up
        reasoning (str): human-readable explanation of strategy choice
    """
    knowledge = scoring.get("knowledge", 5)
    communication = scoring.get("communication", 5)
    relevance = scoring.get("relevance", 5)
    avg_score = (knowledge + communication + relevance) / 3
    word_count = len(answer_text.split())
    answer_lower = answer_text.lower()

    weaknesses = []
    reasoning = ""

    # === Score-based signals (primary) ===
    if relevance < 4:
        weaknesses.append("off_topic")
    if knowledge < 4:
        weaknesses.append("weak_knowledge")
    if communication < 5:
        weaknesses.append("weak_communication")
    if word_count < 40:
        weaknesses.append("too_brief")
    if avg_score >= 7.5:
        weaknesses.append("strong_answer")  # Good answer — escalate harder

    # === Structural signals (secondary) ===
    has_metrics = any(kw in answer_lower for kw in [
        "%", "percent", "ms", "milliseconds", "seconds", "users", "requests",
        "throughput", "latency", "uptime", "million", "thousand", "tps", "rps", "gb", "tb"
    ])
    has_tradeoff = any(kw in answer_lower for kw in [
        "tradeoff", "trade-off", "versus", "vs ", "compared to", "alternative",
        "however", "but the downside", "pros and cons", "on the other hand"
    ])
    has_example = any(kw in answer_lower for kw in [
        "for example", "in my project", "we built", "at my", "when i", "i implemented",
        "for instance", "specifically", "in production"
    ])
    has_failure = any(kw in answer_lower for kw in [
        "failure", "outage", "bottleneck", "failed", "bug", "incident", "crash", "timeout"
    ])

    if not has_metrics:
        weaknesses.append("no_metrics")
    if not has_tradeoff:
        weaknesses.append("no_tradeoff")
    if not has_example:
        weaknesses.append("no_concrete_example")

    # === Layer-forced strategy progression ===
    # Escalation layer ALWAYS moves forward regardless of answer quality.
    # Strategy is chosen by combining layer position with weakest signal.
    next_layer = min(5, escalation_layer)  # current layer determines this follow-up's type

    layer_info = ESCALATION_LAYER_MAP[next_layer]

    # Within the layer, pick the best-fitting strategy based on weakness signals
    if next_layer == 1:
        if "too_brief" in weaknesses or "weak_communication" in weaknesses:
            strategy = "Clarification"
        else:
            strategy = "Depth Probe"
        reasoning = f"Layer 1 (Clarification): Answer was {'too brief' if 'too_brief' in weaknesses else 'unclear in places'}."

    elif next_layer == 2:
        strategy = "Metric Justification"
        reasoning = "Layer 2 (Metrics): Pushing for quantifiable evidence — numbers, benchmarks, or success criteria."

    elif next_layer == 3:
        if "no_tradeoff" in weaknesses:
            strategy = "Tradeoff Challenge"
            reasoning = "Layer 3 (Tradeoff): No tradeoffs mentioned — forcing comparison with alternative approach."
        else:
            strategy = "Edge Case Scenario"
            reasoning = "Layer 3 (Tradeoff): Tradeoffs present — challenging with edge case to deepen analysis."

    elif next_layer == 4:
        if "strong_answer" in weaknesses:
            strategy = "Architecture Breakdown"
            reasoning = "Layer 4 (Scalability): Strong answer — forcing architectural decomposition under scale."
        else:
            strategy = "Scalability Challenge"
            reasoning = "Layer 4 (Scalability): Testing how their solution holds up under 10x load."

    elif next_layer == 5:
        if has_failure:
            strategy = "Security Consideration"
            reasoning = "Layer 5 (Failure): Candidate mentioned failures — pivoting to security implications."
        else:
            strategy = "Failure Simulation"
            reasoning = "Layer 5 (Failure): Maximum pressure — simulating a production incident."

    else:
        strategy = "Depth Probe"
        reasoning = "Default: probing for deeper explanation."

    depth_score = min(10.0, max(0.0, avg_score))

    return {
        "weaknesses": weaknesses,
        "strategy": strategy,
        "next_layer": next_layer,
        "reasoning": reasoning,
        "depth_score": depth_score,
        "layer_name": layer_info["name"],
        "cognitive_pressure": layer_info["cognitive_pressure"],
    }


def generate_adaptive_followup_v2(
    question: str, answer: str, strategy: str,
    escalation_layer: int, role: str, domain: str,
    difficulty: str = "Hard"
) -> str:
    """
    UPGRADED adaptive follow-up generator.

    Key improvements over v1:
    - Uses ESCALATION_LAYER_MAP for precise per-layer instructions
    - Injects domain authority block into follow-up prompts
    - Adds cognitive pressure signal matching the difficulty
    - Hard mode adds explicit pressure framing ("In a live production system...")
    """
    from llm_manager import call_llm

    layer_info = ESCALATION_LAYER_MAP.get(escalation_layer, ESCALATION_LAYER_MAP[3])
    domain_block = build_domain_authority_block(domain, role)

    # Text-interview-optimised strategy instructions.
    # Each follow-up must be answerable in 4-6 paragraphs. No whiteboard scope.
    strategy_instructions = {
        "Clarification": (
            "Ask the candidate to clarify ONE specific statement that was ambiguous or vague. "
            "Reference the exact phrase or claim from their answer."
        ),
        "Depth Probe": (
            "Ask them to explain HOW one specific part of their answer works internally. "
            "Target the mechanism behind ONE claim they made — not the whole answer."
        ),
        "Metric Justification": (
            "Ask them to support ONE specific claim in their answer with concrete numbers or signals. "
            "e.g. 'You mentioned this approach is faster — what metric would you use to validate that, "
            "and what threshold would you consider acceptable?'"
        ),
        "Tradeoff Challenge": (
            "Present ONE specific alternative to the approach they described and ask them to compare: "
            "what does their approach do better, and what does it sacrifice? "
            "Keep the comparison to these two options only — not a full landscape review."
        ),
        "Edge Case Scenario": (
            "Describe ONE specific edge case or unusual input condition their solution might not handle gracefully. "
            "Ask: how would they detect it at runtime, and how would they handle it without breaking normal flow?"
        ),
        "Scalability Challenge": (
            "Inject a single scale constraint — for example, 10x the current load — and ask which ONE component "
            "in their described approach would be the first to break, and what they'd change first. "
            "This is NOT a full re-architecture question."
        ),
        "Constraint Injection": (
            "Add one realistic constraint they didn't mention — e.g. a latency SLA, a budget cap, "
            "or a dependency on a legacy system — and ask how they'd adapt their described approach. "
            "One constraint. One adaptation. Text-answerable."
        ),
        "Failure Simulation": (
            "Describe ONE specific failure mode relevant to what they described — "
            "e.g. the component they mentioned starts returning errors under normal load. "
            "Ask: what are the first three things they'd check, and what would a successful mitigation look like?"
        ),
        "Security Consideration": (
            "Ask about ONE specific security risk in the approach they described. "
            "e.g. 'You mentioned storing X — what's the highest-risk way that could be exploited, "
            "and what's the simplest effective mitigation?'"
        ),
        "Architecture Breakdown": (
            "Ask them to walk through ONE specific data flow or interaction in their described approach — "
            "not the full architecture. e.g. 'Walk me through exactly what happens from the point "
            "the request hits [component] to when the response is returned.'"
        ),
        "Alternative Design Comparison": (
            "Ask them to propose ONE alternative approach to what they described — "
            "just one, not a full survey — and compare the two specifically on reliability and complexity. "
            "Text-answerable: 4-6 paragraphs."
        ),
    }

    strategy_instruction = strategy_instructions.get(strategy, "Ask a deeper technical question that increases cognitive load.")

    # Hard mode: add cognitive pressure but keep it text-answerable.
    # We do NOT ask for full system redesigns. One axis. One focused challenge.
    pressure_framing = ""
    if difficulty == "Hard":
        pressure_framing = f"""
COGNITIVE PRESSURE LEVEL: {layer_info["cognitive_pressure"]}
TEXT-SCOPE ENFORCEMENT: This follow-up must be answerable in 4-6 paragraphs.
Focus on ONE specific gap from their answer — not a full redesign.
Frame with a concrete, realistic scenario (not "millions of users" at scale).
Examples of valid Hard follow-up framings:
  - "You mentioned X — what happens when Y occurs? Walk me through your response."
  - "How would you validate that your approach actually solves the problem? What metric matters most?"
  - "What's the one failure mode in your described approach that you'd lose sleep over?"
Do NOT ask them to redesign the whole system or address every possible failure.
"""

    prompt = f"""You are a senior technical interviewer for {role} in {domain}.

{domain_block}

The candidate just answered this question:
ORIGINAL QUESTION: {question}
CANDIDATE ANSWER: {answer[:600]}

ESCALATION LAYER: {escalation_layer}/5 — {layer_info["name"]}
Layer objective: {layer_info["instruction"]}

STRATEGY: {strategy}
Strategy instruction: {strategy_instruction}
{pressure_framing}

Generate EXACTLY ONE follow-up question. It MUST:
1. Reference ONE specific thing from their answer (name or paraphrase it directly)
2. Target ONE gap, ONE failure mode, ONE metric, or ONE decision point — not everything at once
3. Be answerable in 4-6 paragraphs of text — no whiteboard, no full system redesign
4. Be harder and more targeted than the original question
5. Be 1-3 sentences — concise, precise, not a multi-part essay prompt
6. NOT contain instructions like "design a complete system", "walk through every layer",
   or "handle all possible failure modes"

Output ONLY the follow-up question. No numbering, labels, or explanations.

Follow-up question:"""

    try:
        return call_llm(prompt, session=st.session_state).strip()
    except Exception:
        first_words = " ".join(answer.split()[:6]) if answer and answer.split() else "your described approach"
        return f'You mentioned "{first_words}..." — what is the ONE failure mode in that approach you would be most concerned about, and how would you detect it before it caused user impact?'


# =============================================================================
# DOMAIN-AWARE QUESTION GENERATORS (upgraded wrappers)
# =============================================================================

def generate_resume_based_questions_domain_aware(
    resume_context: dict, role: str, domain: str,
    difficulty: str, num_questions: int = 3, weakness_bias: str = "balanced",
    interview_type: str = "technical"
) -> list:
    """
    Drop-in replacement for generate_resume_based_questions_enhanced.
    Applies Domain Authority Layer + Structured Difficulty Enforcement.
    Supports interview_type ("technical" or "behavioral") for distinct question styles.
    """
    import random
    from llm_manager import call_llm

    # FIX 1: Apply domain filter to resume context
    filtered_context = filter_resume_for_domain(resume_context, domain)

    skills = filtered_context.get("skills", [])
    projects = filtered_context.get("projects", [])
    experience = filtered_context.get("experience", [])
    technologies = filtered_context.get("technologies", [])

    # FIX 1: Domain authority block
    domain_block = build_domain_authority_block(domain, role)

    # FIX 2: Structured difficulty enforcement
    difficulty_block = get_difficulty_instruction_block(difficulty)

    # Interview type block — drives Technical vs Behavioral question framing
    interview_type_block = (
        "⚙️ This is a TECHNICAL interview. Focus on technical depth, implementation details, tradeoffs, and reasoning."
        if interview_type.lower() == "technical"
        else "💬 This is a BEHAVIORAL interview. Focus on past experiences, teamwork, challenges, leadership, decision-making, and communication."
    )

    # Topic variation hint — reduces repetition across resume uploads
    variation_hint = random.choice([
        "Focus more on algorithms and data structures relevant to this resume.",
        "Include one question about troubleshooting or debugging a real issue.",
        "Add one question about collaboration or decision-making under pressure.",
        "Include one scenario-based question referencing tools or skills from the resume.",
        "Add one reflective question about learning or adapting to new technologies.",
    ])

    # Weakness bias instruction
    bias_map = {
        "technical depth": "Prioritize questions that expose gaps in technical depth — ask about internals, edge cases, and implementation specifics.",
        "explanation clarity": "Prioritize questions that require the candidate to explain complex concepts step-by-step.",
        "answer precision": "Prioritize questions that require very specific, targeted answers directly tied to their resume.",
        "balanced": "",
    }
    bias_instruction = bias_map.get(weakness_bias, "")

    prompt = f"""You are a senior technical interviewer.

{interview_type_block}

{domain_block}

{difficulty_block}

RESUME CONTEXT (filtered for domain relevance):
- Skills: {', '.join(skills[:5]) if skills else 'None relevant to ' + domain}
- Projects: {', '.join(projects[:3]) if projects else 'None specified'}
- Experience: {', '.join(experience[:3]) if experience else 'None specified'}
- Technologies: {', '.join(technologies[:5]) if technologies else 'None relevant to ' + domain}

{bias_instruction}

Generate EXACTLY {num_questions} interview questions. Each question MUST:
1. Be about {domain} — not the candidate's previous domain if it differs
2. Reference their resume only if resume content is relevant to {domain}
3. Match the difficulty type specified above (structural enforcement, not just tone)
4. Be a single, clear question (1-2 sentences)
5. Reflect the interview type above — technical questions probe implementation/tradeoffs; behavioral questions probe experience/judgment

Output ONLY the questions, one per line, no numbering or prefixes.

{variation_hint}

Questions:"""

    try:
        response = call_llm(prompt, session=st.session_state)
        raw = [q.strip() for q in response.split("\n") if q.strip()]
        cleaned = []
        for q in raw:
            q = re.sub(r'^[\d\)\.\-•\*]+\s*', '', q).strip()
            if len(q) > 15:
                cleaned.append(q)
            if len(cleaned) >= num_questions:
                break

        # Fallback fill
        while len(cleaned) < num_questions:
            contract = DIFFICULTY_CONTRACTS.get(difficulty, DIFFICULTY_CONTRACTS["Medium"])
            diff_label = contract["label"]
            fallbacks = contract.get("fallback_questions", [])
            if fallbacks:
                import random
                fb = random.choice(fallbacks).replace("{domain}", domain).replace("{role}", role)
                cleaned.append(fb)
            else:
                cleaned.append(
                    f"[{diff_label}] Describe a specific challenge you faced with {domain} and how you resolved it."
                )
        return cleaned[:num_questions]

    except Exception:
        return [f"Explain a core {domain} concept you've worked with recently."] * num_questions


def generate_domain_questions_with_llm(
    domain: str, role: str, interview_type: str,
    num_questions: int, difficulty: str = "Medium"
) -> list:
    """
    Domain-authority-enforced replacement for generate_interview_questions_with_llm.
    Ensures generic questions also respect the selected domain and difficulty contract.
    """
    from llm_manager import call_llm

    domain_block = build_domain_authority_block(domain, role)
    difficulty_block = get_difficulty_instruction_block(difficulty)

    prompt = f"""You are an expert interviewer at a top-tier tech company.

{domain_block}

{difficulty_block}

Generate EXACTLY {num_questions} unique {interview_type} interview questions for a {role} candidate.

RULES:
- Every question MUST be about {domain} — no exceptions
- Match the exact difficulty type defined above (not just tone)
- Avoid duplicates and generic filler questions
- Keep each question concise: 1-2 sentences maximum
- Output ONLY the questions, one per line
- NO numbering, bullets, prefixes, or explanatory text

Generate {num_questions} questions now:"""

    try:
        response = call_llm(prompt, session=st.session_state)
        raw = [q.strip() for q in response.split('\n') if q.strip()]
        cleaned = []
        for q in raw:
            clean_q = re.sub(r'^[\d\)\.\-•\*]+\s*', '', q).strip()
            clean_q = re.sub(r'^Question\s*\d*\s*:?\s*', '', clean_q, flags=re.IGNORECASE).strip()
            if clean_q and len(clean_q) > 15:
                cleaned.append(clean_q)
            if len(cleaned) >= num_questions:
                break

        # Fallback with difficulty-appropriate templates
        while len(cleaned) < num_questions:
            contract = DIFFICULTY_CONTRACTS.get(difficulty, DIFFICULTY_CONTRACTS["Medium"])
            fallbacks = contract.get("fallback_questions", [])
            import random
            if fallbacks:
                fb = random.choice(fallbacks).replace("{domain}", domain).replace("{role}", role)
                cleaned.append(fb)
            else:
                templates = contract.get("question_templates", [])
                template = random.choice(templates) if templates else "Explain a core {domain} concept."
                cleaned.append(template.replace("{domain}", domain).replace("{role}", role).replace("[concept]", domain))

        return cleaned[:num_questions]

    except Exception:
        # Text-answerable fallback questions per difficulty
        import random
        contract = DIFFICULTY_CONTRACTS.get(difficulty, DIFFICULTY_CONTRACTS["Medium"])
        fallbacks = contract.get("fallback_questions", [])
        if fallbacks:
            return [
                random.choice(fallbacks).replace("{domain}", domain).replace("{role}", role)
                for _ in range(num_questions)
            ]
        if difficulty == "Easy":
            return [f"What is {domain} and why is it important? Give a real-world example." for _ in range(num_questions)]
        elif difficulty == "Medium":
            return [f"Describe a specific implementation decision you made in {domain} and why you made it." for _ in range(num_questions)]
        else:
            return [
                f"Describe the most significant tradeoff you've encountered in {domain}. "
                f"What were the two options and what drove your final decision?"
                for _ in range(num_questions)
            ]


# ======================================================
# RESUME TEXT EXTRACTION (pdfplumber + OCR fallback)
# ======================================================
def extract_resume_text_from_pdf(pdf_file):
    """
    Robust resume extraction:
    - pdfplumber for text-based & two-column resumes
    - OCR fallback for scanned/image resumes
    """

    text = ""

    # ---------- PRIMARY: pdfplumber ----------
    try:
        import pdfplumber
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        text = text.strip()
    except Exception:
        text = ""

    # ---------- FALLBACK: OCR ----------
    if len(text.split()) < 120:
        try:
            from pdf2image import convert_from_bytes
            import pytesseract

            images = convert_from_bytes(pdf_file.getvalue())
            ocr_text = ""
            for img in images:
                ocr_text += pytesseract.image_to_string(img)

            if len(ocr_text.split()) > len(text.split()):
                text = ocr_text.strip()
        except Exception:
            st.warning("OCR fallback failed. Resume may be image-heavy.")

    # ---------- FINAL VALIDATION ----------
    if not text or len(text.split()) < 80:
        st.warning("Resume text extraction was weak. Some questions may be generic.")
        return None

    return text


# ======================================================
# RESUME ANALYSIS USING LLM (IMPROVED PROMPT)
# ======================================================
def analyze_resume_with_llm(resume_text):
    """
    Analyze resume using LLM to extract INTERVIEW-RELEVANT structured information
    """

    prompt = f"""
You are a senior technical interviewer and resume screening expert.

Analyze the resume below and extract ONLY the most interview-relevant information.
Focus on technical depth, real-world work, and ownership.
IGNORE generic soft skills unless strongly implied by technical work.

RESUME TEXT:
{resume_text}

Return ONLY a valid JSON object with this exact structure:

{{
  "skills": [
    "Core technical skill clearly demonstrated in projects or experience"
  ],
  "projects": [
    "Project name – what was built, tech used, and key technical challenge solved"
  ],
  "experience": [
    "Role at company – main technical responsibility and impact"
  ],
  "technologies": [
    "Primary technologies actually used (not buzzwords)"
  ]
}}

STRICT RULES:
- Prefer HARD technical skills over soft skills
- Extract ONLY skills clearly demonstrated
- Rank items by importance (most interview-worthy first)
- Avoid generic terms like 'problem solving', 'communication'
- Projects MUST mention tech used
- Experience MUST show ownership or responsibility
- Extract:
  - 4–6 skills
  - 2–4 projects
  - 2–4 experience entries
  - 4–6 technologies
- Keep entries concise but specific
- Output ONLY JSON (no markdown, no explanations)

JSON:
"""

    try:
        response = call_llm(prompt, session=st.session_state).strip()

        # Clean markdown if present
        if response.startswith("```"):
            response = response.split("```")[1]
            if response.lower().startswith("json"):
                response = response[4:]
            response = response.strip()

        resume_data = json.loads(response)

        return {
            "skills": resume_data.get("skills", [])[:6],
            "projects": resume_data.get("projects", [])[:4],
            "experience": resume_data.get("experience", [])[:4],
            "technologies": resume_data.get("technologies", [])[:6]
        }

    except Exception:
        st.warning("Resume analysis failed. Using fallback data.")
        return {
            "skills": ["Basic Programming Knowledge"],
            "projects": ["Personal Technical Project"],
            "experience": ["General Technical Experience"],
            "technologies": ["General Tech Stack"]
        }


# ======================================================
# ⚡ MERGED STARTUP FUNCTION — 3 calls → 1
# ======================================================
def analyze_resume_and_generate_questions(
    resume_text: str,
    role: str,
    domain: str,
    difficulty: str,
    interview_type: str,
    num_resume_qs: int = 2,
    num_generic_qs: int = 4,
    weakness_bias: str = "balanced",
) -> dict:
    """
    Single LLM call that replaces THREE separate startup calls:
      1. analyze_resume_with_llm()           — resume context extraction
      2. generate_resume_based_questions_domain_aware() — resume-based questions
      3. generate_domain_questions_with_llm()           — generic domain questions

    Cost: 1 API call (vs 3 in tab4__2_).
    Quality: Full difficulty contract + domain authority block injected (same depth as tab4__2_).
      - Uses get_difficulty_instruction_block() with structural templates, forbidden patterns,
        cognitive load rules, and scoring calibration.
      - Uses build_domain_authority_block() for domain-specific keyword enforcement.

    Returns dict with keys:
      resume_context   : dict  (skills, projects, experience, technologies)
      resume_questions : list  (num_resume_qs items)
      generic_questions: list  (num_generic_qs items)
    """
    from llm_manager import call_llm
    import json, re, random
    import streamlit as st

    # ── Full difficulty + domain blocks (same as tab4__2_ separate calls) ──────
    diff_block   = get_difficulty_instruction_block(difficulty)
    domain_block = build_domain_authority_block(domain, role)

    # ── Filter resume context for domain relevance (mirrors tab4__2_) ──────────
    filtered_context = filter_resume_for_domain({
        "skills": [],
        "projects": [],
        "experience": [],
        "technologies": [],
    }, domain)  # placeholder — real context parsed after Task 1; used for resume_qs below

    interview_type_block = (
        "⚙️ This is a TECHNICAL interview. Focus on technical depth, implementation details, tradeoffs, and reasoning."
        if interview_type.lower() == "technical"
        else "💬 This is a BEHAVIORAL interview. Focus on past experiences, teamwork, challenges, leadership, decision-making, and communication."
        if interview_type.lower() == "behavioral"
        else "🔀 This is a MIXED interview. Blend technical depth with behavioral judgment."
    )

    bias_map = {
        "technical depth":     "Prioritize questions that expose gaps in technical depth — ask about internals, edge cases, and implementation specifics.",
        "explanation clarity": "Prioritize questions that require the candidate to explain complex concepts step-by-step.",
        "answer precision":    "Prioritize questions that require very specific, targeted answers directly tied to their resume.",
        "balanced":            "",
    }
    bias_note = bias_map.get(weakness_bias, "")

    variation_hint = random.choice([
        "For generic questions, include one scenario referencing real-world constraints.",
        "For generic questions, add one question about debugging or diagnosing a failure.",
        "For generic questions, include one question about a specific tradeoff in this domain.",
        "For generic questions, add one question about collaboration or technical decision-making.",
        "For generic questions, add one reflective question about learning or adapting to new technologies.",
    ])

    prompt = f"""You are a senior technical interviewer. Complete THREE tasks and return ONLY the JSON shown below.

{interview_type_block}

{domain_block}

{diff_block}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TASK 1 — RESUME ANALYSIS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Analyze the resume for INTERVIEW-RELEVANT information only.
Extract: core technical skills (4-6), notable projects (2-4), experience entries (2-4), primary technologies (4-6).
Ignore generic soft skills. Prefer hard technical evidence.

RESUME TEXT (first 3000 chars):
{resume_text[:3000]}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TASK 2 — {num_resume_qs} RESUME-BASED QUESTIONS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Generate EXACTLY {num_resume_qs} interview questions grounded in the resume above.
Target role: {role} | Domain: {domain} | Difficulty: {difficulty}

RULES:
- Every question MUST reference a skill, project, or technology from the resume
- Every question MUST be about {domain} — not the candidate's previous domain if it differs
- Match the DIFFICULTY CONTRACT and structural templates above exactly
- Reflect the interview type above — technical questions probe implementation/tradeoffs; behavioral questions probe experience/judgment
{f"- {bias_note}" if bias_note else ""}
- Each question: 1-2 sentences, self-contained

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TASK 3 — {num_generic_qs} GENERIC DOMAIN QUESTIONS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Generate EXACTLY {num_generic_qs} domain-general interview questions.
Domain: {domain} | Role: {role} | Difficulty: {difficulty}

RULES:
- Every question MUST be about {domain} — no exceptions
- Match the DIFFICULTY CONTRACT and structural templates above exactly
- Reflect the interview type above
- No overlap with resume questions; no numbering or prefixes
- Each question: 1-2 sentences
- {variation_hint}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
OUTPUT FORMAT — return ONLY this JSON, no markdown, no extra text:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{{
  "resume_context": {{
    "skills": ["skill1", "skill2"],
    "projects": ["project1 – tech used – key challenge"],
    "experience": ["Role at Company – main responsibility"],
    "technologies": ["tech1", "tech2"]
  }},
  "resume_questions": [
    "Question referencing resume content 1",
    "Question referencing resume content 2"
  ],
  "generic_questions": [
    "Domain question 1",
    "Domain question 2",
    "Domain question 3",
    "Domain question 4"
  ]
}}
"""

    # ── Fallback data ──────────────────────────────────────────────────────────
    fallback_context = {
        "skills": ["Technical Skills"],
        "projects": ["Personal Technical Project"],
        "experience": ["General Technical Experience"],
        "technologies": ["General Tech Stack"],
    }
    fallback_resume_qs = [
        f"Walk me through your most technically challenging project and the decisions you made.",
        f"How does your experience prepare you for the {role} role?",
    ][:num_resume_qs]
    fallback_generic_qs = [
        f"Explain a core {domain} concept and give a real-world example.",
        f"Describe a specific implementation decision in {domain} and your reasoning.",
        f"What is the most important tradeoff to understand in {domain}?",
        f"How would you debug an unexpected failure in a {domain} system?",
    ][:num_generic_qs]

    try:
        raw = call_llm(prompt, session=st.session_state).strip()

        # Strip markdown fences if present
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.lower().startswith("json"):
                raw = raw[4:]
            raw = raw.strip()

        data = json.loads(raw)

        # ── Parse resume context ───────────────────────────────────────────────
        rc = data.get("resume_context", {})
        resume_context = {
            "skills":       rc.get("skills", fallback_context["skills"])[:6],
            "projects":     rc.get("projects", fallback_context["projects"])[:4],
            "experience":   rc.get("experience", fallback_context["experience"])[:4],
            "technologies": rc.get("technologies", fallback_context["technologies"])[:6],
        }

        # ── Parse questions ────────────────────────────────────────────────────
        def _clean_qs(qs, needed, fallback):
            cleaned = []
            for q in qs:
                q = str(q).strip()
                q = re.sub(r'^[\d\)\.\-•\*]+\s*', '', q).strip()
                if len(q) > 15:
                    cleaned.append(q)
                if len(cleaned) >= needed:
                    break
            # Pad with fallbacks if LLM returned too few
            while len(cleaned) < needed:
                cleaned.append(fallback[len(cleaned) % len(fallback)])
            return cleaned[:needed]

        resume_questions  = _clean_qs(data.get("resume_questions", []),  num_resume_qs,  fallback_resume_qs)
        generic_questions = _clean_qs(data.get("generic_questions", []), num_generic_qs, fallback_generic_qs)

        return {
            "resume_context":    resume_context,
            "resume_questions":  resume_questions,
            "generic_questions": generic_questions,
        }

    except Exception:
        return {
            "resume_context":    fallback_context,
            "resume_questions":  fallback_resume_qs,
            "generic_questions": fallback_generic_qs,
        }


# ======================================================
# RESUME-BASED QUESTION GENERATION
# ======================================================
def generate_resume_based_questions(resume_context, role, domain, difficulty, num_questions=3):
    """
    Generate interview questions strictly based on resume context
    """

    skills = resume_context.get("skills", [])
    projects = resume_context.get("projects", [])
    experience = resume_context.get("experience", [])
    technologies = resume_context.get("technologies", [])

    prompt = f"""
You are a technical interviewer.

Generate EXACTLY {num_questions} interview questions based ONLY on the candidate's resume.

RESUME CONTEXT:
- Skills: {', '.join(skills[:4])}
- Projects: {', '.join(projects[:2])}
- Experience: {', '.join(experience[:2])}
- Technologies: {', '.join(technologies[:4])}

Target Role: {role}
Domain: {domain}
Difficulty: {difficulty}

RULES:
- Every question MUST reference resume content
- Ask like a real interviewer
- Difficulty:
  - Easy: explanation & fundamentals
  - Medium: scenarios & decisions
  - Hard: deep technical trade-offs or design
- Output ONLY questions
- One question per line
- No numbering, no prefixes

Generate now:
"""

    try:
        response = call_llm(prompt, session=st.session_state)
        raw_questions = [q.strip() for q in response.split("\n") if q.strip()]

        cleaned_questions = []
        for q in raw_questions:
            q = re.sub(r'^[\d\)\.\-•\*]+\s*', '', q).strip()
            if len(q) > 15:
                cleaned_questions.append(q)
            if len(cleaned_questions) >= num_questions:
                break

        while len(cleaned_questions) < num_questions:
            cleaned_questions.append(
                f"Explain your most significant project and the technical decisions you made."
            )

        return cleaned_questions[:num_questions]

    except Exception:
        return [
            "Walk us through your most technically challenging project.",
            "What design or implementation decisions did you personally make?",
            "How does your experience prepare you for this role?"
        ]


# ======================================================
# RESUME SCANNING ANIMATION
# ======================================================
def show_resume_scanning_animation():
    """Animated resume scanning UI"""

    status = st.empty()
    progress = st.empty()

    steps = [
        ("📄 Reading resume...", 0.2),
        ("🔍 Extracting key skills...", 0.4),
        ("📊 Evaluating experience...", 0.6),
        ("🧠 Understanding projects...", 0.8),
        ("🎯 Preparing interview questions...", 1.0),
    ]

    for text, value in steps:
        status.markdown(
            f"<h4 style='text-align:center;color:#38bdf8;font-family:-apple-system,BlinkMacSystemFont,\"SF Pro Display\",sans-serif;font-weight:600;letter-spacing:-0.02em;'>{text}</h4>",
            unsafe_allow_html=True
        )
        progress.progress(value)
        time.sleep(0.15)  # FIX: reduced from 0.6s to 0.15s — animation delay, not rerun gate

    status.empty()
    progress.empty()



# =============================================================================
# PART 1-9: UPGRADED ENGINE FUNCTIONS
# =============================================================================

HARD_FOLLOWUP_STRATEGIES = [
    "Depth Probe",
    "Tradeoff Challenge",
    "Edge Case Scenario",
    "Scalability Challenge",
    "Constraint Injection",
    "Failure Simulation",
    "Security Consideration",
    "Architecture Breakdown",
    "Metric Justification",
    "Alternative Design Comparison",
]

DIFFICULTY_MULTIPLIERS = {"Easy": 1.0, "Medium": 1.1, "Hard": 1.25}


def analyze_answer_weaknesses(answer_text: str, scoring: dict) -> dict:
    """
    UPGRADED (Fix 3): Delegates to analyze_answer_weaknesses_smart.
    Backward-compatible — existing callers continue to work.
    Escalation layer is read from session_state to drive layer progression.
    """
    try:
        current_layer = st.session_state.get("escalation_layer", 1)
    except Exception:
        current_layer = 1
    result = analyze_answer_weaknesses_smart(answer_text, scoring, escalation_layer=current_layer)
    # Backward-compat keys
    result["follow_up_count"] = getattr(st.session_state, "follow_up_count", 0) if hasattr(st, "session_state") else 0
    return result


def generate_adaptive_followup(question: str, answer: str, strategy: str, escalation_layer: int, role: str, domain: str) -> str:
    """
    UPGRADED (Fix 3): Delegates to generate_adaptive_followup_v2.
    Backward-compatible wrapper — all existing callers work unchanged.
    """
    diff = getattr(st.session_state, "interview_difficulty", "Hard") if hasattr(st, "session_state") else "Hard"
    return generate_adaptive_followup_v2(
        question=question, answer=answer, strategy=strategy,
        escalation_layer=escalation_layer, role=role, domain=domain, difficulty=diff
    )


@st.cache_data(ttl=60)
def get_user_weakness_history(username: str) -> dict:
    """
    PART 5: Weakness Memory Engine.
    Query ALL past interviews for total count and detect recurring weak skill.
    Uses all interviews for avg score computation and shows true total count.
    Returns dict with weakest_skill and bias recommendation.

    @st.cache_data(ttl=60): result is cached per username for 60 seconds so
    every dropdown change on the setup screen does NOT hit Supabase again.
    """
    import pandas as pd
    try:
        conn = _get_live_conn()
        # Fetch ALL interviews (no LIMIT) so count and averages reflect full history
        df = pd.read_sql_query(
            "SELECT knowledge_avg, communication_avg, relevance_avg FROM interview_results WHERE username=%s ORDER BY id DESC",
            conn, params=(username,)
        )

        if df.empty or len(df) < 1:
            return {"weakest_skill": None, "bias": "balanced"}

        total_count = len(df)

        avgs = {
            "knowledge": df["knowledge_avg"].mean(),
            "communication": df["communication_avg"].mean(),
            "relevance": df["relevance_avg"].mean(),
        }
        weakest = min(avgs, key=avgs.get)
        bias_map = {
            "knowledge": "technical depth",
            "communication": "explanation clarity",
            "relevance": "answer precision",
        }
        return {"weakest_skill": weakest, "bias": bias_map.get(weakest, "balanced"), "averages": avgs, "interview_count": total_count}
    except Exception:
        return {"weakest_skill": None, "bias": "balanced"}


def compute_weighted_score(raw_avg: float, difficulty: str) -> float:
    """PART 3: Apply difficulty multiplier to raw average score."""
    multiplier = DIFFICULTY_MULTIPLIERS.get(difficulty, 1.0)
    return round(min(10.0, raw_avg * multiplier), 2)


def compute_trend_slope(scores: list) -> float:
    """PART 6: Compute linear regression slope over score list."""
    n = len(scores)
    if n < 2:
        return 0.0
    try:
        x = list(range(n))
        x_mean = sum(x) / n
        y_mean = sum(scores) / n
        numerator = sum((x[i] - x_mean) * (scores[i] - y_mean) for i in range(n))
        denominator = sum((x[i] - x_mean) ** 2 for i in range(n))
        return numerator / denominator if denominator != 0 else 0.0
    except Exception:
        return 0.0


def classify_behavior(avg_duration_mins, score_std, hard_delta) -> str:
    """PART 6: Behavioral classification based on performance patterns."""
    if avg_duration_mins is not None and avg_duration_mins < 8:
        return "⚡ Rushed"
    elif avg_duration_mins is not None and avg_duration_mins > 40:
        return "🤔 Overthinking"
    elif score_std < 0.8 and hard_delta is not None and hard_delta > -1:
        return "🎯 Adaptive Learner"
    else:
        return "⚖️ Balanced"


def generate_resume_based_questions_enhanced(resume_context: dict, role: str, domain: str, difficulty: str, num_questions: int = 3, weakness_bias: str = "balanced") -> list:
    """
    PART 7: Enhanced resume intelligence — asks architecture, decisions, tradeoffs, outcomes, and scale.
    """
    from llm_manager import call_llm

    skills = resume_context.get("skills", [])
    projects = resume_context.get("projects", [])
    experience = resume_context.get("experience", [])
    technologies = resume_context.get("technologies", [])

    bias_instruction = ""
    if weakness_bias == "explanation clarity":
        bias_instruction = "Focus on questions that require the candidate to explain complex concepts clearly, describe their reasoning process, and articulate decisions step-by-step."
    elif weakness_bias == "technical depth":
        bias_instruction = "Focus on questions that require deep technical knowledge, architecture reasoning, and internal mechanism understanding."
    elif weakness_bias == "answer precision":
        bias_instruction = "Focus on questions that require precise, targeted answers directly relevant to the role and their stated experience."

    # TEXT-INTERVIEW-OPTIMISED difficulty guidance for resume-based questions.
    # Hard: one focused challenge per question — NOT "design entire system + tradeoffs + scale".
    difficulty_map = {
        "Easy": (
            "Ask about concepts, tools, or technologies visible in their resume. "
            "Questions must be answerable by explaining WHAT something is and giving ONE example. "
            "No architecture, no scaling, no tradeoff decisions."
        ),
        "Medium": (
            "Ask about a specific decision, implementation approach, or challenge from their resume. "
            "Frame it as a small scenario with ONE constraint. "
            "The candidate should explain their approach and justify ONE key choice. "
            "No multi-layer design, no full-system architecture."
        ),
        "Hard": (
            "Ask about a focused technical challenge from their resume. Pick EXACTLY ONE axis: "
            "either (a) a tradeoff between two specific approaches they may have faced, "
            "or (b) a failure or edge case relevant to what they built, "
            "or (c) an optimisation under a specific constraint. "
            "Do NOT combine all three. The question must be answerable in 6-8 text paragraphs."
        ),
    }

    prompt = f"""You are a senior technical interviewer for {role} in {domain}.

Generate EXACTLY {num_questions} interview questions based on the candidate's actual resume.

RESUME CONTEXT:
- Skills: {', '.join(skills[:5])}
- Projects: {', '.join(projects[:3])}
- Experience: {', '.join(experience[:3])}
- Technologies: {', '.join(technologies[:5])}

DIFFICULTY GUIDANCE: {difficulty}
{difficulty_map.get(difficulty, '')}

{bias_instruction}

QUESTION FOCUS RULES:
- Each question must reference ONE specific item from the resume context above
- Each question must test exactly ONE of: concept understanding / decision reasoning / implementation approach / focused tradeoff / single failure mode
- Do NOT combine multiple challenge axes in one question
- Hard questions must be text-answerable (6-8 paragraphs) — no whiteboard system design
- One question per line, no numbering or prefixes

Generate {num_questions} questions:"""

    try:
        response = call_llm(prompt, session=st.session_state)
        raw = [q.strip() for q in response.split("\n") if q.strip()]
        cleaned = []
        for q in raw:
            q = re.sub(r'^[\d\)\.\-•\*]+\s*', '', q).strip()
            if len(q) > 15:
                cleaned.append(q)
            if len(cleaned) >= num_questions:
                break
        while len(cleaned) < num_questions:
            cleaned.append(f"Walk me through your most technically challenging project and the specific decisions you made.")
        return cleaned[:num_questions]
    except Exception:
        return [
            "Walk me through the architecture of your most complex project and the decisions you made.",
            "What tradeoffs did you face and how did you justify your technology choices?",
            "How would you scale your largest project to handle 10x the load?"
        ]


with tab4:
    # Inject CSS styles — Apple-style SaaS dark theme (matching tab1.py HIRELYZER design language)
    st.markdown("""
        <style>
        /* ═══════════════════════════════════════════════════════════════
           HIRELYZER — Premium Apple-Style Dark Theme (Tab 4)
           Font Stack: SF Pro Display → DM Sans → Segoe UI → sans-serif
           Design Language: Glassmorphism · Soft gradients · Refined motion
           ═══════════════════════════════════════════════════════════════ */

        @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&display=swap');

        :root {
            --t4-bg-primary:      #080c12;
            --t4-bg-secondary:    #0e1420;
            --t4-bg-tertiary:     #141c2b;
            --t4-surface-01:      rgba(255,255,255,0.04);
            --t4-surface-02:      rgba(255,255,255,0.07);
            --t4-surface-hover:   rgba(255,255,255,0.10);
            --t4-border-subtle:   rgba(255,255,255,0.07);
            --t4-border-accent:   rgba(99,179,237,0.30);
            --t4-accent-blue:     #4fa3e3;
            --t4-accent-cyan:     #38bdf8;
            --t4-accent-violet:   #818cf8;
            --t4-accent-emerald:  #34d399;
            --t4-accent-amber:    #fbbf24;
            --t4-accent-rose:     #fb7185;
            --t4-text-primary:    #f0f4f8;
            --t4-text-secondary:  #94a3b8;
            --t4-text-muted:      #4a5568;
            --t4-radius-sm:       8px;
            --t4-radius-md:       14px;
            --t4-radius-lg:       20px;
            --t4-radius-xl:       28px;
            --t4-shadow-glow:     0 0 30px rgba(79,163,227,0.15);
            --t4-shadow-card:     0 8px 40px rgba(0,0,0,0.45), 0 1px 0 rgba(255,255,255,0.06) inset;
            --t4-font:            -apple-system, BlinkMacSystemFont, "SF Pro Display", "DM Sans", "Segoe UI", Roboto, sans-serif;
            --t4-ease-fast:       0.18s cubic-bezier(0.4,0,0.2,1);
            --t4-ease-base:       0.28s cubic-bezier(0.4,0,0.2,1);
            --t4-ease-slow:       0.45s cubic-bezier(0.4,0,0.2,1);
        }

        /* ── Animations ── */
        @keyframes t4-fadeSlideUp  { from { opacity:0; transform:translateY(14px); } to { opacity:1; transform:translateY(0); } }
        @keyframes t4-shimmer      { 0% { transform:translateX(-100%) skewX(-12deg); } 100% { transform:translateX(220%) skewX(-12deg); } }
        @keyframes t4-pulseGlow    { 0%,100% { box-shadow: var(--t4-shadow-card); } 50% { box-shadow: var(--t4-shadow-card), var(--t4-shadow-glow); } }
        @keyframes t4-gradientFlow { 0%,100% { background-position:0% 50%; } 50% { background-position:100% 50%; } }
        @keyframes t4-subtlePulse  { 0%,100% { opacity:1; } 50% { opacity:0.82; } }

        /* ── Header Box ── */
        .header-box {
            background: linear-gradient(160deg, rgba(14,20,32,0.97) 0%, rgba(8,12,18,0.99) 100%);
            backdrop-filter: blur(32px) saturate(160%);
            -webkit-backdrop-filter: blur(32px) saturate(160%);
            border: 1px solid rgba(99,179,237,0.20);
            border-radius: var(--t4-radius-xl);
            padding: 32px 28px;
            text-align: center;
            margin-bottom: 32px;
            box-shadow: var(--t4-shadow-card), 0 0 60px rgba(79,163,227,0.07);
            position: relative;
            overflow: hidden;
            animation: t4-fadeSlideUp 0.65s cubic-bezier(0.22,1,0.36,1) forwards;
        }
        .header-box::after {
            content: '';
            position: absolute;
            top: 0; left: -100%;
            width: 60%; height: 100%;
            background: linear-gradient(90deg, transparent, rgba(79,163,227,0.06), transparent);
            animation: t4-shimmer 3.5s ease-in-out infinite;
        }
        .header-box h2 {
            font-family: var(--t4-font) !important;
            font-size: 1.85rem !important;
            font-weight: 700 !important;
            color: var(--t4-text-primary) !important;
            letter-spacing: -0.03em !important;
            margin: 0 !important;
            text-shadow: none !important;
        }

        /* ── Glow Header ── */
        .glow-header {
            font-family: var(--t4-font);
            font-size: 1.25rem;
            text-align: center;
            color: var(--t4-accent-cyan);
            font-weight: 600;
            letter-spacing: -0.02em;
            margin: 20px 0 12px 0;
            animation: t4-subtlePulse 3.5s ease-in-out infinite;
        }

        /* ── Learning Path Container ── */
        .learning-path-container {
            text-align: center;
            margin: 24px 0 18px 0;
            padding: 14px 20px;
            background: var(--t4-surface-01);
            backdrop-filter: blur(16px);
            -webkit-backdrop-filter: blur(16px);
            border-radius: var(--t4-radius-md);
            border: 1px solid var(--t4-border-subtle);
            transition: border-color var(--t4-ease-base);
        }
        .learning-path-container:hover {
            border-color: var(--t4-border-accent);
        }
        .learning-path-text {
            font-family: var(--t4-font);
            color: var(--t4-text-secondary);
            font-weight: 600;
            font-size: 0.875rem;
            text-transform: uppercase;
            letter-spacing: 0.06em;
        }

        /* ── Card ── */
        .card {
            background: var(--t4-surface-01);
            backdrop-filter: blur(24px) saturate(180%);
            -webkit-backdrop-filter: blur(24px) saturate(180%);
            border: 1px solid var(--t4-border-subtle);
            border-radius: var(--t4-radius-lg);
            padding: 20px 24px;
            margin: 10px 0;
            position: relative;
            overflow: hidden;
            transition: transform var(--t4-ease-base), box-shadow var(--t4-ease-base), border-color var(--t4-ease-base);
            box-shadow: var(--t4-shadow-card);
            animation: t4-fadeSlideUp 0.5s ease forwards;
        }
        .card::before {
            content: '';
            position: absolute;
            inset: 0;
            background: linear-gradient(135deg, rgba(255,255,255,0.05) 0%, transparent 60%);
            pointer-events: none;
            border-radius: inherit;
        }
        .card::after {
            content: '';
            position: absolute;
            top: 0; left: -100%;
            width: 50%; height: 100%;
            background: linear-gradient(90deg, transparent, rgba(79,163,227,0.05), transparent);
            transition: left 0.6s ease;
        }
        .card:hover {
            transform: translateY(-4px);
            box-shadow: var(--t4-shadow-card), 0 0 50px rgba(79,163,227,0.10);
            border-color: var(--t4-border-accent);
        }
        .card:hover::after { left: 150%; }
        .card a {
            font-family: var(--t4-font);
            color: var(--t4-accent-cyan);
            font-weight: 600;
            font-size: 0.95rem;
            text-decoration: none;
            display: flex;
            align-items: center;
            gap: 8px;
            transition: all var(--t4-ease-fast);
        }
        .card a:hover {
            color: var(--t4-text-primary);
            text-decoration: none;
            transform: translateX(3px);
        }

        /* ── Course Tile ── */
        .course-tile {
            background: var(--t4-surface-01);
            backdrop-filter: blur(20px);
            -webkit-backdrop-filter: blur(20px);
            border: 1px solid var(--t4-border-subtle);
            border-radius: var(--t4-radius-lg);
            padding: 20px;
            margin: 12px 0;
            transition: all var(--t4-ease-base);
            position: relative;
            overflow: hidden;
            box-shadow: var(--t4-shadow-card);
        }
        .course-tile:hover {
            transform: translateY(-4px);
            border-color: var(--t4-border-accent);
            box-shadow: var(--t4-shadow-card), var(--t4-shadow-glow);
        }
        .course-title {
            font-family: var(--t4-font);
            color: var(--t4-accent-cyan);
            font-size: 1rem;
            font-weight: 600;
            margin-bottom: 8px;
            letter-spacing: -0.01em;
        }
        .course-description {
            font-family: var(--t4-font);
            color: var(--t4-text-secondary);
            font-size: 0.85rem;
            margin-bottom: 14px;
            line-height: 1.55;
        }

        /* ── Difficulty Badges ── */
        .difficulty-badge {
            display: inline-block;
            padding: 3px 10px;
            border-radius: 99px;
            font-size: 0.72rem;
            font-weight: 600;
            letter-spacing: 0.04em;
            text-transform: uppercase;
            margin-bottom: 12px;
            font-family: var(--t4-font);
        }
        .difficulty-beginner    { background: rgba(52,211,153,0.15); color: var(--t4-accent-emerald); border: 1px solid rgba(52,211,153,0.3); }
        .difficulty-intermediate{ background: rgba(251,191,36,0.12); color: var(--t4-accent-amber);   border: 1px solid rgba(251,191,36,0.28); }
        .difficulty-advanced    { background: rgba(251,113,133,0.12); color: var(--t4-accent-rose);   border: 1px solid rgba(251,113,133,0.28); }

        /* ── Course Link Button ── */
        .course-link-btn {
            background: linear-gradient(135deg, rgba(56,189,248,0.18) 0%, rgba(79,163,227,0.12) 100%);
            color: var(--t4-accent-cyan);
            border: 1px solid rgba(56,189,248,0.30);
            padding: 7px 16px;
            border-radius: var(--t4-radius-sm);
            text-decoration: none;
            font-weight: 600;
            font-size: 0.825rem;
            font-family: var(--t4-font);
            display: inline-block;
            transition: all var(--t4-ease-fast);
            backdrop-filter: blur(8px);
        }
        .course-link-btn:hover {
            background: linear-gradient(135deg, rgba(56,189,248,0.28) 0%, rgba(79,163,227,0.22) 100%);
            border-color: rgba(56,189,248,0.55);
            transform: translateY(-1px);
            text-decoration: none;
            color: #e0f6ff;
        }

        /* ── Quiz Card ── */
        .quiz-card {
            background: var(--t4-surface-01);
            backdrop-filter: blur(20px);
            -webkit-backdrop-filter: blur(20px);
            border: 1px solid var(--t4-border-subtle);
            border-radius: var(--t4-radius-lg);
            padding: 20px;
            margin: 14px 0;
            box-shadow: var(--t4-shadow-card);
            transition: all var(--t4-ease-base);
        }
        .quiz-card:hover {
            border-color: var(--t4-border-accent);
        }

        /* ── Badge Container ── */
        .badge-container {
            text-align: center;
            padding: 28px;
            background: var(--t4-surface-01);
            backdrop-filter: blur(24px) saturate(180%);
            -webkit-backdrop-filter: blur(24px) saturate(180%);
            border-radius: var(--t4-radius-lg);
            border: 1px solid var(--t4-border-subtle);
            margin: 18px 0;
            box-shadow: var(--t4-shadow-card);
            animation: t4-fadeSlideUp 0.5s ease forwards;
        }

        /* ── Score Display ── */
        .score-display {
            font-family: var(--t4-font);
            font-size: 4rem;
            font-weight: 700;
            color: var(--t4-accent-cyan);
            letter-spacing: -0.04em;
            line-height: 1;
        }

        /* ── Role Selector ── */
        .role-selector {
            background: var(--t4-surface-01);
            border: 1px solid var(--t4-border-subtle);
            border-radius: var(--t4-radius-md);
            padding: 18px;
            margin: 12px 0;
            backdrop-filter: blur(16px);
            transition: border-color var(--t4-ease-fast);
        }
        .role-selector:hover { border-color: var(--t4-border-accent); }

        /* ── Radar Container ── */
        .radar-container {
            background: var(--t4-surface-01);
            border: 1px solid var(--t4-border-subtle);
            border-radius: var(--t4-radius-lg);
            padding: 20px;
            margin: 18px 0;
            backdrop-filter: blur(16px);
        }

        /* ── Timer ── */
        .timer-container {
            background: linear-gradient(135deg, rgba(251,191,36,0.08) 0%, rgba(251,191,36,0.04) 100%);
            border: 1px solid rgba(251,191,36,0.25);
            border-radius: var(--t4-radius-md);
            padding: 14px;
            margin: 14px 0;
            text-align: center;
            backdrop-filter: blur(16px);
        }
        .timer-display {
            font-family: var(--t4-font);
            font-size: 1.4rem;
            font-weight: 700;
            color: var(--t4-accent-amber);
            letter-spacing: -0.01em;
        }
        .timer-urgent {
            color: var(--t4-accent-rose);
            animation: t4-subtlePulse 1s ease-in-out infinite;
        }

        /* ── Selectbox ── */
        .stSelectbox > div > div {
            background: var(--t4-surface-01) !important;
            border: 1px solid var(--t4-border-subtle) !important;
            border-radius: var(--t4-radius-sm) !important;
            color: var(--t4-text-primary) !important;
            font-family: var(--t4-font) !important;
            transition: border-color var(--t4-ease-fast) !important;
        }
        .stSelectbox > div > div:hover {
            border-color: rgba(79,163,227,0.35) !important;
            box-shadow: 0 0 0 3px rgba(79,163,227,0.08) !important;
        }

        /* ── Subheaders ── */
        .stApp h3 {
            font-family: var(--t4-font) !important;
            color: var(--t4-text-primary) !important;
            font-weight: 600 !important;
            letter-spacing: -0.02em !important;
            margin-bottom: 16px !important;
        }

        /* ── Alert/Info ── */
        .stAlert {
            background: var(--t4-surface-01) !important;
            border: 1px solid var(--t4-border-subtle) !important;
            border-radius: var(--t4-radius-md) !important;
            backdrop-filter: blur(16px) !important;
            font-family: var(--t4-font) !important;
            font-size: 0.875rem !important;
        }

        /* ── Video ── */
        .stVideo {
            border-radius: var(--t4-radius-md);
            overflow: hidden;
            box-shadow: var(--t4-shadow-card);
            transition: transform var(--t4-ease-base);
        }
        .stVideo:hover { transform: scale(1.01); }

        /* ── Radio buttons ── */
        .stRadio > div {
            flex-direction: row !important;
            justify-content: center !important;
            gap: 8px !important;
            flex-wrap: wrap !important;
        }
        .stRadio label {
            background: var(--t4-surface-01) !important;
            border: 1px solid var(--t4-border-subtle) !important;
            color: var(--t4-text-secondary) !important;
            padding: 10px 20px !important;
            border-radius: var(--t4-radius-sm) !important;
            cursor: pointer !important;
            transition: all var(--t4-ease-fast) !important;
            font-family: var(--t4-font) !important;
            font-weight: 500 !important;
            font-size: 0.875rem !important;
            text-align: center !important;
            backdrop-filter: blur(12px) !important;
            box-shadow: 0 2px 8px rgba(0,0,0,0.15) !important;
        }
        .stRadio label:hover {
            background: var(--t4-surface-hover) !important;
            border-color: rgba(79,163,227,0.35) !important;
            color: var(--t4-text-primary) !important;
            transform: translateY(-2px) !important;
        }
        .stRadio input:checked + div > label {
            background: linear-gradient(135deg, rgba(56,189,248,0.18) 0%, rgba(79,163,227,0.12) 100%) !important;
            color: var(--t4-accent-cyan) !important;
            border: 1px solid rgba(56,189,248,0.30) !important;
            font-weight: 600 !important;
            box-shadow: 0 2px 12px rgba(56,189,248,0.12) !important;
        }

        /* ── Score badge classes for table ── */
        .badge-excellent { background:rgba(52,211,153,0.15); color:#34d399; border:1px solid rgba(52,211,153,0.3); padding:2px 8px; border-radius:99px; font-size:12px; font-weight:600; }
        .badge-good      { background:rgba(56,189,248,0.12); color:#38bdf8; border:1px solid rgba(56,189,248,0.28); padding:2px 8px; border-radius:99px; font-size:12px; font-weight:600; }
        .badge-average   { background:rgba(251,191,36,0.12); color:#fbbf24; border:1px solid rgba(251,191,36,0.28); padding:2px 8px; border-radius:99px; font-size:12px; font-weight:600; }
        .badge-weak      { background:rgba(251,113,133,0.10); color:#fb7185; border:1px solid rgba(251,113,133,0.25); padding:2px 8px; border-radius:99px; font-size:12px; font-weight:600; }
        .badge-poor      { background:rgba(100,116,139,0.12); color:#64748b; border:1px solid rgba(100,116,139,0.25); padding:2px 8px; border-radius:99px; font-size:12px; font-weight:600; }

        </style>
    """, unsafe_allow_html=True)

    # Header (keeping existing)
    st.markdown("""
        <div class="header-box">
            <h2>📚 Recommended Learning Hub</h2>
        </div>
    """, unsafe_allow_html=True)

    # Subheader (keeping existing)
    st.markdown('<div class="glow-header">🎓 Explore Career Resources</div>', unsafe_allow_html=True)
    st.markdown("<p style='text-align:center; color:#94a3b8; font-family:-apple-system,BlinkMacSystemFont,\"SF Pro Display\",sans-serif; font-size: 0.95rem; margin-bottom: 22px; letter-spacing:-0.01em;'>Curated courses and videos for your career growth, resume tips, and interview success.</p>", unsafe_allow_html=True)

    # Learning path label (keeping existing)
    st.markdown("""
        <div class="learning-path-container">
            <span class="learning-path-text">
                🧭 Choose Your Learning Path
            </span>
        </div>
    """, unsafe_allow_html=True)

    # Updated Radio buttons with new options
    st.markdown("""
        <div style="display: flex; justify-content: center; width: 100%;">
            <div style="display: flex; justify-content: center; gap: 16px;">
    """, unsafe_allow_html=True)

    # Check if page changed away from AI Interview Coach - stop interview if so
    previous_page = st.session_state.get('previous_page_selection', None)

    page = st.radio(
        label="Select Learning Option",
        options=["Courses by Role", "Resume Videos", "Interview Videos", "AI Interview Coach 🤖", "My Progress 📊"],
        horizontal=True,
        key="page_selection",
        label_visibility="collapsed"
    )

    # STOP INTERVIEW ON TAB CHANGE
    if previous_page == "AI Interview Coach 🤖" and page != "AI Interview Coach 🤖":
        # User switched away from AI Interview Coach - reset interview state
        if st.session_state.get('dynamic_interview_started', False) and not st.session_state.get('dynamic_interview_completed', False):
            st.session_state.dynamic_interview_started = False
            st.session_state.dynamic_interview_completed = True

    # Update previous page for next comparison
    st.session_state.previous_page_selection = page

    st.markdown("</div></div>", unsafe_allow_html=True)

    # NEW: Index-based difficulty function (replaces keyword-based)
    def get_course_difficulty_by_index(index):
        if index == 0:
            return "Beginner"
        elif index in [1, 2]:
            return "Intermediate"
        else:
            return "Advanced"

    # Helper functions for dynamic question generation
    def generate_career_quiz_questions(domain, role):
        """Generate role-specific career quiz questions"""
        questions = []
        
        # Role-specific question templates
        role_templates = {
            "Software Development and Engineering": {
                "Frontend Developer": [
                    {
                        "question": "Which aspect of web development excites you most?",
                        "options": [
                            "Creating beautiful, interactive user interfaces",
                            "Building responsive designs that work on all devices", 
                            "Optimizing website performance and accessibility",
                            "Working with modern JavaScript frameworks"
                        ]
                    },
                    {
                        "question": "What's your preferred approach to styling?",
                        "options": [
                            "Writing custom CSS from scratch",
                            "Using CSS frameworks like Bootstrap or Tailwind",
                            "CSS-in-JS solutions for component-based styling", 
                            "CSS preprocessors like Sass or Less"
                        ]
                    },
                    {
                        "question": "Which tools do you enjoy working with most?",
                        "options": [
                            "React, Vue, or Angular for building SPAs",
                            "HTML5, CSS3, and vanilla JavaScript",
                            "Design tools like Figma or Adobe XD",
                            "Build tools like Webpack, Vite, or Parcel"
                        ]
                    }
                ],
                "Backend Developer": [
                    {
                        "question": "What backend architecture interests you most?",
                        "options": [
                            "RESTful API design and implementation",
                            "Microservices architecture and distributed systems",
                            "Database design and optimization",
                            "Server-side security and authentication"
                        ]
                    },
                    {
                        "question": "Which programming paradigm do you prefer?",
                        "options": [
                            "Object-oriented programming with Java/.NET",
                            "Functional programming with languages like Scala",
                            "Dynamic languages like Python or JavaScript",
                            "Systems programming with Go or Rust"
                        ]
                    },
                    {
                        "question": "What type of backend challenges excite you?",
                        "options": [
                            "Scaling applications to handle millions of users",
                            "Integrating complex third-party services",
                            "Optimizing database queries and performance",
                            "Building robust error handling and monitoring"
                        ]
                    }
                ],
                "Full Stack Developer": [
                    {
                        "question": "What full-stack aspect appeals to you most?",
                        "options": [
                            "Building end-to-end features from UI to database",
                            "Managing the entire application development lifecycle",
                            "Working with both frontend and backend technologies",
                            "Understanding how all system components interact"
                        ]
                    },
                    {
                        "question": "Which tech stack interests you most?",
                        "options": [
                            "MERN (MongoDB, Express, React, Node.js)",
                            "MEAN (MongoDB, Express, Angular, Node.js)",
                            "Django + React/Vue for Python development",
                            "Ruby on Rails with modern frontend frameworks"
                        ]
                    }
                ],
                "Mobile App Developer": [
                    {
                        "question": "What type of mobile development interests you?",
                        "options": [
                            "Native iOS development with Swift",
                            "Native Android development with Kotlin/Java",
                            "Cross-platform development with React Native",
                            "Hybrid app development with Flutter"
                        ]
                    },
                    {
                        "question": "Which mobile development aspect excites you most?",
                        "options": [
                            "Creating intuitive mobile user experiences",
                            "Integrating with device hardware and sensors",
                            "Optimizing app performance and battery usage",
                            "Publishing apps to App Store and Google Play"
                        ]
                    }
                ],
                "Game Developer": [
                    {
                        "question": "What type of game development interests you?",
                        "options": [
                            "3D game development with Unity or Unreal Engine",
                            "2D indie game development and pixel art",
                            "Mobile gaming and casual game mechanics",
                            "VR/AR game development and immersive experiences"
                        ]
                    },
                    {
                        "question": "Which game development aspect excites you most?",
                        "options": [
                            "Game design and player experience",
                            "Graphics programming and visual effects",
                            "Game physics and realistic simulations",
                            "Multiplayer networking and real-time systems"
                        ]
                    }
                ]
            },
            "Data Science and Analytics": {
                "Data Scientist": [
                    {
                        "question": "Which data science task excites you most?",
                        "options": [
                            "Building predictive models and machine learning algorithms",
                            "Exploring large datasets to discover hidden patterns",
                            "Creating data visualizations and storytelling with data",
                            "Designing experiments and A/B testing strategies"
                        ]
                    },
                    {
                        "question": "What's your preferred approach to data analysis?",
                        "options": [
                            "Statistical modeling and hypothesis testing",
                            "Deep learning and neural networks",
                            "Feature engineering and data preprocessing",
                            "Time series analysis and forecasting"
                        ]
                    },
                    {
                        "question": "Which tools do you enjoy working with most?",
                        "options": [
                            "Python with pandas, scikit-learn, and TensorFlow",
                            "R for statistical computing and analysis",
                            "SQL for database querying and data manipulation",
                            "Jupyter notebooks for exploratory data analysis"
                        ]
                    }
                ],
                "Data Analyst": [
                    {
                        "question": "Which type of analysis interests you most?",
                        "options": [
                            "Business intelligence and performance dashboards",
                            "Customer behavior analysis and segmentation",
                            "Financial analysis and risk assessment",
                            "Market research and competitive analysis"
                        ]
                    },
                    {
                        "question": "What's your preferred way to present insights?",
                        "options": [
                            "Interactive dashboards with Tableau or Power BI",
                            "Statistical reports with clear recommendations",
                            "Data visualizations and infographics",
                            "Executive summaries and business presentations"
                        ]
                    }
                ],
                "Machine Learning Engineer": [
                    {
                        "question": "Which ML engineering task excites you most?",
                        "options": [
                            "Deploying models to production at scale",
                            "Building ML pipelines and automation systems",
                            "Optimizing model performance and efficiency",
                            "Implementing MLOps and model monitoring"
                        ]
                    },
                    {
                        "question": "What type of ML problems interest you?",
                        "options": [
                            "Computer vision and image processing",
                            "Natural language processing and text analysis",
                            "Recommendation systems and personalization",
                            "Reinforcement learning and autonomous systems"
                        ]
                    }
                ]
            },
            "Cloud Computing and DevOps": {
                "Cloud Architect": [
                    {
                        "question": "Which cloud architecture aspect interests you most?",
                        "options": [
                            "Designing scalable, fault-tolerant systems",
                            "Multi-cloud and hybrid cloud strategies",
                            "Cloud security and compliance frameworks",
                            "Cost optimization and resource management"
                        ]
                    },
                    {
                        "question": "What type of cloud solutions excite you?",
                        "options": [
                            "Serverless architectures and event-driven systems",
                            "Container orchestration with Kubernetes",
                            "Data lakes and analytics platforms",
                            "AI/ML platforms and managed services"
                        ]
                    }
                ],
                "DevOps Engineer": [
                    {
                        "question": "Which DevOps practice interests you most?",
                        "options": [
                            "Building CI/CD pipelines and automation",
                            "Infrastructure as Code with Terraform/CloudFormation",
                            "Container orchestration and microservices",
                            "Monitoring, logging, and observability"
                        ]
                    },
                    {
                        "question": "What type of automation excites you?",
                        "options": [
                            "Deployment automation and release management",
                            "Infrastructure provisioning and configuration",
                            "Testing automation and quality gates",
                            "Incident response and self-healing systems"
                        ]
                    }
                ],
                "Site Reliability Engineer": [
                    {
                        "question": "Which SRE responsibility interests you most?",
                        "options": [
                            "Maintaining system reliability and uptime",
                            "Performance optimization and capacity planning",
                            "Incident management and post-mortem analysis",
                            "Service level objectives and error budgets"
                        ]
                    },
                    {
                        "question": "What aspect of system reliability excites you?",
                        "options": [
                            "Building robust monitoring and alerting systems",
                            "Designing disaster recovery and backup strategies",
                            "Automating operational tasks and runbooks",
                            "Analyzing system performance and bottlenecks"
                        ]
                    }
                ]
            },
            "Cybersecurity": {
                "Security Analyst": [
                    {
                        "question": "Which security area interests you most?",
                        "options": [
                            "Threat detection and incident response",
                            "Vulnerability assessment and risk management",
                            "Security monitoring and SIEM analysis",
                            "Compliance and security policy development"
                        ]
                    },
                    {
                        "question": "What type of security challenges excite you?",
                        "options": [
                            "Investigating security breaches and forensics",
                            "Analyzing malware and attack patterns",
                            "Network security and firewall management",
                            "Identity and access management systems"
                        ]
                    }
                ],
                "Penetration Tester": [
                    {
                        "question": "Which penetration testing approach interests you?",
                        "options": [
                            "Web application security testing",
                            "Network penetration testing and infrastructure",
                            "Social engineering and phishing simulations",
                            "Mobile application security testing"
                        ]
                    },
                    {
                        "question": "What aspect of ethical hacking excites you?",
                        "options": [
                            "Finding vulnerabilities before malicious actors",
                            "Using creative techniques to bypass security",
                            "Helping organizations improve their defenses",
                            "Staying updated on latest attack methods"
                        ]
                    }
                ]
            },
            "UI/UX Design": {
                "UI Designer": [
                    {
                        "question": "Which UI design aspect interests you most?",
                        "options": [
                            "Creating visually stunning interface designs",
                            "Designing consistent design systems and components",
                            "Working with typography, colors, and visual hierarchy",
                            "Prototyping interactions and micro-animations"
                        ]
                    },
                    {
                        "question": "What type of design work excites you?",
                        "options": [
                            "Mobile app interface design",
                            "Web application and dashboard design",
                            "Icon design and visual asset creation",
                            "Brand identity and visual design systems"
                        ]
                    }
                ],
                "UX Designer": [
                    {
                        "question": "Which UX design activity interests you most?",
                        "options": [
                            "User research and persona development",
                            "Information architecture and user flows",
                            "Wireframing and prototype development",
                            "Usability testing and design validation"
                        ]
                    },
                    {
                        "question": "What aspect of user experience excites you?",
                        "options": [
                            "Solving complex user problems with simple solutions",
                            "Understanding user behavior and psychology",
                            "Designing accessible and inclusive experiences",
                            "Measuring and optimizing user engagement"
                        ]
                    }
                ]
            },
            "Project Management": {
                "Project Manager": [
                    {
                        "question": "Which project management aspect interests you most?",
                        "options": [
                            "Planning and scheduling project timelines",
                            "Coordinating teams and stakeholder communication",
                            "Risk management and problem-solving",
                            "Budget management and resource allocation"
                        ]
                    },
                    {
                        "question": "What type of projects excite you?",
                        "options": [
                            "Large-scale software development projects",
                            "Cross-functional digital transformation initiatives",
                            "Product launches and go-to-market strategies",
                            "Process improvement and organizational change"
                        ]
                    }
                ],
                "Product Manager": [
                    {
                        "question": "Which product management activity interests you most?",
                        "options": [
                            "Product strategy and roadmap development",
                            "User research and market analysis",
                            "Feature prioritization and requirement gathering",
                            "Go-to-market strategy and product launches"
                        ]
                    },
                    {
                        "question": "What aspect of product development excites you?",
                        "options": [
                            "Identifying user needs and pain points",
                            "Defining product vision and strategy",
                            "Working with engineering and design teams",
                            "Analyzing product metrics and user feedback"
                        ]
                    }
                ]
            }
        }

        # Get role-specific questions or generate generic ones
        if domain in role_templates and role in role_templates[domain]:
            questions = role_templates[domain][role]
        else:
            # Generate generic questions based on role name
            questions = [
                {
                    "question": f"How interested are you in pursuing a career as a {role}?",
                    "options": [
                        "Very interested - it's my dream job",
                        "Somewhat interested - I want to learn more",
                        "Moderately interested - it seems challenging",
                        "Not very interested - but I'm curious"
                    ]
                },
                {
                    "question": f"What attracts you most about the {role} role?",
                    "options": [
                        "The technical challenges and problem-solving",
                        "The creative aspects and innovation opportunities", 
                        "The career growth potential and salary",
                        "The impact on users and business outcomes"
                    ]
                }
            ]
        
        return questions

    # Helper function to generate fallback questions
    def self_generate_fallback_questions(role, domain, difficulty, count):
        """Generate fallback questions when LLM doesn't return enough"""
        if difficulty == "Easy":
            base_questions = [
                f"What interests you most about the {role} position?",
                f"Describe your basic understanding of {role} responsibilities.",
                f"What are the fundamental skills needed for {role}?",
                f"How do you stay updated with trends in {domain}?",
                f"Why do you want to work as a {role}?",
                f"What do you know about the {role} role?",
                f"Tell me about yourself and your interest in {role}.",
                f"What motivates you to pursue a career in {domain}?",
                f"Describe a project you've worked on related to {role}.",
                f"What are your career goals as a {role}?"
            ]
        elif difficulty == "Hard":
            # Text-answerable Hard fallbacks — one challenge axis per question.
            base_questions = [
                f"Describe the most significant technical tradeoff you've encountered as a {role}. "
                f"What were the two options, what data drove your decision, and what limitation did you accept?",

                f"Walk through how you would diagnose an unexpected latency spike in a {domain} system you own. "
                f"What signals would you look for first, and what would be your isolation process?",

                f"What is the single most dangerous assumption developers make when working in {domain}, "
                f"and how would you build a guardrail against it?",

                f"You're asked to reduce the memory footprint of a {role} component by 30% "
                f"without changing its public interface. Describe your investigation process and the "
                f"two or three changes you'd prioritise.",

                f"Describe a non-obvious edge case in {domain} that is easy to miss in code review. "
                f"How would you detect it, handle it, and prevent its recurrence?",

                f"You need to choose between two technically valid approaches to implement "
                f"[a core {domain} feature]. What framework do you use to make that call, "
                f"and what would make you revisit the decision later?",

                f"A {domain} service you maintain starts failing intermittently under normal load. "
                f"There are no upstream alerts. Walk through your debugging approach step by step.",

                f"What optimisation would have the highest impact on the reliability of a typical {domain} system? "
                f"Justify your choice with specific reasoning, not just general best practices.",

                f"Describe one {domain} pattern or technology that is frequently misused in production. "
                f"What is the misuse pattern, and how do you recognise it in a codebase?",

                f"You've inherited a {domain} codebase with no tests and unclear ownership. "
                f"What is the first concrete technical action you take, and why that over other options?",
            ]
        else:  # Medium
            base_questions = [
                f"Describe a challenging project you've worked on relevant to {role}.",
                f"How do you approach problem-solving in {domain}?",
                f"What tools and technologies are you most comfortable with for {role}?",
                f"Tell me about a time you had to learn a new skill for {role}.",
                f"How do you prioritize tasks when working as a {role}?",
                f"Describe your experience with {domain} technologies.",
                f"How do you handle tight deadlines as a {role}?",
                f"What's your approach to code quality in {domain}?",
                f"Tell me about a technical challenge you solved as a {role}.",
                f"How do you collaborate with team members in {domain}?"
            ]
        return base_questions[:count]

    # UPDATED: AI-Generated Questions using LLM with DIFFICULTY SUPPORT
    def generate_interview_questions_with_llm(domain, role, interview_type, num_questions, difficulty="Medium"):
        """
        Generate interview questions using LLM based on domain, role, type, and difficulty.

        FIXED: Now difficulty is passed into LLM prompt and affects question complexity
        """
        # Define difficulty-specific instructions
        # TEXT-INTERVIEW-OPTIMISED difficulty specifications.
        # Every level is scoped so the candidate can answer in structured paragraphs — no whiteboard.
        difficulty_instructions = {
            "Easy": (
                "Generate CONCEPT CLARITY questions only. "
                "Each question asks the candidate to define or explain ONE concept, "
                "state why it exists, and give a real-world example of where they'd use it. "
                "Questions must be answerable in 3-5 paragraphs. "
                "FORBIDDEN: system design, scaling, tradeoffs, production failures, architecture."
            ),
            "Medium": (
                "Generate SCENARIO REASONING questions. "
                "Each question presents a small, realistic scenario with ONE constraint or decision point. "
                "The candidate must describe their approach and justify ONE key implementation choice. "
                "Questions must be answerable in 5-6 paragraphs. "
                "FORBIDDEN: full system design, multi-layer architecture, handling 1M+ users, "
                "combined scaling + tradeoff + failure in one question."
            ),
            "Hard": (
                "Generate FOCUSED TECHNICAL DEPTH questions targeting ONE challenge axis. "
                "Choose EXACTLY ONE of: (a) a specific tradeoff between two concrete approaches, "
                "(b) diagnosing and handling one specific failure mode, "
                "or (c) optimising one metric under one constraint. "
                "Questions must be answerable in 6-8 text paragraphs — no diagram, no whiteboard. "
                "FORBIDDEN: 'design the entire system', 'walk through every layer', "
                "'handle X million users AND secure it AND handle failures AND optimise'. "
                "One axis. One decision. Senior depth, text-answerable scope."
            ),
        }

        prompt = f"""You are an expert technical interviewer building a text-based interview simulator.

Generate EXACTLY {num_questions} unique {interview_type} interview questions
for the role of {role} in {domain}.

DIFFICULTY CONTRACT: {difficulty}
{difficulty_instructions.get(difficulty, difficulty_instructions["Medium"])}

GENERATION RULES:
- EXACTLY {num_questions} questions — no more, no less
- Each question must be self-contained (1-3 sentences), answerable in text paragraphs
- Each question must focus on ONE concept, ONE scenario, or ONE challenge axis
- For Hard: do NOT combine design + scale + failure + tradeoff in one question
- Avoid duplicates and generic filler
- Output ONLY the questions, one per line
- NO numbering, NO bullet points, NO prefixes, NO introductory text

Generate {num_questions} questions now:
"""

        try:
            response = call_llm(prompt, session=st.session_state)

            # Split by newlines and clean up
            raw_questions = [q.strip() for q in response.split('\n') if q.strip()]

            # Remove any numbering or bullet points more aggressively
            import re
            cleaned_questions = []
            for q in raw_questions:
                # Remove various prefixes: "1. ", "1) ", "- ", "• ", "* ", "Question 1:", etc.
                clean_q = re.sub(r'^[\d\)\.\-•\*]+\s*', '', q).strip()
                clean_q = re.sub(r'^Question\s*\d*\s*:?\s*', '', clean_q, flags=re.IGNORECASE).strip()

                # Only add if it's a meaningful question
                if clean_q and len(clean_q) > 15 and not clean_q.lower().startswith('generate') and not clean_q.lower().startswith('here'):
                    cleaned_questions.append(clean_q)

                # Stop if we have enough questions
                if len(cleaned_questions) >= num_questions:
                    break

            # If we got fewer questions than requested, try to pad with fallback
            if len(cleaned_questions) < num_questions:
                st.warning(f"Only generated {len(cleaned_questions)} questions, padding with fallback questions...")
                # Add fallback questions to meet the requirement
                fallback_needed = num_questions - len(cleaned_questions)
                fallback_qs = self_generate_fallback_questions(role, domain, difficulty, fallback_needed)
                cleaned_questions.extend(fallback_qs)

            # EXACT QUESTION COUNT: Enforce exact count
            cleaned_questions = cleaned_questions[:num_questions]
            return cleaned_questions

        except Exception as e:
            st.error(f"Failed to generate questions with LLM: {e}")
            # Fallback to static questions appropriate for difficulty
            if difficulty == "Easy":
                fallback_questions = [
                    f"What interests you most about the {role} position?",
                    f"Describe your basic understanding of {role} responsibilities.",
                    f"What are the fundamental skills needed for {role}?",
                    f"How do you stay updated with trends in {domain}?",
                    f"Why do you want to work as a {role}?"
                ]
            elif difficulty == "Hard":
                # Text-answerable Hard fallbacks — one challenge axis each.
                fallback_questions = [
                    f"Describe the most significant tradeoff you've faced in {domain}. "
                    f"What were the options, and what drove your final decision?",

                    f"Walk through how you would diagnose an unexpected performance regression "
                    f"in a {domain} system you own. What would you check first?",

                    f"What is one non-obvious edge case in {domain} that developers frequently "
                    f"miss? How would you detect and handle it?",

                    f"You are asked to reduce latency for a {role} component by 40% "
                    f"without changing its interface. Describe your investigation and top two changes.",

                    f"Describe a {domain} pattern or tool that is often misused in production. "
                    f"How do you recognise the misuse, and what would you do instead?",
                ]
            else:  # Medium
                fallback_questions = [
                    f"Describe a challenging project you've worked on relevant to {role}.",
                    f"How do you approach problem-solving in {domain}?",
                    f"What tools and technologies are you most comfortable with for {role}?",
                    f"Tell me about a time you had to learn a new skill for {role}.",
                    f"How do you prioritize tasks when working as a {role}?"
                ]
            return fallback_questions[:num_questions]

    # Badge system for gamification
    BADGE_CONFIG = {
        "career_quiz": {
            "novice": {"min_score": 0, "max_score": 40, "emoji": "🌱", "title": "Career Explorer"},
            "intermediate": {"min_score": 41, "max_score": 70, "emoji": "📚", "title": "Career Seeker"},
            "advanced": {"min_score": 71, "max_score": 100, "emoji": "🎯", "title": "Career Champion"}
        },
        "interview": {
            "needs_practice": {"min_score": 1.0, "max_score": 2.5, "emoji": "💪", "title": "Keep Practicing"},
            "good": {"min_score": 2.6, "max_score": 3.5, "emoji": "👍", "title": "Good Performer"},
            "excellent": {"min_score": 3.6, "max_score": 4.5, "emoji": "🌟", "title": "Star Performer"},
            "interview_ready": {"min_score": 4.6, "max_score": 5.0, "emoji": "🏆", "title": "Interview Ready"}
        }
    }

    def get_badge_for_score(score_type, score):
        """Get badge based on score type and value"""
        badges = BADGE_CONFIG.get(score_type, {})
        for badge_name, config in badges.items():
            if config["min_score"] <= score <= config["max_score"]:
                return config["emoji"], config["title"]
        return "🎖️", "Participant"

    def create_skill_radar_chart(skills_data):
        """Create a radar chart for skills using Plotly"""
        # Extract skills and values
        skills = list(skills_data.keys())
        values = list(skills_data.values())
        
        # Create radar chart
        fig = go.Figure()
        
        fig.add_trace(go.Scatterpolar(
            r=values,
            theta=skills,
            fill='toself',
            name='Skills',
            line=dict(color='#00c3ff', width=2),
            fillcolor='rgba(0, 195, 255, 0.2)',
            hovertemplate='<b>%{theta}</b><br>Importance: %{r}/10<br><extra></extra>'
        ))
        
        fig.update_layout(
            polar=dict(
                radialaxis=dict(
                    visible=True,
                    range=[0, 10],
                    tickfont=dict(color='white', size=10),
                    gridcolor='rgba(255, 255, 255, 0.2)'
                ),
                angularaxis=dict(
                    tickfont=dict(color='white', size=12),
                    gridcolor='rgba(255, 255, 255, 0.2)'
                ),
                bgcolor='rgba(0, 0, 0, 0)'
            ),
            showlegend=False,
            title=dict(
                text="Skills Importance Radar",
                x=0.5,
                font=dict(color='#00c3ff', size=16)
            ),
            paper_bgcolor='rgba(0, 0, 0, 0)',
            plot_bgcolor='rgba(0, 0, 0, 0)',
            font=dict(color='white'),
            height=400
        )
        
        return fig

    def get_course_description(course_title, role):
        """Generate a short description for the course"""
        descriptions = {
            'Frontend Developer': f"Master modern frontend development with {course_title.split()[0]} and build responsive web applications.",
            'Backend Developer': f"Learn server-side development and API design to become a skilled backend developer.",
            'Full Stack Developer': f"Comprehensive full-stack development course covering both frontend and backend technologies.",
            'Data Scientist': f"Dive deep into data science methodologies, machine learning, and statistical analysis.",
            'Machine Learning Engineer': f"Build and deploy machine learning models at scale with industry best practices.",
            'Cloud Architect': f"Design scalable cloud infrastructure and learn enterprise-grade cloud solutions.",
            'DevOps Engineer': f"Master CI/CD pipelines, containerization, and infrastructure automation.",
            'UI Designer': f"Create stunning user interfaces with modern design principles and tools.",
            'UX Designer': f"Learn user research, wireframing, and create exceptional user experiences."
        }
        
        return descriptions.get(role, f"Comprehensive course to advance your skills in {role} role.")

    def display_courses_by_difficulty(courses, role):
        """Display courses grouped by difficulty using index-based mapping"""
        # Group courses by difficulty
        difficulty_groups = {"Beginner": [], "Intermediate": [], "Advanced": []}
        
        for idx, (title, url) in enumerate(courses):
            difficulty = get_course_difficulty_by_index(idx)
            description = get_course_description(title, role)
            difficulty_groups[difficulty].append((title, url, description))
        
        # Display each difficulty group
        for difficulty in ["Beginner", "Intermediate", "Advanced"]:
            if difficulty_groups[difficulty]:
                st.markdown(f"### 🎯 {difficulty} Level")
                for title, url, description in difficulty_groups[difficulty]:
                    st.markdown(f"""
                        <div class="course-tile">
                            <div class="course-title">{title}</div>
                            <div class="course-description">{description}</div>
                            <span class="difficulty-badge difficulty-{difficulty.lower()}">{difficulty}</span>
                            <br>
                            <a href="{url}" target="_blank" class="course-link-btn">
                                🚀 Start Learning
                            </a>
                        </div>
                    """, unsafe_allow_html=True)

    # UPDATED SECTIONS

    # Section 1: UPDATED Courses by Role with Index-based Difficulty
    if page == "Courses by Role":
        st.subheader("🎯 Courses by Career Role")
        
        col1, col2 = st.columns(2)
        with col1:
            category = st.selectbox(
                "Select Career Category",
                options=list(COURSES_BY_CATEGORY.keys()),
                key="category_selection"
            )
        
        with col2:
            if category:
                roles = list(COURSES_BY_CATEGORY[category].keys())
                role = st.selectbox(
                    "Select Role / Job Title",
                    options=roles,
                    key="role_selection"
                )
            else:
                role = None
        
        if category and role:
            # UPDATED: Add difficulty filter
            difficulty_filter = st.selectbox(
                "Filter by Difficulty Level",
                options=["All Levels", "Beginner", "Intermediate", "Advanced"],
                key="difficulty_filter"
            )
            
            st.subheader(f"📘 Courses for **{role}** in **{category}**:")
            courses = get_courses_for_role(category, role)
            
            if courses:
                # UPDATED: Display courses using index-based difficulty
                filtered_courses = []
                for idx, (title, url) in enumerate(courses):
                    difficulty = get_course_difficulty_by_index(idx)
                    
                    # Apply difficulty filter
                    if difficulty_filter == "All Levels" or difficulty == difficulty_filter:
                        filtered_courses.append((title, url, difficulty, idx))
                
                if filtered_courses:
                    for title, url, difficulty, idx in filtered_courses:
                        description = get_course_description(title, role)
                        
                        # UPDATED: Interactive course tile with index-based difficulty
                        st.markdown(f"""
                            <div class="course-tile">
                                <div class="course-title">{title}</div>
                                <div class="course-description">{description}</div>
                                <span class="difficulty-badge difficulty-{difficulty.lower()}">{difficulty}</span>
                                <br>
                                <a href="{url}" target="_blank" class="course-link-btn">
                                    🚀 Start Learning
                                </a>
                            </div>
                        """, unsafe_allow_html=True)
                else:
                    st.info("🚫 No courses found for this difficulty level.")
            else:
                st.info("🚫 No courses found for this role.")
        
        # Show skill radar chart for selected role
        if category and role:
            st.markdown("---")
            st.markdown('<div class="radar-container">', unsafe_allow_html=True)
            st.subheader("🎯 Skills Radar Chart")
            
            # Generate sample skills data based on role
            role_skills = {
                # ==== Software Development & Engineering ====
                "Frontend Developer": {
                    "JavaScript": 9, "React/Vue": 8, "CSS/HTML": 9,
                    "Responsive Design": 8, "Performance Optimization": 7, "Testing": 6
                },
                "Backend Developer": {
                    "API Design": 9, "Database Management": 8, "Security": 8,
                    "Scalability": 7, "Cloud Services": 7, "Testing": 6
                },
                "Full Stack Developer": {
                    "Frontend": 8, "Backend": 8, "Databases": 7,
                    "API Integration": 8, "DevOps Basics": 6, "Testing": 7
                },
                "Mobile App Developer": {
                    "Flutter/React Native": 8, "Swift/Kotlin": 8, "UI/UX": 8,
                    "APIs": 7, "Performance Optimization": 7, "App Deployment": 7
                },
                "Game Developer": {
                    "Unity/Unreal": 9, "C# / C++": 8, "Game Physics": 7,
                    "Graphics/Rendering": 8, "AI in Games": 6, "Multiplayer Systems": 7
                },
                # ==== Data Science & Analytics ====
                "Data Scientist": {
                    "Python/R": 9, "Machine Learning": 8, "Statistics": 9,
                    "Data Visualization": 7, "SQL": 8, "Domain Knowledge": 6
                },
                "Data Analyst": {
                    "SQL": 9, "Excel/Spreadsheets": 8, "Visualization": 8,
                    "Statistics": 8, "Python/R": 7, "Business Acumen": 7
                },
                "Machine Learning Engineer": {
                    "ML Algorithms": 9, "Deep Learning": 8, "MLOps": 7,
                    "Data Engineering": 8, "Python/Frameworks": 9, "Cloud Deployment": 7
                },
                # ==== Cloud Computing & DevOps ====
                "Cloud Architect": {
                    "AWS/Azure/GCP": 9, "System Design": 8, "Networking": 7,
                    "Security": 8, "Scalability": 9, "Cost Optimization": 7
                },
                "DevOps Engineer": {
                    "CI/CD": 9, "Containerization": 8, "Cloud Platforms": 8,
                    "Monitoring": 7, "Infrastructure as Code": 8, "Security": 7
                },
                "Site Reliability Engineer": {
                    "Reliability Engineering": 9, "Monitoring": 8, "Automation": 8,
                    "Incident Response": 8, "System Design": 7, "Security": 7
                },
                # ==== Cybersecurity ====
                "Security Analyst": {
                    "Threat Detection": 9, "Incident Response": 8, "Networking": 7,
                    "SIEM Tools": 8, "Risk Management": 7, "Compliance": 6
                },
                "Penetration Tester": {
                    "Ethical Hacking": 9, "Web Security": 8, "Exploitation": 8,
                    "Scripting": 7, "Reporting": 6, "Network Security": 7
                },
                # ==== UI/UX Design ====
                "UI Designer": {
                    "Design Tools": 9, "Visual Design": 8, "Typography": 7,
                    "Color Theory": 8, "Prototyping": 7, "User Research": 6
                },
                "UX Designer": {
                    "User Research": 9, "Wireframing": 8, "Prototyping": 8,
                    "Usability Testing": 7, "Accessibility": 8, "Design Thinking": 7
                },
                # ==== Project Management ====
                "Project Manager": {
                    "Planning": 9, "Communication": 8, "Risk Management": 8,
                    "Leadership": 7, "Agile/Scrum": 8, "Budgeting": 7
                },
                "Product Manager": {
                    "Market Research": 9, "Product Strategy": 8, "Analytics": 8,
                    "Communication": 8, "Agile Methods": 7, "User-Centered Design": 7
                }
            }
            
            skills_data = role_skills.get(role, {
                "Technical Skills": 8, "Problem Solving": 7, "Communication": 6,
                "Leadership": 5, "Domain Knowledge": 7, "Continuous Learning": 8
            })
            
            # Create and display radar chart
            radar_fig = create_skill_radar_chart(skills_data)
            st.plotly_chart(radar_fig, use_container_width=True)
            
            # Add hover tooltip information
            st.markdown("""
                <div style="text-align: center; color: #38bdf8; margin-top: 10px;">
                    💡 Hover over the chart points to see skill importance ratings!
                </div>
            """, unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)

    # Section 2: Resume Videos (unchanged)
    elif page == "Resume Videos":
        st.subheader("📄 Resume Writing Videos")
        categories = list(RESUME_VIDEOS.keys())
        selected_cat = st.selectbox(
            "Select Resume Video Category",
            options=categories,
            key="resume_vid_cat"
        )
        if selected_cat:
            st.subheader(f"📂 {selected_cat}")
            videos = RESUME_VIDEOS[selected_cat]
            cols = st.columns(2)
            for idx, (title, url) in enumerate(videos):
                with cols[idx % 2]:
                    st.markdown(f"**{title}**")
                    st.video(url)

    # Section 3: Interview Videos (unchanged)
    elif page == "Interview Videos":
        st.subheader("🗣️ Interview Preparation Videos")
        categories = list(INTERVIEW_VIDEOS.keys())
        selected_cat = st.selectbox(
            "Select Interview Video Category",
            options=categories,
            key="interview_vid_cat"
        )
        if selected_cat:
            st.subheader(f"📂 {selected_cat}")
            videos = INTERVIEW_VIDEOS[selected_cat]
            cols = st.columns(2)
            for idx, (title, url) in enumerate(videos):
                with cols[idx % 2]:
                    st.markdown(f"**{title}**")
                    st.video(url)

    # Section 4: UPDATED AI Interview Coach 🤖 with Resume-Based Interviewing
    elif page == "AI Interview Coach 🤖":
        st.subheader("🤖 AI Interview Coach")
        st.markdown("Upload your resume and practice role-specific interview questions with AI-powered feedback!")

        # Create database tables if not yet done this session (runs once, never on every rerun)
        _ensure_db_initialized()

        # Initialize resume state
        if 'resume_file' not in st.session_state:
            st.session_state.resume_file = None
        if 'resume_context' not in st.session_state:
            st.session_state.resume_context = None
        if 'interview_phase' not in st.session_state:
            st.session_state.interview_phase = "resume"
        if 'resume_questions_answered' not in st.session_state:
            st.session_state.resume_questions_answered = 0

        # RESUME UPLOAD SECTION (MANDATORY)
        st.markdown("---")
        st.markdown("<h3 style='color:#38bdf8;font-family:-apple-system,BlinkMacSystemFont,\"SF Pro Display\",sans-serif;font-weight:600;letter-spacing:-0.02em;'>📄 Step 1: Upload Your Resume</h3>", unsafe_allow_html=True)

        # ── AI Coach quota badge (always visible, even before resume upload) ──
        _ac_username_early = st.session_state.get("username")
        if _ac_username_early:
            _ac_used_early = get_usage_count_last_hour(_ac_username_early, "ai_coach")
            _ac_remaining_early = max(0, 2 - _ac_used_early)
            _ac_color_early = "#34d399" if _ac_remaining_early > 0 else "#fb7185"
            _ac_bg_early = "rgba(52,211,153,0.07)" if _ac_remaining_early > 0 else "rgba(251,113,133,0.07)"
            _ac_border_early = "rgba(52,211,153,0.25)" if _ac_remaining_early > 0 else "rgba(251,113,133,0.25)"
            _ac_icon_early = "🟢" if _ac_remaining_early > 0 else "🔴"
            _ac_status_early = f"{_ac_remaining_early}/2 mock interviews remaining this hour"
            if _ac_remaining_early == 0:
                _ac_status_early = "0/2 — Limit reached. Resets on a rolling 60-minute window."
            st.markdown(
                f'<div style="display:flex;align-items:center;gap:8px;font-size:0.82rem;'
                f'color:{_ac_color_early};background:{_ac_bg_early};'
                f'border:1px solid {_ac_border_early};border-radius:8px;'
                f'padding:9px 14px;margin-bottom:12px;font-family:-apple-system,sans-serif;">'
                f'{_ac_icon_early} <b>AI Coach Quota:</b>&nbsp;{_ac_status_early}'
                f'</div>',
                unsafe_allow_html=True
            )
        # ─────────────────────────────────────────────────────────────────────

        if st.session_state.resume_file is None:
            # ── Quota gate: block upload entirely if limit is reached ──────────
            _upload_quota_user = st.session_state.get("username")
            _upload_blocked = False
            if _upload_quota_user:
                _upload_used = get_usage_count_last_hour(_upload_quota_user, "ai_coach")
                _upload_remaining = max(0, 2 - _upload_used)
                if _upload_remaining == 0:
                    _upload_blocked = True
                    st.markdown(
                        '<div style="display:flex;align-items:center;gap:10px;'
                        'background:rgba(251,113,133,0.08);border:1px solid rgba(251,113,133,0.35);'
                        'border-radius:10px;padding:14px 18px;margin-bottom:12px;">'
                        '<span style="font-size:1.3rem;">🚫</span>'
                        '<div><b style="color:#fb7185;font-size:0.92rem;">Upload Limit Reached</b>'
                        '<p style="color:#fca5a5;font-size:0.82rem;margin:4px 0 0 0;">'
                        'You have used both of your mock interviews for this hour. '
                        'Please try again later.</p></div>'
                        '</div>',
                        unsafe_allow_html=True
                    )
            # ────────────────────────────────────────────────────────────────────

            if not _upload_blocked:
                uploaded_resume = st.file_uploader(
                    "Upload your resume (PDF format)",
                    type=['pdf'],
                    key="resume_uploader"
                )
            else:
                uploaded_resume = None

            if uploaded_resume:
                with st.spinner("Processing your resume..."):
                    # Extract text from PDF
                    resume_text = extract_resume_text_from_pdf(uploaded_resume)

                    if resume_text and len(resume_text.strip()) > 50:
                        st.session_state.resume_file = uploaded_resume.name
                        st.session_state.resume_raw_text = resume_text
                        st.session_state.interview_phase = "resume"
                        st.session_state.resume_questions_answered = 0

                        # Analyze resume immediately so "Key topics in scope" card
                        # is visible during interview setup (before Start Interview).
                        with st.spinner("Analyzing your resume with AI..."):
                            resume_context = analyze_resume_with_llm(resume_text)
                        st.session_state.resume_context = resume_context

                        st.success("✅ Resume uploaded and analyzed successfully!")
                        st.rerun()  # FIX 5: removed time.sleep(1) — blocks server thread
                    else:
                        st.error("Could not extract text from resume. Please ensure it's a valid PDF.")
        else:
            st.success(f"✅ Resume loaded: {st.session_state.resume_file}")
            

            if st.button("🔄 Upload Different Resume"):
                st.session_state.resume_file = None
                st.session_state.resume_context = None
                st.session_state.dynamic_interview_started = False
                st.session_state.dynamic_interview_completed = False
                st.session_state.interview_result_saved = False
                st.session_state.interview_final_duration_seconds = None
                st.session_state.interview_actual_start_time = None
                # Reset usage flag so the next interview is properly counted
                st.session_state._ac_usage_recorded_this_session = False
                st.rerun()

        # Only show domain/role selection if resume is uploaded
        if st.session_state.resume_file is not None:
            st.markdown("---")
            st.markdown("<h3 style='color:#38bdf8;font-family:-apple-system,BlinkMacSystemFont,\"SF Pro Display\",sans-serif;font-weight:600;letter-spacing:-0.02em;'>👔 Step 2: Select Target Role</h3>", unsafe_allow_html=True)

            # Domain and Role selection
            st.markdown('<div class="role-selector">', unsafe_allow_html=True)

            # ── TASK 1 FIX: Domain → Role override bug ───────────────────────
            # Initialize domain/role session_state on first run
            _domain_options = list(COURSES_BY_CATEGORY.keys())
            if "selected_domain" not in st.session_state or st.session_state.selected_domain not in _domain_options:
                st.session_state.selected_domain = _domain_options[0] if _domain_options else None

            def _on_domain_change():
                """Reset role whenever domain changes so stale roles never persist."""
                new_domain = st.session_state._domain_picker
                if new_domain in COURSES_BY_CATEGORY:
                    st.session_state.selected_domain = new_domain
                    _new_roles = list(COURSES_BY_CATEGORY[new_domain].keys())
                    st.session_state.target_role = _new_roles[0] if _new_roles else None
                    # Also reset interview state when domain changes
                    st.session_state.interview_domain = new_domain
                    st.session_state.interview_role = st.session_state.target_role

            col1, col2 = st.columns(2)
            with col1:
                _current_domain_idx = _domain_options.index(st.session_state.selected_domain) if st.session_state.selected_domain in _domain_options else 0
                selected_domain = st.selectbox(
                    "Select Career Domain",
                    options=_domain_options,
                    index=_current_domain_idx,
                    key="_domain_picker",
                    on_change=_on_domain_change
                )
                # Keep selected_domain session_state in sync on initial render
                st.session_state.selected_domain = selected_domain

            with col2:
                if selected_domain:
                    roles = list(COURSES_BY_CATEGORY[selected_domain].keys())
                    # Ensure stored target_role is valid for this domain; reset if not
                    if "target_role" not in st.session_state or st.session_state.target_role not in roles:
                        st.session_state.target_role = roles[0] if roles else None
                    _current_role_idx = roles.index(st.session_state.target_role) if st.session_state.target_role in roles else 0
                    selected_role = st.selectbox(
                        "Select Target Role",
                        options=roles,
                        index=_current_role_idx,
                        key="interview_role_selection"
                    )
                    # Keep target_role session_state in sync
                    st.session_state.target_role = selected_role
                    st.session_state.interview_role = selected_role  # keep fragment in sync
                else:
                    selected_role = None
                    st.session_state.target_role = None
            # ── END TASK 1 FIX ────────────────────────────────────────────────

            st.markdown('</div>', unsafe_allow_html=True)
        else:
            selected_domain = None
            selected_role = None
        
        if selected_domain and selected_role:
            # Initialize interview state
            if 'dynamic_interview_questions' not in st.session_state:
                st.session_state.dynamic_interview_questions = []
            if 'current_dynamic_interview_question' not in st.session_state:
                st.session_state.current_dynamic_interview_question = 0
            if 'dynamic_interview_answers' not in st.session_state:
                st.session_state.dynamic_interview_answers = []
            if 'dynamic_interview_scores' not in st.session_state:
                st.session_state.dynamic_interview_scores = []
            if 'dynamic_interview_feedbacks' not in st.session_state:
                st.session_state.dynamic_interview_feedbacks = []
            if 'dynamic_interview_completed' not in st.session_state:
                st.session_state.dynamic_interview_completed = False
            if 'dynamic_interview_started' not in st.session_state:
                st.session_state.dynamic_interview_started = False
            if 'dynamic_answer_submitted' not in st.session_state:
                st.session_state.dynamic_answer_submitted = False
            if 'current_interview_question_text' not in st.session_state:
                st.session_state.current_interview_question_text = ""
            if 'interview_domain' not in st.session_state or st.session_state.interview_domain != selected_domain:
                st.session_state.interview_domain = selected_domain
                st.session_state.interview_role = selected_role
                st.session_state.dynamic_interview_started = False
                st.session_state.dynamic_interview_completed = False
                st.session_state.interview_result_saved = False
                st.session_state.interview_final_duration_seconds = None
                st.session_state.interview_actual_start_time = None
                # Reset usage flag so the next interview is properly counted
                st.session_state._ac_usage_recorded_this_session = False
                # Clear timer thread keys so auto-submit works on next interview
                for _k in [k for k in st.session_state if k.startswith("_timer_thread_armed_")]:
                    st.session_state.pop(_k, None)
            if 'question_timer_start' not in st.session_state:
                st.session_state.question_timer_start = None
            if 'timer_seconds' not in st.session_state:
                st.session_state.timer_seconds = 120
            if 'interview_difficulty' not in st.session_state:
                st.session_state.interview_difficulty = "Medium"
            if 'interview_mode' not in st.session_state:
                st.session_state.interview_mode = "mixed"
            if 'original_num_questions' not in st.session_state:
                st.session_state.original_num_questions = 6
            if 'resume_based_questions' not in st.session_state:
                st.session_state.resume_based_questions = []
            if 'generic_questions' not in st.session_state:
                st.session_state.generic_questions = []
            if 'current_interview_id' not in st.session_state:
                st.session_state.current_interview_id = None
            # Track DB row ids for parent_question_id linkage: list of row ids per question answered
            if 'question_db_ids' not in st.session_state:
                st.session_state.question_db_ids = []

            # Start interview setup
            if not st.session_state.dynamic_interview_started:
                st.markdown(f"### Practice interview for: {selected_role}")

                # PART 5: Show weakness memory insight
                _username_wm = st.session_state.get("username", "Guest")
                _wm = get_user_weakness_history(_username_wm)
                if _wm.get("weakest_skill"):
                    _wm_avgs = _wm.get("averages", {})
                    _wm_skill = _wm["weakest_skill"].title()
                    _wm_score = _wm_avgs.get(_wm["weakest_skill"], 0)
                    _wm_count = _wm.get("interview_count", 0)
                    _wm_label = f"last {_wm_count} interview{'s' if _wm_count != 1 else ''}"
                    st.info(f"🧠 **Weakness Memory:** Based on your {_wm_label}, your weakest recurring skill is **{_wm_skill}** (avg: {_wm_score:.2f}/10). Questions will be biased toward improving this.")
                else:
                    # Graceful fallback for first-time users with no interview history
                    st.markdown(
                        """
                        <div style="
                            background: linear-gradient(135deg, rgba(79,163,227,0.10) 0%, rgba(56,189,248,0.06) 100%);
                            border: 1px solid rgba(79,163,227,0.25);
                            border-radius: 12px;
                            padding: 14px 18px;
                            margin-bottom: 14px;
                            display: flex;
                            align-items: flex-start;
                            gap: 12px;
                            font-family: -apple-system, BlinkMacSystemFont, 'SF Pro Display', 'DM Sans', sans-serif;
                        ">
                            <span style="font-size:1.4rem; line-height:1;">🎉</span>
                            <div>
                                <p style="margin:0 0 4px 0; font-weight:600; color:#7dd3fc; font-size:0.92rem;">
                                    Welcome to AI Interview Coach!
                                </p>
                                <p style="margin:0; color:#94a3b8; font-size:0.85rem; line-height:1.5;">
                                    This is your first mock interview — great time to start! Complete a session and the coach will automatically
                                    remember your weak areas and personalise future questions to help you improve faster.
                                </p>
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                col1, col2 = st.columns(2)

                with col1:
                    interview_type = st.selectbox(
                        "Interview Type",
                        options=["technical", "behavioral", "mixed"],
                        format_func=lambda x: x.title() + (" (Technical + Behavioral)" if x == "mixed" else ""),
                        key="dynamic_interview_type_select"
                    )

                with col2:
                    interview_difficulty = st.selectbox(
                        "Interview Difficulty",
                        options=["Easy", "Medium", "Hard"],
                        key="interview_difficulty_select",
                        index=1
                    )

                col3, col4 = st.columns(2)
                with col3:
                    num_questions = st.slider("Number of questions:", 5, 10, 6)

                with col4:
                    timer_seconds = st.slider("Time per question (seconds):", 60, 300, 120, step=30)

                # ── DOMAIN AUTHORITY: Show mismatch warning if resume ≠ selected domain ──
                if st.session_state.get("resume_context"):
                    _rc = st.session_state.resume_context
                    _resume_techs = " ".join(_rc.get("technologies", []) + _rc.get("skills", [])).lower()
                    _domain_cfg = get_domain_config(selected_domain)
                    _forbidden = _domain_cfg.get("forbidden_resume_keywords", [])
                    _mandatory = _domain_cfg.get("mandatory_topics", [])
                    _has_mismatch = any(kw.lower() in _resume_techs for kw in _forbidden)
                    _matched_forbidden = [kw for kw in _forbidden if kw.lower() in _resume_techs]

                    # Always show domain scope card; escalate to warning if mismatch detected
                    _context_note = _domain_cfg.get("context_override", "")

                    if _has_mismatch:
                        # Domain override is active — topics come purely from domain config
                        # (resume technologies are suppressed because they conflict with the
                        # selected domain, so we pass an empty resume context to the function).
                        _key_topics = generate_key_topics({}, _domain_cfg, selected_role)

                        # Identify which resume skills are being suppressed
                        _suppressed = list(dict.fromkeys(
                            kw for kw in _matched_forbidden
                            if any(kw.lower() in s.lower() for s in (_rc.get("technologies", []) + _rc.get("skills", [])))
                        ))[:4]
                        _suppressed_str = (
                            "".join(
                                f'<span style="background:rgba(244,67,54,0.15);color:#ef9a9a;'
                                f'border:1px solid rgba(244,67,54,0.3);border-radius:4px;'
                                f'padding:2px 8px;font-size:11px;margin:2px 3px;display:inline-block;">'
                                f'{kw}</span>'
                                for kw in _suppressed
                            )
                            if _suppressed else
                            '<span style="color:#aaa;font-size:12px;">none detected in top skills</span>'
                        )
                        _domain_pills = "".join(
                            f'<span style="background:rgba(56,189,248,0.12);color:#38bdf8;'
                            f'border:1px solid rgba(0,195,255,0.25);border-radius:4px;'
                            f'padding:2px 8px;font-size:11px;margin:2px 3px;display:inline-block;">'
                            f'{t}</span>'
                            for t in _key_topics
                        )
                        st.markdown(f"""
                        <div style="background:linear-gradient(135deg,rgba(255,152,0,0.08) 0%,rgba(255,87,34,0.06) 100%);
                                    border:1px solid rgba(255,152,0,0.35);border-left:4px solid #ff9800;
                                    border-radius:10px;padding:16px 20px;margin:10px 0;">
                            <div style="display:flex;align-items:center;gap:10px;margin-bottom:10px;">
                                <span style="font-size:20px;">🔄</span>
                                <div>
                                    <strong style="color:#ffb74d;font-size:15px;">Domain Override Active</strong>
                                    <span style="color:#aaa;font-size:12px;margin-left:8px;">
                                        Career pivot simulation enabled
                                    </span>
                                </div>
                            </div>
                            <p style="color:#e0e0e0;font-size:13px;margin:0 0 10px 0;line-height:1.6;">
                                Your resume contains skills outside <strong style="color:#ffb74d;">{selected_domain}</strong>.
                                All questions will be strictly scoped to your <em>target domain</em>, regardless of your
                                existing background. This mirrors what a real interviewer would focus on when you apply
                                to a new domain.
                            </p>
                            <div style="margin-bottom:10px;">
                                <span style="color:#ef9a9a;font-size:11px;font-weight:600;text-transform:uppercase;
                                            letter-spacing:0.06em;">Resume skills excluded from question scope:</span><br/>
                                <div style="margin-top:5px;">{_suppressed_str}</div>
                            </div>
                            <div>
                                <span style="color:#38bdf8;font-size:11px;font-weight:600;text-transform:uppercase;
                                            letter-spacing:0.06em;">Questions will draw from these topics:</span><br/>
                                <div style="margin-top:5px;">{_domain_pills}
                                    <span style="color:#aaa;font-size:11px;margin-left:4px;">
                                        + {max(0, len(_mandatory) - len(_key_topics))} more domain topics
                                    </span>
                                </div>
                            </div>
                            <p style="color:#aaa;font-size:11px;margin:10px 0 0 0;font-style:italic;">
                                💡 Treat this as authentic interview prep for breaking into {selected_domain}.
                                Focus on fundamentals, not your existing stack.
                            </p>
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        # Resume aligns with domain — blend resume content with domain topics
                        _key_topics = generate_key_topics(_rc, _domain_cfg, selected_role)
                        _domain_pills = "".join(
                            f'<span style="background:rgba(56,189,248,0.10);color:#38bdf8;'
                            f'border:1px solid rgba(0,195,255,0.2);border-radius:4px;'
                            f'padding:2px 8px;font-size:11px;margin:2px 3px;display:inline-block;">'
                            f'{t}</span>'
                            for t in _key_topics
                        )
                        st.markdown(f"""
                        <div style="background:rgba(0,195,255,0.05);border:1px solid rgba(0,195,255,0.2);
                                    border-left:4px solid #38bdf8;border-radius:10px;
                                    padding:14px 18px;margin:10px 0;">
                            <div style="display:flex;align-items:center;gap:8px;margin-bottom:8px;">
                                <span style="font-size:18px;">✅</span>
                                <strong style="color:#38bdf8;font-size:14px;">
                                    Domain Aligned — {selected_domain}
                                </strong>
                            </div>
                            <p style="color:#ccc;font-size:12px;margin:0 0 8px 0;line-height:1.5;">
                                Your resume aligns with the selected domain. Questions will leverage your
                                background and probe for <strong style="color:#e0e0e0;">depth and decision-making</strong>,
                                not just familiarity.
                            </p>
                            <div>
                                <span style="color:#aaa;font-size:11px;font-weight:600;text-transform:uppercase;
                                            letter-spacing:0.06em;">Key topics in scope:</span><br/>
                                <div style="margin-top:5px;">{_domain_pills}
                                    <span style="color:#666;font-size:11px;margin-left:4px;">
                                        + {max(0, len(_mandatory) - len(_key_topics))} more
                                    </span>
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)

                # ── DIFFICULTY CONTRACT: Show what each level means ──
                _diff_contract = DIFFICULTY_CONTRACTS.get(interview_difficulty, {})
                if _diff_contract:
                    _diff_colors = {"Easy": "#69f0ae", "Medium": "#ffcc02", "Hard": "#f44336"}
                    _diff_icons = {"Easy": "📗", "Medium": "📙", "Hard": "📕"}
                    _dc = _diff_colors.get(interview_difficulty, "#aaa")
                    _di = _diff_icons.get(interview_difficulty, "📋")
                    _scope = _diff_contract.get("answer_scope", "")
                    _cog = _diff_contract.get("cognitive_load_detail", _diff_contract.get("cognitive_load", ""))
                    _desc = _diff_contract.get("description", "")
                    st.markdown(
                        f'<div style="background:rgba(0,195,255,0.07);border-left:4px solid {_dc};'
                        f'padding:12px 16px;border-radius:0 8px 8px 0;margin:8px 0;">'
                        f'<strong style="color:{_dc};font-size:15px;">{_di} {interview_difficulty} Mode — {_diff_contract.get("label","")}</strong><br/>'
                        f'<span style="color:#ddd;font-size:13px;">{_desc}</span><br/>'
                        f'<span style="color:#aaa;font-size:12px;margin-top:4px;display:block;">'
                        f'Expected answer scope: <strong style="color:{_dc}">{_scope}</strong> &nbsp;|&nbsp; {_cog}'
                        f'</span>'
                        f'</div>',
                        unsafe_allow_html=True
                    )

                # ── AI Coach quota badge ──────────────────────────────────────────────
                _ac_username = st.session_state.get("username")
                if _ac_username:
                    _ac_used = get_usage_count_last_hour(_ac_username, "ai_coach")
                    _ac_remaining = max(0, 2 - _ac_used)
                    _ac_color = "#34d399" if _ac_remaining > 0 else "#fb7185"
                    _ac_svg = (
                        '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" '
                        'stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" '
                        'style="display:inline-block;vertical-align:middle;margin-right:5px;">'
                        '<rect x="3" y="3" width="7" height="7"/>'
                        '<rect x="14" y="3" width="7" height="7"/>'
                        '<rect x="14" y="14" width="7" height="7"/>'
                        '<rect x="3" y="14" width="7" height="7"/>'
                        '</svg>'
                    )
                    st.markdown(
                        f'<div style="display:flex;align-items:center;font-size:0.78rem;color:{_ac_color};margin-bottom:8px;font-family:-apple-system,sans-serif;">'
                        f'{_ac_svg} AI Coach: <b style="margin-left:3px;">{_ac_remaining}/2</b>&nbsp;mock interviews remaining this hour</div>',
                        unsafe_allow_html=True
                    )
                # ─────────────────────────────────────────────────────────────────────

                if st.button("🚀 Start Mock Interview"):
                    # ── Usage gate — only CHECK limit here, do NOT record yet ──
                    # Usage is recorded on first answer submission (standard approach).
                    # This means accidental refreshes before answering don't burn a slot.
                    _ac_gate_user = st.session_state.get("username")
                    _ac_already_recorded = st.session_state.get("_ac_usage_recorded_this_session", False)
                    if _ac_gate_user and not _ac_already_recorded:
                        _ac_allowed, _ac_msg = check_and_gate_feature(_ac_gate_user, "ai_coach")
                        if not _ac_allowed:
                            st.markdown(_ac_msg, unsafe_allow_html=True)
                            st.markdown(
                                '<div style="display:flex;align-items:center;font-size:0.88rem;color:#7dd3fc;'
                                'background:rgba(56,189,248,0.08);border:1px solid rgba(56,189,248,0.2);'
                                'border-radius:8px;padding:10px 14px;margin-top:8px;font-family:-apple-system,sans-serif;">'
                                '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" '
                                'stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" '
                                'style="display:inline-block;vertical-align:middle;margin-right:8px;flex-shrink:0;">'
                                '<circle cx="12" cy="12" r="10"/>'
                                '<polyline points="12 6 12 12 16 14"/>'
                                '</svg>'
                                'Your interview limit resets on a rolling 60-minute window.</div>',
                                unsafe_allow_html=True
                            )
                            st.stop()
                    # ─────────────────────────────────────────────────────────────────
                    with st.spinner("Generating personalised interview questions..."):
                        _username_for_bias = st.session_state.get("username", "Guest")
                        _weakness_data = get_user_weakness_history(_username_for_bias)
                        _bias = _weakness_data.get("bias", "balanced")

                        _resume_raw = st.session_state.get("resume_raw_text", "")
                        _num_resume_qs = 2 if _resume_raw else 0
                        _num_generic_qs = num_questions - _num_resume_qs

                        # resume_context already populated at upload time —
                        # pass it directly; no re-analysis needed.
                        merged = analyze_resume_and_generate_questions(
                            resume_text=_resume_raw,
                            role=selected_role,
                            domain=selected_domain,
                            difficulty=interview_difficulty,
                            interview_type=interview_type,
                            num_resume_qs=_num_resume_qs,
                            num_generic_qs=_num_generic_qs,
                            weakness_bias=_bias,
                        )

                        # Keep existing resume_context (set at upload); only update questions
                        if not st.session_state.get("resume_context"):
                            st.session_state.resume_context = merged["resume_context"]
                        resume_based_qs = merged["resume_questions"] if _resume_raw else []
                        generic_qs = merged["generic_questions"]

                        # Combine all questions: resume-based first, then generic
                        all_questions = resume_based_qs + generic_qs
                        all_questions = all_questions[:num_questions]

                        if all_questions:
                            # Reset ALL interview state variables properly
                            import uuid
                            st.session_state.current_interview_id = str(uuid.uuid4())
                            st.session_state.question_db_ids = []
                            st.session_state.dynamic_interview_questions = all_questions
                            st.session_state.resume_based_questions = resume_based_qs
                            st.session_state.generic_questions = generic_qs
                            st.session_state.original_num_questions = num_questions
                            st.session_state.current_dynamic_interview_question = 0
                            st.session_state.dynamic_interview_answers = []
                            st.session_state.dynamic_interview_scores = []
                            st.session_state.dynamic_interview_feedbacks = []
                            st.session_state.dynamic_interview_completed = False
                            st.session_state.dynamic_interview_started = True
                            st.session_state.interview_actual_start_time = time.time()
                            st.session_state.dynamic_answer_submitted = False
                            st.session_state.current_interview_question_text = all_questions[0]
                            # ── TIMER FIX: do NOT start timer here. The timer starts
                            # on the first render of the interview page (below), AFTER
                            # st.rerun() fires and the page is actually shown to the user.
                            # Starting it here causes the animation + rerun latency
                            # (~3-5 s) to be silently consumed before the user sees 5:00.
                            st.session_state.question_timer_start = None
                            st.session_state._timer_needs_reset = True
                            st.session_state.timer_seconds = timer_seconds
                            st.session_state.interview_difficulty = interview_difficulty
                            st.session_state.interview_mode = interview_type
                            st.session_state.interview_phase = "resume" if resume_based_qs else "generic"
                            # PART 4: Escalation ladder tracking
                            st.session_state.escalation_layer = 1
                            st.session_state.follow_up_count = 0
                            st.session_state.follow_up_strategy = "Depth Probe"

                            # Show resume scanning animation if resume questions exist
                            if resume_based_qs:
                                st.info("🎯 Starting with resume-based questions...")
                                show_resume_scanning_animation()

                            st.success("Questions generated! Starting your mock interview...")
                            st.rerun()  # FIX 5: removed time.sleep(1)
                        else:
                            st.error("Failed to generate questions. Please try again.")
            
            # Interview in progress
            elif st.session_state.dynamic_interview_started and not st.session_state.dynamic_interview_completed:
                # CRITICAL FIX: Properly count answered questions
                questions_answered = len(st.session_state.dynamic_interview_answers)
                total_questions = len(st.session_state.dynamic_interview_questions)
                current_index = st.session_state.current_dynamic_interview_question + 1

                # Determine current phase
                num_resume_qs = len(st.session_state.resume_based_questions)
                current_phase = "Resume-Based" if current_index <= num_resume_qs else "Generic Interview"

                # Display progress with correct counts in glassmorphism box
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, rgba(0, 195, 255, 0.08) 0%, rgba(0, 195, 255, 0.04) 100%);
                            backdrop-filter: blur(10px);
                            -webkit-backdrop-filter: blur(10px);
                            border: 1px solid rgba(0, 195, 255, 0.2);
                            border-radius: 12px;
                            padding: 16px 24px;
                            margin: 20px 0;
                            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1), inset 0 1px 0 rgba(255, 255, 255, 0.05);">
                    <p style="color: #ffffff; font-size: 16px; margin: 0; font-weight: 500;">
                        📊 Progress: Answered {questions_answered}/{st.session_state.original_num_questions} questions | Phase: {current_phase}
                    </p>
                </div>
                """, unsafe_allow_html=True)

                # FIX 6: Early completion guard — resolve before fragment renders
                if questions_answered >= st.session_state.original_num_questions and not st.session_state.dynamic_interview_completed:
                    if st.session_state.get('interview_actual_start_time'):
                        st.session_state.interview_final_duration_seconds = int(time.time() - st.session_state.interview_actual_start_time)
                    else:
                        st.session_state.interview_final_duration_seconds = None
                    st.session_state.interview_result_saved = False
                    st.session_state.dynamic_interview_completed = True
                    st.rerun()

                if questions_answered < st.session_state.original_num_questions:
                    question = st.session_state.current_interview_question_text or st.session_state.dynamic_interview_questions[st.session_state.current_dynamic_interview_question]

                    # TIMER FIX: Start timer on first render of this question.
                    # _timer_needs_reset is set True by the "Start Interview" button
                    # so the clock only begins when the page is actually visible to
                    # the user — not during the setup / animation / rerun cycle.
                    if st.session_state.question_timer_start is None or \
                            st.session_state.get("_timer_needs_reset", False):
                        st.session_state.question_timer_start = time.time()
                        st.session_state._timer_needs_reset = False

                    # ── Calculate remaining time (server-side, passed to JS) ──
                    elapsed_time   = time.time() - st.session_state.question_timer_start
                    remaining_time = max(0, st.session_state.timer_seconds - elapsed_time)
                    _q_idx_now     = st.session_state.current_dynamic_interview_question
                    _submitted_now = st.session_state.get("dynamic_answer_submitted", False)

                    # ── Pure-JS browser timer (zero server load per tick) ────────
                    # The countdown runs entirely in the user's browser via JS
                    # setInterval — no server thread wakes up every second.
                    # When it hits zero the JS clicks the hidden button below,
                    # which triggers a normal Streamlit interaction → auto-submit.
                    _render_js_timer(
                        remaining_seconds=remaining_time,
                        total_seconds=st.session_state.timer_seconds,
                        submitted=_submitted_now,
                        q_idx=_q_idx_now,
                    )

                    # ── Question card (rendered by server, NOT inside a fragment) ─
                    _answered_now = len(st.session_state.get("dynamic_interview_answers", []))
                    _total_q_now  = st.session_state.get("original_num_questions", 1)
                    _num_res_now  = len(st.session_state.get("resume_based_questions", []))
                    _phase_badge  = "📄 Resume-Based Question" if (_q_idx_now + 1) <= _num_res_now else "💼 Generic Interview Question"
                    _role_now     = st.session_state.get("interview_role", "")
                    _diff_now     = st.session_state.get("interview_difficulty", "")
                    st.markdown(f"""
                    <div class="quiz-card">
                        <h3 style="color:#38bdf8;font-family:-apple-system,BlinkMacSystemFont,'SF Pro Display',sans-serif;font-weight:600;letter-spacing:-0.02em;">Question {_answered_now + 1} of {_total_q_now}</h3>
                        <div style="background:rgba(56,189,248,0.10);padding:6px 12px;border-radius:99px;margin:10px 0;display:inline-block;border:1px solid rgba(56,189,248,0.22);">
                            <span style="color:#38bdf8;font-weight:600;font-size:0.8rem;letter-spacing:0.03em;text-transform:uppercase;">{_phase_badge}</span>
                        </div>
                        <h4 style="color:#94a3b8;font-family:-apple-system,BlinkMacSystemFont,'SF Pro Display',sans-serif;font-weight:500;font-size:0.875rem;margin:12px 0;letter-spacing:0.02em;">Role: {_role_now} | Difficulty: {_diff_now}</h4>
                        <p style="font-size:1rem;color:#f0f4f8;font-family:-apple-system,BlinkMacSystemFont,'SF Pro Display',sans-serif;line-height:1.6;margin:14px 0;">{question}</p>
                    </div>
                    """, unsafe_allow_html=True)

                    # ── Background-thread rerun trigger ──────────────────────────
                    # All JS-based approaches (button click, postMessage, location.reload)
                    # fail because st.components iframes are cross-origin sandboxed, and
                    # location.reload() navigates away from the app entirely.
                    #
                    # Solution: spawn a daemon thread that sleeps until remaining_time,
                    # then sets _timer_expired in session_state and calls st.rerun()
                    # via Streamlit's runtime API. This is 100% server-side — no JS needed.
                    # The thread is keyed to (_q_idx_now, question_timer_start) so it
                    # spawns only once per question, not on every rerun.
                    _thread_key = f"_timer_thread_armed_{_q_idx_now}"
                    if (not _submitted_now
                            and remaining_time > 0
                            and not st.session_state.get(_thread_key, False)):
                        st.session_state[_thread_key] = True

                        def _expire_timer(sleep_secs, session_id):
                            import time as _t
                            _t.sleep(sleep_secs)
                            try:
                                from streamlit.runtime import get_instance
                                from streamlit.runtime.scriptrunner import add_script_run_ctx
                                runtime = get_instance()
                                session_info = runtime._session_mgr.get_session_info(session_id)
                                if session_info is not None:
                                    session_info.session.request_rerun(None)
                            except Exception:
                                pass  # session may have ended; silently ignore

                        import threading as _threading
                        _sid = st.runtime.scriptrunner.get_script_run_ctx().session_id
                        _t = _threading.Thread(
                            target=_expire_timer,
                            args=(remaining_time + 0.5, _sid),
                            daemon=True,
                        )
                        _t.start()

                    # Edge case: if timer already expired on this render (e.g. user
                    # navigated away and came back) but no thread is running to trigger
                    # the rerun — force one immediately so the auto-submit block fires.
                    elif (not _submitted_now
                            and remaining_time <= 0
                            and not st.session_state.get(_thread_key, False)):
                        st.session_state[_thread_key] = True
                        st.rerun()


                    # Refresh button — always visible, right-aligned, small
                    st.markdown("""
                    <style>
                    div[data-testid="stButton"]:has(button[data-testid="refresh_btn"]) {
                        display: flex; justify-content: flex-end;
                    }
                    button[data-testid="refresh_btn"] {
                        padding: 4px 14px !important;
                        font-size: 0.75rem !important;
                        height: auto !important;
                        min-height: 0 !important;
                        background: rgba(56,189,248,0.08) !important;
                        border: 1px solid rgba(56,189,248,0.25) !important;
                        color: #38bdf8 !important;
                        border-radius: 6px !important;
                    }
                    </style>
                    """, unsafe_allow_html=True)
                    if st.button("🔄 Refresh Interview", key="refresh_btn", help="Restart interview from scratch"):
                        st.session_state.dynamic_interview_questions = []
                        st.session_state.current_dynamic_interview_question = 0
                        st.session_state.dynamic_interview_answers = []
                        st.session_state.dynamic_interview_scores = []
                        st.session_state.dynamic_interview_feedbacks = []
                        st.session_state.dynamic_interview_completed = False
                        st.session_state.dynamic_interview_started = False
                        st.session_state.dynamic_answer_submitted = False
                        st.session_state.current_interview_question_text = ""
                        st.session_state.question_timer_start = None
                        st.session_state.interview_result_saved = False
                        st.session_state.interview_final_duration_seconds = None
                        st.session_state.interview_actual_start_time = None
                        st.session_state.pending_followup_display = ""
                        st.session_state.pending_followup_strategy = ""
                        st.session_state.escalation_layer = 1
                        st.session_state.follow_up_count = 0
                        st.session_state.current_interview_id = None
                        st.session_state.question_db_ids = []
                        st.session_state.pop("_timer_expired", None)
                        st.session_state.pop("_timer_expired_answer", None)
                        # Clear all thread-armed flags on full refresh
                        for _k in [k for k in st.session_state if k.startswith("_timer_thread_armed_")]:
                            st.session_state.pop(_k, None)
                        # Reset usage flag so the next interview is properly counted
                        st.session_state._ac_usage_recorded_this_session = False
                        st.rerun()

                    # Answer input with character limit
                    answer_key = f"dynamic_interview_answer_{st.session_state.current_dynamic_interview_question}"
                    answer = st.text_area(
                        "Your answer:",
                        placeholder="Type your detailed answer here... (Use STAR method: Situation, Task, Action, Result)",
                        height=150,
                        max_chars=2000,
                        key=answer_key,
                        help="Maximum 2000 characters"
                    )

                    # ── SINGLE helper: evaluate + inject follow-up (called from both submit paths) ──
                    def _process_submission(ans_text, q_text, q_idx, n_answered):
                        """
                        Evaluate the answer, store results, and inject the follow-up question
                        into the question list.  The exact same follow-up text is stored in
                        session_state.pending_followup_display so the preview shown to the user
                        is always identical to the question that will appear next.

                        ARCHITECTURE FIX: Every answered question is immediately saved to the
                        interview_questions DB table so the PDF can use it as single source of truth.
                        FIX 8: Idempotency guard — bail out immediately if this question index
                        has already been processed (prevents double-submission on rapid reruns).
                        """
                        # FIX 8: Idempotency check using answered count vs question index
                        if len(st.session_state.dynamic_interview_answers) > q_idx:
                            return  # Already processed this question index — do not re-evaluate

                        # ── Record usage on FIRST answer only (standard approach) ──
                        # Refresh before answering = no usage consumed.
                        # Refresh after answering 1+ questions = usage already counted.
                        if q_idx == 0 and not st.session_state.get("_ac_usage_recorded_this_session", False):
                            _usage_user = st.session_state.get("username")
                            if _usage_user:
                                record_feature_usage(_usage_user, "ai_coach")
                                st.session_state._ac_usage_recorded_this_session = True
                        # ─────────────────────────────────────────────────────────

                        diff = st.session_state.interview_difficulty
                        eval_res = evaluate_interview_answer_for_scores(
                            ans_text, q_text, diff,
                            role=selected_role, domain=selected_domain
                        )

                        st.session_state.dynamic_interview_answers.append(ans_text)
                        st.session_state.dynamic_interview_scores.append(eval_res)
                        st.session_state.dynamic_interview_feedbacks.append(eval_res["feedback"])
                        st.session_state.dynamic_answer_submitted = True
                        st.session_state.pending_followup_display = ""   # reset
                        st.session_state.pending_followup_strategy = ""

                        # ── IMMEDIATELY save to DB (single source of truth for PDF) ──
                        interview_id = st.session_state.get('current_interview_id')
                        parent_db_id = None
                        is_fu = False
                        # Determine if this is a follow-up: index beyond original questions
                        original_count = len(st.session_state.get('resume_based_questions', [])) + len(st.session_state.get('generic_questions', []))
                        if q_idx >= original_count and len(st.session_state.question_db_ids) > 0:
                            # It's a follow-up — find the parent: the main question that triggered it
                            # The parent is the last main question before this follow-up
                            # We store follow-ups linked to the most recent main question db id
                            parent_db_id = st.session_state.question_db_ids[-1]
                            is_fu = True

                        db_row_id = -1
                        if interview_id:
                            score_to_save = dict(eval_res)
                            db_row_id = save_interview_question(
                                interview_id=interview_id,
                                question_text=q_text,
                                answer_text=ans_text,
                                difficulty=diff,
                                is_follow_up=is_fu,
                                parent_question_id=parent_db_id,
                                score_breakdown=score_to_save,
                                question_order=q_idx,
                            )
                        # Track db row id - only for main questions (used as parent for follow-ups)
                        if not is_fu and db_row_id != -1:
                            st.session_state.question_db_ids.append(db_row_id)

                        can_add_followup = n_answered < st.session_state.original_num_questions - 1

                        if diff == "Hard" and can_add_followup:
                            # ── Hard mode: use adaptive engine (single source of truth) ──
                            weakness_data = analyze_answer_weaknesses(ans_text, eval_res)
                            strategy = weakness_data["strategy"]
                            layer = getattr(st.session_state, 'escalation_layer', 1)
                            followup_q = generate_adaptive_followup(
                                q_text, ans_text, strategy, layer, selected_role, selected_domain
                            )
                            followup_q = followup_q.strip() if followup_q else ""
                            if followup_q:
                                st.session_state.dynamic_interview_questions.insert(
                                    q_idx + 1, followup_q
                                )
                                st.session_state.follow_up_count = getattr(st.session_state, 'follow_up_count', 0) + 1
                                st.session_state.escalation_layer = min(5, layer + 1)
                                st.session_state.follow_up_strategy = strategy
                                # ★ Store SAME text for preview ★
                                st.session_state.pending_followup_display = followup_q
                                st.session_state.pending_followup_strategy = strategy

                        elif diff in ("Easy", "Medium") and can_add_followup:
                            # ── Easy/Medium: only inject if LLM returned a valid followup ──
                            # The evaluation prompt does NOT ask for a follow-up for Easy/Medium,
                            # so eval_res["followup"] is always "".  We deliberately do NOT inject
                            # anything — this prevents mismatched questions.
                            pass   # No follow-up for Easy/Medium

                        return eval_res

                    # ── initialise session key on first load ──
                    if 'pending_followup_display' not in st.session_state:
                        st.session_state.pending_followup_display = ""
                    if 'pending_followup_strategy' not in st.session_state:
                        st.session_state.pending_followup_strategy = ""

                    # ── Auto-submit: fires on the thread-triggered rerun ──────────
                    # When the background thread calls session.request_rerun(), this
                    # block runs and _fresh_remaining is <= 0 → auto-submit fires.
                    # Legacy _timer_expired keys cleaned up for safety.
                    st.session_state.pop("_timer_expired", None)
                    st.session_state.pop("_timer_expired_answer", None)
                    _fresh_elapsed   = time.time() - st.session_state.question_timer_start if st.session_state.question_timer_start else 0
                    _fresh_remaining = max(0, st.session_state.timer_seconds - _fresh_elapsed)
                    if _fresh_remaining <= 0 and not st.session_state.dynamic_answer_submitted:
                        st.session_state.dynamic_answer_submitted = True  # set FIRST — prevents double-submission
                        _auto_answer = answer.strip() if answer.strip() else "⚠️ No Answer"
                        with st.spinner("⏰ Time's up! Evaluating your answer..."):
                            _process_submission(
                                _auto_answer, question,
                                st.session_state.current_dynamic_interview_question,
                                questions_answered
                            )
                        st.warning("⏰ Time's up! Answer auto-submitted.")
                        st.rerun()

                    # Submit answer button — shown whenever answer not yet submitted
                    if not st.session_state.dynamic_answer_submitted:
                        if st.button("Submit Answer & Get Feedback"):
                            if answer.strip():
                                with st.spinner("Evaluating your answer..."):
                                    _process_submission(
                                        answer, question,
                                        st.session_state.current_dynamic_interview_question,
                                        questions_answered
                                    )
                                st.rerun()
                            else:
                                st.warning("Please provide an answer before proceeding.")

                    # Show feedback after answer submitted
                    if st.session_state.dynamic_answer_submitted:
                        current_score_dict = st.session_state.dynamic_interview_scores[-1]
                        avg_q_score = (current_score_dict["knowledge"] + current_score_dict["communication"] + current_score_dict["relevance"]) / 3

                        # Format feedback for display
                        feedback_text = current_score_dict["feedback"] if isinstance(current_score_dict["feedback"], str) else chr(10).join(current_score_dict["feedback"])
                        formatted_feedback = format_feedback_text(feedback_text)

                        st.markdown(f"""
                        <div style="background: linear-gradient(135deg, rgba(0, 195, 255, 0.1) 0%, rgba(0, 195, 255, 0.05) 100%);
                                    border: 1px solid rgba(0, 195, 255, 0.3); border-radius: 10px; padding: 15px; margin: 15px 0;">
                            <h4 style="color:#38bdf8;font-family:-apple-system,BlinkMacSystemFont,'SF Pro Display',sans-serif;font-weight:600;letter-spacing:-0.02em;">Immediate Feedback:</h4>
                            <p style="color: #ffffff;">📊 Knowledge: {current_score_dict["knowledge"]}/10 | Communication: {current_score_dict["communication"]}/10 | Relevance: {current_score_dict["relevance"]}/10</p>
                            <p style="color: #ffffff;">⭐ Question Score: {avg_q_score:.2f}/10</p>
                            <div style="color: #ffffff; margin-top: 10px;">
                                {formatted_feedback}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)

                        # ★ Show follow-up preview using SAME text that was injected ★
                        _preview_fq = st.session_state.get('pending_followup_display', '')
                        _preview_strategy = st.session_state.get('pending_followup_strategy', '')
                        if st.session_state.interview_difficulty == "Hard" and _preview_fq:
                            _esc_layer = st.session_state.get("escalation_layer", 1)
                            _layer_info = ESCALATION_LAYER_MAP.get(_esc_layer, {})
                            _layer_name = _layer_info.get("name", "")
                            _pressure = _layer_info.get("cognitive_pressure", "")
                            _pressure_colors = {
                                "LOW": "#69f0ae", "MEDIUM": "#ffcc02",
                                "MEDIUM-HIGH": "#ff9800", "HIGH": "#ff5722", "MAXIMUM": "#f44336"
                            }
                            _pc = _pressure_colors.get(_pressure, "#ffa500")
                            st.markdown(f"""
                            <div style="background: linear-gradient(135deg, rgba(255,165,0,0.12), rgba(255,165,0,0.06));
                                        border: 1px solid rgba(255,165,0,0.4); border-radius: 10px;
                                        padding: 14px 18px; margin: 12px 0;">
                                <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:8px;">
                                    <span style="color: #ffa500; font-weight: 600;">
                                        🔎 Follow-Up — {_preview_strategy}
                                    </span>
                                    <span style="color:{_pc};font-size:12px;font-weight:600;
                                                 background:rgba(0,0,0,0.3);padding:2px 8px;border-radius:12px;">
                                        Layer {_esc_layer}/5: {_layer_name} | Pressure: {_pressure}
                                    </span>
                                </div>
                                <p style="color: #ffffff; margin: 0; font-size: 15px;">{_preview_fq}</p>
                            </div>
                            """, unsafe_allow_html=True)

                        # Continue/Complete button
                        # CRITICAL FIX: Check if we've answered all original questions
                        if questions_answered >= st.session_state.original_num_questions:
                            # All questions answered, mark as complete
                            if st.button("Complete Interview 🏁"):
                                # Capture exact duration at completion moment
                                if st.session_state.get('interview_actual_start_time'):
                                    st.session_state.interview_final_duration_seconds = int(time.time() - st.session_state.interview_actual_start_time)
                                else:
                                    st.session_state.interview_final_duration_seconds = None
                                st.session_state.interview_result_saved = False
                                st.session_state.dynamic_interview_completed = True
                                st.rerun()
                        else:
                            # More questions to go
                            if st.button("Continue to Next Question ➡️"):
                                st.session_state.current_dynamic_interview_question += 1
                                st.session_state.dynamic_answer_submitted = False
                                st.session_state.pending_followup_display = ""
                                st.session_state.pending_followup_strategy = ""
                                st.session_state.pop("_timer_expired", None)
                                st.session_state.pop("_timer_expired_answer", None)
                                # Clear the thread-armed flag so a new thread spawns for next question
                                _prev_idx = st.session_state.current_dynamic_interview_question - 1
                                st.session_state.pop(f"_timer_thread_armed_{_prev_idx}", None)
                                if st.session_state.current_dynamic_interview_question < len(st.session_state.dynamic_interview_questions):
                                    st.session_state.current_interview_question_text = st.session_state.dynamic_interview_questions[st.session_state.current_dynamic_interview_question]
                                else:
                                    # Safety check - if we're out of questions but haven't answered all, generate one
                                    st.session_state.current_interview_question_text = f"Additional question for {selected_role}"
                                # TIMER FIX: Mark timer for reset — it will start on the
                                # next render after rerun, not here on the button click,
                                # so rerun latency does not eat into the question time.
                                st.session_state.question_timer_start = None
                                st.session_state._timer_needs_reset = True
                                st.rerun()

                    # Progress bar for interview completion
                    interview_progress = questions_answered / st.session_state.original_num_questions
                    st.markdown("### Interview Progress")
                    st.progress(interview_progress)

                    # CRITICAL FIX: Review Previous Answers - show all properly
                    if len(st.session_state.dynamic_interview_answers) > 0:
                        with st.expander("📖 Review Previous Answers"):
                            # Show all submitted answers
                            num_to_show = len(st.session_state.dynamic_interview_answers)
                            for i in range(num_to_show):
                                if i < len(st.session_state.dynamic_interview_questions) and i < len(st.session_state.dynamic_interview_scores):
                                    prev_question = st.session_state.dynamic_interview_questions[i]
                                    prev_answer = st.session_state.dynamic_interview_answers[i]
                                    prev_scores = st.session_state.dynamic_interview_scores[i]
                                    prev_avg = (prev_scores["knowledge"] + prev_scores["communication"] + prev_scores["relevance"]) / 3

                                    # Show full answer (up to 500 chars in review, full in final)
                                    answer_preview = prev_answer[:500]
                                    if len(prev_answer) > 500:
                                        answer_preview += "..."

                                    st.markdown(f"**Question {i+1}:** {prev_question}")
                                    st.markdown(f"**Your Answer:** {answer_preview}")
                                    st.markdown(f"**Score:** {prev_avg:.2f}/10")
                                    if i < num_to_show - 1:  # Don't add separator after last item
                                        st.markdown("---")

                    # NOTE: No more time.sleep(1) + st.rerun() here.
                    # The JS timer inside the components.html block above handles
                    # the visual countdown entirely in the browser. Auto-submit
                    # is triggered by the hidden __TIMER_EXPIRED__ button click.
                else:
                    # FIX 6 (fallback): early guard above handles this; this is a safety net
                    if not st.session_state.dynamic_interview_completed:
                        if st.session_state.get('interview_actual_start_time'):
                            st.session_state.interview_final_duration_seconds = int(time.time() - st.session_state.interview_actual_start_time)
                        else:
                            st.session_state.interview_final_duration_seconds = None
                        st.session_state.interview_result_saved = False
                        st.session_state.dynamic_interview_completed = True
                        st.rerun()
            
            # UNIFIED: Interview completed + Course Recommendations + DB + PDF
            elif st.session_state.dynamic_interview_completed:
                # Calculate average scores for each dimension
                knowledge_scores = [s["knowledge"] for s in st.session_state.dynamic_interview_scores]
                communication_scores = [s["communication"] for s in st.session_state.dynamic_interview_scores]
                relevance_scores = [s["relevance"] for s in st.session_state.dynamic_interview_scores]

                avg_knowledge = sum(knowledge_scores) / len(knowledge_scores)
                avg_communication = sum(communication_scores) / len(communication_scores)
                avg_relevance = sum(relevance_scores) / len(relevance_scores)
                overall_avg = (avg_knowledge + avg_communication + avg_relevance) / 3

                # PART 3: Compute weighted score using difficulty multiplier
                _raw_avg = overall_avg
                _weighted_avg = compute_weighted_score(_raw_avg, st.session_state.interview_difficulty)
                _follow_up_count = getattr(st.session_state, 'follow_up_count', 0)
                _depth_score = (avg_knowledge + avg_relevance) / 2

                # Determine badge based on overall average
                if overall_avg >= 8.5:
                    badge = "Interview Ready"
                    badge_emoji = "🏆"
                elif overall_avg >= 7.0:
                    badge = "Excellent"
                    badge_emoji = "🌟"
                elif overall_avg >= 5.0:
                    badge = "Good"
                    badge_emoji = "👍"
                else:
                    badge = "Needs Practice"
                    badge_emoji = "💪"

                st.markdown(f"""
                <div class="badge-container">
                    <h2 style="margin: 0; color: #ffffff; font-size: 28px; font-weight: 600;">🎉 Mock Interview Complete!</h2>
                    <div style="margin: 30px 0;">
                        <div class="score-display">{overall_avg:.2f}/10</div>
                        <h3 style="color: #ffffff; margin: 15px 0; font-size: 24px; font-weight: 500;">{badge_emoji} {badge}</h3>
                    </div>
                    <p style="color: rgba(255, 255, 255, 0.85); font-size: 16px; margin: 8px 0;">Role: {selected_role} in {selected_domain}</p>
                    <p style="color: rgba(255, 255, 255, 0.85); font-size: 16px; margin: 8px 0;">Difficulty: {st.session_state.interview_difficulty}</p>
                    <p style="color: rgba(0, 195, 255, 0.9); font-size: 15px; margin: 8px 0;">⚡ Weighted Score: {_weighted_avg:.2f}/10 (×{DIFFICULTY_MULTIPLIERS.get(st.session_state.interview_difficulty, 1.0)} difficulty multiplier)</p>
                    <p style="color: rgba(255, 255, 255, 0.7); font-size: 14px; margin: 4px 0;">Follow-up Probes: {_follow_up_count} | Depth Score: {_depth_score:.2f}/10</p>
                </div>
                """, unsafe_allow_html=True)

                # Create radar chart for skills
                st.markdown('<div class="radar-container">', unsafe_allow_html=True)
                st.subheader("📊 Performance Radar Chart")

                radar_data = {
                    "Communication": avg_communication,
                    "Knowledge": avg_knowledge,
                    "Confidence": avg_relevance
                }

                fig = go.Figure()
                fig.add_trace(go.Scatterpolar(
                    r=list(radar_data.values()),
                    theta=list(radar_data.keys()),
                    fill='toself',
                    name='Performance',
                    line=dict(color='#00c3ff', width=2),
                    fillcolor='rgba(0, 195, 255, 0.2)'
                ))

                fig.update_layout(
                    polar=dict(
                        radialaxis=dict(
                            visible=True,
                            range=[0, 10],
                            tickfont=dict(color='white', size=10),
                            gridcolor='rgba(255, 255, 255, 0.2)'
                        ),
                        angularaxis=dict(
                            tickfont=dict(color='white', size=12),
                            gridcolor='rgba(255, 255, 255, 0.2)'
                        ),
                        bgcolor='rgba(0, 0, 0, 0)'
                    ),
                    showlegend=False,
                    title=dict(
                        text="Interview Performance Metrics",
                        x=0.5,
                        font=dict(color='#00c3ff', size=16)
                    ),
                    paper_bgcolor='rgba(0, 0, 0, 0)',
                    plot_bgcolor='rgba(0, 0, 0, 0)',
                    font=dict(color='white'),
                    height=400
                )

                st.plotly_chart(fig, use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)

                # Strengths and Weaknesses
                st.subheader("💡 Performance Analysis")
                col1, col2 = st.columns(2)

                metrics = [("Communication", avg_communication), ("Knowledge", avg_knowledge), ("Confidence", avg_relevance)]
                metrics_sorted = sorted(metrics, key=lambda x: x[1], reverse=True)

                with col1:
                    st.markdown("**🌟 Strengths:**")
                    for name, score in metrics_sorted[:2]:
                        st.markdown(f"- {name}: {score:.2f}/10")

                with col2:
                    st.markdown("**📈 Areas to Improve:**")
                    for name, score in metrics_sorted[-2:]:
                        st.markdown(f"- {name}: {score:.2f}/10")

                # FIXED: Show detailed Q&A results with full answers and proper matching
                st.markdown("---")
                st.subheader("📋 Detailed Q&A Review:")

                # Ensure we only show as many Q&A pairs as we have complete data for
                num_complete_qa = min(
                    len(st.session_state.dynamic_interview_scores),
                    len(st.session_state.dynamic_interview_answers),
                    len(st.session_state.dynamic_interview_feedbacks),
                    len(st.session_state.dynamic_interview_questions)
                )

                for i in range(num_complete_qa):
                    score_dict = st.session_state.dynamic_interview_scores[i]
                    answer = st.session_state.dynamic_interview_answers[i]
                    feedback = st.session_state.dynamic_interview_feedbacks[i]
                    question = st.session_state.dynamic_interview_questions[i]

                    q_avg = (score_dict["knowledge"] + score_dict["communication"] + score_dict["relevance"]) / 3

                    with st.expander(f"Question {i+1}: Score {q_avg:.2f}/10"):
                        st.write(f"**Question:** {question}")
                        st.write(f"**Your Answer:** {answer}")  # Show full answer
                        st.write(f"**Scores:** Knowledge: {score_dict['knowledge']}/10 | Communication: {score_dict['communication']}/10 | Relevance: {score_dict['relevance']}/10")

                        # Format and display feedback as bullet points
                        feedback_text = "\n".join(feedback) if isinstance(feedback, list) else feedback
                        formatted_feedback = format_feedback_text(feedback_text)
                        st.markdown(formatted_feedback, unsafe_allow_html=True)

                # Save to database — guarded by flag so it only runs ONCE
                username = st.session_state.get("username", "Guest")
                feedback_summary = f"Strengths: {metrics_sorted[0][0]}, {metrics_sorted[1][0]}. Weaknesses: {metrics_sorted[-1][0]}, {metrics_sorted[-2][0]}."

                if not st.session_state.get('interview_result_saved', False):
                    # Capture duration at the exact moment of first save, not on reruns
                    _interview_duration = st.session_state.get('interview_final_duration_seconds', None)
                    _interview_mode = st.session_state.get('interview_mode', None)
                    # PART 6: Compute behavior class
                    _dur_mins = (_interview_duration / 60.0) if _interview_duration else None
                    _b_class = classify_behavior(_dur_mins, 0.0, None)
                    if save_interview_result(username, selected_role, selected_domain, overall_avg, st.session_state.original_num_questions, feedback_summary,
                                             knowledge_avg=avg_knowledge, communication_avg=avg_communication, relevance_avg=avg_relevance,
                                             difficulty=st.session_state.interview_difficulty, duration_seconds=_interview_duration,
                                             interview_mode=_interview_mode,
                                             weighted_score=_weighted_avg, raw_avg_score=_raw_avg,
                                             follow_up_count=_follow_up_count, depth_score=_depth_score, behavior_class=_b_class):
                        st.session_state.interview_result_saved = True
                        log_user_action(username, "completed_interview")

                # Generate PDF report
                st.markdown("---")
                st.subheader("📄 Download Interview Report")

                completed_on = get_ist_time()

                # CRITICAL FIX: Ensure all arrays have same length for PDF generation
                num_complete = min(
                    len(st.session_state.dynamic_interview_questions),
                    len(st.session_state.dynamic_interview_answers),
                    len(st.session_state.dynamic_interview_scores),
                    len(st.session_state.dynamic_interview_feedbacks)
                )

                pdf_bytes = generate_interview_pdf_report(
                    username,
                    selected_role,
                    selected_domain,
                    completed_on,
                    st.session_state.dynamic_interview_questions[:num_complete],
                    st.session_state.dynamic_interview_answers[:num_complete],
                    st.session_state.dynamic_interview_scores[:num_complete],
                    st.session_state.dynamic_interview_feedbacks[:num_complete],
                    overall_avg,
                    badge,
                    difficulty=st.session_state.interview_difficulty,
                    interview_id=st.session_state.get('current_interview_id')
                )

                if pdf_bytes:
                    st.download_button(
                        label="📄 Download Interview Report",
                        data=pdf_bytes,
                        file_name=f"interview_report_{username}_{selected_role.replace(' ', '_')}_{completed_on.split()[0]}.pdf",
                        mime="application/pdf"
                    )

                    # ── Auto-email the report to the user's registered address ──
                    _SVG_MAIL = (
                        '<svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" '
                        'viewBox="0 0 24 24" fill="none" stroke="#1a7f37" stroke-width="2" '
                        'style="vertical-align:middle;margin-right:6px;">'
                        '<path d="M4 4h16v16H4z"/><path d="m22 6-10 7L2 6"/></svg>'
                    )
                    _SVG_WARN = (
                        '<svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" '
                        'viewBox="0 0 24 24" fill="none" stroke="#b45309" stroke-width="2" '
                        'style="vertical-align:middle;margin-right:6px;">'
                        '<path d="M10.29 3.86 1.82 18a2 2 0 0 0 1.71 3h16.94a2 2 0 0 0 '
                        '1.71-3L13.71 3.86a2 2 0 0 0-3.42 0z"/><line x1="12" y1="9" x2="12" y2="13"/>'
                        '<line x1="12" y1="17" x2="12.01" y2="17"/></svg>'
                    )
                    _SVG_INFO = (
                        '<svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" '
                        'viewBox="0 0 24 24" fill="none" stroke="#0369a1" stroke-width="2" '
                        'style="vertical-align:middle;margin-right:6px;">'
                        '<circle cx="12" cy="12" r="10"/><line x1="12" y1="16" x2="12" y2="12"/>'
                        '<line x1="12" y1="8" x2="12.01" y2="8"/></svg>'
                    )

                    # Guard so the email only fires ONCE per completed interview —
                    # without this, any Streamlit rerun (widget clicks, expander
                    # toggles, etc.) re-executes this block and re-sends the mail.
                    _email_flag_key = f"interview_email_sent_{st.session_state.get('current_interview_id', 'na')}"

                    if st.session_state.get(_email_flag_key, False):
                        st.markdown(
                            f'<div style="background:#e6f4ea;border:1px solid #b7dfc0;'
                            f'border-radius:6px;padding:10px 14px;margin-top:8px;'
                            f'color:#1a7f37;">'
                            f'{_SVG_MAIL}Report already sent to your registered email.</div>',
                            unsafe_allow_html=True
                        )
                    else:
                        try:
                            from user_login import get_user_email_by_username, send_interview_report_email

                            _recipient_email = get_user_email_by_username(username)
                            if _recipient_email:
                                _email_sent = send_interview_report_email(
                                    to_email=_recipient_email,
                                    candidate_name=username,
                                    pdf_bytes=pdf_bytes,
                                    role=selected_role,
                                    domain=selected_domain,
                                    difficulty=st.session_state.interview_difficulty,
                                    overall_score=overall_avg,
                                )
                                if _email_sent:
                                    st.session_state[_email_flag_key] = True
                                    st.markdown(
                                        f'<div style="background:#e6f4ea;border:1px solid #b7dfc0;'
                                        f'border-radius:6px;padding:10px 14px;margin-top:8px;'
                                        f'color:#1a7f37;">'
                                        f'{_SVG_MAIL}Report also sent to your registered email '
                                        f'({_recipient_email}).</div>',
                                        unsafe_allow_html=True
                                    )
                                else:
                                    st.markdown(
                                        f'<div style="background:#fef3e2;border:1px solid #f2d29b;'
                                        f'border-radius:6px;padding:10px 14px;margin-top:8px;'
                                        f'color:#b45309;">'
                                        f'{_SVG_WARN}Could not send the report by email, but you can '
                                        f'still download it above.</div>',
                                        unsafe_allow_html=True
                                    )
                            else:
                                st.markdown(
                                    f'<div style="background:#e8f4fb;border:1px solid #b7d9ec;'
                                    f'border-radius:6px;padding:10px 14px;margin-top:8px;'
                                    f'color:#0369a1;">'
                                    f'{_SVG_INFO}No registered email found on your account — '
                                    f'download the report above instead.</div>',
                                    unsafe_allow_html=True
                                )
                        except Exception as _email_err:
                            st.markdown(
                                f'<div style="background:#fef3e2;border:1px solid #f2d29b;'
                                f'border-radius:6px;padding:10px 14px;margin-top:8px;'
                                f'color:#b45309;">'
                                f'{_SVG_WARN}Email delivery skipped due to an error: {_email_err}</div>',
                                unsafe_allow_html=True
                            )
                else:
                    st.warning("PDF generation failed. You can still review your results above.")

                # UNIFIED: Display recommended courses by difficulty
                st.markdown("---")
                st.subheader("📚 Recommended Courses for Your Career Growth")
                st.markdown(f"Based on your interview practice for **{selected_role}** in **{selected_domain}**, here are our course recommendations organized by difficulty level:")

                courses = get_courses_for_role(selected_domain, selected_role)
                if courses:
                    display_courses_by_difficulty(courses, selected_role)
                else:
                    st.info("No specific courses found for this role. Explore our course categories to find relevant learning resources!")

                # FIXED: Restart button - properly resets ALL interview state
                if st.button("🔄 Practice Again"):
                    # Reset all interview-related session state variables
                    st.session_state.dynamic_interview_started = False
                    st.session_state.dynamic_interview_completed = False
                    st.session_state.dynamic_interview_questions = []
                    st.session_state.current_dynamic_interview_question = 0
                    st.session_state.dynamic_interview_answers = []
                    st.session_state.dynamic_interview_scores = []
                    st.session_state.dynamic_interview_feedbacks = []
                    st.session_state.dynamic_answer_submitted = False
                    st.session_state.current_interview_question_text = ""
                    st.session_state.question_timer_start = None
                    st.session_state.timer_seconds = 120
                    st.session_state.interview_difficulty = "Medium"
                    st.session_state.original_num_questions = 6
                    st.session_state.resume_based_questions = []
                    st.session_state.generic_questions = []
                    st.session_state.interview_phase = "resume"
                    st.session_state.interview_result_saved = False
                    st.session_state.interview_final_duration_seconds = None
                    st.session_state.interview_actual_start_time = None
                    st.session_state.interview_mode = "mixed"
                    st.session_state.pending_followup_display = ""
                    st.session_state.pending_followup_strategy = ""
                    st.session_state.escalation_layer = 1
                    st.session_state.follow_up_count = 0
                    st.session_state.current_interview_id = None
                    st.session_state.question_db_ids = []
                    # ── Reset usage flag so next interview is properly gated ──
                    st.session_state._ac_usage_recorded_this_session = False
                    # ── Clear timer thread keys so auto-submit works on next interview ──
                    for _k in [k for k in st.session_state if k.startswith("_timer_thread_armed_")]:
                        st.session_state.pop(_k, None)
                    st.rerun()
        else:
            st.info("Please select both a career domain and target role to start the interview practice.")
    # Section 5: My Progress 📊
    elif page == "My Progress 📊":
        import pandas as pd
        import numpy as np
        import matplotlib.pyplot as plt
        import matplotlib
        matplotlib.use('Agg')
        import plotly.graph_objects as go
        import plotly.express as px
        from plotly.subplots import make_subplots

        # ── Dashboard CSS ──────────────────────────────────────────────────────
        st.markdown("""
        <style>
        /* Metric cards */
        .metric-card {
            background: rgba(255,255,255,0.04);
            backdrop-filter: blur(16px);
            -webkit-backdrop-filter: blur(16px);
            border: 1px solid rgba(255,255,255,0.07);
            border-radius: 14px;
            padding: 18px 20px;
            margin: 6px 0;
            transition: transform 0.18s cubic-bezier(0.4,0,0.2,1), box-shadow 0.18s cubic-bezier(0.4,0,0.2,1), border-color 0.18s cubic-bezier(0.4,0,0.2,1);
        }
        .metric-card:hover {
            transform: translateY(-3px);
            border-color: rgba(99,179,237,0.30);
            box-shadow: 0 8px 40px rgba(0,0,0,0.45), 0 0 30px rgba(79,163,227,0.15);
        }
        .metric-card .metric-label {
            color: #94a3b8;
            font-size: 0.72rem;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.06em;
            margin: 0 0 6px 0;
            font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
        }
        .metric-card .metric-value {
            color: #38bdf8;
            font-size: 1.75rem;
            font-weight: 700;
            margin: 0;
            line-height: 1.2;
            letter-spacing: -0.03em;
            font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
        }
        .metric-card .metric-sub {
            color: rgba(148,163,184,0.6);
            font-size: 0.72rem;
            margin: 4px 0 0 0;
            font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
        }
        /* Score badges — Apple SaaS style */
        .badge-excellent { background:rgba(52,211,153,0.15); color:#34d399; border:1px solid rgba(52,211,153,0.30); border-radius:99px; padding:2px 10px; font-weight:600; font-size:12px; }
        .badge-good      { background:rgba(56,189,248,0.12); color:#38bdf8; border:1px solid rgba(56,189,248,0.28); border-radius:99px; padding:2px 10px; font-weight:600; font-size:12px; }
        .badge-average   { background:rgba(251,191,36,0.12); color:#fbbf24; border:1px solid rgba(251,191,36,0.28); border-radius:99px; padding:2px 10px; font-weight:600; font-size:12px; }
        .badge-weak      { background:rgba(251,113,133,0.10); color:#fb7185; border:1px solid rgba(251,113,133,0.25); border-radius:99px; padding:2px 10px; font-weight:600; font-size:12px; }
        .badge-poor      { background:rgba(100,116,139,0.12); color:#64748b; border:1px solid rgba(100,116,139,0.25); border-radius:99px; padding:2px 10px; font-weight:600; font-size:12px; }
        /* Highlighted best row */
        .best-row { background: rgba(0,230,118,0.12) !important; }
        /* Section divider */
        .section-header {
            font-size: 1.1rem; font-weight: 700; color: #38bdf8;
            border-left: 4px solid #38bdf8; padding-left: 12px;
            margin: 24px 0 4px 0;
        }
        </style>
        """, unsafe_allow_html=True)

        st.subheader("📊 My Progress Dashboard")
        st.markdown("Track how you're improving over time, spot your strengths, and find exactly what to work on next.")

        username = st.session_state.get("username", "Guest")

        # Ensure DB and columns exist (runs once per session, not on every rerun)
        _ensure_db_initialized()

        # ── Load dashboard data with session_state caching ──────────────────────
        # Only re-query the DB when the user navigates to this page fresh, or when
        # a new interview has been saved (signalled by clearing _dashboard_cache_key).
        # This prevents a full DB round-trip (and visible flicker) on every widget
        # interaction that triggers a Streamlit rerun.
        _cache_key = f"_dashboard_df_{username}"
        _cache_dirty_key = f"_dashboard_dirty_{username}"

        if st.session_state.get(_cache_dirty_key, True) or _cache_key not in st.session_state:
            try:
                conn = _get_live_conn()
                df = pd.read_sql_query(
                    "SELECT * FROM interview_results WHERE username = %s ORDER BY id ASC",
                    conn, params=(username,)
                )
            except Exception as e:
                _err_str = str(e).lower()
                # Table doesn't exist yet (fresh deployment / first-time user) —
                # treat exactly the same as "no interviews yet"; no raw SQL shown.
                if "does not exist" in _err_str or "no such table" in _err_str or "undefined table" in _err_str:
                    df = pd.DataFrame()
                else:
                    # Genuine unexpected DB error — log a friendly message only
                    st.warning("⚠️ We couldn't load your dashboard right now. Please try refreshing in a moment.")
                    df = pd.DataFrame()
            st.session_state[_cache_key] = df
            st.session_state[_cache_dirty_key] = False
        else:
            df = st.session_state[_cache_key]

        # Refresh button — invalidates cache without a full page rerun
        if st.button("🔄 Refresh Dashboard", key="_dashboard_refresh_btn"):
            st.session_state[_cache_dirty_key] = True
            st.rerun()

        if df.empty:
            st.info("👋 You haven't completed any interviews yet. Head over to the **AI Interview Coach** tab, do your first practice session, and come back here to see your results!")
        else:
            # Ensure numeric types
            for col in ['avg_score', 'knowledge_avg', 'communication_avg', 'relevance_avg', 'duration_seconds', 'total_questions', 'weighted_score', 'raw_avg_score', 'depth_score', 'follow_up_count']:
                if col in df.columns:
                    df[col] = pd.to_numeric(df[col], errors='coerce')

            if 'difficulty' not in df.columns:
                df['difficulty'] = 'Unknown'
            df['difficulty'] = df['difficulty'].fillna('Unknown')

            # Backfill weighted_score if missing
            if 'weighted_score' not in df.columns or df['weighted_score'].isna().all():
                df['weighted_score'] = df['avg_score']
            else:
                df['weighted_score'] = df['weighted_score'].fillna(df['avg_score'])

            # =====================================================
            # SECTION A — EXECUTIVE SUMMARY METRICS
            # =====================================================
            st.markdown("---")
            st.markdown("### 🏆 Your Progress at a Glance")
            st.caption("Here's a quick overview of everything you've accomplished so far.")

            total_interviews = len(df)
            highest_score = df['avg_score'].max()
            lowest_score = df['avg_score'].min()
            overall_avg = df['avg_score'].mean()
            total_questions = int(df['total_questions'].fillna(0).sum()) if 'total_questions' in df.columns else 0

            # Improvement %
            if total_interviews >= 2:
                try:
                    first_score = float(df['avg_score'].dropna().iloc[0])
                    latest_score = float(df['avg_score'].dropna().iloc[-1])
                    improvement_pct = ((latest_score - first_score) / first_score) * 100 if first_score > 0 else 0.0
                except Exception:
                    improvement_pct = 0.0
            else:
                improvement_pct = 0.0

            # Consistency score based on std deviation
            score_std = df['avg_score'].std() if total_interviews > 1 else 0.0
            if score_std < 0.5:
                consistency_label = "🟢 Very Consistent"
            elif score_std < 1.5:
                consistency_label = "🟡 Fairly Consistent"
            else:
                consistency_label = "🔴 Varies a Lot"

            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.markdown(f"""<div class="metric-card">
                    <p class="metric-label">Interviews Completed</p>
                    <p class="metric-value">{total_interviews}</p>
                    <p class="metric-sub">Total sessions</p>
                </div>""", unsafe_allow_html=True)
            with col2:
                best_val = f"{format_score(highest_score)}/10" if not pd.isna(highest_score) else "N/A"
                st.markdown(f"""<div class="metric-card">
                    <p class="metric-label">Best Score Ever</p>
                    <p class="metric-value">{best_val}</p>
                    <p class="metric-sub">Personal best</p>
                </div>""", unsafe_allow_html=True)
            with col3:
                low_val = f"{format_score(lowest_score)}/10" if not pd.isna(lowest_score) else "N/A"
                st.markdown(f"""<div class="metric-card">
                    <p class="metric-label">Lowest Score</p>
                    <p class="metric-value" style="color:#ff9800;">{low_val}</p>
                    <p class="metric-sub">Room to grow</p>
                </div>""", unsafe_allow_html=True)
            with col4:
                avg_val = f"{format_score(overall_avg)}/10" if not pd.isna(overall_avg) else "N/A"
                st.markdown(f"""<div class="metric-card">
                    <p class="metric-label">Average Score</p>
                    <p class="metric-value">{avg_val}</p>
                    <p class="metric-sub">All-time average</p>
                </div>""", unsafe_allow_html=True)

            st.markdown("<div style='margin-top:10px'></div>", unsafe_allow_html=True)

            col5, col6, col7 = st.columns(3)
            with col5:
                st.markdown(f"""<div class="metric-card">
                    <p class="metric-label">Total Questions Answered</p>
                    <p class="metric-value">{total_questions}</p>
                    <p class="metric-sub">Real practice time</p>
                </div>""", unsafe_allow_html=True)
            with col6:
                sign = "+" if improvement_pct >= 0 else ""
                imp_color = "#00e676" if improvement_pct >= 0 else "#f44336"
                st.markdown(f"""<div class="metric-card">
                    <p class="metric-label">How Much You've Improved</p>
                    <p class="metric-value" style="color:{imp_color};">{sign}{improvement_pct:.1f}%</p>
                    <p class="metric-sub">vs. your first interview</p>
                </div>""", unsafe_allow_html=True)
            with col7:
                cons_color = "#00e676" if "Very" in consistency_label else ("#ffcc02" if "Fairly" in consistency_label else "#f44336")
                st.markdown(f"""<div class="metric-card">
                    <p class="metric-label">Score Consistency</p>
                    <p class="metric-value" style="color:{cons_color};font-size:18px;">{consistency_label}</p>
                    <p class="metric-sub">Std dev: {score_std:.2f}</p>
                </div>""", unsafe_allow_html=True)

            # =====================================================
            # SECTION B — SCORE TREND INTELLIGENCE
            # =====================================================
            st.markdown("---")
            st.markdown("### 📈 Are You Getting Better Over Time?")
            st.caption("This chart shows how your scores have changed across every interview you've done. The smoother line helps filter out one-off good or bad days.")

            trend_df = df[['avg_score', 'weighted_score']].copy().reset_index(drop=True)
            trend_df.index = trend_df.index + 1
            trend_df.index.name = "Interview #"

            # 3-point moving average
            trend_df['Smoothed Performance Trend'] = trend_df['avg_score'].rolling(window=3, min_periods=1).mean()
            trend_df = trend_df.rename(columns={
                'avg_score': 'Your Score',
                'weighted_score': 'Adjusted Score (Hard Interviews Count More)'
            })

            # ── Interactive Plotly trend chart ───────────────────────────────
            _x_vals = list(trend_df.index)
            _raw_scores = trend_df['Your Score'].tolist()
            _adj_scores = trend_df['Adjusted Score (Hard Interviews Count More)'].tolist()
            _smooth_scores = trend_df['Smoothed Performance Trend'].tolist()

            # Find best and worst interview indices
            _best_idx = int(np.argmax(_raw_scores))
            _worst_idx = int(np.argmin(_raw_scores))

            # Build difficulty labels for hover if available
            _diff_labels = df['difficulty'].tolist() if 'difficulty' in df.columns else [''] * len(_x_vals)
            _role_labels = df['role'].tolist() if 'role' in df.columns else [''] * len(_x_vals)
            _date_labels = df['completed_on'].tolist() if 'completed_on' in df.columns else [''] * len(_x_vals)

            _hover_text = [
                f"<b>Interview #{x}</b><br>Score: {float(s):.2f}/10<br>Role: {r}<br>Difficulty: {d}<br>Date: {dt}"
                for x, s, r, d, dt in zip(_x_vals, _raw_scores, _role_labels, _diff_labels, _date_labels)
            ]

            fig_trend = go.Figure()

            # Adjusted score area fill
            fig_trend.add_trace(go.Scatter(
                x=_x_vals, y=_adj_scores,
                name='Adjusted Score',
                mode='lines',
                line=dict(color='rgba(102,187,106,0.7)', width=1.5, dash='dot'),
                fill='tozeroy',
                fillcolor='rgba(102,187,106,0.05)',
                hovertemplate='Interview #%{x}<br>Adjusted: %{y:.2f}/10<extra></extra>'
            ))

            # Raw score line
            fig_trend.add_trace(go.Scatter(
                x=_x_vals, y=_raw_scores,
                name='Your Score',
                mode='lines+markers',
                line=dict(color='#00c3ff', width=2.5),
                marker=dict(size=7, color='#00c3ff', line=dict(width=1.5, color='white')),
                hovertext=_hover_text,
                hoverinfo='text',
            ))

            # Smoothed trend
            fig_trend.add_trace(go.Scatter(
                x=_x_vals, y=_smooth_scores,
                name='3-Interview Trend',
                mode='lines',
                line=dict(color='#ff9800', width=2, dash='dash'),
                hovertemplate='Interview #%{x}<br>Trend: %{y:.2f}/10<extra></extra>'
            ))

            # Best interview marker
            fig_trend.add_trace(go.Scatter(
                x=[_x_vals[_best_idx]], y=[_raw_scores[_best_idx]],
                name='🏆 Best',
                mode='markers+text',
                marker=dict(size=14, color='#00e676', symbol='star', line=dict(width=1.5, color='white')),
                text=[f" Best: {_raw_scores[_best_idx]:.2f}"],
                textposition='top right',
                textfont=dict(color='#00e676', size=11),
                hovertemplate=f'<b>🏆 Best Interview!</b><br>Score: {_raw_scores[_best_idx]:.2f}/10<extra></extra>'
            ))

            # Worst interview marker
            fig_trend.add_trace(go.Scatter(
                x=[_x_vals[_worst_idx]], y=[_raw_scores[_worst_idx]],
                name='⚠️ Lowest',
                mode='markers+text',
                marker=dict(size=14, color='#f44336', symbol='x', line=dict(width=2, color='white')),
                text=[f" Low: {_raw_scores[_worst_idx]:.2f}"],
                textposition='bottom right',
                textfont=dict(color='#f44336', size=11),
                hovertemplate=f'<b>⚠️ Lowest Interview</b><br>Score: {_raw_scores[_worst_idx]:.2f}/10<extra></extra>'
            ))

            # Average reference line
            fig_trend.add_hline(
                y=float(np.mean(_raw_scores)),
                line_dash='dot', line_color='rgba(255,255,255,0.25)',
                annotation_text=f'  Avg: {float(np.mean(_raw_scores)):.2f}',
                annotation_font_color='rgba(255,255,255,0.5)',
                annotation_position='right'
            )

            fig_trend.update_layout(
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(15,20,25,0.8)',
                font=dict(color='white', family='Inter, sans-serif'),
                legend=dict(
                    bgcolor='rgba(15,20,35,0.85)',
                    bordercolor='rgba(0,195,255,0.3)',
                    borderwidth=1,
                    orientation='h',
                    yanchor='bottom', y=1.02, xanchor='right', x=1
                ),
                xaxis=dict(
                    title='Interview #',
                    gridcolor='rgba(255,255,255,0.07)',
                    tickmode='linear',
                    dtick=max(1, len(_x_vals) // 20),   # max ~20 ticks visible at once
                    tickangle=-45 if len(_x_vals) > 20 else 0,
                    automargin=True,
                    showline=True, linecolor='rgba(0,195,255,0.3)'
                ),
                yaxis=dict(
                    title='Score (/10)',
                    range=[0, 10.5],
                    gridcolor='rgba(255,255,255,0.07)',
                    showline=True, linecolor='rgba(0,195,255,0.3)'
                ),
                hovermode='x unified',
                margin=dict(l=10, r=10, t=30, b=60),
                height=380
            )
            st.plotly_chart(fig_trend, use_container_width=True)
            st.caption("💡 **Adjusted Score** gives a little extra credit for completing harder interviews. **Smoothed Trend** is the average of your last 3 interviews — it shows your real direction without single-interview spikes.")

            # Detect trend direction using linear regression slope
            if total_interviews >= 3:
                _scores_list = df['avg_score'].dropna().tolist()
                _slope = compute_trend_slope(_scores_list)
                if _slope > 0.15:
                    trend_badge = "🟢 **You're Improving!** Your scores are going up across your recent interviews. Keep it up!"
                elif _slope < -0.15:
                    trend_badge = "🔴 **Scores Are Slipping.** Your recent interviews scored lower than earlier ones. Try reviewing feedback from your past sessions."
                else:
                    trend_badge = "🟡 **Holding Steady.** Your scores are staying about the same. Try harder difficulty levels to push your growth."
                # Stagnation detection
                if abs(_slope) < 0.05 and total_interviews >= 5:
                    trend_badge += " — ⚠️ **You may be in a plateau.** Switch to Hard mode or try a new topic to break through."
            else:
                _slope = 0.0
                trend_badge = "ℹ️ **Complete at least 3 interviews** to see your improvement trend here."
            st.markdown(trend_badge)

            # =====================================================
            # SECTION C — DOMAIN & ROLE ANALYTICS
            # =====================================================
            st.markdown("---")
            st.markdown("### 🌐 Where Are You Strongest?")
            st.caption("See which career areas and job roles you score highest in — and which ones need more practice.")

            if 'domain' in df.columns:
                col_l, col_r = st.columns(2)

                domain_counts = df.groupby('domain').size().rename('Interviews')
                domain_avg = df.groupby('domain')['avg_score'].mean().rename('Avg Score')

                # ── Dynamic identity color map ──────────────────────────────
                # sorted() ensures stable assignment — same area = same color
                # always, regardless of data order or how many areas exist.
                # % len(_CA_PALETTE) cycles gracefully for any number of areas.
                _CA_PALETTE = [
                    '#00c3ff', '#00e676', '#ff6b6b', '#ffd93d',
                    '#c77dff', '#ff9a3c', '#06d6a0', '#ff4d6d',
                    '#4cc9f0', '#f72585', '#3a86ff', '#a8dadc'
                ]
                _ca_list = sorted(df['domain'].dropna().unique().tolist())
                _ca_color_map = {
                    ca: _CA_PALETTE[i % len(_CA_PALETTE)]
                    for i, ca in enumerate(_ca_list)
                }

                with col_l:
                    st.markdown("**Interviews Done per Career Area**")
                    _dc_labels = domain_counts.index.tolist()
                    _dc_vals   = domain_counts.values.tolist()
                    _dc_colors = [_ca_color_map.get(ca, '#00c3ff') for ca in _dc_labels]
                    _fig_dc = go.Figure(go.Bar(
                        x=_dc_labels,
                        y=_dc_vals,
                        marker=dict(
                            color=_dc_colors,
                            line=dict(color='rgba(0,0,0,0.35)', width=1)
                        ),
                        text=_dc_vals,
                        textposition='outside',
                        textfont=dict(color='white', size=12),
                        hovertemplate='<b>%{x}</b><br>Interviews: %{y}<extra></extra>'
                    ))
                    _fig_dc.update_layout(
                        paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(15,20,25,0.8)',
                        font=dict(color='white'),
                        xaxis=dict(
                            title='Career Area', gridcolor='rgba(255,255,255,0.06)',
                            tickangle=-35, automargin=True
                        ),
                        yaxis=dict(
                            title='Interviews',
                            gridcolor='rgba(255,255,255,0.06)',
                            range=[0, max(_dc_vals) * 1.2]
                        ),
                        margin=dict(l=10, r=10, t=30, b=80), height=340
                    )
                    st.plotly_chart(_fig_dc, use_container_width=True)

                with col_r:
                    st.markdown("**Average Score per Career Area**")
                    # Same identity colors as left chart — same area = same color
                    # so both charts are instantly cross-referenceable visually
                    _da_labels = domain_avg.index.tolist()
                    _da_vals   = domain_avg.values.tolist()
                    _da_colors = [_ca_color_map.get(ca, '#00c3ff') for ca in _da_labels]
                    _fig_da = go.Figure(go.Bar(
                        x=_da_labels,
                        y=_da_vals,
                        marker=dict(
                            color=_da_colors,
                            line=dict(color='rgba(0,0,0,0.35)', width=1)
                        ),
                        text=[f"{v:.2f}" for v in _da_vals],
                        textposition='outside',
                        textfont=dict(color='white', size=12),
                        hovertemplate='<b>%{x}</b><br>Avg Score: %{y:.2f}/10<extra></extra>'
                    ))
                    _fig_da.update_layout(
                        paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(15,20,25,0.8)',
                        font=dict(color='white'),
                        xaxis=dict(
                            title='Career Area', gridcolor='rgba(255,255,255,0.06)',
                            tickangle=-35, automargin=True
                        ),
                        yaxis=dict(
                            title='Avg Score', range=[0, 10.5],
                            gridcolor='rgba(255,255,255,0.06)'
                        ),
                        margin=dict(l=10, r=10, t=30, b=80), height=340
                    )
                    st.plotly_chart(_fig_da, use_container_width=True)

                # Strongest / Weakest Domain
                if len(domain_avg) >= 1:
                    strongest_domain = domain_avg.idxmax()
                    weakest_domain = domain_avg.idxmin()
                    st.markdown(f"🏆 **You shine in:** {strongest_domain} — avg score {domain_avg[strongest_domain]:.2f}/10")
                    st.markdown(f"📌 **Room to grow in:** {weakest_domain} — avg score {domain_avg[weakest_domain]:.2f}/10. Spend more time practising here.")

            # Role breakdown — bar chart + pie chart + styled table
            if 'role' in df.columns:
                role_perf = df.groupby('role').agg(
                    Attempts=('avg_score', 'count'),
                    Avg_Score=('avg_score', 'mean'),
                    Best_Score=('avg_score', 'max'),
                    Latest_Score=('avg_score', 'last')
                ).reset_index()
                role_perf.columns = ['Role', 'Times Practised', 'Avg Score', 'Best Score', 'Last Score']
                role_perf = role_perf.round(2)

                st.markdown("**Role Performance Analytics**")
                col_rb1, col_rb2 = st.columns(2)

                # ── Dynamic role identity color map ─────────────────────────────────────
                # Each career area has its own color family (start RGB → end RGB).
                # Roles within a family are evenly interpolated across that gradient.
                # Works for ANY number of roles — no hardcoding, no cycling, no cutoff.
                # Add a role to DOMAIN_ROLES → it gets a shade automatically.
                # Fallback grey family handles any future domain not listed here.
                _ROLE_COLOR_FAMILIES = {
                    "Software Development and Engineering": ((30,  80,  220), (130, 180, 255)),
                    "Data Science and Analytics":           ((140, 40,  220), (210, 130, 255)),
                    "Cloud Computing and DevOps":           ((0,   160, 160), (100, 230, 210)),
                    "Cybersecurity":                        ((220, 40,  60),  (255, 130, 100)),
                    "UI/UX Design":                         ((220, 60,  160), (255, 160, 210)),
                    "Project Management":                   ((180, 180, 0),   (240, 230, 80)),
                }
                _role_color_map = {}
                for _dom, _roles in DOMAIN_ROLES.items():
                    _sorted_roles = sorted(_roles)
                    _n = len(_sorted_roles)
                    if _dom in _ROLE_COLOR_FAMILIES:
                        (_r0,_g0,_b0), (_r1,_g1,_b1) = _ROLE_COLOR_FAMILIES[_dom]
                    else:
                        # Fallback: grey gradient for any future unknown domain
                        (_r0,_g0,_b0), (_r1,_g1,_b1) = (100,100,100), (200,200,200)
                    for _i, _role in enumerate(_sorted_roles):
                        _t = _i / max(_n - 1, 1)
                        _role_color_map[_role] = (
                            f'rgb({int(_r0+_t*(_r1-_r0))},'
                            f'{int(_g0+_t*(_g1-_g0))},'
                            f'{int(_b0+_t*(_b1-_b0))})'
                        )

                with col_rb1:
                    # Avg Score by Role — identity color per role
                    _colors_bar = [
                        _role_color_map.get(r, 'rgb(120,120,120)')
                        for r in role_perf['Role']
                    ]
                    _fig_rb = go.Figure(go.Bar(
                        x=role_perf['Role'],
                        y=role_perf['Avg Score'],
                        marker=dict(
                            color=_colors_bar,
                            line=dict(color='rgba(0,0,0,0.35)', width=1)
                        ),
                        text=[f"{v:.2f}" for v in role_perf['Avg Score']],
                        textposition='outside',
                        textfont=dict(color='white', size=11),
                        hovertemplate='<b>%{x}</b><br>Avg Score: %{y:.2f}/10<extra></extra>'
                    ))
                    _fig_rb.update_layout(
                        title=dict(text='Avg Score by Role', font=dict(color='#00c3ff', size=14)),
                        paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(15,20,25,0.8)',
                        font=dict(color='white'),
                        xaxis=dict(
                            gridcolor='rgba(255,255,255,0.06)',
                            tickangle=-40,
                            automargin=True
                        ),
                        yaxis=dict(range=[0,10.5], gridcolor='rgba(255,255,255,0.06)'),
                        margin=dict(l=5, r=5, t=40, b=80),
                        height=max(380, len(role_perf) * 28 + 120)
                    )
                    st.plotly_chart(_fig_rb, use_container_width=True)

                with col_rb2:
                    # Interview Distribution by Role — same identity colors, sorted largest at top
                    _total_pie = role_perf['Times Practised'].sum()
                    _rd = role_perf.copy()
                    _rd['Pct'] = (_rd['Times Practised'] / _total_pie * 100).round(1)
                    _rd = _rd.sort_values('Times Practised', ascending=True)
                    _bar_colors = [
                        _role_color_map.get(r, 'rgb(120,120,120)')
                        for r in _rd['Role']
                    ]
                    _fig_rdist = go.Figure(go.Bar(
                        x=_rd['Times Practised'],
                        y=_rd['Role'],
                        orientation='h',
                        marker=dict(
                            color=_bar_colors,
                            line=dict(color='rgba(0,0,0,0.35)', width=1)
                        ),
                        text=[f"{int(v)}  ({p}%)" for v, p in zip(_rd['Times Practised'], _rd['Pct'])],
                        textposition='outside',
                        textfont=dict(color='white', size=11),
                        hovertemplate='<b>%{y}</b><br>Interviews: %{x}<extra></extra>',
                        cliponaxis=False
                    ))
                    # Dynamic height: grows with number of roles, never cuts off
                    _dyn_height = max(380, len(_rd) * 40 + 100)
                    _fig_rdist.update_layout(
                        title=dict(text='Interview Distribution by Role', font=dict(color='#00c3ff', size=14)),
                        paper_bgcolor='rgba(0,0,0,0)',
                        plot_bgcolor='rgba(15,20,25,0.8)',
                        font=dict(color='white'),
                        xaxis=dict(
                            title='Interviews',
                            gridcolor='rgba(255,255,255,0.07)',
                            showline=True, linecolor='rgba(0,195,255,0.3)',
                            range=[0, _rd['Times Practised'].max() * 1.35]
                        ),
                        yaxis=dict(gridcolor='rgba(255,255,255,0.04)', automargin=True),
                        margin=dict(l=10, r=20, t=45, b=40),
                        height=_dyn_height,
                        showlegend=False
                    )
                    st.plotly_chart(_fig_rdist, use_container_width=True)

                # Styled role table
                st.markdown("**Your Scores by Job Role**")
                _rp_styled = role_perf.copy()
                def _score_badge(v):
                    if v >= 8.5: return f'<span class="badge-excellent">{v:.2f}</span>'
                    elif v >= 7.0: return f'<span class="badge-good">{v:.2f}</span>'
                    elif v >= 5.5: return f'<span class="badge-average">{v:.2f}</span>'
                    elif v >= 4.0: return f'<span class="badge-weak">{v:.2f}</span>'
                    else: return f'<span class="badge-poor">{v:.2f}</span>'
                _best_role_idx = role_perf['Avg Score'].idxmax()
                _table_rows = ""
                for i, row in role_perf.iterrows():
                    _row_style = 'background:rgba(0,230,118,0.08);' if i == _best_role_idx else ''
                    _crown = ' 🏆' if i == _best_role_idx else ''
                    _table_rows += f"""<tr style="{_row_style}">
                        <td style="padding:8px 12px;color:#fff;">{row['Role']}{_crown}</td>
                        <td style="padding:8px 12px;color:#aaa;text-align:center;">{int(row['Times Practised'])}</td>
                        <td style="padding:8px 12px;text-align:center;">{_score_badge(row['Avg Score'])}</td>
                        <td style="padding:8px 12px;text-align:center;">{_score_badge(row['Best Score'])}</td>
                        <td style="padding:8px 12px;text-align:center;">{_score_badge(row['Last Score'])}</td>
                    </tr>"""
                st.markdown(f"""
                <div style="overflow-x:auto;border-radius:10px;border:1px solid rgba(0,195,255,0.2);">
                <table style="width:100%;border-collapse:collapse;background:rgba(15,20,25,0.8);">
                  <thead>
                    <tr style="border-bottom:1px solid rgba(0,195,255,0.3);">
                      <th style="padding:10px 12px;color:#38bdf8;text-align:left;font-size:12px;text-transform:uppercase;letter-spacing:0.07em;">Role</th>
                      <th style="padding:10px 12px;color:#38bdf8;text-align:center;font-size:12px;text-transform:uppercase;letter-spacing:0.07em;">Times</th>
                      <th style="padding:10px 12px;color:#38bdf8;text-align:center;font-size:12px;text-transform:uppercase;letter-spacing:0.07em;">Avg Score</th>
                      <th style="padding:10px 12px;color:#38bdf8;text-align:center;font-size:12px;text-transform:uppercase;letter-spacing:0.07em;">Best</th>
                      <th style="padding:10px 12px;color:#38bdf8;text-align:center;font-size:12px;text-transform:uppercase;letter-spacing:0.07em;">Last</th>
                    </tr>
                  </thead>
                  <tbody>{_table_rows}</tbody>
                </table></div>
                """, unsafe_allow_html=True)

            # =====================================================
            # SECTION D — DIFFICULTY PERFORMANCE
            # =====================================================
            st.markdown("---")
            st.markdown("### 🎯 How You Handle Different Difficulty Levels")
            st.caption("Easy interviews build confidence. Medium tests your thinking. Hard interviews push your limits — and show real growth.")

            if 'difficulty' in df.columns:
                # Only show rows where difficulty is known
                df_diff = df[df['difficulty'].notna() & (df['difficulty'] != 'Unknown') & (df['difficulty'] != '')]
                if df_diff.empty:
                    st.info("⚠️ No difficulty data yet. Complete a few more interviews and this section will fill up!")
                else:
                    diff_counts = df_diff.groupby('difficulty').size().rename('Attempts')
                    diff_avg = df_diff.groupby('difficulty')['avg_score'].mean().rename('Avg Score')

                    col_dl, col_dr = st.columns(2)
                    # Difficulty color map
                    _diff_colors = {'Easy': '#69f0ae', 'Medium': '#ffcc02', 'Hard': '#f44336'}
                    with col_dl:
                        st.markdown("**How Many Times You Tried Each Level**")
                        _fig_dfc = go.Figure(go.Bar(
                            x=diff_counts.index.tolist(), y=diff_counts.values.tolist(),
                            marker_color=[_diff_colors.get(d, '#00c3ff') for d in diff_counts.index],
                            text=diff_counts.values.tolist(), textposition='outside',
                            hovertemplate='<b>%{x}</b><br>Attempts: %{y}<extra></extra>'
                        ))
                        _fig_dfc.update_layout(
                            paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(15,20,25,0.8)',
                            font=dict(color='white'),
                            xaxis=dict(gridcolor='rgba(255,255,255,0.06)'),
                            yaxis=dict(gridcolor='rgba(255,255,255,0.06)'),
                            margin=dict(l=5,r=5,t=10,b=5), height=250
                        )
                        st.plotly_chart(_fig_dfc, use_container_width=True)
                    with col_dr:
                        st.markdown("**Your Average Score at Each Level**")
                        _fig_dfa = go.Figure(go.Bar(
                            x=diff_avg.index.tolist(), y=diff_avg.values.tolist(),
                            marker_color=[_diff_colors.get(d, '#00c3ff') for d in diff_avg.index],
                            text=[f"{v:.2f}" for v in diff_avg.values], textposition='outside',
                            hovertemplate='<b>%{x}</b><br>Avg Score: %{y:.2f}/10<extra></extra>'
                        ))
                        _fig_dfa.update_layout(
                            paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(15,20,25,0.8)',
                            font=dict(color='white'),
                            xaxis=dict(gridcolor='rgba(255,255,255,0.06)'),
                            yaxis=dict(range=[0,10.5], gridcolor='rgba(255,255,255,0.06)'),
                            margin=dict(l=5,r=5,t=10,b=5), height=250
                        )
                        st.plotly_chart(_fig_dfa, use_container_width=True)

                    # Analysis
                    hard_count = int(diff_counts.get('Hard', 0))
                    total_count = int(diff_counts.sum())
                    if total_count > 0 and hard_count / total_count < 0.2:
                        st.warning("⚠️ You haven't tried many Hard interviews yet. Pushing yourself to Hard level is one of the fastest ways to improve!")

                    hard_avg = float(diff_avg['Hard']) if 'Hard' in diff_avg.index else None
                    medium_avg = float(diff_avg['Medium']) if 'Medium' in diff_avg.index else None
                    if hard_avg is not None and medium_avg is not None:
                        if hard_avg >= medium_avg - 0.5:
                            st.success("✅ You're holding up well even in Hard interviews — that's a great sign of real progress!")
                        else:
                            st.info("💡 Your Hard interview scores are a bit lower than Medium, which is totally normal. Keep practising Hard mode to close the gap.")

            # =====================================================
            # SECTION E — SKILL INTELLIGENCE (RADAR CHART)
            # =====================================================
            st.markdown("---")
            st.markdown("### 🕸️ Your Skill Strengths")
            st.caption("This chart shows how you're performing across three key interview skills. The bigger the shape, the stronger you are overall.")

            skill_cols = ['knowledge_avg', 'communication_avg', 'relevance_avg']
            skill_labels = ['Knowledge', 'Communication', 'Relevance']

            # Use actual columns if available, else fallback to avg_score
            skill_avgs = []
            for col in skill_cols:
                if col in df.columns and df[col].notna().any():
                    skill_avgs.append(df[col].mean())
                else:
                    skill_avgs.append(df['avg_score'].mean())

            # Draw radar with matplotlib
            categories = skill_labels + [skill_labels[0]]
            values = skill_avgs + [skill_avgs[0]]
            angles = np.linspace(0, 2 * np.pi, len(skill_labels), endpoint=False).tolist()
            angles += angles[:1]

            fig_radar, ax_radar = plt.subplots(figsize=(5, 5), subplot_kw=dict(polar=True))
            fig_radar.patch.set_facecolor('#0f1419')
            ax_radar.set_facecolor('#1a2332')
            ax_radar.plot(angles, values, color='#00c3ff', linewidth=2)
            ax_radar.fill(angles, values, color='#00c3ff', alpha=0.25)
            ax_radar.set_xticks(angles[:-1])
            ax_radar.set_xticklabels(skill_labels, color='white', size=12)
            ax_radar.set_ylim(0, 10)
            ax_radar.set_yticks([2, 4, 6, 8, 10])
            ax_radar.set_yticklabels(['2', '4', '6', '8', '10'], color='gray', size=8)
            ax_radar.tick_params(colors='white')
            ax_radar.spines['polar'].set_color('#00c3ff')
            ax_radar.grid(color='gray', alpha=0.3)
            ax_radar.set_title("Skill Radar", color='#00c3ff', pad=20, size=14)

            col_radar, col_skill_info = st.columns([1, 1])
            with col_radar:
                st.pyplot(fig_radar)
            plt.close(fig_radar)

            with col_skill_info:
                weakest_skill_idx = skill_avgs.index(min(skill_avgs))
                weakest_skill = skill_labels[weakest_skill_idx]
                strongest_skill_idx = skill_avgs.index(max(skill_avgs))
                strongest_skill = skill_labels[strongest_skill_idx]

                st.markdown(f"🌟 **You're best at:** {strongest_skill} ({skill_avgs[strongest_skill_idx]:.2f}/10)")
                st.markdown(f"📌 **Focus area:** {weakest_skill} ({skill_avgs[weakest_skill_idx]:.2f}/10) — this is where more practice will help the most")
                st.markdown("")
                for lbl, val in zip(skill_labels, skill_avgs):
                    st.markdown(f"**{lbl}:** {val:.2f}/10")
                    st.progress(val / 10.0)

            # =====================================================
            # SECTION F — BEHAVIORAL ANALYTICS
            # =====================================================
            st.markdown("---")
            st.markdown("### 🧠 Your Interview Style")
            st.caption("This section looks at how you behave during interviews — how long you spend, how that affects your score, and what kind of interviewer you are.")

            col_b1, col_b2, col_b3 = st.columns(3)

            dur_available = 'duration_seconds' in df.columns and df['duration_seconds'].notna().any()
            _dur_series = df['duration_seconds'].dropna() if dur_available else None
            avg_duration_mins = (float(_dur_series.mean()) / 60.0) if (dur_available and len(_dur_series) > 0) else None
            avg_score_per_q = float((df['avg_score'] / df['total_questions'].replace(0, 1)).mean()) if ('total_questions' in df.columns and df['total_questions'].notna().any()) else None

            with col_b1:
                if avg_duration_mins is not None:
                    st.markdown(f"""<div class="metric-card">
                        <p class="metric-label">Average Time Per Interview</p>
                        <p class="metric-value">{avg_duration_mins:.1f}<span style="font-size:16px;color:#aaa"> min</span></p>
                        <p class="metric-sub">Typical session length</p>
                    </div>""", unsafe_allow_html=True)
                else:
                    st.markdown("""<div class="metric-card">
                        <p class="metric-label">Average Time Per Interview</p>
                        <p class="metric-value" style="font-size:18px;color:#666;">N/A</p>
                    </div>""", unsafe_allow_html=True)

            with col_b2:
                if avg_score_per_q is not None:
                    st.markdown(f"""<div class="metric-card">
                        <p class="metric-label">Score Per Question</p>
                        <p class="metric-value">{avg_score_per_q:.2f}</p>
                        <p class="metric-sub">Avg per individual question</p>
                    </div>""", unsafe_allow_html=True)
                else:
                    st.markdown("""<div class="metric-card">
                        <p class="metric-label">Score Per Question</p>
                        <p class="metric-value" style="font-size:18px;color:#666;">N/A</p>
                    </div>""", unsafe_allow_html=True)

            with col_b3:
                # Score vs duration correlation — convert to human badge
                if dur_available and len(df) >= 3:
                    corr = df[['avg_score', 'duration_seconds']].dropna().corr().iloc[0, 1]
                    if corr > 0.4:
                        corr_badge = "⚡ Yes — more time = better"
                    elif corr < -0.2:
                        corr_badge = "🤔 No — time isn't helping"
                    else:
                        corr_badge = "⚖️ Not much difference"
                    st.markdown(f"""<div class="metric-card">
                        <p class="metric-label">Does More Time Help?</p>
                        <p class="metric-value" style="font-size:16px;">{corr_badge}</p>
                        <p class="metric-sub">Based on all your interviews</p>
                    </div>""", unsafe_allow_html=True)
                else:
                    st.markdown("""<div class="metric-card">
                        <p class="metric-label">Does More Time Help?</p>
                        <p class="metric-value" style="font-size:16px;color:#666;">Need 3+ interviews</p>
                    </div>""", unsafe_allow_html=True)

            # Candidate type classification
            if dur_available and avg_duration_mins is not None:
                if avg_duration_mins < 10:
                    candidate_type = "⚡ **You tend to answer quickly.** That's great for pace, but try spending a bit more time structuring your answers — quality over speed!"
                elif avg_duration_mins > 35:
                    candidate_type = "🤔 **You take your time — sometimes too much.** Try to be more concise and direct. Interviewers appreciate clear, structured answers."
                else:
                    candidate_type = "⚖️ **Great balance!** You're pacing your interviews well — not too rushed, not too slow."
                st.info(candidate_type)

            # PART 6: Enhanced behavior classification using stored data
            if 'behavior_class' in df.columns and df['behavior_class'].notna().any():
                _bc_counts = df['behavior_class'].value_counts()
                _dominant_class = _bc_counts.index[0] if len(_bc_counts) > 0 else None
                if _dominant_class:
                    st.markdown(f"**🎭 Your Typical Interview Style:** {_dominant_class}")

            # Hard mode delta analysis
            if 'difficulty' in df.columns and 'Hard' in df['difficulty'].values and 'Medium' in df['difficulty'].values:
                _hard_avg_b = df[df['difficulty'] == 'Hard']['avg_score'].mean()
                _med_avg_b = df[df['difficulty'] == 'Medium']['avg_score'].mean()
                _hard_delta = _hard_avg_b - _med_avg_b
                st.markdown("#### 💪 How You Perform in Hard Interviews")
                st.caption("Hard interviews are more demanding — it's normal to score a little lower. Here's how you're doing.")
                col_hd1, col_hd2 = st.columns(2)
                with col_hd1:
                    st.markdown(f"""<div class="metric-card">
                        <p class="metric-label">Your Hard Interview Score</p>
                        <p class="metric-value">{_hard_avg_b:.2f}<span style="font-size:16px;color:#aaa">/10</span></p>
                        <p class="metric-sub">Average on Hard difficulty</p>
                    </div>""", unsafe_allow_html=True)
                with col_hd2:
                    if _hard_delta >= 0:
                        _delta_display = f"⬆️ {abs(_hard_delta):.1f} pts above Medium"
                        _dc = "#00e676"
                    elif _hard_delta >= -1.0:
                        _delta_display = f"Slightly below Medium (–{abs(_hard_delta):.1f} pts)"
                        _dc = "#ffcc02"
                    else:
                        _delta_display = f"Below Medium (–{abs(_hard_delta):.1f} pts)"
                        _dc = "#f44336"
                    st.markdown(f"""<div class="metric-card">
                        <p class="metric-label">Compared to Medium</p>
                        <p class="metric-value" style="color:{_dc};font-size:16px;">{_delta_display}</p>
                        <p class="metric-sub">Hard vs Medium gap</p>
                    </div>""", unsafe_allow_html=True)
                if _hard_delta < -1.5:
                    st.warning("⚠️ Hard interviews are noticeably tougher for you right now. That's okay — keep practising Hard mode and you'll build the muscle for it.")
                elif _hard_delta >= -0.5:
                    st.success("✅ You're doing great under pressure! Your Hard interview scores are close to your Medium ones — a real strength.")

            # =====================================================
            # SECTION G — CLASSIFICATION ENGINE
            # =====================================================
            st.markdown("---")
            st.markdown("### 🎖️ Where Do You Stand Right Now?")
            st.caption("Based on all your interviews, here's an honest picture of where you are today — and where you're headed.")

            if not pd.isna(overall_avg):
                if overall_avg < 5:
                    classification = "🔵 Just Getting Started"
                    cls_color = "#4fc3f7"
                    cls_desc = "Every expert was once a beginner. Focus on understanding the basics and practise regularly — you'll improve fast!"
                elif overall_avg < 6.5:
                    classification = "🟡 Building Momentum"
                    cls_color = "#ffcc02"
                    cls_desc = "You're making real progress! Work on giving more detailed answers and communicating your ideas more clearly."
                elif overall_avg < 7.5:
                    classification = "🟠 Looking Strong"
                    cls_color = "#ff9800"
                    cls_desc = "Solid work! You're getting there. Keep sharpening your answers and push yourself with harder interview levels."
                elif overall_avg < 8.5:
                    classification = "🟢 Almost There!"
                    cls_color = "#66bb6a"
                    cls_desc = "You're performing at a high level. A little more polish and you'll be fully interview-ready!"
                else:
                    classification = "🏆 Interview Ready!"
                    cls_color = "#00e676"
                    cls_desc = "Outstanding! You're ready to walk into real interviews with confidence. Go get that job!"

                st.markdown(f"""
                <div style="background: linear-gradient(135deg, rgba(0,195,255,0.1), rgba(0,195,255,0.05));
                            border: 2px solid {cls_color}; border-radius: 12px; padding: 20px; text-align: center; margin: 10px 0;">
                    <h2 style="color: {cls_color}; margin: 0;">{classification}</h2>
                    <p style="color: #ffffff; margin: 10px 0 0 0;">{cls_desc}</p>
                    <p style="color: #aaaaaa; margin: 5px 0 0 0;">Overall Average: {overall_avg:.2f}/10</p>
                </div>
                """, unsafe_allow_html=True)

            # =====================================================
            # SECTION H — AI GENERATED PERFORMANCE SUMMARY
            # =====================================================
            st.markdown("---")
            st.markdown("### 📝 Your Personal Progress Report")
            st.caption("Here's a plain-English summary of everything your data is telling us about your interview journey so far.")

            # Generate programmatic summary from real data
            summary_parts = []

            _domain_avg_safe = df.groupby('domain')['avg_score'].mean() if 'domain' in df.columns else None
            if _domain_avg_safe is not None and len(_domain_avg_safe) >= 1:
                _s_domain = _domain_avg_safe.idxmax()
                _w_domain = _domain_avg_safe.idxmin()
                summary_parts.append(f"You perform best in **{_s_domain}** — that's where your confidence and knowledge really shows, with an average score of {_domain_avg_safe[_s_domain]:.2f}/10.")
                if len(_domain_avg_safe) > 1:
                    summary_parts.append(f"**{_w_domain}** is the area that needs the most attention right now ({_domain_avg_safe[_w_domain]:.2f}/10). A little focused practice there will go a long way.")

            summary_parts.append(f"Across all your interviews, **{strongest_skill}** is your strongest skill ({skill_avgs[strongest_skill_idx]:.2f}/10). **{weakest_skill}** is the skill to focus on next ({skill_avgs[weakest_skill_idx]:.2f}/10) — even small improvements here will lift your overall scores.")

            # Trend direction — fully plain English, no slope values shown
            if total_interviews >= 3:
                _scores_for_summary = df['avg_score'].dropna().tolist()
                _slope_summary = compute_trend_slope(_scores_for_summary)
                if _slope_summary > 0.15:
                    summary_parts.append("The great news? **Your scores are going up** across your recent interviews. Whatever you're doing, keep doing it — it's working!")
                elif _slope_summary < -0.15:
                    summary_parts.append("Your recent scores have dipped a little compared to earlier interviews. Don't worry — this is normal. Try revisiting the feedback from your past sessions and focus on one skill at a time.")
                else:
                    summary_parts.append("Your scores have been fairly steady. That's a stable foundation to build on. To move to the next level, try bumping up to a harder difficulty or exploring a new topic area.")

            summary_parts.append(f"So far, you've completed **{total_interviews} interview{'s' if total_interviews != 1 else ''}** and answered **{total_questions} questions** in total — that's real practice time that adds up!")

            # Weighted score — explained simply
            _w_avg = df['weighted_score'].mean() if 'weighted_score' in df.columns else overall_avg
            summary_parts.append(f"Your adjusted score — which gives a little extra credit for harder interviews — is **{_w_avg:.2f}/10**. Hard interviews count more because they're more demanding.")

            if improvement_pct > 5:
                summary_parts.append(f"Since your very first interview, you've improved by **{improvement_pct:.1f}%**. That's a meaningful jump — you should feel great about that progress!")
            elif improvement_pct > 0:
                summary_parts.append(f"You're up **{improvement_pct:.1f}%** since your first interview. You're moving in the right direction — keep the momentum going.")
            elif improvement_pct < 0:
                summary_parts.append(f"Your score has dipped **{abs(improvement_pct):.1f}%** since your first interview. A small setback is part of learning. Try revisiting easier difficulty levels to rebuild your confidence, then push back up.")

            # Performance under pressure — plain English
            if 'difficulty' in df.columns and 'Hard' in df['difficulty'].values:
                _hard_avg_s = df[df['difficulty'] == 'Hard']['avg_score'].mean()
                if _hard_avg_s < overall_avg - 1.0:
                    summary_parts.append(f"Hard interviews are a challenge for you right now — you average {_hard_avg_s:.2f}/10 there, which is lower than your overall average. That's completely normal. The more you practise Hard mode, the more comfortable you'll get with tough questions.")
                else:
                    summary_parts.append(f"You're handling Hard interviews really well — averaging {_hard_avg_s:.2f}/10 even under pressure. That kind of resilience is exactly what real interviews reward.")

            # Behavior class — explained naturally
            if 'behavior_class' in df.columns and df['behavior_class'].notna().any():
                _bc = df['behavior_class'].mode().iloc[0] if not df['behavior_class'].dropna().empty else None
                _bc_descriptions = {
                    "⚡ Rushed": "You tend to answer quickly. Slowing down a little and structuring your thoughts before speaking can really lift your scores.",
                    "🤔 Overthinking": "You tend to take more time than needed. Practise giving focused, direct answers — interviewers love clarity.",
                    "⚖️ Balanced": "You have a great natural rhythm in interviews — not too fast, not too slow. That's a real skill.",
                    "🎯 Adaptive Learner": "You're adapting well as interviews get harder. That's a sign of someone who learns fast under pressure.",
                }
                if _bc:
                    _bc_desc = _bc_descriptions.get(_bc, f"Your typical style is: {_bc}.")
                    summary_parts.append(_bc_desc)

            full_summary = " ".join(summary_parts)
            st.markdown(f"""
            <div style="background: linear-gradient(135deg, rgba(0,195,255,0.08), rgba(0,195,255,0.03));
                        border: 1px solid rgba(0,195,255,0.3); border-radius: 12px; padding: 20px; margin: 10px 0;">
                <p style="color: #ffffff; font-size: 15px; line-height: 1.8; margin: 0;">{full_summary}</p>
            </div>
            """, unsafe_allow_html=True)

            # =====================================================
            # SECTION I — RECOMMENDATION ENGINE
            # =====================================================
            st.markdown("---")
            st.markdown("### 💡 What You Should Do Next")
            st.caption("These suggestions are personalised based on your actual interview history. Follow them and you'll see real improvement.")

            recommendations = []

            # Skill-based recommendations
            if weakest_skill == "Communication":
                recommendations.append("🗣️ **Work on explaining yourself more clearly.** Your communication scores are your lowest right now. Try practising with the STAR method: describe the Situation, your Task, the Action you took, and the Result. Even better — record yourself answering a question out loud and listen back.")
            elif weakest_skill == "Knowledge":
                recommendations.append("📚 **Deepen your technical knowledge.** Your knowledge scores suggest there are some topic gaps. Go back to basics in your target field, review common interview questions for your role, and spend time on real-world concepts like system design and best practices.")
            elif weakest_skill == "Relevance":
                recommendations.append("🎯 **Stay on-topic when you answer.** Your answers sometimes drift away from what was asked. Before you respond, mentally note the 2–3 key points that directly answer the question — then expand from there.")

            # Difficulty-based recommendations
            if 'difficulty' in df.columns:
                _diff_vals = df['difficulty'].dropna().values
                hard_avg_val = float(df[df['difficulty'] == 'Hard']['avg_score'].mean()) if 'Hard' in _diff_vals else None
                medium_avg_val = float(df[df['difficulty'] == 'Medium']['avg_score'].mean()) if 'Medium' in _diff_vals else None
                if hard_avg_val is not None and medium_avg_val is not None and hard_avg_val < medium_avg_val - 1.0:
                    recommendations.append("💪 **Practise more Hard interviews.** There's a noticeable gap between your Medium and Hard scores. The best way to close it is to get comfortable with the discomfort — book a few Hard mode sessions and treat each one as a learning experience, not a test.")
                hard_c = int((df['difficulty'] == 'Hard').sum())
                if total_interviews >= 3 and hard_c == 0:
                    recommendations.append("🔥 **Try your first Hard interview!** You haven't attempted Hard level yet. It's challenging, but one Hard interview teaches you more than three Easy ones. Give it a go — you're ready.")

            # Stagnation detection
            if total_interviews >= 5 and abs(improvement_pct) < 5:
                recommendations.append("📖 **Your scores have plateaued — it's time to shake things up.** Try a structured 2-week plan: spend week one revisiting technical concepts, and week two on behavioural questions. Finish each week with a full mock interview to test yourself.")

            # More interviews
            if total_interviews < 3:
                recommendations.append("📅 **Complete at least 5 interviews to unlock full insights.** Right now you don't have enough data for detailed trend analysis. The more you practise, the more personalised your recommendations become.")

            if recommendations:
                for rec in recommendations:
                    st.markdown(f"""
                    <div style="background: rgba(56,189,248,0.07); border-left: 4px solid #38bdf8;
                                padding: 12px 16px; margin: 8px 0; border-radius: 0 8px 8px 0;">
                        <p style="color: #ffffff; margin: 0;">{rec}</p>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.success("🎉 You're on track! Keep practising consistently and the results will keep coming.")

            # Raw data expander
            # Mode breakdown if available
            if 'interview_mode' in df.columns and df['interview_mode'].notna().any():
                st.markdown("---")
                st.markdown("### 🎮 Which Interview Type Do You Prefer?")
                st.caption("See how you perform across technical, behavioural, and mixed interview formats.")
                _mode_df = df[df['interview_mode'].notna() & (df['interview_mode'] != '')]
                if not _mode_df.empty:
                    col_m1, col_m2 = st.columns(2)
                    with col_m1:
                        st.markdown("**How Many Times You Tried Each Format**")
                        _mode_cnt = _mode_df.groupby('interview_mode').size().rename('Times Tried')
                        _fig_mc = go.Figure(go.Bar(
                            x=_mode_cnt.index.tolist(), y=_mode_cnt.values.tolist(),
                            marker_color='#00c3ff',
                            text=_mode_cnt.values.tolist(), textposition='outside',
                            hovertemplate='<b>%{x}</b><br>Times: %{y}<extra></extra>'
                        ))
                        _fig_mc.update_layout(
                            paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(15,20,25,0.8)',
                            font=dict(color='white'),
                            xaxis=dict(gridcolor='rgba(255,255,255,0.06)'),
                            yaxis=dict(gridcolor='rgba(255,255,255,0.06)'),
                            margin=dict(l=5,r=5,t=10,b=5), height=250
                        )
                        st.plotly_chart(_fig_mc, use_container_width=True)
                    with col_m2:
                        st.markdown("**Your Average Score by Format**")
                        _mode_avg = _mode_df.groupby('interview_mode')['avg_score'].mean().rename('Avg Score')
                        _fig_ma = go.Figure(go.Bar(
                            x=_mode_avg.index.tolist(), y=_mode_avg.values.tolist(),
                            marker_color=[f'rgba(0,195,255,{0.5 + 0.5*(v/10)})' for v in _mode_avg.values],
                            text=[f"{v:.2f}" for v in _mode_avg.values], textposition='outside',
                            hovertemplate='<b>%{x}</b><br>Avg Score: %{y:.2f}/10<extra></extra>'
                        ))
                        _fig_ma.update_layout(
                            paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(15,20,25,0.8)',
                            font=dict(color='white'),
                            xaxis=dict(gridcolor='rgba(255,255,255,0.06)'),
                            yaxis=dict(range=[0,10.5], gridcolor='rgba(255,255,255,0.06)'),
                            margin=dict(l=5,r=5,t=10,b=5), height=250
                        )
                        st.plotly_chart(_fig_ma, use_container_width=True)

            with st.expander("📋 See All Your Interview Records"):
                # Exclude raw DB 'id' — inject a clean per-user sequential # instead
                display_cols = [c for c in ['role', 'domain', 'avg_score', 'weighted_score', 'knowledge_avg', 'communication_avg',
                                             'relevance_avg', 'difficulty', 'interview_mode', 'total_questions', 'duration_seconds',
                                             'follow_up_count', 'depth_score', 'behavior_class', 'completed_on']
                                if c in df.columns]
                rename_map = {
                    'avg_score': 'Score', 'weighted_score': 'Adjusted Score', 'knowledge_avg': 'Knowledge',
                    'communication_avg': 'Communication', 'relevance_avg': 'Relevance',
                    'difficulty': 'Level', 'interview_mode': 'Format',
                    'total_questions': 'Questions', 'duration_seconds': 'Duration (s)',
                    'completed_on': 'Date', 'role': 'Role', 'domain': 'Career Area',
                    'follow_up_count': 'Follow-ups', 'depth_score': 'Depth', 'behavior_class': 'Style'
                }
                display_df = df[display_cols].rename(columns=rename_map)
                # Per-user sequential numbering: always starts at 1 regardless of DB id
                display_df.insert(0, '#', range(1, len(display_df) + 1))

                # Build enhanced HTML table with score badges, trend arrows, best-row highlight
                _score_col = 'Score'
                _scores_list_disp = display_df[_score_col].tolist() if _score_col in display_df.columns else []
                _best_score_val = max(_scores_list_disp) if _scores_list_disp else None

                def _badge(v):
                    if pd.isna(v): return '<span style="color:#666">N/A</span>'
                    v = float(v)
                    if v >= 8.5: return f'<span class="badge-excellent">{v:.2f}</span>'
                    elif v >= 7.0: return f'<span class="badge-good">{v:.2f}</span>'
                    elif v >= 5.5: return f'<span class="badge-average">{v:.2f}</span>'
                    elif v >= 4.0: return f'<span class="badge-weak">{v:.2f}</span>'
                    else: return f'<span class="badge-poor">{v:.2f}</span>'

                def _trend_arrow(current, prev):
                    if prev is None or pd.isna(prev): return ''
                    delta = float(current) - float(prev)
                    if delta > 0.3: return f'<span style="color:#00e676;font-size:14px;" title="+{delta:.2f}">▲</span>'
                    elif delta < -0.3: return f'<span style="color:#f44336;font-size:14px;" title="{delta:.2f}">▼</span>'
                    else: return f'<span style="color:#ffcc02;font-size:14px;" title="~{delta:.2f}">●</span>'

                _th_style = "padding:9px 12px;color:#38bdf8;text-align:left;font-size:11px;text-transform:uppercase;letter-spacing:0.07em;border-bottom:1px solid rgba(0,195,255,0.3);white-space:nowrap;"
                _td_style = "padding:8px 12px;color:#e0e0e0;font-size:13px;white-space:nowrap;"

                _headers = list(display_df.columns)
                _header_row = "".join([f'<th style="{_th_style}">{h}</th>' for h in _headers]) + f'<th style="{_th_style}">Trend</th>'

                _body_rows = ""
                _prev_score = None
                for i, row in display_df.iterrows():
                    _cur_score = row.get('Score', None)
                    _is_best = (not pd.isna(_cur_score) and not pd.isna(_best_score_val) and float(_cur_score) == float(_best_score_val))
                    _row_bg = 'background:rgba(0,230,118,0.10);' if _is_best else ('background:rgba(255,255,255,0.02);' if i % 2 == 0 else '')
                    _cells = ""
                    for col_name in _headers:
                        val = row[col_name]
                        if col_name in ('Score', 'Adjusted Score', 'Knowledge', 'Communication', 'Relevance'):
                            _cells += f'<td style="{_td_style}text-align:center;">{_badge(val)}</td>'
                        elif col_name == 'Level':
                            _lc = {'Easy':'#69f0ae','Medium':'#ffcc02','Hard':'#f44336'}.get(str(val), '#aaa')
                            _cells += f'<td style="{_td_style}"><span style="color:{_lc};font-weight:600;">{val}</span></td>'
                        elif col_name == '#':
                            _crown = ' 🏆' if _is_best else ''
                            _cells += f'<td style="{_td_style}font-weight:600;">{val}{_crown}</td>'
                        else:
                            _disp_val = str(val) if not pd.isna(val) else '—'
                            _cells += f'<td style="{_td_style}">{_disp_val}</td>'
                    _arrow = _trend_arrow(_cur_score, _prev_score) if not pd.isna(_cur_score) else ''
                    _cells += f'<td style="{_td_style}text-align:center;">{_arrow}</td>'
                    _body_rows += f'<tr style="{_row_bg}">{_cells}</tr>'
                    if not pd.isna(_cur_score):
                        _prev_score = _cur_score

                _total_records = len(display_df)
                _html_table = (
                    """<!DOCTYPE html><html><head><meta charset="utf-8"><style>"""
                    """body{margin:0;background:transparent;font-family:Inter,sans-serif;}"""
                    """.sw{display:flex;align-items:center;gap:10px;padding:10px 12px 8px;"""
                    """background:rgba(15,20,25,0.95);border:1px solid rgba(0,195,255,0.2);"""
                    """border-bottom:none;border-radius:10px 10px 0 0;position:sticky;top:0;z-index:10;}"""
                    """#si{flex:1;padding:7px 12px;border-radius:7px;border:1px solid rgba(0,195,255,0.3);"""
                    """background:rgba(255,255,255,0.06);color:#e0e0e0;font-size:13px;outline:none;}"""
                    """#si::placeholder{color:rgba(255,255,255,0.3);}"""
                    """#si:focus{border-color:rgba(0,195,255,0.7);}"""
                    """#cl{font-size:12px;color:rgba(255,255,255,0.4);white-space:nowrap;min-width:100px;text-align:right;}"""
                    """.sc{overflow-y:auto;overflow-x:auto;max-height:480px;border:1px solid rgba(0,195,255,0.2);border-radius:0 0 10px 10px;}"""
                    """table{width:100%;border-collapse:collapse;background:rgba(15,20,25,0.85);}"""
                    """thead tr{position:sticky;top:0;z-index:5;background:rgba(10,15,22,0.98);}"""
                    """tr.hidden{display:none;}"""
                    """.lg{color:rgba(255,255,255,0.35);font-size:11px;padding:6px 2px 0;}"""
                    """.badge-excellent{display:inline-block;padding:2px 8px;border-radius:12px;background:rgba(0,230,118,0.15);color:#00e676;font-weight:600;font-size:12px;}"""
                    """.badge-good{display:inline-block;padding:2px 8px;border-radius:12px;background:rgba(0,195,255,0.12);color:#00c3ff;font-weight:600;font-size:12px;}"""
                    """.badge-average{display:inline-block;padding:2px 8px;border-radius:12px;background:rgba(255,204,2,0.12);color:#ffcc02;font-weight:600;font-size:12px;}"""
                    """.badge-weak{display:inline-block;padding:2px 8px;border-radius:12px;background:rgba(255,152,0,0.12);color:#ff9800;font-weight:600;font-size:12px;}"""
                    """.badge-poor{display:inline-block;padding:2px 8px;border-radius:12px;background:rgba(244,67,54,0.12);color:#f44336;font-weight:600;font-size:12px;}"""
                    """</style></head><body>"""
                    f"""<div class="sw">"""
                    """<input id="si" type="text" placeholder="&#128269; Filter by role, career area, level, format..." />"""
                    f"""<span id="cl">All {_total_records} records</span>"""
                    """</div>"""
                    f"""<div class="sc"><table><thead><tr>{_header_row}</tr></thead>"""
                    f"""<tbody id="tb">{_body_rows}</tbody></table></div>"""
                    """<div class="lg">&#127942; Gold = personal best &nbsp;|&nbsp; &#9650; improved &nbsp;&#9660; dipped &nbsp;&#9679; steady vs previous</div>"""
                    f"""<script>"""
                    """(function(){{"""
                    """var inp=document.getElementById('si');"""
                    """var lbl=document.getElementById('cl');"""
                    """var rows=document.querySelectorAll('#tb tr');"""
                    f"""var total={_total_records};"""
                    """inp.addEventListener('input',function(){{"""
                    """var q=this.value.toLowerCase().trim();"""
                    """var vis=0;"""
                    """rows.forEach(function(r){{"""
                    """if(!q||r.textContent.toLowerCase().includes(q)){{r.classList.remove('hidden');vis++;}}"""
                    """else{{r.classList.add('hidden');}}"""
                    """}});"""
                    """lbl.textContent=q?(vis+' of '+total+' records'):('All '+total+' records');"""
                    """}});"""
                    """}})();"""
                    """</script></body></html>"""
                )
                st.components.v1.html(_html_table, height=600, scrolling=False)
with tab_scam:
    render_job_scam_detector_tab(call_llm)
if tab5:
	with tab5:
		# sqlite3 removed — using Supabase PostgreSQL via db_manager
		import pandas as pd
		import matplotlib.pyplot as plt
		import numpy as np
		import streamlit as st
		from datetime import datetime, timedelta
		import plotly.express as px
		import plotly.graph_objects as go
		from plotly.subplots import make_subplots
		import time
		import glob, os

		# Import enhanced database manager functions
		from db_manager import (
			get_top_domains_by_score,
			get_resume_count_by_day,
			get_average_ats_by_domain,
			get_domain_distribution,
			get_bias_distribution,
			filter_candidates_by_date,
			delete_candidate_by_id,
			get_all_candidates,
			get_candidate_by_id,
			get_domain_performance_stats,
			get_daily_ats_stats,
			get_flagged_candidates,
			get_database_stats,
			analyze_domain_transitions,
			export_to_csv,
			cleanup_old_records,
			DatabaseManager
		)

		# Initialize enhanced database manager
		@st.cache_resource
		def get_db_manager():
			return DatabaseManager()

		db_manager = get_db_manager()

		def create_enhanced_pie_chart(df, values_col, labels_col, title):
			"""Create an enhanced pie chart with better styling"""
			fig = px.pie(
				df, 
				values=values_col, 
				names=labels_col,
				title=title,
				color_discrete_sequence=px.colors.qualitative.Set3
			)
			fig.update_traces(
				textposition='inside', 
				textinfo='percent+label',
				hovertemplate='<b>%{label}</b><br>Count: %{value}<br>Percentage: %{percent}<extra></extra>'
			)
			fig.update_layout(
				showlegend=True,
				legend=dict(orientation="v", yanchor="middle", y=0.5, xanchor="left", x=1.01),
				margin=dict(t=50, b=50, l=50, r=150)
			)
			return fig

		def create_enhanced_bar_chart(df, x_col, y_col, title, orientation='v'):
			"""Create enhanced bar chart with better interactivity"""
			if orientation == 'v':
				fig = px.bar(df, x=x_col, y=y_col, title=title, 
							color=y_col, color_continuous_scale='viridis')
				fig.update_xaxes(tickangle=45)
			else:
				fig = px.bar(df, x=y_col, y=x_col, title=title, orientation='h',
							color=y_col, color_continuous_scale='viridis')
			
			fig.update_traces(
				hovertemplate='<b>%{y if orientation == "v" else x}</b><br>Value: %{x if orientation == "v" else y}<extra></extra>'
			)
			fig.update_layout(showlegend=False, margin=dict(t=50, b=50, l=50, r=50))
			return fig

		def load_domain_distribution():
			"""Enhanced domain distribution loading with error handling"""
			try:
				df = get_domain_distribution()
				if not df.empty:
					df = df.sort_values(by="count", ascending=False).reset_index(drop=True)
					return df
			except Exception as e:
				st.error(f"Error loading domain distribution: {e}")
			return pd.DataFrame()

		# Enhanced Data Loading with Caching
		@st.cache_data(ttl=300)  # Cache for 5 minutes
		def load_all_candidates():
			try:
				return get_all_candidates()
			except Exception as e:
				st.error(f"Error loading candidates: {e}")
				return pd.DataFrame()

		# -------- Glassmorphism Styles with Shimmer --------
		st.markdown("""
		<style>
		.glass-box {
			background: rgba(10, 20, 40, 0.55);
			border-radius: 18px;
			padding: 2rem;
			backdrop-filter: blur(14px);
			border: 1px solid rgba(0, 200, 255, 0.35);
			box-shadow: 0 8px 32px rgba(0, 200, 255, 0.25);
			position: relative;
			overflow: hidden;
			text-align: center;
			margin-bottom: 2rem;
		}
		.glass-box::before {
			content: "";
			position: absolute;
			top: -50%;
			left: -50%;
			width: 200%;
			height: 200%;
			background: linear-gradient(
				120deg,
				rgba(255,255,255,0.15) 0%,
				rgba(255,255,255,0.05) 40%,
				transparent 60%
			);
			transform: rotate(25deg);
			animation: shimmer 6s infinite;
		}
		@keyframes shimmer {
			0% { top: -50%; left: -50%; }
			50% { top: 100%; left: 100%; }
			100% { top: -50%; left: -50%; }
		}
		.glass-box h1, .glass-box h2 {
			color: #4da6ff;
			text-shadow: 0 0 12px rgba(0,200,255,0.7);
			margin: 0 0 0.5rem 0;
			font-weight: 600;
		}
		.glass-box p {
			color: #cce6ff;
			margin: 0;
			font-size: 0.95rem;
		}

		/* Glassy input fields */
		.stTextInput > div > div > input {
			background: rgba(255, 255, 255, 0.08) !important;
			border: 1px solid rgba(0, 200, 255, 0.3) !important;
			border-radius: 12px !important;
			padding: 10px !important;
			color: #e6f7ff !important;
			font-weight: 500 !important;
			backdrop-filter: blur(10px) !important;
		}
		.stTextInput > div > div > input:focus {
			border: 1px solid rgba(0, 200, 255, 0.8) !important;
			box-shadow: 0 0 12px rgba(0, 200, 255, 0.6) !important;
			outline: none !important;
		}

		/* Glassy button */
		.stButton > button {
			background: rgba(0, 200, 255, 0.15);
			border: 1px solid rgba(0, 200, 255, 0.4);
			border-radius: 12px;
			color: #e6f7ff;
			padding: 0.6rem 1.2rem;
			font-weight: bold;
			backdrop-filter: blur(8px);
			transition: all 0.3s ease;
		}
		.stButton > button:hover {
			background: rgba(0, 200, 255, 0.3);
			box-shadow: 0 0 16px rgba(0, 200, 255, 0.7);
			transform: translateY(-2px);
		}
		</style>
		""", unsafe_allow_html=True)

		# ---------------- Enhanced Authentication System ----------------
		if "admin_logged_in" not in st.session_state:
			st.session_state.admin_logged_in = False

		if not st.session_state.admin_logged_in:
			st.markdown("""
			<div class="glass-box">
				<h2>🔐 Admin Authentication Required</h2>
				<p>Please enter your email and password to access the admin dashboard</p>
			</div>
			""", unsafe_allow_html=True)
			
			col1, col2, col3 = st.columns([1, 2, 1])
			with col2:
				email = st.text_input("📧 Enter Admin Email", placeholder="Enter email...")
				password = st.text_input("🔑 Enter Admin Password", type="password", placeholder="Enter password...")
				login_clicked = st.button("🚀 Login", use_container_width=True)

				if login_clicked:
					valid_email = "admin@example.com"
					valid_password = "Swagato@2002"

					if email == valid_email and password == valid_password:
						st.session_state.admin_logged_in = True
						st.success("✅ Authentication successful! Redirecting to dashboard...")
						st.rerun()
					else:
						msg_placeholder = st.empty()
						msg_placeholder.markdown("""
							<div style='
								background-color: #ff4d4d;
								color: white;
								padding: 10px 15px;
								border-radius: 10px;
								text-align: center;
								animation: slideDown 0.5s ease-in-out;
							'>❌ Invalid credentials. Please try again.</div>
							<style>
							@keyframes slideDown {
								0% {transform: translateY(-50px); opacity: 0;}
								100% {transform: translateY(0); opacity: 1;}
							}
							</style>
						""", unsafe_allow_html=True)
						time.sleep(3)
						msg_placeholder.empty()

			st.stop()

		# ---------------- Enhanced Header with Database Stats ----------------
		st.markdown("""
		<div class="glass-box">
			<h1>🛡️ Enhanced Admin Database Panel</h1>
			<p>Advanced Resume Analysis System Dashboard</p>
		</div>
		""", unsafe_allow_html=True)

		# Enhanced Control Panel
		col1, col2, col3, col4 = st.columns(4)
		with col1:
			if st.button("🔄 Refresh All Data", use_container_width=True):
				st.cache_data.clear()
				st.rerun()
		with col2:
			if st.button("📊 Database Stats", use_container_width=True):
				st.session_state.show_db_stats = True
		with col3:
			if st.button("🧹 Cleanup Old Records", use_container_width=True):
				st.session_state.show_cleanup = True
		with col4:
			if st.button("🚪 Secure Logout", use_container_width=True):
				st.session_state.admin_logged_in = False
				st.success("👋 Logged out successfully.")
				st.rerun()

		# Database Statistics Panel
		if st.session_state.get('show_db_stats', False):
			with st.expander("📈 Database Statistics", expanded=True):
				try:
					stats = get_database_stats()
					if stats:
						col1, col2, col3, col4 = st.columns(4)
						with col1:
							st.metric("Total Candidates", stats.get('total_candidates', 0))
						with col2:
							st.metric("Average ATS Score", f"{stats.get('avg_ats_score', 0):.2f}")
						with col3:
							st.metric("Unique Domains", stats.get('unique_domains', 0))
						with col4:
							st.metric("Database Size", f"{stats.get('database_size_mb', 0):.2f} MB")
						
						col5, col6 = st.columns(2)
						with col5:
							st.metric("Earliest Record", stats.get('earliest_date', 'N/A'))
						with col6:
							st.metric("Latest Record", stats.get('latest_date', 'N/A'))
				except Exception as e:
					st.error(f"Error loading database statistics: {e}")

		# Cleanup Panel
		if st.session_state.get('show_cleanup', False):
			with st.expander("🧹 Database Cleanup", expanded=True):
				days_to_keep = st.slider("Days to Keep", 30, 730, 365)
				if st.button("⚠️ Cleanup Old Records"):
					try:
						deleted_count = cleanup_old_records(days_to_keep)
						if deleted_count > 0:
							st.success(f"✅ Cleaned up {deleted_count} old records")
						else:
							st.info("ℹ️ No old records found to cleanup")
					except Exception as e:
						st.error(f"Error during cleanup: {e}")

		st.markdown("<hr style='border-top: 2px solid #bbb; margin: 2rem 0;'>", unsafe_allow_html=True)

		df = load_all_candidates()

		# Enhanced Search and Filter Section
		st.markdown("### 🔍 Advanced Search & Filters")
		
		col1, col2 = st.columns(2)
		with col1:
			search = st.text_input("🔍 Search by Candidate Name", placeholder="Enter candidate name...")
			if search:
				df = df[df["candidate_name"].str.contains(search, case=False, na=False)]
		
		with col2:
			domain_filter = st.selectbox("🏢 Filter by Domain", 
									options=["All Domains"] + list(df["domain"].unique()) if not df.empty else ["All Domains"])
			if domain_filter != "All Domains":
				df = df[df["domain"] == domain_filter]

		# Enhanced Date Filter
		st.markdown("#### 📅 Date Range Filter")
		col1, col2, col3 = st.columns(3)
		with col1:
			start_date = st.date_input("📅 Start Date", value=datetime.now() - timedelta(days=30))
		with col2:
			end_date = st.date_input("📅 End Date", value=datetime.now())
		with col3:
			if st.button("🎯 Apply Filters", use_container_width=True):
				try:
					df = filter_candidates_by_date(str(start_date), str(end_date))
					if domain_filter != "All Domains":
						df = df[df["domain"] == domain_filter]
					if search:
						df = df[df["candidate_name"].str.contains(search, case=False, na=False)]
					st.success(f"✅ Filters applied. Found {len(df)} candidates.")
				except Exception as e:
					st.error(f"Error applying filters: {e}")

		# Enhanced Candidates Display
		if df.empty:
			st.info("ℹ️ No candidate data available with current filters.")
		else:
			st.markdown(f"### 📋 Candidates Overview ({len(df)} records)")
			
			# Enhanced metrics
			col1, col2, col3, col4 = st.columns(4)
			with col1:
				st.metric("Total Candidates", len(df))
			with col2:
				st.metric("Avg ATS Score", f"{df['ats_score'].mean():.2f}")
			with col3:
				st.metric("Avg Bias Score", f"{df['bias_score'].mean():.3f}")
			with col4:
				st.metric("Unique Domains", df['domain'].nunique())

			# Enhanced data display with sorting
			sort_column = st.selectbox("📊 Sort by", 
								options=['timestamp', 'ats_score', 'bias_score', 'candidate_name', 'domain'])
			sort_order = st.radio("Sort Order", ["Descending", "Ascending"], horizontal=True)
			
			df_sorted = df.sort_values(by=sort_column, ascending=(sort_order == "Ascending"))
			
			# Display with enhanced formatting
			st.dataframe(
				df_sorted.style.format({
					'ats_score': '{:.0f}',
					'edu_score': '{:.0f}',
					'exp_score': '{:.0f}',
					'skills_score': '{:.0f}',
					'lang_score': '{:.0f}',
					'keyword_score': '{:.0f}',
					'format_score': '{:.0f}',
					'bias_score': '{:.3f}'
				}),
				use_container_width=True,
				height=400
			)

			# Enhanced Export Options
			col1, col2 = st.columns(2)
			with col1:
				csv_data = df_sorted.to_csv(index=False)
				st.download_button(
					label="📥 Download Filtered Data (CSV)",
					data=csv_data,
					file_name=f"candidates_filtered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
					mime="text/csv",
					use_container_width=True
				)
			with col2:
				if st.button("📤 Export All Data", use_container_width=True):
					try:
						filename = f"full_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
						if export_to_csv(filename):
							st.success(f"✅ Data exported to {filename}")
						else:
							st.error("❌ Export failed")
					except Exception as e:
						st.error(f"Export error: {e}")

			st.markdown("### 📂 Export Archive")
			export_files = sorted(glob.glob("full_export_*.csv"), reverse=True)

			if export_files:
				for file in export_files:
					with open(file, "rb") as f:
						st.download_button(
							label=f"⬇️ Download {os.path.basename(file)}",
							data=f,
							file_name=os.path.basename(file),
							mime="text/csv",
							use_container_width=True
						)
			else:
				st.info("📭 No export files found yet.")

			# Enhanced Delete Functionality
			with st.expander("🗑️ Delete Candidate", expanded=False):
				st.warning("⚠️ This action cannot be undone!")
				delete_id = st.number_input("Enter Candidate ID", min_value=1, step=1, key="delete_id")
				
				if delete_id in df["id"].values:
					candidate_info = get_candidate_by_id(delete_id)
					if not candidate_info.empty:
						st.info("📄 Candidate to be deleted:")
						st.dataframe(candidate_info, use_container_width=True)
						
						if st.button("❌ Confirm Delete", type="primary"):
							try:
								if delete_candidate_by_id(delete_id):
									st.success(f"✅ Candidate with ID {delete_id} deleted successfully.")
									st.cache_data.clear()
									st.rerun()
								else:
									st.error("❌ Failed to delete candidate.")
							except Exception as e:
								st.error(f"Delete error: {e}")
				elif delete_id > 0:
					st.error("❌ Candidate ID not found.")

		# Enhanced Analytics Section
		st.markdown("<hr style='border-top: 2px solid #bbb; margin: 2rem 0;'>", unsafe_allow_html=True)
		st.markdown("## 📊 Advanced Analytics Dashboard")

		# Enhanced Top Domains Analysis
		st.markdown("### 🏆 Top Performing Domains")
		
		try:
			top_domains = get_top_domains_by_score(limit=10)
			if top_domains:
				df_top = pd.DataFrame(top_domains, columns=["domain", "avg_ats", "count"])
				
				col1, col2 = st.columns([1, 2])
				with col1:
					sort_order = st.radio("📊 Sort by ATS", ["⬆️ High to Low", "⬇️ Low to High"], horizontal=True)
					limit = st.slider("Show Top N Domains", 1, len(df_top), value=min(8, len(df_top)))
				
				ascending = sort_order == "⬇️ Low to High"
				df_sorted = df_top.sort_values(by="avg_ats", ascending=ascending).head(limit)
				
				# Interactive chart
				fig = create_enhanced_bar_chart(df_sorted, "domain", "avg_ats", 
										"Average ATS Score by Domain", orientation='h')
				st.plotly_chart(fig, use_container_width=True)
				
				# Enhanced domain cards with glassmorphism
				st.markdown("""
				<style>
				@keyframes tab5-shimmer {
					0% { background-position: -200% 0; }
					100% { background-position: 200% 0; }
				}
				.tab5-domain-card {
					background: rgba(10, 20, 40, 0.3);
					backdrop-filter: blur(10px);
					border: 1px solid rgba(0, 200, 255, 0.2);
					box-shadow: 0 4px 12px rgba(0, 0, 0, 0.3);
					border-radius: 15px;
					padding: 15px;
					margin-bottom: 15px;
					transition: all 0.3s ease;
					cursor: pointer;
					position: relative;
					overflow: hidden;
				}
				.tab5-domain-card::before {
					content: "";
					position: absolute;
					top: 0;
					left: 0;
					width: 100%;
					height: 100%;
					background: linear-gradient(
						120deg,
						transparent 0%,
						rgba(255, 255, 255, 0.08) 50%,
						transparent 100%
					);
					background-size: 200% 100%;
					opacity: 0;
					transition: opacity 0.3s ease;
				}
				.tab5-domain-card:hover::before {
					opacity: 1;
					animation: tab5-shimmer 1.5s ease-in-out infinite;
				}
				.tab5-domain-card:hover {
					transform: translateY(-2px);
					border-color: rgba(0, 200, 255, 0.35);
					background: rgba(10, 20, 40, 0.4);
				}
				</style>
				""", unsafe_allow_html=True)

				for i, row in df_sorted.iterrows():
					progress_value = row['avg_ats'] / 100
					st.markdown(f"""
					<div class="tab5-domain-card">
						<div style="display: flex; justify-content: space-between; align-items: center; position: relative; z-index: 1;">
							<h4 style="margin: 0; color: #5eb8ff;">📁 {row['domain']}</h4>
							<span style="
								background: rgba(0, 200, 255, 0.1);
								border: 1px solid rgba(0, 200, 255, 0.25);
								color: #5eb8ff;
								padding: 5px 10px;
								border-radius: 20px;
								font-size: 12px;
								font-weight: bold;
								backdrop-filter: blur(8px);
							">
								Rank #{i+1}
							</span>
						</div>
						<div style="margin: 10px 0; position: relative; z-index: 1;">
							<div style="
								background: rgba(255, 255, 255, 0.05);
								border-radius: 10px;
								height: 8px;
								overflow: hidden;
							">
								<div style="
									background: linear-gradient(90deg, rgba(0, 200, 255, 0.4), rgba(0, 255, 200, 0.5));
									height: 100%;
									width: {progress_value*100}%;
									transition: width 0.3s ease;
								"></div>
							</div>
						</div>
						<div style="display: flex; justify-content: space-between; margin-top: 10px; position: relative; z-index: 1;">
							<span style="color: #cce6ff;"><b>🧠 Avg ATS:</b> <span style="color: #5eb8ff; font-weight: bold;">{row['avg_ats']:.2f}</span></span>
							<span style="color: #cce6ff;"><b>📄 Resumes:</b> <span style="color: #5eb8ff; font-weight: bold;">{row['count']}</span></span>
						</div>
					</div>
					""", unsafe_allow_html=True)
			else:
				st.info("ℹ️ No domain performance data available.")
		except Exception as e:
			st.error(f"Error loading top domains: {e}")

		# Enhanced Domain Distribution
		st.markdown("### 📊 Domain Distribution Analysis")

		try:
			df_domain_dist = load_domain_distribution()
			if not df_domain_dist.empty:
				col1, col2 = st.columns(2)
				with col1:
					chart_type = st.radio(
						"📊 Visualization Type:",
						["📈 Interactive Bar Chart", "🥧 Interactive Pie Chart"],
						horizontal=True
					)
				with col2:
					max_val = len(df_domain_dist)
					if max_val <= 5:
						show_top_n = max_val  # No slider, just show all available domains
					else:
						show_top_n = st.slider(
							"Show Top N Domains",
							min_value=5,
							max_value=max_val,
							value=min(10, max_val)
						)

				df_top_domains = df_domain_dist.head(show_top_n)

				if chart_type == "📈 Interactive Bar Chart":
					fig = create_enhanced_bar_chart(df_top_domains, "domain", "count", 
											"Resume Count by Domain")
					st.plotly_chart(fig, use_container_width=True)
				else:
					fig = create_enhanced_pie_chart(df_top_domains, "count", "domain", 
											"Domain Distribution")
					st.plotly_chart(fig, use_container_width=True)

				# Summary statistics
				with st.expander("📋 Domain Statistics Summary"):
					st.dataframe(
						df_domain_dist.style.format({'percentage': '{:.2f}%'}),
						use_container_width=True
					)
			else:
				st.info("ℹ️ No domain distribution data available.")
		except Exception as e:
			st.error(f"Error loading domain distribution: {e}")

		# Enhanced ATS Performance Analysis
		st.markdown("### 📈 ATS Performance Analysis")
		
		try:
			df_ats = get_average_ats_by_domain()
			if not df_ats.empty:
				col1, col2 = st.columns(2)
				with col1:
					chart_orientation = st.radio("Chart Style", ["Vertical", "Horizontal"], horizontal=True)
				with col2:
					color_scheme = st.selectbox("Color Scheme", 
										["plasma", "viridis", "inferno", "magma", "turbo"])
				
				orientation = 'v' if chart_orientation == "Vertical" else 'h'
				fig = px.bar(df_ats, 
							x="domain" if orientation == 'v' else "avg_ats_score",
							y="avg_ats_score" if orientation == 'v' else "domain",
							title="Average ATS Score by Domain",
							orientation=orientation,
							color="avg_ats_score",
							color_continuous_scale=color_scheme,
							text="avg_ats_score",
							template="plotly_dark")  # Use dark theme for better readability
				
				fig.update_traces(texttemplate='%{text:.1f}', textposition='outside')
				if orientation == 'v':
					fig.update_xaxes(tickangle=45)
				
				# Enhanced layout for better readability
				fig.update_layout(
					showlegend=False,
					plot_bgcolor='rgba(0,0,0,0.1)',
					paper_bgcolor='rgba(0,0,0,0.05)',
					font=dict(color='white', size=12),
					title=dict(font=dict(size=16, color='white')),
					xaxis=dict(
						gridcolor='rgba(255,255,255,0.2)',
						tickfont=dict(color='white')
					),
					yaxis=dict(
						gridcolor='rgba(255,255,255,0.2)',
						tickfont=dict(color='white')
					),
					margin=dict(t=60, b=80, l=80, r=50)
				)
				
				st.plotly_chart(fig, use_container_width=True)
			else:
				st.info("ℹ️ No ATS performance data available.")
		except Exception as e:
			st.error(f"Error loading ATS performance data: {e}")

		# Enhanced Timeline Analysis
		st.markdown("### 📈 Resume Upload Timeline & Trends")
		
		try:
			df_timeline = get_resume_count_by_day()
			df_daily_ats = get_daily_ats_stats()  # no limit — fetch full history to match upload timeline
			
			if not df_timeline.empty:
				df_timeline = df_timeline.sort_values("day")
				df_timeline["7_day_avg"] = df_timeline["count"].rolling(window=7, min_periods=1).mean()
				df_timeline["30_day_avg"] = df_timeline["count"].rolling(window=30, min_periods=1).mean()
				
				# Create subplot with proper spacing and formatting
				fig = make_subplots(
					rows=2, cols=1,
					subplot_titles=('Daily Upload Count with Moving Averages', 'Daily Average ATS Score Trend'),
					vertical_spacing=0.25,
					specs=[[{"secondary_y": False}], [{"secondary_y": False}]]
				)
				
				# Convert day column to datetime for proper spacing
				df_timeline['day'] = pd.to_datetime(df_timeline['day'])
				
				# Upload count plot
				fig.add_trace(
					go.Scatter(x=df_timeline["day"], y=df_timeline["count"], 
								mode='lines+markers', name='Daily Uploads',
								line=dict(color='#1f77b4', width=2),
								marker=dict(size=6)),
					row=1, col=1
				)
				
				fig.add_trace(
					go.Scatter(x=df_timeline["day"], y=df_timeline["7_day_avg"], 
								mode='lines', name='7-Day Average',
								line=dict(color='#ff7f0e', width=2, dash='dash')),
					row=1, col=1
				)
				
				fig.add_trace(
					go.Scatter(x=df_timeline["day"], y=df_timeline["30_day_avg"], 
								mode='lines', name='30-Day Average',
								line=dict(color='#2ca02c', width=2, dash='dot')),
					row=1, col=1
				)
				
				# ATS trend plot
				if not df_daily_ats.empty:
					df_daily_ats['date'] = pd.to_datetime(df_daily_ats['date'])
					fig.add_trace(
						go.Scatter(x=df_daily_ats["date"], y=df_daily_ats["avg_ats"],
									mode='lines+markers', name='Daily Avg ATS',
									line=dict(color='#d62728', width=2),
									marker=dict(size=6)),
						row=2, col=1
					)

				# Compute a SHARED date range across both datasets so row2 starts
				# from the same date as row1 (not just from when ATS data begins)
				all_dates = list(df_timeline["day"])
				if not df_daily_ats.empty:
					all_dates += list(df_daily_ats["date"])
				global_min = min(all_dates)
				global_max = max(all_dates)
				pad = pd.Timedelta(hours=12)
				x_range = [global_min - pad, global_max + pad]

				# Pick tick density based on total date span
				total_days = (global_max - global_min).days
				MS_PER_DAY = 86400000
				if total_days <= 14:
					dtick = MS_PER_DAY
					tickfmt = "%b %d"
				elif total_days <= 60:
					dtick = 7 * MS_PER_DAY
					tickfmt = "%b %d"
				elif total_days <= 180:
					dtick = 14 * MS_PER_DAY
					tickfmt = "%b %d '%y"
				else:
					dtick = "M1"
					tickfmt = "%b '%y"

				# Update layout for better spacing and readability
				fig.update_layout(
					height=800,
					showlegend=True,
					legend=dict(
						orientation="h",
						yanchor="bottom",
						y=1.02,
						xanchor="right",
						x=1
					),
					margin=dict(t=80, b=70, l=50, r=50)
				)

				# Apply identical range + ticks to BOTH subplots
				shared_xaxis = dict(
					range=x_range,
					tickformat=tickfmt,
					tickangle=45,
					dtick=dtick,
					ticklabelmode="period",
				)
				fig.update_xaxes(title_text="Date", row=2, col=1)
				fig.update_xaxes(**shared_xaxis, row=1, col=1)
				fig.update_xaxes(**shared_xaxis, row=2, col=1)
				
				fig.update_yaxes(title_text="Upload Count", row=1, col=1)
				fig.update_yaxes(title_text="Average ATS Score", row=2, col=1)
				
				st.plotly_chart(fig, use_container_width=True)
				
				# Timeline statistics
				col1, col2, col3, col4 = st.columns(4)
				with col1:
					st.metric("Total Days", len(df_timeline))
				with col2:
					st.metric("Peak Daily Uploads", df_timeline["count"].max())
				with col3:
					st.metric("Avg Daily Uploads", f"{df_timeline['count'].mean():.1f}")
				with col4:
					if not df_daily_ats.empty:
						st.metric("Avg ATS Trend", f"{df_daily_ats['avg_ats'].mean():.2f}")
			else:
				st.info("ℹ️ No timeline data available.")
		except Exception as e:
			st.error(f"Error loading timeline data: {e}")

		# Enhanced Bias Analysis
		st.markdown("### 🧠 Advanced Bias Analysis")
		
		col1, col2 = st.columns(2)
		with col1:
			bias_threshold_pie = st.slider("Bias Detection Threshold", 
									min_value=0.0, max_value=1.0, value=0.6, step=0.05)
		with col2:
			analysis_type = st.radio("Analysis Type", ["Distribution", "Flagged Candidates"], horizontal=True)
		
		try:
			if analysis_type == "Distribution":
				df_bias = get_bias_distribution(threshold=bias_threshold_pie)
				if not df_bias.empty and "bias_category" in df_bias.columns:
					fig = create_enhanced_pie_chart(df_bias, "count", "bias_category", 
											f"Bias Distribution (Threshold: {bias_threshold_pie})")
					st.plotly_chart(fig, use_container_width=True)
					
					# Bias statistics
					col1, col2 = st.columns(2)
					with col1:
						total_candidates = df_bias["count"].sum()
						biased_count = df_bias[df_bias["bias_category"] == "Biased"]["count"].iloc[0] if len(df_bias[df_bias["bias_category"] == "Biased"]) > 0 else 0
						st.metric("Total Analyzed", total_candidates)
					with col2:
						bias_percentage = (biased_count / total_candidates * 100) if total_candidates > 0 else 0
						st.metric("Bias Percentage", f"{bias_percentage:.1f}%")
				else:
					st.info("📭 No bias distribution data available.")
			
			else:  # Flagged Candidates
				flagged_df = get_flagged_candidates(threshold=bias_threshold_pie)
				if not flagged_df.empty:
					st.markdown(f"**🚩 {len(flagged_df)} candidates flagged with bias score > {bias_threshold_pie}**")
					
					# Enhanced flagged candidates display
					display_df = flagged_df.copy()
					display_df = display_df.sort_values('bias_score', ascending=False)
					
					st.dataframe(
						display_df.style.format({'bias_score': '{:.3f}', 'ats_score': '{:.0f}'}),
						use_container_width=True,
						height=300
					)
					
					# Flagged candidates statistics
					col1, col2, col3 = st.columns(3)
					with col1:
						st.metric("Flagged Count", len(flagged_df))
					with col2:
						st.metric("Avg Bias Score", f"{flagged_df['bias_score'].mean():.3f}")
					with col3:
						st.metric("Avg ATS Score", f"{flagged_df['ats_score'].mean():.1f}")
				else:
					st.success("✅ No candidates flagged above the selected threshold.")
		except Exception as e:
			st.error(f"Error in bias analysis: {e}")

		# Enhanced Domain Performance Deep Dive
		with st.expander("🔍 Domain Performance Deep Dive", expanded=False):
			try:
				df_performance = get_domain_performance_stats()
				if not df_performance.empty:
					st.markdown("#### Comprehensive Domain Performance Metrics")
					
					# Performance heatmap
					performance_cols = ['avg_ats_score', 'avg_edu_score', 'avg_exp_score',
								'avg_skills_score', 'avg_lang_score', 'avg_keyword_score',
								'avg_format_score']
					
					if all(col in df_performance.columns for col in performance_cols):
						heatmap_data = df_performance[['domain'] + performance_cols].set_index('domain')
						
						fig = px.imshow(heatmap_data.T, 
									title="Domain Performance Heatmap",
									color_continuous_scale="RdYlGn",
									aspect="auto")
						fig.update_layout(height=400)
						st.plotly_chart(fig, use_container_width=True)
					
					# Detailed performance table
					st.dataframe(
						df_performance.style.format({
							col: '{:.2f}' for col in performance_cols + ['avg_bias_score']
						}),
						use_container_width=True
					)
				else:
					st.info("ℹ️ No detailed performance data available.")
			except Exception as e:
				st.error(f"Error loading performance deep dive: {e}")

		# Footer with system information
		st.markdown("<hr style='border-top: 1px solid #ddd; margin: 2rem 0;'>", unsafe_allow_html=True)
		st.markdown("""
		<div style='text-align: center; color: #666; font-size: 0.9em; padding: 1rem;'>
			<p>🛡️ Enhanced Admin Dashboard | Powered by Advanced Database Manager</p>
			<p>Last updated: {}</p>
		</div>
		""".format(datetime.now().strftime("%Y-%m-%d %H:%M:%S")), unsafe_allow_html=True)







