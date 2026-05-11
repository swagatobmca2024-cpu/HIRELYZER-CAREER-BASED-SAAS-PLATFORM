def _sanitize_html_for_pdf(html_string):
    """
    Strip / replace CSS properties that xhtml2pdf (pisa) does not support.

    xhtml2pdf uses a very limited CSS 2.1 subset and crashes with a
    CSSParseError on any property it cannot parse — including all of the
    modern layout primitives used by the resume templates for the browser
    preview:  display:flex, flex-wrap, gap, align-items, justify-content,
    flex-direction, flex-shrink, object-fit, object-position, letter-spacing,
    text-transform (partial), border-radius on some elements, box-shadow,
    background-clip, linear-gradient(), rgba() in some positions, etc.

    Strategy:
      1. Use regex to find every inline style="..." attribute.
      2. Within each style block, strip individual property:value pairs that
         are known to crash the parser, while keeping PDF-safe ones intact.
      3. Replace display:flex → display:block  so containers still render.
      4. Replace gap:... → margin-bottom on the container (best-effort).

    This sanitisation is applied ONLY for PDF export — the browser preview
    still uses the full modern CSS.
    """
    import re as _re

    # Properties whose entire declaration should be dropped for PDF safety.
    # Each entry is a regex that matches  "property-name:value"  (no semicolon).
    STRIP_PROPS = [
        r'flex-wrap\s*:[^;]*',
        r'flex-direction\s*:[^;]*',
        r'flex-shrink\s*:[^;]*',
        r'flex-grow\s*:[^;]*',
        r'flex\s*:[^;]*',           # shorthand — must come after flex-*
        r'align-items\s*:[^;]*',
        r'align-self\s*:[^;]*',
        r'justify-content\s*:[^;]*',
        r'justify-self\s*:[^;]*',
        r'gap\s*:[^;]*',
        r'row-gap\s*:[^;]*',
        r'column-gap\s*:[^;]*',
        r'grid[^:]*:[^;]*',         # any grid-* property
        r'object-fit\s*:[^;]*',
        r'object-position\s*:[^;]*',
        r'box-shadow\s*:[^;]*',
        r'text-shadow\s*:[^;]*',
        r'background-clip\s*:[^;]*',
        r'-webkit-[^:]*:[^;]*',     # any vendor prefix
        r'-moz-[^:]*:[^;]*',
        r'-ms-[^:]*:[^;]*',
        r'transition\s*:[^;]*',
        r'transform\s*:[^;]*',
        r'animation[^:]*:[^;]*',
        r'will-change\s*:[^;]*',
        r'pointer-events\s*:[^;]*',
        r'resize\s*:[^;]*',
        r'cursor\s*:[^;]*',
        r'overflow-x\s*:[^;]*',
        r'overflow-y\s*:[^;]*',
        r'overflow\s*:\s*(?!hidden)[^;]*',   # keep overflow:hidden, drop others
        r'white-space\s*:[^;]*',
        r'word-break\s*:[^;]*',
        r'overflow-wrap\s*:[^;]*',
        r'text-overflow\s*:[^;]*',
    ]

    # background: linear-gradient(...) crashes pisa — replace with a flat colour
    GRADIENT_RE = _re.compile(
        r'background\s*:\s*linear-gradient\([^)]*\)\s*(?:;|$)', _re.IGNORECASE
    )

    # display:flex → display:block  (pisa only knows block/inline/table/none)
    FLEX_DISPLAY_RE = _re.compile(
        r'display\s*:\s*flex\b', _re.IGNORECASE
    )
    INLINE_FLEX_RE = _re.compile(
        r'display\s*:\s*inline-flex\b', _re.IGNORECASE
    )

    # Compile all strip patterns once
    strip_re_list = [_re.compile(p, _re.IGNORECASE) for p in STRIP_PROPS]

    def _clean_style(style_value):
        """Clean a single style="..." value string."""
        # Replace gradients with a neutral background
        style_value = GRADIENT_RE.sub('background:#f5f5f5;', style_value)
        # Replace flex displays
        style_value = FLEX_DISPLAY_RE.sub('display:block', style_value)
        style_value = INLINE_FLEX_RE.sub('display:inline-block', style_value)
        # Strip unsupported properties
        for pat in strip_re_list:
            style_value = pat.sub('', style_value)
        # Clean up leftover semicolons / whitespace
        # e.g. ";;  ;" → ";"
        style_value = _re.sub(r'\s*;\s*;+', ';', style_value)
        style_value = _re.sub(r'^\s*;+', '', style_value)
        style_value = style_value.strip().strip(';')
        return style_value

    # Match every  style="..."  or  style='...'  attribute in the HTML
    def _replace_style_attr(m):
        quote = m.group(1)           # ' or "
        style_content = m.group(2)
        cleaned = _clean_style(style_content)
        if not cleaned.strip():
            return ''                 # remove empty style attributes entirely
        return f'style={quote}{cleaned}{quote}'

    sanitized = _re.sub(
        r'''style=(['"])(.*?)\1''',
        _replace_style_attr,
        html_string,
        flags=_re.DOTALL | _re.IGNORECASE,
    )

    # Also strip any <style> blocks that contain flex/grid — pisa parses
    # embedded <style> tags through its CSS engine and will also crash there.
    def _clean_style_tag(m):
        css_text = m.group(1)
        css_text = GRADIENT_RE.sub('background:#f5f5f5;', css_text)
        css_text = FLEX_DISPLAY_RE.sub('display:block', css_text)
        css_text = INLINE_FLEX_RE.sub('display:inline-block', css_text)
        for pat in strip_re_list:
            css_text = pat.sub('', css_text)
        return f'<style>{css_text}</style>'

    sanitized = _re.sub(
        r'<style[^>]*>(.*?)</style>',
        _clean_style_tag,
        sanitized,
        flags=_re.DOTALL | _re.IGNORECASE,
    )

    return sanitized


def html_to_pdf_bytes(html_string):
    # NOTE: Do NOT use f-string here. Template HTML contains CSS variables like
    # {C_PRIMARY} which Python re-evaluates as f-string placeholders → CSSParseError crash.
    # Use plain string + .replace() to safely inject html_string.

    # Sanitise modern CSS that xhtml2pdf cannot parse before wrapping.
    safe_html = _sanitize_html_for_pdf(html_string)

    wrapper = """
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {
                size: A4 portrait;
                margin: 0;
            }
            html, body {
                margin: 0;
                padding: 0;
                font-size: 12pt;
                font-family: "Segoe UI", "Helvetica", sans-serif;
                line-height: 1.5;
                color: #000;
                background: #fff;
            }
            table {
                width: 100%;
                border-collapse: collapse;
            }
        </style>
    </head>
    <body>
        __HTML_CONTENT__
    </body>
    </html>
    """
    styled_html = wrapper.replace("__HTML_CONTENT__", safe_html)

    pdf_io = BytesIO()
    pisa.CreatePDF(styled_html, dest=pdf_io)
    pdf_io.seek(0)
    return pdf_io
# ══════════════════════════════════════════════════════════════════════════════
# IMPORTS — modular resume & cover letter engines
# ══════════════════════════════════════════════════════════════════════════════
from resume_builder import (
    render_template_default, render_template_modern, render_template_sidebar,
    render_template_classic, render_template_executive, render_template_timeline,
    render_template_corporate, render_template_creative_green,
    render_template_terracotta, render_template_navy_prestige,
    render_template_slate_gray, render_template_teal_impact,
    render_template_burgundy_classic, render_template_indigo_tech,
    render_template_forest_green,
    render_template_pure_white, render_template_midnight_black,
    render_template_soft_lavender, render_template_warm_sand,
    render_template_ice_blue,
    RESUME_TEMPLATES, render_resume,
    _fmt_desc, _cert_name_html,
)
from cover_letter import (
    render_cover_letter_professional, render_cover_letter_modern,
    render_cover_letter_creative, render_cover_letter_executive,
    render_cover_letter_entry_level, render_cover_letter_ats,
    COVER_LETTER_TEMPLATES, render_cover_letter,
    generate_cover_letter_from_resume_builder,
)

from collections import Counter as _Counter

# ── Filler phrase / word constants ────────────────────────────────────────────
_FILLER_PHRASES = {
    "lorem ipsum", "test test", "sample text", "placeholder",
    "your text here", "enter here", "tbd", "todo",
    "fill in", "coming soon", "to be added", "add here",
    "description here", "write here", "dummy text",
    "hello world", "foo bar", "asdf", "qwerty",
}
_FILLER_WORDS = {"placeholder", "tbd", "todo", "dummy", "example", "test", "asdf", "qwerty"}


def detect_garbage_text(text: str) -> bool:
    """
    Returns True when the text is meaningless / garbage — should earn 0 XP.

    Detection layers (cheapest first):
      1. Empty / whitespace-only
      2. Pure numeric string  ("123456789", "+91 98765")
      3. Known filler/placeholder phrases (multi-word substring match)
      4. All-same character repeated  ("aaaaaaa", "--------")
      5. Keyboard-mash: any token > 18 chars with < 12% vowels
      6. Single long token (> 20 chars) with < 15% vowels
      7. Repeated-token spam: dominant word > 55% of all tokens (≥ 3 tokens)
      8. Very low unique-word ratio: < 35% unique among 4+ word inputs
      9. Pathological token length: avg < 2.0 chars with > 60% short tokens
    """
    t = str(text).strip()
    if not t:
        return True

    # 1. Pure numeric (phone-safe: strip spaces, +, -, .)
    stripped_num = t.replace(" ", "").replace("-", "").replace("+", "").replace(".", "")
    if stripped_num.isdigit() and len(stripped_num) > 5:
        return True

    t_lower = t.lower()

    # 2. Known filler phrases (substring match)
    for phrase in _FILLER_PHRASES:
        if phrase in t_lower:
            return True

    # 3. "n/a" exact or sole token
    if t_lower.strip() in {"n/a", "na", "n.a.", "n.a"}:
        return True

    # 4. Single-word filler: whole-token match only
    t_tokens_lower = [tok.lower() for tok in t.split() if tok]
    if len(t_tokens_lower) <= 2 and any(w in _FILLER_WORDS for w in t_tokens_lower):
        return True

    # 5. All-same character ("aaaaaaa", "-------")
    stripped_chars = t.replace(" ", "")
    if len(stripped_chars) >= 3 and len(set(stripped_chars.lower())) == 1:
        return True

    tokens = [tok for tok in t.replace(",", " ").replace(";", " ").split() if tok]
    if not tokens:
        return True

    # 5b. Per-token all-same-character check: if majority of tokens are
    #     single-repeated-char words (e.g. "aaaaaa bbbb cccc"), reject.
    #     Threshold: ≥ 60% of tokens are all-same-char and len ≥ 3.
    _same_char_count = sum(
        1 for tok in tokens
        if len(tok) >= 3 and len(set(tok.lower())) == 1
    )
    if len(tokens) >= 2 and (_same_char_count / len(tokens)) >= 0.60:
        return True

    # 6. Any token > 18 chars with < 12% vowels → keyboard mash
    for tok in tokens:
        if len(tok) > 18:
            vowels = sum(1 for c in tok.lower() if c in "aeiou")
            if vowels / len(tok) < 0.12:
                return True

    # 7. Single mega-token keyboard mash (> 20 chars, < 15% vowels)
    if len(tokens) == 1 and len(tokens[0]) > 20:
        vowels = sum(1 for c in tokens[0].lower() if c in "aeiou")
        if vowels / len(tokens[0]) < 0.15:
            return True

    # 8. Repeated-token spam: one word > 55% of all tokens
    lowered = [tok.lower() for tok in tokens]
    if len(tokens) >= 3:
        most_common_count = _Counter(lowered).most_common(1)[0][1]
        if most_common_count / len(tokens) > 0.55:
            return True

    # 9. Low unique-word ratio: for 4+ word inputs, require ≥ 35% unique
    if len(tokens) >= 4:
        if calculate_unique_word_ratio(t) < 0.35:
            return True

    # 10. Pathological token length pattern
    avg_len = sum(len(tok) for tok in tokens) / len(tokens)
    short_ratio = sum(1 for tok in tokens if len(tok) <= 2) / len(tokens)
    if avg_len < 2.0 and short_ratio > 0.60:
        return True

    return False


def is_meaningful_text(text: str, min_words: int = 5) -> bool:
    """
    Returns True when text passes quality gates and meets the minimum word count.
    Inverse of detect_garbage_text with an additional word-count floor.

    Parameters
    ----------
    text      : the raw string to validate
    min_words : minimum number of space-separated tokens required (default 5)

    Usage
    -----
    Suitable for experience / project / education description fields.
    Short-form fields (skills, cert name) should use detect_garbage_text directly.
    """
    t = str(text).strip()
    if detect_garbage_text(t):
        return False
    tokens = [tok for tok in t.split() if tok]
    return len(tokens) >= min_words


def detect_repeated_words(text: str, threshold: float = 0.55) -> bool:
    """
    Returns True when a single word dominates more than `threshold` fraction
    of all tokens — indicating repetitive / spam input.

    Parameters
    ----------
    text      : raw input string
    threshold : fraction above which dominant word is considered spam (default 0.55)
    """
    tokens = [tok.lower() for tok in str(text).split() if tok]
    if len(tokens) < 3:
        return False
    most_common_count = _Counter(tokens).most_common(1)[0][1]
    return (most_common_count / len(tokens)) > threshold


def calculate_unique_word_ratio(text: str) -> float:
    """
    Returns the ratio of unique lowercase words to total words (0.0 – 1.0).
    A ratio of 1.0 means every word is unique; 0.0 means only one distinct word.
    Returns 0.0 for empty / single-token input.

    Example
    -------
    "hello world hello" → 2 unique / 3 total → 0.667
    """
    tokens = [tok.lower() for tok in str(text).split() if tok]
    if not tokens:
        return 0.0
    return round(len(set(tokens)) / len(tokens), 4)


# ── Internal quality scorer (used by all section scorers) ─────────────────────

def _text_quality_score(text: str, min_words: int = 3) -> float:
    """
    Returns a quality multiplier 0.0–1.0 for free-text fields.
    Garbage inputs always score 0.0.
    Rewards: sufficient word count + unique-word diversity.
    """
    t = str(text).strip()
    if not t:
        return 0.0
    if detect_garbage_text(t):
        return 0.0

    tokens = [tok for tok in t.split() if tok]
    word_count = len(tokens)

    if word_count < min_words:
        return 0.0

    unique_ratio = calculate_unique_word_ratio(t)
    diversity_mult = 1.0
    if word_count >= 6 and unique_ratio < 0.50:
        diversity_mult = max(0.30, min(1.0, unique_ratio * 1.5))

    if word_count < 5:
        base = 0.25
    elif word_count < 15:
        base = 0.55
    elif word_count < 30:
        base = 0.80
    elif word_count < 50:
        base = 0.92
    else:
        base = 1.0

    return round(min(base * diversity_mult, 1.0), 3)


def _desc_score(text: str) -> float:
    """
    Score a long-form description field (experience / project / education / cert).
    Combines word-count quality with character-length signal.
    Garbage → 0.0 always.
    """
    t = str(text).strip()
    if not t or detect_garbage_text(t):
        return 0.0

    char_len = len(t)
    word_score = _text_quality_score(t, min_words=4)

    if char_len < 30:
        char_tier = 0.20
    elif char_len < 80:
        char_tier = 0.55
    elif char_len < 180:
        char_tier = 0.82
    else:
        char_tier = 1.0

    return round(min((word_score * 0.60) + (char_tier * 0.40), 1.0), 3)


def _count_valid_tokens(raw_str: str) -> int:
    """Return count of unique, non-empty, non-garbage comma-separated tokens."""
    if not raw_str:
        return 0
    seen = set()
    count = 0
    for tok in raw_str.split(","):
        tok = tok.strip()
        if not tok:
            continue
        tok_norm = tok.lower()
        if tok_norm in seen:
            continue
        seen.add(tok_norm)
        if not detect_garbage_text(tok):
            count += 1
    return count


# ── Industry-standard XP weights ──────────────────────────────────────────────
# Inspired by LinkedIn Profile Strength, Indeed Resume Score, and Jobscan:
#   Experience    → 30 XP  (highest — core employability signal)
#   Projects      → 22 XP  (second — demonstrates practical skills)
#   Skills & More → 16 XP  (medium — keyword matching & breadth)
#   Education     → 14 XP  (medium — credential verification)
#   Summary       →  7 XP  (differentiator — first impression text)
#   Certificates  →  5 XP  (bonus — validated expertise)
#   Personal Info →  4 XP  (baseline — completeness check)
#   Contact       →  2 XP  (small — reachability signal)
#                  ─────
#   TOTAL MAX     → 100 XP
XP_WEIGHTS = {
    "Experience":    30,
    "Projects":      22,
    "Skills & More": 16,
    "Education":     14,
    "Summary":        7,
    "Certificates":   5,
    "Personal Info":  4,
    "Contact":        2,
}
XP_TOTAL_MAX = sum(XP_WEIGHTS.values())  # 100


def score_experience_section(experience_entries: list) -> float:
    """
    Score the Experience section. Returns float 0.0–1.0.

    Scoring per entry:
      Required (40%): job title + company — both validated (non-garbage)
      Quality  (60%): duration 20% + rich description 80%
    Cap: 1 entry → max 0.80; 2+ entries can reach 1.0.
    """
    if not experience_entries:
        return 0.0
    n = len(experience_entries)
    total = 0.0
    for e in experience_entries:
        title   = str(e.get("title",       "")).strip()
        company = str(e.get("company",     "")).strip()
        dur     = str(e.get("duration",    "")).strip()
        desc    = str(e.get("description", "")).strip()
        if not title and not company:
            continue
        title_ok   = bool(title)   and not detect_garbage_text(title)
        company_ok = bool(company) and not detect_garbage_text(company)
        req   = (float(title_ok) + float(company_ok)) / 2
        qual  = (float(bool(dur)) * 0.20) + (_desc_score(desc) * 0.80)
        total += (req * 0.40) + (qual * 0.60)
    avg = total / n
    if n == 1:
        avg = min(avg, 0.80)
    return round(min(avg, 1.0), 3)


def score_education_section(education_entries: list) -> float:
    """
    Score the Education section. Returns float 0.0–1.0.

    Scoring per entry:
      Required (40%): institution + degree — both validated
      Quality  (60%): year 20% + academic details 80%
    Cap: 1 entry → max 0.85; 2+ can reach 1.0.
    """
    if not education_entries:
        return 0.0
    n = len(education_entries)
    total = 0.0
    for e in education_entries:
        inst   = str(e.get("institution", "")).strip()
        degree = str(e.get("degree",      "")).strip()
        if isinstance(degree, list):
            degree = ", ".join(degree)
        year   = str(e.get("year",    "")).strip()
        det    = str(e.get("details", "")).strip()
        if not inst and not degree:
            continue
        inst_ok   = bool(inst)   and not detect_garbage_text(inst)
        degree_ok = bool(degree) and not detect_garbage_text(degree)
        req  = (float(inst_ok) + float(degree_ok)) / 2
        qual = (float(bool(year)) * 0.20) + (_desc_score(det) * 0.80)
        total += (req * 0.40) + (qual * 0.60)
    avg = total / n
    if n == 1:
        avg = min(avg, 0.85)
    return round(min(avg, 1.0), 3)


def score_project_section(project_entries: list, project_links: list = None) -> float:
    """
    Score the Projects section. Returns float 0.0–1.0.

    Scoring per entry:
      Required (35%): title (validated) + tech stack (≥ 1 valid token)
      Quality  (65%): duration 15% + rich description 85%
    Cap: 1 project → max 0.75; 2 → max 0.90; 3+ → 1.0.
    """
    if not project_entries:
        return 0.0
    n = len(project_entries)
    total = 0.0
    for e in project_entries:
        title = str(e.get("title",       "")).strip()
        tech  = str(e.get("tech",        "")).strip()
        dur   = str(e.get("duration",    "")).strip()
        desc  = str(e.get("description", "")).strip()
        if not title:
            continue
        title_ok  = not detect_garbage_text(title)
        tech_cnt  = _count_valid_tokens(tech)
        tech_sc   = min(tech_cnt / 2.0, 1.0)
        req  = (float(title_ok) + tech_sc) / 2
        qual = (float(bool(dur)) * 0.15) + (_desc_score(desc) * 0.85)
        total += (req * 0.35) + (qual * 0.65)
    avg = total / n
    if n == 1:
        avg = min(avg, 0.75)
    elif n == 2:
        avg = min(avg, 0.90)
    return round(min(avg, 1.0), 3)


def score_certificate_section(certificate_entries: list) -> float:
    """
    Score the Certificates section. Returns float 0.0–1.0.

    Scoring per entry:
      Required (30%): certificate name (validated — non-garbage)
      Quality  (70%): link 20% + duration 20% + description 60%
    1 complete cert CAN reach 1.0 (certs are optional, 1 complete = excellent).
    """
    if not certificate_entries:
        return 0.0
    n = len(certificate_entries)
    total = 0.0
    for e in certificate_entries:
        name = str(e.get("name",        "")).strip()
        link = str(e.get("link",        "")).strip()
        dur  = str(e.get("duration",    "")).strip()
        desc = str(e.get("description", "")).strip()
        if not name:
            continue
        name_ok = not detect_garbage_text(name)
        req  = float(name_ok)
        qual = (float(bool(link)) * 0.20) + (float(bool(dur)) * 0.20) + (_desc_score(desc) * 0.60)
        total += (req * 0.30) + (qual * 0.70)
    return round(min(total / n, 1.0), 3)


def score_skills_section(
    skills: str,
    soft_skills: str = "",
    interests: str = "",
    languages: str = "",
) -> float:
    """
    Score the Skills & More composite section. Returns float 0.0–1.0.

    Sub-weights: Skills 50% | Soft Skills 20% | Interests 20% | Languages 10%
    Each sub-score uses _count_valid_tokens for garbage rejection.
    Skills: 0 → 0.0, 1 → 0.20, 2 → 0.45, 3 → 0.65, 4 → 0.82, 5+ → 1.0
    Soft/Interests: 0 → 0.0, 1 → 0.35, 2 → 0.70, 3+ → 1.0
    Languages: 0 → 0.0, 1 → 0.50, 2+ → 1.0
    """
    skill_count = _count_valid_tokens(skills)
    if skill_count == 0:    sub_skills = 0.0
    elif skill_count == 1:  sub_skills = 0.20
    elif skill_count == 2:  sub_skills = 0.45
    elif skill_count == 3:  sub_skills = 0.65
    elif skill_count == 4:  sub_skills = 0.82
    else:                   sub_skills = 1.0

    interest_count = _count_valid_tokens(interests)
    sub_interests = 0.0 if interest_count == 0 else (0.35 if interest_count == 1 else (0.70 if interest_count == 2 else 1.0))

    soft_count = _count_valid_tokens(soft_skills)
    sub_soft = 0.0 if soft_count == 0 else (0.35 if soft_count == 1 else (0.70 if soft_count == 2 else 1.0))

    lang_count = _count_valid_tokens(languages)
    sub_lang = 0.0 if lang_count == 0 else (0.50 if lang_count == 1 else 1.0)

    return round(
        (sub_skills * 0.50) + (sub_interests * 0.20) + (sub_soft * 0.20) + (sub_lang * 0.10),
        3
    )


def calculate_resume_xp(section_fills: dict, weights: dict = None) -> tuple:
    """
    Compute total resume XP from section fill scores and weights.
    XP is capped at XP_TOTAL_MAX (100).

    Parameters
    ----------
    section_fills : dict[section_name → float 0.0–1.0]
    weights       : dict[section_name → int] — defaults to module XP_WEIGHTS

    Returns
    -------
    (raw_xp: float, xp: int, pct: int)

    Usage
    -----
    xp_score = calculate_resume_xp(st.session_state)  # legacy call
    raw, xp, pct = calculate_resume_xp(fills, XP_WEIGHTS)
    """
    if weights is None:
        weights = XP_WEIGHTS
    raw_xp = sum(section_fills.get(k, 0.0) * weights.get(k, 0) for k in weights)
    raw_xp = min(raw_xp, XP_TOTAL_MAX)
    xp_int = int(round(raw_xp))
    pct    = int(round((raw_xp / XP_TOTAL_MAX) * 100))
    return raw_xp, xp_int, pct


# ── Public aliases used in external imports / tests ───────────────────────────
score_experience_section.__module__ = __name__
score_project_section.__module__    = __name__
score_education_section.__module__  = __name__
score_certificate_section.__module__ = __name__
score_skills_section.__module__     = __name__

# ══════════════════════════════════════════════════════════════════════════════
# END OF MODULE-LEVEL XP ENGINE
# ══════════════════════════════════════════════════════════════════════════════

# Import necessary modules first
import streamlit as st
import time

# Tab setup (assuming this is within a tab2 context)
with tab2:
    st.session_state.active_tab = "Resume Builder"

    # ---------- Title with Blue Glassmorphism + Shine ----------
    st.markdown("""
    <style>
    .glass-title {
        background: rgba(10, 20, 40, 0.5);
        border-radius: 20px;
        padding: 20px;
        backdrop-filter: blur(14px);
        box-shadow: 0 8px 32px rgba(0, 200, 255, 0.25);
        border: 1px solid rgba(0, 200, 255, 0.3);
        text-align: center;
        position: relative;
        overflow: hidden;
    }
    .glass-title h2 {
        color: #4da6ff;
        margin: 0;
        text-shadow: 0 0 12px rgba(0,200,255,0.7);
        font-weight: 600;
    }
    .glass-title::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.18) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .glass-title:hover::before {
        left: 100%;
        top: 100%;
    }
    </style>

    <div class="glass-title">
        <h2>🧾 Advanced Resume Builder</h2>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<hr style='border-top: 2px solid rgba(0,200,255,0.4);'>", unsafe_allow_html=True)

    # ---------- Anti-Flicker / Smooth Rerun CSS ----------
    st.markdown("""
        <style>
        /* Prevent white flash and blinking on Streamlit reruns */
        [data-testid="stAppViewContainer"],
        [data-testid="stVerticalBlock"],
        [data-testid="stForm"],
        [data-testid="stSidebar"],
        section.main > div {
            transition: opacity 0.15s ease-in-out !important;
        }

        /* Suppress the brief layout jump when widgets remount */
        iframe, [data-testid="stIFrame"] {
            transition: none !important;
        }

        /* Prevent scrollbar flicker */
        html {
            overflow-y: scroll !important;
            scrollbar-gutter: stable !important;
        }

        /* Smooth button press — no jump */
        button[kind="formSubmit"],
        button[kind="secondary"],
        .stButton > button {
            transition: background-color 0.2s ease, box-shadow 0.2s ease, transform 0.1s ease !important;
        }
        .stButton > button:active {
            transform: scale(0.98) !important;
        }

        /* Prevent input field flicker on focus/blur */
        input, textarea, select {
            transition: border-color 0.15s ease, box-shadow 0.15s ease !important;
        }

        /* Prevent expander flicker */
        details summary {
            transition: background 0.2s ease !important;
        }

        /* Prevent layout shift during widget remounts */
        [data-testid="stVerticalBlock"] > div {
            min-height: 0 !important;
        }

        /* Smooth fade-in for newly rendered blocks */
        @keyframes fadeInBlock {
            from { opacity: 0.6; }
            to   { opacity: 1; }
        }
        [data-testid="stVerticalBlock"] {
            animation: fadeInBlock 0.12s ease-out !important;
        }
        </style>

        <script>
        (function() {
            var lastScrollY = 0;
            var ticking = false;
            var observer = new MutationObserver(function() {
                if (!ticking) {
                    requestAnimationFrame(function() {
                        if (Math.abs(window.scrollY - lastScrollY) > 200) {
                            window.scrollTo({ top: lastScrollY, behavior: 'instant' });
                        }
                        ticking = false;
                    });
                    ticking = true;
                }
            });
            observer.observe(document.body, { childList: true, subtree: false });
            window.addEventListener('scroll', function() {
                lastScrollY = window.scrollY;
            }, { passive: true });
        })();
        </script>
    """, unsafe_allow_html=True)

    # ---------- Global Styles (Glassmorphism + Glow + Shine) ----------
    st.markdown("""
        <style>
        /* File uploader */
        .uploadedFile { 
            background: rgba(10, 20, 40, 0.6) !important;
            border: 1px solid rgba(0,200,255,0.5) !important;
            border-radius: 14px !important;
            color: #cce6ff !important;
            box-shadow: 0 0 12px rgba(0,200,255,0.3) !important;
        }

        /* Sidebar expander style */
        .streamlit-expanderHeader {
            background: rgba(10, 20, 40, 0.45);
            border-radius: 12px;
            color: #4da6ff !important;
            font-weight: bold;
            backdrop-filter: blur(12px);
            box-shadow: 0 4px 12px rgba(0,200,255,0.25);
            transition: all 0.3s ease-in-out;
        }
        .streamlit-expanderHeader:hover {
            background: rgba(0, 200, 255, 0.12);
            box-shadow: 0 0 16px rgba(0,200,255,0.4);
        }
        .streamlit-expanderContent {
            background: rgba(10, 20, 40, 0.45);
            border-radius: 10px;
            padding: 8px;
            color: #e6f7ff;
        }

        /* Selectbox */
        div[data-baseweb="select"] {
            background: rgba(10, 20, 40, 0.35);
            border: 1px solid rgba(0, 200, 255, 0.6);
            border-radius: 12px;
            color: #e6f7ff;
            backdrop-filter: blur(14px);
            box-shadow: 0 0 10px rgba(0,200,255,0.3);
        }

        /* Buttons with Shine Effect */
        div.stButton > button {
            position: relative;
            background: rgba(10, 20, 40, 0.35);
            border: 1px solid rgba(0, 200, 255, 0.6);
            color: #e6f7ff;
            border-radius: 14px;
            padding: 10px 20px;
            font-size: 15px;
            font-weight: 500;
            backdrop-filter: blur(16px);
            box-shadow: 0 0 12px rgba(0, 200, 255, 0.35),
                        inset 0 0 20px rgba(0, 200, 255, 0.05);
            overflow: hidden;
            transition: all 0.3s ease-in-out;
        }
        div.stButton > button::before {
            content: "";
            position: absolute;
            top: -50%;
            left: -50%;
            width: 200%;
            height: 200%;
            background: linear-gradient(
                120deg,
                rgba(255,255,255,0.15) 0%,
                rgba(255,255,255,0.05) 40%,
                transparent 60%
            );
            transform: rotate(25deg);
            transition: all 0.6s;
        }
        div.stButton > button:hover::before {
            left: 100%;
            top: 100%;
        }
        div.stButton > button:hover {
            background: rgba(0, 200, 255, 0.12);
            box-shadow: 0 0 20px rgba(0, 200, 255, 0.65),
                        inset 0 0 25px rgba(0, 200, 255, 0.15);
            transform: translateY(-2px);
        }
        div.stButton > button:active {
            transform: scale(0.95);
            box-shadow: 0 0 10px rgba(0, 200, 255, 0.45);
        }
        </style>
    """, unsafe_allow_html=True)

    # 🎨 Template Selection — visual card grid
    st.markdown("""
    <style>
    div[data-testid="stHorizontalBlock"] .tpl-card-wrap { padding: 4px; }
    </style>
    """, unsafe_allow_html=True)

    TEMPLATE_META = [
        ("Default (Professional)",        "#2f4f6f", "#e8f0fe"),
        ("Modern Minimal",                 "#0d9488", "#f0fdfa"),
        ("Elegant Sidebar",                "#7c3aed", "#f5f3ff"),
        ("Classic Clean (Single Column)",  "#374151", "#f9fafb"),
        ("Executive (Single Column)",      "#1e3a5f", "#eff6ff"),
        ("Timeline (Single Column)",       "#b45309", "#fffbeb"),
        ("Corporate Blue (Two Column)",    "#1d4ed8", "#eff6ff"),
        ("Creative Green (Two Column)",    "#166534", "#f0fdf4"),
        ("Warm Terracotta (Two Column)",   "#c2410c", "#fff7ed"),
        ("Navy Prestige (Two Column)",     "#1e3a5f", "#f0f4ff"),
        ("Slate Gray (Single Column)",     "#475569", "#f8fafc"),
        ("Teal Impact (Two Column)",       "#0f766e", "#f0fdfa"),
        ("Burgundy Classic (Single Column)","#881337","#fff1f2"),
        ("Indigo Tech (Two Column)",       "#4338ca", "#eef2ff"),
        ("Forest Green (Single Column)",   "#14532d", "#f0fdf4"),
        # ── 6 new premium templates ──────────────────────────────────────────
        ("Pure White (Single Column)",     "linear-gradient(to right,#ffffff 70%,#111111 100%)", "#ffffff"),
        ("Midnight Black (Single Column)", "#f59e0b", "#111827"),
        ("Soft Lavender (Single Column)",  "#6366f1", "#f5f3ff"),
        ("Warm Sand (Single Column)",      "#b45309", "#fdf8f0"),
        ("Ice Blue (Single Column)",       "#0369a1", "#eff9ff"),
    ]
    TEMPLATE_NAMES = [t[0] for t in TEMPLATE_META]

    if "selected_template_name" not in st.session_state:
        st.session_state["selected_template_name"] = TEMPLATE_NAMES[0]

    st.markdown("<div style='margin:18px 0 8px;font-size:14px;font-weight:600;color:#93c5fd;'>🎨 Choose Resume Template</div>", unsafe_allow_html=True)

    # Show 5 cards per row
    _tpl_rows = [TEMPLATE_META[i:i+5] for i in range(0, len(TEMPLATE_META), 5)]
    for _row in _tpl_rows:
        _cols = st.columns(len(_row))
        for _ci, (_tname, _color, _light) in enumerate(_row):
            with _cols[_ci]:
                _is_sel = st.session_state["selected_template_name"] == _tname
                _border = "2px solid #4da6ff" if _is_sel else "1px solid rgba(0,180,255,0.15)"
                _glow   = "box-shadow: 0 0 12px rgba(77,166,255,0.45);" if _is_sel else ""
                _is_light_swatch = _color.startswith("linear-gradient") or _color in ("#fdf8f0", "#eff9ff", "#fff0f6", "#f5f3ff", "#f8fafc", "#fffbeb")
                _swatch_extra = "border:1px solid rgba(255,255,255,0.25);" if _is_light_swatch else ""
                st.markdown(
                    f"<div style='background:rgba(13,20,40,0.6);border:{_border};border-radius:10px;"
                    f"padding:8px 6px 6px;text-align:center;{_glow}'>"
                    f"<div style='height:28px;border-radius:6px;background:{_color};margin-bottom:6px;{_swatch_extra}'></div>"
                    f"<div style='font-size:9.5px;color:{'#93c5fd' if _is_sel else '#6b7280'};font-weight:{'700' if _is_sel else '500'};line-height:1.3;'>{_tname}</div>"
                    f"</div>",
                    unsafe_allow_html=True,
                )
                if st.button("✓" if _is_sel else "Select", key=f"tpl_btn_{_tname}", use_container_width=True):
                    if st.session_state["selected_template_name"] != _tname:
                        st.session_state["selected_template_name"] = _tname
                        st.rerun()

    selected_template = st.session_state["selected_template_name"]

    # 📸 Upload profile photo
    st.markdown("<div style='margin:18px 0 6px;font-size:14px;font-weight:600;color:#93c5fd;'>📸 Profile Photo</div>", unsafe_allow_html=True)
    _img_col1, _img_col2 = st.columns([3, 1])
    with _img_col1:
        uploaded_image = st.file_uploader("Upload a Profile Image (PNG/JPG, square preferred)", type=["png", "jpg", "jpeg"], key="profile_img_upload", label_visibility="collapsed")
        # ── FIX: Encode and store as soon as a new file is uploaded, then rerun
        # so the preview column (already rendered above) refreshes immediately.
        if uploaded_image is not None:
            import base64 as _base64
            _new_encoded = _base64.b64encode(uploaded_image.read()).decode()
            if _new_encoded != st.session_state.get("encoded_profile_image"):
                st.session_state["encoded_profile_image"] = _new_encoded
                st.rerun()
        # ── FIX: "Remove Photo" button clears session state so image disappears.
        if st.session_state.get("encoded_profile_image"):
            if st.button("🗑️ Remove Photo", key="remove_profile_photo"):
                st.session_state.pop("encoded_profile_image", None)
                st.rerun()
    with _img_col2:
        if st.session_state.get("encoded_profile_image"):
            st.markdown(
                f"<img src='data:image/png;base64,{st.session_state['encoded_profile_image']}' "
                f"class='photo-preview' />",
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                "<div style='width:72px;height:72px;border-radius:50%;background:#1e2535;"
                "border:2px dashed #374151;display:flex;align-items:center;justify-content:center;"
                "font-size:22px;margin:4px auto;'>👤</div>",
                unsafe_allow_html=True,
            )
    profile_img_html = ""

    if st.session_state.get("encoded_profile_image"):
        encoded_image = st.session_state["encoded_profile_image"]
        profile_img_html = f"""
        <div style="display: flex; justify-content: flex-end; margin-top: 20px;">
            <img src="data:image/png;base64,{encoded_image}" alt="Profile Photo"
                 style="
                    width: 140px;
                    height: 140px;
                    border-radius: 50%;
                    object-fit: cover;
                    object-position: center;
                    border: 4px solid rgba(255,255,255,0.6);
                    box-shadow:
                        0 0 0 3px #4da6ff,
                        0 8px 25px rgba(77, 166, 255, 0.3),
                        0 4px 15px rgba(0, 0, 0, 0.15);
                    transition: transform 0.3s ease-in-out;
                "
                onmouseover="this.style.transform='scale(1.07)'"
                onmouseout="this.style.transform='scale(1)'"
             />
        </div>
        """
    else:
        st.markdown("<div style='font-size:12px;color:#4b5563;margin-top:4px;'>📸 Upload a clear, front-facing photo (square or portrait preferred)</div>", unsafe_allow_html=True)

    # ---------------- Session State Defaults ----------------
    fields = ["name", "email", "phone", "linkedin", "location", "portfolio", "summary",
              "skills", "languages", "interests", "Softskills", "job_title"]
    for f in fields:
        st.session_state.setdefault(f, "")

    st.session_state.setdefault("experience_entries", [{"title": "", "company": "", "duration": "", "description": ""}])
    st.session_state.setdefault("education_entries", [{"degree": "", "institution": "", "year": "", "details": ""}])
    st.session_state.setdefault("project_entries", [{"title": "", "tech": "", "duration": "", "description": ""}])
    st.session_state.setdefault("project_links", [])
    st.session_state.setdefault("certificate_links", [{"name": "", "link": "", "duration": "", "description": ""}])
    st.session_state.setdefault("form_key_counter", 0)

    # ─────────────────────────────────────────────────────────────────────────
    # GAMIFIED SIDEBAR
    # ─────────────────────────────────────────────────────────────────────────
    def render_gamified_sidebar(ss, fk):
        """
        Renders a fully gamified sidebar with:
        - XP counter + rank badge
        - Master progress bar
        - Streak dot row
        - Per-section rows with inline SVG icons, mini-bars, and check circles
        - Stats footer (Done / XP / % Complete)
        - Section add/delete controls (Experience, Education, Projects, Certificates)
        All icons are inline SVG — no emojis.
        """

        # ══════════════════════════════════════════════════════════════════════
        # PRE-SCORING SYNC: push live widget values into entry dicts BEFORE
        # scoring runs. This eliminates the double-press/stale-score bug that
        # occurs because the sidebar renders BEFORE the form widgets, so the
        # entry dicts still hold the previous run's committed values.
        # By reading directly from session_state widget keys here, we always
        # score the text the user has typed RIGHT NOW.
        # ══════════════════════════════════════════════════════════════════════
        def _sync_entries():
            # Sync simple scalar fields
            for widget_key, ss_key in [
                (f"name_input_{fk}",    "name"),
                (f"email_input_{fk}",   "email"),
                (f"phone_input_{fk}",   "phone"),
                (f"loc_input_{fk}",     "location"),
                (f"job_input_{fk}",     "job_title"),
                (f"summary_input_{fk}", "summary"),
                (f"skills_input_{fk}",  "skills"),
                (f"lang_input_{fk}",    "languages"),
                (f"int_input_{fk}",     "interests"),
                (f"soft_input_{fk}",    "Softskills"),
                (f"ln_input_{fk}",      "linkedin"),
                (f"phone_input_{fk}",   "phone"),
            ]:
                if widget_key in ss:
                    ss[ss_key] = ss[widget_key]

            # Sync experience entries
            entries = ss.get("experience_entries", [])
            n = len(entries)
            for i, e in enumerate(entries):
                for widget_suffix, dict_key in [
                    (f"title_{i}_{n}_{fk}",       "title"),
                    (f"company_{i}_{n}_{fk}",     "company"),
                    (f"duration_{i}_{n}_{fk}",    "duration"),
                    (f"description_{i}_{n}_{fk}", "description"),
                ]:
                    if widget_suffix in ss:
                        e[dict_key] = ss[widget_suffix]

            # Sync education entries
            entries = ss.get("education_entries", [])
            n = len(entries)
            for i, e in enumerate(entries):
                for widget_suffix, dict_key in [
                    (f"degree_{i}_{n}_{fk}",      "degree"),
                    (f"institution_{i}_{n}_{fk}", "institution"),
                    (f"edu_year_{i}_{n}_{fk}",    "year"),
                    (f"edu_details_{i}_{n}_{fk}", "details"),
                ]:
                    if widget_suffix in ss:
                        e[dict_key] = ss[widget_suffix]

            # Sync project entries
            entries = ss.get("project_entries", [])
            n = len(entries)
            for i, e in enumerate(entries):
                for widget_suffix, dict_key in [
                    (f"proj_title_{i}_{n}_{fk}",    "title"),
                    (f"proj_tech_{i}_{n}_{fk}",     "tech"),
                    (f"proj_duration_{i}_{n}_{fk}", "duration"),
                    (f"proj_desc_{i}_{n}_{fk}",     "description"),
                ]:
                    if widget_suffix in ss:
                        e[dict_key] = ss[widget_suffix]

            # Sync certificate entries
            entries = ss.get("certificate_links", [])
            n = len(entries)
            for i, e in enumerate(entries):
                for widget_suffix, dict_key in [
                    (f"cert_name_{i}_{n}_{fk}",        "name"),
                    (f"cert_link_{i}_{n}_{fk}",        "link"),
                    (f"cert_duration_{i}_{n}_{fk}",    "duration"),
                    (f"cert_description_{i}_{n}_{fk}", "description"),
                ]:
                    if widget_suffix in ss:
                        e[dict_key] = ss[widget_suffix]

        _sync_entries()

        # ── SVG icon library ──────────────────────────────────────────────────
        SVG = {
            "personal": '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><circle cx="8" cy="5" r="2.5"/><path d="M2.5 13.5c0-3 2.5-4.5 5.5-4.5s5.5 1.5 5.5 4.5"/></svg>',
            "summary":  '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><rect x="2" y="2" width="12" height="12" rx="2"/><line x1="5" y1="6" x2="11" y2="6"/><line x1="5" y1="9" x2="9" y2="9"/></svg>',
            "exp":      '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><rect x="2" y="6" width="12" height="8" rx="1.5"/><path d="M5 6V4.5A2.5 2.5 0 0 1 11 4.5V6"/><line x1="8" y1="9" x2="8" y2="11"/></svg>',
            "edu":      '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><path d="M8 2L14 5.5 8 9 2 5.5Z"/><path d="M4.5 7.5V11.5c0 0 1.5 1.5 3.5 1.5s3.5-1.5 3.5-1.5V7.5"/><line x1="14" y1="5.5" x2="14" y2="9"/></svg>',
            "projects": '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><polyline points="2,12 6,7 9,10 12,5 14,7"/><circle cx="14" cy="4" r="1.2" fill="currentColor" stroke="none"/></svg>',
            "skills":   '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><path d="M10.5 2.5l1 1-7 7-1-1z"/><path d="M12 4l1.5-1.5-1-1L11 3"/><path d="M3 11l-0.5 2 2-0.5"/></svg>',
            "certs":    '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><circle cx="8" cy="7" r="3.5"/><path d="M5.5 10L4 14l4-2 4 2-1.5-4"/></svg>',
            "contact":  '<svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"><path d="M2 3h12v9a1 1 0 01-1 1H3a1 1 0 01-1-1V3z"/><polyline points="2,3 8,8.5 14,3"/></svg>',
            "add":      '<svg width="12" height="12" viewBox="0 0 12 12" fill="none" stroke="currentColor" stroke-width="1.6" stroke-linecap="round"><line x1="6" y1="2" x2="6" y2="10"/><line x1="2" y1="6" x2="10" y2="6"/></svg>',
            "remove":   '<svg width="12" height="12" viewBox="0 0 12 12" fill="none" stroke="currentColor" stroke-width="1.6" stroke-linecap="round"><line x1="2" y1="6" x2="10" y2="6"/></svg>',
        }

        # ══════════════════════════════════════════════════════════════════════
        # PARTIAL-FILL SCORING
        # Each section returns a float 0.0–1.0 based on how many of its
        # required fields the user has actually filled.
        # The mini-bar width, XP, and master bar all reflect partial progress.
        # A section is "complete" (check circle on) only at 1.0.
        # ══════════════════════════════════════════════════════════════════════

        def _wv(widget_key, fallback_key=""):
            """Live widget value → stored session value → empty string."""
            v = ss.get(widget_key, "")
            if not v and fallback_key:
                v = ss.get(fallback_key, "")
            return str(v).strip()

        def _filled(*values):
            """Count how many of the given values are non-empty."""
            return sum(1 for v in values if str(v).strip())

        # ── Personal Info: 5 fields, each worth 0.2 ───────────────────────────
        pi_name     = _wv(f"name_input_{fk}",  "name")
        pi_email    = _wv(f"email_input_{fk}",  "email")
        pi_phone    = _wv(f"phone_input_{fk}",  "phone")
        pi_location = _wv(f"loc_input_{fk}",    "location")
        pi_jobtitle = _wv(f"job_input_{fk}",    "job_title")
        _fill_personal = round(_filled(pi_name, pi_email, pi_phone, pi_location, pi_jobtitle) / 5, 2)

        # ── Summary: quality-gated length scoring ─────────────────────────────
        # Garbage text earns 0 regardless of length. Valid text is scored by
        # character length tiers, rewarding rich professional summaries.
        _summary_text = _wv(f"summary_input_{fk}", "summary")
        _summary_len  = len(_summary_text)
        if _summary_len == 0 or detect_garbage_text(_summary_text):
            _fill_summary = 0.0
        elif _summary_len < 40:
            _fill_summary = 0.25   # started but too short
        elif _summary_len < 100:
            _fill_summary = 0.60   # decent but brief
        elif _summary_len < 200:
            _fill_summary = 0.85   # good
        else:
            _fill_summary = 1.0    # full — rich summary

        # ══════════════════════════════════════════════════════════════════════
        # TEXT QUALITY VALIDATION ENGINE — delegates to module-level functions
        # All logic lives at module scope for reuse; these are local aliases.
        # ══════════════════════════════════════════════════════════════════════
        _is_low_quality_text  = detect_garbage_text      # module-level
        _is_gibberish         = detect_garbage_text      # backwards-compat alias

        # ══════════════════════════════════════════════════════════════════════
        # LIVE VALUE READER
        # ══════════════════════════════════════════════════════════════════════

        def _get_val(ss, widget_key, entry, stored_key, fk):
            """
            Read the most up-to-date value for a field.
            Priority:
              1. Live Streamlit widget value (ss[widget_key])
              2. Stored entry dict value (entry[stored_key])
            """
            live = ss.get(widget_key, "")
            if live:
                return str(live).strip()
            stored = entry.get(stored_key, "")
            return str(stored).strip()

        # XP_WEIGHTS and XP_TOTAL_MAX are defined at module level — use them directly.

        # ── Skills & More — delegates to module-level score_skills_section ──────
        _skills_raw    = _wv(f"skills_input_{fk}", "skills")
        _interests_raw = _wv(f"int_input_{fk}",    "interests")
        _soft_raw      = _wv(f"soft_input_{fk}",   "Softskills")
        _lang_raw      = _wv(f"lang_input_{fk}",   "languages")
        # Keep individual counts for feedback tips
        _skill_count   = _count_valid_tokens(_skills_raw)
        _soft_count    = _count_valid_tokens(_soft_raw)
        _fill_skills   = score_skills_section(
            skills=_skills_raw,
            soft_skills=_soft_raw,
            interests=_interests_raw,
            languages=_lang_raw,
        )

        # ── Contact: phone + linkedin, each worth 0.5 ─────────────────────────
        pi_phone2   = _wv(f"phone_input_{fk}", "phone")
        pi_linkedin = _wv(f"ln_input_{fk}",    "linkedin")
        _fill_contact = round(_filled(pi_phone2, pi_linkedin) / 2, 3)

        # ── Experience scoring — delegates to module-level score_experience_section ──
        def _score_experience():
            entries = ss.get("experience_entries", [])
            if not entries:
                return 0.0
            n = len(entries)
            # Merge live widget values into a temporary list for scoring
            merged = []
            for i, e in enumerate(entries):
                merged.append({
                    "title":       _get_val(ss, f"title_{i}_{n}_{fk}",       e, "title",       fk),
                    "company":     _get_val(ss, f"company_{i}_{n}_{fk}",     e, "company",     fk),
                    "duration":    _get_val(ss, f"duration_{i}_{n}_{fk}",    e, "duration",    fk),
                    "description": _get_val(ss, f"description_{i}_{n}_{fk}", e, "description", fk),
                })
            return score_experience_section(merged)  # module-level

        _fill_exp = _score_experience()

        # ── Education scoring — delegates to module-level score_education_section ──
        def _score_education():
            entries = ss.get("education_entries", [])
            if not entries:
                return 0.0
            n = len(entries)
            merged = []
            for i, e in enumerate(entries):
                merged.append({
                    "institution": _get_val(ss, f"institution_{i}_{n}_{fk}", e, "institution", fk),
                    "degree":      _get_val(ss, f"degree_{i}_{n}_{fk}",      e, "degree",      fk),
                    "year":        _get_val(ss, f"edu_year_{i}_{n}_{fk}",    e, "year",        fk),
                    "details":     _get_val(ss, f"edu_details_{i}_{n}_{fk}", e, "details",     fk),
                })
            return score_education_section(merged)  # module-level

        _fill_edu = _score_education()

        # ── Projects scoring — delegates to module-level score_project_section ──
        def _score_projects():
            entries = ss.get("project_entries", [])
            if not entries:
                return 0.0
            n = len(entries)
            merged = []
            for i, e in enumerate(entries):
                merged.append({
                    "title":       _get_val(ss, f"proj_title_{i}_{n}_{fk}",    e, "title",       fk),
                    "tech":        _get_val(ss, f"proj_tech_{i}_{n}_{fk}",     e, "tech",        fk),
                    "duration":    _get_val(ss, f"proj_duration_{i}_{n}_{fk}", e, "duration",    fk),
                    "description": _get_val(ss, f"proj_desc_{i}_{n}_{fk}",     e, "description", fk),
                })
            return score_project_section(merged)  # module-level

        _fill_proj = _score_projects()

        # ── Certificates scoring — delegates to module-level score_certificate_section ──
        def _score_certificates():
            entries = ss.get("certificate_links", [])
            if not entries:
                return 0.0
            n = len(entries)
            merged = []
            for i, e in enumerate(entries):
                merged.append({
                    "name":        _get_val(ss, f"cert_name_{i}_{n}_{fk}",        e, "name",        fk),
                    "link":        _get_val(ss, f"cert_link_{i}_{n}_{fk}",        e, "link",        fk),
                    "duration":    _get_val(ss, f"cert_duration_{i}_{n}_{fk}",    e, "duration",    fk),
                    "description": _get_val(ss, f"cert_description_{i}_{n}_{fk}", e, "description", fk),
                })
            return score_certificate_section(merged)  # module-level

        _fill_cert = _score_certificates()

        # ── Aggregate XP via module-level calculate_resume_xp() — idempotent ──
        # Order matches XP_WEIGHTS for consistent rendering.
        SECTIONS = {
            "Experience":     _fill_exp,
            "Projects":       _fill_proj,
            "Skills & More":  _fill_skills,
            "Education":      _fill_edu,
            "Summary":        _fill_summary,
            "Certificates":   _fill_cert,
            "Personal Info":  _fill_personal,
            "Contact":        _fill_contact,
        }
        ICON_KEYS = ["exp", "projects", "skills", "edu", "summary", "certs", "personal", "contact"]

        # ── Done thresholds — matched to new scoring formulas ─────────────────
        DONE_THRESHOLD = {
            "Personal Info":  1.0,   # all 5 fields filled
            "Summary":        0.85,  # rich summary with good word quality
            "Experience":     0.72,  # role + company + quality description
            "Education":      0.68,  # institution + degree + year + details
            "Projects":       0.65,  # title + techs + quality description
            "Skills & More":  0.57,  # 3+ tech skills + soft skills + interests
            "Certificates":   0.62,  # name + link + duration + description
            "Contact":        1.0,   # both phone AND linkedin
        }

        # ── Aggregate XP via calculate_resume_xp() — idempotent, no accumulation ──
        _raw_xp, xp, pct = calculate_resume_xp(SECTIONS, XP_WEIGHTS)
        max_xp    = XP_TOTAL_MAX   # always 100
        total     = len(SECTIONS)
        fully_done = sum(
            1 for (k, v) in SECTIONS.items()
            if v >= DONE_THRESHOLD.get(k, 1.0)
        )

        # Store computed XP in session state for external access
        ss["resume_xp"]  = xp
        ss["resume_pct"] = pct

        if   pct == 0:    rank, rank_color, rank_bg, rank_border = "Unranked",   "#6b7280", "#1e2535", "#374151"
        elif pct <= 20:   rank, rank_color, rank_bg, rank_border = "Beginner",   "#d97706", "#2a1f12", "#92400e"
        elif pct <= 40:   rank, rank_color, rank_bg, rank_border = "Builder",    "#94a3b8", "#1a2133", "#475569"
        elif pct <= 60:   rank, rank_color, rank_bg, rank_border = "Proficient", "#22d3ee", "#0c2233", "#0e4f60"
        elif pct <= 80:   rank, rank_color, rank_bg, rank_border = "Advanced",   "#f59e0b", "#2a2410", "#92700e"
        elif pct < 100:   rank, rank_color, rank_bg, rank_border = "Expert",     "#a78bfa", "#1a1a2e", "#6d28d9"
        else:             rank, rank_color, rank_bg, rank_border = "Pro Resume", "#34d399", "#0a2318", "#065f46"

        # ── helper: section row HTML — now takes fill float 0.0–1.0 ──────────
        def _section_row(label, icon_key, fill):
            done      = fill >= DONE_THRESHOLD.get(label, 1.0)
            partial   = 0.0 < fill < 1.0
            bar_pct   = f"{int(fill * 100)}%"
            # Per-section XP earned vs max for this section
            _sec_max  = XP_WEIGHTS.get(label, 0)
            _sec_earned = int(round(fill * _sec_max))
            # colour ramp: empty=dark, partial=amber, done=blue
            if done:
                icon_bg, icon_col, name_col = "#1d3a6e", "#93c5fd", "#93c5fd"
                bar_col = "#3b82f6"
                row_bg  = "#131c33"
                row_bdr = "#1d4ed8"
            elif partial:
                icon_bg, icon_col, name_col = "#2a1f12", "#f59e0b", "#d4a017"
                bar_col = "#f59e0b"
                row_bg  = "#1a1600"
                row_bdr = "#78450a"
            else:
                icon_bg, icon_col, name_col = "#1e2535", "#6b7280", "#6b7280"
                bar_col = "#374151"
                row_bg  = "#161b27"
                row_bdr = "#1e2535"
            chk_bg  = "#2563eb"    if done    else "transparent"
            chk_bdr = "#2563eb"    if done    else ("#78450a" if partial else "#374151")
            chk_op  = "1"          if done    else "0"
            # XP badge: shows "earned/max XP" — always visible, colour-coded
            if done:
                xp_badge_col = "#93c5fd"
            elif partial:
                xp_badge_col = "#f59e0b"
            else:
                xp_badge_col = "#4b5563"
            xp_badge = (
                "<div style='font-size:9px;color:" + xp_badge_col + ";font-weight:600;"
                "white-space:nowrap;'>"
                + str(_sec_earned) + "<span style='opacity:0.55;font-weight:400;'>/"
                + str(_sec_max) + "</span></div>"
            )
            return (
                "<div style='display:flex;align-items:center;gap:10px;padding:9px 10px;"
                "border-radius:9px;background:" + row_bg + ";border:0.5px solid " + row_bdr + ";"
                "margin-bottom:7px;'>"
                "<div style='width:28px;height:28px;border-radius:7px;background:" + icon_bg + ";"
                "display:flex;align-items:center;justify-content:center;flex-shrink:0;color:" + icon_col + ";'>"
                + SVG[icon_key] +
                "</div>"
                "<div style='flex:1;min-width:0;'>"
                "<div style='display:flex;justify-content:space-between;align-items:center;'>"
                "<div style='font-size:12px;font-weight:500;color:" + name_col + ";'>" + label + "</div>"
                + xp_badge +
                "</div>"
                "<div style='height:3px;background:#1e2535;border-radius:3px;margin-top:4px;overflow:hidden;'>"
                "<div style='height:100%;width:" + bar_pct + ";background:" + bar_col + ";border-radius:3px;'></div>"
                "</div>"
                "</div>"
                "<div style='width:18px;height:18px;border-radius:50%;background:" + chk_bg + ";"
                "border:1.5px solid " + chk_bdr + ";display:flex;align-items:center;"
                "justify-content:center;flex-shrink:0;'>"
                "<svg width='9' height='9' viewBox='0 0 10 10' fill='none' stroke='#fff' "
                "stroke-width='1.8' stroke-linecap='round' stroke-linejoin='round' "
                "style='opacity:" + chk_op + ";'>"
                "<polyline points='2,5 4.5,7.5 8.5,2.5'/>"
                "</svg>"
                "</div>"
                "</div>"
            )

        # ── streak dots — one dot per section, lit when section is 100% ─────────
        dots_html = "".join(
            f"<div style='flex:1;height:4px;border-radius:3px;"
            f"background:{'#3b82f6' if i < fully_done else '#1e2535'};'></div>"
            for i in range(total)
        )

        # ── render into sidebar ────────────────────────────────────────────────
        if st.session_state.get("username") != "admin":
            with st.sidebar:

                # ── XP header ─────────────────────────────────────────────────────
                st.markdown(f"""
<div style='margin-bottom:6px;display:flex;justify-content:space-between;align-items:center;'>
  <span style='font-size:10px;letter-spacing:1.2px;text-transform:uppercase;
               color:#6b7280;font-weight:500;'>Resume XP</span>
  <span style='font-size:11px;color:#9ca3af;font-weight:500;'>{xp} / {max_xp} XP</span>
</div>
<div style='margin-bottom:8px;'>
  <span style='font-size:11px;font-weight:500;padding:2px 10px;border-radius:20px;
               background:{rank_bg};border:0.5px solid {rank_border};color:{rank_color};'>{rank}</span>
</div>
<div style='width:100%;height:7px;background:#1e2535;border-radius:6px;overflow:hidden;margin-bottom:10px;'>
  <div style='height:100%;width:{pct}%;background:#3b82f6;border-radius:6px;'></div>
</div>
<div style='display:flex;gap:4px;margin-bottom:18px;'>{dots_html}</div>
<div style='font-size:10px;letter-spacing:1.4px;text-transform:uppercase;
            color:#4b5563;margin-bottom:10px;font-weight:500;'>Sections</div>
""", unsafe_allow_html=True)

                # ── section rows ──────────────────────────────────────────────────
                for (label, fill), icon_key in zip(SECTIONS.items(), ICON_KEYS):
                    st.markdown(_section_row(label, icon_key, fill), unsafe_allow_html=True)

                # ── UX Feedback Panel ─────────────────────────────────────────────
                # Generate actionable, ranked feedback tips based on current state.
                # Tips are prioritised by XP weight (highest-impact sections first).
                _feedback_tips = []

                # Experience feedback
                if _fill_exp == 0.0:
                    _feedback_tips.append("Add work experience to strengthen your resume (+30 XP potential)")
                elif _fill_exp < DONE_THRESHOLD["Experience"]:
                    _exp_entries = ss.get("experience_entries", [])
                    _has_desc = any(
                        len(str(e.get("description", "")).strip()) > 40
                        for e in _exp_entries
                    )
                    if not _has_desc:
                        _feedback_tips.append("Add detailed descriptions to your work experience to unlock more XP")
                    else:
                        _feedback_tips.append("Enrich your experience descriptions with specific achievements and metrics")

                # Projects feedback
                if _fill_proj == 0.0:
                    _feedback_tips.append("Add a project to gain up to 22 XP — projects are highly valued")
                elif _fill_proj < DONE_THRESHOLD["Projects"]:
                    _proj_entries = ss.get("project_entries", [])
                    _has_tech = any(
                        _count_valid_tokens(e.get("tech", "")) >= 1
                        for e in _proj_entries if e.get("title")
                    )
                    if not _has_tech:
                        _feedback_tips.append("List the technologies used in your projects to improve your score")
                    else:
                        _feedback_tips.append("Improve your project descriptions — explain the problem solved and your impact")
                elif _fill_proj < 0.90:
                    n_projs = len([e for e in ss.get("project_entries", []) if e.get("title")])
                    if n_projs < 2:
                        _feedback_tips.append("Add another project to increase your Projects score further")

                # Skills feedback
                if _fill_skills < DONE_THRESHOLD["Skills & More"]:
                    if _skill_count < 3:
                        _feedback_tips.append(f"Add more technical skills — you have {_skill_count}, aim for 5+ (comma-separated)")
                    if _soft_count < 2:
                        _feedback_tips.append("Add 2–3 soft skills (e.g. Leadership, Communication) to boost your score")

                # Education feedback
                if _fill_edu == 0.0:
                    _feedback_tips.append("Add your education details to build a complete resume")
                elif _fill_edu < DONE_THRESHOLD["Education"]:
                    _feedback_tips.append("Add graduation year and academic details/achievements to complete education")

                # Summary feedback
                if _fill_summary == 0.0:
                    _feedback_tips.append("Write a professional summary — it's your first impression on recruiters")
                elif _fill_summary < DONE_THRESHOLD["Summary"]:
                    _feedback_tips.append("Expand your summary with more specific skills, experience, and career goals")

                # Certificate feedback
                if _fill_cert == 0.0 and pct >= 40:
                    _feedback_tips.append("Add a certification to differentiate yourself from other candidates")
                elif 0.0 < _fill_cert < DONE_THRESHOLD["Certificates"]:
                    _feedback_tips.append("Add a verification link and description to your certificates for full credit")

                # Contact feedback
                if _fill_contact < 1.0:
                    if not pi_phone2:
                        _feedback_tips.append("Add your phone number to make your resume complete")
                    if not pi_linkedin:
                        _feedback_tips.append("Add your LinkedIn profile URL — recruiters always check it")

                # Show up to 3 tips (highest priority = highest XP weight = listed first)
                if _feedback_tips:
                    _tips_to_show = _feedback_tips[:3]
                    _tip_html_items = "".join(
                        f"<div style='display:flex;align-items:flex-start;gap:8px;margin-bottom:8px;'>"
                        f"<span style='color:#f59e0b;font-size:12px;flex-shrink:0;margin-top:1px;'>&#9654;</span>"
                        f"<span style='font-size:11px;color:#cbd5e1;line-height:1.5;'>{tip}</span>"
                        f"</div>"
                        for tip in _tips_to_show
                    )
                    st.markdown(
                        f"<div style='margin:12px 0 6px;padding:10px 12px;background:#111827;"
                        f"border-radius:8px;border:0.5px solid #374151;'>"
                        f"<div style='font-size:9px;letter-spacing:1.2px;text-transform:uppercase;"
                        f"color:#6b7280;font-weight:600;margin-bottom:8px;'>Tips to Boost XP</div>"
                        f"{_tip_html_items}"
                        f"</div>",
                        unsafe_allow_html=True
                    )

                # ── divider + stats footer ─────────────────────────────────────────
                st.markdown(f"""
<hr style='border:none;border-top:0.5px solid #1e2535;margin:14px 0;'>
<div style='display:flex;justify-content:space-between;text-align:center;margin-bottom:18px;'>
  <div>
    <span style='font-size:16px;font-weight:500;color:#e2e8f0;display:block;'>{fully_done}</span>
    <span style='font-size:10px;color:#4b5563;letter-spacing:0.8px;text-transform:uppercase;'>Done</span>
  </div>
  <div>
    <span style='font-size:16px;font-weight:500;color:#e2e8f0;display:block;'>{xp}<span style='font-size:10px;color:#4b5563;'>/{max_xp}</span></span>
    <span style='font-size:10px;color:#4b5563;letter-spacing:0.8px;text-transform:uppercase;'>XP</span>
  </div>
  <div>
    <span style='font-size:16px;font-weight:500;color:#e2e8f0;display:block;'>{pct}%</span>
    <span style='font-size:10px;color:#4b5563;letter-spacing:0.8px;text-transform:uppercase;'>Complete</span>
  </div>
</div>
<hr style='border:none;border-top:0.5px solid #1e2535;margin:0 0 14px;'>
<div style='font-size:10px;letter-spacing:1.4px;text-transform:uppercase;
                color:#4b5563;margin-bottom:10px;font-weight:500;'>Manage Sections</div>
""", unsafe_allow_html=True)

                # ── section add/delete controls ────────────────────────────────────
                if "edit_mode" not in ss:
                    ss.edit_mode = "Add"

                mode = st.selectbox(
                    "Mode",
                    ["Add", "Delete"],
                    index=0,
                    key="mode_dropdown",
                    label_visibility="collapsed",
                )
                ss.edit_mode = mode

                st.markdown("<div style='height:8px;'></div>", unsafe_allow_html=True)

                # Experience
                with st.expander("Experience", expanded=False):
                    if st.button(("Add" if mode == "Add" else "Delete") + " Experience", key="exp_btn", use_container_width=True):
                        if mode == "Add":
                            ss.experience_entries.append({"title": "", "company": "", "duration": "", "description": ""})
                        elif mode == "Delete" and len(ss.experience_entries) > 1:
                            ss.experience_entries.pop()

                # Education
                with st.expander("Education", expanded=False):
                    if st.button(("Add" if mode == "Add" else "Delete") + " Education", key="edu_btn", use_container_width=True):
                        if mode == "Add":
                            ss.education_entries.append({"degree": "", "institution": "", "year": "", "details": ""})
                        elif mode == "Delete" and len(ss.education_entries) > 1:
                            ss.education_entries.pop()

                # Projects
                with st.expander("Projects", expanded=False):
                    if st.button(("Add" if mode == "Add" else "Delete") + " Project", key="proj_btn", use_container_width=True):
                        if mode == "Add":
                            ss.project_entries.append({"title": "", "tech": "", "duration": "", "description": ""})
                        elif mode == "Delete" and len(ss.project_entries) > 1:
                            ss.project_entries.pop()

                # Certificates
                with st.expander("Certificates", expanded=False):
                    if st.button(("Add" if mode == "Add" else "Delete") + " Certificate", key="cert_btn", use_container_width=True):
                        if mode == "Add":
                            ss.certificate_links.append({"name": "", "link": "", "duration": "", "description": ""})
                        elif mode == "Delete" and len(ss.certificate_links) > 1:
                            ss.certificate_links.pop()

    # ── call gamified sidebar AFTER fk is known so widget keys resolve correctly ──
    fk = st.session_state["form_key_counter"]
    render_gamified_sidebar(st.session_state, fk)
    mode = st.session_state.get("edit_mode", "Add")

    # ── Shared section-header style injected once ────────────────────────────
    st.markdown("""
    <style>
    .sec-header {
        display: flex; align-items: center; gap: 10px;
        background: linear-gradient(90deg, rgba(0,180,255,0.10) 0%, rgba(0,180,255,0.03) 100%);
        border-left: 3px solid #4da6ff;
        border-radius: 0 10px 10px 0;
        padding: 9px 14px; margin: 18px 0 10px;
    }
    .sec-header .sec-icon { font-size: 18px; line-height: 1; }
    .sec-header .sec-title {
        font-size: 15px; font-weight: 700;
        color: #93c5fd; letter-spacing: 0.4px; margin: 0;
    }
    .sec-header .sec-badge {
        margin-left: auto; font-size: 10px; font-weight: 600;
        color: #4b5563; background: #1e2535;
        padding: 2px 8px; border-radius: 20px;
        border: 0.5px solid #374151;
    }
    .entry-card {
        background: rgba(13,20,40,0.55);
        border: 0.5px solid rgba(0,180,255,0.18);
        border-radius: 12px; padding: 14px 16px; margin-bottom: 12px;
    }
    .entry-card-label {
        font-size: 12px; font-weight: 600; color: #4da6ff;
        margin-bottom: 10px; letter-spacing: 0.3px;
    }
    .field-hint { font-size: 11px; color: #4b5563; margin-top: -8px; margin-bottom: 6px; }
    .tag-chip {
        display: inline-block; background: rgba(0,180,255,0.12);
        border: 0.5px solid rgba(0,180,255,0.35); color: #93c5fd;
        font-size: 12px; font-weight: 500;
        padding: 3px 10px; border-radius: 20px; margin: 3px 3px 3px 0;
    }
    .confirm-warn {
        background: rgba(239,68,68,0.10); border: 0.5px solid rgba(239,68,68,0.35);
        border-radius: 10px; padding: 10px 14px; margin-top: 6px;
        font-size: 13px; color: #fca5a5;
    }
    .photo-preview {
        width: 90px; height: 90px; border-radius: 50%; object-fit: cover;
        border: 3px solid #4da6ff;
        box-shadow: 0 0 14px rgba(77,166,255,0.4);
        display: block; margin: 8px auto 0;
    }
    @media (max-width: 768px) {
        [data-testid="column"] { min-width: 100% !important; }
    }
    </style>
    """, unsafe_allow_html=True)

    def _sec_hdr(icon, title, badge=None):
        badge_html = f"<span class='sec-badge'>{badge}</span>" if badge else ""
        st.markdown(
            f"<div class='sec-header'>"
            f"<span class='sec-icon'>{icon}</span>"
            f"<span class='sec-title'>{title}</span>"
            f"{badge_html}</div>",
            unsafe_allow_html=True,
        )

    def _hint(text):
        st.markdown(f"<div class='field-hint'>💡 {text}</div>", unsafe_allow_html=True)

    def _tag_chips(raw, label):
        items = [s.strip() for s in raw.split(",") if s.strip()]
        if not items:
            return
        chips = "".join(f"<span class='tag-chip'>{t}</span>" for t in items)
        st.markdown(
            f"<div style='margin-bottom:4px;font-size:11px;color:#6b7280;'>{label}</div>"
            f"<div style='margin-bottom:10px;'>{chips}</div>",
            unsafe_allow_html=True,
        )

    # ---------------- Resume Form ----------------
    with st.form(f"resume_form_{fk}", clear_on_submit=False):
        _sec_hdr("👤", "Personal Information")
        col1, col2 = st.columns(2)
        with col1:
            # FIX: Do NOT assign back to session_state inside the form.
            # Streamlit batches form widgets — writing to session_state here
            # triggers an immediate rerun on every keystroke, defeating the form.
            # Instead, just render the widget with `value=` for pre-fill.
            # Values are committed to session_state only when submitted=True below.
            st.text_input("👤 Full Name", value=st.session_state.name, placeholder="e.g., Arjun Sharma", key=f"name_input_{fk}")
            st.text_input("📞 Phone Number", value=st.session_state.phone, placeholder="e.g., +91 98765 43210", key=f"phone_input_{fk}")
            st.text_input("📍 Location", value=st.session_state.location, placeholder="e.g., Kolkata, West Bengal", key=f"loc_input_{fk}")
        with col2:
            st.text_input("📧 Email", value=st.session_state.email, placeholder="e.g., arjun@gmail.com", key=f"email_input_{fk}")
            st.text_input("🔗 LinkedIn", value=st.session_state.linkedin, placeholder="e.g., linkedin.com/in/arjun", key=f"ln_input_{fk}")
            st.text_input("🌐 Portfolio", value=st.session_state.portfolio, placeholder="e.g., arjun.dev or github.com/arjun", key=f"port_input_{fk}")
            st.text_input("💼 Job Title / Target Role", value=st.session_state.job_title, placeholder="e.g., Full Stack Developer", key=f"job_input_{fk}")

        _sec_hdr("📝", "Professional Summary")
        st.text_area(
            "Summary",
            value=st.session_state.summary,
            placeholder="Write 3–5 sentences about your career goals, key strengths, and what makes you stand out. E.g., 'Results-driven software engineer with 3+ years building scalable web apps...'",
            height=120,
            key=f"summary_input_{fk}",
        )
        _hint("Aim for 80–200 characters. Recruiters read this first — make it count.")

        _sec_hdr("🛠️", "Skills, Languages, Interests & Soft Skills")
        st.text_area(
            "Technical Skills (comma-separated)",
            value=st.session_state.skills,
            placeholder="e.g., Python, React, Node.js, PostgreSQL, Docker, AWS",
            height=70,
            key=f"skills_input_{fk}",
        )
        _hint("List 5+ skills for best score. Separate each with a comma.")
        # FIX: _tag_chips now reads from the widget key directly (live value),
        # not from session_state.skills which lags by one submit cycle.
        _tag_chips(st.session_state.get(f"skills_input_{fk}", st.session_state.skills), "Preview:")

        st.text_area(
            "Languages (comma-separated)",
            value=st.session_state.languages,
            placeholder="e.g., English, Bengali, Hindi",
            height=60,
            key=f"lang_input_{fk}",
        )
        _tag_chips(st.session_state.get(f"lang_input_{fk}", st.session_state.languages), "Preview:")
        st.text_area(
            "Interests / Hobbies (comma-separated)",
            value=st.session_state.interests,
            placeholder="e.g., Open Source, Machine Learning, Chess, Blogging",
            height=60,
            key=f"int_input_{fk}",
        )
        _tag_chips(st.session_state.get(f"int_input_{fk}", st.session_state.interests), "Preview:")
        st.text_area(
            "Soft Skills (comma-separated)",
            value=st.session_state.Softskills,
            placeholder="e.g., Leadership, Communication, Problem Solving, Teamwork",
            height=60,
            key=f"soft_input_{fk}",
        )
        _tag_chips(st.session_state.get(f"soft_input_{fk}", st.session_state.Softskills), "Preview:")

        _sec_hdr("🧱", "Work Experience", badge=f"{len(st.session_state.experience_entries)} entr{'y' if len(st.session_state.experience_entries)==1 else 'ies'}")
        for idx, exp in enumerate(st.session_state.experience_entries):
            _entry_label = exp.get("title", "") or f"Experience #{idx+1}"
            _entry_company = exp.get("company", "")
            _display = f"{_entry_label} @ {_entry_company}" if _entry_company else _entry_label
            with st.expander(f"🏢 {_display}", expanded=True):
                st.markdown(f"<div class='entry-card-label'>Entry #{idx+1}</div>", unsafe_allow_html=True)
                # FIX: Do not assign back to exp dict here — that mutates session_state
                # inside the form, causing a rerun on every keystroke.
                # Widget keys are unique and Streamlit persists their values automatically.
                # _sync_entries() in the sidebar and the submit handler below read them.
                st.text_input("Job Title", value=exp.get("title", ""), placeholder="e.g., Software Engineer", key=f"title_{idx}_{len(st.session_state.experience_entries)}_{fk}")
                st.text_input("Company", value=exp.get("company", ""), placeholder="e.g., Infosys, TCS, Google", key=f"company_{idx}_{len(st.session_state.experience_entries)}_{fk}")
                st.text_input("Duration", value=exp.get("duration", ""), placeholder="e.g., Jun 2022 – Present", key=f"duration_{idx}_{len(st.session_state.experience_entries)}_{fk}")
                st.text_area("Description", value=exp.get("description", ""), placeholder="• Developed REST APIs using Node.js that reduced response time by 35%\n• Led a team of 4 engineers to deliver the project 2 weeks ahead of schedule", height=100, key=f"description_{idx}_{len(st.session_state.experience_entries)}_{fk}")
                _hint("Use bullet points starting with action verbs. Include metrics where possible.")

        _sec_hdr("🎓", "Education", badge=f"{len(st.session_state.education_entries)} entr{'y' if len(st.session_state.education_entries)==1 else 'ies'}")
        for idx, edu in enumerate(st.session_state.education_entries):
            _edu_label = edu.get("degree", "") or f"Education #{idx+1}"
            _edu_inst = edu.get("institution", "")
            _edu_display = f"{_edu_label} — {_edu_inst}" if _edu_inst else _edu_label
            with st.expander(f"🏫 {_edu_display}", expanded=True):
                st.markdown(f"<div class='entry-card-label'>Entry #{idx+1}</div>", unsafe_allow_html=True)
                st.text_input("Degree / Qualification", value=edu.get("degree", ""), placeholder="e.g., B.Tech in Computer Science", key=f"degree_{idx}_{len(st.session_state.education_entries)}_{fk}")
                st.text_input("Institution", value=edu.get("institution", ""), placeholder="e.g., Jadavpur University", key=f"institution_{idx}_{len(st.session_state.education_entries)}_{fk}")
                st.text_input("Year / Duration", value=edu.get("year", ""), placeholder="e.g., 2019 – 2023", key=f"edu_year_{idx}_{len(st.session_state.education_entries)}_{fk}")
                st.text_area("Academic Details", value=edu.get("details", ""), placeholder="e.g., CGPA: 8.7/10 | Relevant: Data Structures, OS, DBMS | Dean's List 2022", height=80, key=f"edu_details_{idx}_{len(st.session_state.education_entries)}_{fk}")

        _sec_hdr("🚀", "Projects", badge=f"{len(st.session_state.project_entries)} entr{'y' if len(st.session_state.project_entries)==1 else 'ies'}")
        for idx, proj in enumerate(st.session_state.project_entries):
            _proj_label = proj.get("title", "") or f"Project #{idx+1}"
            with st.expander(f"📌 {_proj_label}", expanded=True):
                st.markdown(f"<div class='entry-card-label'>Project #{idx+1}</div>", unsafe_allow_html=True)
                st.text_input("Project Title", value=proj.get("title", ""), placeholder="e.g., AI Resume Builder", key=f"proj_title_{idx}_{len(st.session_state.project_entries)}_{fk}")
                st.text_input("Tech Stack", value=proj.get("tech", ""), placeholder="e.g., Python, Streamlit, OpenAI API, PostgreSQL", key=f"proj_tech_{idx}_{len(st.session_state.project_entries)}_{fk}")
                st.text_input("Duration", value=proj.get("duration", ""), placeholder="e.g., Jan 2024 – Mar 2024  (or  2 months)", key=f"proj_duration_{idx}_{len(st.session_state.project_entries)}_{fk}")
                st.text_area("Description", value=proj.get("description", ""), placeholder="• Built a full-stack resume builder with AI-powered cover letter generation\n• Reduced resume creation time by 70% compared to manual methods", height=100, key=f"proj_desc_{idx}_{len(st.session_state.project_entries)}_{fk}")
                _hint("Describe the problem solved, your role, and the impact or outcome.")

        _sec_hdr("🔗", "Project Links")
        st.text_area(
            "Enter one project link per line:",
            value="\n".join(st.session_state.project_links),
            placeholder="https://github.com/yourname/project1\nhttps://yourproject.netlify.app",
            height=80,
            key=f"proj_links_input_{fk}",
        )
        # FIX: Do NOT write session_state.project_links here — that's a mutation
        # inside the form. It will be committed on submit below.

        _sec_hdr("🏅", "Certificates", badge=f"{len(st.session_state.certificate_links)} entr{'y' if len(st.session_state.certificate_links)==1 else 'ies'}")
        for idx, cert in enumerate(st.session_state.certificate_links):
            _cert_label = cert.get("name", "") or f"Certificate #{idx+1}"
            with st.expander(f"🎖️ {_cert_label}", expanded=True):
                st.markdown(f"<div class='entry-card-label'>Certificate #{idx+1}</div>", unsafe_allow_html=True)
                st.text_input("Certificate Name", value=cert.get("name", ""), placeholder="e.g., AWS Certified Solutions Architect", key=f"cert_name_{idx}_{len(st.session_state.certificate_links)}_{fk}")
                st.text_input("Verification Link", value=cert.get("link", ""), placeholder="e.g., https://credly.com/badges/...", key=f"cert_link_{idx}_{len(st.session_state.certificate_links)}_{fk}")
                st.text_input("Issued Date", value=cert.get("duration", ""), placeholder="e.g., March 2024", key=f"cert_duration_{idx}_{len(st.session_state.certificate_links)}_{fk}")
                st.text_area("Description", value=cert.get("description", ""), placeholder="e.g., Demonstrates expertise in designing distributed systems on AWS. Covers EC2, S3, RDS, and networking.", height=80, key=f"cert_description_{idx}_{len(st.session_state.certificate_links)}_{fk}")

        st.markdown("<div style='height:16px;'></div>", unsafe_allow_html=True)
        btn_col1, btn_col2 = st.columns([1, 1])
        with btn_col1:
            submitted = st.form_submit_button(
                "📑 Generate Resume",
                use_container_width=True,
                type="primary",
            )
        with btn_col2:
            clear_clicked = st.form_submit_button(
                "🗑️ Clear All",
                use_container_width=True,
            )

        if submitted:
            st.session_state["_resume_generated_msg"] = True
            st.session_state["_resume_generating"] = True
            # ── Commit all form widget values to session_state on submit ──────
            # This is the ONLY place we write widget values back — not during typing.
            ss = st.session_state
            ss.name      = ss.get(f"name_input_{fk}",    ss.name)
            ss.email     = ss.get(f"email_input_{fk}",   ss.email)
            ss.phone     = ss.get(f"phone_input_{fk}",   ss.phone)
            ss.location  = ss.get(f"loc_input_{fk}",     ss.location)
            ss.linkedin  = ss.get(f"ln_input_{fk}",      ss.linkedin)
            ss.portfolio = ss.get(f"port_input_{fk}",    ss.portfolio)
            ss.job_title = ss.get(f"job_input_{fk}",     ss.job_title)
            ss.summary   = ss.get(f"summary_input_{fk}", ss.summary)
            ss.skills    = ss.get(f"skills_input_{fk}",  ss.skills)
            ss.languages = ss.get(f"lang_input_{fk}",    ss.languages)
            ss.interests = ss.get(f"int_input_{fk}",     ss.interests)
            ss.Softskills = ss.get(f"soft_input_{fk}",   ss.Softskills)
            # Project links
            _pl_raw = ss.get(f"proj_links_input_{fk}", "")
            ss.project_links = [lnk.strip() for lnk in _pl_raw.splitlines() if lnk.strip()]
            # Sync experience entry dicts
            _n_exp = len(ss.experience_entries)
            for _i, _e in enumerate(ss.experience_entries):
                _e["title"]       = ss.get(f"title_{_i}_{_n_exp}_{fk}",       _e.get("title", ""))
                _e["company"]     = ss.get(f"company_{_i}_{_n_exp}_{fk}",     _e.get("company", ""))
                _e["duration"]    = ss.get(f"duration_{_i}_{_n_exp}_{fk}",    _e.get("duration", ""))
                _e["description"] = ss.get(f"description_{_i}_{_n_exp}_{fk}", _e.get("description", ""))
            # Sync education entry dicts
            _n_edu = len(ss.education_entries)
            for _i, _e in enumerate(ss.education_entries):
                _e["degree"]      = ss.get(f"degree_{_i}_{_n_edu}_{fk}",      _e.get("degree", ""))
                _e["institution"] = ss.get(f"institution_{_i}_{_n_edu}_{fk}", _e.get("institution", ""))
                _e["year"]        = ss.get(f"edu_year_{_i}_{_n_edu}_{fk}",    _e.get("year", ""))
                _e["details"]     = ss.get(f"edu_details_{_i}_{_n_edu}_{fk}", _e.get("details", ""))
            # Sync project entry dicts
            _n_proj = len(ss.project_entries)
            for _i, _e in enumerate(ss.project_entries):
                _e["title"]       = ss.get(f"proj_title_{_i}_{_n_proj}_{fk}",    _e.get("title", ""))
                _e["tech"]        = ss.get(f"proj_tech_{_i}_{_n_proj}_{fk}",     _e.get("tech", ""))
                _e["duration"]    = ss.get(f"proj_duration_{_i}_{_n_proj}_{fk}", _e.get("duration", ""))
                _e["description"] = ss.get(f"proj_desc_{_i}_{_n_proj}_{fk}",     _e.get("description", ""))
            # Sync certificate entry dicts
            _n_cert = len(ss.certificate_links)
            for _i, _e in enumerate(ss.certificate_links):
                _e["name"]        = ss.get(f"cert_name_{_i}_{_n_cert}_{fk}",        _e.get("name", ""))
                _e["link"]        = ss.get(f"cert_link_{_i}_{_n_cert}_{fk}",        _e.get("link", ""))
                _e["duration"]    = ss.get(f"cert_duration_{_i}_{_n_cert}_{fk}",    _e.get("duration", ""))
                _e["description"] = ss.get(f"cert_description_{_i}_{_n_cert}_{fk}", _e.get("description", ""))

        if clear_clicked:
            st.session_state["_confirm_clear"] = True

    # ── Clear confirmation (outside form so it can render fresh buttons) ──────
    if st.session_state.get("_confirm_clear"):
        st.markdown(
            "<div class='confirm-warn'>⚠️ <strong>This will erase all entered data.</strong> "
            "This cannot be undone.</div>",
            unsafe_allow_html=True,
        )
        cc1, cc2 = st.columns([1, 1])
        with cc1:
            if st.button("✅ Yes, Clear", key="confirm_clear_yes", use_container_width=True):
                _new_counter = st.session_state.get("form_key_counter", 0) + 1
                resume_fields = ["name", "email", "phone", "linkedin", "location",
                                 "portfolio", "summary", "skills", "languages",
                                 "interests", "Softskills", "job_title"]
                for _f in resume_fields:
                    st.session_state[_f] = ""
                st.session_state["experience_entries"] = [{"title": "", "company": "", "duration": "", "description": ""}]
                st.session_state["education_entries"] = [{"degree": "", "institution": "", "year": "", "details": ""}]
                st.session_state["project_entries"] = [{"title": "", "tech": "", "duration": "", "description": ""}]
                st.session_state["project_links"] = []
                st.session_state["certificate_links"] = [{"name": "", "link": "", "duration": "", "description": ""}]
                for _key in ["generated_html", "ai_output", "cover_letter",
                             "cover_letter_html", "encoded_profile_image"]:
                    st.session_state.pop(_key, None)
                st.session_state["form_key_counter"] = _new_counter
                st.session_state.pop("_confirm_clear", None)
                st.rerun()
        with cc2:
            if st.button("❌ Cancel", key="confirm_clear_no", use_container_width=True):
                st.session_state.pop("_confirm_clear", None)
                st.rerun()

    st.markdown("""
    <style>
        .heading-large {
            font-size: 36px;
            font-weight: bold;
            color: #336699;
        }
        .subheading-large {
            font-size: 30px;
            font-weight: bold;
            color: #336699;
        }
        .tab-section {
            margin-top: 20px;
        }
    </style>
    """, unsafe_allow_html=True)

    # ── Loading button pulse animation CSS ────────────────────────────────────
    st.markdown("""
    <style>
    /* Pulse animation for buttons during processing */
    @keyframes btn-pulse {
        0%   { box-shadow: 0 0 0 0 rgba(0,200,255,0.55); }
        70%  { box-shadow: 0 0 0 10px rgba(0,200,255,0); }
        100% { box-shadow: 0 0 0 0 rgba(0,200,255,0); }
    }
    /* Spinner overlay for the stSpinner */
    [data-testid="stSpinner"] > div {
        background: rgba(10, 20, 40, 0.75) !important;
        border: 1px solid rgba(0,200,255,0.4) !important;
        border-radius: 12px !important;
        padding: 14px 20px !important;
        backdrop-filter: blur(10px) !important;
        color: #93c5fd !important;
        font-size: 14px !important;
        font-weight: 500 !important;
    }
    [data-testid="stSpinner"] svg {
        color: #4da6ff !important;
        stroke: #4da6ff !important;
    }
    /* Download button loading feel */
    [data-testid="stDownloadButton"] > button {
        position: relative;
        overflow: hidden;
    }
    [data-testid="stDownloadButton"] > button:active::after {
        content: "";
        position: absolute;
        inset: 0;
        background: rgba(0,200,255,0.18);
        animation: btn-pulse 0.6s ease-out;
    }
    /* Form submit button active state — pulse */
    button[kind="formSubmit"]:active,
    button[data-testid="baseButton-primary"]:active {
        animation: btn-pulse 0.5s ease-out !important;
    }
    /* Disabled state for buttons during loading */
    .stButton > button:disabled,
    button[disabled] {
        opacity: 0.55 !important;
        cursor: not-allowed !important;
        animation: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # --- Visual Resume Preview Section (only shown after form is submitted) ---
    if st.session_state.get("_resume_generated_msg"):
        st.success("✅ Resume Generated Successfully! Scroll down to preview or download.")
        st.session_state["_resume_generated_msg"] = False  # show only once per submit

    if "generated_html" in st.session_state:
        st.markdown("## 🧾 <span style='color:#336699;'>Resume Preview</span>", unsafe_allow_html=True)
        st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

        left, right = st.columns([1, 2])

        with left:
            st.markdown(f"""
                <h2 style='color:#2f2f2f;margin-bottom:0;'>{st.session_state['name']}</h2>
                <h4 style='margin-top:5px;color:#444;'>{st.session_state['job_title']}</h4>
                <p style='font-size:14px;'>
                📍 {st.session_state['location']}<br>
                📞 {st.session_state['phone']}<br>
                📧 <a href="mailto:{st.session_state['email']}">{st.session_state['email']}</a><br>
                🔗 <a href="{st.session_state['linkedin']}" target="_blank">LinkedIn</a><br>
                🌐 <a href="{st.session_state['portfolio']}" target="_blank">Portfolio</a>
                </p>
            """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Skills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            _skill_items = [s.strip() for s in st.session_state["skills"].split(",") if s.strip()]
            if _skill_items:
                _chips = "".join(f"<span class='tag-chip'>{s}</span>" for s in _skill_items)
                st.markdown(f"<div style='margin-bottom:8px;'>{_chips}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Languages</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            _lang_items = [l.strip() for l in st.session_state["languages"].split(",") if l.strip()]
            if _lang_items:
                _lang_chips = "".join(f"<span class='tag-chip'>{l}</span>" for l in _lang_items)
                st.markdown(f"<div style='margin-bottom:8px;'>{_lang_chips}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Interests</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            _int_items = [i.strip() for i in st.session_state["interests"].split(",") if i.strip()]
            if _int_items:
                _int_chips = "".join(f"<span class='tag-chip'>{t}</span>" for t in _int_items)
                st.markdown(f"<div style='margin-bottom:8px;'>{_int_chips}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Soft Skills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            _soft_items = [i.strip() for i in st.session_state["Softskills"].split(",") if i.strip()]
            if _soft_items:
                _soft_chips = "".join(f"<span class='tag-chip'>{t}</span>" for t in _soft_items)
                st.markdown(f"<div style='margin-bottom:8px;'>{_soft_chips}</div>", unsafe_allow_html=True)

        with right:
            st.markdown("<h4 style='color:#336699;'>Summary</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            summary_text = st.session_state["summary"].replace("\n", "<br>")
            st.markdown(f"<p style='font-size:17px;'>{summary_text}</p>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Experience</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for exp in st.session_state.experience_entries:
                if exp["company"] or exp["title"]:
                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px; border-radius:8px;'>
                        <div style='display:flex; justify-content:space-between;'>
                            <b>🏢 {exp['company']}</b><span style='color:gray;'>📆 {exp['duration']}</span>
                        </div>
                        <div style='font-size:14px;'>💼 <i>{exp['title']}</i></div>
                        <div style='font-size:17px;'>📝 {exp['description']}</div>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>🎓 Education</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for edu in st.session_state.education_entries:
                if edu["institution"] or edu["degree"]:
                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px 15px; border-radius:8px;'>
                        <div style='display:flex; justify-content:space-between; font-size:16px; font-weight:bold;'>
                            <span>🏫 {edu['institution']}</span>
                            <span style='color:gray;'>📅 {edu['year']}</span>
                        </div>
                        <div style='font-size:14px;'>🎓 <i>{edu['degree']}</i></div>
                        <div style='font-size:14px;'>📄 {edu['details']}</div>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Projects</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for proj in st.session_state.project_entries:
                if proj.get("title"):
                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px;'>
                        <strong style='font-size:16px;'>{proj['title']}</strong><br>
                        <span style='font-size:14px;'>🛠️ <strong>Tech Stack:</strong> {proj['tech']}</span><br>
                        <span style='font-size:14px;'>⏳ <strong>Duration:</strong> {proj['duration']}</span><br>
                        <span style='font-size:17px;'>📝 <strong>Description:</strong> {proj['description']}</span>
                    </div>
                    """, unsafe_allow_html=True)

            if st.session_state.project_links:
                st.markdown("<h4 style='color:#336699;'>Project Links</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for i, link in enumerate(st.session_state.project_links):
                    st.markdown(f"[🔗 Project {i+1}]({link})", unsafe_allow_html=True)

            if st.session_state.certificate_links:
                st.markdown("<h4 style='color:#336699;'>Certificates</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for cert in st.session_state.certificate_links:
                    if cert["name"]:
                        name_html = (f"<a href=\"{cert['link']}\" target=\"_blank\"><b>\U0001f4c4 {cert['name']}</b></a>"
                                     if cert.get("link") else f"<b>\U0001f4c4 {cert['name']}</b>")
                        st.markdown(f"""
                        <div style='display:flex; justify-content:space-between;'>
                            {name_html}
                            <span style='color:gray;'>{cert['duration']}</span>
                        </div>
                        <div style='margin-bottom:10px; font-size:14px;'>{cert['description']}</div>
                        """, unsafe_allow_html=True)

import re

with tab2:
    st.markdown("## ✨ <span style='color:#336699;'>Enhanced AI Resume Preview</span>", unsafe_allow_html=True)
    st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

    col1, spacer, col2 = st.columns([1, 0.2, 1])

    with col1:
        if st.button("🔁 Clear Preview"):
            st.session_state.pop("ai_output", None)
            st.toast("🗑️ Preview cleared!")

    with col2:
        if st.button("🚀 Generate AI Resume Preview"):

            # ── Job Title Spell-Correction ─────────────────────────────────────────
            # Fuzzy-matches the user's input against the canonical role list.
            # If a close match is found (≥ 70% similarity) it silently corrects the
            # title in session_state before anything else runs.
            import difflib as _difflib

            _CANONICAL_JOB_TITLES = [
                # ── Software Engineering ──────────────────────────────────────
                "Software Engineer", "Senior Software Engineer", "Staff Software Engineer",
                "Frontend Developer", "Backend Developer", "Full Stack Developer",
                "React Developer", "Angular Developer", "Vue.js Developer",
                "Node.js Developer", "Python Developer", "Java Developer",
                "Go Developer", "Rust Developer", "C++ Developer",
                "PHP Developer", "Ruby on Rails Developer", ".NET Developer",
                "iOS Developer", "Android Developer", "Mobile App Developer",
                "Flutter Developer", "React Native Developer",
                "Embedded Systems Engineer", "Firmware Engineer",
                # ── Data & AI ─────────────────────────────────────────────────
                "Data Scientist", "Senior Data Scientist",
                "Data Analyst", "Senior Data Analyst",
                "Data Engineer", "Senior Data Engineer",
                "Machine Learning Engineer", "ML Engineer",
                "AI Engineer", "Generative AI Engineer",
                "LLM Engineer", "Prompt Engineer",
                "Computer Vision Engineer", "NLP Engineer",
                "Business Intelligence Analyst", "BI Developer",
                "Analytics Engineer", "Quantitative Analyst",
                # ── Infrastructure & Cloud ────────────────────────────────────
                "DevOps Engineer", "Senior DevOps Engineer",
                "Cloud Engineer", "AWS Engineer", "Azure Engineer", "GCP Engineer",
                "Site Reliability Engineer", "Platform Engineer",
                "Infrastructure Engineer", "Systems Administrator",
                "Network Engineer", "Network Administrator",
                "Database Administrator", "Database Engineer",
                # ── Security ──────────────────────────────────────────────────
                "Cybersecurity Analyst", "Information Security Analyst",
                "Security Engineer", "Penetration Tester",
                "SOC Analyst", "Cloud Security Engineer",
                # ── QA & Testing ──────────────────────────────────────────────
                "QA Engineer", "QA Analyst", "SDET",
                "Automation Test Engineer", "Performance Test Engineer",
                # ── Architecture & Leadership ─────────────────────────────────
                "Solutions Architect", "Cloud Architect", "Enterprise Architect",
                "Technical Lead", "Engineering Manager",
                "Chief Technology Officer", "VP of Engineering",
                # ── Product & Design ──────────────────────────────────────────
                "Product Manager", "Senior Product Manager",
                "Product Owner", "Technical Product Manager",
                "UI/UX Designer", "UX Designer", "UI Designer",
                "Product Designer", "Graphic Designer",
                "Visual Designer", "Motion Designer",
                "Interaction Designer", "Design Lead",
                # ── Project & Delivery ────────────────────────────────────────
                "Project Manager", "Senior Project Manager",
                "Scrum Master", "Agile Coach",
                "Program Manager", "Delivery Manager",
                "IT Project Manager",
                # ── Business & Analysis ───────────────────────────────────────
                "Business Analyst", "Senior Business Analyst",
                "Systems Analyst", "Functional Consultant",
                "ERP Consultant", "Salesforce Developer",
                "Salesforce Administrator",
                # ── Marketing & Growth ────────────────────────────────────────
                "Digital Marketing Specialist", "Digital Marketing Manager",
                "SEO Specialist", "SEM Specialist",
                "Content Strategist", "Content Writer",
                "Social Media Manager", "Growth Hacker",
                "Performance Marketing Manager", "Email Marketing Specialist",
                "Brand Manager", "Marketing Analyst",
                "E-commerce Specialist", "E-commerce Manager",
                # ── Finance & Accounting ──────────────────────────────────────
                "Financial Analyst", "Senior Financial Analyst",
                "Chartered Accountant", "Cost Accountant",
                "Investment Analyst", "Equity Research Analyst",
                "Risk Analyst", "Credit Analyst",
                "Fintech Developer", "Quantitative Developer",
                "Accounts Manager", "Tax Consultant",
                "Audit Manager", "CFO",
                # ── HR & People ───────────────────────────────────────────────
                "HR Manager", "HR Business Partner",
                "Talent Acquisition Specialist", "Recruiter",
                "Technical Recruiter", "HR Analyst",
                "Learning and Development Manager", "Compensation Analyst",
                # ── Sales & CRM ───────────────────────────────────────────────
                "Sales Manager", "Account Executive",
                "Business Development Manager", "Sales Engineer",
                "Technical Sales Engineer", "Pre-Sales Consultant",
                "Customer Success Manager",
                # ── Operations & Supply Chain ─────────────────────────────────
                "Operations Manager", "Supply Chain Analyst",
                "Logistics Manager", "Procurement Manager",
                # ── Niche Tech ────────────────────────────────────────────────
                "Game Developer", "Blockchain Developer",
                "AR/VR Developer", "IoT Engineer",
                "EdTech Developer", "HealthTech Developer",
                "Healthcare Software Engineer",
                # ── Writing & Documentation ───────────────────────────────────
                "Technical Writer", "API Documentation Specialist",
                # ── Support & Admin ───────────────────────────────────────────
                "IT Support Engineer", "Help Desk Analyst",
                "Systems Engineer",
            ]

            def _normalize_job_title(raw: str) -> tuple[str, bool]:
                """
                Returns (corrected_title, was_corrected).
                Tries an exact case-insensitive match first, then fuzzy.
                Threshold: 0.70 similarity — catches typos but won't misfire on
                completely unrelated inputs (e.g. 'Chef').
                """
                if not raw or not raw.strip():
                    return raw, False
                raw_stripped = raw.strip()
                raw_lower = raw_stripped.lower()

                # 1. Exact case-insensitive match → silently fix capitalisation only
                #    was_corrected = False here because only casing changed, not spelling
                for title in _CANONICAL_JOB_TITLES:
                    if title.lower() == raw_lower:
                        return title, False  # no toast — user spelled it right

                # 2. Fuzzy match against canonical list (case-insensitive compare)
                lower_map = {t.lower(): t for t in _CANONICAL_JOB_TITLES}
                matches = _difflib.get_close_matches(
                    raw_lower,
                    lower_map.keys(),
                    n=1,
                    cutoff=0.70,
                )
                if matches:
                    corrected = lower_map[matches[0]]
                    return corrected, True

                # 3. No confident match — return original unchanged
                return raw_stripped, False

            _raw_title = st.session_state.get("job_title", "").strip()
            _corrected_title, _was_corrected = _normalize_job_title(_raw_title)
            if _was_corrected:
                st.session_state["job_title"] = _corrected_title
                st.toast(f"✏️ Job title corrected: \"{_raw_title}\" → \"{_corrected_title}\"", icon="✅")
            # ── End Job Title Spell-Correction ────────────────────────────────────

            # ── Helper: detect if a field has real user-entered content ──
            def _has_real_content(value, min_len=4):
                if not value:
                    return False
                v = str(value).strip().lower()
                placeholders = {"placeholder", "sample", "n/a", "na", "none", "tbd", ""}
                return len(v) >= min_len and v not in placeholders

            # ── Normalize experience entries + detect if user provided real data ──
            experience_entries = st.session_state.get('experience_entries', [])
            normalized_experience_entries = []
            user_has_real_experience = False

            for entry in experience_entries:
                if isinstance(entry, dict):
                    title    = entry.get("title", "").strip()
                    company  = entry.get("company", "").strip()
                    duration = entry.get("duration", "").strip()
                    desc     = entry.get("description", "").strip()
                    if _has_real_content(company) or _has_real_content(desc):
                        user_has_real_experience = True
                    parts = []
                    if title:    parts.append(f"Role: {title}")
                    if company:  parts.append(f"Company: {company}")
                    if duration: parts.append(f"Duration: {duration}")
                    if desc:     parts.append(f"Description: {desc}")
                    formatted = "\n".join(parts)
                else:
                    formatted = entry.strip()
                    if _has_real_content(formatted, min_len=10):
                        user_has_real_experience = True
                if formatted:
                    normalized_experience_entries.append(formatted)

            # ── Normalize project entries + detect if user provided real data ──
            project_entries = st.session_state.get('project_entries', [])
            normalized_project_entries = []
            user_has_real_projects = False

            for entry in project_entries:
                if isinstance(entry, dict):
                    title    = entry.get("title", "").strip()
                    tech     = entry.get("tech", "").strip()
                    duration = entry.get("duration", "").strip()
                    desc     = entry.get("description", "").strip()
                    if _has_real_content(title) or _has_real_content(desc):
                        user_has_real_projects = True
                    parts = []
                    if title:    parts.append(f"Title: {title}")
                    if tech:     parts.append(f"Tech Stack: {tech}")
                    if duration: parts.append(f"Duration: {duration}")
                    if desc:     parts.append(f"Description: {desc}")
                    formatted = "\n".join(parts)
                else:
                    formatted = entry.strip()
                    if _has_real_content(formatted, min_len=10):
                        user_has_real_projects = True
                if formatted:
                    normalized_project_entries.append(formatted)

            # ── Build experience prompt section based on whether user has real data ──
            if user_has_real_experience:
                experience_instruction = f"""2. EXPERIENCE (USER HAS PROVIDED REAL DATA — LIGHT POLISH ONLY):
               The user has already entered their experience. Your ONLY job is to lightly polish the bullet descriptions.
               STRICT RULES — violating any of these is an error:
               - DO NOT change, rename, or reword any company name. Copy it exactly as given.
               - DO NOT change any role/job title. Copy it exactly as given.
               - DO NOT change any date or duration. Copy it exactly as given.
               - DO NOT add new entries that the user did not provide.
               - DO NOT restructure or reorder entries.
               - ONLY rewrite existing description bullets to be cleaner, more ATS-friendly, and results-oriented.
               - If a description is already strong, make minimal or no edits.
               Present as A., B., C. using the user's original data below:
               {normalized_experience_entries}"""
            else:
                experience_instruction = f"""2. EXPERIENCE (NO USER DATA PROVIDED — GENERATE REALISTIC DUMMY DATA):
               The user's target job title is: "{st.session_state['job_title']}".
               Generate 2–3 realistic experience entries showing a natural career progression toward this SPECIFIC role.

               COMPANY NAMING RULES:
               - Choose company names that are REALISTIC and DOMAIN-APPROPRIATE for "{st.session_state['job_title']}".
               - For tech/software roles: use companies like Google, Microsoft, Amazon, Flipkart, Razorpay, Zomato, Swiggy, PhonePe, Freshworks, Zoho, Paytm, Byju's, Ola, or similar product/tech companies.
               - For finance/banking roles: use HDFC Bank, ICICI Bank, Axis Bank, Kotak, JPMorgan, Goldman Sachs, KPMG, Deloitte, EY, or similar.
               - For data/analytics roles: use Mu Sigma, Fractal Analytics, ThoughtWorks, Tiger Analytics, or product companies with data teams.
               - For marketing/design/HR roles: use relevant advertising agencies, startups, or consumer brands.
               - NEVER use the same company list for every job title. Pick what makes sense for THIS domain.
               - Each entry MUST have a DIFFERENT company name.

               DATE RULES:
               - Dates must show logical progression (most recent first, oldest last).
               - NEVER repeat the same date range across entries.
               - Use realistic tenure lengths (1–3 years per role).

               CONTENT RULES:
               - Role titles must be SPECIFIC to "{st.session_state['job_title']}" — not generic.
               - Each entry must have 3–4 achievement bullets with measurable metrics relevant to this role.
               - Tools, technologies, and responsibilities must match what someone in "{st.session_state['job_title']}" actually does.
               Present as A., B., C. format."""

            # ── Build projects prompt section based on whether user has real data ──
            if user_has_real_projects:
                projects_instruction = f"""3. PROJECTS (USER HAS PROVIDED REAL DATA — LIGHT POLISH ONLY):
               The user has already entered their projects. Your ONLY job is to lightly improve the descriptions.
               STRICT RULES — violating any of these is an error:
               - DO NOT change, rename, or reword any project title. Copy it exactly as given.
               - DO NOT change any tech stack. Copy it exactly as given.
               - DO NOT change any duration/date. Copy it exactly as given.
               - DO NOT add new projects that the user did not provide.
               - DO NOT restructure or reorder entries.
               - ONLY rewrite existing description bullets to be more technical, impactful, and measurable.
               - If a description is already strong, make minimal or no edits.
               Present as A., B., C. using the user's original data below:
               {normalized_project_entries}"""
            else:
                projects_instruction = f"""3. PROJECTS (NO USER DATA PROVIDED — GENERATE REALISTIC DUMMY DATA):
               The user's target job title is: "{st.session_state['job_title']}".
               Generate 2–3 realistic, industry-standard projects that someone in THIS SPECIFIC ROLE would actually build.

               PROJECT NAMING RULES:
               - Project names MUST be derived directly from the domain of "{st.session_state['job_title']}".
               - Do NOT use any generic or pre-existing example names. Think from scratch for this role.
               - The name should sound like a real internal tool, product feature, or system — not a textbook exercise.
               - For example: a "Digital Marketing Manager" would NOT build a Kubernetes migration; they'd build a "Campaign Attribution Analytics Dashboard" or "SEO Content Performance Tracker".
               - A "Chartered Accountant" would NOT build a fraud detection pipeline; they'd build a "GST Reconciliation Automation Tool" or "Audit Trail Ledger System".
               - Generate names that ONLY make sense for "{st.session_state['job_title']}" — if the name could apply to a different role, discard it and try again.

               TECH STACK RULES:
               - Use ONLY tools, frameworks, and technologies that are standard for "{st.session_state['job_title']}".
               - Do NOT use backend/data engineering stacks (Kafka, Spark, Flink) for non-technical roles.
               - Do NOT use frontend stacks (React, TypeScript) for data or infrastructure roles.
               - Match the stack precisely to what this role uses day-to-day.

               DATE RULES:
               - All project dates must be DIFFERENT from each other.
               - Ordered most recent first. Use realistic 3–5 month project durations.

               CONTENT RULES:
               - Each project must have 3–5 strong technical/functional bullets with measurable outcomes.
               - Bullets must use vocabulary and actions that belong to "{st.session_state['job_title']}" — not generic software engineering language unless the role demands it.
               - NEVER use names like "Sample Project", "My Project", "Project 1", or any name from unrelated domains.
               Present as A., B., C. format."""

            # ── Build certificates prompt section ────────────────────────────────
            _cert_entries = st.session_state.get("certificate_links", [])
            _user_has_real_certs = any(
                _has_real_content(c.get("name", "")) for c in _cert_entries
            )
            if _user_has_real_certs:
                # Build structured representation passing ALL fields to the AI
                _cert_lines = []
                for c in _cert_entries:
                    _cname = c.get("name", "").strip()
                    _clink = c.get("link", "").strip()
                    _cdur  = c.get("duration", "").strip()
                    _cdesc = c.get("description", "").strip()
                    if not _cname:
                        continue
                    parts = [f"Name: {_cname}"]
                    if _clink:  parts.append(f"Link: {_clink}")
                    if _cdur:   parts.append(f"Date: {_cdur}")
                    if _cdesc:  parts.append(f"Description: {_cdesc}")
                    _cert_lines.append("\n".join(parts))
                _cert_data_str = "\n\n".join(_cert_lines)
                certificates_instruction = f"""8. CERTIFICATES (USER HAS PROVIDED REAL DATA — PRESERVE EVERYTHING):
               STRICT RULES — violating any of these is an error:
               - DO NOT change the certificate name. Copy it exactly as given.
               - DO NOT change, alter, or reformat the issued date in ANY way. Copy it character-for-character.
               - DO NOT change the verification link. Copy it exactly as given.
               - ONLY lightly polish the description if one was provided; otherwise leave it as-is.
               - DO NOT invent or add certificates the user did not provide.
               Output each certificate on one line as: [Name] - [Issuer] ([Date])
               where [Date] is EXACTLY the date the user entered — no reformatting, no substitution.
               User's certificate data:
               {_cert_data_str}"""
            else:
                certificates_instruction = f"""8. CERTIFICATES (NO USER DATA — GENERATE):
               Generate 3 realistic, industry-recognized certifications for {st.session_state['job_title']} with provider name."""

            enhance_prompt = f"""
            You are a professional Resume Optimization Specialist with deep expertise in ATS systems,
            industry hiring standards, and professional resume writing.
            Target role: "{st.session_state['job_title']}"

            ⚠️ CRITICAL DOMAIN RULE — READ BEFORE GENERATING ANYTHING:
            Every single piece of content you generate (projects, experience, skills, summary, certificates)
            MUST be tailored specifically and exclusively to the role: "{st.session_state['job_title']}".
            - A "Digital Marketing Manager" should NEVER have Java/Kafka/Kubernetes projects.
            - A "Java Backend Developer" should NEVER have Excel/VLOOKUP or GST reconciliation content.
            - A "Chartered Accountant" should NEVER have React/TypeScript or ML pipeline content.
            Before writing any section, ask yourself: "Would someone hiring a {st.session_state['job_title']} 
            care about this?" If no → discard and generate something domain-appropriate.
            The output must look like it was written BY a {st.session_state['job_title']} FOR a {st.session_state['job_title']} role.

            LANGUAGE & TONE:
            - Neutral, professional, ATS-optimized tone throughout.
            - NO first-person (I, me, my, we, our). NO gendered pronouns.
            - NO marketing terms (rockstar, guru, ninja).
            - Concise, quantifiable, outcome-focused language only.
            - Each section must use a DISTINCT verb set — no verb/phrase/action repeated across sections.

            SECTION LANGUAGE RULES:
            - SUMMARY: Third-person PRESENT tense. Strategic positioning bullets only.
            - EXPERIENCE: PAST tense. Ownership, delivery, accountability language.
            - PROJECTS: PAST tense. Use vocabulary natural to the {st.session_state['job_title']} domain.
            - SKILLS / SOFTSKILLS: Nouns only. Comma-separated list.
            - INTERESTS: Professional domain-engagement language.

            SECTION INSTRUCTIONS:

            1. SUMMARY:
               {"Enhance the provided summary." if _has_real_content(st.session_state.get('summary','')) else f"Generate a 3-4 bullet professional summary for a {st.session_state['job_title']}."}
               Write 3-4 bullets defining professional identity, specialization, and measurable strengths.

            {experience_instruction}

            {projects_instruction}

            4. SKILLS:
               {"Enhance and expand the provided skills list to be more specific and ATS-optimized for this role." if _has_real_content(st.session_state.get('skills','')) else f"Generate 6-8 current, highly specific technical/functional skills that are EXCLUSIVELY relevant to a {st.session_state['job_title']}. Do NOT list generic skills that apply to every role."}
               List only - no sentences.

            5. SOFTSKILLS:
               {"Enhance the provided soft skills to align with this role's demands." if _has_real_content(st.session_state.get('Softskills','')) else f"Generate 5-6 soft skills that are most valued specifically for a {st.session_state['job_title']} — not a generic list."}
               List only - no sentences.

            6. LANGUAGES:
               {"Use provided languages." if _has_real_content(st.session_state.get('languages','')) else "Generate 2-3 relevant languages (include English)."}

            7. INTERESTS:
               {"Enhance provided interests." if _has_real_content(st.session_state.get('interests','')) else f"Generate 3-5 professional interests aligned with {st.session_state['job_title']}."}

            {certificates_instruction}

            OUTPUT FORMAT (FOLLOW EXACTLY):

            Summary:
            * [bullet]
            * [bullet]
            * [bullet]

            Experience:
            A. [Company Name] ([Start Month Year - End Month Year or Present])
               * [Role Title]
               * [Achievement with metric]
               * [Ownership/delivery bullet]
               * [Impact/improvement bullet]

            B. [Company Name] ([Start Month Year - End Month Year])
               * [Role Title]
               * [Achievement]
               * [Responsibility]

            Projects:
            A. [Unique Realistic Project Name]
               * Tech Stack: [tools]
               * Duration: [Start Month Year - End Month Year]
               * Description:
                 - [bullet]
                 - [bullet]
                 - [bullet]

            B. [Unique Realistic Project Name]
               * Tech Stack: [tools]
               * Duration: [Start Month Year - End Month Year]
               * Description:
                 - [bullet]
                 - [bullet]
                 - [bullet]

            Skills:
            [skill1], [skill2], [skill3], [skill4], [skill5], [skill6]

            SoftSkills:
            [soft1], [soft2], [soft3], [soft4], [soft5]

            Languages:
            [lang1], [lang2]

            Interests:
            [interest1], [interest2], [interest3]

            Certificates:
            [Certificate] - [Provider] ([Year/Level])
            [Certificate] - [Provider] ([Year/Level])

            SOURCE DATA TO ENHANCE (where provided):
            Summary: {st.session_state.get('summary', '')}
            Skills: {st.session_state.get('skills', '')}
            SoftSkills: {st.session_state.get('Softskills', '')}
            Languages: {st.session_state.get('languages', '')}
            Interests: {st.session_state.get('interests', '')}
            Certificates: {[{"name": c.get("name",""), "date": c.get("duration",""), "link": c.get("link","")} for c in st.session_state.get('certificate_links', []) if c.get('name')]}

            CRITICAL RULES:
            - Output ONLY the formatted resume content. No explanations, no preamble.
            - NEVER use "Sample Project", "Previous Company", "Placeholder", or any generic names.
            - ALL dates must be DIFFERENT across experience entries AND across project entries.
            - Experience dates must show logical career progression (most recent first, oldest last).
            - Project dates must all be different and logically ordered (most recent first).
            - If user provided real experience/project data, PRESERVE every company name, project title, tech stack, and date EXACTLY as written. Only polish the description bullets.
            - If user provided real certificate data, PRESERVE the certificate name, issued date, and link EXACTLY as written. Never substitute or reformat the date.
            - UNIQUENESS RULE: Every generation must produce fresh, original content. Never repeat the same project names, company names, or bullet phrasing across different runs. Treat each generation as a brand-new resume for a brand-new person.
            - DOMAIN LOCK: Every project name, tech stack, skill, and certificate must be something a real "{st.session_state['job_title']}" would have. Cross-domain content is forbidden.
            """





            import uuid as _uuid
            import datetime as _datetime
            _unique_seed = _uuid.uuid4().hex[:8]
            _timestamp = _datetime.datetime.now().strftime("%H%M%S")
            enhance_prompt += f"\n[Generation ID: {_unique_seed}-{_timestamp} — produce content unique to this exact run]"

            with st.spinner("🧠 Thinking..."):
                ai_output = call_llm(enhance_prompt, session=st.session_state)
                st.session_state["ai_output"] = ai_output

    # ------------------------- PARSE + RENDER -------------------------
    if "ai_output" in st.session_state:
        ai_output = st.session_state["ai_output"]

        def extract_section(label, output, default=""):
            pattern = rf"(?m)^{re.escape(label)}:\s*\n?(.*?)(?=\n[A-Za-z][A-Za-z\s]*:\s*\n?|\Z)"
            match = re.search(pattern, output, re.DOTALL)
            return match.group(1).strip() if match else default

        summary_enhanced  = extract_section("Summary",      ai_output, st.session_state["summary"])
        experience_raw    = extract_section("Experience",   ai_output)
        projects_raw      = extract_section("Projects",     ai_output)
        skills_list       = extract_section("Skills",       ai_output, st.session_state["skills"])
        softskills_list   = extract_section("SoftSkills",   ai_output, st.session_state["Softskills"])
        languages_list    = extract_section("Languages",    ai_output, st.session_state["languages"])
        interests_list    = extract_section("Interests",    ai_output, st.session_state["interests"])
        certificates_list = extract_section("Certificates", ai_output)

        experience_blocks = [b.strip() for b in re.split(r"\n(?=[A-Z]\. )", experience_raw.strip()) if b.strip()]
        projects_blocks   = [b.strip() for b in re.split(r"\n(?=[A-Z]\. )", projects_raw.strip())   if b.strip()]

        left, right = st.columns([1, 2])

        with left:
            st.markdown(
                f"<h2 style='color:#2f2f2f;margin-bottom:0;'>{st.session_state['name']}</h2>"
                f"<h4 style='margin-top:5px;color:#444;'>{st.session_state['job_title']}</h4>"
                f"<p style='font-size:14px;'>"
                f"📍 {st.session_state['location']}<br>"
                f"📞 {st.session_state['phone']}<br>"
                f"📧 <a href='mailto:{st.session_state['email']}'>{st.session_state['email']}</a><br>"
                f"🔗 <a href='{st.session_state['linkedin']}' target='_blank'>LinkedIn</a><br>"
                f"🌐 <a href='{st.session_state['portfolio']}' target='_blank'>Portfolio</a></p>",
                unsafe_allow_html=True
            )

            def render_bullet_section(title, items):
                st.markdown(f"<h4 style='color:#336699;'>{title}</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for item in [i.strip() for i in items.split(",") if i.strip()]:
                    st.markdown(f"<div style='margin-left:10px;'>• {item}</div>", unsafe_allow_html=True)

            render_bullet_section("Skills",      skills_list)
            render_bullet_section("Languages",   languages_list)
            render_bullet_section("Interests",   interests_list)
            render_bullet_section("Soft Skills", softskills_list)

        with right:
            formatted_summary = summary_enhanced.replace("\n• ", "<br>• ").replace("\n* ", "<br>• ").replace("\n", "<br>")
            st.markdown("<h4 style='color:#336699;'>Summary</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            st.markdown(f"<p style='font-size:17px;'>{formatted_summary}</p>", unsafe_allow_html=True)

            if experience_blocks:
                st.markdown("<h4 style='color:#336699;'>Experience</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for idx, exp_block in enumerate(experience_blocks):
                    lines_exp = [l for l in exp_block.strip().split("\n") if l.strip()]
                    if not lines_exp:
                        continue
                    heading = lines_exp[0]
                    m = re.match(r"[A-Z]\.\s*(.+?)\s*\((.+?)\)\s*$", heading)
                    if m:
                        company, duration = m.group(1).strip(), m.group(2).strip()
                    else:
                        m2 = re.match(r"[A-Z]\.\s*(.+?)\s*[\u2014\u2013-]+\s*(.+)$", heading)
                        if m2:
                            company, duration = m2.group(1).strip(), m2.group(2).strip()
                        else:
                            company  = re.sub(r"^[A-Z]\.\s*", "", heading).strip()
                            duration = ""
                    role = ""
                    bullet_lines = []
                    for line in lines_exp[1:]:
                        stripped = line.strip().lstrip("\u2022\u00b7*-\u2013\u2014 ").strip()
                        if not role and len(stripped) <= 60 and not re.search(r"\d+%|\d+ [a-z]", stripped):
                            role = stripped
                        else:
                            bullet_lines.append(line.strip())
                    if not role:
                        ss_entries = st.session_state.get("experience_entries", [])
                        role = ss_entries[idx].get("title", "") if idx < len(ss_entries) else ""
                    formatted_exp = "".join(
                        f"<div style='margin-left:12px;margin-bottom:4px;'>• {l.lstrip('\u2022\u00b7*-\u2013\u2014 ').strip()}</div>"
                        for l in bullet_lines if l.strip()
                    )
                    st.markdown(
                        f"<div style='margin-bottom:15px;padding:10px;border-radius:8px;border-left:3px solid #336699;'>"
                        f"<div style='display:flex;justify-content:space-between;flex-wrap:wrap;gap:4px;'>"
                        f"<b style='font-size:15px;'>🏢 {company}</b>"
                        f"<span style='color:gray;font-size:13px;'>📆 {duration}</span></div>"
                        f"<div style='font-size:14px;margin-top:3px;'>💼 <i>{role}</i></div>"
                        f"<div style='font-size:14px;margin-top:6px;'>{formatted_exp}</div></div>",
                        unsafe_allow_html=True
                    )

            edu_to_show = [e for e in st.session_state.education_entries if e.get("institution") or e.get("degree")]
            if edu_to_show:
                st.markdown("<h4 style='color:#336699;'>🎓 Education</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for edu in edu_to_show:
                    degree_val = edu.get("degree", "")
                    if isinstance(degree_val, list):
                        degree_val = ", ".join(degree_val)
                    st.markdown(
                        f"<div style='margin-bottom:15px;padding:10px 15px;border-radius:8px;border-left:3px solid #336699;'>"
                        f"<div style='display:flex;justify-content:space-between;font-size:16px;font-weight:bold;flex-wrap:wrap;gap:4px;'>"
                        f"<span>🏫 {edu.get('institution','')}</span>"
                        f"<span style='color:gray;font-size:13px;'>📅 {edu.get('year','')}</span></div>"
                        f"<div style='font-size:14px;'>🎓 <i>{degree_val}</i></div>"
                        f"<div style='font-size:14px;color:#555;'>📄 {edu.get('details','')}</div></div>",
                        unsafe_allow_html=True
                    )

            if projects_blocks:
                st.markdown("<h4 style='color:#336699;'>Projects</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for idx, proj_block in enumerate(projects_blocks):
                    plines = proj_block.strip().split("\n")
                    label  = chr(65 + idx)
                    ai_title = re.sub(r"^[A-Z]\.\s*", "", plines[0]).strip() if plines else ""
                    ai_tech = ai_duration = ""
                    desc_lines = []
                    in_desc = False
                    for line in plines[1:]:
                        stripped = line.strip()
                        tl = stripped.lstrip("\u2022\u00b7*-\u2013\u2014 ").strip()
                        if re.match(r"Tech\s*Stack\s*:", tl, re.I):
                            ai_tech = re.sub(r"(?i)^Tech\s*Stack\s*:\s*", "", tl).strip()
                        elif re.match(r"Duration\s*:", tl, re.I):
                            ai_duration = re.sub(r"(?i)^Duration\s*:\s*", "", tl).strip()
                        elif re.match(r"Description\s*:", tl, re.I):
                            in_desc = True
                            remainder = re.sub(r"(?i)^Description\s*:\s*", "", tl).strip()
                            if remainder:
                                desc_lines.append(remainder)
                        elif in_desc and stripped:
                            desc_lines.append(stripped)
                    ss_proj = st.session_state.project_entries[idx] if idx < len(st.session_state.project_entries) else {}
                    final_title    = ai_title    or ss_proj.get("title", "")
                    final_tech     = ai_tech     or ss_proj.get("tech", "")
                    final_duration = ai_duration or ss_proj.get("duration", "")
                    formatted_proj = "".join(
                        f"<div style='margin-left:12px;margin-bottom:4px;'>• {dl.lstrip('\u2022\u00b7*-\u2013\u2014 ').strip()}</div>"
                        for dl in desc_lines if dl.strip()
                    )
                    st.markdown(
                        f"<div style='margin-bottom:15px;padding:10px;border-radius:8px;border-left:3px solid #336699;'>"
                        f"<strong style='font-size:16px;'>📌 <span style='color:#444;'>{label}. </span>{final_title}</strong><br>"
                        f"<span style='font-size:13px;color:#555;'>🛠️ <strong>Tech Stack:</strong> {final_tech}</span><br>"
                        f"<span style='font-size:13px;color:#555;'>⏳ <strong>Duration:</strong> {final_duration}</span><br>"
                        f"<div style='font-size:14px;margin-top:6px;'>{formatted_proj}</div></div>",
                        unsafe_allow_html=True
                    )

            # ── Certificate rendering — strip AI-invented "Unknown" placeholders ──
            # Build a name→date lookup from session state so we can fall back to
            # the user's real data if the AI mangles or omits it.
            _ss_cert_lookup = {
                c.get("name", "").strip(): c.get("duration", "").strip()
                for c in st.session_state.get("certificate_links", [])
                if c.get("name", "").strip()
            }

            _cert_lines_raw = [c.strip() for c in certificates_list.split("\n") if c.strip()] if certificates_list else []

            # Fall back to session state when AI produced nothing or only garbage
            if not _cert_lines_raw and _ss_cert_lookup:
                _cert_lines_raw = [
                    f"{name} ({date})" if date else name
                    for name, date in _ss_cert_lookup.items()
                ]

            if _cert_lines_raw:
                st.markdown("<h4 style='color:#336699;'>📜 Certificates</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for _cert_line in _cert_lines_raw:
                    # Remove any " - Unknown" or "Unknown - " the AI injected
                    _cleaned = re.sub(r'\s*-\s*Unknown\b', '', _cert_line, flags=re.IGNORECASE).strip()
                    _cleaned = re.sub(r'\bUnknown\s*-\s*', '', _cleaned, flags=re.IGNORECASE).strip()
                    _cleaned = re.sub(r'\bUnknown\b', '', _cleaned, flags=re.IGNORECASE).strip(" -–—").strip()

                    # If AI dropped the date, re-inject it from session state
                    for _ss_name, _ss_date in _ss_cert_lookup.items():
                        if _ss_name.lower() in _cleaned.lower() and _ss_date and _ss_date not in _cleaned:
                            _cleaned = f"{_cleaned} ({_ss_date})"
                            break

                    if _cleaned:
                        st.markdown(f"<div style='margin-left:10px;margin-bottom:4px;'>• {_cleaned}</div>", unsafe_allow_html=True)

            if st.session_state.project_links:
                st.markdown("<h4 style='color:#336699;'>Project Links</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for i, link in enumerate(st.session_state.project_links):
                    st.markdown(f"[🔗 Project {i+1}]({link})", unsafe_allow_html=True)

    # Generate HTML content based on selected template — only on submit, stored in session_state
    if submitted:
        with st.spinner("⚙️ Generating your resume... please wait"):
            # Render selected resume template via the registry dispatcher (resume_builder.py)
            html_content = render_resume(selected_template, st.session_state, profile_img_html)

            # Store the generated content and invalidate cached PDF so it's recomputed fresh
            # NOTE: Use direct assignment instead of .pop() — .pop() on an existing key
            # triggers an extra Streamlit rerun which causes visible page blinking.
            st.session_state["generated_html"] = html_content
            st.session_state["pdf_resume_bytes"] = None   # invalidate cache without extra rerun
            st.session_state["show_template_preview"] = False
        st.session_state.pop("_resume_generating", None)

with tab2:
    # ==========================
    # 📥 Resume Download Header
    # ==========================
    if "generated_html" in st.session_state:
        st.markdown(
            """
            <div style='text-align: center; margin-top: 20px; margin-bottom: 30px;'>
                <h2 style='color: #2f4f6f; font-family: Arial, sans-serif; font-size: 24px;'>
                    📥 Download Your Resume
                </h2>
                <p style="color:#555; font-size:14px;">
                    Choose your preferred format below
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )

        # Cache PDF bytes in session_state to avoid expensive recomputation on every rerun
        if not st.session_state.get("pdf_resume_bytes"):
            st.session_state["pdf_resume_bytes"] = html_to_pdf_bytes(
                st.session_state["generated_html"]
            ).read()

        col1, spacer, col2 = st.columns([1, 0.15, 0.85])

        # HTML Resume Download Button
        with col1:
            html_bytes = st.session_state["generated_html"].encode("utf-8")
            html_file = BytesIO(html_bytes)

            st.download_button(
                label="⬇️ Download as Template",
                data=html_file,
                file_name=f"{st.session_state['name'].replace(' ', '_')}_Resume.html",
                mime="text/html",
                key="download_resume_html"
            )

        # Preview Template Button — smart toggle: spinner only when opening, instant when closing
        with col2:
            is_previewing = st.session_state.get("show_template_preview", False)
            if st.button("👁️ Preview Template", key="preview_template_btn"):
                if not is_previewing:
                    # Opening — show spinner since we're loading the iframe
                    with st.spinner("Loading template preview..."):
                        time.sleep(2)
                        st.session_state["show_template_preview"] = True
                else:
                    # Closing — instant, no spinner
                    st.session_state["show_template_preview"] = False

        # Show/hide the template preview iframe
        if st.session_state.get("show_template_preview", False):
            import streamlit.components.v1 as components
            st.markdown(
                "<p style='color:#555; font-size:13px; margin-top:8px;'>"
                "📄 Template Preview (scroll to explore):</p>",
                unsafe_allow_html=True,
            )
            components.html(
                st.session_state["generated_html"],
                height=600,
                scrolling=True,
            )

        # PDF Resume Download Button — use cached bytes
        pdf_resume_bytes = BytesIO(st.session_state["pdf_resume_bytes"])
        
        # ✅ Extra Help Note
        st.markdown("""
        ✅ After downloading your HTML resume, you can 
        <a href="https://www.sejda.com/html-to-pdf" target="_blank" style="color:#2f4f6f; text-decoration:none;">
        convert it to PDF using Sejda's free online tool</a>.
        """, unsafe_allow_html=True)

        # ==========================
        # 📩 Cover Letter Expander
        # ==========================
        with st.expander("📩 Generate Cover Letter from This Resume"):
            generate_cover_letter_from_resume_builder()

        # ==========================
        # ✉️ Generated Cover Letter Downloads (NO PREVIEW HERE)
        # ==========================
        if "cover_letter" in st.session_state:
            st.markdown(
                """
                <div style="margin-top: 30px; margin-bottom: 20px;">
                    <h3 style="color: #003366;">✉️ Generated Cover Letter</h3>
                    <p style="color:#555; font-size:14px;">
                        You can download your generated cover letter in multiple formats.
                    </p>
                </div>
                """,
                unsafe_allow_html=True
            )

            # ✅ Use already-rendered HTML from session (don't show again)
            styled_cover_letter = st.session_state.get("cover_letter_html", "")

            # ✅ Generate PDF from styled HTML
            pdf_file = html_to_pdf_bytes(styled_cover_letter)

            # ✅ DOCX Generator (preserves line breaks)
            def create_docx_from_text(text, filename="cover_letter.docx"):
                from docx import Document
                bio = BytesIO()
                doc = Document()
                doc.add_heading("Cover Letter", 0)

                for line in text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)
                    else:
                        doc.add_paragraph("")  # preserve empty lines

                doc.save(bio)
                bio.seek(0)
                return bio

            # ==========================
            # 📥 Cover Letter Download Buttons
            # ==========================
            st.markdown("""
            <div style="margin-top: 25px; margin-bottom: 15px;">
                <strong>⬇️ Download Your Cover Letter:</strong>
            </div>
            """, unsafe_allow_html=True)

            col1,col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📥 Download Cover Letter (.docx)",
                    data=create_docx_from_text(st.session_state["cover_letter"]),
                    file_name=f"{st.session_state['name'].replace(' ', '_')}_Cover_Letter.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_coverletter_docx"
                )
            
            with col2:
                st.download_button(
                    label="📥 Download Cover Letter (Template)",
                    data=styled_cover_letter.encode("utf-8"),
                    file_name=f"{st.session_state['name'].replace(' ', '_')}_Cover_Letter.html",
                    mime="text/html",
                    key="download_coverletter_html"
                )

            # ✅ Helper note
            st.markdown("""
            ✅ If the HTML cover letter doesn't display properly, you can 
            <a href="https://www.sejda.com/html-to-pdf" target="_blank" style="color:#2f4f6f; text-decoration:none;">
            convert it to PDF using Sejda's free online tool</a>.
            """, unsafe_allow_html=True)
