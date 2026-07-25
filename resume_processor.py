import os
os.environ["STREAMLIT_WATCHDOG"] = "false"
import json
import random
import string
import re
import asyncio
import io
from concurrent.futures import ThreadPoolExecutor, as_completed
import urllib.parse
import base64
from io import BytesIO
from collections import Counter
from datetime import datetime
import time

import streamlit as st
import streamlit.components.v1 as components
from base64 import b64encode
import requests
import fitz
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import altair as alt
from PIL import Image
from pdf2image import convert_from_path
from dotenv import load_dotenv
from nltk.stem import WordNetLemmatizer
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from xhtml2pdf import pisa
from pydantic import BaseModel
from streamlit_pdf_viewer import pdf_viewer
import torch
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from llm_manager import (
    call_llm, load_groq_api_keys, get_healthy_keys, increment_key_usage,
    mark_key_failure, _mem_record_failure, _mem_clear_failure,
    _mem_increment_usage, _async_mark_failure, _async_increment_usage,
    _async_clear_failure,
)
from db_manager import (
    db_manager, insert_candidate, get_top_domains_by_score,
    get_database_stats, detect_domain_from_title_and_description,
    get_domain_similarity
)
from user_login import (
    create_user_table, add_user, complete_registration, verify_user,
    get_logins_today, get_total_registered_users, log_user_action,
    username_exists, email_exists, is_valid_email, get_all_user_logs,
    generate_otp, send_email_otp,
    get_user_by_email, update_password_by_email, is_strong_password,
    domain_has_mx_record, send_login_link, verify_login_token,
    cleanup_expired_login_tokens, check_and_gate_feature,
    record_feature_usage, get_usage_count_last_hour, check_brute_force,
)

# ── resume_processor.py ─────────────────────────────────────────────────────
def get_easyocr_reader():
    import easyocr
    return easyocr.Reader(["en"], gpu=torch.cuda.is_available())

@st.cache_data(show_spinner=False)
def ensure_nltk():
    import nltk
    nltk.download('wordnet', quiet=True)
    return WordNetLemmatizer()

lemmatizer = ensure_nltk()
# EasyOCR reader is intentionally NOT loaded here at module startup.
# It is loaded lazily inside extract_text_from_images() only when a
# partially-scanned PDF is encountered. This prevents the ~500 MB model
# from being pulled on every cold start and avoids OOM on Streamlit Cloud.

def generate_docx(text, filename="bias_free_resume.docx"):
    doc = Document()

    # ── Page margins (standard resume: 1 inch all sides) ──
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

    # ── Document title heading ──
    title = doc.add_heading('Bias-Free Resume', 0)
    title.alignment = 1  # center
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0x2F, 0x4F, 0x6F)
    title_run.font.size = Pt(18)

    doc.add_paragraph()  # spacer

    # ── Process text: detect section headers and bullet points ──
    lines = text.strip().split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Section headers (emoji + CAPS or all-caps lines)
        if (stripped.isupper() and len(stripped) > 3) or \
           any(stripped.startswith(e) for e in ['🏷️','📞','📧','📍','🔗','🌐','✍️','🛠️','💼','🧑‍💼','📂','🎓','🏫','🤝','🌟','🎯']):
            p = doc.add_heading(stripped, level=2)
            p.runs[0].font.color.rgb = RGBColor(0x2F, 0x4F, 0x6F)
            p.runs[0].font.size = Pt(12)
            continue

        # Bullet points
        if stripped.startswith(('•', '-', '*')):
            content = stripped.lstrip('•-* ').strip()
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(content)
            run.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(3)
            continue

        # Regular paragraph
        p = doc.add_paragraph(stripped)
        p.runs[0].font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(4)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Extract text from PDF
def _extract_page_text_smart(page) -> str:
    """
    Extract text from a single fitz page in correct reading order,
    handling both simple single-column and complex multi-column / graphic
    resume layouts (like sidebar designs, blue-header templates, etc.).

    Strategy:
      1. Pull raw text blocks with full bounding-box coordinates.
      2. Detect whether the page has a multi-column layout by checking
         whether meaningful content exists in both the left (<45%) and
         right (>52%) horizontal zones.
      3. Single-column  → sort blocks top-to-bottom (y0), then left-to-right.
      4. Multi-column   → split blocks into left/right columns,
                          sort each column top-to-bottom independently,
                          then concatenate left column first (the main body
                          on most sidebar resume designs sits on the right,
                          but the name/header almost always spans full width
                          or sits at the very top — so we sort by y0 first
                          for the header region, then by column).
    This ensures the candidate name — which is always the topmost block —
    is the very first text we see regardless of layout.
    """
    blocks = page.get_text("blocks")  # (x0, y0, x1, y1, text, block_no, block_type)
    page_width = page.rect.width

    # Filter: only text blocks (block_type == 0) with real content
    text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]

    if not text_blocks:
        return ""

    # ── Detect multi-column layout ────────────────────────────────────────
    x_starts = [b[0] for b in text_blocks if len(b[4].strip()) > 10]
    left_zone  = [x for x in x_starts if x < page_width * 0.45]
    right_zone = [x for x in x_starts if x > page_width * 0.52]
    is_multicolumn = len(left_zone) >= 3 and len(right_zone) >= 3

    if not is_multicolumn:
        # ── Single-column: simple top-to-bottom sort ──────────────────────
        sorted_blocks = sorted(text_blocks, key=lambda b: (round(b[1] / 10) * 10, b[0]))
        return "\n".join(b[4].strip() for b in sorted_blocks)

    # ── Multi-column: split into header + left + right zones ─────────────
    # Blocks in the top 15% of page height are "header" — name, title, etc.
    # They are sorted purely by y0 so the name always comes first.
    page_height  = page.rect.height
    header_zone  = page_height * 0.15

    header_blocks = [b for b in text_blocks if b[1] < header_zone]
    body_blocks   = [b for b in text_blocks if b[1] >= header_zone]

    # Sort header top-to-bottom
    header_sorted = sorted(header_blocks, key=lambda b: (b[1], b[0]))

    # Split body into left / right columns and sort each top-to-bottom
    left_blocks  = sorted(
        [b for b in body_blocks if b[0] < page_width * 0.48],
        key=lambda b: b[1]
    )
    right_blocks = sorted(
        [b for b in body_blocks if b[0] >= page_width * 0.48],
        key=lambda b: b[1]
    )

    # Concatenate: header → left column → right column
    all_sorted = header_sorted + left_blocks + right_blocks
    return "\n".join(b[4].strip() for b in all_sorted)


# ============================================================
# 🔍 Scanned PDF Detection — Multi-Signal Engine
# ============================================================

# Sentinel returned by safe_extract_text when file is image-only
_SCANNED_SENTINEL = "__SCANNED_PDF__"

def _classify_pdf(file_path: str) -> dict:
    """
    Multi-signal classifier that determines whether a PDF is text-based
    or image/scan-based. Returns a dict with keys:
      - is_scanned     : bool   — True if PDF is image-only
      - confidence     : str    — 'definite' | 'likely' | 'uncertain'
      - page_count     : int
      - text_pages     : int    — pages with meaningful text (≥40 words)
      - image_only_pages: int   — pages with embedded images but no text
      - total_words    : int    — total word count across all pages
      - file_size_mb   : float
      - signals        : list   — human-readable evidence list

    Detection uses FOUR independent signals so no single edge case can
    misclassify a real text-based resume as scanned:

    Signal 1 — Text density per page via PyMuPDF get_text()
      A real resume page has ≥40 recognisable words. Fewer means the
      page either has no text layer or the text is in images.

    Signal 2 — Image-block ratio via get_text("dict") block analysis
      Counts how many blocks on a page are image-type (type == 1) vs
      text-type (type == 0). Scanned PDFs are 100% image blocks.

    Signal 3 — Character confidence heuristic
      OCR-produced or truly embedded text has normal char distribution.
      Random symbol dumps (mojibake from scanned-then-OCR'd PDFs) have
      high ratio of non-ASCII or non-printable chars → flagged.

    Signal 4 — Embedded font check
      Text-based PDFs always embed at least one font. A PDF with zero
      embedded fonts across all pages has no real text layer.
    """
    signals = []
    result = {
        "is_scanned": False,
        "confidence": "uncertain",
        "page_count": 0,
        "text_pages": 0,
        "image_only_pages": 0,
        "total_words": 0,
        "file_size_mb": 0.0,
        "signals": signals,
    }

    try:
        result["file_size_mb"] = round(os.path.getsize(file_path) / (1024 * 1024), 2)
        doc = fitz.open(file_path)
        page_count = doc.page_count
        result["page_count"] = page_count

        text_pages       = 0
        image_only_pages = 0
        total_words      = 0
        total_fonts      = 0
        high_noise_pages = 0

        for page in doc:
            # ── Signal 1: text density ────────────────────────────────────
            raw_text  = page.get_text("text") or ""
            words     = [w for w in raw_text.split() if len(w) > 1]
            word_count = len(words)
            total_words += word_count

            if word_count >= 40:
                text_pages += 1
            else:
                # ── Signal 2: image-block ratio ───────────────────────────
                page_dict   = page.get_text("dict")
                all_blocks  = page_dict.get("blocks", [])
                img_blocks  = [b for b in all_blocks if b.get("type") == 1]
                text_blocks = [b for b in all_blocks if b.get("type") == 0
                               and any(
                                   span.get("text", "").strip()
                                   for line in b.get("lines", [])
                                   for span in line.get("spans", [])
                               )]
                if img_blocks and not text_blocks:
                    image_only_pages += 1
                    signals.append(f"Page {page.number + 1}: image-only block (no text layer)")

            # ── Signal 3: character noise ratio ──────────────────────────
            if word_count > 5:
                total_chars  = len(raw_text)
                noise_chars  = sum(
                    1 for c in raw_text
                    if ord(c) > 127 or (ord(c) < 32 and c not in "\n\r\t")
                )
                noise_ratio  = noise_chars / max(total_chars, 1)
                if noise_ratio > 0.15:
                    high_noise_pages += 1

            # ── Signal 4: embedded font check ─────────────────────────────
            font_list = page.get_fonts(full=False)
            total_fonts += len(font_list)

        doc.close()

        result["text_pages"]        = text_pages
        result["image_only_pages"]  = image_only_pages
        result["total_words"]       = total_words

        # ── Decision logic (multi-signal voting) ─────────────────────────
        text_page_ratio  = text_pages / max(page_count, 1)
        image_page_ratio = image_only_pages / max(page_count, 1)

        # DEFINITE scanned: zero readable text pages + image blocks present
        if text_pages == 0 and image_only_pages > 0:
            result["is_scanned"]  = True
            result["confidence"]  = "definite"
            signals.append(f"Zero text pages found; {image_only_pages}/{page_count} page(s) are pure image blocks")

        # DEFINITE scanned: no text at all and no embedded fonts
        elif total_words < 30 and total_fonts == 0:
            result["is_scanned"]  = True
            result["confidence"]  = "definite"
            signals.append(f"Only {total_words} words found and zero embedded fonts — no text layer")

        # LIKELY scanned: very low text coverage across majority of pages
        elif text_page_ratio < 0.30 and image_page_ratio >= 0.50:
            result["is_scanned"]  = True
            result["confidence"]  = "likely"
            signals.append(
                f"Only {text_pages}/{page_count} pages have ≥40 words; "
                f"{image_only_pages} page(s) are image-only"
            )

        # UNCERTAIN: some text but very low total word count with images present
        elif total_words < 80 and image_only_pages > 0:
            result["is_scanned"]  = True
            result["confidence"]  = "uncertain"
            signals.append(
                f"Low word count ({total_words} total) with image-only pages — "
                "likely partially scanned"
            )

        else:
            result["is_scanned"]  = False
            result["confidence"]  = "definite"
            signals.append(
                f"Text-based PDF confirmed: {text_pages}/{page_count} pages readable, "
                f"{total_words} total words"
            )

    except Exception as e:
        signals.append(f"Classifier error: {e}")
        result["confidence"] = "uncertain"

    return result


def _render_scanned_rejection_card(filename: str, classification: dict, container=None):
    """
    Renders an industry-standard rejection card for scanned/image PDFs.
    Matches the app's existing glassmorphism dark theme exactly.
    Shows classification evidence so the user understands why their file failed.
    """
    # SVG dot icon per confidence tier — replaces colored circle emojis
    _dot_svg = {
        "definite":  '<svg width="9" height="9" viewBox="0 0 9 9" fill="none" xmlns="http://www.w3.org/2000/svg" style="flex-shrink:0;margin-top:1px;"><circle cx="4.5" cy="4.5" r="4.5" fill="#fb7185"/></svg>',
        "likely":    '<svg width="9" height="9" viewBox="0 0 9 9" fill="none" xmlns="http://www.w3.org/2000/svg" style="flex-shrink:0;margin-top:1px;"><circle cx="4.5" cy="4.5" r="4.5" fill="#fb923c"/></svg>',
        "uncertain": '<svg width="9" height="9" viewBox="0 0 9 9" fill="none" xmlns="http://www.w3.org/2000/svg" style="flex-shrink:0;margin-top:1px;"><circle cx="4.5" cy="4.5" r="4.5" fill="#fbbf24"/></svg>',
    }

    confidence_label = {
        "definite":  ("Unreadable — Image-Only PDF",        "#fb7185", "rgba(251,113,133,0.18)", "rgba(251,113,133,0.35)"),
        "likely":    ("Likely Scanned — Low Text Coverage",  "#fb923c", "rgba(251,146, 60,0.18)", "rgba(251,146, 60,0.35)"),
        "uncertain": ("Partially Scanned — Low Text Quality","#fbbf24", "rgba(251,191, 36,0.18)", "rgba(251,191, 36,0.35)"),
    }.get(
        classification["confidence"],
        ("Scan Quality Issue", "#fbbf24", "rgba(251,191,36,0.18)", "rgba(251,191,36,0.35)")
    )

    dot_svg     = _dot_svg.get(classification["confidence"], _dot_svg["uncertain"])
    badge_text, text_color, bg_color, border_color = confidence_label
    page_count  = classification.get("page_count", "?")
    total_words = classification.get("total_words", 0)
    file_size   = classification.get("file_size_mb", 0)
    signals     = classification.get("signals", [])

    # Build evidence bullet list — SVG bullet marker instead of default disc
    _bullet = '<svg width="6" height="6" viewBox="0 0 6 6" fill="none" xmlns="http://www.w3.org/2000/svg" style="flex-shrink:0;margin-top:5px;"><circle cx="3" cy="3" r="3" fill="#475569"/></svg>'
    evidence_html = "".join(
        f"<li style='display:flex;align-items:flex-start;gap:8px;margin-bottom:6px;list-style:none;'>"
        f"{_bullet}"
        f"<span>{s}</span></li>"
        for s in signals[:4]
    )

    # SVG icons for stat pills
    _svg_pages = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>'
    _svg_text  = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;"><line x1="17" y1="10" x2="3" y2="10"/><line x1="21" y1="6" x2="3" y2="6"/><line x1="21" y1="14" x2="3" y2="14"/><line x1="17" y1="18" x2="3" y2="18"/></svg>'
    _svg_disk  = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;"><ellipse cx="12" cy="5" rx="9" ry="3"/><path d="M21 12c0 1.66-4 3-9 3s-9-1.34-9-3"/><path d="M3 5v14c0 1.66 4 3 9 3s9-1.34 9-3V5"/></svg>'

    # SVG icon for header — image/scan concept (image frame with slash)
    _svg_header = '<svg width="28" height="28" viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="3" width="18" height="18" rx="2"/><circle cx="8.5" cy="8.5" r="1.5"/><polyline points="21 15 16 10 5 21"/><line x1="2" y1="2" x2="22" y2="22" stroke="{c}" stroke-width="1.5"/></svg>'.replace("{c}", text_color)

    # SVG icons for "How to Fix" section labels
    _svg_fix_word  = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/><polyline points="10 9 9 9 8 9"/></svg>'
    _svg_fix_export= '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><path d="M21 15v4a2 2 0 01-2 2H5a2 2 0 01-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>'
    _svg_fix_scan  = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><line x1="2" y1="2" x2="22" y2="22"/><path d="M10.58 10.58A2 2 0 0013 13"/><path d="M17.94 17.94A10 10 0 013.34 7.34"/></svg>'
    _svg_fix_acro  = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><circle cx="11" cy="11" r="8"/><line x1="21" y1="21" x2="16.65" y2="16.65"/></svg>'

    fix_items = [
        (_svg_fix_word,   "Open your resume in <strong>Microsoft Word or Google Docs</strong>"),
        (_svg_fix_export, "Export / Download as <strong>PDF</strong> — this creates a text-layer PDF"),
        (_svg_fix_scan,   "Avoid scanning a printed resume — use the original digital file"),
        (_svg_fix_acro,   "If you only have a scanned copy, use <strong>Adobe Acrobat → OCR → Save as PDF</strong>"),
    ]
    fix_html = "".join(
        f"<li style='display:flex;align-items:flex-start;gap:8px;margin-bottom:8px;list-style:none;'>"
        f"{icon}<span>{text}</span></li>"
        for icon, text in fix_items
    )

    # SVG icons for section header labels (evidence / how to fix)
    _svg_label_evidence = '<svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#64748b" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="8"/><line x1="21" y1="21" x2="16.65" y2="16.65"/></svg>'
    _svg_label_fix      = '<svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>'

    card_html = (
        f'<div style="background:linear-gradient(135deg,{bg_color} 0%,rgba(0,0,0,0) 100%);border:1px solid {border_color};border-radius:16px;padding:22px 24px;margin:14px 0;backdrop-filter:blur(20px);-webkit-backdrop-filter:blur(20px);box-shadow:0 8px 32px rgba(0,0,0,0.3),inset 0 1px 0 rgba(255,255,255,0.06);font-family:-apple-system,BlinkMacSystemFont,sans-serif;position:relative;overflow:hidden;">'
        f'<div style="position:absolute;top:0;left:0;right:0;height:3px;background:linear-gradient(90deg,transparent,{text_color},transparent);opacity:0.6;"></div>'
        f'<div style="display:flex;align-items:flex-start;gap:14px;margin-bottom:16px;">'
        f'<div style="width:44px;height:44px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:10px;display:flex;align-items:center;justify-content:center;flex-shrink:0;">{_svg_header}</div>'
        f'<div style="flex:1;">'
        f'<div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;color:{text_color};margin-bottom:4px;">{dot_svg}{badge_text}</div>'
        f'<div style="font-size:1rem;font-weight:600;color:#f0f4f8;word-break:break-all;">{filename}</div>'
        f'</div></div>'
        f'<div style="display:flex;gap:10px;flex-wrap:wrap;margin-bottom:16px;">'
        f'<div style="display:flex;align-items:center;gap:6px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:8px;padding:6px 12px;font-size:0.78rem;color:#94a3b8;">{_svg_pages} {page_count} page{"s" if page_count != 1 else ""}</div>'
        f'<div style="display:flex;align-items:center;gap:6px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:8px;padding:6px 12px;font-size:0.78rem;color:#94a3b8;">{_svg_text} {total_words} words detected</div>'
        f'<div style="display:flex;align-items:center;gap:6px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:8px;padding:6px 12px;font-size:0.78rem;color:#94a3b8;">{_svg_disk} {file_size} MB</div>'
        f'</div>'
        f'<div style="background:rgba(0,0,0,0.2);border:1px solid rgba(255,255,255,0.06);border-radius:10px;padding:12px 16px;margin-bottom:16px;">'
        f'<div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;font-weight:700;letter-spacing:0.07em;text-transform:uppercase;color:#64748b;margin-bottom:10px;">{_svg_label_evidence} Detection Evidence</div>'
        f'<ul style="margin:0;padding:0;color:#94a3b8;font-size:0.82rem;line-height:1.6;">{evidence_html}</ul>'
        f'</div>'
        f'<div style="background:rgba(56,189,248,0.07);border:1px solid rgba(56,189,248,0.18);border-radius:10px;padding:12px 16px;">'
        f'<div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;font-weight:700;letter-spacing:0.07em;text-transform:uppercase;color:#38bdf8;margin-bottom:10px;">{_svg_label_fix} How to Fix This</div>'
        f'<ul style="margin:0;padding:0;color:#7dd3fc;font-size:0.82rem;line-height:1.8;">{fix_html}</ul>'
        f'</div>'
        f'</div>'
    )
    _render_target = container if container is not None else st
    _render_target.markdown(card_html, unsafe_allow_html=True)


def extract_text_from_pdf(file_path: str):
    """
    Extracts text from a PDF using PyMuPDF smart extraction.
    Returns list of page text strings, or empty list if nothing found.
    Does NOT fall through to OCR — scanned detection is handled upstream
    by _classify_pdf() + safe_extract_text().
    """
    try:
        doc = fitz.open(file_path)
        text_list = []
        for page in doc:
            page_text = _extract_page_text_smart(page)
            if page_text.strip():
                text_list.append(page_text)
        doc.close()
        return text_list
    except Exception as e:
        st.error(f"⚠ Error extracting text: {e}")
        return []


def extract_text_from_images(pdf_path):
    """
    OCR fallback using EasyOCR — only called when explicitly requested
    and classification confidence is 'uncertain' (partial scan).
    EasyOCR reader is loaded lazily here, not at module startup.
    """
    try:
        _reader = get_easyocr_reader()
        images = convert_from_path(pdf_path, dpi=120, first_page=1, last_page=3)
        results = []
        for img in images:
            try:
                page_text = "\n".join(_reader.readtext(np.array(img), detail=0))
                if page_text.strip():
                    results.append(page_text)
            except Exception:
                continue
        return results
    except Exception as e:
        st.error(f"⚠ Error extracting from image: {e}")
        return []



# ============================================================
# 🌐 Non-English Resume Detection & Rejection Card
# ============================================================

# Sentinel returned when resume is non-English
_NON_ENGLISH_SENTINEL = "__NON_ENGLISH_RESUME__"

def _detect_non_english(text: str) -> dict:
    """
    Detects whether resume text is non-English using character-level analysis.
    Strategy:
      - Count ASCII alphabetic characters vs total alphabetic characters
      - Real English resumes are overwhelmingly ASCII (>85%)
      - Also check for high-frequency non-Latin Unicode blocks (Devanagari,
        Arabic, CJK, Cyrillic, etc.) as a secondary signal
    Returns dict with: is_non_english, confidence, ascii_ratio, script_detected
    """
    if not text or len(text.strip()) < 50:
        return {"is_non_english": False, "confidence": "uncertain",
                "ascii_ratio": 1.0, "script_detected": "unknown"}

    alpha_chars = [c for c in text if c.isalpha()]
    if not alpha_chars:
        return {"is_non_english": False, "confidence": "uncertain",
                "ascii_ratio": 1.0, "script_detected": "unknown"}

    ascii_alpha  = sum(1 for c in alpha_chars if ord(c) < 128)
    ascii_ratio  = ascii_alpha / len(alpha_chars)

    # Detect dominant non-Latin script
    script_detected = "Unknown Script"
    script_ranges = [
        ((0x0900, 0x097F), "Devanagari (Hindi/Marathi/Sanskrit)"),
        ((0x0980, 0x09FF), "Bengali"),
        ((0x0A00, 0x0A7F), "Gurmukhi (Punjabi)"),
        ((0x0A80, 0x0AFF), "Gujarati"),
        ((0x0B00, 0x0B7F), "Odia"),
        ((0x0B80, 0x0BFF), "Tamil"),
        ((0x0C00, 0x0C7F), "Telugu"),
        ((0x0C80, 0x0CFF), "Kannada"),
        ((0x0D00, 0x0D7F), "Malayalam"),
        ((0x0600, 0x06FF), "Arabic"),
        ((0x0400, 0x04FF), "Cyrillic (Russian/Ukrainian)"),
        ((0x4E00, 0x9FFF), "Chinese (CJK)"),
        ((0x3040, 0x30FF), "Japanese (Hiragana/Katakana)"),
        ((0xAC00, 0xD7AF), "Korean (Hangul)"),
        ((0x0E00, 0x0E7F), "Thai"),
    ]
    for (start, end), name in script_ranges:
        count = sum(1 for c in text if start <= ord(c) <= end)
        if count > 20:
            script_detected = name
            break

    is_non_english = ascii_ratio < 0.70
    confidence = "definite" if ascii_ratio < 0.50 else "likely" if ascii_ratio < 0.70 else "uncertain"

    return {
        "is_non_english": is_non_english,
        "confidence":     confidence,
        "ascii_ratio":    round(ascii_ratio, 3),
        "script_detected": script_detected,
    }


def _render_non_english_card(filename: str, detection: dict, container=None):
    """
    Renders a rejection card for non-English resumes.
    Matches the app's existing glassmorphism dark theme exactly.
    """
    confidence  = detection.get("confidence", "likely")
    ascii_ratio = detection.get("ascii_ratio", 0.0)
    script      = detection.get("script_detected", "Non-Latin Script")
    english_pct = round(ascii_ratio * 100)

    _dot_svg = {
        "definite": '<svg width="9" height="9" viewBox="0 0 9 9" fill="none" xmlns="http://www.w3.org/2000/svg" style="flex-shrink:0;margin-top:1px;"><circle cx="4.5" cy="4.5" r="4.5" fill="#fb7185"/></svg>',
        "likely":   '<svg width="9" height="9" viewBox="0 0 9 9" fill="none" xmlns="http://www.w3.org/2000/svg" style="flex-shrink:0;margin-top:1px;"><circle cx="4.5" cy="4.5" r="4.5" fill="#fb923c"/></svg>',
    }
    dot_svg = _dot_svg.get(confidence, _dot_svg["likely"])

    text_color   = "#fb7185" if confidence == "definite" else "#fb923c"
    bg_color     = "rgba(251,113,133,0.18)" if confidence == "definite" else "rgba(251,146,60,0.18)"
    border_color = "rgba(251,113,133,0.35)" if confidence == "definite" else "rgba(251,146,60,0.35)"
    badge_text   = "Non-English Resume Detected"

    _svg_header = f'''<svg width="28" height="28" viewBox="0 0 24 24" fill="none" stroke="{text_color}" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><line x1="2" y1="12" x2="22" y2="12"/><path d="M12 2a15.3 15.3 0 010 20M12 2a15.3 15.3 0 000 20"/></svg>'''

    _svg_lang  = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;"><circle cx="12" cy="12" r="10"/><line x1="2" y1="12" x2="22" y2="12"/><path d="M12 2a15.3 15.3 0 010 20M12 2a15.3 15.3 0 000 20"/></svg>'
    _svg_pct   = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;"><line x1="19" y1="5" x2="5" y2="19"/><circle cx="6.5" cy="6.5" r="2.5"/><circle cx="17.5" cy="17.5" r="2.5"/></svg>'

    # How to fix items
    _svg_translate = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><path d="M5 8l6 6"/><path d="M4 14l6-6 2-3"/><path d="M2 5h12"/><path d="M7 2h1"/><path d="M22 22l-5-10-5 10"/><path d="M14 18h6"/></svg>'
    _svg_word      = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/></svg>'
    _svg_tip       = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>'

    fix_items = [
        (_svg_translate, "Translate your resume to <strong>English</strong> using Google Translate or DeepL"),
        (_svg_word,      "Open the translated text in <strong>Microsoft Word or Google Docs</strong> and format it as a resume"),
        (_svg_tip,       "Re-export as PDF and upload again — English resumes get significantly better ATS scores"),
    ]
    fix_html = "".join(
        f"<li style='display:flex;align-items:flex-start;gap:8px;margin-bottom:8px;list-style:none;'>"
        f"{icon}<span>{text}</span></li>"
        for icon, text in fix_items
    )

    _svg_label_why = '<svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#64748b" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="8"/><line x1="21" y1="21" x2="16.65" y2="16.65"/></svg>'
    _svg_label_fix = '<svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>'

    card_html = (
        f'<div style="background:linear-gradient(135deg,{bg_color} 0%,rgba(0,0,0,0) 100%);border:1px solid {border_color};border-radius:16px;padding:22px 24px;margin:14px 0;backdrop-filter:blur(20px);-webkit-backdrop-filter:blur(20px);box-shadow:0 8px 32px rgba(0,0,0,0.3),inset 0 1px 0 rgba(255,255,255,0.06);font-family:-apple-system,BlinkMacSystemFont,sans-serif;position:relative;overflow:hidden;">'
        f'<div style="position:absolute;top:0;left:0;right:0;height:3px;background:linear-gradient(90deg,transparent,{text_color},transparent);opacity:0.6;"></div>'
        f'<div style="display:flex;align-items:flex-start;gap:14px;margin-bottom:16px;">'
        f'<div style="width:44px;height:44px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:10px;display:flex;align-items:center;justify-content:center;flex-shrink:0;">{_svg_header}</div>'
        f'<div style="flex:1;">'
        f'<div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;color:{text_color};margin-bottom:4px;">{dot_svg}{badge_text}</div>'
        f'<div style="font-size:1rem;font-weight:600;color:#f0f4f8;word-break:break-all;">{filename}</div>'
        f'</div></div>'
        f'<div style="display:flex;gap:10px;flex-wrap:wrap;margin-bottom:16px;">'
        f'<div style="display:flex;align-items:center;gap:6px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:8px;padding:6px 12px;font-size:0.78rem;color:#94a3b8;">{_svg_lang} {script}</div>'
        f'<div style="display:flex;align-items:center;gap:6px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:8px;padding:6px 12px;font-size:0.78rem;color:#94a3b8;">{_svg_pct} {english_pct}% English characters</div>'
        f'</div>'
        f'<div style="background:rgba(0,0,0,0.2);border:1px solid rgba(255,255,255,0.06);border-radius:10px;padding:12px 16px;margin-bottom:16px;">'
        f'<div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;font-weight:700;letter-spacing:0.07em;text-transform:uppercase;color:#64748b;margin-bottom:10px;">{_svg_label_why} Why This Happened</div>'
        f'<p style="margin:0;color:#94a3b8;font-size:0.82rem;line-height:1.6;">This tool is optimized for <strong style="color:#cbd5e1;">English-language resumes</strong>. Non-English content cannot be accurately parsed, rewritten, or ATS-scored. The detected script is <strong style="color:#cbd5e1;">{script}</strong> ({100 - english_pct}% non-English characters).</p>'
        f'</div>'
        f'<div style="background:rgba(56,189,248,0.07);border:1px solid rgba(56,189,248,0.18);border-radius:10px;padding:12px 16px;">'
        f'<div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;font-weight:700;letter-spacing:0.07em;text-transform:uppercase;color:#38bdf8;margin-bottom:10px;">{_svg_label_fix} How to Fix This</div>'
        f'<ul style="margin:0;padding:0;color:#7dd3fc;font-size:0.82rem;line-height:1.8;">{fix_html}</ul>'
        f'</div>'
        f'</div>'
    )
    _render_target = container if container is not None else st
    _render_target.markdown(card_html, unsafe_allow_html=True)


# ============================================================
# 📏 Long Resume Warning Card
# ============================================================

# Character threshold — beyond this, bottom sections may be silently cut
# (lowered from 8000 → 4000 to match resume_engine.py's prompt truncation,
#  which was reduced for GPT-OSS's smaller TPM budget vs llama-3.3)
_LONG_RESUME_THRESHOLD = 4000

def _render_long_resume_warning(filename: str, char_count: int, container=None):
    """
    Renders a non-blocking warning card for resumes that exceed 8000 characters.
    Does NOT block processing — just informs the user.
    Matches the app's existing glassmorphism dark theme exactly.
    """
    _svg_header = '<svg width="28" height="28" viewBox="0 0 24 24" fill="none" stroke="#fbbf24" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M10.29 3.86L1.82 18a2 2 0 001.71 3h16.94a2 2 0 001.71-3L13.71 3.86a2 2 0 00-3.42 0z"/><line x1="12" y1="9" x2="12" y2="13"/><line x1="12" y1="17" x2="12.01" y2="17"/></svg>'
    _dot_svg    = '<svg width="9" height="9" viewBox="0 0 9 9" fill="none" xmlns="http://www.w3.org/2000/svg" style="flex-shrink:0;margin-top:1px;"><circle cx="4.5" cy="4.5" r="4.5" fill="#fbbf24"/></svg>'
    _svg_chars  = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;"><polyline points="4 7 4 4 20 4 20 7"/><line x1="9" y1="20" x2="15" y2="20"/><line x1="12" y1="4" x2="12" y2="20"/></svg>'

    _svg_fix1 = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>'
    _svg_fix2 = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><line x1="18" y1="6" x2="6" y2="18"/><line x1="6" y1="6" x2="18" y2="18"/></svg>'
    _svg_fix3 = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#7dd3fc" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;margin-top:3px;"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg>'

    fix_items = [
        (_svg_fix1, "Keep your resume to <strong>1–2 pages</strong> — ATS systems and recruiters prefer concise resumes"),
        (_svg_fix2, "Remove outdated roles, redundant bullets, or overly detailed project descriptions"),
        (_svg_fix3, "Processing will continue — but sections beyond ~8,000 characters may be partially analysed"),
    ]
    fix_html = "".join(
        f"<li style='display:flex;align-items:flex-start;gap:8px;margin-bottom:8px;list-style:none;'>"
        f"{icon}<span>{text}</span></li>"
        for icon, text in fix_items
    )

    _svg_label_fix = '<svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>'

    card_html = (
        '<div style="background:linear-gradient(135deg,rgba(251,191,36,0.12) 0%,rgba(0,0,0,0) 100%);border:1px solid rgba(251,191,36,0.30);border-radius:16px;padding:22px 24px;margin:14px 0;backdrop-filter:blur(20px);-webkit-backdrop-filter:blur(20px);box-shadow:0 8px 32px rgba(0,0,0,0.3),inset 0 1px 0 rgba(255,255,255,0.06);font-family:-apple-system,BlinkMacSystemFont,sans-serif;position:relative;overflow:hidden;">'
        '<div style="position:absolute;top:0;left:0;right:0;height:3px;background:linear-gradient(90deg,transparent,#fbbf24,transparent);opacity:0.6;"></div>'
        f'<div style="display:flex;align-items:flex-start;gap:14px;margin-bottom:16px;">'
        f'<div style="width:44px;height:44px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:10px;display:flex;align-items:center;justify-content:center;flex-shrink:0;">{_svg_header}</div>'
        f'<div style="flex:1;">'
        f'<div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;color:#fbbf24;margin-bottom:4px;">{_dot_svg}Resume Too Long</div>'
        f'<div style="font-size:1rem;font-weight:600;color:#f0f4f8;word-break:break-all;">{filename}</div>'
        f'</div></div>'
        f'<div style="display:flex;gap:10px;flex-wrap:wrap;margin-bottom:16px;">'
        f'<div style="display:flex;align-items:center;gap:6px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:8px;padding:6px 12px;font-size:0.78rem;color:#94a3b8;">{_svg_chars} {char_count:,} characters detected</div>'
        f'<div style="display:flex;align-items:center;gap:6px;background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.09);border-radius:8px;padding:6px 12px;font-size:0.78rem;color:#94a3b8;">{_svg_chars} Recommended limit: 8,000 characters</div>'
        f'</div>'
        f'<div style="background:rgba(56,189,248,0.07);border:1px solid rgba(56,189,248,0.18);border-radius:10px;padding:12px 16px;">'
        f'<div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;font-weight:700;letter-spacing:0.07em;text-transform:uppercase;color:#38bdf8;margin-bottom:10px;">{_svg_label_fix} What This Means</div>'
        f'<ul style="margin:0;padding:0;color:#7dd3fc;font-size:0.82rem;line-height:1.8;">{fix_html}</ul>'
        f'</div>'
        '</div>'
    )
    _render_target = container if container is not None else st
    _render_target.markdown(card_html, unsafe_allow_html=True)

def safe_extract_text(uploaded_file, container=None):
    """
    Main entry point for PDF text extraction.

    Flow:
      1. Write file to /tmp
      2. Run _classify_pdf() — multi-signal scanned detector
      3. If DEFINITE scanned  → render rejection card, return sentinel
      4. If LIKELY scanned    → render rejection card, return sentinel
      5. If UNCERTAIN         → attempt OCR, return text if usable or sentinel
      6. If text-based        → extract with PyMuPDF, return text
      7. If extracted text is still empty after all attempts → warn, return None

    Returns:
      str   — extracted resume text (usable)
      _SCANNED_SENTINEL — file is image-based, rejection card already shown
      None  — file is unreadable for an unknown reason
    """
    try:
        temp_path = f"/tmp/{uploaded_file.name}"
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # ── Step 1: classify the PDF ──────────────────────────────────────
        classification = _classify_pdf(temp_path)

        # ── Step 2: handle scanned/image PDFs ────────────────────────────
        if classification["is_scanned"]:
            confidence = classification["confidence"]

            if confidence in ("definite", "likely"):
                # Hard reject — no OCR attempt, show full rejection card
                _render_scanned_rejection_card(uploaded_file.name, classification, container=container)
                return _SCANNED_SENTINEL

            else:
                # Uncertain — try OCR as last resort
                ocr_text_list = extract_text_from_images(temp_path)
                if ocr_text_list:
                    ocr_text = "\n".join(ocr_text_list)
                    ocr_words = len([w for w in ocr_text.split() if len(w) > 1])
                    if ocr_words >= 60:
                        # OCR gave enough signal — usable but warn the user
                        _render_target = container if container is not None else st
                        _render_target.markdown(f"""
                        <div class='slide-message warn-msg'>
                            <svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="flex-shrink:0;vertical-align:middle;"><path d="M10.29 3.86L1.82 18a2 2 0 001.71 3h16.94a2 2 0 001.71-3L13.71 3.86a2 2 0 00-3.42 0z"/><line x1="12" y1="9" x2="12" y2="13"/><line x1="12" y1="17" x2="12.01" y2="17"/></svg>
                            <b>{uploaded_file.name}</b> appears partially scanned.
                            Analysis is based on OCR text ({ocr_words} words) — accuracy may be reduced.
                            For best results, upload a text-based PDF.
                        </div>
                        """, unsafe_allow_html=True)
                        return ocr_text

                # OCR gave too little — show rejection card
                _render_scanned_rejection_card(uploaded_file.name, classification, container=container)
                return _SCANNED_SENTINEL

        # ── Step 3: text-based PDF — normal extraction ────────────────────
        text_list = extract_text_from_pdf(temp_path)

        if not text_list or all(len(t.strip()) == 0 for t in text_list):
            _render_target = container if container is not None else st
            _render_target.warning("⚠️ This file doesn't look like a resume or contains no readable text.")
            return None

        full_text = "\n".join(text_list)

        # ── Step 4: Non-English detection — block if resume is non-English ──
        lang_result = _detect_non_english(full_text)
        if lang_result["is_non_english"] and lang_result["confidence"] in ("definite", "likely"):
            _render_non_english_card(uploaded_file.name, lang_result, container=container)
            return _NON_ENGLISH_SENTINEL

        # ── Step 5: Long resume warning — non-blocking, informational only ──
        if len(full_text) > _LONG_RESUME_THRESHOLD:
            _render_long_resume_warning(uploaded_file.name, len(full_text), container=container)
            # Processing continues — warning is informational only

        return full_text

    except Exception as e:
        st.error(f"⚠️ Could not process this file: {e}")
        return None


# ============================================================
# 🏷️ Deterministic Candidate Name Extractor
# ============================================================



# ============================================================
# 📐 Industry-Standard Resume Format Checker (v2 — Enhanced)
# ============================================================

def _detect_multicolumn_pdf(pdf_path: str) -> bool:
    """
    Detect multi-column layout by analysing raw text-block x-coordinates
    from the first page of the PDF via PyMuPDF.
    Returns True if two or more distinct horizontal content zones are found.
    """
    try:
        doc = fitz.open(pdf_path)
        page = doc[0]
        blocks = page.get_text("blocks")   # (x0, y0, x1, y1, text, block_no, block_type)
        page_width = page.rect.width
        doc.close()

        # Only consider blocks with meaningful text content
        x_starts = [b[0] for b in blocks if len(b[4].strip()) > 10]
        if len(x_starts) < 6:
            return False

        # Split the page width into left zone (< 45 %) and right zone (> 52 %)
        left_zone  = [x for x in x_starts if x < page_width * 0.45]
        right_zone = [x for x in x_starts if x > page_width * 0.52]

        # Multi-column confirmed when both zones carry real content
        return len(left_zone) >= 3 and len(right_zone) >= 3
    except Exception:
        return False


def check_resume_format(text: str, num_pages: int = 1, pdf_path: str = None, session=None) -> dict:
    """
    Hybrid ATS format checker:
      • Regex  — fast, deterministic checks for email, phone, URLs, word count,
                 employment dates, encoding, page count, multi-column layout.
      • LLM    — intelligent checks for section presence, action verb quality,
                 quantified achievements, buzzword stuffing, and ATS red flags.
                 The LLM reads the actual resume content so it never misses a section
                 due to an unusual heading like "Internship Experience" or "Tech Stack".

    Scoring model (100 pts total, deduction-based):
      Contact completeness   — up to −21 pts  (regex)
      Resume length/pages    — up to −14 pts  (regex)
      ATS structural flags   — up to −7  pts  (regex: multi-column, encoding, dates)
      LLM-assessed content   — up to −58 pts  (sections, verbs, achievements, flags)
      Bonus credits          — up to +4  pts  (certifications, portfolio — LLM)

    Returns a dict compatible with all existing callers (same keys as v1).
    """
    try:
        import streamlit as _st
        _session = session or getattr(_st, "session_state", None)
    except Exception:
        _session = session

    issues     = []
    passes     = []
    deductions = 0
    bonuses    = 0

    text_lower = text.lower() if text else ""

    # ══════════════════════════════════════════════════════════════════════
    # PART A — REGEX CHECKS  (fast, always run, no LLM needed)
    # ══════════════════════════════════════════════════════════════════════

    # A1. Contact / Email  (−15 if missing)
    has_email = bool(re.search(r'[\w.+-]+@[\w.-]+\.[a-z]{2,}', text or ""))
    if has_email:
        passes.append("Section present: Contact / Email")
    else:
        issues.append("Missing section: 'Contact / Email' — ATS will likely reject without it")
        deductions += 15

    # A2. Phone number  (−6 if missing)
    has_phone = bool(re.search(r'(\+?\d[\d\s\-\(\)]{7,}\d)', text or ""))
    if has_phone:
        passes.append("Section present: Phone Number")
    else:
        issues.append("Missing section: 'Phone Number' — ATS will likely reject without it")
        deductions += 6

    # A3. LinkedIn URL  (−5 if missing)
    if re.search(r'linkedin\.com/in/[\w\-]+', text_lower):
        passes.append("LinkedIn profile URL detected")
    else:
        issues.append("No LinkedIn URL — recruiters expect it; many ATS rank it as a signal")
        deductions += 5

    # A4. GitHub / Portfolio  (−3 if missing, +1 bonus if present)
    if re.search(r'github\.com/[\w\-]+', text_lower):
        passes.append("GitHub profile URL detected")
        bonuses += 1
    elif re.search(r'(portfolio|behance\.net|dribbble\.com|leetcode\.com|kaggle\.com)', text_lower):
        passes.append("Portfolio / professional profile URL detected")
        bonuses += 1
    else:
        issues.append("No GitHub or portfolio URL — especially important for technical roles")
        deductions += 3

    # A5. Word count / length  (−2 to −12)
    word_count = len(text.split()) if text else 0
    if word_count < 150:
        issues.append(
            f"Resume critically short ({word_count} words) — "
            "ATS expects 400–900 words; this will likely be filtered out"
        )
        deductions += 12
    elif word_count < 400:
        issues.append(
            f"Resume too short ({word_count} words) — "
            "aim for 400–900 words with detailed experience and skills"
        )
        deductions += 7
    elif word_count > 1400:
        issues.append(
            f"Resume too long ({word_count} words) — "
            "trim to under 1,000 words; ATS and recruiters prefer concise resumes"
        )
        deductions += 5
    elif word_count > 1000:
        issues.append(
            f"Resume slightly long ({word_count} words) — "
            "consider tightening to under 1,000 words"
        )
        deductions += 2
    else:
        passes.append(f"Optimal length ({word_count} words — within 400–1,000 word sweet spot)")

    # A6. Page count  (−4 if > 2 pages)
    if num_pages > 2:
        issues.append(
            f"Resume is {num_pages} pages — "
            "ATS industry standard is 1–2 pages; longer resumes are often truncated"
        )
        deductions += 4
    elif num_pages == 2:
        passes.append("Page count acceptable (2 pages — standard for 5+ years experience)")
    else:
        passes.append("Page count ideal (1 page — strong for early-career candidates)")

    # A7. Multi-column layout  (−7 if detected)
    multicolumn_detected = False
    if pdf_path:
        try:
            multicolumn_detected = _detect_multicolumn_pdf(pdf_path)
        except Exception:
            multicolumn_detected = False
    if not multicolumn_detected and text:
        multicolumn_detected = (text.count('\t') > 8 or text.count('|') > 12)
    if multicolumn_detected:
        issues.append(
            "Multi-column or table layout detected — "
            "many ATS parsers read columns out of order, scrambling your resume content; "
            "use a single-column layout"
        )
        deductions += 7
    else:
        passes.append("Single-column layout detected — ATS-safe structure")

    # A8. Employment dates  (−5 if missing)
    has_dates = bool(re.search(r'\b(19|20)\d{2}\b', text or ""))
    if not has_dates:
        issues.append(
            "No employment dates detected — "
            "ATS requires dates to build a timeline; "
            "add month/year ranges (e.g., 'Jan 2021 – Mar 2023')"
        )
        deductions += 5
    else:
        passes.append("Employment dates detected — ATS can parse your timeline")

    # A9. Consistent Month-Year dates  (+1 bonus)
    month_year_dates = re.findall(
        r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*(19|20)\d{2}\b',
        text_lower
    )
    if len(month_year_dates) >= 2:
        passes.append("Consistent Month-Year date format detected — preferred by ATS parsers")
        bonuses += 1

    # A10. Non-ASCII encoding  (−3 if high density)
    if text:
        special_char_count = len(re.findall(r'[^\x00-\x7F]', text))
        ratio = special_char_count / max(len(text), 1)
        if ratio > 0.04:
            issues.append(
                f"High non-ASCII character density ({special_char_count} chars) — "
                "special characters from stylised fonts or copy-paste can corrupt ATS parsing"
            )
            deductions += 3
        else:
            passes.append("Character encoding looks ATS-safe (low non-ASCII density)")

    # ══════════════════════════════════════════════════════════════════════
    # PART B — LLM CHECKS  (intelligent content analysis)
    # Covers: section presence, action verbs, quantified achievements,
    #         buzzword stuffing, ATS red flags — all in ONE LLM call.
    # Falls back to keyword heuristics if LLM is unavailable.
    # ══════════════════════════════════════════════════════════════════════

    # Metric counters — populated by LLM or fallback
    verb_count   = 0
    metric_count = 0

    llm_prompt = f"""You are a senior ATS (Applicant Tracking System) specialist with 15+ years of experience evaluating resumes for enterprise hiring systems.

Analyze the resume text below and return a structured JSON evaluation. Read the CONTENT carefully — do not rely on exact heading names. A resume with an "Internship" section HAS work experience. A resume with "Tech Stack" HAS a skills section. A resume that opens with a professional paragraph HAS a summary.

════════════════════════════════════════════════════
EVALUATION CRITERIA
════════════════════════════════════════════════════

1. SECTION PRESENCE — Determine if each section EXISTS based on content, not heading names:
   - Experience: ANY work, internship, apprenticeship, freelance, contract, or volunteer role with company/org name and dates
   - Education: ANY degree, diploma, school, college, university, certification program, or board exam
   - Skills: ANY list of technical skills, tools, languages, frameworks, platforms, or competencies
   - Summary / Profile: A professional summary paragraph OR objective statement OR career profile at the top

2. ACTION VERBS — Count distinct strong action verbs used (Developed, Led, Engineered, Reduced, Implemented, etc.)
   Strong = past-tense active verbs that show ownership. Weak = "responsible for", "helped", "worked on".

3. QUANTIFIED ACHIEVEMENTS — Count distinct metrics/numbers that show impact:
   Examples: "reduced latency by 30%", "served 10K users", "cut costs by $5K", "3 projects delivered"
   Do NOT count CGPA/grades, dates, phone numbers, or years of experience as achievements.

4. ATS RED FLAGS — identify any of these issues:
   - Uses "Objective" heading instead of "Summary" (outdated format)
   - Excessive buzzwords repeated 2+ times: passionate, hardworking, synergy, go-getter, results-driven, team player, dynamic, proactive, detail-oriented
   - No quantified impact anywhere in the resume

5. BONUS SIGNALS — identify any of these positive signals:
   - Has a certifications or credentials section (not just a skill listed as "certified in X")
   - Has a projects section with described project work

════════════════════════════════════════════════════
RETURN ONLY THIS EXACT JSON — NO PREAMBLE, NO MARKDOWN FENCES:
{{
  "sections": {{
    "Experience": true_or_false,
    "Education": true_or_false,
    "Skills": true_or_false,
    "Summary / Profile": true_or_false
  }},
  "verb_count": integer,
  "metric_count": integer,
  "red_flags": {{
    "uses_objective_not_summary": true_or_false,
    "buzzword_stuffing": true_or_false,
    "buzzwords_found": ["word1", "word2"],
    "no_quantified_achievements": true_or_false
  }},
  "bonuses": {{
    "has_certifications": true_or_false,
    "has_projects": true_or_false
  }}
}}

RESUME TEXT:
\"\"\"{text[:4000]}\"\"\"
"""

    # Penalties for LLM-assessed items (mirrors original scoring weights)
    _SECTION_PENALTIES = {
        "Experience":       12,
        "Education":        4,
        "Skills":           3,
        "Summary / Profile": 2,
    }

    llm_ok = False
    try:
        _raw = call_llm(llm_prompt, session=_session)
        _ERROR_PREFIXES = ("❌", "⚠️", "Error", "LLM unavailable", "No healthy", "rate limit", "quota")
        if _raw and not any(_raw.strip().startswith(p) for p in _ERROR_PREFIXES):
            # Strip markdown fences if LLM added them
            _clean = re.sub(r'^```(?:json)?\s*', '', _raw.strip(), flags=re.IGNORECASE)
            _clean = re.sub(r'\s*```$', '', _clean).strip()
            _j = json.loads(_clean)

            # ── Section presence ──────────────────────────────────────────
            for sec_name, penalty in _SECTION_PENALTIES.items():
                present = bool(_j.get("sections", {}).get(sec_name, False))
                if present:
                    passes.append(f"Section present: {sec_name}")
                else:
                    issues.append(
                        f"Missing section: '{sec_name}' — "
                        f"ATS {'will likely reject' if penalty >= 10 else 'may penalise'} without it"
                    )
                    deductions += penalty

            # ── Action verbs ──────────────────────────────────────────────
            verb_count = int(_j.get("verb_count", 0))
            if verb_count == 0:
                issues.append(
                    "No strong action verbs found — ATS and recruiters expect bullet points "
                    "starting with verbs like 'Engineered', 'Led', 'Optimized'"
                )
                deductions += 8
            elif verb_count < 3:
                issues.append(
                    f"Weak action verb usage ({verb_count} found) — "
                    "aim for 5+ distinct strong verbs across experience bullet points"
                )
                deductions += 5
            elif verb_count < 5:
                issues.append(
                    f"Limited action verb variety ({verb_count} found) — "
                    "diversify verbs to better demonstrate range of contributions"
                )
                deductions += 2
            else:
                passes.append(f"Strong action verb usage ({verb_count} distinct verbs detected)")

            # ── Quantified achievements ───────────────────────────────────
            metric_count = int(_j.get("metric_count", 0))
            if metric_count == 0:
                issues.append(
                    "No quantified achievements detected — add measurable impact "
                    "(e.g., 'reduced latency by 35%', 'served 10K users', 'saved $50K annually')"
                )
                deductions += 8
            elif metric_count < 3:
                issues.append(
                    f"Few quantified achievements ({metric_count} found) — "
                    "aim for 4+ metrics across your experience to demonstrate concrete impact"
                )
                deductions += 4
            else:
                passes.append(f"Quantified achievements present ({metric_count} metrics detected)")

            # ── Red flags ─────────────────────────────────────────────────
            _flags = _j.get("red_flags", {})
            if _flags.get("uses_objective_not_summary"):
                issues.append(
                    "Uses 'Objective' section — this is outdated; "
                    "replace with a modern 'Professional Summary' (2–3 targeted sentences)"
                )
                deductions += 3
            if _flags.get("buzzword_stuffing"):
                _bw = ", ".join(_flags.get("buzzwords_found", []))
                issues.append(
                    f"Possible keyword stuffing detected ({_bw}) — "
                    "overused buzzwords reduce credibility; replace with concrete examples"
                )
                deductions += 2

            # ── Bonus signals ─────────────────────────────────────────────
            _bon = _j.get("bonuses", {})
            if _bon.get("has_certifications"):
                passes.append("Certifications / credentials detected — strong ATS positive signal")
                bonuses += 1
            if _bon.get("has_projects"):
                passes.append("Projects section detected — demonstrates initiative beyond job roles")
                bonuses += 1

            llm_ok = True
    except Exception:
        llm_ok = False

    # ── Fallback: keyword heuristics if LLM unavailable ──────────────────
    if not llm_ok:
        _SECTION_KEYWORDS = {
            "Experience": [
                "experience", "employment", "work history", "career",
                "professional experience", "work experience", "positions held",
                "relevant experience", "professional background",
                "internship", "intern", "industrial training", "apprenticeship",
                "trainee", "placement", "freelance", "volunteer",
            ],
            "Education": [
                "education", "university", "college", "degree",
                "bachelor", "master", "b.tech", "b.sc", "m.sc",
                "mca", "bca", "phd", "diploma", "high school",
                "graduated", "pursuing", "b.e", "m.tech",
                "cbse", "icse", "hsc", "ssc", "12th", "10th",
            ],
            "Skills": [
                "skills", "technologies", "tech stack", "competencies",
                "proficiencies", "tools", "technical skills", "core competencies",
                "expertise", "technical proficiency", "programming languages",
                "languages & tools", "frameworks", "key skills", "skillset",
            ],
            "Summary / Profile": [
                "summary", "objective", "profile", "about me", "overview",
                "professional summary", "career objective", "personal statement",
                "professional profile", "executive summary", "career summary",
            ],
        }
        _SECTION_PENALTIES = {"Experience": 12, "Education": 4, "Skills": 3, "Summary / Profile": 2}
        for sec_name, penalty in _SECTION_PENALTIES.items():
            present = any(w in text_lower for w in _SECTION_KEYWORDS.get(sec_name, []))
            if not present and sec_name == "Summary / Profile":
                # Implicit summary: resume opens with 2+ sentences before the first section
                _top = (text or "")[:400].strip()
                present = len(re.findall(r'[.!?]', _top)) >= 2
            if present:
                passes.append(f"Section present: {sec_name}")
            else:
                issues.append(
                    f"Missing section: '{sec_name}' — "
                    f"ATS {'will likely reject' if penalty >= 10 else 'may penalise'} without it"
                )
                deductions += penalty

        # Keyword-based verb and metric fallback
        strong_verbs = [
            "architected", "engineered", "designed", "deployed", "optimized", "automated",
            "built", "launched", "developed", "implemented", "integrated", "configured",
            "migrated", "refactored", "debugged", "scaled", "maintained", "upgraded",
            "tested", "validated", "led", "managed", "directed", "oversaw", "supervised",
            "coordinated", "spearheaded", "mentored", "trained", "guided", "facilitated",
            "reduced", "increased", "improved", "accelerated", "streamlined", "transformed",
            "negotiated", "established", "executed", "delivered", "created",
            "resolved", "analyzed", "collaborated", "authored", "published",
            "researched", "evaluated", "identified", "forecasted", "presented",
        ]
        found_verbs = [v for v in strong_verbs if re.search(rf'\b{v}\b', text_lower)]
        verb_count = len(found_verbs)
        if verb_count == 0:
            issues.append(
                "No strong action verbs found — ATS and recruiters expect bullet points "
                "starting with verbs like 'Engineered', 'Led', 'Optimized'"
            )
            deductions += 8
        elif verb_count < 3:
            issues.append(f"Weak action verb usage ({verb_count} found) — aim for 5+ distinct strong verbs")
            deductions += 5
        elif verb_count < 5:
            issues.append(f"Limited action verb variety ({verb_count} found) — diversify verbs")
            deductions += 2
        else:
            passes.append(f"Strong action verb usage ({verb_count} distinct verbs detected)")

        _quant = []
        _quant += re.findall(r'\b\d+[\.,]?\d*\s*(%|percent)\b', text_lower)
        _quant += re.findall(r'\b\d+[,.]?\d*\s*(k|m)?\s*(users|clients|customers|projects|tickets|requests|transactions)\b', text_lower)
        _quant += re.findall(r'\$\s*\d+[\d,.]*\s*[kKmMbB]?\b', text_lower)
        _quant += re.findall(r'\b(doubled|tripled|halved)\b', text_lower)
        metric_count = len(_quant)
        if metric_count == 0:
            issues.append(
                "No quantified achievements detected — add measurable impact "
                "(e.g., 'reduced latency by 35%', 'served 10K users', 'saved $50K annually')"
            )
            deductions += 8
        elif metric_count < 3:
            issues.append(f"Few quantified achievements ({metric_count} found) — aim for 4+ metrics")
            deductions += 4
        else:
            passes.append(f"Quantified achievements present ({metric_count} metrics detected)")

        if any(w in text_lower for w in [
            "certification", "certified", "certificate", "aws certified",
            "google certified", "microsoft certified", "pmp", "cissp", "comptia", "coursera", "udemy", "edx",
        ]):
            passes.append("Certifications / credentials detected — strong ATS positive signal")
            bonuses += 1
        if any(w in text_lower for w in ["projects", "personal projects", "open source", "hackathon"]):
            passes.append("Projects section detected — demonstrates initiative beyond job roles")
            bonuses += 1
        if "objective" in text_lower and "summary" not in text_lower:
            issues.append(
                "Uses 'Objective' section — this is outdated; "
                "replace with a modern 'Professional Summary'"
            )
            deductions += 3

    # ══════════════════════════════════════════════════════════════════════
    # FINAL SCORE CALCULATION
    # ══════════════════════════════════════════════════════════════════════
    raw_score = max(0, min(100, 100 - deductions + bonuses))

    if raw_score >= 90:
        letter_grade, label = "A+", "ATS-Optimized"
    elif raw_score >= 80:
        letter_grade, label = "A", "Excellent Format"
    elif raw_score >= 70:
        letter_grade, label = "B+", "Good Format"
    elif raw_score >= 60:
        letter_grade, label = "B", "Acceptable"
    elif raw_score >= 45:
        letter_grade, label = "C", "Needs Work"
    else:
        letter_grade, label = "D", "Poor — Major Issues"

    return {
        "format_score":  raw_score,
        "letter_grade":  letter_grade,
        "label":         label,
        "issues":        issues,
        "passes":        passes,
        "word_count":    word_count,
        "deductions":    deductions,
        "bonuses":       bonuses,
        "verb_count":    verb_count,
        "metric_count":  metric_count,
        "multicolumn":   multicolumn_detected,
        "llm_assessed":  llm_ok,   # callers can check if LLM was used
    }

# Detect bias in resume
# Predefined gender-coded word listss
