# resume_docx_builder.py
# ══════════════════════════════════════════════════════════════════════════════
# DOCX Resume & Cover Letter Generator
#
# Produces properly formatted A4 DOCX files directly — no Sejda, no external service.
# Every template maps to a native python-docx implementation with correct:
#   - A4 page size (11906 × 16838 DXA)
#   - 20mm margins (top/bottom), 18mm margins (left/right)
#   - Proper line-break logic, no content overlap
#   - Section dividers, spacing, alignment
#
# Usage (in main.py / TAAB_2_RESUME_BUILDER.PY):
#
#   from resume_docx_builder import build_resume_docx, build_cover_letter_docx
#
#   docx_bytes = build_resume_docx(selected_template, st.session_state)
#   st.download_button("⬇️ Download DOCX", data=docx_bytes, file_name="resume.docx",
#                      mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
#
#   cl_bytes = build_cover_letter_docx(cover_letter_template, cl_data)
#   st.download_button("⬇️ Download Cover Letter DOCX", data=cl_bytes, ...)
# ══════════════════════════════════════════════════════════════════════════════

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re as _re


# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS — A4 in twips (1 inch = 1440 twips)
# ─────────────────────────────────────────────────────────────────────────────

A4_WIDTH_TWIPS  = 11906   # 210 mm
A4_HEIGHT_TWIPS = 16838   # 297 mm

# Margins: top/bottom = 20 mm, left/right = 18 mm
MARGIN_TOP    = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)
MARGIN_LEFT   = Cm(1.8)
MARGIN_RIGHT  = Cm(1.8)

# Content width in cm (page - margins)
CONTENT_WIDTH_CM = 21.0 - 1.8 - 1.8   # ≈ 17.4 cm


# ─────────────────────────────────────────────────────────────────────────────
# XML / DOCX HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _set_a4(doc: Document) -> None:
    """Set document to A4 portrait with defined margins."""
    section = doc.sections[0]
    section.page_width  = A4_WIDTH_TWIPS
    section.page_height = A4_HEIGHT_TWIPS
    section.top_margin    = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin   = MARGIN_LEFT
    section.right_margin  = MARGIN_RIGHT


def _set_para_spacing(para, before: float = 0, after: float = 0,
                      line_rule: str = "auto", line_val: int = 240) -> None:
    """Set paragraph spacing (before/after in points, line as twips)."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(before * 20)))   # pt → twips
    spacing.set(qn("w:after"),  str(int(after  * 20)))
    spacing.set(qn("w:line"),   str(line_val))
    spacing.set(qn("w:lineRule"), line_rule)
    pPr.append(spacing)


def _shading(cell, fill_hex: str) -> None:
    """Apply solid background colour to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.lstrip("#"))
    tcPr.append(shd)


def _cell_margins(cell, top: int = 80, bottom: int = 80,
                  start: int = 100, end: int = 100) -> None:
    """Set internal cell padding (in twips)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", start), ("right", end)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _no_borders(table) -> None:
    """Remove all borders from a table."""
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _set_col_width(cell, width_cm: float) -> None:
    """Force column width on a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    twips = int(width_cm * 567)     # 1 cm = 567 twips
    tcW.set(qn("w:w"),    str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _hr(doc: Document, color_hex: str = "CCCCCC", thickness: int = 6) -> None:
    """Add a horizontal rule paragraph."""
    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=2)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thickness))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(bot)
    pPr.append(pBdr)
    return para


def _section_heading(doc: Document, title: str,
                     color_hex: str = "2F4F6F",
                     font_name: str = "Calibri",
                     font_pt:   float = 11.0,
                     uppercase: bool = True) -> None:
    """Add a section heading with underline rule."""
    para = doc.add_paragraph()
    _set_para_spacing(para, before=8, after=2, line_val=276)
    run  = para.add_run(title.upper() if uppercase else title)
    run.bold      = True
    run.font.size = Pt(font_pt)
    run.font.name = font_name
    run.font.color.rgb = RGBColor.from_string(color_hex)
    _hr(doc, color_hex, thickness=8)


def _add_run(para, text: str, bold: bool = False, italic: bool = False,
             font_name: str = "Calibri", font_pt: float = 10.0,
             color_hex: str | None = None) -> None:
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def _para(doc: Document, text: str = "", bold: bool = False,
          italic: bool = False, font_name: str = "Calibri",
          font_pt: float = 10.0, color_hex: str | None = None,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          before: float = 0, after: float = 2,
          line_val: int = 252) -> None:
    para = doc.add_paragraph()
    para.alignment = align
    _set_para_spacing(para, before=before, after=after, line_val=line_val)
    if text:
        _add_run(para, text, bold=bold, italic=italic,
                 font_name=font_name, font_pt=font_pt, color_hex=color_hex)
    return para


def _bullet_para(doc: Document, text: str,
                 font_name: str = "Calibri", font_pt: float = 9.5,
                 color_hex: str = "374151") -> None:
    """Add a single bullet-point paragraph."""
    para = doc.add_paragraph(style="List Bullet")
    _set_para_spacing(para, before=0, after=1, line_val=240)
    run  = para.add_run(text.strip())
    run.font.name  = font_name
    run.font.size  = Pt(font_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _parse_description_bullets(text: str) -> list[tuple[str, bool]]:
    """
    Parse a description string into (text, is_bullet) tuples.
    Lines beginning with - • * · > – — are bullet items.
    Other non-blank lines are plain text.
    Returns list of (line_text, is_bullet).
    """
    PREFIXES = ("-", "•", "*", "·", ">", "–", "—")
    result = []
    for raw in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()
        is_bullet = False
        for p in PREFIXES:
            if stripped.startswith(p):
                line     = stripped[len(p):].strip()
                is_bullet = True
                break
        if not is_bullet:
            line = stripped
        result.append((line, is_bullet))
    return result


def _write_description(doc: Document, text: str,
                       font_name: str = "Calibri",
                       font_pt:   float = 9.5,
                       color_hex: str = "374151") -> None:
    """Write a description field: bullets as list items, plain text as para."""
    if not text or not text.strip():
        return
    for line, is_bullet in _parse_description_bullets(text):
        if is_bullet:
            _bullet_para(doc, line, font_name=font_name,
                         font_pt=font_pt, color_hex=color_hex)
        else:
            p = _para(doc, line, font_name=font_name, font_pt=font_pt,
                      color_hex=color_hex, before=0, after=1)


def _tab_pair(doc: Document, left_text: str, right_text: str,
              font_name: str = "Calibri",
              left_pt:  float = 10.5, left_bold:  bool = True,
              right_pt: float = 9.5,  right_bold: bool = False,
              left_color: str = "1F2937", right_color: str = "6B7280") -> None:
    """Single line: left-aligned text + right-aligned text via tab stop."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=0, line_val=240)

    # Set right tab stop at max page content width
    pPr   = para._p.get_or_add_pPr()
    tabs  = OxmlElement("w:tabs")
    tab   = OxmlElement("w:tab")
    # content width in twips
    content_w = int(CONTENT_WIDTH_CM * 567)
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(content_w))
    tabs.append(tab)
    pPr.append(tabs)

    r_left = para.add_run(left_text)
    r_left.bold        = left_bold
    r_left.font.name   = font_name
    r_left.font.size   = Pt(left_pt)
    r_left.font.color.rgb = RGBColor.from_string(left_color)

    r_tab = para.add_run("\t")
    r_tab.font.size = Pt(right_pt)

    r_right = para.add_run(right_text)
    r_right.bold        = right_bold
    r_right.font.name   = font_name
    r_right.font.size   = Pt(right_pt)
    r_right.font.color.rgb = RGBColor.from_string(right_color)


def _contact_line(doc: Document, ss: dict,
                  font_name: str = "Calibri",
                  font_pt:   float = 9.0,
                  color_hex: str = "374151",
                  separator: str = "  |  ") -> None:
    """Build a contact info line from session state fields."""
    parts = []
    for key in ("email", "phone", "location", "linkedin", "portfolio"):
        val = ss.get(key, "").strip()
        if val:
            parts.append(val)
    if not parts:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=0, after=2)
    full = separator.join(parts)
    run  = p.add_run(full)
    run.font.name  = font_name
    run.font.size  = Pt(font_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _skills_line(doc: Document, raw: str,
                 font_name: str = "Calibri",
                 font_pt:   float = 9.5,
                 color_hex: str = "374151") -> None:
    """Comma-separated skills as a wrapped paragraph."""
    items = [s.strip() for s in raw.split(",") if s.strip()]
    if not items:
        return
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=2)
    run = p.add_run("  ·  ".join(items))
    run.font.name  = font_name
    run.font.size  = Pt(font_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)


# ─────────────────────────────────────────────────────────────────────────────
# SHARED RESUME BODY BUILDER (used by all single-column templates)
# ─────────────────────────────────────────────────────────────────────────────

def _write_resume_body(doc: Document, ss: dict,
                       accent:     str = "2F4F6F",
                       body_color: str = "374151",
                       font_name:  str = "Calibri",
                       heading_pt: float = 11.0,
                       body_pt:    float = 9.5) -> None:
    """
    Write all main resume sections: Summary, Experience, Education, Projects, Certificates.
    Called by every single-column template to keep logic DRY.
    """
    # ── SUMMARY ──────────────────────────────────────────────────────────────
    summary = ss.get("summary", "").strip()
    if summary:
        _section_heading(doc, "Professional Summary", color_hex=accent,
                         font_name=font_name, font_pt=heading_pt)
        _write_description(doc, summary, font_name=font_name,
                           font_pt=body_pt, color_hex=body_color)
        _para(doc, before=0, after=4)

    # ── WORK EXPERIENCE ───────────────────────────────────────────────────────
    exp_entries = [e for e in ss.get("experience_entries", [])
                   if e.get("company") or e.get("title")]
    if exp_entries:
        _section_heading(doc, "Work Experience", color_hex=accent,
                         font_name=font_name, font_pt=heading_pt)
        for exp in exp_entries:
            company  = exp.get("company",  "").strip()
            title    = exp.get("title",    "").strip()
            duration = exp.get("duration", "").strip()
            desc     = exp.get("description", "").strip()
            _tab_pair(doc, company, duration,
                      font_name=font_name,
                      left_pt=body_pt + 0.5, left_bold=True, left_color=accent,
                      right_pt=body_pt - 0.5, right_color="6B7280")
            if title:
                _para(doc, title, italic=True, font_name=font_name,
                      font_pt=body_pt, color_hex="4B5563", before=0, after=1)
            if desc:
                _write_description(doc, desc, font_name=font_name,
                                   font_pt=body_pt - 0.5, color_hex=body_color)
            _para(doc, before=0, after=4)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    edu_entries = [e for e in ss.get("education_entries", [])
                   if e.get("institution") or e.get("degree")]
    if edu_entries:
        _section_heading(doc, "Education", color_hex=accent,
                         font_name=font_name, font_pt=heading_pt)
        for edu in edu_entries:
            inst   = edu.get("institution", "").strip()
            degree = edu.get("degree",      "").strip()
            if isinstance(degree, list):
                degree = ", ".join(degree)
            year    = edu.get("year",    "").strip()
            details = edu.get("details", "").strip()
            _tab_pair(doc, inst, year,
                      font_name=font_name,
                      left_pt=body_pt + 0.5, left_bold=True, left_color=accent,
                      right_pt=body_pt - 0.5, right_color="6B7280")
            if degree:
                _para(doc, degree, italic=True, font_name=font_name,
                      font_pt=body_pt, color_hex="4B5563", before=0, after=1)
            if details:
                _para(doc, details, font_name=font_name,
                      font_pt=body_pt - 0.5, color_hex=body_color, before=0, after=1)
            _para(doc, before=0, after=4)

    # ── PROJECTS ──────────────────────────────────────────────────────────────
    proj_entries  = [p for p in ss.get("project_entries", []) if p.get("title")]
    proj_links    = ss.get("project_links", [])
    if proj_entries:
        _section_heading(doc, "Projects", color_hex=accent,
                         font_name=font_name, font_pt=heading_pt)
        for idx, proj in enumerate(proj_entries):
            title    = proj.get("title",       "").strip()
            tech     = proj.get("tech",        "").strip()
            duration = proj.get("duration",    "").strip()
            desc     = proj.get("description", "").strip()
            link     = proj_links[idx] if idx < len(proj_links) else ""
            _tab_pair(doc, title, duration,
                      font_name=font_name,
                      left_pt=body_pt + 0.5, left_bold=True, left_color=accent,
                      right_pt=body_pt - 0.5, right_color="6B7280")
            if tech:
                p = doc.add_paragraph()
                _set_para_spacing(p, before=0, after=1)
                r1 = p.add_run("Tech: ")
                r1.bold = True
                r1.font.name = font_name
                r1.font.size = Pt(body_pt - 0.5)
                r1.font.color.rgb = RGBColor.from_string("4B5563")
                r2 = p.add_run(tech)
                r2.font.name = font_name
                r2.font.size = Pt(body_pt - 0.5)
                r2.font.color.rgb = RGBColor.from_string(body_color)
            if desc:
                _write_description(doc, desc, font_name=font_name,
                                   font_pt=body_pt - 0.5, color_hex=body_color)
            if link:
                _para(doc, f"Link: {link}", italic=True, font_name=font_name,
                      font_pt=body_pt - 1, color_hex="2563EB", before=0, after=1)
            _para(doc, before=0, after=4)

    # ── CERTIFICATES ──────────────────────────────────────────────────────────
    certs = [c for c in ss.get("certificate_links", []) if c.get("name")]
    if certs:
        _section_heading(doc, "Certifications", color_hex=accent,
                         font_name=font_name, font_pt=heading_pt)
        for cert in certs:
            name = cert.get("name",        "").strip()
            link = cert.get("link",        "").strip()
            dur  = cert.get("duration",    "").strip()
            desc = cert.get("description", "").strip()
            _tab_pair(doc, name, dur,
                      font_name=font_name,
                      left_pt=body_pt, left_bold=True, left_color=accent,
                      right_pt=body_pt - 0.5, right_color="6B7280")
            if link:
                _para(doc, link, italic=True, font_name=font_name,
                      font_pt=body_pt - 1, color_hex="2563EB", before=0, after=1)
            if desc:
                _para(doc, desc, font_name=font_name,
                      font_pt=body_pt - 1, color_hex=body_color, before=0, after=1)
            _para(doc, before=0, after=3)


# ─────────────────────────────────────────────────────────────────────────────
# TEMPLATE 1: DEFAULT PROFESSIONAL (two-column sidebar)
# ─────────────────────────────────────────────────────────────────────────────

def _build_default(doc: Document, ss: dict) -> None:
    """
    Dark sidebar (30% width) + white main content (70% width).
    Sidebar: photo placeholder, contact, skills, soft skills, languages, interests, certs.
    Main: name/title header, summary, experience, education, projects.
    """
    _set_a4(doc)

    SIDEBAR_BG  = "374151"
    SIDEBAR_TXT = "FFFFFF"
    ACCENT      = "9CA3AF"
    MAIN_ACCENT = "374151"
    FONT        = "Calibri"

    # Total content width in twips ≈ 17.4 cm * 567 ≈ 9865
    total_w = int(CONTENT_WIDTH_CM * 567)
    sidebar_w = int(total_w * 0.30)
    main_w    = total_w - sidebar_w

    table = doc.add_table(rows=1, cols=2)
    _no_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    side_cell = table.rows[0].cells[0]
    main_cell = table.rows[0].cells[1]
    _set_col_width(side_cell, CONTENT_WIDTH_CM * 0.30)
    _set_col_width(main_cell, CONTENT_WIDTH_CM * 0.70)
    _shading(side_cell, SIDEBAR_BG)
    _cell_margins(side_cell, top=200, bottom=200, start=150, end=120)
    _cell_margins(main_cell, top=200, bottom=200, start=180, end=140)

    # ── Sidebar content ───────────────────────────────────────────────────────
    def _sb_heading(cell, title):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=10, after=3)
        r = p.add_run(title.upper())
        r.bold = True
        r.font.name  = FONT
        r.font.size  = Pt(8)
        r.font.color.rgb = RGBColor.from_string("E5E7EB")
        # underline via border
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "9CA3AF")
        pBdr.append(bot)
        pPr.append(pBdr)

    def _sb_line(cell, text, italic=False):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=0, after=2)
        r = p.add_run(text.strip())
        r.italic = italic
        r.font.name  = FONT
        r.font.size  = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("E5E7EB")

    # Name + title in sidebar top
    p_name = side_cell.add_paragraph()
    _set_para_spacing(p_name, before=0, after=4)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_name.add_run(ss.get("name", ""))
    r.bold = True
    r.font.name  = FONT
    r.font.size  = Pt(13)
    r.font.color.rgb = RGBColor.from_string("FFFFFF")

    if ss.get("job_title"):
        p_title = side_cell.add_paragraph()
        _set_para_spacing(p_title, before=0, after=12)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_title.add_run(ss.get("job_title", "").upper())
        r.font.name  = FONT
        r.font.size  = Pt(7.5)
        r.font.color.rgb = RGBColor.from_string("D1D5DB")

    # Contact
    contact_fields = [
        ("Email",     ss.get("email",     "")),
        ("Phone",     ss.get("phone",     "")),
        ("Location",  ss.get("location",  "")),
        ("LinkedIn",  ss.get("linkedin",  "")),
        ("Portfolio", ss.get("portfolio", "")),
    ]
    has_contact = any(v for _, v in contact_fields)
    if has_contact:
        _sb_heading(side_cell, "Contact")
        for label, val in contact_fields:
            if val:
                _sb_line(side_cell, val)

    # Skills
    for section_title, key in [
        ("Technical Skills", "skills"),
        ("Soft Skills",       "Softskills"),
        ("Languages",         "languages"),
        ("Interests",         "interests"),
    ]:
        raw = ss.get(key, "").strip()
        if raw:
            _sb_heading(side_cell, section_title)
            for item in [s.strip() for s in raw.split(",") if s.strip()]:
                _sb_line(side_cell, f"• {item}")

    # Certificates
    certs = [c for c in ss.get("certificate_links", []) if c.get("name")]
    if certs:
        _sb_heading(side_cell, "Certifications")
        for c in certs:
            _sb_line(side_cell, c.get("name", ""), italic=True)
            if c.get("duration"):
                _sb_line(side_cell, c.get("duration", ""))

    # ── Main content ──────────────────────────────────────────────────────────
    def _mc_section(cell, title):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=10, after=2)
        r = p.add_run(title.upper())
        r.bold = True
        r.font.name  = FONT
        r.font.size  = Pt(10)
        r.font.color.rgb = RGBColor.from_string(MAIN_ACCENT)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "8")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "9CA3AF")
        pBdr.append(bot)
        pPr.append(pBdr)

    def _mc_para(cell, text, bold=False, italic=False, pt=9.5, color="374151", before=0, after=2):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=before, after=after)
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.name  = FONT
        r.font.size  = Pt(pt)
        r.font.color.rgb = RGBColor.from_string(color)
        return p

    def _mc_bullet(cell, text, pt=9.0, color="374151"):
        p = cell.add_paragraph(style="List Bullet")
        _set_para_spacing(p, before=0, after=1)
        r = p.add_run(text.strip())
        r.font.name  = FONT
        r.font.size  = Pt(pt)
        r.font.color.rgb = RGBColor.from_string(color)

    def _mc_tab(cell, left, right, left_bold=True, left_pt=10.5, right_pt=9.0,
                left_color=MAIN_ACCENT, right_color="6B7280"):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=0, after=0)
        pPr   = p._p.get_or_add_pPr()
        tabs  = OxmlElement("w:tabs")
        tab   = OxmlElement("w:tab")
        cw    = int(CONTENT_WIDTH_CM * 0.70 * 567)
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(cw))
        tabs.append(tab)
        pPr.append(tabs)
        r1 = p.add_run(left)
        r1.bold      = left_bold
        r1.font.name = FONT
        r1.font.size = Pt(left_pt)
        r1.font.color.rgb = RGBColor.from_string(left_color)
        r2 = p.add_run("\t" + right)
        r2.font.name = FONT
        r2.font.size = Pt(right_pt)
        r2.font.color.rgb = RGBColor.from_string(right_color)

    # Summary
    summary = ss.get("summary", "").strip()
    if summary:
        _mc_section(main_cell, "Professional Summary")
        for line, is_b in _parse_description_bullets(summary):
            if is_b:
                _mc_bullet(main_cell, line)
            else:
                _mc_para(main_cell, line, pt=9.5)
        _mc_para(main_cell, "", before=0, after=4)

    # Experience
    exps = [e for e in ss.get("experience_entries", [])
            if e.get("company") or e.get("title")]
    if exps:
        _mc_section(main_cell, "Work Experience")
        for exp in exps:
            _mc_tab(main_cell, exp.get("company", ""), exp.get("duration", ""))
            if exp.get("title"):
                _mc_para(main_cell, exp.get("title", ""), italic=True, pt=9.5, color="4B5563")
            for line, is_b in _parse_description_bullets(exp.get("description", "")):
                if is_b:
                    _mc_bullet(main_cell, line)
                else:
                    _mc_para(main_cell, line, pt=9.0)
            _mc_para(main_cell, "", before=0, after=4)

    # Education
    edus = [e for e in ss.get("education_entries", [])
            if e.get("institution") or e.get("degree")]
    if edus:
        _mc_section(main_cell, "Education")
        for edu in edus:
            degree = edu.get("degree", "")
            if isinstance(degree, list):
                degree = ", ".join(degree)
            _mc_tab(main_cell, edu.get("institution", ""), edu.get("year", ""))
            if degree:
                _mc_para(main_cell, degree, italic=True, pt=9.5, color="4B5563")
            if edu.get("details"):
                _mc_para(main_cell, edu.get("details", ""), pt=9.0)
            _mc_para(main_cell, "", before=0, after=4)

    # Projects
    projs = [p for p in ss.get("project_entries", []) if p.get("title")]
    links = ss.get("project_links", [])
    if projs:
        _mc_section(main_cell, "Projects")
        for idx, proj in enumerate(projs):
            _mc_tab(main_cell, proj.get("title", ""), proj.get("duration", ""))
            if proj.get("tech"):
                _mc_para(main_cell, f"Tech: {proj['tech']}", italic=True, pt=9.0, color="4B5563")
            for line, is_b in _parse_description_bullets(proj.get("description", "")):
                if is_b:
                    _mc_bullet(main_cell, line)
                else:
                    _mc_para(main_cell, line, pt=9.0)
            if idx < len(links) and links[idx]:
                _mc_para(main_cell, f"Link: {links[idx]}", italic=True, pt=8.5, color="2563EB")
            _mc_para(main_cell, "", before=0, after=4)


# ─────────────────────────────────────────────────────────────────────────────
# TEMPLATE 2: MODERN MINIMAL (single column, teal accents)
# ─────────────────────────────────────────────────────────────────────────────

def _build_modern(doc: Document, ss: dict) -> None:
    _set_a4(doc)
    ACCENT = "0D9488"
    FONT   = "Calibri"

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(p, before=0, after=2)
    r = p.add_run(ss.get("name", ""))
    r.bold = True
    r.font.name  = FONT
    r.font.size  = Pt(22)
    r.font.color.rgb = RGBColor.from_string("0F172A")

    if ss.get("job_title"):
        p2 = doc.add_paragraph()
        _set_para_spacing(p2, before=0, after=4)
        r2 = p2.add_run(ss.get("job_title", ""))
        r2.font.name  = FONT
        r2.font.size  = Pt(12)
        r2.font.color.rgb = RGBColor.from_string(ACCENT)

    _contact_line(doc, ss, font_name=FONT, font_pt=9.0, color_hex="6B7280",
                  separator="  ·  ")
    _hr(doc, ACCENT, thickness=12)

    # Skills summary line
    if ss.get("skills"):
        _skills_line(doc, ss.get("skills", ""), font_name=FONT,
                     font_pt=9.0, color_hex="374151")
        _hr(doc, "E5E7EB", thickness=4)

    _write_resume_body(doc, ss, accent=ACCENT, body_color="374151",
                       font_name=FONT, heading_pt=10.5, body_pt=9.5)


# ─────────────────────────────────────────────────────────────────────────────
# TEMPLATE 3: ELEGANT SIDEBAR (purple)
# ─────────────────────────────────────────────────────────────────────────────

def _build_sidebar(doc: Document, ss: dict) -> None:
    """Two-column: narrow purple sidebar (28%) + wide main (72%)."""
    _set_a4(doc)
    ACCENT = "7C3AED"
    FONT   = "Calibri"

    # Reuse default two-column logic with purple palette
    _build_default.__wrapped__ if hasattr(_build_default, "__wrapped__") else None

    total_w   = int(CONTENT_WIDTH_CM * 567)
    sidebar_w = int(total_w * 0.28)
    main_w    = total_w - sidebar_w

    table = doc.add_table(rows=1, cols=2)
    _no_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    sc = table.rows[0].cells[0]
    mc = table.rows[0].cells[1]
    _set_col_width(sc, CONTENT_WIDTH_CM * 0.28)
    _set_col_width(mc, CONTENT_WIDTH_CM * 0.72)
    _shading(sc, "F5F3FF")
    _cell_margins(sc, top=200, bottom=200, start=140, end=110)
    _cell_margins(mc, top=200, bottom=200, start=180, end=140)

    def _sb(cell, title):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=10, after=2)
        r = p.add_run(title.upper())
        r.bold = True
        r.font.name  = FONT
        r.font.size  = Pt(8)
        r.font.color.rgb = RGBColor.from_string(ACCENT)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"),  "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), ACCENT)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _sl(cell, text, italic=False):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=0, after=2)
        r = p.add_run(text.strip())
        r.italic = italic
        r.font.name  = FONT
        r.font.size  = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("374151")

    # Sidebar name
    pn = sc.add_paragraph()
    _set_para_spacing(pn, before=0, after=2)
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = pn.add_run(ss.get("name", ""))
    rn.bold = True
    rn.font.name  = FONT
    rn.font.size  = Pt(12)
    rn.font.color.rgb = RGBColor.from_string(ACCENT)

    if ss.get("job_title"):
        pjt = sc.add_paragraph()
        _set_para_spacing(pjt, before=0, after=10)
        pjt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rjt = pjt.add_run(ss.get("job_title", ""))
        rjt.font.name  = FONT
        rjt.font.size  = Pt(7.5)
        rjt.font.color.rgb = RGBColor.from_string("6B7280")

    for label, val in [
        ("Email",    ss.get("email",    "")),
        ("Phone",    ss.get("phone",    "")),
        ("Location", ss.get("location", "")),
        ("LinkedIn", ss.get("linkedin", "")),
    ]:
        if val:
            _sb(sc, label) if not label else None
    has_c = any(ss.get(k) for k in ("email","phone","location","linkedin"))
    if has_c:
        _sb(sc, "Contact")
        for k in ("email","phone","location","linkedin","portfolio"):
            if ss.get(k):
                _sl(sc, ss.get(k))

    for sec, key in [("Skills", "skills"), ("Soft Skills", "Softskills"),
                     ("Languages", "languages"), ("Interests", "interests")]:
        raw = ss.get(key, "").strip()
        if raw:
            _sb(sc, sec)
            for item in [s.strip() for s in raw.split(",") if s.strip()]:
                _sl(sc, f"• {item}")

    certs = [c for c in ss.get("certificate_links", []) if c.get("name")]
    if certs:
        _sb(sc, "Certifications")
        for c in certs:
            _sl(sc, c.get("name", ""), italic=True)

    # Main column: same as default but purple accent
    def _ms(cell, title):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=10, after=2)
        r = p.add_run(title.upper())
        r.bold = True
        r.font.name  = FONT
        r.font.size  = Pt(10)
        r.font.color.rgb = RGBColor.from_string(ACCENT)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"),  "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), ACCENT)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _ml(cell, text, bold=False, italic=False, pt=9.5, color="374151",
            before=0, after=2):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=before, after=after)
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.name  = FONT
        r.font.size  = Pt(pt)
        r.font.color.rgb = RGBColor.from_string(color)

    def _mb(cell, text, pt=9.0, color="374151"):
        p = cell.add_paragraph(style="List Bullet")
        _set_para_spacing(p, before=0, after=1)
        r = p.add_run(text.strip())
        r.font.name  = FONT
        r.font.size  = Pt(pt)
        r.font.color.rgb = RGBColor.from_string(color)

    def _mt(cell, left, right, left_bold=True, left_pt=10.5, right_pt=9.0):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=0, after=0)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab  = OxmlElement("w:tab")
        cw   = int(CONTENT_WIDTH_CM * 0.72 * 567)
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(cw))
        tabs.append(tab)
        pPr.append(tabs)
        r1 = p.add_run(left)
        r1.bold      = left_bold
        r1.font.name = FONT
        r1.font.size = Pt(left_pt)
        r1.font.color.rgb = RGBColor.from_string(ACCENT)
        r2 = p.add_run("\t" + right)
        r2.font.name = FONT
        r2.font.size = Pt(right_pt)
        r2.font.color.rgb = RGBColor.from_string("6B7280")

    summary = ss.get("summary", "").strip()
    if summary:
        _ms(mc, "Summary")
        for line, is_b in _parse_description_bullets(summary):
            if is_b:
                _mb(mc, line)
            else:
                _ml(mc, line, pt=9.5)
        _ml(mc, "", before=0, after=4)

    exps = [e for e in ss.get("experience_entries", [])
            if e.get("company") or e.get("title")]
    if exps:
        _ms(mc, "Experience")
        for e in exps:
            _mt(mc, e.get("company",""), e.get("duration",""))
            if e.get("title"):
                _ml(mc, e.get("title",""), italic=True, pt=9.5, color="4B5563")
            for ln, ib in _parse_description_bullets(e.get("description","")):
                if ib:
                    _mb(mc, ln)
                else:
                    _ml(mc, ln, pt=9.0)
            _ml(mc, "", before=0, after=4)

    edus = [e for e in ss.get("education_entries", [])
            if e.get("institution") or e.get("degree")]
    if edus:
        _ms(mc, "Education")
        for edu in edus:
            dg = edu.get("degree","")
            if isinstance(dg, list):
                dg = ", ".join(dg)
            _mt(mc, edu.get("institution",""), edu.get("year",""))
            if dg:
                _ml(mc, dg, italic=True, pt=9.5, color="4B5563")
            if edu.get("details"):
                _ml(mc, edu.get("details",""), pt=9.0)
            _ml(mc, "", before=0, after=4)

    projs = [p for p in ss.get("project_entries", []) if p.get("title")]
    links = ss.get("project_links", [])
    if projs:
        _ms(mc, "Projects")
        for idx, proj in enumerate(projs):
            _mt(mc, proj.get("title",""), proj.get("duration",""))
            if proj.get("tech"):
                _ml(mc, f"Tech: {proj['tech']}", italic=True, pt=9.0, color="4B5563")
            for ln, ib in _parse_description_bullets(proj.get("description","")):
                if ib:
                    _mb(mc, ln)
                else:
                    _ml(mc, ln, pt=9.0)
            if idx < len(links) and links[idx]:
                _ml(mc, f"Link: {links[idx]}", italic=True, pt=8.5, color="2563EB")
            _ml(mc, "", before=0, after=4)

    certs2 = [c for c in ss.get("certificate_links", []) if c.get("name")]
    if certs2:
        _ms(mc, "Certifications")
        for c in certs2:
            _mt(mc, c.get("name",""), c.get("duration",""), left_pt=9.5)
            if c.get("description"):
                _ml(mc, c.get("description",""), pt=9.0)
            _ml(mc, "", before=0, after=3)


# ─────────────────────────────────────────────────────────────────────────────
# SINGLE-COLUMN TEMPLATES (Classic, Executive, Timeline, Slate, Burgundy, etc.)
# All share _write_resume_body with different accent colours / font choices.
# ─────────────────────────────────────────────────────────────────────────────

def _build_single_column(doc: Document, ss: dict,
                         accent:     str   = "1E3A5F",
                         name_color: str   = "1E3A5F",
                         contact_color: str = "374151",
                         font_name:  str   = "Calibri",
                         name_pt:    float = 22.0,
                         title_pt:   float = 12.0,
                         heading_pt: float = 10.5,
                         body_pt:    float = 9.5,
                         header_border_color: str = "1E3A5F") -> None:
    """Generic single-column resume builder shared by many templates."""
    _set_a4(doc)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=0, after=2)
    r = p.add_run(ss.get("name", ""))
    r.bold = True
    r.font.name  = font_name
    r.font.size  = Pt(name_pt)
    r.font.color.rgb = RGBColor.from_string(name_color)

    if ss.get("job_title"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p2, before=0, after=3)
        r2 = p2.add_run(ss.get("job_title", ""))
        r2.font.name  = font_name
        r2.font.size  = Pt(title_pt)
        r2.font.color.rgb = RGBColor.from_string(accent)

    _contact_line(doc, ss, font_name=font_name, font_pt=8.5,
                  color_hex=contact_color, separator="  |  ")
    _hr(doc, header_border_color, thickness=10)

    # Skills summary
    if ss.get("skills"):
        p_sk = doc.add_paragraph()
        _set_para_spacing(p_sk, before=2, after=2)
        p_sk.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sk = p_sk.add_run(
            "  ·  ".join(s.strip() for s in ss.get("skills","").split(",") if s.strip())
        )
        r_sk.font.name  = font_name
        r_sk.font.size  = Pt(8.5)
        r_sk.font.color.rgb = RGBColor.from_string("4B5563")
        _hr(doc, "E5E7EB", thickness=4)

    _write_resume_body(doc, ss, accent=accent, body_color="374151",
                       font_name=font_name, heading_pt=heading_pt, body_pt=body_pt)

    # Soft skills / languages / interests footer section
    for sec, key in [("Soft Skills", "Softskills"),
                     ("Languages",   "languages"),
                     ("Interests",   "interests")]:
        raw = ss.get(key, "").strip()
        if raw:
            _section_heading(doc, sec, color_hex=accent,
                             font_name=font_name, font_pt=heading_pt)
            _skills_line(doc, raw, font_name=font_name, font_pt=body_pt - 0.5)
            _para(doc, before=0, after=4)


# ─────────────────────────────────────────────────────────────────────────────
# TWO-COLUMN CORPORATE TEMPLATES (Corporate Blue, Teal Impact, Navy Prestige,
#                                  Creative Green, Warm Terracotta, Indigo Tech)
# Thin left column (35%) for skills sidebar + wide right column (65%) for body.
# ─────────────────────────────────────────────────────────────────────────────

def _build_two_column(doc: Document, ss: dict,
                      header_bg:  str = "1D4ED8",
                      accent:     str = "1D4ED8",
                      sidebar_bg: str = "EFF6FF",
                      font_name:  str = "Calibri") -> None:
    """Two-column layout: sidebar left (35%) + body right (65%)."""
    _set_a4(doc)

    FONT = font_name

    # Full-width header band via a 1-col table
    hdr_tbl = doc.add_table(rows=1, cols=1)
    _no_borders(hdr_tbl)
    hc = hdr_tbl.rows[0].cells[0]
    _shading(hc, header_bg)
    _cell_margins(hc, top=180, bottom=140, start=180, end=180)

    p_name = hc.add_paragraph()
    _set_para_spacing(p_name, before=0, after=3)
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_name.add_run(ss.get("name", ""))
    r.bold = True
    r.font.name  = FONT
    r.font.size  = Pt(20)
    r.font.color.rgb = RGBColor.from_string("FFFFFF")

    if ss.get("job_title"):
        p_jt = hc.add_paragraph()
        _set_para_spacing(p_jt, before=0, after=6)
        r2 = p_jt.add_run(ss.get("job_title","").upper())
        r2.font.name  = FONT
        r2.font.size  = Pt(9.5)
        r2.font.color.rgb = RGBColor.from_string("DBEAFE")

    # Contact row in header
    parts = []
    for k in ("email","phone","location","linkedin"):
        v = ss.get(k,"").strip()
        if v:
            parts.append(v)
    if parts:
        p_c = hc.add_paragraph()
        _set_para_spacing(p_c, before=0, after=0)
        rc = p_c.add_run("  |  ".join(parts))
        rc.font.name  = FONT
        rc.font.size  = Pt(8.5)
        rc.font.color.rgb = RGBColor.from_string("BFDBFE")

    _para(doc, before=0, after=4)

    # Body table: sidebar + main
    body_tbl = doc.add_table(rows=1, cols=2)
    _no_borders(body_tbl)
    sc = body_tbl.rows[0].cells[0]
    mc = body_tbl.rows[0].cells[1]
    _set_col_width(sc, CONTENT_WIDTH_CM * 0.33)
    _set_col_width(mc, CONTENT_WIDTH_CM * 0.67)
    _shading(sc, sidebar_bg)
    _cell_margins(sc, top=160, bottom=160, start=140, end=120)
    _cell_margins(mc, top=160, bottom=160, start=160, end=140)

    def _ssec(cell, title):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=10, after=2)
        r = p.add_run(title.upper())
        r.bold = True
        r.font.name  = FONT
        r.font.size  = Pt(8)
        r.font.color.rgb = RGBColor.from_string(accent)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"),  "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), accent)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _sline(cell, text, italic=False, color="374151"):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=0, after=2)
        r = p.add_run(text.strip())
        r.italic = italic
        r.font.name  = FONT
        r.font.size  = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(color)

    # Sidebar sections
    for sec, key in [("Technical Skills", "skills"),
                     ("Soft Skills",       "Softskills"),
                     ("Languages",         "languages"),
                     ("Interests",         "interests")]:
        raw = ss.get(key,"").strip()
        if raw:
            _ssec(sc, sec)
            for item in [s.strip() for s in raw.split(",") if s.strip()]:
                _sline(sc, f"• {item}")

    certs = [c for c in ss.get("certificate_links",[]) if c.get("name")]
    if certs:
        _ssec(sc, "Certifications")
        for c in certs:
            _sline(sc, c.get("name",""), italic=True)
            if c.get("duration"):
                _sline(sc, c.get("duration",""), color="6B7280")

    # Main column
    def _msec(cell, title):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=10, after=2)
        r = p.add_run(title.upper())
        r.bold = True
        r.font.name  = FONT
        r.font.size  = Pt(10)
        r.font.color.rgb = RGBColor.from_string(accent)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"),  "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), accent)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _ml(cell, text, bold=False, italic=False, pt=9.5, color="374151", before=0, after=2):
        p = cell.add_paragraph()
        _set_para_spacing(p, before=before, after=after)
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.name  = FONT
        r.font.size  = Pt(pt)
        r.font.color.rgb = RGBColor.from_string(color)

    def _mb(cell, text, pt=9.0, color="374151"):
        p = cell.add_paragraph(style="List Bullet")
        _set_para_spacing(p, before=0, after=1)
        r = p.add_run(text.strip())
        r.font.name  = FONT
        r.font.size  = Pt(pt)
        r.font.color.rgb = RGBColor.from_string(color)

    def _mt(cell, left, right, lpt=10.5, rpt=9.0, lcol=None, rcol="6B7280"):
        if lcol is None:
            lcol = accent
        p = cell.add_paragraph()
        _set_para_spacing(p, before=0, after=0)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab  = OxmlElement("w:tab")
        cw   = int(CONTENT_WIDTH_CM * 0.67 * 567)
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(cw))
        tabs.append(tab)
        pPr.append(tabs)
        r1 = p.add_run(left)
        r1.bold      = True
        r1.font.name = FONT
        r1.font.size = Pt(lpt)
        r1.font.color.rgb = RGBColor.from_string(lcol)
        r2 = p.add_run("\t" + right)
        r2.font.name = FONT
        r2.font.size = Pt(rpt)
        r2.font.color.rgb = RGBColor.from_string(rcol)

    summary = ss.get("summary","").strip()
    if summary:
        _msec(mc, "Summary")
        for ln, ib in _parse_description_bullets(summary):
            if ib:
                _mb(mc, ln)
            else:
                _ml(mc, ln, pt=9.5)
        _ml(mc, "", before=0, after=4)

    exps = [e for e in ss.get("experience_entries",[])
            if e.get("company") or e.get("title")]
    if exps:
        _msec(mc, "Work Experience")
        for e in exps:
            _mt(mc, e.get("company",""), e.get("duration",""))
            if e.get("title"):
                _ml(mc, e.get("title",""), italic=True, pt=9.5, color="4B5563")
            for ln, ib in _parse_description_bullets(e.get("description","")):
                if ib:
                    _mb(mc, ln)
                else:
                    _ml(mc, ln, pt=9.0)
            _ml(mc, "", before=0, after=4)

    edus = [e for e in ss.get("education_entries",[])
            if e.get("institution") or e.get("degree")]
    if edus:
        _msec(mc, "Education")
        for edu in edus:
            dg = edu.get("degree","")
            if isinstance(dg, list): dg = ", ".join(dg)
            _mt(mc, edu.get("institution",""), edu.get("year",""))
            if dg:
                _ml(mc, dg, italic=True, pt=9.5, color="4B5563")
            if edu.get("details"):
                _ml(mc, edu.get("details",""), pt=9.0)
            _ml(mc, "", before=0, after=4)

    projs = [p for p in ss.get("project_entries",[]) if p.get("title")]
    links = ss.get("project_links", [])
    if projs:
        _msec(mc, "Projects")
        for idx, proj in enumerate(projs):
            _mt(mc, proj.get("title",""), proj.get("duration",""))
            if proj.get("tech"):
                _ml(mc, f"Tech: {proj['tech']}", italic=True, pt=9.0, color="4B5563")
            for ln, ib in _parse_description_bullets(proj.get("description","")):
                if ib:
                    _mb(mc, ln)
                else:
                    _ml(mc, ln, pt=9.0)
            if idx < len(links) and links[idx]:
                _ml(mc, f"Link: {links[idx]}", italic=True, pt=8.5, color="2563EB")
            _ml(mc, "", before=0, after=4)


# ─────────────────────────────────────────────────────────────────────────────
# TEMPLATE REGISTRY — maps template name → (builder_fn, kwargs)
# ─────────────────────────────────────────────────────────────────────────────

_TEMPLATE_BUILDERS: dict[str, tuple] = {
    "Default (Professional)": (
        _build_default, {}
    ),
    "Modern Minimal": (
        _build_modern, {}
    ),
    "Elegant Sidebar": (
        _build_sidebar, {}
    ),
    "Classic Clean (Single Column)": (
        _build_single_column, {
            "accent": "374151", "name_color": "111827",
            "header_border_color": "374151",
        }
    ),
    "Executive (Single Column)": (
        _build_single_column, {
            "accent": "1E3A5F", "name_color": "1E3A5F",
            "header_border_color": "1E3A5F", "font_name": "Georgia",
        }
    ),
    "Timeline (Single Column)": (
        _build_single_column, {
            "accent": "B45309", "name_color": "92400E",
            "header_border_color": "B45309",
        }
    ),
    "Corporate Blue (Two Column)": (
        _build_two_column, {
            "header_bg": "1D4ED8", "accent": "1D4ED8",
            "sidebar_bg": "EFF6FF",
        }
    ),
    "Creative Green (Two Column)": (
        _build_two_column, {
            "header_bg": "166534", "accent": "166534",
            "sidebar_bg": "F0FDF4",
        }
    ),
    "Warm Terracotta (Two Column)": (
        _build_two_column, {
            "header_bg": "C2410C", "accent": "C2410C",
            "sidebar_bg": "FFF7ED",
        }
    ),
    "Navy Prestige (Two Column)": (
        _build_two_column, {
            "header_bg": "1E3A5F", "accent": "1E3A5F",
            "sidebar_bg": "F0F4FF",
        }
    ),
    "Slate Gray (Single Column)": (
        _build_single_column, {
            "accent": "475569", "name_color": "1E293B",
            "header_border_color": "475569",
        }
    ),
    "Teal Impact (Two Column)": (
        _build_two_column, {
            "header_bg": "0F766E", "accent": "0F766E",
            "sidebar_bg": "F0FDFA",
        }
    ),
    "Burgundy Classic (Single Column)": (
        _build_single_column, {
            "accent": "881337", "name_color": "4C0519",
            "header_border_color": "881337",
        }
    ),
    "Indigo Tech (Two Column)": (
        _build_two_column, {
            "header_bg": "4338CA", "accent": "4338CA",
            "sidebar_bg": "EEF2FF",
        }
    ),
    "Forest Green (Single Column)": (
        _build_single_column, {
            "accent": "14532D", "name_color": "14532D",
            "header_border_color": "14532D",
        }
    ),
    "Pure White (Single Column)": (
        _build_single_column, {
            "accent": "111111", "name_color": "111111",
            "header_border_color": "CCCCCC",
            "contact_color": "555555",
        }
    ),
    "Midnight Black (Single Column)": (
        _build_single_column, {
            "accent": "F59E0B", "name_color": "111827",
            "header_border_color": "F59E0B",
        }
    ),
    "Soft Lavender (Single Column)": (
        _build_single_column, {
            "accent": "6366F1", "name_color": "312E81",
            "header_border_color": "6366F1",
        }
    ),
    "Warm Sand (Single Column)": (
        _build_single_column, {
            "accent": "B45309", "name_color": "78350F",
            "header_border_color": "B45309",
        }
    ),
    "Ice Blue (Single Column)": (
        _build_single_column, {
            "accent": "0369A1", "name_color": "0C4A6E",
            "header_border_color": "0369A1",
        }
    ),
}



# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC API: build_resume_docx
# ─────────────────────────────────────────────────────────────────────────────

def build_resume_docx(template_name: str, session_state: Any) -> bytes:
    """
    Build a native A4 DOCX resume.

    Parameters
    ----------
    template_name  : One of the 20 template names in TEMPLATE_META
    session_state  : Streamlit session_state (dict-like) with all resume fields

    Returns
    -------
    bytes  — raw .docx bytes ready for st.download_button
    """
    fn, kwargs = _TEMPLATE_BUILDERS.get(
        template_name,
        _TEMPLATE_BUILDERS["Default (Professional)"]
    )
    doc = Document()

    # Remove default empty paragraph that Word adds
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Ensure List Bullet style exists
    try:
        _ = doc.styles["List Bullet"]
    except KeyError:
        style = doc.styles.add_style("List Bullet", 1)
        style.base_style = doc.styles["Normal"]

    fn(doc, session_state, **kwargs)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# COVER LETTER DOCX BUILDERS
# ─────────────────────────────────────────────────────────────────────────────

def _cl_header_simple(doc: Document, data: dict,
                      font_name: str = "Calibri",
                      name_color: str = "1E3A5F",
                      accent: str = "1E3A5F",
                      contact_color: str = "374151") -> None:
    """Shared header: name, title, contact line."""
    name      = data.get("name", "")
    job_title = data.get("job_title", "")
    parts     = []
    for k in ("email", "phone", "location", "linkedin"):
        v = data.get(k, "").strip()
        if v:
            parts.append(v)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(p, before=0, after=2)
    r = p.add_run(name)
    r.bold = True
    r.font.name  = font_name
    r.font.size  = Pt(20)
    r.font.color.rgb = RGBColor.from_string(name_color)

    if job_title:
        pjt = doc.add_paragraph()
        _set_para_spacing(pjt, before=0, after=3)
        rjt = pjt.add_run(job_title)
        rjt.font.name  = font_name
        rjt.font.size  = Pt(11)
        rjt.font.color.rgb = RGBColor.from_string(accent)

    if parts:
        pc = doc.add_paragraph()
        _set_para_spacing(pc, before=0, after=4)
        rc = pc.add_run("  |  ".join(parts))
        rc.font.name  = font_name
        rc.font.size  = Pt(9)
        rc.font.color.rgb = RGBColor.from_string(contact_color)


def _cl_recipient_block(doc: Document, data: dict,
                        font_name: str = "Calibri") -> None:
    date_str = data.get("date", "")
    hiring   = data.get("hiring_manager", "Hiring Manager")
    company  = data.get("company", "")
    role     = data.get("role", "")

    if date_str:
        pd = doc.add_paragraph()
        _set_para_spacing(pd, before=0, after=3)
        rd = pd.add_run(date_str)
        rd.font.name  = font_name
        rd.font.size  = Pt(9)
        rd.font.color.rgb = RGBColor.from_string("6B7280")

    ph = doc.add_paragraph()
    _set_para_spacing(ph, before=0, after=1)
    rh = ph.add_run(hiring)
    rh.bold = True
    rh.font.name  = font_name
    rh.font.size  = Pt(10)
    rh.font.color.rgb = RGBColor.from_string("111827")

    if company:
        pco = doc.add_paragraph()
        _set_para_spacing(pco, before=0, after=6)
        rco = pco.add_run(company)
        rco.font.name  = font_name
        rco.font.size  = Pt(10)
        rco.font.color.rgb = RGBColor.from_string("374151")

    # Greeting
    pg = doc.add_paragraph()
    _set_para_spacing(pg, before=0, after=8)
    rg = pg.add_run(f"Dear {hiring},")
    rg.font.name  = font_name
    rg.font.size  = Pt(10.5)
    rg.font.color.rgb = RGBColor.from_string("111827")


def _cl_body_paragraphs(doc: Document, paragraphs: list,
                        font_name: str = "Calibri",
                        color: str = "374151") -> None:
    for para_text in paragraphs:
        if not para_text.strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_para_spacing(p, before=0, after=10, line_val=276)
        r = p.add_run(para_text.strip())
        r.font.name  = font_name
        r.font.size  = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string(color)


def _cl_closing(doc: Document, name: str, job_title: str,
                sign_off: str = "Sincerely,",
                font_name: str = "Calibri",
                name_color: str = "1E3A5F") -> None:
    _para(doc, before=0, after=4)
    pc = doc.add_paragraph()
    _set_para_spacing(pc, before=0, after=2)
    rc = pc.add_run(sign_off)
    rc.font.name  = font_name
    rc.font.size  = Pt(10.5)
    rc.font.color.rgb = RGBColor.from_string("374151")

    pn = doc.add_paragraph()
    _set_para_spacing(pn, before=0, after=2)
    rn = pn.add_run(name)
    rn.bold = True
    rn.font.name  = font_name
    rn.font.size  = Pt(12)
    rn.font.color.rgb = RGBColor.from_string(name_color)

    if job_title:
        pjt = doc.add_paragraph()
        _set_para_spacing(pjt, before=0, after=0)
        rjt = pjt.add_run(job_title)
        rjt.font.name  = font_name
        rjt.font.size  = Pt(9.5)
        rjt.font.color.rgb = RGBColor.from_string("6B7280")


# ── Individual cover letter template builders ─────────────────────────────────

def _cl_professional(doc: Document, data: dict) -> None:
    _set_a4(doc)
    ACCENT = "1E3A5F"
    FONT   = "Georgia"
    _cl_header_simple(doc, data, font_name=FONT,
                      name_color=ACCENT, accent=ACCENT)
    _hr(doc, ACCENT, thickness=10)
    _cl_recipient_block(doc, data, font_name=FONT)
    _cl_body_paragraphs(doc, data.get("body_paragraphs", []),
                        font_name=FONT, color="1A1A1A")
    pc = doc.add_paragraph()
    _set_para_spacing(pc, before=0, after=8)
    rc = pc.add_run(f"I would welcome the opportunity to discuss how my experience aligns with the needs of {data.get('company', 'your organization')}. Thank you for your time and consideration.")
    rc.font.name  = FONT
    rc.font.size  = Pt(10.5)
    rc.font.color.rgb = RGBColor.from_string("1A1A1A")
    _cl_closing(doc, data.get("name",""), data.get("job_title",""),
                sign_off="Sincerely,", font_name=FONT, name_color=ACCENT)


def _cl_modern(doc: Document, data: dict) -> None:
    _set_a4(doc)
    ACCENT = "0D9488"
    FONT   = "Calibri"
    _cl_header_simple(doc, data, font_name=FONT,
                      name_color="0F172A", accent=ACCENT)
    _hr(doc, ACCENT, thickness=12)
    _cl_recipient_block(doc, data, font_name=FONT)
    _cl_body_paragraphs(doc, data.get("body_paragraphs", []),
                        font_name=FONT, color="374151")
    pc = doc.add_paragraph()
    _set_para_spacing(pc, before=0, after=8)
    rc = pc.add_run(f"I'd love the chance to discuss how I can contribute to {data.get('company', 'your team')}. Thank you for considering my application.")
    rc.font.name  = FONT
    rc.font.size  = Pt(10.5)
    rc.font.color.rgb = RGBColor.from_string("374151")
    _cl_closing(doc, data.get("name",""), data.get("job_title",""),
                sign_off="Best regards,", font_name=FONT, name_color=ACCENT)


def _cl_creative(doc: Document, data: dict) -> None:
    _set_a4(doc)
    ACCENT = data.get("accent_color", "#7C3AED").lstrip("#")
    FONT   = "Calibri"

    # Coloured top bar
    hdr = doc.add_table(rows=1, cols=1)
    _no_borders(hdr)
    hc = hdr.rows[0].cells[0]
    _shading(hc, ACCENT)
    _cell_margins(hc, top=160, bottom=140, start=180, end=180)

    pn = hc.add_paragraph()
    _set_para_spacing(pn, before=0, after=2)
    rn = pn.add_run(data.get("name",""))
    rn.bold = True
    rn.font.name  = FONT
    rn.font.size  = Pt(20)
    rn.font.color.rgb = RGBColor.from_string("FFFFFF")

    if data.get("job_title"):
        pjt = hc.add_paragraph()
        _set_para_spacing(pjt, before=0, after=6)
        rjt = pjt.add_run(data.get("job_title",""))
        rjt.font.name  = FONT
        rjt.font.size  = Pt(11)
        rjt.font.color.rgb = RGBColor.from_string("E9D5FF")

    parts = []
    for k in ("email","phone","location","linkedin"):
        v = data.get(k,"").strip()
        if v:
            parts.append(v)
    if parts:
        pc = hc.add_paragraph()
        _set_para_spacing(pc, before=0, after=0)
        rc = pc.add_run("  ·  ".join(parts))
        rc.font.name  = FONT
        rc.font.size  = Pt(9)
        rc.font.color.rgb = RGBColor.from_string("F3E8FF")

    _para(doc, before=0, after=8)
    _cl_recipient_block(doc, data, font_name=FONT)
    _cl_body_paragraphs(doc, data.get("body_paragraphs",[]),
                        font_name=FONT, color="1F2937")
    pc2 = doc.add_paragraph()
    _set_para_spacing(pc2, before=0, after=8)
    rc2 = pc2.add_run("I would be thrilled to discuss this further. Thank you for your time — I look forward to hearing from you.")
    rc2.font.name  = FONT
    rc2.font.size  = Pt(10.5)
    rc2.font.color.rgb = RGBColor.from_string("374151")

    pclose = doc.add_paragraph()
    _set_para_spacing(pclose, before=0, after=2)
    rclose = pclose.add_run("Warmly,")
    rclose.font.name  = FONT
    rclose.font.size  = Pt(10.5)
    rclose.font.color.rgb = RGBColor.from_string("374151")

    pname = doc.add_paragraph()
    _set_para_spacing(pname, before=0, after=0)
    rname = pname.add_run(data.get("name",""))
    rname.bold = True
    rname.font.name  = FONT
    rname.font.size  = Pt(14)
    rname.font.color.rgb = RGBColor.from_string(ACCENT)


def _cl_executive(doc: Document, data: dict) -> None:
    _set_a4(doc)
    ACCENT = "0D1B2A"
    GOLD   = "D4AF37"
    FONT   = "Georgia"

    # Dark header
    hdr = doc.add_table(rows=1, cols=1)
    _no_borders(hdr)
    hc  = hdr.rows[0].cells[0]
    _shading(hc, ACCENT)
    _cell_margins(hc, top=200, bottom=160, start=200, end=200)

    pn = hc.add_paragraph()
    _set_para_spacing(pn, before=0, after=3)
    rn = pn.add_run(data.get("name",""))
    rn.bold = True
    rn.font.name  = FONT
    rn.font.size  = Pt(22)
    rn.font.color.rgb = RGBColor.from_string("FFFFFF")

    if data.get("job_title"):
        pjt = hc.add_paragraph()
        _set_para_spacing(pjt, before=0, after=8)
        rjt = pjt.add_run(data.get("job_title","").upper())
        rjt.font.name  = FONT
        rjt.font.size  = Pt(9)
        rjt.font.color.rgb = RGBColor.from_string(GOLD)

    parts = []
    for k in ("email","phone","location","linkedin"):
        v = data.get(k,"").strip()
        if v:
            parts.append(v)
    if parts:
        pc = hc.add_paragraph()
        _set_para_spacing(pc, before=0, after=0)
        rc = pc.add_run("  |  ".join(parts))
        rc.font.name  = FONT
        rc.font.size  = Pt(9)
        rc.font.color.rgb = RGBColor.from_string("ADB5BD")

    # Gold separator line
    _hr(doc, GOLD, thickness=8)
    _para(doc, before=0, after=4)
    _cl_recipient_block(doc, data, font_name=FONT)
    _cl_body_paragraphs(doc, data.get("body_paragraphs",[]),
                        font_name=FONT, color="1A1A1A")
    pc2 = doc.add_paragraph()
    _set_para_spacing(pc2, before=0, after=8)
    rc2 = pc2.add_run("I welcome the opportunity to explore this further at your convenience. Please find my resume enclosed for your review.")
    rc2.font.name  = FONT
    rc2.font.size  = Pt(10.5)
    rc2.font.color.rgb = RGBColor.from_string("374151")
    _cl_closing(doc, data.get("name",""), data.get("job_title",""),
                sign_off="Respectfully yours,", font_name=FONT, name_color=ACCENT)


def _cl_entry_level(doc: Document, data: dict) -> None:
    _set_a4(doc)
    ACCENT = "1D4ED8"
    FONT   = "Calibri"

    # Blue accent left-bar box header
    hdr = doc.add_table(rows=1, cols=1)
    _no_borders(hdr)
    hc  = hdr.rows[0].cells[0]
    _shading(hc, "EFF6FF")
    _cell_margins(hc, top=160, bottom=140, start=180, end=180)

    # Left blue bar via paragraph border
    pn = hc.add_paragraph()
    _set_para_spacing(pn, before=0, after=2)
    pPr  = pn._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    lft  = OxmlElement("w:left")
    lft.set(qn("w:val"),   "single")
    lft.set(qn("w:sz"),    "20")
    lft.set(qn("w:space"), "8")
    lft.set(qn("w:color"), ACCENT)
    pBdr.append(lft)
    pPr.append(pBdr)
    rn = pn.add_run(data.get("name",""))
    rn.bold = True
    rn.font.name  = FONT
    rn.font.size  = Pt(20)
    rn.font.color.rgb = RGBColor.from_string("1E3A8A")

    if data.get("job_title"):
        pjt = hc.add_paragraph()
        _set_para_spacing(pjt, before=0, after=4)
        rjt = pjt.add_run(data.get("job_title",""))
        rjt.font.name  = FONT
        rjt.font.size  = Pt(11)
        rjt.font.color.rgb = RGBColor.from_string("3B82F6")

    parts = []
    for k in ("email","phone","location","linkedin"):
        v = data.get(k,"").strip()
        if v:
            parts.append(v)
    if parts:
        pc = hc.add_paragraph()
        _set_para_spacing(pc, before=0, after=0)
        rc = pc.add_run("  |  ".join(parts))
        rc.font.name  = FONT
        rc.font.size  = Pt(9)
        rc.font.color.rgb = RGBColor.from_string("6B7280")

    _para(doc, before=0, after=8)
    _cl_recipient_block(doc, data, font_name=FONT)
    _cl_body_paragraphs(doc, data.get("body_paragraphs",[]),
                        font_name=FONT, color="374151")
    pc2 = doc.add_paragraph()
    _set_para_spacing(pc2, before=0, after=8)
    rc2 = pc2.add_run("I would be grateful for the opportunity to interview and learn more about this role. Thank you for your time and consideration.")
    rc2.font.name  = FONT
    rc2.font.size  = Pt(10.5)
    rc2.font.color.rgb = RGBColor.from_string("374151")
    _cl_closing(doc, data.get("name",""), data.get("job_title",""),
                sign_off="Sincerely,", font_name=FONT, name_color="1E3A8A")


def _cl_ats(doc: Document, data: dict) -> None:
    _set_a4(doc)
    FONT = "Arial"

    # Plain ATS header
    pn = doc.add_paragraph()
    _set_para_spacing(pn, before=0, after=2)
    rn = pn.add_run(data.get("name",""))
    rn.bold = True
    rn.font.name  = FONT
    rn.font.size  = Pt(18)
    rn.font.color.rgb = RGBColor.from_string("111827")

    if data.get("job_title"):
        pjt = doc.add_paragraph()
        _set_para_spacing(pjt, before=0, after=2)
        rjt = pjt.add_run(data.get("job_title",""))
        rjt.font.name  = FONT
        rjt.font.size  = Pt(11)
        rjt.font.color.rgb = RGBColor.from_string("374151")

    parts = []
    for k in ("email","phone","location","linkedin","portfolio"):
        v = data.get(k,"").strip()
        if v:
            parts.append(v)
    if parts:
        pc = doc.add_paragraph()
        _set_para_spacing(pc, before=0, after=4)
        rc = pc.add_run("  |  ".join(parts))
        rc.font.name  = FONT
        rc.font.size  = Pt(9)
        rc.font.color.rgb = RGBColor.from_string("374151")

    _hr(doc, "D1D5DB", thickness=4)
    _cl_recipient_block(doc, data, font_name=FONT)

    # Subject line
    psub = doc.add_paragraph()
    _set_para_spacing(psub, before=0, after=8)
    rsub = psub.add_run(f"Re: Application for {data.get('role', data.get('job_title',''))} — {data.get('name','')}")
    rsub.bold = True
    rsub.font.name  = FONT
    rsub.font.size  = Pt(10.5)
    rsub.font.color.rgb = RGBColor.from_string("111827")

    _cl_body_paragraphs(doc, data.get("body_paragraphs",[]),
                        font_name=FONT, color="111827")

    # Key skills
    key_skills = data.get("key_skills","").strip()
    if key_skills:
        psk = doc.add_paragraph()
        _set_para_spacing(psk, before=0, after=8)
        rsk_bold = psk.add_run("Core Technical Skills: ")
        rsk_bold.bold = True
        rsk_bold.font.name  = FONT
        rsk_bold.font.size  = Pt(10.5)
        rsk_bold.font.color.rgb = RGBColor.from_string("111827")
        rsk = psk.add_run(key_skills)
        rsk.font.name  = FONT
        rsk.font.size  = Pt(10.5)
        rsk.font.color.rgb = RGBColor.from_string("111827")

    contact_ref = data.get("email") or data.get("phone") or "the contact details above"
    pclose = doc.add_paragraph()
    _set_para_spacing(pclose, before=0, after=8)
    rclose = pclose.add_run(f"I have attached my resume for your review. I am available for an interview at your earliest convenience and can be reached at {contact_ref}.")
    rclose.font.name  = FONT
    rclose.font.size  = Pt(10.5)
    rclose.font.color.rgb = RGBColor.from_string("374151")

    _cl_closing(doc, data.get("name",""), data.get("job_title",""),
                sign_off="Sincerely,", font_name=FONT, name_color="111827")


# ── Cover letter template registry ───────────────────────────────────────────

_CL_BUILDERS: dict[str, Any] = {
    "Professional / Corporate":     _cl_professional,
    "Modern Minimal":               _cl_modern,
    "Creative":                     _cl_creative,
    "Executive":                    _cl_executive,
    "Entry-Level / Fresher":        _cl_entry_level,
    "Technical / ATS-Optimized":    _cl_ats,
}


# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC API: build_cover_letter_docx
# ─────────────────────────────────────────────────────────────────────────────

def build_cover_letter_docx(template_name: str, data: dict) -> bytes:
    """
    Build a native A4 DOCX cover letter.

    Parameters
    ----------
    template_name  : One of the COVER_LETTER_TEMPLATES keys
    data           : dict with keys: name, job_title, email, phone, location,
                     linkedin, portfolio, company, hiring_manager, role, date,
                     body_paragraphs (list[str]), key_skills (str), accent_color (str)

    Returns
    -------
    bytes — raw .docx bytes
    """
    fn = _CL_BUILDERS.get(template_name, _cl_professional)
    doc = Document()

    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    try:
        _ = doc.styles["List Bullet"]
    except KeyError:
        style = doc.styles.add_style("List Bullet", 1)
        style.base_style = doc.styles["Normal"]

    fn(doc, data)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
